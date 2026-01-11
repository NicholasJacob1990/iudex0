#!/usr/bin/env python3
"""
Script para processar transcrições de aulas usando Gemini via OpenRouter
Formata conforme diretrizes específicas para concursos de procuradorias
VERSÃO COM CHUNKS MENORES E PROMPT ANTI-RESUMO
"""

import os
import sys
import time
import re
import requests
from pathlib import Path
from docx import Document
from typing import List

# Configuração
OPENROUTER_API_KEY = "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"
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL_NAME = "google/gemini-2.0-flash-001"
MAX_OUTPUT_TOKENS = 65_536
# Chunks menores para garantir que o modelo formate sem resumir
MAX_CHARS_PER_CHUNK = 30_000  # ~7.5k tokens
OUTPUT_DIR = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ')

PROMPT_FORMATACAO = """TAREFA: REVISAR E FORMATAR TRANSCRIÇÃO - NÃO RESUMIR!

Você é um revisor de transcrições para concursos de procuradorias.

⚠️ REGRA CRÍTICA: NÃO RESUMA! O TEXTO DE SAÍDA DEVE TER TAMANHO SIMILAR OU MAIOR QUE O DE ENTRADA!

Sua tarefa é APENAS:
1. Corrigir erros gramaticais e de pontuação
2. Melhorar a formatação e legibilidade
3. Organizar em tópicos/subtópicos
4. Adicionar tabelas de síntese ao final de cada tópico principal

DIRETRIZES OBRIGATÓRIAS:
- Mantenha TODAS as ideias, exemplos, explicações, dicas e casos
- Mantenha o modo em primeira pessoa
- Corrija linguagem coloquial para português padrão formal
- Elimine vícios de oralidade MAS MANTENHA TODO O CONTEÚDO
- Preserve a sequência exata das ideias
- Use parágrafos bem definidos (evite parágrafos muito longos)
- Enumere tópicos e subtópicos (1, 1.1, 2, 2.1...)
- Use negrito para conceitos importantes
- USE TEXTO CORRIDO - mantenha todo o conteúdo explicativo

⚠️ PROIBIDO:
- NÃO resuma
- NÃO omita informações
- NÃO pule exemplos ou casos práticos
- NÃO remova dicas de prova
- NÃO corte explicações

Ao final de CADA tópico principal, adicione uma tabela:

| Conceito | Definição | Fundamento Legal | Observações |
|----------|-----------|------------------|-------------|
| [Termo] | [Explicação] | [Art./Lei] | [Dicas] |

{contexto_numeracao}

<transcrição>
{texto}
</transcrição>"""


def extrair_texto(caminho_arquivo: str) -> str:
    """Extrai texto de arquivo txt ou docx"""
    path = Path(caminho_arquivo)
    if path.suffix.lower() == '.docx':
        doc = Document(caminho_arquivo)
        return '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        with open(caminho_arquivo, 'r', encoding='utf-8') as f:
            return f.read()


def dividir_texto_em_chunks(texto: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> List[str]:
    """Divide o texto em chunks em pontos naturais (parágrafos/quebras duplas)."""
    paragrafos = texto.split('\n\n')
    chunks = []
    chunk_atual = []
    tamanho_atual = 0
    
    for paragrafo in paragrafos:
        tamanho_paragrafo = len(paragrafo)
        
        if tamanho_atual + tamanho_paragrafo > max_chars and chunk_atual:
            chunks.append('\n\n'.join(chunk_atual))
            chunk_atual = [paragrafo]
            tamanho_atual = tamanho_paragrafo
        else:
            chunk_atual.append(paragrafo)
            tamanho_atual += tamanho_paragrafo
    
    if chunk_atual:
        chunks.append('\n\n'.join(chunk_atual))
    
    return chunks


def processar_com_openrouter(texto: str, parte_num: int, total_partes: int, ultimo_topico: int) -> str:
    """Processa um chunk de texto com Gemini via OpenRouter"""
    
    if parte_num == 1:
        contexto_numeracao = "Esta é a PRIMEIRA parte. Inicie a numeração em 1."
    else:
        contexto_numeracao = f"PARTE {parte_num}/{total_partes}. Continue a numeração a partir do tópico {ultimo_topico + 1}."
    
    prompt_completo = PROMPT_FORMATACAO.format(
        texto=texto,
        contexto_numeracao=contexto_numeracao
    )
    
    print(f"\n{'='*60}")
    print(f"📝 Processando parte {parte_num}/{total_partes}")
    print(f"   Entrada: {len(texto):,} chars")
    print(f"{'='*60}")
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://iudex.app",
        "X-Title": "Iudex Formatter"
    }
    
    payload = {
        "model": MODEL_NAME,
        "messages": [{"role": "user", "content": prompt_completo}],
        "temperature": 0.1,
        "max_tokens": MAX_OUTPUT_TOKENS,
        "top_p": 0.95
    }
    
    response = requests.post(
        OPENROUTER_BASE_URL,
        headers=headers,
        json=payload,
        timeout=600
    )
    response.raise_for_status()
    
    result = response.json()['choices'][0]['message']['content']
    print(f"   Saída: {len(result):,} chars")
    
    return result


def extrair_ultimo_topico(texto: str) -> int:
    """Extrai o número do último tópico principal"""
    matches = re.findall(r'(?:^|\n)(?:#+ |\*\*)?(\d+)\.', texto)
    return max(int(m) for m in matches) if matches else 0


def salvar_markdown(texto: str, caminho: Path):
    """Salva como Markdown"""
    with open(caminho, 'w', encoding='utf-8') as f:
        f.write(texto)


def salvar_docx(texto: str, caminho: Path):
    """Salva como DOCX simples"""
    doc = Document()
    for linha in texto.split('\n'):
        if linha.strip():
            doc.add_paragraph(linha)
    doc.save(str(caminho))


def main():
    if len(sys.argv) > 1:
        arquivo_entrada = Path(sys.argv[1])
    else:
        arquivo_entrada = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/Administrativo.txt')
    
    if not arquivo_entrada.exists():
        print(f"❌ Arquivo não encontrado: {arquivo_entrada}")
        return
    
    print(f"📄 Lendo: {arquivo_entrada}")
    texto_original = extrair_texto(str(arquivo_entrada))
    print(f"✅ {len(texto_original):,} caracteres")
    
    chunks = dividir_texto_em_chunks(texto_original)
    print(f"📊 Dividido em {len(chunks)} partes")
    
    resultados = []
    ultimo_topico = 0
    
    for i, chunk in enumerate(chunks, 1):
        print(f"\n🔄 Parte {i}/{len(chunks)}...")
        
        try:
            resultado = processar_com_openrouter(chunk, i, len(chunks), ultimo_topico)
            resultados.append(resultado)
            ultimo_topico = extrair_ultimo_topico(resultado)
            print(f"✅ Concluído (último tópico: {ultimo_topico})")
            
            if i < len(chunks):
                time.sleep(1)
                
        except Exception as e:
            print(f"❌ Erro: {e}")
            continue
    
    texto_final = '\n\n'.join(resultados)
    nome_base = arquivo_entrada.stem
    
    print(f"\n💾 Salvando...")
    
    arquivo_md = OUTPUT_DIR / f'{nome_base}_GEMINI_V2.md'
    salvar_markdown(texto_final, arquivo_md)
    print(f"✅ Markdown: {arquivo_md}")
    
    arquivo_docx = OUTPUT_DIR / f'{nome_base}_GEMINI_V2.docx'
    salvar_docx(texto_final, arquivo_docx)
    print(f"✅ DOCX: {arquivo_docx}")
    
    print(f"\n{'='*60}")
    print(f"🎉 CONCLUÍDO!")
    print(f"   Entrada: {len(texto_original):,} chars")
    print(f"   Saída: {len(texto_final):,} chars")
    print(f"   Razão: {len(texto_final)/len(texto_original)*100:.1f}%")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
