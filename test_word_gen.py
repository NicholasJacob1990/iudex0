from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re
import os

def _format_inline_markdown(paragraph, text):
    """Formata markdown inline (negrito, itálico) em um parágrafo existente usando Regex"""
    paragraph.clear()
    
    # Regex que captura **bold**, *italic* ou texto normal
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for token in tokens:
        if not token: continue
        
        run = paragraph.add_run()
        if token.startswith('**') and token.endswith('**'):
            run.bold = True
            run.text = token[2:-2] # Remove os **
        elif token.startswith('*') and token.endswith('*'):
            run.italic = True
            run.text = token[1:-1] # Remove os *
        else:
            run.text = token

def save_as_word(formatted_text, video_name, output_folder):
    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = formatted_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
            i += 1
            continue
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
            i += 1
            continue
            
        # Bullets
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            _format_inline_markdown(p, line[2:])
            i += 1
            continue
            
        # Normal
        p = doc.add_paragraph()
        _format_inline_markdown(p, line)
        i += 1
    
    output_path = os.path.join(output_folder, f"{video_name}_TEST.docx")
    doc.save(output_path)
    print(f"Saved to {output_path}")

# Test Data
text = """# Título Principal
## Subtítulo
Texto normal com **negrito** e *itálico*.
- Item de lista
- Item com **negrito**
"""

save_as_word(text, "Teste_Formatacao", ".")
