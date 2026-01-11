#!/usr/bin/env python3
"""
Script para processar transcrições de aulas usando Gemini 3 Pro Preview
Formata conforme diretrizes específicas para concursos de procuradorias
"""

import os
import sys
import time
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from typing import List
from colorama import Fore, init
from dotenv import load_dotenv
import re

init(autoreset=True)
load_dotenv(override=True)

# Configuração
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
MODEL_NAME = 'gemini-2.0-flash-exp'  # Modelo disponível atualmente
MAX_OUTPUT_TOKENS = 65_536
MAX_CHARS_PER_CHUNK = 50_000  # ~12.5k tokens, conservador para saída
OUTPUT_DIR = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ')

PROMPT_FORMATACAO = """Você é um especialista em revisão de transcrições para concursos de procuradorias. Sua tarefa é revisar a transcrição fornecida de uma aula, corrigindo erros de português, erros gramaticais e de pontuação, melhorando a formatação para facilitar a leitura, e mantendo o conteúdo original. Não resuma, não parafraseie e não adicione informações que não estejam na transcrição original. Siga estas diretrizes:

- Mantenha o modo em primeira pessoa
- Corrija erros gramaticais, ortográficos e de pontuação, tornando o texto gramaticalmente correto e claro
- Mantenha todo o conteúdo original, incluindo ideias, exemplos, explicações, pausas, hesitações e ideias incompletas, fazendo o uso apropriado de aspas, parênteses. Não resuma, não omita informações nem altere o significado
- Melhore a formatação para facilitar a leitura
- Mantenha todo o conteúdo original, mas corrija erros da linguagem coloquial para torná-la mais clara, didática e legível
- Ajuste a linguagem coloquial para um português padrão, mantendo o significado original
- Elimine vícios da oralidade e gírias
- Preserve a sequência exata das falas e ideias apresentadas
- Utilize formatação e estrutura com parágrafos bem definidos, facilitando a leitura e compreensão, para melhorar a legibilidade, seguindo o fluxo natural do discurso. Evite parágrafos longos
- Reproduza fielmente as informações, apenas melhorando a clareza e a legibilidade
- Utilize conectivos necessários para tornar o texto mais fluido. Aplique a pontuação devida para deixar o texto coeso e coerente
- Corrija vícios de linguagem, como repetições desnecessárias, uso excessivo de advérbios, linguagem vaga ou imprecisa, gírias, expressões redundantes, e outros erros que afetem a clareza e a eficácia da comunicação, sem alterar o significado do texto
- NÃO mencione os falantes (SPEAKER, Professor:, etc.) - converta para texto impessoal didático
- Divida a aula em tópicos e subtópicos para melhor organização e visualização do conteúdo
- Enumere os tópicos e subtópicos sequencialmente (1, 1.1, 1.2, 2, 2.1, etc.), use negrito quando mais apropriado
- Intercale parágrafos curtos e longos conforme a ideia neles contida, tornando a leitura menos cansativa
- Seja didático sem perder detalhes e conteúdo
- USE TEXTO CORRIDO NA MEDIDA DO POSSÍVEL
- **Ao final de cada tópico/capítulo, sintetize/resume o assunto de forma esquematizada, preferencialmente por tabelas

**Ao final de CADA tópico principal, crie uma tabela de síntese:**

| Conceito/Instituto | Definição | Fundamento Legal | Observações |
|-------------------|-----------|-----------------|-------------|
| [Preencher] | [Resumo] | [Art. X, Lei Y] | [Dicas/Exceções] |

IMPORTANTE: 
- Retorne APENAS o texto formatado em Markdown
- NÃO adicione comentários como "Continuação...", "Parte X...", "[Fim]", etc.
- Mantenha a numeração sequencial dos tópicos
- NÃO mencione "Professor", "SPEAKER", "Aluno" - converta tudo para texto impessoal

{contexto_numeracao}

<transcrição>
{texto}
</transcrição>
"""


def extrair_texto_txt(caminho_arquivo: str) -> str:
    """Extrai texto de um arquivo .txt"""
    with open(caminho_arquivo, 'r', encoding='utf-8') as f:
        return f.read()


def dividir_texto_em_chunks(texto: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> List[str]:
    """
    Divide o texto em chunks para não exceder o limite de tokens do modelo.
    Tenta dividir em pontos naturais (parágrafos).
    """
    paragrafos = texto.split('\n\n')
    chunks = []
    chunk_atual = []
    tamanho_atual = 0
    
    for paragrafo in paragrafos:
        tamanho_paragrafo = len(paragrafo)
        
        if tamanho_atual + tamanho_paragrafo > max_chars and chunk_atual:
            chunks.append('\n\n'.join(chunk_atual))
            chunk_atual = [paragrafo]
            tamanho_atual = tamanho_paragrafo
        else:
            chunk_atual.append(paragrafo)
            tamanho_atual += tamanho_paragrafo
    
    if chunk_atual:
        chunks.append('\n\n'.join(chunk_atual))
    
    return chunks


def processar_com_gemini(texto: str, parte_num: int = 1, total_partes: int = 1, ultimo_topico: int = 0) -> str:
    """Processa um chunk de texto com Gemini"""
    
    genai.configure(api_key=GEMINI_API_KEY)
    
    generation_config = {
        "temperature": 0.2,
        "top_p": 0.95,
        "top_k": 40,
        "max_output_tokens": MAX_OUTPUT_TOKENS,
    }
    
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]
    
    model = genai.GenerativeModel(
        model_name=MODEL_NAME,
        generation_config=generation_config,
        safety_settings=safety_settings
    )
    
    # Contexto para manter numeração contínua
    if parte_num == 1:
        contexto_numeracao = "Esta é a PRIMEIRA parte. Inicie a numeração dos tópicos em 1."
    else:
        contexto_numeracao = f"Esta é a PARTE {parte_num} de {total_partes}. Continue a numeração a partir do tópico {ultimo_topico + 1}. NÃO reinicie a numeração."
    
    prompt_completo = PROMPT_FORMATACAO.format(
        texto=texto,
        contexto_numeracao=contexto_numeracao
    )
    
    print(f"\n{Fore.CYAN}{'='*60}")
    print(f"{Fore.CYAN}📝 Processando parte {parte_num}/{total_partes}")
    print(f"{Fore.GREEN}   Tamanho do texto: {len(texto):,} caracteres")
    print(f"{Fore.CYAN}{'='*60}\n")
    
    try:
        response = model.generate_content(prompt_completo)
        return response.text
    except Exception as e:
        print(f"{Fore.RED}⚠️  Erro ao processar com Gemini: {e}")
        raise


def extrair_ultimo_topico(texto: str) -> int:
    """Extrai o número do último tópico principal do texto formatado"""
    # Procura por padrões como "## 5." ou "# 12." ou "**5."
    matches = re.findall(r'(?:^|\n)(?:#+ |\*\*)?(\d+)\.', texto)
    if matches:
        return max(int(m) for m in matches)
    return 0


def salvar_resultado_markdown(texto: str, caminho: Path) -> Path:
    """Salva o resultado em um arquivo .md"""
    with open(caminho, 'w', encoding='utf-8') as f:
        f.write(texto)
    return caminho


def salvar_resultado_docx(texto_markdown: str, caminho: Path) -> Path:
    """Converte Markdown para DOCX e salva"""
    doc = Document()
    
    # Margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)
    
    lines = texto_markdown.split('\n')
    current_table_data = []
    in_table = False
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            if in_table and current_table_data:
                _add_table(doc, current_table_data)
                current_table_data = []
                in_table = False
            continue
        
        # Tabelas
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            in_table = True
            if not re.match(r'^[\|\-\:\s]+$', line_stripped):
                cells = [cell.strip() for cell in line_stripped.split('|')[1:-1]]
                if cells and any(c for c in cells):
                    current_table_data.append(cells)
            continue
        
        if in_table and current_table_data:
            _add_table(doc, current_table_data)
            current_table_data = []
            in_table = False
        
        # Headings
        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=1)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith('#### '):
            doc.add_heading(line_stripped[5:], level=4)
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line_stripped[2:])
        elif re.match(r'^\d+\.\s', line_stripped):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s*', '', line_stripped)
            _add_formatted_text(p, text)
        elif line_stripped.startswith('---'):
            pass  # Ignora separadores
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line_stripped)
    
    if in_table and current_table_data:
        _add_table(doc, current_table_data)
    
    doc.save(str(caminho))
    return caminho


def _add_table(doc, table_data):
    """Adiciona tabela ao documento"""
    if not table_data:
        return
    
    num_cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                row.cells[j].text = cell_text
                if i == 0:
                    for para in row.cells[j].paragraphs:
                        for run in para.runs:
                            run.bold = True
    
    doc.add_paragraph()


def _add_formatted_text(paragraph, text):
    """Adiciona texto com formatação"""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def main():
    """Função principal"""
    
    if not GEMINI_API_KEY:
        print(f"{Fore.RED}❌ ERRO: Configure a variável de ambiente GEMINI_API_KEY")
        print(f"{Fore.YELLOW}   Execute: export GEMINI_API_KEY='sua-chave-aqui'")
        return
    
    # Arquivo de entrada
    if len(sys.argv) > 1:
        arquivo_entrada = Path(sys.argv[1])
    else:
        arquivo_entrada = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/Administrativo.txt')
    
    if not arquivo_entrada.exists():
        print(f"{Fore.RED}❌ Arquivo não encontrado: {arquivo_entrada}")
        return
    
    print(f"{Fore.CYAN}📄 Lendo arquivo: {arquivo_entrada}")
    texto_original = extrair_texto_txt(str(arquivo_entrada))
    
    print(f"{Fore.GREEN}✅ Texto extraído: {len(texto_original):,} caracteres")
    print(f"{Fore.GREEN}   Estimativa: ~{len(texto_original)//4:,} tokens")
    
    # Dividir em chunks
    chunks = dividir_texto_em_chunks(texto_original)
    print(f"\n{Fore.CYAN}📊 Documento dividido em {len(chunks)} parte(s)")
    
    # Processar cada chunk
    resultados = []
    ultimo_topico = 0
    
    for i, chunk in enumerate(chunks, 1):
        print(f"\n{Fore.YELLOW}🔄 Processando parte {i}/{len(chunks)}...")
        
        try:
            resultado = processar_com_gemini(
                texto=chunk,
                parte_num=i,
                total_partes=len(chunks),
                ultimo_topico=ultimo_topico
            )
            resultados.append(resultado)
            
            # Atualiza o último tópico para manter numeração
            ultimo_topico = extrair_ultimo_topico(resultado)
            
            print(f"{Fore.GREEN}✅ Parte {i} processada com sucesso (último tópico: {ultimo_topico})")
            
            if i < len(chunks):
                print(f"{Fore.CYAN}⏳ Aguardando 2 segundos...")
                time.sleep(2)
                
        except Exception as e:
            print(f"{Fore.RED}❌ Erro ao processar parte {i}: {e}")
            continue
    
    # Combinar resultados
    texto_final = '\n\n'.join(resultados)
    
    print(f"\n{Fore.CYAN}💾 Salvando resultados...")
    
    # Nome base para os arquivos
    nome_base = arquivo_entrada.stem
    
    # Salvar Markdown
    arquivo_md = OUTPUT_DIR / f'{nome_base}_GEMINI_APOSTILA.md'
    salvar_resultado_markdown(texto_final, arquivo_md)
    print(f"{Fore.GREEN}✅ Markdown salvo: {arquivo_md}")
    
    # Salvar DOCX
    try:
        arquivo_docx = OUTPUT_DIR / f'{nome_base}_GEMINI_APOSTILA.docx'
        salvar_resultado_docx(texto_final, arquivo_docx)
        print(f"{Fore.GREEN}✅ DOCX salvo: {arquivo_docx}")
    except Exception as e:
        print(f"{Fore.YELLOW}⚠️  Erro ao salvar DOCX: {e}")
    
    print(f"\n{Fore.GREEN}{'='*60}")
    print(f"{Fore.GREEN}🎉 PROCESSAMENTO CONCLUÍDO!")
    print(f"{Fore.GREEN}   Total de partes processadas: {len(resultados)}/{len(chunks)}")
    print(f"{Fore.GREEN}   Tamanho final: {len(texto_final):,} caracteres")
    print(f"{Fore.GREEN}{'='*60}")


if __name__ == '__main__':
    main()
