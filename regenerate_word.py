#!/usr/bin/env python3
"""
Script para regenerar o documento Word a partir do Markdown já formatado.
Versão robusta com tratamento de erros.
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from colorama import Fore, init
from datetime import datetime

init(autoreset=True)


def sanitize_text(text):
    """Remove caracteres problemáticos para XML"""
    if not text:
        return ""
    # Remove caracteres de controle (exceto newline, tab)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Substitui caracteres Unicode problemáticos
    text = text.replace('\u2028', ' ')  # Line separator
    text = text.replace('\u2029', '\n')  # Paragraph separator
    return text


def add_page_number(paragraph):
    """Adiciona número de página ao rodapé"""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.text = "PAGE"
    
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_cell_shading(cell, color_hex):
    """Define a cor de fundo de uma célula"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_formatted_text(paragraph, text):
    """Adiciona texto com formatação inline (negrito, itálico)"""
    text = sanitize_text(text)
    
    # Padrão para **negrito**, *itálico* e `código`
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        try:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                paragraph.add_run(part)
        except Exception as e:
            # Em caso de erro, adiciona texto limpo
            paragraph.add_run(re.sub(r'[*`]', '', part))


def add_professional_table(doc, table_data):
    """Adiciona uma tabela profissional ao documento"""
    if not table_data or not any(table_data):
        return
    
    # Filtra linhas vazias
    table_data = [row for row in table_data if row and any(cell.strip() for cell in row)]
    
    if not table_data:
        return
    
    try:
        num_cols = max(len(row) for row in table_data)
        if num_cols == 0:
            return
            
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell_text = sanitize_text(cell_text)
                    cell.text = cell_text
                    
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(4)
                        paragraph.paragraph_format.space_after = Pt(4)
                        
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                            
                            if i == 0:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    
                    if i == 0:
                        set_cell_shading(cell, "1A3C6E")
                    elif i % 2 == 0:
                        set_cell_shading(cell, "F0F4F8")
        
        doc.add_paragraph()
    except Exception as e:
        print(f"{Fore.YELLOW}⚠️ Erro ao criar tabela: {e}")


def markdown_to_word(markdown_text, output_path, title="Apostila de Direito Administrativo"):
    """Converte Markdown para documento Word profissional"""
    
    print(f"{Fore.CYAN}📄 Gerando documento Word profissional...")
    
    doc = Document()
    
    # ============================================
    # CONFIGURAÇÃO DE ESTILOS
    # ============================================
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
    
    # Estilo do corpo do texto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(8)
    
    # Configurar estilos de Heading
    for i in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Arial'
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
            if i == 1:
                heading_style.font.size = Pt(18)
            elif i == 2:
                heading_style.font.size = Pt(14)
            elif i == 3:
                heading_style.font.size = Pt(12)
            else:
                heading_style.font.size = Pt(11)
        except:
            pass
    
    # ============================================
    # CAPA
    # ============================================
    
    for _ in range(6):
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Material de Estudo para Concursos")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    doc.add_paragraph()
    
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run("_" * 50)
    line_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    
    for _ in range(4):
        doc.add_paragraph()
    
    concurso_para = doc.add_paragraph()
    concurso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    concurso_run = concurso_para.add_run("PGM/PGE - Procuradorias")
    concurso_run.font.name = 'Arial'
    concurso_run.font.size = Pt(16)
    concurso_run.font.bold = True
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B de %Y").title())
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(12)
    date_run.font.italic = True
    
    doc.add_page_break()
    
    # ============================================
    # SUMÁRIO
    # ============================================
    
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run("SUMÁRIO")
    toc_run.font.name = 'Arial'
    toc_run.font.size = Pt(16)
    toc_run.font.bold = True
    toc_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    
    doc.add_paragraph()
    
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("(Atualize o sumário no Word: Ctrl+A, F9)")
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    
    doc.add_page_break()
    
    # ============================================
    # CABEÇALHO E RODAPÉ
    # ============================================
    
    section = doc.sections[0]
    
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(title)
    header_run.font.name = 'Arial'
    header_run.font.size = Pt(9)
    header_run.font.italic = True
    header_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)
    
    # ============================================
    # CONTEÚDO
    # ============================================
    
    lines = markdown_text.split('\n')
    current_table_data = []
    in_table = False
    line_count = 0
    total_lines = len(lines)
    
    print(f"{Fore.GREEN}   Processando {total_lines:,} linhas...")
    
    for line in lines:
        line_count += 1
        if line_count % 1000 == 0:
            print(f"   Progresso: {line_count:,}/{total_lines:,} linhas...")
        
        line_stripped = line.strip()
        
        if not line_stripped:
            if in_table and current_table_data:
                add_professional_table(doc, current_table_data)
                current_table_data = []
                in_table = False
            continue
        
        # Detecta tabelas markdown
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            in_table = True
            # Ignora linha de separação
            if not re.match(r'^[\|\-\:\s]+$', line_stripped):
                cells = [sanitize_text(cell.strip()) for cell in line_stripped.split('|')[1:-1]]
                if cells and any(c for c in cells):
                    current_table_data.append(cells)
            continue
        
        if in_table and current_table_data:
            add_professional_table(doc, current_table_data)
            current_table_data = []
            in_table = False
        
        try:
            # Headings
            if line_stripped.startswith('# '):
                doc.add_heading(sanitize_text(line_stripped[2:]), level=1)
            elif line_stripped.startswith('## '):
                doc.add_heading(sanitize_text(line_stripped[3:]), level=2)
            elif line_stripped.startswith('### '):
                doc.add_heading(sanitize_text(line_stripped[4:]), level=3)
            elif line_stripped.startswith('#### '):
                doc.add_heading(sanitize_text(line_stripped[5:]), level=4)
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, line_stripped[2:])
            elif re.match(r'^\d+\.\s', line_stripped):
                p = doc.add_paragraph(style='List Number')
                text = re.sub(r'^\d+\.\s*', '', line_stripped)
                add_formatted_text(p, text)
            elif line_stripped.startswith('>'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                add_formatted_text(p, line_stripped[1:].strip())
                for run in p.runs:
                    run.font.italic = True
            elif line_stripped.startswith('---'):
                # Linha horizontal - ignora
                pass
            elif line_stripped.startswith('[Fim'):
                # Marcadores de chunk - ignora
                pass
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, line_stripped)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Erro na linha {line_count}: {e}")
            try:
                p = doc.add_paragraph()
                p.add_run(sanitize_text(line_stripped))
            except:
                pass
    
    # Finaliza última tabela
    if in_table and current_table_data:
        add_professional_table(doc, current_table_data)
    
    # Salva o documento
    doc.save(output_path)
    print(f"{Fore.GREEN}✅ Documento Word salvo: {output_path}")


def main():
    if len(sys.argv) < 2:
        # Default: processa o arquivo Administrativo
        md_file = "/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/Administrativo_FORMATADO.md"
    else:
        md_file = sys.argv[1]
    
    if not os.path.exists(md_file):
        print(f"{Fore.RED}❌ Arquivo não encontrado: {md_file}")
        sys.exit(1)
    
    print(f"{Fore.CYAN}📂 Lendo arquivo Markdown: {md_file}")
    
    with open(md_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    print(f"{Fore.GREEN}   ✅ {len(markdown_text):,} caracteres lidos")
    
    # Gera caminho de saída
    output_docx = md_file.replace('.md', '.docx')
    
    markdown_to_word(markdown_text, output_docx)
    
    print(f"\n{Fore.GREEN}{'='*50}")
    print(f"{Fore.GREEN}✅ CONCLUÍDO!")
    print(f"{Fore.GREEN}{'='*50}")


if __name__ == "__main__":
    main()
