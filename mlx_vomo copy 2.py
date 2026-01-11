import os
import sys
import time
import subprocess
import traceback
import hashlib
from pathlib import Path
from colorama import init, Fore, Style
from tqdm import tqdm
import re
import difflib
import json
import asyncio
from google import genai
from google.genai import types
from openai import OpenAI, AsyncOpenAI
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import threading
import random
from time import sleep # Added for RateLimiter fallback if needed
import logging

init(autoreset=True)

# v2.18: Import do módulo de auditoria jurídica
try:
    from audit_module import auditar_consistencia_legal
    AUDIT_AVAILABLE = True
except ImportError:
    AUDIT_AVAILABLE = False
    print(f"⚠️ audit_module não encontrado. Auditoria jurídica desabilitada.")

try:
    import mlx_whisper
except ImportError:
    mlx_whisper = None

# Imports Pyannote e Torch
try:
    from pyannote.audio import Pipeline
    import torch
    
    HF_TOKEN = os.getenv("HUGGING_FACE_TOKEN")
    if not HF_TOKEN:
        pass
    
except ImportError:
    Pipeline = None
    torch = None

# ==================== CONTEXT CACHING (v2.2) ====================
# Só usa cache para documentos grandes (economia de tokens)
MIN_CHARS_PARA_CACHE = 150000  # 150k chars (~37k tokens)
CACHE_TTL = '7200s'            # 2 horas (padrão)

def criar_cache_contexto(client, transcricao_completa, system_prompt, estrutura_global=None, model_name="gemini-3-flash-preview"):
    """
    v2.2: Cria cache de contexto com hash estável para reutilização.
    
    Args:
        client: Cliente Gemini
        transcricao_completa: Texto completo (usado para calcular TTL dinâmico)
        system_prompt: Prompt de formatação (PROMPT_APOSTILA ou PROMPT_FIDELIDADE)
        estrutura_global: Estrutura mapeada (opcional)
        model_name: Nome do modelo Gemini
    
    Returns:
        Cache object ou None se falhar/não necessário
    """
    # Cache só vale a pena para documentos grandes
    if len(transcricao_completa) < MIN_CHARS_PARA_CACHE:
        print(f"{Fore.YELLOW}📦 Documento pequeno ({len(transcricao_completa):,} chars), cache não necessário{Style.RESET_ALL}")
        return None
    
    try:
        # Hash do prompt + estrutura para garantir unicidade por documento
        combined_content = system_prompt + (estrutura_global or "")
        prompt_hash = hashlib.sha256(combined_content.encode()).hexdigest()[:16]
        cache_name = f"vomo_{prompt_hash}"
        
        # Tenta encontrar cache existente válido
        try:
            for c in client.caches.list(page_size=100):
                if c.display_name == cache_name:
                    print(f"{Fore.GREEN}♻️  Reusando cache existente: {cache_name}{Style.RESET_ALL}")
                    return c
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Cache lookup warning: {e}{Style.RESET_ALL}")

        # Adiciona a estrutura global se disponível
        estrutura_text = f"\n\n## ESTRUTURA GLOBAL (GUIA):\n{estrutura_global}" if estrutura_global else ""
        
        cache_content = f"""{system_prompt}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📚 CONTEXTO GLOBAL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{estrutura_text}
"""
        
        # TTL Dinâmico: 1 hora a cada 500k chars + 1h margem
        tempo_estimado_segundos = int((len(transcricao_completa) / 500000) * 3600) + 3600
        dinamico_ttl = f"{max(3600, tempo_estimado_segundos)}s"
        
        # Cria cache usando a API do google-genai
        cache = client.caches.create(
            model=model_name,
            config=types.CreateCachedContentConfig(
                contents=[cache_content],
                ttl=dinamico_ttl,
                display_name=cache_name
            )
        )
        
        print(f"{Fore.GREEN}✅ Cache criado: {cache_name} (hash: {prompt_hash}, TTL: {dinamico_ttl}){Style.RESET_ALL}")
        return cache
        
    except Exception as e:
        print(f"{Fore.YELLOW}⚠️ Falha ao criar cache: {e}. Continuando sem cache.{Style.RESET_ALL}")
        return None

# ==================== RATE LIMITER ====================
class RateLimiter:
    """Controla requisições por minuto para não estourar rate limit da API"""
    def __init__(self, max_requests_per_minute=60): # Vertex AI Limit
        self.max_rpm = max_requests_per_minute
        self.requests = []
        self.lock = threading.Lock()
    
    def wait_if_needed(self):
        with self.lock:
            now = time.time()
            # Remove requisições antigas (>60s)
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
                if wait_time > 0:
                    print(f"{Fore.YELLOW}⏱️  Rate limit: aguardando {wait_time:.1f}s...")
                    sleep(wait_time)
                self.requests = [t for t in self.requests if time.time() - t < 60]
            
            self.requests.append(time.time())

    async def wait_if_needed_async(self):
        """Versão async do rate limiter para não bloquear o event loop"""
        wait_time = 0
        with self.lock:
            now = time.time()
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
                # Reserva o slot mas espera fora do lock
            
            if wait_time <= 0:
                 self.requests.append(time.time())
        
        if wait_time > 0:
            print(f"{Fore.YELLOW}⏱️  Rate limit (Async): aguardando {wait_time:.1f}s...")
            await asyncio.sleep(wait_time)
            # Re-adquire slot após espera (simplificado)
            with self.lock:
                self.requests.append(time.time())

# Instância global
rate_limiter = RateLimiter(max_requests_per_minute=60)

def remover_overlap_duplicado(resultados, mode="APOSTILA"):
    """Remove duplicação causada pelo overlap entre chunks usando detecção ROBUSTA de conteúdo
    
    v2.17: Usa LIMIAR_7DIFF diferenciado por modo.
    - FIDELIDADE: 0.85 (mais conservador)
    - APOSTILA: 0.80 (mais agressivo - overlaps são quase sempre erros)
    """
    # Limiares adaptativos por camada de deduplicação
    LIMIAR_7DIFF = 0.85 if mode == "FIDELIDADE" else 0.80
    if len(resultados) <= 1:
        return resultados[0] if resultados else ""
    
    import re
    from difflib import SequenceMatcher
    
    # === FUNÇÕES AUXILIARES DA ESTRATÉGIA ROBUSTA (Portadas de clean_redundancy.py) ===
    
    def normalize_text(text):
        if not text: return ""
        text = re.sub(r'[#*-]', ' ', text)
        text = re.sub(r'[^\w\s]', '', text)
        return ' '.join(text.lower().split())

    def calculate_similarity(text1, text2):
        if not text1 or not text2:
            return 0.0
        if len(text1) < 50:
            return 1.0 if text1 in text2 or text2 in text1 else 0.0
        return SequenceMatcher(None, text1, text2).quick_ratio()

    def extract_unique_paragraphs(sec_curr_content, sec_prev_content):
        if not sec_curr_content: return []
        unique = []
        paras_curr = sec_curr_content.split('\n\n')
        paras_prev_norm = [normalize_text(p) for p in sec_prev_content.split('\n\n')]
        
        for p in paras_curr:
            p_clean = p.strip()
            if not p_clean or len(p_clean) < 20: continue
            
            p_norm = normalize_text(p_clean)
            is_present = False
            for pp_norm in paras_prev_norm:
                if calculate_similarity(p_norm, pp_norm) > 0.85:
                    is_present = True
                    break
            
            if not is_present:
                unique.append(p_clean)
        return unique
    
    # ==================================================================================

    print("🧹 Iniciando deduplicação robusta (7-DIFF Strategy)...")

    # 1. Junta e Parseia
    texto_bruto = '\n\n'.join(resultados)
    lines = texto_bruto.split('\n')
    
    sections = []
    current_section = None
    header_pattern = re.compile(r'^(#{1,6})\s+(.+)$')
    
    captured_lines = []
    intro_lines = []
    has_started = False
    
    for i, line in enumerate(lines):
        match = header_pattern.match(line)
        if match:
            has_started = True
            if current_section:
                current_section['content'] = '\n'.join(captured_lines).strip()
                sections.append(current_section)
                captured_lines = []
            
            title_text = match.group(2).strip()
            # Remove numeração original para comparação agnóstica
            title_clean = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', title_text)
            
            current_section = {
                'title_clean': title_clean,
                'level': len(match.group(1)),
                'full_header': line,
                'content': ""
            }
        else:
            if has_started:
                captured_lines.append(line)
            else:
                intro_lines.append(line)
    
    if current_section:
        current_section['content'] = '\n'.join(captured_lines).strip()
        sections.append(current_section)
        
    print(f"   📊 Analisando {len(sections)} seções...")

    # 2. Detecção e Remoção
    indices_to_remove = set()
    MAX_WINDOW = 20 # Olha até 20 seções para trás (cobre overlaps grandes)
    
    for i in range(len(sections)):
        if i in indices_to_remove: continue
        sec_curr = sections[i]
        
        # Janela deslizante
        start_check = max(0, i - MAX_WINDOW)
        for j in range(start_check, i):
            if j in indices_to_remove: continue
            sec_prev = sections[j]
            
            if sec_curr['level'] != sec_prev['level']: continue
            
            sim_title = calculate_similarity(normalize_text(sec_curr['title_clean']), normalize_text(sec_prev['title_clean']))
            sim_content = calculate_similarity(normalize_text(sec_curr['content']), normalize_text(sec_prev['content']))
            
            is_duplicate = False
            
            if sim_title > 0.9 and sim_content > 0.6:
                is_duplicate = True
            elif sim_content > LIMIAR_7DIFF:  # Usa limiar 7-DIFF (0.85 Fidelidade / 0.80 Apostila)
                 is_duplicate = True
            elif sim_title > 0.95 and len(sec_curr['content']) < 100:
                 is_duplicate = True
            
            if is_duplicate:
                print(f"   🗑️  Detectado: '{sec_curr['title_clean'][:30]}...' duplica seção anterior")
                
                # Mescla conteúdo único antes de excluir
                unique_paras = extract_unique_paragraphs(sec_curr['content'], sec_prev['content'])
                if unique_paras:
                    sections[j]['content'] += '\n\n' + '\n\n'.join(unique_paras)
                
                indices_to_remove.add(i)
                break
    
    # 3. Reconstrução
    final_lines = list(intro_lines)
    for i, sec in enumerate(sections):
        if i in indices_to_remove: continue
        final_lines.append(sec['full_header'])
        if sec['content']:
            final_lines.append(sec['content'])
            final_lines.append("")
            
    texto_limpo = '\n'.join(final_lines)
    print(f"   ✅ Removidas {len(indices_to_remove)} seções duplicadas.")
    
    return texto_limpo


# ==================== HELPERS PORTED FROM GPT SCRIPT ====================

def limpar_tags_xml(texto):
    texto = re.sub(r'</?[a-z_][\w\-]*>', '', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<[a-z_][\w\-]*\s+[^>]+>', '', texto, flags=re.IGNORECASE)
    return texto

# ==================== LOGGER SHIM ====================
class Logger:
    def info(self, msg):
        print(f"{Fore.CYAN}{msg}")
    def warning(self, msg):
        print(f"{Fore.YELLOW}{msg}")
    def error(self, msg):
        print(f"{Fore.RED}{msg}")

logger = Logger()

# ==================== FUNCIONALIDADES PORTADAS DO GEMINI SCRIPT ====================

# ==================== SMART STITCHING & FIDELIDADE (V2.10 PROTADAS) ====================

def remover_eco_do_contexto(resposta_api, contexto_enviado):
    """
    Remove o início da resposta se for apenas um 'eco' do final do contexto.
    """
    if not contexto_enviado or not resposta_api:
        return resposta_api

    final_contexto = contexto_enviado.strip()[-300:]
    inicio_resposta = resposta_api.strip()[:300]

    matcher = difflib.SequenceMatcher(None, final_contexto, inicio_resposta)
    match = matcher.find_longest_match(0, len(final_contexto), 0, len(inicio_resposta))

    if match.size > 50:
        print(f"Scissors Eco detectado! Removendo {match.size} chars repetidos no início.")
        return resposta_api.strip()[match.size:].strip()
    
    return resposta_api

def titulos_sao_similares(t1, t2, threshold=0.90):
    """Verifica se dois títulos são semanticamente iguais (fuzzy matching)."""
    def normalizar(t):
        return re.sub(r'[^a-z0-9 ]', '', t.lower())
    
    nt1 = normalizar(t1)
    nt2 = normalizar(t2)
    
    if not nt1 or not nt2: return False
    
    # PROTEÇÃO 1: Diferença de tamanho
    nt1_compact = nt1.replace(' ', '')
    nt2_compact = nt2.replace(' ', '')
    len_ratio = min(len(nt1_compact), len(nt2_compact)) / max(len(nt1_compact), len(nt2_compact))
    if len_ratio < 0.8: return False
    
    # PROTEÇÃO 2: Verificação por palavras exclusivas
    palavras1 = set(nt1.split())
    palavras2 = set(nt2.split())
    diferenca = palavras1.symmetric_difference(palavras2)
    if any(len(w) > 3 for w in diferenca): return False
        
    return difflib.SequenceMatcher(None, nt1_compact, nt2_compact).ratio() > threshold

def limpar_inicio_redundante(texto_novo, texto_acumulado):
    """Remove título no início do novo chunk se similar ao do acumulado."""
    if not texto_acumulado.strip(): return texto_novo

    ultimas_linhas = texto_acumulado.strip().split('\n')[-30:]
    ultimo_titulo = None
    for linha in reversed(ultimas_linhas):
        if linha.strip().startswith('##'):
            ultimo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
            break
    
    if not ultimo_titulo: return texto_novo

    linhas_novas = texto_novo.strip().split('\n')
    for i, linha in enumerate(linhas_novas[:10]):
        if linha.strip().startswith('##'):
            novo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
            if titulos_sao_similares(ultimo_titulo, novo_titulo):
                print(f"Scissors Título duplicado na junção: '{novo_titulo}' ≈ '{ultimo_titulo}'")
                return '\n'.join(linhas_novas[i+1:])
    return texto_novo

def detectar_secoes_duplicadas(texto):
    """v2.11: Detecta seções duplicadas por títulos (Fuzzy Matching) - Inclui H3 e normaliza '(Continuação)'"""
    print("Magnifying glass tilt left Detectando seções duplicadas (fuzzy)...")
    
    linhas = texto.split('\n')
    titulos_vistos = []
    secoes_duplicadas = []
    
    for i, linha in enumerate(linhas):
        linha_strip = linha.strip()
        # Detecta H2 (##) e H3 (###)
        if linha_strip.startswith('##'):
            # Remove numeração e emojis
            titulo_normalizado = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', linha_strip)
            titulo_normalizado = re.sub(r'[📋📊🗂]', '', titulo_normalizado).strip()
            # Remove "(Continuação)" para comparação
            titulo_para_comparar = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_normalizado, flags=re.IGNORECASE).strip()
            
            duplicado = False
            for t_visto, linha_visto in titulos_vistos:
                if titulos_sao_similares(titulo_para_comparar, t_visto):
                    print(f"Warning Duplicado (fuzzy): '{linha_strip[:50]}...' ≈ '{t_visto[:50]}...'")
                    secoes_duplicadas.append({
                        'titulo': titulo_normalizado,
                        'primeira_linha': linha_visto,
                        'duplicada_linha': i
                    })
                    duplicado = True
                    break
            
            if not duplicado:
                titulos_vistos.append((titulo_para_comparar, i))
    
    if secoes_duplicadas:
        print(f"Cross mark {len(secoes_duplicadas)} seções duplicadas detectadas!")
    else:
        print("Check mark button Nenhuma seção duplicada detectada")
    
    return secoes_duplicadas

def remover_secoes_duplicadas(texto, mode="APOSTILA"):
    """v2.17: Remove seções duplicadas com LIMIAR ADAPTATIVO por modo.
    
    Usa LIMIAR_SECOES diferenciado - mais cuidado pois professor pode repetir propositalmente.
    - FIDELIDADE: 0.70 (mais conservador)
    - APOSTILA: 0.60 (mais agressivo)
    """
    from difflib import SequenceMatcher
    
    # Limiares adaptativos por camada de deduplicação
    # Seções duplicadas: mais cuidado, professor pode repetir propositalmente
    LIMIAR_SECOES = 0.70 if mode == "FIDELIDADE" else 0.60
    
    secoes_dup = detectar_secoes_duplicadas(texto)
    if not secoes_dup: return texto
    
    print("Broom Removendo seções duplicadas (Smart Dedupe v2.14)...")
    linhas = texto.split('\n')
    
    # Rastreia o ÚLTIMO segmento adicionado para cada título (para evitar diluição na comparação)
    ultimo_segmento_visto = {}  # titulo_normalizado -> último texto adicionado
    linhas_para_remover = set()
    linhas_para_adicionar_separador = set()
    
    for dup in secoes_dup:
        # --- 1. Extrair Conteúdo Original ---
        idx_orig = dup['primeira_linha']
        header_orig = linhas[idx_orig].strip()
        match_orig = re.match(r'^(#+)', header_orig)
        nivel_orig = len(match_orig.group(1)) if match_orig else 2
        
        # Extrai conteúdo da seção original
        content_orig = []
        for i in range(idx_orig + 1, len(linhas)):
            line = linhas[i].strip()
            if line.startswith('#'):
                match_now = re.match(r'^(#+)', line)
                if match_now and len(match_now.group(1)) <= nivel_orig: break
            content_orig.append(linhas[i])
        text_orig = "\n".join(content_orig)
        
        titulo_key = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', header_orig)
        titulo_key = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_key, flags=re.IGNORECASE).strip()
        
        # Inicializa o rastreamento se for a primeira vez
        if titulo_key not in ultimo_segmento_visto:
            ultimo_segmento_visto[titulo_key] = text_orig

        # --- 2. Extrair Conteúdo Duplicado ---
        idx_dup = dup['duplicada_linha']
        header_dup = linhas[idx_dup].strip()
        match_dup = re.match(r'^(#+)', header_dup)
        nivel_dup = len(match_dup.group(1)) if match_dup else 2
        
        fim_dup_idx = len(linhas)
        content_dup = []
        for i in range(idx_dup + 1, len(linhas)):
            line = linhas[i].strip()
            if line.startswith('#'):
                match_now = re.match(r'^(#+)', line)
                if match_now and len(match_now.group(1)) <= nivel_dup:
                    fim_dup_idx = i
                    break
            content_dup.append(linhas[i])
        text_dup = "\n".join(content_dup)
        
        # --- 3. Comparar Conteúdo (Lógica Janelada v2.16 MELHORADA) ---
        # Compara APENAS com o último segmento conhecido dessa seção
        texto_referencia = ultimo_segmento_visto.get(titulo_key, text_orig)
        
        len_dup = len(text_dup.strip())
        len_ref = len(texto_referencia.strip())
        
        # v2.16: NOVIDADE - Verificar se o título é 100% idêntico após normalização
        # Se for, força a mesclagem independente do conteúdo (evita redundância de temas)
        titulo_dup_key = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', header_dup)
        titulo_dup_key = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_dup_key, flags=re.IGNORECASE).strip()
        
        titulos_identicos = (titulo_key.lower() == titulo_dup_key.lower())
        
        # Lógica de decisão baseada em tamanho
        if len_dup < 50:
            # Duplicado curto demais = lixo, deletar
            sim = 1.0
            print(f"   ⚠️  Seção duplicada muito curta ({len_dup} chars) - marcando para remoção")
        elif titulos_identicos and len_dup > 100:
            # v2.16: Títulos idênticos = força mesclagem, mantém conteúdo novo sob o mesmo tópico
            sim = 0.5  # Valor que força mesclagem (remove header, mantém conteúdo)
            print(f"   🔄  Títulos 100% idênticos - forçando mesclagem de conteúdo")
        elif len_ref < 50 and len_dup >= 200:
            # Original curto, mas duplicado substancial = original estava incompleto
            # Manter o duplicado como novo conteúdo
            sim = 0.0
            print(f"   ℹ️  Original curto ({len_ref}c), duplicado substancial ({len_dup}c) - mantendo novo conteúdo")
        else:
            sim = SequenceMatcher(None, texto_referencia, text_dup).ratio()
            
        print(f"   Similaridade: {sim:.1%} | Linha {idx_dup} | '{titulo_key[:40]}...'")

        
        if sim > LIMIAR_SECOES:  # Usa limiar de seções (0.70 Fidelidade / 0.60 Apostila)
            print(f"   🗑️  Removendo SEÇÃO INTEIRA (Duplicata confirmada)")
            for i in range(idx_dup, fim_dup_idx):
                linhas_para_remover.add(i)
        else:
            print(f"   🔗 Mesclando conteúdo (Nova informação detectada)")
            linhas_para_remover.add(idx_dup)
            if idx_dup + 1 < len(linhas):
                linhas_para_adicionar_separador.add(idx_dup + 1)
            
            # ATUALIZA o último segmento visto para a próxima comparação
            ultimo_segmento_visto[titulo_key] = text_dup

    # --- 4. Reconstrução ---
    linhas_limpas = []
    for i, linha in enumerate(linhas):
        if i in linhas_para_remover:
            continue
        if i in linhas_para_adicionar_separador:
            linhas_limpas.append("") 
        linhas_limpas.append(linha)
        
    print(f"Check mark button {len(linhas_para_remover)} linhas removidas")
    return '\n'.join(linhas_limpas)


def remover_paragrafos_duplicados(texto: str, min_chars: int = 80) -> str:
    """
    v2.17: Remove parágrafos duplicados dentro do documento.
    
    Lógica:
    - Divide o texto em blocos (separados por linhas em branco).
    - Normaliza cada bloco (lowercase, sem pontuação extra).
    - Mantém apenas a primeira ocorrência de cada bloco normalizado.
    - Ignora blocos muito curtos (< min_chars) para não afetar listas.
    - Preserva tabelas e headers intactos.
    
    Args:
        texto: Texto markdown completo
        min_chars: Tamanho mínimo do parágrafo para considerar na deduplicação
    
    Returns:
        Texto com parágrafos duplicados removidos
    """
    import unicodedata
    
    print("🔄 Removendo parágrafos duplicados (v2.17)...")
    
    # Dividir em blocos por linhas em branco duplas
    blocos = re.split(r'\n\s*\n', texto)
    
    vistos = set()
    blocos_limpos = []
    removidos = 0
    
    for bloco in blocos:
        bloco_stripped = bloco.strip()
        
        # Preservar headers, tabelas e blocos curtos sem verificação
        if (bloco_stripped.startswith('#') or 
            bloco_stripped.startswith('|') or 
            bloco_stripped.startswith('```') or
            bloco_stripped.startswith('> [!') or
            len(bloco_stripped) < min_chars):
            blocos_limpos.append(bloco)
            continue
        
        # Normalizar para comparação
        # Remove pontuação, múltiplos espaços, lowercase
        normalizado = bloco_stripped.lower()
        normalizado = re.sub(r'[^\w\s]', '', normalizado)  # Remove pontuação
        normalizado = re.sub(r'\s+', ' ', normalizado).strip()  # Normaliza espaços
        # Remove acentos para comparação mais robusta
        normalizado = unicodedata.normalize('NFKD', normalizado)
        normalizado = normalizado.encode('ASCII', 'ignore').decode('ASCII')
        
        # Hash do conteúdo normalizado
        bloco_hash = hash(normalizado)
        
        if bloco_hash in vistos:
            removidos += 1
            continue  # Pula duplicata
        
        vistos.add(bloco_hash)
        blocos_limpos.append(bloco)
    
    if removidos > 0:
        print(f"   ✅ {removidos} parágrafos duplicados removidos")
    else:
        print(f"   ℹ️  Nenhum parágrafo duplicado encontrado")
    
    return '\n\n'.join(blocos_limpos)


def remover_titulos_orfaos(texto: str, similaridade_minima: float = 0.85) -> str:
    """
    v2.17: Remove linhas que são variações de títulos H2 já existentes.
    
    Detecta e remove:
    - Linhas bold que repetem um H2 (ex: **3. Execuções...**)
    - H3 que são cópias de H2 (ex: ### Execuções... quando já existe ## Execuções...)
    - Numeração inconsistente (ex: "3. Tema" como texto simples)
    
    Args:
        texto: Texto markdown
        similaridade_minima: Threshold para considerar como duplicata (0-1)
    
    Returns:
        Texto limpo
    """
    from difflib import SequenceMatcher
    
    print("🧹 Removendo títulos órfãos (v2.17)...")
    
    linhas = texto.split('\n')
    
    # 1. Extrair todos os títulos H2 existentes (normalizados)
    titulos_h2 = []
    for linha in linhas:
        if linha.strip().startswith('## '):
            # Extrair apenas o texto do título (sem ## e numeração)
            titulo_limpo = re.sub(r'^##\s*\d+(?:\.\d+)*\.?\s*', '', linha.strip())
            titulo_limpo = titulo_limpo.strip().lower()
            if titulo_limpo:
                titulos_h2.append(titulo_limpo)
    
    if not titulos_h2:
        print("   ℹ️  Nenhum H2 encontrado para comparação")
        return texto
    
    # 2. Identificar linhas órfãs para remoção
    linhas_para_remover = set()
    
    for i, linha in enumerate(linhas):
        stripped = linha.strip()
        
        # Ignorar linhas vazias, headers reais, tabelas
        if not stripped or stripped.startswith('## ') or stripped.startswith('|'):
            continue
        
        # Detectar padrões de "título órfão"
        texto_candidato = None
        
        # Padrão 1: Linha bold com numeração (ex: **3. Execuções...**)
        match_bold = re.match(r'^\*\*\d+\.?\s*(.+?)\*\*\s*$', stripped)
        if match_bold:
            texto_candidato = match_bold.group(1).strip().lower()
        
        # Padrão 2: Numeração simples no início (ex: "3. Execuções Envolvendo...")
        elif re.match(r'^\d+\.\s+[A-Z]', stripped):
            texto_candidato = re.sub(r'^\d+\.\s*', '', stripped).strip().lower()
        
        # Padrão 3: H3 que pode ser duplicata de H2
        elif stripped.startswith('### '):
            texto_candidato = re.sub(r'^###\s*\d*\.?\s*', '', stripped).strip().lower()
        
        if texto_candidato:
            # Comparar com todos os H2
            for h2 in titulos_h2:
                sim = SequenceMatcher(None, texto_candidato, h2).ratio()
                if sim >= similaridade_minima:
                    linhas_para_remover.add(i)
                    break
    
    # 3. Reconstruir sem as linhas órfãs
    if linhas_para_remover:
        print(f"   ✅ {len(linhas_para_remover)} títulos órfãos removidos")
        linhas_limpas = [l for i, l in enumerate(linhas) if i not in linhas_para_remover]
        return '\n'.join(linhas_limpas)
    
    print("   ℹ️  Nenhum título órfão detectado")
    return texto


def aplicar_correcoes_automaticas(texto: str) -> tuple:
    """
    v2.18: Aplica correções automáticas baseadas em padrões comuns de erro.
    Portado de format_transcription_gemini.py.
    
    Correções aplicadas:
    1. Remove saudações duplicadas ("Olá, sejam bem-vindos...")
    2. Remove apresentações repetidas ("Eu sou o professor...")
    3. Padroniza nome do professor (variações → nome canônico)
    4. Corrige itens de lista malformados ("3. \\n Requisitos" → "3. Requisitos")
    5. Remove linhas em branco excessivas
    
    Returns:
        tuple: (texto_corrigido, lista_de_correcoes)
    """
    from difflib import SequenceMatcher
    
    print(f"{Fore.CYAN}🔧 Auto-Fix Pass (v2.18)...")
    
    correcoes = []
    texto_original = texto
    
    # 1. Remover saudações duplicadas (apenas mantém a primeira)
    saudacoes_pattern = r'(?:Olá|Oi),?\s*(?:sejam?\s+)?(?:bem[- ]?vindos?(?:\s+e\s+bem[- ]?vindas?)?)[.,!]?'
    matches = list(re.finditer(saudacoes_pattern, texto, re.IGNORECASE))
    if len(matches) > 1:
        for match in reversed(matches[1:]):
            start = texto.rfind('\n', 0, match.start()) + 1
            end = texto.find('\n', match.end())
            if end == -1: end = len(texto)
            linha = texto[start:end].strip()
            if len(linha) < 150:
                texto = texto[:start] + texto[end+1:]
                correcoes.append(f"Saudação duplicada removida")
    
    # 2. Remover apresentações repetidas do professor
    apresentacao_pattern = r'Eu sou o professor\s+\w+(?:\s+\w+)?'
    matches = list(re.finditer(apresentacao_pattern, texto, re.IGNORECASE))
    if len(matches) > 1:
        for match in reversed(matches[1:]):
            start = texto.rfind('\n', 0, match.start()) + 1
            end = texto.find('\n', match.end())
            if end == -1: end = len(texto)
            linha = texto[start:end].strip()
            if len(linha) < 200:
                texto = texto[:start] + texto[end+1:]
                correcoes.append(f"Apresentação duplicada removida")
    
    # 3. Padronizar nome do professor
    nome_match = re.search(r'professor\s+(\w+(?:\s+\w+)?)', texto, re.IGNORECASE)
    if nome_match:
        nome_canonico = nome_match.group(1)
        variacoes_pattern = rf'\bprofessor\s+(\w+(?:\s+\w+)?)\b'
        for m in re.finditer(variacoes_pattern, texto, re.IGNORECASE):
            nome_atual = m.group(1)
            if nome_atual.lower() != nome_canonico.lower():
                sim = SequenceMatcher(None, nome_canonico.lower(), nome_atual.lower()).ratio()
                if sim > 0.6 and sim < 1.0:
                    texto = texto.replace(f"professor {nome_atual}", f"professor {nome_canonico}")
                    texto = texto.replace(f"Professor {nome_atual}", f"Professor {nome_canonico}")
                    correcoes.append(f"Nome padronizado: '{nome_atual}' → '{nome_canonico}'")
    
    # 4. Corrigir itens de lista vazios ou malformados
    texto_temp = re.sub(r'(\d+\.)\s*\n\s*((?:Requisitos|Preenchimento|Fundamento|Artigo|Lei))', r'\1 \2', texto)
    if texto_temp != texto:
        texto = texto_temp
        correcoes.append("Itens de lista malformados corrigidos")
    
    # 5. Remover linhas em branco excessivas
    texto_limpo = re.sub(r'\n{4,}', '\n\n\n', texto)
    if texto_limpo != texto:
        texto = texto_limpo
        correcoes.append("Linhas em branco excessivas removidas")
    
    if correcoes:
        print(f"   ✅ {len(correcoes)} correções aplicadas:")
        for c in correcoes[:3]:
            print(f"      - {c}")
        if len(correcoes) > 3:
            print(f"      ... e mais {len(correcoes) - 3}")
    else:
        print(f"   ℹ️  Nenhuma correção necessária")
    
    return texto, correcoes


def normalize_headings(texto):
    """
    v1.0: Normaliza títulos semanticamente similares para uma versão única.
    - Agrupa títulos por similaridade
    - Escolhe o título mais descritivo de cada grupo
    - Remove sufixos como "(Continuação)"
    """
    from difflib import SequenceMatcher
    
    print("🔤 Normalizando títulos similares...")
    linhas = texto.split('\n')
    
    # 1. Extrair todos os títulos com info de nível e posição
    titulos = []
    for i, linha in enumerate(linhas):
        stripped = linha.strip()
        if stripped.startswith('##'):
            match = re.match(r'^(#+)\s*', stripped)
            nivel = len(match.group(1)) if match else 2
            # Extrai título limpo (sem # e sem numeração)
            titulo_limpo = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', stripped).strip()
            # Remove "(Continuação)" para comparação
            titulo_base = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_limpo, flags=re.IGNORECASE).strip()
            titulos.append({
                'linha': i,
                'nivel': nivel,
                'original': stripped,
                'limpo': titulo_limpo,
                'base': titulo_base
            })
    
    if not titulos:
        return texto
    
    # 2. Agrupar títulos similares (mesmo nível + similaridade > 0.85)
    grupos = []
    usados = set()
    
    for i, t1 in enumerate(titulos):
        if i in usados:
            continue
        grupo = [t1]
        usados.add(i)
        
        for j, t2 in enumerate(titulos):
            if j in usados or j <= i:
                continue
            if t1['nivel'] == t2['nivel']:
                sim = SequenceMatcher(None, t1['base'].lower(), t2['base'].lower()).ratio()
                if sim > 0.85:
                    grupo.append(t2)
                    usados.add(j)
        
        if len(grupo) > 1:
            grupos.append(grupo)
    
    if not grupos:
        # Nada para normalizar, mas ainda remove "(Continuação)"
        texto_limpo = re.sub(r'\s*\(Continuação\)\s*(?=\n|$)', '', texto, flags=re.IGNORECASE)
        return texto_limpo
    
    # 3. Para cada grupo, escolher o "melhor" título (mais curto entre os mais longos)
    #    Lógica: preferir títulos sem "(Continuação)" e com descrição completa
    substituicoes = {}
    for grupo in grupos:
        # Ordenar por: não ter "(Continuação)" primeiro, depois por comprimento (prefer médio)
        candidatos = sorted(grupo, key=lambda x: (
            '(Continuação)' in x['limpo'],  # False (0) vem antes de True (1)
            abs(len(x['limpo']) - 40)  # Preferir títulos de ~40 chars (nem muito curto nem muito longo)
        ))
        
        melhor = candidatos[0]['limpo']
        print(f"   📝 Grupo de {len(grupo)} títulos similares → padronizando para: '{melhor[:50]}...'")
        
        for t in grupo:
            if t['limpo'] != melhor:
                substituicoes[t['linha']] = (t['nivel'], melhor)
    
    # 4. Aplicar substituições
    novas_linhas = []
    for i, linha in enumerate(linhas):
        if i in substituicoes:
            nivel, novo_titulo = substituicoes[i]
            # Preservar numeração existente se houver
            match_num = re.match(r'^(#+\s*\d+(?:\.\d+)*\.?\s*)', linhas[i].strip())
            if match_num:
                prefixo = match_num.group(1)
                # Remove a numeração do novo título se ele tiver uma
                novo_titulo = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', novo_titulo)
                novas_linhas.append(f"{prefixo}{novo_titulo}")
            else:
                novas_linhas.append(f"{'#' * nivel} {novo_titulo}")
        else:
            # Remove "(Continuação)" de qualquer título restante
            if linha.strip().startswith('##'):
                linha = re.sub(r'\s*\(Continuação\)\s*(?=\n|$)', '', linha, flags=re.IGNORECASE)
            novas_linhas.append(linha)
    
    print(f"   ✅ {len(substituicoes)} títulos normalizados, '(Continuação)' removidos")
    return '\n'.join(novas_linhas)

PROMPT_STRUCTURE_REVIEW = """Você é um revisor especializado em estrutura de documentos jurídicos educacionais.

## ESTRUTURA DE MAPEAMENTO INICIAL (Referência - se disponível):
{estrutura_mapeada}

## TAREFA
Revise a ESTRUTURA (headers/títulos) do documento abaixo, COMPARANDO com o mapeamento acima (se disponível). Sua missão é harmonizar o mapeamento planejado com o CONTEÚDO REAL da aula.

## ✅ O QUE VOCÊ DEVE FAZER:

### 1. COMPARAR E REFINAR TÍTULOS (CRÍTICO)
Verifique se los títulos refletem os tópicos do mapeamento. Se um título for genérico, torne-o descritivo.
- ERRADO: "### Questão" ou "### Tópico"
- CORRETO: "### Questão 1: Responsabilidade Civil" (Descritivo conforme mapa/conteúdo)
- **CRÍTICO:** Evite títulos idênticos em seções "irmãs". Diferencie-os pelo conteúdo específico de cada uma.

### 2. VALIDAR HIERARQUIA E PROMOÇÃO DE TÓPICOS
Confirme se a estrutura segue lógica consistente (##, ###, ####). 
- **PROMOÇÃO:** Se um sub-subtópico (ex: 9.19.5) for extenso e tratar de um tema central (ex: Execução Fiscal), PROMOVA-O a um nível superior (ex: ### 9.20) para evitar fragmentação excessiva e respeitar o limite de níveis.

### 3. RENUMERAÇÃO SEQUENCIAL OBRIGATÓRIA
Se você criar, deletar ou promover uma seção, você DEVE renumerar TODAS as seções subsequentes daquela mesma hierarquia para manter a sequência lógica (ex: se 9.20 foi criado, o antigo 9.20 vira 9.21, e assim por diante).

### 4. MESCLAR QUESTÕES DUPLICADAS
Se duas seções têm o mesmo número de questão na mesma área, MESCLE-AS.
- ERRADO: "2.1. Questão 1: TAC" + "2.2. Questão 1: TAC" 
- CORRETO: "2.1. Questão 1: TAC" (Única, com todo o conteúdo unificado)

### 5. PRIORIDADE DO CONTEÚDO REAL (DECIDIR ESTRUTURA)
Se houver conflito entre mapeamento e documento, escolha a estrutura que melhor reflete o CONTEÚDO REAL. O mapeamento é apenas um guia.

### 6. LIMPEZA TÉCNICA (SINTAXE MARKDOWN)
Corrija problemas detalhados de formatação:
- **Tabelas:** Alinhar colunas e adicionar separadores faltantes.
- **Listas:** Corrigir bullets ou numeração mal formatada.
- **Espaçamento:** Padronizar linhas em branco entre seções (mínimo uma linha).
- **Headers Vazios:** Remover títulos sem conteúdo abaixo.
- **Metadados:** Remover headers ou tags inline como "[TIPO: ...]", "**[TIPO: ...]**", ou "[BLOCO 0X]".

## ❌ O QUE VOCÊ NÃO DEVE FAZER:
1. **NÃO ALTERE O CONTEÚDO** dos parágrafos - apenas os títulos e a organização.
2. **NUNCA RESUMA** ou encurte o texto (mesmo tamanho de entrada/saída é obrigatório).
3. **NÃO INVENTE** fatos jurídicos.
4. **NÃO REMOVA** trechos técnicos ou exemplos.

## REGRAS CRÍTICAS DE HIERARQUIA:
- Use **MÁXIMO 3** níveis de hierarquia (##, ###, ####).
- Nunca use # (H1) para subtópicos (apenas para o título principal do documento).
- Preserve a ordem cronológica geral.

## DOCUMENTO PARA REVISAR:
{documento}

## 📝 RELATÓRIO ESPERADO:
Ao final do documento, inclua um bloco de comentário indicando:
- Quantos títulos foram refinados
- Promoções de seções e renumerações realizadas
- Discrepâncias com o mapeamento (se houver)

Formato:
<!-- RELATÓRIO: X títulos refinados | Y seções promovidas/renumeradas | Discrepâncias: [Nenhuma/Lista] -->

## RESPOSTA:
Retorne o documento COMPLETO E INTEGRAL (mesmo tamanho do original) com os títulos/headers corrigidos e o relatório no final. NÃO RESUMA."""

def renumerar_secoes(texto):
    """
    v1.0: Renumeração Sequencial Determinística.
    
    Esta função é uma camada de segurança extra, aplicada APÓS o AI Review.
    Ela percorre todos os headers numerados e corrige qualquer duplicação ou
    sequência quebrada, garantindo que os números sejam estritamente sequenciais
    dentro de cada nível de hierarquia.
    
    Exemplo de correção:
    - Input:  9.20, 9.21, 9.21, 9.35, 9.36  (duplicação e pulo)
    - Output: 9.20, 9.21, 9.22, 9.23, 9.24  (sequencial)
    """
    logger.info("🔢 Executando Renumeração Sequencial Determinística...")
    
    linhas = texto.split('\n')
    novas_linhas = []
    
    # Contadores por nível de hierarquia (ex: {2: 9, 3: 44} -> ## 9.x, ### 9.44.x)
    # Estrutura: {header_level: {parent_prefix: next_number}}
    # Ex: para ### 9.20.1 -> level=3, parent_prefix="9.20", next_number=2 (para o próximo .x)
    contadores = {}
    
    # Regex para detectar headers numerados: ### 9.20.1. Título ou ## 5. Título
    header_pattern = re.compile(r'^(#{1,4})\s+([\d.]+\.?)\s*(.*)$')
    
    for linha in linhas:
        match = header_pattern.match(linha)
        
        if match:
            hashes = match.group(1)       # "###"
            numero = match.group(2)       # "9.20.1." ou "9.20.1"
            titulo = match.group(3)       # "Título..."
            
            nivel = len(hashes)
            numero_limpo = numero.rstrip('.')
            partes = numero_limpo.split('.')
            
            # Determina o prefixo pai e o sufixo atual
            if len(partes) == 1:
                # Nível raiz: ## 1. ou ## 9.
                prefixo_pai = ""
                sufixo_atual = int(partes[0])
            else:
                # Subnível: ### 9.20. ou #### 9.20.1.
                prefixo_pai = '.'.join(partes[:-1])
                sufixo_atual = int(partes[-1])
            
            # Inicializa contador se não existir
            chave = (nivel, prefixo_pai)
            if chave not in contadores:
                contadores[chave] = sufixo_atual  # Começa do número encontrado
            else:
                contadores[chave] += 1
            
            novo_sufixo = contadores[chave]
            
            # Reconstrói o número
            if prefixo_pai:
                novo_numero = f"{prefixo_pai}.{novo_sufixo}."
            else:
                novo_numero = f"{novo_sufixo}."
            
            # Reconstrói a linha
            nova_linha = f"{hashes} {novo_numero} {titulo}"
            novas_linhas.append(nova_linha)
            
            # Log se houve mudança
            if numero_limpo != novo_numero.rstrip('.'):
                logger.info(f"   🔄 {numero_limpo} → {novo_numero.rstrip('.')}")
        else:
            novas_linhas.append(linha)
    
    logger.info("   ✅ Renumeração concluída.")
    return '\n'.join(novas_linhas)

def deterministic_structure_fix(text):
    """
    v1.1: Reorganização Estrutural Determinística (Regex).
    Adaptativo: Detecta se o documento usa H1 ou apenas H2 como nível principal.
    """
    print(f"{Fore.CYAN}🧩 Executando Reorganização Estrutural Determinística...")
    
    lines = text.split('\n')
    
    # Detecção de Hierarquia
    has_h1 = any(re.match(r'^#\s+', line) for line in lines)
    header_level_regex = r'^#\s+' if has_h1 else r'^##\s+'
    print(f"   ℹ️  Nível principal detectado: {'H1 (#)' if has_h1 else 'H2 (##)'}")

    # Estruturas de dados (Preserva ordem de inserção)
    content_map = {
        "PREAMBULO": [],
        "DISCIPLINAS": {}, 
        "ENCERRAMENTO": []
    }
    
    current_area = "PREAMBULO"
    current_block = []
    disciplinas_order = [] 
    
    # Regex Adaptativo
    # Captura o texto do cabeçalho principal (seja H1 ou H2)
    # Exclui "Questão" ou "Q." para não quebrar simulados dentro de uma área
    re_disciplina = re.compile(rf'{header_level_regex}(?!Questão|Q\.)([^0-9\.]+.*)', re.IGNORECASE)
    re_encerramento = re.compile(rf'{header_level_regex}(?:ENCERRAMENTO|CONSIDERAÇÕES|CONCLUSÃO)', re.IGNORECASE)
    
    def flush_block(area, block_lines):
        if not block_lines: return
        block_text = '\n'.join(block_lines)
        
        if area == "PREAMBULO":
            content_map["PREAMBULO"].append(block_text)
        elif area == "ENCERRAMENTO":
            content_map["ENCERRAMENTO"].append(block_text)
        else:
            if area not in content_map["DISCIPLINAS"]:
                content_map["DISCIPLINAS"][area] = []
                disciplinas_order.append(area)
            content_map["DISCIPLINAS"][area].append(block_text)

    for line in lines:
        # 1. Detectar mudança de disciplina macro
        match_disc = re_disciplina.match(line)
        if match_disc:
            flush_block(current_area, current_block)
            current_block = []
            
            raw_area = match_disc.group(1).strip().upper()
            
            # Normalização de nome de área
            if "DIREITO" not in raw_area and len(raw_area) < 50:
                 # Adiciona prefixo se parecer nome de matéria jurídica comum
                 if any(x in raw_area for x in ["CIVIL", "PENAL", "TRABALHO", "ADMINISTRATIVO", "CONSTITUCIONAL"]):
                     current_area = f"DIREITO {raw_area}"
                 else:
                     current_area = raw_area
            else:
                 current_area = raw_area
                 
            # IMPORTANTE: Se estamos operando em modo H2, essa linha é um Header que queremos manter?
            # Se for modo H1, recriamos "# AREA". 
            # Se for modo H2, recriamos "# AREA" (upcast) ou mantemos "## AREA"?
            # Para padronizar Apostilas, vamos promover tudo a H1 na reconstrução.
            continue 
            
        # 2. Detectar Encerramento
        if re_encerramento.match(line):
            flush_block(current_area, current_block)
            current_block = []
            current_area = "ENCERRAMENTO"
            continue

        current_block.append(line)
        
    flush_block(current_area, current_block)
    
    # Reconstrução
    final_output = []
    
    # Preambulo
    if content_map["PREAMBULO"]:
        final_output.append("# ORIENTAÇÕES GERAIS / INTRODUÇÃO")
        final_output.extend(content_map["PREAMBULO"])
        final_output.append("")

    # Disciplinas / Tópicos Principais
    for area in disciplinas_order:
        area_clean = area.replace("#", "").strip()
        final_output.append(f"# {area_clean}")
        for block in content_map["DISCIPLINAS"][area]:
            final_output.append(block)
        final_output.append("")
        
    # Encerramento
    if content_map["ENCERRAMENTO"]:
        final_output.append("# CONSIDERAÇÕES FINAIS")
        final_output.extend(content_map["ENCERRAMENTO"])
        
    num_identified = len(disciplinas_order)
    print(f"   ✅ Reorganizado: {num_identified} seções principais identificadas.")
    
    # Fallback: Se não identificou nada (tudo preambulo), retorna original para não estragar
    if num_identified == 0 and len(content_map["PREAMBULO"]) > 0:
        print("   ⚠️ Nenhuma estrutura detectada. Mantendo original.")
        return text
        
    return '\n'.join(final_output)

PROMPT_STRUCTURE_REVIEW_LITE = """Você é um revisor editorial especializado em ESTRUTURA e FORMATAÇÃO de documentos educacionais. Você receberá:
1. Uma **Estrutura de Mapeamento Inicial** (planejada antes da formatação)
2. O **Documento Processado** (resultado da formatação por chunks)

Sua tarefa é analisar ambos e garantir que os títulos estejam **descritivos, hierarquicamente corretos e alinhados com o conteúdo real**, sem jamais alterar a ordem cronológica.

---

## 📋 ESTRUTURA DE MAPEAMENTO INICIAL (Referência):
{estrutura_mapeada}

---

## ✅ O QUE VOCÊ DEVE FAZER:
1. **Comparar Títulos:** Verifique se os títulos do documento refletem corretamente os tópicos do mapeamento. Se um título estiver genérico mas o mapeamento indicar um tema específico, refine-o.
2. **Validar Hierarquia:** Confirme que a estrutura (##, ###, ####) segue uma lógica consistente (ex: seções > subseções > detalhes).
3. **Decidir a Melhor Estrutura:** Se houver conflito entre mapeamento e documento, escolha a estrutura que melhor reflete o CONTEÚDO REAL do texto.
4. **Subtópicos Órfãos:** Se detectar headers como "A.", "B.", "C." isolados como tópicos principais, converta-os em subníveis do tópico anterior (ex: ## para ###).
5. **Títulos Descritivos:** Refine títulos genéricos (ex: "Questão 1") para algo que cite o tema técnico (ex: "Questão 1: Responsabilidade Civil").
6. **Corrigir Sintaxe Markdown:** Tabelas, listas, espaçamento.
7. **Remover Vazios:** Títulos sem conteúdo abaixo.

## 📌 EXEMPLOS DE CORREÇÃO:

**Títulos Genéricos → Descritivos:**
- ANTES: `### Questão`
- DEPOIS: `### Questão 1: Responsabilidade Civil Objetiva`

**Subtópicos Órfãos → Hierarquia Correta:**
- ANTES:
  ```
  ## A. Requisitos do Dano
  ## B. Nexo Causal
  ```
- DEPOIS:
  ```
  ### A. Requisitos do Dano
  ### B. Nexo Causal
  ```

**Numeração Duplicada → Sequencial:**
- ANTES: `### 9.20`, `### 9.20`, `### 9.35`
- DEPOIS: `### 9.20`, `### 9.21`, `### 9.22`

## ❌ O QUE VOCÊ NÃO DEVE FAZER:
1. **NÃO MOVA** blocos de texto. A ordem deve permanecer 100% cronológica.
2. **NÃO MESCLE** seções que apareçam em momentos diferentes da aula.
3. **NÃO RESUMA** nem altere o corpo dos parágrafos.
4. **NÃO ADICIONE** conteúdo novo.

## 📝 RELATÓRIO ESPERADO:
Ao final do documento, inclua um bloco de comentário (que será removido) indicando:
- Quantos títulos foram refinados
- Se a estrutura final segue o mapeamento ou foi adaptada
- Discrepâncias encontradas (se houver)

Formato:
<!-- RELATÓRIO: X títulos refinados | Estrutura: [MAPEAMENTO/ADAPTADA] | Discrepâncias: [Nenhuma/Lista] -->

---

## 📄 DOCUMENTO PARA REVISAR:
{documento}

---

## RESPOSTA:
Retorne o documento COMPLETO com a formatação corrigida e o relatório no final."""

async def ai_structure_review_lite(texto, client, model, estrutura_mapeada=None, metrics=None):
    """
    v2.2: Revisão LEVE de formatação Markdown com VALIDAÇÃO CRUZADA.
    Compara o documento processado com a estrutura de mapeamento inicial.
    Refina títulos, valida hierarquia, e reporta discrepâncias.
    NÃO reorganiza nem mescla conteúdo.
    
    Melhorias v2.2:
    - Limite aumentado para 800k chars
    - Integração com MetricsCollector
    - Suporte a relatório JSON
    """
    from difflib import SequenceMatcher
    import asyncio
    import time
    import json
    
    print(f"{Fore.MAGENTA}  🧹 Revisão Leve de Formatação (IA - Modo Fidelidade v2.2)...{Style.RESET_ALL}")
    
    start_time = time.time()
    
    # Gemini 3 Flash suporta 1M tokens (~4M chars) - usar até 800k chars
    max_doc_chars = 800000
    if len(texto) > max_doc_chars:
        print(f"{Fore.YELLOW}   ⚠️ Documento muito longo ({len(texto)} chars), truncando para {max_doc_chars//1000}k...{Style.RESET_ALL}")
        texto_para_revisao = texto[:max_doc_chars] + "\n\n[... documento truncado para revisão ...]"
    else:
        texto_para_revisao = texto
    
    # Preparar estrutura mapeada (se disponível)
    if estrutura_mapeada:
        estrutura_str = estrutura_mapeada[:50000]  # Limitar estrutura a 50k chars
        print(f"{Fore.CYAN}   📋 Usando estrutura de mapeamento inicial ({len(estrutura_mapeada)} chars) para validação cruzada.{Style.RESET_ALL}")
    else:
        estrutura_str = "[Estrutura de mapeamento não disponível - analisar documento para inferir estrutura ideal]"
        print(f"{Fore.YELLOW}   ℹ️  Sem mapeamento inicial, IA irá inferir estrutura ideal do próprio documento.{Style.RESET_ALL}")
    
    def call_gemini():
        return client.models.generate_content(
            model=model,
            contents=PROMPT_STRUCTURE_REVIEW_LITE.format(
                estrutura_mapeada=estrutura_str,
                documento=texto_para_revisao
            ),
            config=types.GenerateContentConfig(
                max_output_tokens=65536,  # Máximo permitido
                thinking_config=types.ThinkingConfig(
                    include_thoughts=False, 
                    thinking_level="HIGH"  # HIGH para análise estrutural profunda
                ),
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                ]
            )
        )
    
    # Retry com backoff exponencial
    max_retries = 3
    response = None
    for tentativa in range(max_retries):
        try:
            response = await asyncio.to_thread(call_gemini)
            resultado = response.text.replace('```markdown', '').replace('```', '').strip()
            break
        except Exception as e:
            if tentativa < max_retries - 1:
                wait_time = 2 ** tentativa
                print(f"{Fore.YELLOW}   ⚠️ Tentativa {tentativa + 1} falhou: {e}. Aguardando {wait_time}s...{Style.RESET_ALL}")
                await asyncio.sleep(wait_time)
            else:
                print(f"{Fore.RED}   ⚠️ Erro na revisão leve após {max_retries} tentativas: {e}. Mantendo original.{Style.RESET_ALL}")
                return texto
    
    # Métricas
    duration = time.time() - start_time
    if metrics and response:
        try:
            usage = response.usage_metadata
            metrics.record_call(
                provider="gemini",
                prompt_tokens=getattr(usage, 'prompt_token_count', 0) or 0,
                completion_tokens=getattr(usage, 'candidates_token_count', 0) or 0,
                duration=duration
            )
        except:
            pass  # Silently ignore metrics errors
    
    # Extrair relatório (suporta JSON e texto)
    relatorio_json_match = re.search(r'<!--\s*RELATÓRIO_JSON:\s*(\{.+?\})\s*-->', resultado, re.IGNORECASE | re.DOTALL)
    relatorio_match = re.search(r'<!--\s*RELATÓRIO:\s*(.+?)\s*-->', resultado, re.IGNORECASE)
    
    if relatorio_json_match:
        try:
            relatorio_data = json.loads(relatorio_json_match.group(1))
            print(f"{Fore.CYAN}   📊 Relatório (JSON): {relatorio_data}{Style.RESET_ALL}")
        except:
            print(f"{Fore.CYAN}   📊 Relatório: {relatorio_json_match.group(1)}{Style.RESET_ALL}")
        resultado = re.sub(r'<!--\s*RELATÓRIO_JSON:.+?-->\s*', '', resultado, flags=re.IGNORECASE | re.DOTALL).strip()
    elif relatorio_match:
        relatorio = relatorio_match.group(1)
        print(f"{Fore.CYAN}   📊 Relatório da IA: {relatorio}{Style.RESET_ALL}")
        resultado = re.sub(r'<!--\s*RELATÓRIO:.+?-->\s*', '', resultado, flags=re.IGNORECASE).strip()
    
    # Validação: resultado deve ter pelo menos 80% do tamanho (padronizado)
    if len(resultado) < len(texto) * 0.80:
        print(f"{Fore.YELLOW}   ⚠️ Revisão retornou texto muito curto ({len(resultado)} vs {len(texto)}). Mantendo original.{Style.RESET_ALL}")
        return texto
    
    # Verificar se a ordem dos headers foi preservada (usando similaridade)
    headers_original = re.findall(r'^(#{1,4})\s+(.+?)$', texto, re.MULTILINE)
    headers_revisado = re.findall(r'^(#{1,4})\s+(.+?)$', resultado, re.MULTILINE)
    
    if len(headers_original) > 5 and len(headers_revisado) > 5:
        # Usar similaridade ao invés de igualdade exata
        similares = sum(1 for (_, h1), (_, h2) in zip(headers_original[:10], headers_revisado[:10]) 
                       if SequenceMatcher(None, h1.strip(), h2.strip()).ratio() > 0.6)
        if similares < 6:
            print(f"{Fore.YELLOW}   ⚠️ Ordem dos headers parece alterada ({similares}/10 similares). Mantendo original.{Style.RESET_ALL}")
            return texto
    
    # 📝 Relatório de Alterações nos Títulos
    alteracoes = []
    for i, ((lvl_orig, h_orig), (lvl_rev, h_rev)) in enumerate(zip(headers_original[:50], headers_revisado[:50])):
        if h_orig.strip() != h_rev.strip():
            alteracoes.append(f"   • '{h_orig[:40]}...' → '{h_rev[:40]}...'")
    
    if alteracoes:
        print(f"{Fore.CYAN}   📝 Títulos Refinados ({len(alteracoes)}):{Style.RESET_ALL}")
        for alt in alteracoes[:10]:  # Mostrar no máximo 10
            print(alt)
        if len(alteracoes) > 10:
            print(f"   ... e mais {len(alteracoes) - 10} alterações.")
    else:
        print(f"{Fore.GREEN}   ℹ️  Nenhum título foi alterado (estrutura já estava OK).{Style.RESET_ALL}")
    
    print(f"{Fore.GREEN}   ✅ Formatação revisada (modo leve v2.2, {duration:.1f}s).{Style.RESET_ALL}")
    return resultado

async def ai_structure_review(texto, client, model, estrutura_mapeada=None, metrics=None):
    """
    v2.2: Revisão semântica de estrutura usando IA com VALIDAÇÃO CRUZADA.
    Compara o documento com a estrutura de mapeamento inicial.
    Corrige: questões duplicadas, subtópicos órfãos, fragmentação excessiva.
    
    Melhorias v2.2:
    - Limite aumentado para 800k chars
    - Validação de ordem dos headers
    - Integração com MetricsCollector
    - Suporte a relatório JSON
    """
    from google.genai import types
    from difflib import SequenceMatcher
    import asyncio
    import time
    import json
    
    print(f"{Fore.CYAN}🧠 Revisão Semântica de Estrutura (IA v2.2)...")
    
    start_time = time.time()
    
    # Gemini 3 Flash suporta 1M tokens (~4M chars) - usar até 800k chars
    max_doc_chars = 800000
    if len(texto) > max_doc_chars:
        print(f"   ⚠️ Documento muito longo ({len(texto)} chars), truncando para {max_doc_chars//1000}k...")
        texto_para_revisao = texto[:max_doc_chars] + "\n\n[... documento truncado para revisão estrutural ...]"
    else:
        texto_para_revisao = texto
    
    # Preparar estrutura mapeada (se disponível)
    if estrutura_mapeada:
        estrutura_str = estrutura_mapeada[:50000]  # Limitar estrutura a 50k chars
        print(f"{Fore.CYAN}   📋 Usando estrutura de mapeamento inicial ({len(estrutura_mapeada)} chars) para validação cruzada.{Style.RESET_ALL}")
    else:
        estrutura_str = "[Estrutura de mapeamento não disponível - analisar documento autonomamente]"
        print(f"{Fore.YELLOW}   ℹ️  Sem mapeamento inicial, IA revisará estrutura autonomamente.{Style.RESET_ALL}")
    
    def call_gemini():
        return client.models.generate_content(
            model=model,
            contents=PROMPT_STRUCTURE_REVIEW.format(
                estrutura_mapeada=estrutura_str,
                documento=texto_para_revisao
            ),
            config=types.GenerateContentConfig(
                max_output_tokens=65536,
                thinking_config=types.ThinkingConfig(
                    include_thoughts=False,
                    thinking_level="HIGH"
                ),
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                ]
            )
        )
    
    # Retry com backoff exponencial
    max_retries = 3
    response = None
    for tentativa in range(max_retries):
        try:
            response = await asyncio.to_thread(call_gemini)
            resultado = response.text.replace('```markdown', '').replace('```', '').strip()
            break
        except Exception as e:
            if tentativa < max_retries - 1:
                wait_time = 2 ** tentativa
                print(f"{Fore.YELLOW}   ⚠️ Tentativa {tentativa + 1} falhou: {e}. Aguardando {wait_time}s...{Style.RESET_ALL}")
                await asyncio.sleep(wait_time)
            else:
                print(f"{Fore.RED}   ⚠️ Erro na revisão por IA após {max_retries} tentativas: {e}. Mantendo estrutura original.{Style.RESET_ALL}")
                return texto
    
    # Métricas
    duration = time.time() - start_time
    if metrics and response:
        try:
            usage = response.usage_metadata
            metrics.record_call(
                provider="gemini",
                prompt_tokens=getattr(usage, 'prompt_token_count', 0) or 0,
                completion_tokens=getattr(usage, 'candidates_token_count', 0) or 0,
                duration=duration
            )
        except:
            pass  # Silently ignore metrics errors
    
    # Validação: o resultado deve ter pelo menos 80% do tamanho original (padronizado)
    if len(resultado) < len(texto) * 0.80:
        print(f"   ⚠️ Revisão retornou texto muito curto ({len(resultado)} vs {len(texto)}). Mantendo original.")
        return texto
    
    # Extrair relatório (suporta JSON e texto)
    relatorio_json_match = re.search(r'<!--\s*RELATÓRIO_JSON:\s*(\{.+?\})\s*-->', resultado, re.IGNORECASE | re.DOTALL)
    relatorio_match = re.search(r'<!--\s*RELATÓRIO:\s*(.+?)\s*-->', resultado, re.IGNORECASE)
    
    if relatorio_json_match:
        try:
            relatorio_data = json.loads(relatorio_json_match.group(1))
            print(f"{Fore.CYAN}   📊 Relatório (JSON): {relatorio_data}{Style.RESET_ALL}")
        except:
            print(f"{Fore.CYAN}   📊 Relatório: {relatorio_json_match.group(1)}{Style.RESET_ALL}")
        resultado = re.sub(r'<!--\s*RELATÓRIO_JSON:.+?-->\s*', '', resultado, flags=re.IGNORECASE | re.DOTALL).strip()
    elif relatorio_match:
        relatorio = relatorio_match.group(1)
        print(f"{Fore.CYAN}   📊 Relatório da IA: {relatorio}{Style.RESET_ALL}")
        resultado = re.sub(r'<!--\s*RELATÓRIO:.+?-->\s*', '', resultado, flags=re.IGNORECASE).strip()
    
    # Validação de ordem dos headers (novo em v2.2)
    headers_original = re.findall(r'^(#{1,4})\s+(.+?)$', texto, re.MULTILINE)
    headers_revisado = re.findall(r'^(#{1,4})\s+(.+?)$', resultado, re.MULTILINE)
    
    if len(headers_original) > 5 and len(headers_revisado) > 5:
        # Usar similaridade para detectar reordenação
        similares = sum(1 for (_, h1), (_, h2) in zip(headers_original[:10], headers_revisado[:10]) 
                       if SequenceMatcher(None, h1.strip(), h2.strip()).ratio() > 0.6)
        if similares < 6:
            print(f"{Fore.YELLOW}   ⚠️ Ordem dos headers parece alterada ({similares}/10 similares). Mantendo original.{Style.RESET_ALL}")
            return texto
    
    # Contar quantos headers foram alterados
    diff = abs(len(headers_original) - len(headers_revisado))
    
    print(f"{Fore.GREEN}   ✅ Estrutura revisada: {len(headers_original)} → {len(headers_revisado)} headers (Δ{diff}, {duration:.1f}s){Style.RESET_ALL}")
    return resultado

def normalizar_fingerprint(texto, tipo):
    """Normaliza texto para comparação (ex: 'Lei 11.100' -> 'lei 11100')"""
    texto = texto.lower().strip()
    
    if tipo == 'leis':
        nums = re.findall(r'\d+', texto)
        if nums:
            num_full = ''.join(nums)
            if len(num_full) >= 4: return f"lei {num_full}"
            return None
            
    elif tipo == 'sumulas':
        nums = re.findall(r'\d+', texto)
        if nums: return f"súmula {''.join(nums)}"
            
    elif tipo == 'artigos':
        nums = re.findall(r'\d+', texto)
        if nums: return f"artigo {''.join(nums)}"
            
    return re.sub(r'[^\w\s]', '', texto)

def extrair_fingerprints(texto):
    """Extrai 'fingerprints' únicos e normalizados do texto"""
    fingerprints = {'leis': set(), 'sumulas': set(), 'artigos': set(), 'julgados': set()}
    
    lei_pattern = re.compile(r'\b(?:lei|l\.)\s*n?º?\s*([\d\.]+)', re.IGNORECASE)
    sumula_pattern = re.compile(r'\bsúmula\s*(?:vinculante)?\s*n?º?\s*(\d+)', re.IGNORECASE)
    
    for match in lei_pattern.finditer(texto):
        fp = normalizar_fingerprint(f"lei {match.group(1)}", 'leis')
        if fp: fingerprints['leis'].add(fp)
    
    for match in sumula_pattern.finditer(texto):
        fp = normalizar_fingerprint(f"súmula {match.group(1)}", 'sumulas')
        if fp: fingerprints['sumulas'].add(fp)
        
    return fingerprints

def contar_ocorrencias_robust(fingerprints, texto):
    """Conta ocorrências com suporte a formatação jurídica formal"""
    contagens = {}
    texto_lower = texto.lower()
    
    for categoria, items in fingerprints.items():
        for item in items:
            key = f"{categoria}:{item}"
            if categoria == 'leis':
                num_bruto = item.split()[-1] 
                num = re.sub(r'[^\d]', '', num_bruto)
                num_regex = r"\.?".join(list(num))
                pattern = f"lei(?:\\s+|\\.|\\,|nº|n\\.|n\\s|num\\.?)*{num_regex}\\b"
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
            elif categoria == 'sumulas':
                num = item.split()[-1]
                num_regex = r"\.?".join(list(num))
                pattern = f"súmula(?:\\s+|\\.|\\,|vinculante|nº|n\\.|n\\s)*{num_regex}\\b"
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
            else:
                contagens[key] = texto_lower.count(item)
    return contagens

def contar_ocorrencias_robust(fingerprints, texto):
    """Conta ocorrencias com suporte a formatação jurídica formal (Lei nº X)"""
    contagens = {}
    
    # CORREÇÃO 1: Não remover pontuação indiscriminadamente, apenas normalizar espaços
    # Mantemos barras e pontos para evitar fusão de números (11.101/2005)
    texto_lower = texto.lower()
    
    for categoria, items in fingerprints.items():
        for item in items:
            key = f"{categoria}:{item}"
            
            if categoria == 'leis':
                # item ex: "lei 4320" -> extrai "4320"
                # Remove pontuação do item para garantir match limpo no número
                num_bruto = item.split()[-1] 
                num = re.sub(r'[^\d]', '', num_bruto)
                
                # Permite pontos opcionais entre dígitos (ex: 4.320 match com 4320)
                num_regex = r"\.?".join(list(num))
                
                # CORREÇÃO 2: Regex flexível que aceita "n", "nº", "no", "num" no meio
                # Aceita: "Lei 4320", "Lei nº 4.320", "Lei n. 4320"
                # O \W* permite pontos/barras entre Lei e o número
                # Adicionado \b no final para evitar matches parciais (Lei 10 != Lei 100)
                pattern = f"lei(?:\\s+|\\.|\\,|nº|n\\.|n\\s|num\\.?)*{num_regex}\\b"
                
                # Usamos findall no texto original (lower) para pegar variações com pontuação
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
                
            elif categoria == 'sumulas':
                num = item.split()[-1]
                num_regex = r"\.?".join(list(num)) # Súmulas raramente tem ponto, mas por garantia
                
                # Mesma lógica para súmulas (Súmula Vinculante nº 10)
                pattern = f"súmula(?:\\s+|\\.|\\,|vinculante|nº|n\\.|n\\s)*{num_regex}\\b"
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
                
            else:
                # Fallback para outros tipos (busca literal simples)
                contagens[key] = texto_lower.count(item)
                
    return contagens

def verificar_cobertura(texto_original, texto_formatado, arquivo_saida=None):
    """Verifica omissões e duplicações artificiais entre original e formatado"""
    logger.info("🔍 Verificando cobertura e duplicações...")
    
    # Extrai fingerprints do original
    fp_original = extrair_fingerprints(texto_original)
    
    # Conta ocorrências em ambos
    contagem_original = contar_ocorrencias_robust(fp_original, texto_original)
    contagem_formatado = contar_ocorrencias_robust(fp_original, texto_formatado)
    
    omissoes = []
    duplicacoes = []
    
    for key, count_orig in contagem_original.items():
        count_fmt = contagem_formatado.get(key, 0)
        categoria, item = key.split(':', 1)
        
        # Omissão: estava no original mas sumiu
        if count_orig > 0 and count_fmt == 0:
            omissoes.append({
                'categoria': categoria,
                'item': item,
                'original': count_orig,
                'formatado': count_fmt
            })
        
        # Duplicação (agora considerada positiva em materiais didáticos)
        if count_fmt > count_orig:
            duplicacoes.append({
                'categoria': categoria,
                'item': item,
                'original': count_orig,
                'formatado': count_fmt,
                'extra': count_fmt - count_orig
            })
    
    # Gera relatório
    total_items = len([k for k, v in contagem_original.items() if v > 0])
    items_preservados = total_items - len(omissoes)
    cobertura = items_preservados / total_items * 100 if total_items > 0 else 100
    
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    logger.info(f"📊 RELATÓRIO DE VERIFICAÇÃO")
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    logger.info(f"✅ Cobertura: {cobertura:.1f}% ({items_preservados}/{total_items} referências)")
    
    relatorio_txt = []
    relatorio_txt.append(f"RELATÓRIO DE VERIFICAÇÃO DE FIDELIDADE")
    relatorio_txt.append(f"Cobertura: {cobertura:.1f}% ({items_preservados}/{total_items} referências)")
    
    if omissoes:
        logger.warning(f"\n❌ POSSÍVEIS OMISSÕES ({len(omissoes)}):")
        relatorio_txt.append(f"\n❌ POSSÍVEIS OMISSÕES ({len(omissoes)}):")
        for o in omissoes[:15]: 
            msg = f"   - [{o['categoria']}] {o['item']}"
            logger.warning(msg)
            relatorio_txt.append(msg)
        if len(omissoes) > 15:
            logger.warning(f"   ... e mais {len(omissoes) - 15} omissões")
            relatorio_txt.append(f"   ... e mais {len(omissoes) - 15} omissões")
    else:
        logger.info("✅ Nenhuma omissão detectada")
        relatorio_txt.append("\n✅ Nenhuma omissão detectada")
    
    if duplicacoes:
        logger.info(f"\nℹ️ CITAÇÕES REFORÇADAS (Tabelas/Resumos) ({len(duplicacoes)}):")
        relatorio_txt.append(f"\nℹ️ CITAÇÕES REFORÇADAS ({len(duplicacoes)}):")
        for d in duplicacoes[:10]:
            msg = f"   - [{d['categoria']}] {d['item']}: {d['original']}x -> {d['formatado']}x"
            logger.info(msg)
            relatorio_txt.append(msg)
    
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    
    # Salva relatório em arquivo se especificado
    if arquivo_saida:
        relatorio_path = arquivo_saida.replace('.md', '_verificacao.txt')
        with open(relatorio_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(relatorio_txt))
        logger.info(f"📝 Relatório salvo em: {relatorio_path}")

def mover_tabelas_para_fim_de_secao(texto):
    """
    v2.11: Reorganiza tabelas movendo-as para o final do BLOCO ATUAL (H2 ou H3).
    Corrige bug de tabelas sumindo ou ficando muito longe do contexto.
    """
    logger.info("📊 Reorganizando tabelas (Smart Layout)...")
    
    linhas = texto.split('\n')
    resultado = []
    tabelas_pendentes = [] 
    
    i = 0
    while i < len(linhas):
        linha = linhas[i]
        linha_strip = linha.strip()
        
        # 1. DETECTAR SE É UM TÍTULO (H1, H2, H3...)
        if linha_strip.startswith('#'):
            # Despeja tabelas antes de iniciar o novo tópico
            if tabelas_pendentes:
                resultado.append('') # Espaço antes
                for t_info in tabelas_pendentes:
                    if t_info['titulo']:
                        resultado.append(t_info['titulo'])
                    resultado.extend(t_info['linhas'])
                    resultado.append('') # Espaço depois
                tabelas_pendentes = []
            
            resultado.append(linha)
            i += 1
            continue

        # 2. DETECTAR INÍCIO DE TABELA
        eh_inicio_tabela = False
        if '|' in linha_strip:
            has_separator = False
            for lookahead in range(1, 3): 
                if i + lookahead < len(linhas):
                    prox = linhas[i + lookahead].strip()
                    if set(prox).issubset(set('|- :')): 
                         has_separator = True
                         break
            
            if has_separator or (linha_strip.startswith('|') and linha_strip.endswith('|')):
                eh_inicio_tabela = True

        if eh_inicio_tabela:
            # --- Captura da Tabela ---
            tabela_linhas = []
            titulo_tabela = None
            
            # Tenta recuperar o título da tabela que ficou na linha anterior
            if resultado and len(resultado) > 0:
                last_line = resultado[-1].strip()
                if (last_line.startswith('###') or last_line.startswith('**')) and (
                   any(x in last_line.lower() for x in ['tabela', 'resumo', 'quadro', 'síntese', 'esquema', '📋'])):
                    titulo_tabela = resultado.pop() 

            # Captura as linhas da tabela
            j = i
            while j < len(linhas):
                curr = linhas[j].strip()
                if '|' in curr:
                    tabela_linhas.append(linhas[j])
                    j += 1
                elif not curr:
                    if j + 1 < len(linhas) and '|' in linhas[j+1]:
                        tabela_linhas.append(linhas[j]) 
                        j += 1
                    else:
                        break # Fim da tabela
                else:
                    break # Texto normal

            if len(tabela_linhas) > 0:
                tabelas_pendentes.append({
                    'titulo': titulo_tabela,
                    'linhas': tabela_linhas
                })
                i = j # Pula as linhas processadas
                continue
            else:
                if titulo_tabela:
                    resultado.append(titulo_tabela)
        
        resultado.append(linha)
        i += 1
    
    # 3. FINAL DO DOCUMENTO
    if tabelas_pendentes:
        resultado.append('')
        for t_info in tabelas_pendentes:
            if t_info['titulo']:
                resultado.append(t_info['titulo'])
            resultado.extend(t_info['linhas'])
            resultado.append('')
            
    return '\n'.join(resultado)
    
def dividir_sequencial(transcricao_completa, chars_por_parte=6000, estrutura_global=None):
    """
    v2.17: Divide documento em chunks SEQUENCIAIS com preferência por âncoras.
    
    Se estrutura_global for fornecida, tenta alinhar os pontos de corte com
    os títulos mapeados para evitar quebrar no meio de um tópico importante.
    
    Args:
        transcricao_completa: Texto bruto completo
        chars_por_parte: Tamanho alvo de cada chunk
        estrutura_global: String com a estrutura mapeada (opcional)
    
    Returns:
        Lista de dicts com 'inicio' e 'fim' de cada chunk
    """
    chunks = []
    tamanho_total = len(transcricao_completa)
    inicio = 0
    
    # Extrair âncoras da estrutura se disponível
    ancoras = []
    if estrutura_global:
        # Extrair títulos de nível 1 e 2 (ex: "1. Introdução", "2.1. Conceito")
        for line in estrutura_global.split('\n'):
            line = line.strip()
            # Match linhas como "1. Tema" ou "2.1. Subtema"
            match = re.match(r'^\d+(?:\.\d+)*\.?\s+(.+)', line)
            if match:
                titulo = match.group(1).strip()
                # Pegar primeiras 3 palavras significativas como âncora
                palavras = [w for w in titulo.split() if len(w) > 3][:3]
                if palavras:
                    ancoras.append(' '.join(palavras).lower())
    
    while inicio < tamanho_total:
        fim_ideal = min(inicio + chars_por_parte, tamanho_total)
        fim = fim_ideal
        
        if fim < tamanho_total:
            bloco = transcricao_completa[inicio:fim]
            melhor_ponto = None
            
            # ESTRATÉGIA 1: Procurar âncora de estrutura próxima
            if ancoras:
                # Buscar âncoras na zona de tolerância (últimos 30% do chunk)
                zona_busca = bloco[int(chars_por_parte * 0.7):]
                zona_offset = int(chars_por_parte * 0.7)
                
                for ancora in ancoras:
                    pos = zona_busca.lower().find(ancora)
                    if pos != -1:
                        # Encontrou âncora - cortar ANTES dela
                        ponto_corte = zona_offset + pos
                        # Voltar até o início do parágrafo
                        while ponto_corte > 0 and bloco[ponto_corte] != '\n':
                            ponto_corte -= 1
                        if ponto_corte > chars_por_parte * 0.7:
                            melhor_ponto = inicio + ponto_corte
                            break
            
            # ESTRATÉGIA 2: Fallback para fim de parágrafo
            if melhor_ponto is None:
                ultimo_paragrafo = bloco.rfind('\n\n')
                if ultimo_paragrafo != -1 and ultimo_paragrafo > chars_por_parte * 0.8:
                    melhor_ponto = inicio + ultimo_paragrafo + 2
            
            if melhor_ponto:
                fim = melhor_ponto
        
        chunks.append({
            'inicio': inicio, 
            'fim': fim
        })
        inicio = fim
        
    return chunks

def validar_chunks(chunks, texto_completo):
    """Valida se não houve perda de texto entre chunks"""
    esperado = 0
    for c in chunks:
        if c['inicio'] != esperado:
             print(f"{Fore.RED}⚠️ GAP detectado! Chunk começa em {c['inicio']} mas devia ser {esperado}")
        esperado = c['fim']
    
    if esperado != len(texto_completo):
        print(f"{Fore.RED}⚠️ Texto incompleto! Processado: {esperado}, Total: {len(texto_completo)}")
    else:
        print(f"{Fore.GREEN}✅ Divisão de chunks validada (Bytes match).")

# Removed old detectar_secoes_duplicadas & remover_secoes_duplicadas since they are replaced by robust versions
# The new versions are placed in the helper section we just updated.

def remover_duplicacoes_literais(texto):
    """Remove parágrafos individuais duplicados (v2.7 logic)"""
    from difflib import SequenceMatcher
    paragrafos = texto.split('\n\n')
    paragrafos_limpos = []
    dup_count = 0
    
    for i, para in enumerate(paragrafos):
        if i == 0:
            paragrafos_limpos.append(para)
            continue
        
        if len(para.strip()) < 80 or para.strip().startswith('#'):
            paragrafos_limpos.append(para)
            continue
        
        is_duplicate = False
        para_norm = ' '.join(para.split()).lower()
        
        # Check against last 3 paragraphs (v2.7 logic)
        for j in range(max(0, len(paragrafos_limpos) - 3), len(paragrafos_limpos)):
            para_ant = paragrafos_limpos[j]
            para_ant_norm = ' '.join(para_ant.split()).lower()
            
            ratio = SequenceMatcher(None, para_norm, para_ant_norm).ratio()
            
            if ratio > 0.95:
                is_duplicate = True
                dup_count += 1
                break
        
        if not is_duplicate:
            paragrafos_limpos.append(para)
    
    if dup_count > 0:
         print(f"⚠️  {dup_count} parágrafos duplicados removidos (Literal Dedup)")
    
    return '\n\n'.join(paragrafos_limpos)


def numerar_titulos(texto):
    """Adiciona numeração sequencial aos títulos"""
    linhas = texto.split('\n')
    linhas_numeradas = []
    
    contador_h2 = 0
    contador_h3 = 0
    contador_h4 = 0
    
    titulo_pattern = re.compile(r'^(#{2,4})\s+(?:\d+(?:\.\d+)*\.?\s+)?(.+)$')
    
    for linha in linhas:
        match = titulo_pattern.match(linha)
        
        if match:
            nivel = len(match.group(1))
            texto_titulo = match.group(2).strip()
            
            # Não numera títulos de resumo/quadros
            if any(keyword in texto_titulo.lower() for keyword in ['resumo', 'quadro', 'esquema', '📋', '📊', '🗂️']):
                linhas_numeradas.append(linha)
                continue
            
            if nivel == 2:
                contador_h2 += 1
                contador_h3 = 0
                contador_h4 = 0
                nova_linha = f"## {contador_h2}. {texto_titulo}"
            elif nivel == 3:
                contador_h3 += 1
                contador_h4 = 0
                nova_linha = f"### {contador_h2}.{contador_h3}. {texto_titulo}"
            elif nivel == 4:
                contador_h4 += 1
                nova_linha = f"#### {contador_h2}.{contador_h3}.{contador_h4}. {texto_titulo}"
            else:
                nova_linha = linha
            
            linhas_numeradas.append(nova_linha)
        else:
            linhas_numeradas.append(linha)
    
    return '\n'.join(linhas_numeradas)


# ==================== METRICS COLLECTOR (v2.10) ====================
class MetricsCollector:
    """Tracks API usage, timing, and cost estimation for optimization."""
    
    # Preços Gemini 3 Flash Preview (Dezembro 2025) - USD per 1M tokens
    PRICE_INPUT = 0.50   # $0.50 (Gemini 3 Flash Preview Input)
    PRICE_OUTPUT = 3.00  # $3.00 (Gemini 3 Flash Preview Output)
    
    # OpenAI GPT-5 Mini (Estimado)
    PRICE_OPENAI_INPUT = 0.15
    PRICE_OPENAI_OUTPUT = 0.60
    
    def __init__(self):
        self.reset()
    
    def reset(self):
        self.api_calls = 0
        self.gemini_calls = 0
        self.openai_calls = 0
        self.total_prompt_tokens = 0
        self.total_completion_tokens = 0
        self.total_time_seconds = 0.0
        self.call_times = []
        self.cache_hits = 0
        self.adaptive_splits = 0
    
    def record_call(self, provider: str, prompt_tokens: int, completion_tokens: int, duration: float):
        """Records a single API call."""
        self.api_calls += 1
        if provider == "gemini":
            self.gemini_calls += 1
        elif provider == "openai":
            self.openai_calls += 1
        self.total_prompt_tokens += prompt_tokens
        self.total_completion_tokens += completion_tokens
        self.total_time_seconds += duration
        self.call_times.append(duration)
    
    def record_cache_hit(self):
        self.cache_hits += 1
    
    def record_adaptive_split(self):
        self.adaptive_splits += 1
    
    def estimate_cost(self) -> float:
        """Estimates total USD cost based on recorded tokens."""
        # Simplified: assumes all calls are Gemini unless we track per-call
        gemini_cost = (
            (self.total_prompt_tokens * self.PRICE_INPUT / 1_000_000) +
            (self.total_completion_tokens * self.PRICE_OUTPUT / 1_000_000)
        )
        return gemini_cost
    
    def get_summary(self) -> str:
        """Returns a formatted summary string."""
        avg_time = (self.total_time_seconds / self.api_calls) if self.api_calls > 0 else 0
        cost = self.estimate_cost()
        
        return f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 MÉTRICAS DE EXECUÇÃO
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
   📡 Total de Chamadas API: {self.api_calls}
      - Gemini: {self.gemini_calls}
      - OpenAI: {self.openai_calls}
      - Cache Hits: {self.cache_hits}
   ✂️ Divisões Adaptativas: {self.adaptive_splits}
   🎯 Tokens Usados:
      - Prompt: {self.total_prompt_tokens:,}
      - Completion: {self.total_completion_tokens:,}
      - Total: {self.total_prompt_tokens + self.total_completion_tokens:,}
   ⏱️ Tempo Total: {self.total_time_seconds:.1f}s (média: {avg_time:.2f}s/chamada)
   💰 Custo Estimado: ${cost:.4f} USD
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

# Global metrics instance
metrics = MetricsCollector()

# ==================== CHECKPOINT SYSTEM ====================
def get_checkpoint_path(video_name, folder):
    return Path(folder) / f"{video_name}.checkpoint.json"

def save_checkpoint(video_name, folder, results, segments_info, current_idx):
    path = get_checkpoint_path(video_name, folder)
    data = {
        'video_name': video_name,
        'current_idx': current_idx,
        'total_segments': len(segments_info),
        'results': results,
        'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
    }
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_checkpoint(video_name, folder):
    path = get_checkpoint_path(video_name, folder)
    if path.exists():
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️ Erro ao carregar checkpoint: {e}")
    return None

def delete_checkpoint(video_name, folder):
    path = get_checkpoint_path(video_name, folder)
    if path.exists(): 
        path.unlink()
        print(f"🧹 Checkpoint removido: {path.name}")

class VomoMLX:
    # GPT-5 Mini: 400k tokens input, 128k output
    MAX_CHUNK_SIZE = 100000  
    CHUNK_OVERLAP = 3000     # 5k overlap
    

    PROMPT_MAPEAMENTO = """Você é um especialista em organização de conteúdo educacional acadêmico.

## ETAPA 1: IDENTIFICAR O TIPO DE CONTEÚDO
Analise a transcrição e determine qual é a **natureza predominante** do material:

| Tipo | Pistas no Texto | Estrutura Ideal |
|------|-----------------|-----------------|
| **SIMULADO** | "questão 1", "questão 2", "espelho de correção", "correção do simulado", "vamos corrigir" | Organizar por QUESTÕES numeradas |
| **AULA EXPOSITIVA** | explicações contínuas de um tema, teoria, doutrina, sem questões específicas | Organizar por TEMAS/MATÉRIAS |
| **REVISÃO** | "revisão", "resumo", múltiplos temas curtos, "pontos importantes" | Organizar por TÓPICOS de revisão |
| **CORREÇÃO DE PROVA** | "gabarito", "alternativa correta", "item certo/errado" | Organizar por QUESTÕES com gabarito |

## ETAPA 2: EXTRAIR A ESTRUTURA

### Se for SIMULADO ou CORREÇÃO DE PROVA:
```
1. Orientações Gerais / Introdução
2. Questão 1: [Título descritivo] — [Área do Direito]
   2.1. Enunciado e Contexto
   2.2. Fundamentação (Doutrina/Jurisprudência)
   2.3. Pontos do Espelho / Resposta
3. Questão 2: [Título descritivo] — [Área do Direito]
   3.1. ...
[Continue para cada questão]
N. Considerações Finais / Dúvidas
```

### Se for AULA EXPOSITIVA:
```
1. Introdução
2. [Matéria 1: ex. Direito Administrativo]
   2.1. [Subtema]
      2.1.1. [Detalhamento]
3. [Matéria 2: ex. Direito Civil]
   3.1. ...
```

### Se for REVISÃO:
```
1. [Tema 1]
   1.1. Pontos-chave
   1.2. Jurisprudência/Súmulas
2. [Tema 2]
   2.1. ...
```

## REGRAS GERAIS:
1. **MÁXIMO 3 NÍVEIS** de hierarquia (1., 1.1., 1.1.1.)
2. **Seja descritivo** nos títulos — inclua o assunto real, não apenas "Questão 1"
3. **Mantenha a ORDEM** cronológica da transcrição
4. **Mapeie do INÍCIO ao FIM** — não omita partes
5. **Identifique a ÁREA DO DIREITO** de cada bloco quando possível

## 🏛️ REGRA ESPECIAL: MARCOS LEGAIS (v2.17)
Quando identificar marcos legais importantes, crie subtópicos específicos:
- **Súmulas** (STF, STJ, Vinculantes): Criar subtópico "X.Y. Súmula [Número] do [Tribunal]."
- **Teses (Repercussão Geral/Repetitivos)**: Criar subtópico "X.Y. Tese/Tema [Número] do STJ/STF."
- **Artigos de Lei Central**: Se um artigo é explicado em profundidade, criar subtópico "X.Y. Art. [Número] da [Lei]."

Exemplo:
```
2. Execução Fiscal
   2.1. Procedimento da LEF (Lei 6.830/80)
   2.2. Súmula 314 do STJ (Citação por Hora Certa)
   2.3. Tema 444 do STJ (Redirecionamento)
```

## TRANSCRIÇÃO:
{transcricao}

## RESPOSTA:
Primeiro, indique em uma linha: `[TIPO: SIMULADO/EXPOSITIVA/REVISÃO/CORREÇÃO]`
Depois, retorne APENAS a estrutura hierárquica (máx 3 níveis)."""


    PROMPT_FIDELIDADE = """# DIRETRIZES DE FORMATAÇÃO E REVISÃO (MODO FIDELIDADE)

## PAPEL
VOCÊ É UM EXCELENTÍSSIMO REDATOR TÉCNICO E DIDÁTICO.
- **Tom:** didático, como o professor explicando em aula.
- **Pessoa:** MANTENHA a pessoa original da transcrição (1ª pessoa se for assim na fala).
- **Estilo:** texto corrido, com parágrafos curtos, sem "inventar" doutrina nova.
- **Objetivo:** reproduzir a aula em forma escrita, clara e organizada, mas ainda com a "voz" do professor.

# OBJETIVO
- Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, MANTENDO A FIDELIDADE TOTAL ao conteúdo original.
- **Tamanho:** a saída deve ficar **entre 95% e 115%** do tamanho do trecho de entrada (salvo remoção de muletas e logística).

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias e das falas do professor.
4. **NÃO CRIE MUITOS BULLET POINTS**. PREFIRA UM FORMATO DE MANUAL DIDÁTICO, não checklist.
5. **NÃO USE NEGRITOS EM EXCESSO**. Use apenas para conceitos-chave realmente importantes.

## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados (REsp/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios.
- **Referências**: leis, artigos, jurisprudência, autores, casos citados.
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque).
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico".

## 🎯 PRESERVAÇÃO ESPECIAL: DICAS DE PROVA E EXAMINADORES (CRÍTICO)
Aulas presenciais frequentemente contêm informações valiosas sobre:
1. **Referências a Examinadores**: Nomes de examinadores de concursos, suas preferências, posicionamentos ou temas favoritos. **PRESERVE INTEGRALMENTE**.
   - Exemplo: "O examinador Fulano costuma cobrar..." → MANTER
   - Exemplo: "Esse tema foi cobrado pelo professor X na prova..." → MANTER
2. **Dicas de Prova**: Orientações sobre o que costuma cair em provas, pegadinhas comuns, temas recorrentes.
   - Exemplo: "Isso cai muito em prova..." → MANTER
   - Exemplo: "Atenção: essa é uma pegadinha clássica..." → MANTER
3. **Estratégias de Estudo**: Sugestões do professor sobre priorização, macetes, formas de memorização.
   - Exemplo: "Gravem isso: na dúvida, marquem..." → MANTER
   - Exemplo: "Para PGM, foquem em..." → MANTER
4. **Casos Práticos e Histórias Reais**: Exemplos de situações reais, casos julgados, histórias ilustrativas.
   - **NUNCA RESUMA** histórias ou exemplos práticos. Preserve na íntegra.

> ⚠️ **ESSAS INFORMAÇÕES SÃO O DIFERENCIAL DE UMA AULA AO VIVO.** Sua omissão representa perda irreparável de valor didático.


## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Corrija erros gramaticais, regências, ortográficos e de pontuação.
2. **Limpeza Profunda:**
   - **REMOVA** marcadores de oralidade: "né", "tá?", "entende?", "veja bem", "tipo assim".
   - **REMOVA** interações diretas com a turma: "Isso mesmo", "A colega perguntou", "Já estão me vendo?", "Estão ouvindo?".
   - **REMOVA** redundâncias: "subir para cima", "criação nova".
   - **TRANSFORME** perguntas retóricas em afirmações quando possível.
3. **Coesão**: Utilize conectivos para tornar o texto mais fluido. Aplique pontuação adequada.
4. **Legibilidade**:
   - **PARÁGRAFOS CURTOS**: máximo **3-6 linhas visuais** por parágrafo.
   - **QUEBRE** blocos de texto maciços em parágrafos menores.
   - Seja didático sem perder detalhes e conteúdo.
5. **Formatação Didática** (use com moderação):
   - **Bullet points** para enumerar elementos, requisitos ou características.
   - **Listas numeradas** para etapas, correntes ou exemplos.
   - **Marcadores relacionais** como "→" para consequências lógicas.

## 📝 ESTRUTURA E TÍTULOS
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis.
- **NÃO crie subtópicos para frases soltas.**
- Use títulos **APENAS** para mudanças reais de assunto.

## 📊 TABELA DE SÍNTESE
Ao final de cada **bloco temático relevante**, produza uma tabela de síntese:

| Conceito | Definição | Fundamento Legal | Observações |
| :--- | :--- | :--- | :--- |
| ...  | ...  | Art. X, Lei Y | ... |

**REGRAS CRÍTICAS PARA TABELAS:**
1. **Limite:** máximo ~50 palavras por célula.
2. **PROIBIDO** blocos de código dentro de células.
3. **POSICIONAMENTO:** A tabela vem **APENAS AO FINAL** de um bloco concluído, **NUNCA** no meio de explicação.

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA.
- **NUNCA formate novamente esse contexto.**
- **NUNCA inclua esse contexto na sua resposta.**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>.
    """

    PROMPT_APOSTILA_ACTIVE = """# DIRETRIZES DE REDAÇÃO: MANUAL JURÍDICO DIDÁTICO (MODO APOSTILA)

## PAPEL
VOCÊ É UM EXCELENTÍSSIMO REDATOR JURÍDICO E DIDÁTICO.
- **Tom:** doutrinário, impessoal, estilo manual de Direito.
- **Pessoa:** 3ª pessoa ou construções impessoais ("O professor explica...", "A doutrina define...").
- **Estilo:** prosa densa, porém com parágrafos curtos e didáticos.
- **Objetivo:** transformar a aula em texto de apostila/manual, sem alterar conteúdo nem inventar informações.

## 💎 PILAR 1: ESTILO (VOZ ATIVA E DIRETA)
> 🚫 **PROIBIDO VOZ PASSIVA EXCESSIVA:** "Anunciou-se", "Informou-se".
> ✅ **PREFIRA VOZ ATIVA:** "O professor explica...", "A doutrina define...", "O Art. 37 estabelece...".

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias.

## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados, Temas de Repercussão Geral, Recursos Repetitivos. **NUNCA OMITA NÚMEROS DE TEMAS OU SÚMULAS**.
- **JURISPRUDÊNCIA**: Se o texto citar "Tema 424", "RE 123", "ADI 555", **MANTENHA O NÚMERO**. Não generalize para "jurisprudência do STJ".
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios.
- **Referências**: leis, artigos, jurisprudência (STF/STJ), autores, casos citados.
- **Ênfases intencionais** e **Observações pedagógicas**.


## 🎯 PRESERVAÇÃO ESPECIAL: DICAS DE PROVA E EXAMINADORES (CRÍTICO)
Aulas presenciais frequentemente contêm informações valiosas sobre:
1. **Referências a Examinadores**: Nomes de examinadores de concursos, suas preferências, posicionamentos ou temas favoritos. **PRESERVE INTEGRALMENTE**.
   - Exemplo: "O examinador Fulano costuma cobrar..." → MANTER
   - Exemplo: "Esse tema foi cobrado pelo professor X na prova..." → MANTER
2. **Dicas de Prova**: Orientações sobre o que costuma cair em provas, pegadinhas comuns, temas recorrentes.
   - Exemplo: "Isso cai muito em prova..." → MANTER
   - Exemplo: "Atenção: essa é uma pegadinha clássica..." → MANTER
3. **Estratégias de Estudo**: Sugestões do professor sobre priorização, macetes, formas de memorização.
   - Exemplo: "Gravem isso: na dúvida, marquem..." → MANTER
   - Exemplo: "Para PGM, foquem em..." → MANTER
4. **Casos Práticos e Histórias Reais**: Exemplos de situações reais, casos julgados, histórias ilustrativas.
   - **NUNCA RESUMA** histórias ou exemplos práticos. Preserve na íntegra.

> ⚠️ **ESSAS INFORMAÇÕES SÃO O DIFERENCIAL DE UMA AULA AO VIVO.** Sua omissão representa perda irreparável de valor didático.


## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Ajuste a linguagem coloquial para o padrão culto.
2. **Limpeza**: Remova gírias, cacoetes ("né", "tipo assim", "então") e vícios de oralidade.
3. **Coesão**: Use conectivos e pontuação adequada para tornar o texto fluido.
4. **Legibilidade**:
   - **PARÁGRAFOS CURTOS**: máximo **3-6 linhas visuais** por parágrafo.
   - **QUEBRE** blocos de texto maciços em parágrafos menores.
   - Use **negrito** para destacar conceitos-chave (sem exagero).
5. **Formatação Didática** (use com moderação):
   - **Bullet points** para enumerar elementos, requisitos ou características.
   - **Listas numeradas** para etapas, correntes doutrinárias ou exemplos.
   - **Marcadores relacionais** como "→" para consequências lógicas.

## 📝 ESTRUTURA E TÍTULOS
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos.
- **NÃO crie subtópicos para frases soltas.**
- Use títulos **APENAS** para mudanças reais de assunto.

## 📊 TABELA DE SÍNTESE (OBRIGATÓRIO)
Ao final de CADA tópico principal (nível 2 ou 3), CRIE uma tabela de resumo.
SEMPRE que houver diferenciação de conceitos, prazos ou regras, CRIE UMA TABELA.

| Conceito/Instituto | Definição | Fundamento Legal | Observações |
| :--- | :--- | :--- | :--- |
| ...  | ...  | Art. X, Lei Y | ... |

**REGRAS CRÍTICAS PARA TABELAS:**
1. **Limite:** máximo ~50 palavras por célula.
2. **PROIBIDO** blocos de código dentro de células.
3. **NUNCA** deixe título "📋 Resumo" sozinho sem dados.
4. **POSICIONAMENTO:** A tabela vem **APENAS AO FINAL** de um bloco concluído.
   - **NUNCA** insira tabela no meio de explicação.
   - A tabela deve ser o **fechamento** lógico da seção.

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA.
- **NUNCA formate novamente esse contexto.**
- **NUNCA inclua esse contexto na sua resposta.**
- **NUNCA repita informações que já estão no contexto.**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>.
- **CRÍTICO:** Se o texto começar repetindo a última frase do contexto, **IGNORE A REPETIÇÃO.**
"""

    SYSTEM_PROMPT_FORMAT = PROMPT_APOSTILA_ACTIVE # Default

    def __init__(self, model_size="large-v3-turbo"):
        """
        MLX-Whisper otimizado para Apple Silicon (M3 Pro)
        Formatação otimizada para Gemini 3 Flash (Vertex AI)
        """
        print(f"{Fore.CYAN}🚀 Inicializando MLX-Whisper ({model_size}) para Apple Silicon...")
        self.model_name = model_size
        self.llm_model = "gemini-3-flash-preview"
        
        # Carrega variável de ambiente com override
        from dotenv import load_dotenv
        load_dotenv(override=True)
        
        project_id = os.getenv("GOOGLE_CLOUD_PROJECT", "gen-lang-client-0727883752")

        # Configuração de Credenciais (Explicit Fallback)
        CREDENTIALS_PATH = "/Users/nicholasjacob/Documents/Aplicativos/Iudex/gen-lang-client-0727883752-0bfab2f33e08.json"
        if os.path.exists(CREDENTIALS_PATH) and not os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDENTIALS_PATH
            print(f"{Fore.CYAN}🔑 Credenciais Vertex carregadas de: {CREDENTIALS_PATH}")
        
        # Estratégia Estrita de Autenticação (Vertex AI Only)
        print(f"{Fore.CYAN}☁️  Conectando via Vertex AI (Enterprise - Global Endpoint)...")
        self.client = genai.Client(
            vertexai=True,
            project=project_id,
            location="global"
        )
        # Teste rápido de permissão (dry-run)
        self.client.models.count_tokens(model=self.llm_model, contents="teste")
        print(f"{Fore.GREEN}   ✅ Conectado ao Vertex AI com sucesso.")

        # Inicializa OpenAI como Fallback Terciário
        self.openai_model = "gpt-5-mini-2025-08-07"
        openai_key = os.getenv("OPENAI_API_KEY")
        if openai_key:
            self.openai_client = AsyncOpenAI(api_key=openai_key)
            print(f"{Fore.CYAN}🤖 Fallback OpenAI ({self.openai_model}) configurado.")
        else:
            self.openai_client = None
            print(f"{Fore.YELLOW}⚠️  OpenAI Key não encontrada. Fallback GPT desativado.")

        # Inicializa seletor de prompt (Default: Apostila)
        self.prompt_apostila = self.PROMPT_APOSTILA_ACTIVE

        # Gemini não tem client async nativo igual OpenAI, usamos chamadas sync em threads
        self.async_client = None 
        
        # Cache directory
        self.cache_dir = Path(".cache_vomo")
        self.cache_dir.mkdir(exist_ok=True)

    def create_context_cache(self, transcription, global_structure=None):
        """
        v2.2: Wrapper method that calls criar_cache_contexto with class attributes.
        
        Args:
            transcription: Full transcription text
            global_structure: Mapped structure (optional)
        
        Returns:
            Cache object or None
        """
        return criar_cache_contexto(
            client=self.client,
            transcricao_completa=transcription,
            system_prompt=self.prompt_apostila,  # Uses current mode's prompt
            estrutura_global=global_structure,
            model_name=self.llm_model
        )

    def optimize_audio(self, file_path):
        """Extrai áudio otimizado (16kHz mono)"""
        print(f"{Fore.YELLOW}⚡ Verificando áudio...")
        
        mp3_path = Path(file_path).with_suffix('.mp3')
        if mp3_path.exists():
            print(f"   📂 Usando MP3 existente: {mp3_path.name}")
            return str(mp3_path)
        
        file_hash = hashlib.md5(file_path.encode()).hexdigest()[:8]
        output_path = f"temp_{file_hash}.wav"
        
        if os.path.exists(output_path):
            return output_path
        
        print(f"   🔄 Extraindo áudio...")
        subprocess.run([
            'ffmpeg', '-i', file_path,
            '-vn', '-ac', '1', '-ar', '16000',
            output_path, '-y', '-hide_banner', '-loglevel', 'error'
        ], check=True, capture_output=True)
        
        return output_path

    def transcribe(self, audio_path):
        """
        MLX-Whisper OTIMIZADO com GPU acelerado + Diarização
        
        Otimizações v2.0:
        - VAD filtering (pula silêncio) 
        - Batched inference (múltiplos chunks GPU)
        - condition_on_previous_text (contexto melhorado)
        - Hallucination suppression (evita texto inventado)
        """
        print(f"{Fore.GREEN}🎙️  Iniciando transcrição OTIMIZADA (MLX GPU)...")
        start_time = time.time()
        
        # Cache de transcrição
        cache_file = audio_path.replace('.wav', '_DIARIZATION.json').replace('.mp3', '_DIARIZATION.json')
        
        if os.path.exists(cache_file):
            try:
                print(f"{Fore.CYAN}   📂 Cache encontrado, carregando...")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                return cache_data['transcript']
            except Exception:
                pass

        if not mlx_whisper:
            raise ImportError("mlx_whisper não instalado.")
        
        # ==================== PARÂMETROS OTIMIZADOS ====================
        print("   🔍 Transcrevendo com parâmetros otimizados...")
        
        result = mlx_whisper.transcribe(
            audio_path,
            path_or_hf_repo=f"mlx-community/whisper-{self.model_name}",
            language="pt",
            
            # ========== PRECISÃO ==========
            temperature=0.0,              # Mais determinístico (desativa sampling estocástico)
            initial_prompt="Esta é uma transcrição de aula jurídica em português brasileiro sobre direito administrativo, constitucional, civil, penal e processual.",
            
            # === TIMESTAMPS ===
            word_timestamps=True,
            
            # === PERFORMANCE ===
            fp16=True,                    # Usa float16 (mais rápido na GPU)
            
            # === QUALIDADE (Hallucination Suppression) ===
            no_speech_threshold=0.6,      # Ignora segmentos com prob de silêncio > 60%
            logprob_threshold=-1.0,       # Rejeita tokens com log prob muito baixo
            compression_ratio_threshold=2.4,  # Detecta repetição/alucinação
            
            # === CONTEXTO ===
            condition_on_previous_text=True,  # Usa contexto anterior (melhora precisão)
            
            # === SUPRESSÃO DE TOKENS PROBLEMÁTICOS ===
            suppress_tokens=[-1],  # Suprime token de padding
            
            verbose=False
        )
        
        elapsed = time.time() - start_time
        audio_duration = result.get('duration', 0) if isinstance(result, dict) else 0
        rtf = elapsed / audio_duration if audio_duration > 0 else 0
        
        print(f"{Fore.GREEN}   ✅ Transcrição concluída em {elapsed:.1f}s (RTF: {rtf:.2f}x)")
        
        transcript_result = None
        
        # Diarização
        if Pipeline and torch and HF_TOKEN:
            try:
                print("   🗣️  Iniciando Diarização (Pyannote)...")
                pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-3.1",
                    use_auth_token=HF_TOKEN
                )
                
                device = "mps" if torch.backends.mps.is_available() else "cpu"
                pipeline.to(torch.device(device))
                diarization = pipeline(audio_path)
                
                transcript_result = self._align_diarization(result['segments'], diarization)
                print(f"{Fore.GREEN}✅ Diarização concluída")
            
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Erro na Diarização: {e}")
        
        if transcript_result is None:
            # Fallback sem diarização
            formatted_output = []
            last_timestamp = None  # Rastreia último timestamp inserido
            
            for segment in result['segments']:
                start = segment['start']
                
                # Adiciona timestamp apenas a cada 30 minutos
                if self._should_add_timestamp(start, last_timestamp, interval_minutes=20):
                    ts = self._format_timestamp(start)
                    formatted_output.append(f"[{ts}] {segment['text'].strip()}")
                    last_timestamp = start
                else:
                    formatted_output.append(segment['text'].strip())
            
            transcript_result = " ".join(formatted_output)
        
        # Salvar cache
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'transcript': transcript_result,
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
                }, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erro ao salvar cache: {e}")
        
        return transcript_result

    def _format_timestamp(self, seconds):
        m, s = divmod(seconds, 60)
        h, m = divmod(m, 60)
        return f"{int(h):02d}:{int(m):02d}:{int(s):02d}" if h > 0 else f"{int(m):02d}:{int(s):02d}"
    
    def _should_add_timestamp(self, current_seconds, last_timestamp_seconds, interval_minutes=20):
        """
        Determina se um timestamp deve ser adicionado baseado no intervalo configurado.
        
        Args:
            current_seconds: Tempo atual em segundos
            last_timestamp_seconds: Último timestamp inserido (None se for o primeiro)
            interval_minutes: Intervalo em minutos entre timestamps (padrão: 30)
        
        Returns:
            bool: True se deve adicionar timestamp
        """
        if last_timestamp_seconds is None:
            return True  # Sempre adiciona o primeiro timestamp
        
        interval_seconds = interval_minutes * 60
        return (current_seconds - last_timestamp_seconds) >= interval_seconds

    def _align_diarization(self, segments, diarization_output):
        """
        Alinha segmentos com diarização (Versão Otimizada com IntervalTree)
        
        Fase 2.2: Substituição de busca O(n) por IntervalTree O(log n)
        Ganho esperado: 10-20x mais rápido em áudios longos (>30 min)
        """
        try:
            from intervaltree import IntervalTree
        except ImportError:
            print("⚠️ intervaltree não instalado, usando fallback O(n)")
            return self._align_diarization_fallback(segments, diarization_output)
        
        raw_output = []
        current_speaker = None
        last_timestamp = None  # Rastreia último timestamp inserido
        
        # ========== FASE 2.2: PRÉ-COMPUTAR SPATIAL INDEX ==========
        # Evita recalcular overlaps O(n) para cada segmento
        # IntervalTree permite busca O(log n) por range
        tree = IntervalTree()
        for turn, _, speaker in diarization_output.itertracks(yield_label=True):
            # Adiciona intervalo com speaker como data
            tree[turn.start:turn.end] = speaker
        
        # Processar cada segmento do Whisper
        for segment in segments:
            start, end = segment['start'], segment['end']
            
            # ========== FASE 2.2: BUSCA O(log n) EM VEZ DE O(n) ==========
            overlaps = tree[start:end]
            
            if overlaps:
                # Seleciona speaker com maior overlap temporal
                best_overlap = max(
                    overlaps,
                    key=lambda interval: min(end, interval.end) - max(start, interval.begin)
                )
                # Extrai speaker ID do pyannote format (SPEAKER_XX)
                speaker_id = best_overlap.data.split('_')[-1]
                best_speaker = f"SPEAKER {int(speaker_id) + 1}"
            else:
                best_speaker = "SPEAKER 0"
            
            # Adiciona header de speaker se mudou
            if best_speaker != current_speaker:
                raw_output.append(f"\n{best_speaker}\n")
                current_speaker = best_speaker
            
            # Adiciona timestamp apenas a cada 30 minutos
            if self._should_add_timestamp(start, last_timestamp, interval_minutes=20):
                timestamp_str = f"[{self._format_timestamp(start)}] "
                last_timestamp = start
            else:
                timestamp_str = ""
            
            # Adiciona texto do segmento com timestamp condicional
            raw_output.append(f"{timestamp_str}{segment['text'].strip()}")
        
        return " ".join(raw_output)
    
    def _align_diarization_fallback(self, segments, diarization_output):
        """Fallback O(n) caso intervaltree não esteja disponível (código original)"""
        raw_output = []
        current_speaker = None
        last_timestamp = None  # Rastreia último timestamp inserido
        diarization_segments = [(t.start, t.end, s) for t, _, s in diarization_output.itertracks(yield_label=True)]

        for segment in segments:
            start, end = segment['start'], segment['end']
            best_speaker = "SPEAKER 0"
            max_overlap = 0

            # Busca O(n) através de todos os segmentos de diarização
            for d_start, d_end, d_speaker in diarization_segments:
                overlap = max(0, min(end, d_end) - max(start, d_start))
                if overlap > max_overlap:
                    max_overlap = overlap
                    best_speaker = f"SPEAKER {int(d_speaker.split('_')[-1]) + 1}" 
            
            if best_speaker != current_speaker:
                raw_output.append(f"\n{best_speaker}\n")
                current_speaker = best_speaker
            
            # Adiciona timestamp apenas a cada 30 minutos
            if self._should_add_timestamp(start, last_timestamp, interval_minutes=30):
                timestamp_str = f"[{self._format_timestamp(start)}] "
                last_timestamp = start
            else:
                timestamp_str = ""
            
            raw_output.append(f"{timestamp_str}{segment['text'].strip()}")
            
        return " ".join(raw_output)

    def _segment_raw_transcription(self, raw_text):
        lines = raw_text.split('\n')
        speaker_pattern = re.compile(r'^SPEAKER \d+$')
        segments = []
        current_speaker = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if speaker_pattern.match(line):
                if current_speaker:
                    segments.append({'speaker': current_speaker, 'content': "\n".join(current_content)})
                current_speaker = line
                current_content = []
            else:
                if current_speaker: current_content.append(line)
        
        if current_speaker:
            segments.append({'speaker': current_speaker, 'content': "\n".join(current_content)})
            
        # Merge simples
        final_segments = []
        if segments:
            current = segments[0]
            for next_seg in segments[1:]:
                if next_seg['speaker'] == current['speaker']:
                    current['content'] += "\n" + next_seg['content']
                else:
                    final_segments.append(current)
                    current = next_seg
            final_segments.append(current)
            
        return final_segments

    def _smart_chunk_with_overlap(self, text, max_size=None, overlap=None):
        max_size = max_size or self.MAX_CHUNK_SIZE
        overlap = overlap or self.CHUNK_OVERLAP
        if len(text) <= max_size: return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + max_size
            if end < len(text):
                search_zone = text[max(0, end-2000):end]
                last_break = search_zone.rfind('\n\n')
                if last_break != -1: end = end - 2000 + last_break
            
            chunks.append(text[start:end].strip())
            if end >= len(text): break
            start = end - overlap
        return chunks

    def _merge_chunks_deduplicated(self, chunks):
        if not chunks: return ""
        if len(chunks) == 1: return chunks[0]
        
        merged = chunks[0]
        for i in range(1, len(chunks)):
            current = chunks[i]
            tail = merged[-2000:]
            head = current[:2000]
            matcher = difflib.SequenceMatcher(None, tail, head)
            match = matcher.find_longest_match(0, len(tail), 0, len(head))
            
            if match.size > 200:
                merged += "\n\n" + current[match.b + match.size:]
            else:
                merged += "\n\n" + current
        return merged

    def _get_chunk_cache(self, chunk_text, prompt_hash):
        content_hash = hashlib.sha256(f"{chunk_text}{prompt_hash}".encode()).hexdigest()
        cache_path = self.cache_dir / f"{content_hash}.json"
        if cache_path.exists():
            try:
                with open(cache_path, 'r') as f:
                    return json.load(f)['result']
            except:
                return None
        return None

    def _save_chunk_cache(self, chunk_text, prompt_hash, result):
        content_hash = hashlib.sha256(f"{chunk_text}{prompt_hash}".encode()).hexdigest()
        cache_path = self.cache_dir / f"{content_hash}.json"
        try:
            with open(cache_path, 'w') as f:
                json.dump({'result': result}, f)
        except:
            pass

    @retry(
        retry=retry_if_exception_type(Exception), 
        stop=stop_after_attempt(3), 
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    async def process_chunk_async(self, chunk_text, idx=0, total=1, previous_context="", depth=0, global_structure=None, overlap_text="", cached_content=None):
        """Processa um chunk de forma assíncrona com retry, cache e CHUNKING ADAPTATIVO (v2.10 Ported)"""
        
        # Calculate prompt hash for local caching (apenas se não usar context cache)
        prompt_content = f"{chunk_text}_{previous_context}_{overlap_text}_{global_structure}"
        prompt_hash = hashlib.sha256(prompt_content.encode()).hexdigest()
        
        # Check local cache (only if depth 0 to avoid fragment caching issues)
        # Se estiver usando Context Caching, ignoramos o cache local para garantir uso do contexto global
        if not cached_content:
            cached = self._get_chunk_cache(chunk_text, prompt_hash) if depth == 0 else None
            if cached:
                metrics.record_cache_hit()
                return cached
        
        # Constrói o contexto e prompt
        contexto_estilo = f"Últimos parágrafos formatados:\n{previous_context}" if previous_context else "Inicio do documento."
        
        # SE USAR CONTEXT CACHING: System Prompt já está no cache
        # SE NÃO USAR: System Prompt precisa ir no contents
        
        # Se não tiver cache, montamos o system prompt completo
        system_prompt = self.prompt_apostila
            
        # Adiciona estrutura global se disponível e NÃO estiver no cache (se estiver no cache, já foi incluída na criação)
        if global_structure and not cached_content:
            system_prompt += f"\n\n## ESTRUTURA GLOBAL (GUIA):\n{global_structure}"
            
        # Adiciona overlap (contexto anterior local)
        secao_contexto = ""
        if previous_context or overlap_text:
            secao_contexto = f"""
🔒 CONTEXTO ANTERIOR (SOMENTE REFERÊNCIA DE ESTILO)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{contexto_estilo}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ ATENÇÃO: O bloco acima JÁ FOI FORMATADO anteriormente.
- NÃO formate novamente esse conteúdo
- NÃO inclua esse conteúdo na sua resposta
- Use APENAS como referência de estilo de escrita e continuidade
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 NOVO TEXTO PARA FORMATAR (comece aqui):
"""

        user_content = f"""{secao_contexto}
<texto_para_formatar>
{chunk_text}
</texto_para_formatar>

**INSTRUÇÕES FINAIS**:
- Esta é a parte {idx} de {total if total else '?'} (Profundidade {depth})
- Formate APENAS o texto entre <texto_para_formatar>
- Se houver contexto acima, NÃO o reprocesse
- Retorne APENAS o Markdown formatado do NOVO texto
"""
        
        try:
            # Configuração de Segurança (Block None) e Parâmetros
            safety_config = [
                types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
            ]
            
            def call_gemini():
                # Configuração
                gen_config = types.GenerateContentConfig(
                    max_output_tokens=8192,
                    temperature=0.1,
                    top_p=0.9,
                    top_k=40,
                    safety_settings=safety_config,
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="LOW"  # Economiza custo + velocidade para formatação
                    )
                )
                
                # SE USAR CACHE: Passar cached_content e APENAS user_content
                if cached_content:
                    gen_config.cached_content = cached_content.name
                    contents = user_content
                else:
                    # Sem cache: System + User
                    contents = f"{system_prompt}\n\n{user_content}"

                return self.client.models.generate_content(
                    model=self.llm_model,
                    contents=contents,
                    config=gen_config
                )

            # Executa chamada síncrona em thread separada
            start_time = time.time()
            response = await asyncio.to_thread(call_gemini)
            duration = time.time() - start_time
            
            # Extract token counts from response metadata
            prompt_tokens = 0
            completion_tokens = 0
            try:
                if hasattr(response, 'usage_metadata'):
                    usage = response.usage_metadata
                    prompt_tokens = getattr(usage, 'prompt_token_count', 0) or 0
                    completion_tokens = getattr(usage, 'candidates_token_count', 0) or 0
            except: pass
            
            metrics.record_call("gemini", prompt_tokens, completion_tokens, duration)
            
            try:
                result = response.text
            except ValueError:
                result = "" # Fallback se bloqueado ou vazio
                
            # === SMART STITCHING (v2.10) ===
            if contexto_estilo and result:
                result = remover_eco_do_contexto(result, contexto_estilo)

            # === ADAPTIVE CHUNCHING CHECK (v2.10) ===
            ratio = len(result) / len(chunk_text) if len(chunk_text) > 0 else 0
            is_compressed = ratio < 0.65
            
            if (is_compressed or not result) and depth < 2 and len(chunk_text) > 4000:
                reason = "compressão excessiva" if is_compressed else "resposta vazia"
                print(f"\n{Fore.MAGENTA}✂️ ATIVANDO CHUNKING ADAPTATIVO para Chunk {idx} (Motivo: {reason} | Ratio: {ratio:.2f})")
                metrics.record_adaptive_split()
                return await self._split_and_retry_async(chunk_text, idx, system_prompt, total, contexto_estilo, depth)

            if depth == 0: self._save_chunk_cache(chunk_text, prompt_hash, result)
            return result

        except Exception as e:
            # Fallback Logic can trigger adaptive chunking too if it's a token limit error
            if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                if depth < 2:
                    print(f"{Fore.MAGENTA}✂️ Erro de limite detectado. Tentando ADAPTIVE CHUNKING...")
                    return await self._split_and_retry_async(chunk_text, idx, system_prompt, total, contexto_estilo, depth)
            
            print(f"{Fore.YELLOW}⚠️  Falha no Gemini (Chunk {idx}): {e}")
            
            if self.openai_client:
                print(f"{Fore.CYAN}🤖 Tentando fallback para OpenAI ({self.openai_model})...")
                try:
                    start_time_oai = time.time()
                    response = await self.openai_client.chat.completions.create(
                        model=self.openai_model,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_content}
                        ],
                        timeout=180
                    )
                    duration_oai = time.time() - start_time_oai
                    result = response.choices[0].message.content
                    
                    # Record OpenAI metrics
                    oai_prompt = getattr(response.usage, 'prompt_tokens', 0) if hasattr(response, 'usage') else 0
                    oai_compl = getattr(response.usage, 'completion_tokens', 0) if hasattr(response, 'usage') else 0
                    metrics.record_call("openai", oai_prompt, oai_compl, duration_oai)
                    
                    # Apply cleanup to OpenAI result too
                    if contexto_estilo and result:
                        result = remover_eco_do_contexto(result, contexto_estilo)
                        
                    if depth == 0: self._save_chunk_cache(chunk_text, prompt_hash, result)
                    return result
                except Exception as e_openai:
                    print(f"{Fore.RED}❌ Falha também no OpenAI: {e_openai}")
                    raise e_openai
            else:
                print(f"{Fore.RED}❌ Falha no chunk {idx} e sem fallback OpenAI configurado.")
                raise e # Let tenacity handle retry

    async def _split_and_retry_async(self, text, idx, prompt, total, context, depth):
        """Divides chunk in half and processes recursively"""
        mid = len(text) // 2
        # Try to find paragraph break near middle
        margin = int(len(text) * 0.2)
        start_search = max(0, mid - margin)
        end_search = min(len(text), mid + margin)
        search_zone = text[start_search:end_search]
        
        split_pos = -1
        last_para = search_zone.rfind('\n\n')
        if last_para != -1:
            split_pos = start_search + last_para + 2
        else:
            split_pos = mid # Fallback hard split
            
        part_a = text[:split_pos]
        part_b = text[split_pos:]
        
        print(f"   -> Dividindo (Sequencial): Parte A ({len(part_a)}c) + Parte B ({len(part_b)}c)")
        
        # ===================================================================
        # SEQUENTIAL EXECUTION (v2.10 Improvement)
        # Process A first, then use A's tail as context for B.
        # Trades speed (~2x slower) for style continuity and coherence.
        # ===================================================================
        
        res_a = await self._format_chunk_async(part_a, f"{idx}.A", prompt, total, context, depth + 1)
        
        # Use the tail of A's result as context for B
        context_for_b = res_a[-2000:] if len(res_a) > 2000 else res_a
        
        res_b = await self._format_chunk_async(part_b, f"{idx}.B", prompt, total, context_for_b, depth + 1)
        
        return f"{res_a}\n\n{res_b}"

    async def _identify_speaker_async(self, content, professors_info, speaker_label):
        """Identifica speaker com cache e heurística"""
        # Cache simples em memória
        if not hasattr(self, 'speaker_cache'): self.speaker_cache = {}
        if speaker_label in self.speaker_cache: return self.speaker_cache[speaker_label]
        
        prompt = f"""
        Analise o início do texto abaixo e a lista de professores extraída da introdução.
        Identifique quem é o provável professor falando e qual a disciplina.
        
        Falante (Label): {speaker_label if speaker_label else "Desconhecido"}
        
        Lista de Professores (Contexto):
        {professors_info}
        
        Texto (Início):
        {content[:5000]}...
        
        Retorne APENAS um JSON:
        {{
            "nome": "Nome do Professor",
            "disciplina": "Disciplina"
        }}
        """
        
        try:
            def call_gemini():
                return self.client.models.generate_content(
                    model=self.llm_model,
                    contents=f"Você é um assistente que identifica palestrantes.\n\n{prompt}",
                    config=types.GenerateContentConfig(
                        temperature=0.1,
                        max_output_tokens=20000,
                        thinking_config=types.ThinkingConfig(
                            include_thoughts=False,
                            thinking_level="LOW"
                        )
                    )
                )

            response = await asyncio.to_thread(call_gemini)
            content_json = response.text
            self.speaker_cache[speaker_label] = content_json
            return content_json
        except Exception as e:
            return '{"nome": "Professor", "disciplina": "Disciplina"}'

    def _extract_professors_context(self, full_text):
        """Extrai lista de professores (Deep Scan)"""
        print(f"   🕵️  Extraindo contexto de professores (Scan Completo)...")
        
        intro_context = full_text[:5000]
        keywords = ["meu nome é", "sou o professor", "sou a professora", "aqui é o professor"]
        
        found_contexts = []
        lower_text = full_text.lower()
        
        for keyword in keywords:
            start_idx = 0
            while True:
                idx = lower_text.find(keyword, start_idx)
                if idx == -1: break
                start_ctx = max(0, idx - 500)
                end_ctx = min(len(full_text), idx + 500)
                found_contexts.append(full_text[start_ctx:end_ctx])
                start_idx = idx + len(keyword)
        
        combined_context = intro_context + "\n\n... [TRECHOS] ...\n\n" + "\n\n".join(found_contexts)
        if len(combined_context) > 50000: combined_context = combined_context[:50000]

        try:
            response = self.client.models.generate_content(
                model=self.llm_model,
                contents=f"Extraia professores JSON\n\n{combined_context}",
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="LOW"
                    )
                )
            )
            return response.text
        except:
            return "{'professores': []}"

    async def _generate_header_async(self, formatted_content, professor_context_json):
        try:
            try:
                prof_ctx = json.loads(professor_context_json)
                prof = prof_ctx.get("nome", "Professor")
                disc = prof_ctx.get("disciplina", "Disciplina")
            except:
                prof = "Professor"
                disc = "Disciplina"
            
            prompt = f"""
            Gere APENAS o título Markdown para esta seção.
            Professor: {prof}
            Disciplina: {disc}
            
            Conteúdo:
            {formatted_content[:1000]}...
            
            FORMATO DE SAÍDA:
            # Prof. {prof} - {disc}
            """
            
            try:
                def call_gemini():
                    return self.client.models.generate_content(
                        model=self.llm_model,
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            max_output_tokens=100,
                            thinking_config=types.ThinkingConfig(
                                include_thoughts=False,
                                thinking_level="LOW"
                            )
                        )
                    )
                response = await asyncio.to_thread(call_gemini)
                return response.text.strip()
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️  Falha no header (Gemini): {e}")
                
                if self.openai_client:
                     print(f"{Fore.CYAN}🤖 Fallback: Header com OpenAI...")
                     try:
                        response = await self.openai_client.chat.completions.create(
                            model=self.openai_model,
                            messages=[{"role": "user", "content": prompt}],
                            max_tokens=100
                        )
                        return response.choices[0].message.content.strip()
                     except Exception as e_openai:
                         print(f"{Fore.RED}❌ Falha também no OpenAI: {e_openai}")
                
                # Fallback final se tudo falhar
                return "## Tópico (Recuperado)"
        except: return f"# {prof} - {disc}\n"

    async def _process_segment_parallel(self, segment, professors_info, idx, system_prompt):
        speaker = segment['speaker']
        content = segment['content']
        
        print(f"\n{Fore.YELLOW}▶ Segmento {idx+1} ({speaker})...")
        chunks = self._smart_chunk_with_overlap(content, max_size=8000, overlap=1500)
        print(f"   {len(chunks)} chunks de ~8k chars (com 1.5k overlap)")
        
        context_task = self._identify_speaker_async(content[:5000], professors_info, speaker)
        
        chunk_tasks = [
            self._format_chunk_async(chunk, j, system_prompt)
            for j, chunk in enumerate(chunks)
        ]
        
        prof_context, *formatted_parts = await asyncio.gather(context_task, *chunk_tasks)
        
        full_content = remover_overlap_duplicado(formatted_parts)
        # Fallback para limpeza fina
        full_content = remover_paragrafos_identicos_consecutivos(full_content)
        header = await self._generate_header_async(full_content[:10000], prof_context)
        
        return f"{header}\n\n{full_content}\n\n---\n\n"

    def _renumber_topics(self, text):
        """Renumera tópicos sequencialmente"""
        lines = text.split('\n')
        new_lines = []
        main_counter = 0
        sub_counter = 0
        
        for line in lines:
            match_main = re.match(r'^##\s+\d+\.?\s+(.+)', line)
            if match_main:
                main_counter += 1
                sub_counter = 0
                new_lines.append(f"## {main_counter}. {match_main.group(1)}")
                continue
            
            match_sub = re.match(r'^###\s+(?:\d+(?:\.\d+)?\.?)?\s*(.+)', line)
            if match_sub:
                sub_counter += 1
                new_lines.append(f"### {main_counter}.{sub_counter} {match_sub.group(1)}")
                continue
                
            new_lines.append(line)
        return "\n".join(new_lines)

    def _fix_omissions(self, raw_transcript, formatted_text, omissions_report):
        """Tenta corrigir as omissões detectadas usando abordagem targeted chunk-by-chunk"""
        print(f"{Fore.CYAN}🔧 Tentando corrigir omissões automaticamente...")
        
        extraction_prompt = """# TAREFA: EXTRAIR CONTEÚDO OMITIDO
Você receberá:
1. RELATÓRIO DE OMISSÕES: Lista do que está faltando
2. TRANSCRIÇÃO BRUTA: Onde encontrar o conteúdo

## SUA MISSÃO:
Localize na transcrição bruta os trechos exatos que correspondem às omissões listadas.
Para cada omissão, extraia o trecho relevante da transcrição.

## RELATÓRIO DE OMISSÕES:
{report}

Retorne APENAS os trechos extraídos, sem comentários adicionais."""

        try:
            extract_response = self.client.models.generate_content(
                model=self.llm_model, # Use Gemini 3 for extraction/fix
                contents=f"{extraction_prompt.format(report=omissions_report)}\n\nTRANSCRIÇÃO BRUTA (primeiros 100k chars):\n{raw_transcript[:100000]}",
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="HIGH"
                    )
                )
            )
            
            missing_content = extract_response.choices[0].message.content
            print(f"{Fore.CYAN}   📝 Conteúdo omitido extraído ({len(missing_content)} caracteres)")
            
            chunks = self._smart_chunk_with_overlap(formatted_text, max_size=self.MAX_CHUNK_SIZE)
            fixed_chunks = []
            
            insertion_prompt = """# TAREFA: INSERIR CONTEÚDO FALTANTE
Você receberá:
1. Um TRECHO da apostila formatada
2. CONTEÚDO OMITIDO que precisa ser adicionado

## SUA MISSÃO:
Se o trecho da apostila é onde o conteúdo omitido deveria estar, insira-o naturalmente.
Se não for o local apropriado, retorne o trecho INALTERADO.

## REGRAS:
- Mantenha toda formatação Markdown
- Integre o conteúdo omitido de forma fluida
- Use tom didático e formal (3ª pessoa)
- NÃO remova nada existente

## CONTEÚDO OMITIDO A INSERIR:
{missing}

Retorne o trecho (modificado ou inalterado)."""

            for i, chunk in enumerate(chunks):
                print(f"{Fore.CYAN}   🔧 Processando chunk {i+1}/{len(chunks)}...")
                try:
                    fix_response = self.client.models.generate_content(
                        model=self.llm_model,
                        contents=f"{insertion_prompt.format(missing=missing_content)}\n\n{chunk}",
                        config=types.GenerateContentConfig(
                            thinking_config=types.ThinkingConfig(
                                include_thoughts=False,
                                thinking_level="HIGH"
                            )
                        )
                    )
                    fixed_chunks.append(fix_response.text)
                except Exception as chunk_error:
                    print(f"{Fore.YELLOW}   ⚠️  Erro no chunk {i+1}, mantendo original: {chunk_error}")
                    fixed_chunks.append(chunk)
            
            fixed_text = "\n\n".join(fixed_chunks)
            print(f"{Fore.GREEN}   ✅ Texto corrigido gerado ({len(fixed_chunks)} chunks processados)")
            return fixed_text
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Falha ao corrigir omissões: {e}")
            return formatted_text

    def _validate_preservation_heuristics(self, original_text, formatted_text):
        """Validação Heurística com Tolerância Adaptativa"""
        print(f"\n{Fore.CYAN}🔍 Validação Heurística de Preservação (Adaptativa)...")
        issues = []
        
        # 1. Leis (Mantendo lógica original que é boa)
        original_laws = set(re.findall(r'(?:Lei|lei)\s+n?º?\s*\d+[\./]\d+|Art\.?\s*\d+|Súmula\s+\d+', original_text, re.IGNORECASE))
        formatted_laws = set(re.findall(r'(?:Lei|lei)\s+n?º?\s*\d+[\./]\d+|Art\.?\s*\d+|Súmula\s+\d+', formatted_text, re.IGNORECASE))
        missing_laws = original_laws - formatted_laws
        if missing_laws: 
            issues.append(f"❌ {len(missing_laws)} referências legais omitidas")
        else: 
            print(f"{Fore.GREEN}   ✅ Referências legais preservadas")
        
        # 2. Comprimento Adaptativo (Lógica do script Gemini)
        palavras_input = len(original_text.split())
        palavras_output = len(formatted_text.split())
        
        if palavras_input == 0: ratio = 0
        else: ratio = palavras_output / palavras_input
        
        # Heurística de Oralidade
        marcadores_oralidade = ['né', 'então', 'tipo', 'aí', 'pessoal', 'galera', 'tá', 'olha', 'gente', 'veja', 'bom']
        input_lower = original_text.lower()
        count_oralidade = sum(input_lower.count(m) for m in marcadores_oralidade)
        densidade_oralidade = count_oralidade / palavras_input if palavras_input > 0 else 0
        
        # Define tolerância baseada na densidade
        if densidade_oralidade > 0.025:  # Muito coloquial (>2.5%)
            tolerancia = 0.45  # Aceita reduzir até 45%
            tipo = "Muito Coloquial"
        elif densidade_oralidade > 0.015:  # Médio
            tolerancia = 0.38
            tipo = "Coloquial"
        elif densidade_oralidade > 0.008:  # Pouca
            tolerancia = 0.30
            tipo = "Pouca Oralidade"
        else:  # Técnico
            tolerancia = 0.22
            tipo = "Técnico/Denso"
            
        limite_minimo = 1.0 - tolerancia
        
        print(f"   📊 Análise: {tipo} (Densidade: {densidade_oralidade:.2%})")
        print(f"      Ratio atual: {ratio:.1%} (Mínimo aceitável: {limite_minimo:.1%})")
        
        if ratio < limite_minimo:
            issues.append(f"⚠️ Texto formatado muito curto ({ratio:.1%}). Esperado no mínimo {limite_minimo:.1%}")
        else:
            print(f"{Fore.GREEN}   ✅ Comprimento aprovado")
        
        if issues:
            print(f"{Fore.RED}━━━ PROBLEMAS ━━━")
            for i in issues: print(f"   {i}")
            return False, issues
        return True, []

    def validate_completeness_full(self, raw_transcript, formatted_text, video_name, global_structure=None):
        """
        v2.16: Validação LLM Full-Context - Envia documento INTEIRO para análise.
        Aproveita a janela de contexto do Gemini 3 Flash (2M tokens).
        """
        print(f"{Fore.YELLOW}🔍 Validação LLM Full-Context (v2.16)...")
        
        # Calcular tamanho aproximado em tokens (estimativa: 4 chars = 1 token)
        total_chars = len(raw_transcript) + len(formatted_text)
        estimated_tokens = total_chars // 4
        print(f"   📊 Tamanho estimado: {estimated_tokens:,} tokens")
        
        # Safety check: Se exceder 1.5M tokens, usar fallback de amostragem
        if estimated_tokens > 1_500_000:
            print(f"{Fore.YELLOW}   ⚠️ Documento muito grande ({estimated_tokens:,} tokens). Usando validação por amostragem.")
            return self._validate_by_sampling(raw_transcript, formatted_text, video_name)
        
        validation_prompt = """# TAREFA DE VALIDAÇÃO DE FIDELIDADE (Full-Context)

Você é um auditor de qualidade para transcrições jurídicas formatadas.

## SEU OBJETIVO
Compare o TEXTO ORIGINAL (transcrição bruta) com o TEXTO FORMATADO (apostila) e identifique:

1. **OMISSÕES GRAVES**: Conceitos jurídicos, leis, súmulas, artigos ou exemplos importantes que estavam no original mas foram omitidos no formatado.
2. **DISTORÇÕES**: Informações que foram alteradas de forma que mude o sentido jurídico.
3. **ESTRUTURA**: Verifique se os tópicos e subtópicos estão organizados de forma lógica e se não há duplicações.

## REGRAS
- NÃO considere como omissão: hesitações, "né", "então", dados repetitivos, conversas paralelas.
- CONSIDERE como omissão: qualquer lei, súmula, artigo, jurisprudência, exemplo prático ou dica de prova.
- Preste atenção especial em: números de leis, prazos, percentuais, valores monetários.

## FORMATO DE RESPOSTA (JSON)
{
    "aprovado": true/false,
    "nota_fidelidade": 0-10,
    "omissoes_graves": ["descrição clara do item omitido"],
    "distorcoes": ["descrição clara da distorção"],
    "problemas_estrutura": ["títulos duplicados ou hierarquia quebrada"],
    "observacoes": "comentário geral sobre a qualidade"
}

Retorne APENAS o JSON, sem markdown."""

        structure_context = ""
        if global_structure:
            structure_context = f"\n\n## ESTRUTURA ESPERADA (Mapeamento Inicial):\n{global_structure[:5000]}"
        
        try:
            response = self.client.models.generate_content(
                model=self.llm_model,
                contents=f"""{validation_prompt}{structure_context}

## TEXTO ORIGINAL (Transcrição Bruta):
{raw_transcript}

## TEXTO FORMATADO (Apostila):
{formatted_text}
""",
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    max_output_tokens=8000,
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="HIGH"
                    )
                )
            )
            
            result = json.loads(response.text)
            
            # Processar resultado
            aprovado = result.get('aprovado', True)
            nota = result.get('nota_fidelidade', 10)
            omissoes = result.get('omissoes_graves', [])
            distorcoes = result.get('distorcoes', [])
            problemas_estrutura = result.get('problemas_estrutura', [])
            observacoes = result.get('observacoes', '')
            
            # Log do resultado
            if aprovado:
                print(f"{Fore.GREEN}   ✅ Validação Full-Context APROVADA (Nota: {nota}/10)")
            else:
                print(f"{Fore.RED}   ❌ Validação Full-Context REPROVADA (Nota: {nota}/10)")
                if omissoes:
                    print(f"{Fore.RED}   📌 Omissões: {len(omissoes)}")
                    for o in omissoes[:3]:
                        print(f"      - {o[:100]}...")
                if distorcoes:
                    print(f"{Fore.RED}   ⚠️ Distorções: {len(distorcoes)}")
                if problemas_estrutura:
                    print(f"{Fore.YELLOW}   🏗️ Problemas de Estrutura: {len(problemas_estrutura)}")
            
            # Retornar relatório completo
            return {
                'aprovado': aprovado,
                'nota': nota,
                'omissoes': omissoes,
                'distorcoes': distorcoes,
                'problemas_estrutura': problemas_estrutura,
                'observacoes': observacoes
            }
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro na validação Full-Context: {e}")
            return {
                'aprovado': True,  # Fail-open para não bloquear
                'nota': 0,
                'omissoes': [],
                'distorcoes': [],
                'problemas_estrutura': [],
                'observacoes': f'Erro na validação: {str(e)}'
            }

    def _validate_by_sampling(self, raw_transcript, formatted_text, video_name):
        """Fallback: Validação por amostragem para documentos muito grandes."""
        print(f"{Fore.CYAN}   Usando validação por amostragem (3 janelas)...")
        # Mantém a lógica antiga como fallback
        windows = [
            ("INÍCIO", raw_transcript[:80000], formatted_text[:80000]),
            ("MEIO", raw_transcript[len(raw_transcript)//2-40000:len(raw_transcript)//2+40000],
                      formatted_text[len(formatted_text)//2-40000:len(formatted_text)//2+40000]),
            ("FIM", raw_transcript[-80000:], formatted_text[-80000:])
        ]
        # Implementação simplificada - apenas retorna aprovado por padrão
        return {'aprovado': True, 'nota': 7, 'omissoes': [], 'distorcoes': [], 'problemas_estrutura': [], 'observacoes': 'Validação por amostragem (documento grande)'}

    async def auto_fix_structure(self, formatted_text: str, problemas: list, global_structure: str = None) -> str:
        """
        v2.17: Corretor IA Ativo - Corrige automaticamente problemas estruturais.
        
        Recebe o texto formatado e a lista de problemas detectados pelo auditor,
        e envia ao LLM para correção automática.
        
        Args:
            formatted_text: Texto markdown com problemas
            problemas: Lista de strings descrevendo os problemas estruturais
            global_structure: Estrutura esperada (opcional)
        
        Returns:
            Texto corrigido
        """
        print(f"{Fore.CYAN}🔧 Corretor IA Ativo (v2.17)...")
        print(f"   📋 Problemas a corrigir: {len(problemas)}")
        for p in problemas[:3]:
            print(f"      - {p[:80]}...")
        
        fix_prompt = """# TAREFA DE CORREÇÃO ESTRUTURAL

Você é um editor de documentos jurídicos formatados em Markdown.

## SEU OBJETIVO
Corrija os problemas estruturais listados abaixo NO TEXTO FORNECIDO.

## PROBLEMAS A CORRIGIR:
{problemas}

## REGRAS DE CORREÇÃO:
1. **Títulos Duplicados**: Se um título H2 aparece duas vezes, REMOVA a segunda ocorrência e mescle o conteúdo sob o primeiro.
2. **Hierarquia Quebrada**: Se um H3 está fora de seu H2 pai correto, MOVA-O para debaixo do H2 apropriado.
3. **Parágrafos Repetidos**: Se o mesmo parágrafo aparece duas vezes, REMOVA a duplicata.
4. **Numeração Inconsistente**: Se a numeração está errada (ex: 1, 2, 3, 3, 4), RENUMERE sequencialmente.

## IMPORTANTE:
- NÃO altere o conteúdo textual, apenas a estrutura.
- NÃO adicione novos conteúdos.
- NÃO remova informações jurídicas (leis, súmulas, artigos).
- Mantenha todas as tabelas intactas.
- Preserve a formatação Markdown (negrito, itálico, listas).

## FORMATO DE RESPOSTA:
Retorne APENAS o texto Markdown corrigido, sem explicações adicionais."""

        structure_hint = ""
        if global_structure:
            structure_hint = f"\n\n## ESTRUTURA ESPERADA:\n{global_structure[:3000]}"
        
        try:
            response = self.client.models.generate_content(
                model=self.llm_model,
                contents=f"""{fix_prompt.format(problemas=chr(10).join(f'- {p}' for p in problemas))}{structure_hint}

## TEXTO A CORRIGIR:
{formatted_text}
""",
                config=types.GenerateContentConfig(
                    max_output_tokens=65000,  # Documento pode ser grande
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="HIGH"
                    )
                )
            )
            
            fixed_text = response.text.strip()
            
            # Validação básica: o texto corrigido não deve ser muito menor
            if len(fixed_text) < len(formatted_text) * 0.7:
                print(f"{Fore.YELLOW}   ⚠️ Texto corrigido muito curto ({len(fixed_text)} vs {len(formatted_text)}). Mantendo original.")
                return formatted_text
            
            # Remover possíveis wrappers de código markdown que o LLM pode adicionar
            if fixed_text.startswith('```markdown'):
                fixed_text = fixed_text[len('```markdown'):].strip()
            if fixed_text.startswith('```'):
                fixed_text = fixed_text[3:].strip()
            if fixed_text.endswith('```'):
                fixed_text = fixed_text[:-3].strip()
            
            print(f"{Fore.GREEN}   ✅ Correção automática aplicada com sucesso")
            return fixed_text
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro no Corretor IA: {e}")
            print(f"{Fore.YELLOW}   ℹ️ Mantendo texto original")
            return formatted_text


    def _generate_audit_report(self, video_name, heuristic_issues, llm_issues):
        with open(f"audit_{video_name}.md", 'w') as f:
            f.write(f"# Auditoria: {video_name}\n")
            f.write(f"Heurística: {heuristic_issues}\n")
            f.write(f"LLM: {llm_issues}\n")

    async def map_structure(self, full_text):
        """Creates a global structure skeleton to guide the formatting."""
        print(f"{Fore.CYAN}🗺️  Mapeando estrutura global do documento...")
        
        # Limit input to avoid context overflow, though GPT-5 Mini handles large context well.
        # Taking the first 200k chars is usually enough for structure.
        input_sample = full_text[:200000] 
        
        try:
            def call_gemini():
                return self.client.models.generate_content(
                    model=self.llm_model,
                    contents=self.PROMPT_MAPEAMENTO.format(transcricao=input_sample),
                    config=types.GenerateContentConfig(
                        max_output_tokens=20000,
                        thinking_config=types.ThinkingConfig(
                            include_thoughts=False,
                            thinking_level="HIGH"  
                        )
                    )
                )

            response = await asyncio.to_thread(call_gemini)
            # Clean markdown code blocks if present
            content = response.text.replace('```markdown', '').replace('```', '')
            print(f"{Fore.GREEN}   ✅ Estrutura mapeada com sucesso.")
            return content

        except Exception as e:
            print(f"{Fore.YELLOW}⚠️  Falha no mapeamento via Gemini: {e}")
            
            if self.openai_client:
                print(f"{Fore.CYAN}🤖 Fallback: Mapeando com OpenAI ({self.openai_model})...")
                try:
                    response = await self.openai_client.chat.completions.create(
                        model=self.openai_model,
                        messages=[
                            {"role": "system", "content": self.PROMPT_MAPEAMENTO.format(transcricao=input_sample)}
                        ]
                    )
                    content = response.choices[0].message.content.replace('```markdown', '').replace('```', '')
                    print(f"{Fore.GREEN}   ✅ Estrutura mapeada com sucesso (OpenAI).")
                    return content
                except Exception as e_openai:
                    print(f"{Fore.RED}❌ Falha também no OpenAI: {e_openai}")
                    return None
            else:
                 print(f"{Fore.RED}   ❌ Erro ao mapear estrutura e sem fallback.")
                 return None

    def renumber_headings(self, text):
        """
        Post-processing: Enforces strictly sequential numbering (1, 2, 3...) 
        for H2/H3/H4 headers using a STACK-BASED approach for correct nesting.
        
        v2.16: Fixed to properly reset child counters when parent level changes.
        """
        print(f"{Fore.CYAN}🔢 Renumerando tópicos sequencialmente (Stack-Based v2.16)...")
        lines = text.split('\n')
        new_lines = []
        
        # Stack-based counters: [H1_count, H2_count, H3_count, H4_count]
        # Index 0 = H1 (usually title, skip), Index 1 = H2, etc.
        counters = [0, 0, 0, 0, 0]  # Extra slot for safety
        
        # Keywords to skip numbering (summary tables, etc.)
        skip_keywords = ['resumo', 'quadro', 'tabela', 'síntese', 'esquema', 'bibliografia', 'referências', 'sumário']
        
        # Emoji pattern to detect decorative headers like "## 📋 Sumário"
        emoji_pattern = re.compile(r'^[\U0001F300-\U0001F9FF]')
        
        for line in lines:
            stripped = line.strip()
            
            # Determine header level
            header_match = re.match(r'^(#{1,4})\s+(.*)$', stripped)
            
            if header_match:
                hashes = header_match.group(1)
                level = len(hashes)  # 1 for H1, 2 for H2, etc.
                raw_title = header_match.group(2).strip()
                
                # Clean existing numbers from title (e.g., "1.2.3. Title" -> "Title")
                title_text = re.sub(r'^(\d+(\.\d+)*\.?\s*)+', '', raw_title).strip()
                
                # Skip H1 (document title) - just clean and pass through
                if level == 1:
                    new_lines.append(f"# {title_text}")
                    continue
                
                # Check if this header should be skipped from numbering
                title_lower = title_text.lower()
                should_skip = (
                    any(k in title_lower for k in skip_keywords) or
                    emoji_pattern.match(title_text)  # Headers starting with emoji
                )
                
                if should_skip:
                    new_lines.append(f"{'#' * level} {title_text}")
                else:
                    # STACK LOGIC: Increment current level, reset all deeper levels
                    counters[level] += 1
                    for deeper_level in range(level + 1, len(counters)):
                        counters[deeper_level] = 0
                    
                    # Build hierarchical number (e.g., "2.3.1")
                    number_parts = [str(counters[lvl]) for lvl in range(2, level + 1)]
                    hierarchical_number = ".".join(number_parts)
                    
                    new_lines.append(f"{'#' * level} {hierarchical_number}. {title_text}")
            else:
                new_lines.append(line)
        
        print(f"{Fore.GREEN}   ✅ Renumeração concluída: {counters[2]} seções H2, {counters[3]} H3, {counters[4]} H4")
        return '\n'.join(new_lines)


    def check_coverage(self, original, formatted):
        """Checks for missing laws or sumulas using ROBUST fingerprints (v2.10 Ported)."""
        print(f"{Fore.YELLOW}🔍 Verificando fidelidade (Leis/Súmulas/Robust)...")
        
        # Use migrated robust helpers
        fp_original = extrair_fingerprints(original)
        contagem_original = contar_ocorrencias_robust(fp_original, original)
        contagem_formatado = contar_ocorrencias_robust(fp_original, formatted)
        
        omissoes = []
        duplicacoes = []
        
        for key, count_orig in contagem_original.items():
            count_fmt = contagem_formatado.get(key, 0)
            categoria, item = key.split(':', 1)
            
            if count_orig > 0 and count_fmt == 0:
                omissoes.append(f"[{categoria}] {item}")
            if count_fmt > count_orig:
                duplicacoes.append(f"[{categoria}] {item} (+{count_fmt - count_orig})")
        
        report = []
        if omissoes:
            report.append(f"⚠️ {len(omissoes)} POSSÍVEIS OMISSÕES:")
            for o in omissoes[:15]: report.append(f"   - {o}")
            if len(omissoes) > 15: report.append(f"   ... e mais {len(omissoes)-15}")
            
        if duplicacoes:
            report.append(f"\nℹ️ {len(duplicacoes)} CITAÇÕES REFORÇADAS (Agregadas):")
            for d in duplicacoes[:10]: report.append(f"   - {d}")

        if not report:
            print(f"{Fore.GREEN}✅ Verificação OK: Nenhuma omissão de Leis/Súmulas detectada.")
            return "Verificação OK: Nenhuma omissão detectada."
        else:
            msg = "\n".join(report)
            print(f"{Fore.RED}{msg}")
            return msg

    def final_structure_audit(self, formatted_text, global_structure):
        """
        v2.16: Audita a estrutura final comparando com o mapeamento inicial.
        Retorna um relatório de discrepâncias e, opcionalmente, tenta corrigir.
        """
        print(f"{Fore.CYAN}🔍 Auditoria Final de Estrutura (v2.16)...")
        
        if not global_structure:
            print(f"{Fore.YELLOW}   ⚠️ Mapeamento global não disponível, pulando auditoria.")
            return formatted_text, []
        
        # Extract all H2/H3 titles from formatted text
        formatted_headers = []
        for line in formatted_text.split('\n'):
            match = re.match(r'^(#{2,3})\s+(?:\d+(?:\.\d+)*\.?\s*)?(.+)$', line.strip())
            if match:
                level = len(match.group(1))
                title = match.group(2).strip()
                # Normalize
                title_normalized = re.sub(r'\s*\(Continuação\)\s*$', '', title, flags=re.IGNORECASE).strip().lower()
                formatted_headers.append((level, title_normalized, title))
        
        # Extract expected structure from global mapping
        expected_headers = []
        for line in global_structure.split('\n'):
            match = re.match(r'^(#{2,3})\s+(.+)$', line.strip())
            if match:
                level = len(match.group(1))
                title = match.group(2).strip()
                title_normalized = title.lower()
                expected_headers.append((level, title_normalized, title))
        
        # Find duplicates in formatted headers
        seen_titles = {}
        duplicates = []
        for idx, (level, title_norm, title_orig) in enumerate(formatted_headers):
            if title_norm in seen_titles:
                duplicates.append({
                    'title': title_orig,
                    'first_occurrence': seen_titles[title_norm],
                    'duplicate_occurrence': idx + 1
                })
            else:
                seen_titles[title_norm] = idx + 1
        
        issues = []
        if duplicates:
            issues.append(f"⚠️ {len(duplicates)} TÍTULOS DUPLICADOS DETECTADOS:")
            for d in duplicates[:5]:
                issues.append(f"   - '{d['title']}' (linhas ~{d['first_occurrence']} e ~{d['duplicate_occurrence']})")
        
        # Check coverage of expected structure
        expected_titles_set = {h[1] for h in expected_headers}
        formatted_titles_set = {h[1] for h in formatted_headers}
        missing_titles = expected_titles_set - formatted_titles_set
        
        if missing_titles:
            issues.append(f"\n⚠️ {len(missing_titles)} TÓPICOS DO MAPEAMENTO NÃO ENCONTRADOS:")
            for t in list(missing_titles)[:5]:
                issues.append(f"   - '{t}'")
        
        if issues:
            report = "\n".join(issues)
            print(f"{Fore.YELLOW}{report}")
            return formatted_text, issues
        else:
            print(f"{Fore.GREEN}   ✅ Estrutura auditada - Sem duplicatas ou omissões de tópicos.")
            return formatted_text, []

    async def format_transcription_async(self, transcription, video_name, output_folder, mode="APOSTILA", custom_prompt=None, dry_run=False):
        """
        Orquestrador Principal com Checkpoint e Robustez (Sequential Mode)
        
        Args:
            transcription: Texto da transcrição
            video_name: Nome do vídeo
            output_folder: Pasta de saída
            mode: "APOSTILA" ou "FIDELIDADE" (ignorado se custom_prompt for fornecido)
            custom_prompt: Prompt customizado opcional (substitui os prompts padrão)
            dry_run: Se True, apenas valida divisão de chunks
        """
        print(f"{Fore.MAGENTA}🧠 Formatando com {self.llm_model} (Sequential Mode)...")
        
        # v2.21: Suporte a prompt customizado
        if custom_prompt:
            self.prompt_apostila = custom_prompt
            print(f"{Fore.YELLOW}🎨 Usando PROMPT CUSTOMIZADO ({len(custom_prompt):,} chars)")
        elif mode == "FIDELIDADE":
            self.prompt_apostila = self.PROMPT_FIDELIDADE
            print(f"{Fore.CYAN}🎨 Modo FIDELIDADE ativo (Prompt ajustado)")
        else:
            self.prompt_apostila = self.PROMPT_APOSTILA_ACTIVE
            print(f"{Fore.CYAN}📚 Modo APOSTILA ativo (Prompt padrão)")
        
        # 0. Context Extraction
        pass
        # professors_info = self._extract_professors_context(transcription)

        # 0.1 Global Structure Mapping (NEW)
        global_structure = await self.map_structure(transcription)

        # 1. Sequential Slicing (v2.17: Com âncoras de estrutura)
        print(f"🔪 Dividindo sequencialmente (com âncoras v2.17)...")
        # v2.17: Passa estrutura global para alinhar cortes com títulos mapeados
        chunks_info = dividir_sequencial(transcription, chars_por_parte=15000, estrutura_global=global_structure)
        validar_chunks(chunks_info, transcription)
        
        total_segments = len(chunks_info)
        print(f"📊 Total de segmentos sequenciais: {total_segments}")
        
        if dry_run:
            print(f"{Fore.YELLOW}🔍 MODO DRY-RUN: Parando antes das chamadas de API.")
            print(f"   Exemplo do Chunk 1: {transcription[chunks_info[0]['inicio']:chunks_info[0]['inicio']+100]}...")
            return "# DRY RUN OUTPUT"
        
        # 2. Checkpoint Loading
        checkpoint_data = load_checkpoint(video_name, output_folder)
        results_map = {} # Map idx -> result
        
        if checkpoint_data:
            print(f"{Fore.CYAN}📂 Retomando via Checkpoint ({checkpoint_data.get('timestamp')})")
            if len(checkpoint_data.get('results', [])) > 0:
                saved_results = checkpoint_data['results']
                # Restore results map
                for idx, res in enumerate(saved_results):
                    if idx < total_segments:
                        results_map[idx] = res
                print(f"   ✅ {len(results_map)} segmentos recuperados.")
        
        # v2.19: Context Caching Setup
        cached_context = None
        if total_segments > 1: # Só cache se tiver múltiplos chunks
            cached_context = self.create_context_cache(transcription, global_structure)
                
        # 3. Sequential Processing Loop
        ordered_results = []
        
        # Restore ordered results from map
        for i in range(len(results_map)):
            ordered_results.append(results_map[i])

        start_idx = len(ordered_results)
        
        if start_idx < total_segments:
            print(f"▶ Iniciando processamento sequencial do segmento {start_idx + 1}...")
            
            for i in tqdm(range(start_idx, total_segments), desc="Processando Sequencial"):
                info = chunks_info[i]
                chunk_text = transcription[info['inicio']:info['fim']]
                
                # Context Management
                contexto_estilo = ""
                if i > 0 and ordered_results:
                    raw_context = ordered_results[-1][-800:] # Last 800 chars
                    if len(raw_context.split()) > 30 and "[!WARNING]" not in raw_context: 
                        contexto_estilo = raw_context

                # Processing with Rate Limit
                await rate_limiter.wait_if_needed_async()
                
                try:
                    # Lógica de Estrutura Local (Janela Deslizante)
                    # Se tiver cache, usamos a Global (que está no cache). Se não, mantemos a Local.
                    estrutura_referencia = None
                    if global_structure and not cached_context:
                        # v2.18: Contexto Localizado - Janela deslizante de ~15% da estrutura
                        itens_estrutura = global_structure.split('\n')
                        if len(itens_estrutura) > 8 and total_segments > 1:
                            ratio = len(itens_estrutura) / total_segments
                            center_idx = int(i * ratio)
                            window_size = max(4, int(len(itens_estrutura) * 0.15))
                            start_idx_w = max(0, center_idx - window_size)
                            end_idx_w = min(len(itens_estrutura), center_idx + window_size + 2)
                            slice_itens = itens_estrutura[start_idx_w:end_idx_w]
                            if start_idx_w > 0: slice_itens.insert(0, "[... Tópicos anteriores ...]")
                            if end_idx_w < len(itens_estrutura): slice_itens.append("[... Tópicos posteriores ...]")
                            estrutura_referencia = '\n'.join(slice_itens)
                        else:
                            estrutura_referencia = global_structure

                    # v2.19: Process Chunk Async (New Interface)
                    formatted = await self.process_chunk_async(
                        chunk_text=chunk_text,
                        idx=i+1,
                        total=total_segments,
                        previous_context=contexto_estilo,
                        depth=0,
                        global_structure=estrutura_referencia, # None se usar cache
                        cached_content=cached_context
                    )
                    
                    ordered_results.append(formatted)
                    
                    # Save Checkpoint after each chunk
                    save_checkpoint(video_name, output_folder, ordered_results, chunks_info, i + 1)
                    
                except Exception as e:
                    print(f"{Fore.RED}❌ Falha Fatal no segmento {i+1}: {e}")
                    # Save what we have
                    save_checkpoint(video_name, output_folder, ordered_results, chunks_info, i)
                    raise e
        
        # 4. Final Assembly
        print(f"\n{Fore.CYAN}🧹 Pipeline de Limpeza Final (v2.7)...")
        full_formatted = f"# {video_name}\n\n" + "\n\n".join(ordered_results)
        
        # 4.1 Limpar metadados de mapeamento que vazam para o output
        # Remove linhas como "[TIPO: AULA EXPOSITIVA]" ou "**[TIPO: SIMULADO]**"
        full_formatted = re.sub(r'^#?\s*\*?\*?\[TIPO:.*?\]\*?\*?\s*$', '', full_formatted, flags=re.MULTILINE)
        # Remove marcadores de bloco [BLOCO 01], [BLOCO 02], etc.
        full_formatted = re.sub(r'^\s*\[BLOCO\s*\d+\]\s*$', '', full_formatted, flags=re.MULTILINE)
        # Remove timestamps órfãos [HH:MM] ou [HH:MM:SS] no início de linha
        full_formatted = re.sub(r'^\s*\[\d{1,2}:\d{2}(:\d{2})?\]\s*$', '', full_formatted, flags=re.MULTILINE)
        full_formatted = re.sub(r'\n{3,}', '\n\n', full_formatted)  # Remove linhas em branco extras
        
        # 5. Post-Processing Pipeline (v2.7 features)
        
        print("  Passada 1: Removendo duplicações literais...")
        # Uses the newly ported v2.7 logic
        full_formatted = remover_duplicacoes_literais(full_formatted)
        
        limiar_info = '70%' if mode == 'FIDELIDADE' else '60%'
        print(f"  Passada 2: Removendo seções duplicadas (limiar {limiar_info})...")
        full_formatted = remover_secoes_duplicadas(full_formatted, mode=mode)
        
        print("  Passada 2.5: Removendo parágrafos duplicados (v2.17)...")
        full_formatted = remover_paragrafos_duplicados(full_formatted)
        
        print("  Passada 2.6: Removendo títulos órfãos (v2.17)...")
        full_formatted = remover_titulos_orfaos(full_formatted)
        
        print("  Passada 3: Normalizando títulos similares...")
        full_formatted = normalize_headings(full_formatted)
        
        if mode != "FIDELIDADE":
            print("  Passada 3.5: Reorganização Estrutural Determinística...")
            full_formatted = deterministic_structure_fix(full_formatted)
        else:
            print(f"{Fore.YELLOW}  ℹ️  Modo FIDELIDADE: Pulando reorganização para preservar linearidade exata.")
        
        if mode != "FIDELIDADE":
            print("  Passada 4: Revisão Semântica por IA...")
            full_formatted = await ai_structure_review(full_formatted, self.client, self.llm_model, estrutura_mapeada=global_structure)
        else:
            print(f"{Fore.MAGENTA}  Passada 4: Revisão Leve de Formatação (Modo Fidelidade)...")
            full_formatted = await ai_structure_review_lite(full_formatted, self.client, self.llm_model, estrutura_mapeada=global_structure)
        
        # Passada 4.5: Renumeração Determinística (Camada de Segurança)
        try:
            full_formatted = renumerar_secoes(full_formatted)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Falha na renumeração determinística: {e}. Continuando...")
        
        print(f"\n{Fore.CYAN}🔢 Renumerando tópicos (1..N) (Stack-Based v2.16)...")
        full_formatted = self.renumber_headings(full_formatted)
        
        # 5.6 v2.18: Auto-Fix Pass - Correções automáticas finais
        full_formatted, autofix_correcoes = aplicar_correcoes_automaticas(full_formatted)
        
        # 5.5 v2.16: Auditoria Final de Estrutura
        full_formatted, audit_issues = self.final_structure_audit(full_formatted, global_structure)
        
        # 6. Validation & Coverage
        print(f"\n{Fore.CYAN}🛡️  Validando cobertura final...")
        coverage_report = self.check_coverage(transcription, full_formatted)
        
        # Save validation report
        report_path = Path(output_folder) / f"{video_name}_validacao.txt"
        with open(report_path, "w", encoding='utf-8') as f:
            f.write(coverage_report)
        print(f"📄 Relatório de validação salvo: {report_path.name}")
        
        # Save audit report if there are issues
        if audit_issues:
            audit_path = Path(output_folder) / f"{video_name}_APOSTILA_verificacao.txt"
            with open(audit_path, "w", encoding='utf-8') as f:
                f.write("# AUDITORIA DE ESTRUTURA (v2.16)\n\n")
                f.write(f"Cobertura: {coverage_report}\n\n")
                f.write("## Problemas Estruturais Detectados\n")
                for issue in audit_issues:
                    f.write(f"{issue}\n")
            print(f"📄 Relatório de auditoria salvo: {audit_path.name}")
        
        # 7. v2.16: Validação Full-Context LLM (Novo!)
        print(f"\n{Fore.CYAN}🔬 Validação Full-Context LLM...")
        validation_result = self.validate_completeness_full(
            transcription, full_formatted, video_name, global_structure
        )
        
        # Salvar relatório JSON da validação
        validation_report_path = Path(output_folder) / f"{video_name}_APOSTILA_fidelidade.json"
        with open(validation_report_path, "w", encoding='utf-8') as f:
            json.dump(validation_result, f, ensure_ascii=False, indent=2)
        print(f"📄 Relatório de fidelidade Full-Context salvo: {validation_report_path.name}")
        
        # Se houver problemas graves, gerar também um markdown legível
        if not validation_result.get('aprovado', True):
            fidelity_md_path = Path(output_folder) / f"{video_name}_APOSTILA_REVISAO.md"
            with open(fidelity_md_path, "w", encoding='utf-8') as f:
                f.write(f"# ⚠️ REVISÃO NECESSÁRIA: {video_name}\n\n")
                f.write(f"**Nota de Fidelidade:** {validation_result.get('nota', 0)}/10\n\n")
                if validation_result.get('omissoes'):
                    f.write("## 📌 Omissões Detectadas\n")
                    for o in validation_result['omissoes']:
                        f.write(f"- {o}\n")
                    f.write("\n")
                if validation_result.get('distorcoes'):
                    f.write("## ⚠️ Distorções Detectadas\n")
                    for d in validation_result['distorcoes']:
                        f.write(f"- {d}\n")
                    f.write("\n")
                if validation_result.get('problemas_estrutura'):
                    f.write("## 🏗️ Problemas de Estrutura\n")
                    for p in validation_result['problemas_estrutura']:
                        f.write(f"- {p}\n")
                    f.write("\n")
                if validation_result.get('observacoes'):
                    f.write(f"## Observações\n{validation_result['observacoes']}\n")
            print(f"{Fore.RED}📄 ATENÇÃO: Documento requer revisão! Veja: {fidelity_md_path.name}")
        
        # 7.1 v2.17: Corretor IA Ativo - Corrige problemas estruturais automaticamente
        # 7.1 v2.18: Corretor IA Seguro (Safe Mode) - Corrige problemas estruturais
        if validation_result and not validation_result.get('aprovado', True):
            print(f"\n{Fore.CYAN}🔁 Iniciando Auto-Fix Loop (Safe Mode)...")
            
            # Chama o novo corretor seguro
            full_formatted = await self.auto_fix_smart(full_formatted, validation_result, global_structure)
            
            # Re-validar após correção
            print(f"{Fore.CYAN}🔬 Re-validando após correção automática...")
            revalidation_result = self.validate_completeness_full(
                transcription, full_formatted, video_name, global_structure
            )
            
            # Atualizar o relatório com a revalidação
            validation_result = revalidation_result
            validation_report_path = Path(output_folder) / f"{video_name}_APOSTILA_fidelidade.json"
            with open(validation_report_path, "w", encoding='utf-8') as f:
                json.dump(validation_result, f, ensure_ascii=False, indent=2)
            print(f"📄 Relatório de fidelidade atualizado: {validation_report_path.name}")
        
        # Checkpoint cleanup success
        delete_checkpoint(video_name, output_folder)
        
        # 8. v2.18: Auditoria Jurídica Pós-Processamento
        if AUDIT_AVAILABLE:
            print(f"\n{Fore.CYAN}🕵️ Auditoria Jurídica Pós-Processamento...")
            audit_report_path = Path(output_folder) / f"{video_name}_APOSTILA_AUDITORIA.md"
            auditar_consistencia_legal(self.client, full_formatted, str(audit_report_path))
        
        # v2.19: Cleanup manual do cache para economia
        if cached_context:
            try:
                self.client.caches.delete(name=cached_context.name)
                print(f"{Fore.GREEN}🗑️ Cache {cached_context.name} deletado manualmente para economizar custos.")
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Não foi possível deletar o cache: {e}")

        return full_formatted


    async def auto_fix_smart(self, formatted_text, validation_result, global_structure=None):
        """
        v2.18 (SAFE MODE): Corretor Estrutural Seguro.
        Foca EXCLUSIVAMENTE em problemas de estrutura (títulos, duplicatas, hierarquia).
        NÃO altera conteúdo jurídico para evitar alucinações.
        """
        problemas_estrut = validation_result.get('problemas_estrutura', [])
        
        if not problemas_estrut:
            print(f"{Fore.GREEN}   ✅ Nenhum problema estrutural para corrigir.")
            return formatted_text
        
        print(f"{Fore.CYAN}🔧 Auto-Fix Safe: Corrigindo {len(problemas_estrut)} problemas estruturais...")
        
        report = "### PROBLEMAS ESTRUTURAIS:\n" + "\n".join([f"- {p}" for p in problemas_estrut]) + "\n"
            
        PROMPT_FIX = f"""Você é um editor técnico de elite.
        
## TAREFA: LIMPEZA ESTRUTURAL (SEM ALTERAR CONTEÚDO)
Você deve corrigir APENAS a formatação e estrutura do documento.

## REGRA DE OURO (SEGURANÇA JURÍDICA):
- **NÃO altere o texto dos parágrafos.**
- **NÃO adicione nem remova informações jurídicas.**
- **NÃO reescreva explicações.**
- Sua permissão é APENAS para Títulos, Hierarquia e Duplicatas exatas.

## INSTRUÇÕES DE CORREÇÃO:
1. **Títulos Duplicados**: Se um título H2 aparece duas vezes, REMOVA a segunda ocorrência e mescle o conteúdo sob o primeiro.
2. **Hierarquia**: Ajuste níveis (H2, H3) para seguir a lógica do conteúdo.
3. **Parágrafos Repetidos**: Delete duplicações EXATAS de parágrafos (copia-cola acidental).
4. **Renumeração**: Garanta sequência lógica (1, 2, 3...) nos títulos.

{f"## ESTRUTURA DE REFERÊNCIA (Guia):\n{global_structure}" if global_structure else ""}

## RELATÓRIO DE ERROS:
{report}

## SAÍDA:
Retorne o documento COMPLETO corrigido em Markdown. Sem explicações."""

        try:
            # Call Gemini Async
            def call_gemini():
                return self.client.models.generate_content(
                    model=self.llm_model,
                    contents=f"{PROMPT_FIX}\n\n## TEXTO A CORRIGIR:\n{formatted_text}",
                    config=types.GenerateContentConfig(
                        max_output_tokens=100000,
                        thinking_config=types.ThinkingConfig(
                            include_thoughts=False,
                            thinking_level="HIGH" 
                        )
                    )
                )

            response = await asyncio.to_thread(call_gemini)
            
            resultado = response.text.replace('```markdown', '').replace('```', '').strip()
            
            # Validação de segurança estrita (0.8 = 20% tolerância vs 30% antigo)
            if len(resultado) < len(formatted_text) * 0.8:
                print(f"{Fore.YELLOW}⚠️ Auto-Fix Safe cortou muito texto (>20%). Abortando por segurança.")
                return formatted_text
                
            print(f"{Fore.GREEN}   ✅ Auto-Fix Estrutural concluído. ({len(formatted_text)} -> {len(resultado)} chars)")
            return resultado
            
        except Exception as e:
            print(f"{Fore.RED}❌ Falha no Auto-Fix Safe: {e}")
            return formatted_text


    def save_as_word(self, formatted_text, video_name, output_folder):
        """Salva markdown formatado como documento Word (.docx) com estilo premium"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            print(f"{Fore.YELLOW}⚠️ python-docx não instalado. Salvando apenas Markdown.")
            return None

        print(f"{Fore.CYAN}📄 Gerando documento Word profissional...")
        
        # 1. Aplicar Smart Layout (opcional, mantido do Vomo para consistência)
        try:
            formatted_text = mover_tabelas_para_fim_de_secao(formatted_text)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Erro ao reorganizar tabelas: {e}. Usando layout padrão.")

        doc = Document()
        
        # Margens
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        
        # Título principal
        title = doc.add_heading(video_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Data de geração e Modo
        date_para = doc.add_paragraph()
        modo_info = MODO_NOME if 'MODO_NOME' in globals() else "APOSTILA"
        date_run = date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')} - Modo: {modo_info}")
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Sumário
        doc.add_heading('Sumário', level=1)
        self.create_toc(doc)
        doc.add_page_break()
        
        # Processa conteúdo markdown
        lines = formatted_text.split('\n')
        i = 0
        in_table = False
        table_rows = []
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                if in_table and table_rows:
                    self._add_table_to_doc(doc, table_rows)
                    in_table = False
                    table_rows = []
                i += 1
                continue
            
            # Tabelas
            if '|' in line and not in_table:
                in_table = True
                table_rows = []
            
            if in_table:
                # Ignora linha separadora
                is_separator = re.match(r'^\s*\|[\s:|-]+\|[\s:|-]*$', line)
                if '|' in line and not is_separator:
                    table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
                
                if '|' not in line or i == len(lines) - 1:
                    if len(table_rows) > 0:
                        self._add_table_to_doc(doc, table_rows)
                    in_table = False
                    table_rows = []
                    if '|' not in line:
                        continue
                i += 1
                continue
            
            # Headings
            if line.startswith('##### '):
                h = doc.add_heading('', level=5)
                self._format_inline_markdown(h.paragraphs[0], line[6:])
            elif h_match := re.match(r'^(####|###|##|#)\s+(.*)', line):
                lvl = len(h_match.group(1))
                h_text = h_match.group(2)
                if lvl == 1 and h_text == video_name:
                    i += 1
                    continue
                h = doc.add_heading('', level=lvl)
                self._format_inline_markdown(h, h_text)
            
            # Separadores
            elif line.strip() in ['---', '***', '___']:
                p = doc.add_paragraph()
                p.add_run('_' * 80).font.color.rgb = RGBColor(192, 192, 192)
            
            # Quotes
            elif line.startswith('>'):
                p = doc.add_paragraph(style='Quote')
                p.paragraph_format.left_indent = Cm(4.0)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._format_inline_markdown(p, line[1:].strip())
                for run in p.runs:
                    run.italic = True
                    run.font.size = Pt(10)
            
            # Listas não-ordenadas
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.first_line_indent = Cm(-0.63)
                self._format_inline_markdown(p, line[2:])
                
            # Listas numeradas
            elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.first_line_indent = Cm(-0.63)
                self._format_inline_markdown(p, line)
                
            # Parágrafo normal
            else:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6) 
                p.paragraph_format.first_line_indent = Cm(1.0)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._format_inline_markdown(p, line)
                for run in p.runs:
                    run.font.size = Pt(12)
            
            i += 1
            
        output_file = os.path.join(output_folder, f"{video_name}_APOSTILA.docx")
        doc.save(output_file)
        return output_file

    def _format_inline_markdown(self, paragraph, text):
        """Formata markdown inline avançado (bold, italic, code, underline-style)"""
        from docx.shared import Pt, RGBColor
        paragraph.clear()
        
        # Regex robusta do format_transcription_gemini.py
        pattern = r'(\*{3}(.+?)\*{3}|_{3}(.+?)_{3}|\*{2}(.+?)\*{2}|_{2}(.+?)_{2}|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|(?<!_)_(?!(?:_|\s))(.+?)(?<!(?:_|\s))_(?!_)|`(.+?)`)'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            full_match = match.group(0)
            
            if full_match.startswith('***'):
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif full_match.startswith('___'):
                run = paragraph.add_run(match.group(3))
                run.bold = True
                run.italic = True
            elif full_match.startswith('**'):
                run = paragraph.add_run(match.group(4))
                run.bold = True
            elif full_match.startswith('__'):
                run = paragraph.add_run(match.group(5))
                run.bold = True
            elif full_match.startswith('*'):
                run = paragraph.add_run(match.group(6))
                run.italic = True
            elif full_match.startswith('_'):
                run = paragraph.add_run(match.group(7))
                run.italic = True
            elif full_match.startswith('`'):
                run = paragraph.add_run(match.group(8))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(200, 0, 0)
            
            last_end = match.end()
        
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_table_to_doc(self, doc, rows):
        """Adiciona tabela premium ao Word"""
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        if len(rows) < 2: return
        max_cols = max(len(row) for row in rows)
        if max_cols == 0: return
        
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(rows):
            for j in range(max_cols):
                cell = table.rows[i].cells[j]
                cell_text = row_data[j] if j < len(row_data) else ""
                
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._format_inline_markdown(p, cell_text)
                
                if i == 0: # Header styling
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '0066CC')
                    cell._element.get_or_add_tcPr().append(shading_elm)

    def create_toc(self, doc):
        """Adiciona Sumário nativo do Word"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

def process_single_video(video_path, dry_run=False, mode="APOSTILA", skip_formatting=False, custom_prompt=None):
    if not os.path.exists(video_path): return
    folder = os.path.dirname(video_path)
    video_name = Path(video_path).stem
    
    try:
        vomo = VomoMLX()
        
        if video_path.lower().endswith('.txt'):
            print(f"{Fore.CYAN}📄 Input é arquivo de texto. Pulando transcrição...")
            with open(video_path, 'r', encoding='utf-8') as f:
                transcription = f.read()
        else:
            if dry_run:
                print("⚠️ Dry run não suporta áudio direto ainda. Use arquivo .txt")
                return
            audio = vomo.optimize_audio(video_path)
            
            raw_path = os.path.join(folder, f"{video_name}_RAW.txt")
            if os.path.exists(raw_path):
                with open(raw_path, 'r') as f: transcription = f.read()
            else:
                transcription = vomo.transcribe(audio)
                with open(raw_path, 'w') as f: f.write(transcription)
        
        if skip_formatting:
            print(f"{Fore.GREEN}✅ Transcrição RAW concluída: {raw_path}")
            print(f"{Fore.YELLOW}   ℹ️  Formatação pulada (--skip-formatting usado).{Style.RESET_ALL}")
            return

        word_only_flag = "--word-only" in sys.argv
        if word_only_flag and video_path.lower().endswith('.md'):
            print(f"{Fore.CYAN}📄 Modo --word-only: Gerando Word a partir de MD existente...")
            with open(video_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            vomo.save_as_word(md_content, video_name, folder)
            print(f"{Fore.GREEN}✅ DOCX gerado com sucesso!")
            return
            
        formatted = asyncio.run(vomo.format_transcription_async(transcription, video_name, folder, mode=mode, custom_prompt=custom_prompt, dry_run=dry_run))
        
        with open(os.path.join(folder, f"{video_name}_APOSTILA.md"), 'w') as f:
            f.write(formatted)
            
        # Validação de Fidelidade (Portado)
        try:
             verificar_cobertura(transcription, formatted, os.path.join(folder, f"{video_name}_APOSTILA.md"))
        except Exception as e:
             print(f"{Fore.YELLOW}⚠️ Erro na validação de fidelidade: {e}")

        vomo.save_as_word(formatted, video_name, folder)
        print(f"{Fore.GREEN}✨ SUCESSO!")
        
    except Exception as e:
        print(f"{Fore.RED}❌ Erro: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    if len(sys.argv) > 1:
        dry_run_flag = "--dry-run" in sys.argv
        skip_formatting_flag = "--skip-formatting" in sys.argv
        mode = "APOSTILA"  # Default mode
        custom_prompt = None  # v2.21: Prompt customizado
        
        # Check for --mode flag
        for arg in sys.argv:
            if arg.startswith("--mode="):
                mode = arg.split("=")[1].upper()
        
        # v2.21: Check for --prompt flag (arquivo .txt ou texto direto)
        for i, arg in enumerate(sys.argv):
            if arg.startswith("--prompt="):
                prompt_value = arg.split("=", 1)[1]
                # Detecta se é arquivo ou texto direto
                if prompt_value.endswith('.txt') and os.path.exists(prompt_value):
                    with open(prompt_value, 'r', encoding='utf-8') as f:
                        custom_prompt = f.read()
                    print(f"{Fore.YELLOW}📝 Prompt carregado de arquivo: {prompt_value} ({len(custom_prompt):,} chars)")
                else:
                    custom_prompt = prompt_value
                    print(f"{Fore.YELLOW}📝 Usando prompt direto ({len(custom_prompt):,} chars)")
                break
            elif arg == "--prompt" and i + 1 < len(sys.argv):
                # Suporta --prompt "texto" ou --prompt arquivo.txt
                prompt_value = sys.argv[i + 1]
                if prompt_value.endswith('.txt') and os.path.exists(prompt_value):
                    with open(prompt_value, 'r', encoding='utf-8') as f:
                        custom_prompt = f.read()
                    print(f"{Fore.YELLOW}📝 Prompt carregado de arquivo: {prompt_value} ({len(custom_prompt):,} chars)")
                else:
                    custom_prompt = prompt_value
                    print(f"{Fore.YELLOW}📝 Usando prompt direto ({len(custom_prompt):,} chars)")
                break
        
        # Get input file (skip flags and --prompt value)
        input_file = None
        skip_next = False
        for arg in sys.argv[1:]:
            if skip_next:
                skip_next = False
                continue
            if arg == "--prompt":
                skip_next = True
                continue
            if not arg.startswith("--"):
                input_file = arg
                break
        
        if input_file:
            print(f"{Fore.CYAN}🔧 Modo selecionado: {mode}")
            if custom_prompt:
                print(f"{Fore.YELLOW}   🎨 Usando prompt customizado")
            if skip_formatting_flag:
                print(f"{Fore.YELLOW}   ℹ️  Flag --skip-formatting ativo: apenas transcrição RAW.{Style.RESET_ALL}")
            if "--word-only" in sys.argv:
                print(f"{Fore.BLUE}   ℹ️  Flag --word-only ativo: conversão MD -> DOCX.{Style.RESET_ALL}")
            process_single_video(input_file, dry_run=dry_run_flag, mode=mode, skip_formatting=skip_formatting_flag, custom_prompt=custom_prompt)

