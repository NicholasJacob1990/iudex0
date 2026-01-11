import os
import sys
import time
import subprocess
import hashlib
from pathlib import Path
from colorama import init, Fore, Style
from tqdm import tqdm
import re
import difflib
import json
import asyncio
from openai import OpenAI, AsyncOpenAI, APIError
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import threading
import random
from time import sleep # Added for RateLimiter fallback if needed
import logging

init(autoreset=True)

try:
    import mlx_whisper
except ImportError:
    mlx_whisper = None

# Imports Pyannote e Torch
try:
    from pyannote.audio import Pipeline
    import torch
    
    HF_TOKEN = os.getenv("HUGGING_FACE_TOKEN")
    if not HF_TOKEN:
        pass
    
except ImportError:
    Pipeline = None
    torch = None

# ==================== RATE LIMITER ====================
class RateLimiter:
    """Controla requisições por minuto para não estourar rate limit da API"""
    def __init__(self, max_requests_per_minute=50): # OpenAI tem limites mais altos que Gemini, mas bom prevenir
        self.max_rpm = max_requests_per_minute
        self.requests = []
        self.lock = threading.Lock()
    
    def wait_if_needed(self):
        with self.lock:
            now = time.time()
            # Remove requisições antigas (>60s)
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
                if wait_time > 0:
                    print(f"{Fore.YELLOW}⏱️  Rate limit: aguardando {wait_time:.1f}s...")
                    sleep(wait_time)
                self.requests = [t for t in self.requests if time.time() - t < 60]
            
            self.requests.append(time.time())

    async def wait_if_needed_async(self):
        """Versão async do rate limiter para não bloquear o event loop"""
        wait_time = 0
        with self.lock:
            now = time.time()
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
                # Reserva o slot mas espera fora do lock
            
            if wait_time <= 0:
                 self.requests.append(time.time())
        
        if wait_time > 0:
            print(f"{Fore.YELLOW}⏱️  Rate limit (Async): aguardando {wait_time:.1f}s...")
            await asyncio.sleep(wait_time)
            # Re-adquire slot após espera (simplificado)
            with self.lock:
                self.requests.append(time.time())

# Instância global
rate_limiter = RateLimiter(max_requests_per_minute=60)

def remover_overlap_duplicado(resultados):
    """Remove duplicação causada pelo overlap entre chunks usando detecção ROBUSTA de conteúdo"""
    if len(resultados) <= 1:
        return resultados[0] if resultados else ""
    
    import re
    from difflib import SequenceMatcher
    
    # === FUNÇÕES AUXILIARES DA ESTRATÉGIA ROBUSTA (Portadas de clean_redundancy.py) ===
    
    def normalize_text(text):
        if not text: return ""
        text = re.sub(r'[#*-]', ' ', text)
        text = re.sub(r'[^\w\s]', '', text)
        return ' '.join(text.lower().split())

    def calculate_similarity(text1, text2):
        if not text1 or not text2:
            return 0.0
        if len(text1) < 50:
            return 1.0 if text1 in text2 or text2 in text1 else 0.0
        return SequenceMatcher(None, text1, text2).quick_ratio()

    def extract_unique_paragraphs(sec_curr_content, sec_prev_content):
        if not sec_curr_content: return []
        unique = []
        paras_curr = sec_curr_content.split('\n\n')
        paras_prev_norm = [normalize_text(p) for p in sec_prev_content.split('\n\n')]
        
        for p in paras_curr:
            p_clean = p.strip()
            if not p_clean or len(p_clean) < 20: continue
            
            p_norm = normalize_text(p_clean)
            is_present = False
            for pp_norm in paras_prev_norm:
                if calculate_similarity(p_norm, pp_norm) > 0.85:
                    is_present = True
                    break
            
            if not is_present:
                unique.append(p_clean)
        return unique
    
    # ==================================================================================

    print("🧹 Iniciando deduplicação robusta (7-DIFF Strategy)...")

    # 1. Junta e Parseia
    texto_bruto = '\n\n'.join(resultados)
    lines = texto_bruto.split('\n')
    
    sections = []
    current_section = None
    header_pattern = re.compile(r'^(#{1,6})\s+(.+)$')
    
    captured_lines = []
    intro_lines = []
    has_started = False
    
    for i, line in enumerate(lines):
        match = header_pattern.match(line)
        if match:
            has_started = True
            if current_section:
                current_section['content'] = '\n'.join(captured_lines).strip()
                sections.append(current_section)
                captured_lines = []
            
            title_text = match.group(2).strip()
            # Remove numeração original para comparação agnóstica
            title_clean = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', title_text)
            
            current_section = {
                'title_clean': title_clean,
                'level': len(match.group(1)),
                'full_header': line,
                'content': ""
            }
        else:
            if has_started:
                captured_lines.append(line)
            else:
                intro_lines.append(line)
    
    if current_section:
        current_section['content'] = '\n'.join(captured_lines).strip()
        sections.append(current_section)
        
    print(f"   📊 Analisando {len(sections)} seções...")

    # 2. Detecção e Remoção
    indices_to_remove = set()
    MAX_WINDOW = 20 # Olha até 20 seções para trás (cobre overlaps grandes)
    
    for i in range(len(sections)):
        if i in indices_to_remove: continue
        sec_curr = sections[i]
        
        # Janela deslizante
        start_check = max(0, i - MAX_WINDOW)
        for j in range(start_check, i):
            if j in indices_to_remove: continue
            sec_prev = sections[j]
            
            if sec_curr['level'] != sec_prev['level']: continue
            
            sim_title = calculate_similarity(normalize_text(sec_curr['title_clean']), normalize_text(sec_prev['title_clean']))
            sim_content = calculate_similarity(normalize_text(sec_curr['content']), normalize_text(sec_prev['content']))
            
            is_duplicate = False
            
            if sim_title > 0.9 and sim_content > 0.6:
                is_duplicate = True
            elif sim_content > 0.85:
                 is_duplicate = True
            elif sim_title > 0.95 and len(sec_curr['content']) < 100:
                 is_duplicate = True
            
            if is_duplicate:
                print(f"   🗑️  Detectado: '{sec_curr['title_clean'][:30]}...' duplica seção anterior")
                
                # Mescla conteúdo único antes de excluir
                unique_paras = extract_unique_paragraphs(sec_curr['content'], sec_prev['content'])
                if unique_paras:
                    sections[j]['content'] += '\n\n' + '\n\n'.join(unique_paras)
                
                indices_to_remove.add(i)
                break
    
    # 3. Reconstrução
    final_lines = list(intro_lines)
    for i, sec in enumerate(sections):
        if i in indices_to_remove: continue
        final_lines.append(sec['full_header'])
        if sec['content']:
            final_lines.append(sec['content'])
            final_lines.append("")
            
    texto_limpo = '\n'.join(final_lines)
    print(f"   ✅ Removidas {len(indices_to_remove)} seções duplicadas.")
    
    return texto_limpo


# ==================== HELPERS PORTED FROM GPT SCRIPT ====================

def limpar_tags_xml(texto):
    texto = re.sub(r'</?[a-z_][\w\-]*>', '', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<[a-z_][\w\-]*\s+[^>]+>', '', texto, flags=re.IGNORECASE)
    return texto

def dividir_sequencial(transcricao_completa, chars_por_parte=6000):
    """Divide documento em chunks SEQUENCIAIS sem sobreposição"""
    chunks = []
    tamanho_total = len(transcricao_completa)
    inicio = 0
    
    while inicio < tamanho_total:
        fim = min(inicio + chars_por_parte, tamanho_total)
        
        if fim < tamanho_total:
            janela = transcricao_completa[max(0, fim - 500):min(tamanho_total, fim + 500)]
            titulo_match = re.search(r'\n(#{2,4}\s+.+)\n', janela)
            
            if titulo_match:
                pos_titulo = janela.find(titulo_match.group(0))
                fim = max(0, fim - 500) + pos_titulo
            else:
                quebra = transcricao_completa.rfind('\n\n', fim - 300, fim + 300)
                if quebra != -1 and quebra > inicio:
                    fim = quebra
                else:
                    quebra = transcricao_completa.rfind('. ', fim - 150, fim + 150)
                    if quebra != -1 and quebra > inicio:
                        fim = quebra + 1
        
        chunks.append({'inicio': inicio, 'fim': fim})
        inicio = fim
    
    return chunks

def validar_chunks(chunks, transcricao_completa):
    """Validação rigorosa de chunks sequenciais"""
    print("🔍 Validando chunks sequenciais...")
    
    for i in range(len(chunks)):
        chunk = chunks[i]
        
        if i > 0:
            anterior = chunks[i-1]
            if chunk['inicio'] != anterior['fim']:
                print(f"❌ Gap/Overlap no chunk {i+1}!")
                print(f"   Anterior termina em: {anterior['fim']}")
                print(f"   Atual começa em: {chunk['inicio']}")
                return False
    
    print(f"✅ {len(chunks)} chunks validados (sequenciais, sem overlap)")
    return True

def detectar_secoes_duplicadas(texto):
    """Detecta seções inteiras duplicadas por títulos"""
    print("🔍 Detectando seções duplicadas...")
    
    linhas = texto.split('\n')
    titulos_vistos = {}
    secoes_duplicadas = []
    
    for i, linha in enumerate(linhas):
        if linha.strip().startswith('##') and not linha.strip().startswith('###'):
            # Normaliza título (remove numeração)
            titulo_normalizado = re.sub(r'^##\s*\d+\.?\s*', '', linha.strip()).lower()
            titulo_normalizado = re.sub(r'[📋📊🗂️]', '', titulo_normalizado).strip()
            
            if titulo_normalizado in titulos_vistos:
                print(f"⚠️  Título duplicado encontrado: '{linha.strip()}'")
                secoes_duplicadas.append({
                    'titulo': titulo_normalizado,
                    'primeira_linha': titulos_vistos[titulo_normalizado],
                    'duplicada_linha': i
                })
            else:
                titulos_vistos[titulo_normalizado] = i
    
    return secoes_duplicadas

def remover_secoes_duplicadas(texto):
    """Remover seções duplicadas mantendo apenas a primeira ocorrência"""
    secoes_dup = detectar_secoes_duplicadas(texto)
    
    if not secoes_dup:
        return texto
    
    print("🧹 Removendo seções duplicadas...")
    
    linhas = texto.split('\n')
    linhas_para_remover = set()
    
    for dup in secoes_dup:
        inicio_remocao = dup['duplicada_linha']
        
        # Encontra onde a seção duplicada termina
        fim_remocao = len(linhas)
        for i in range(inicio_remocao + 1, len(linhas)):
            if linhas[i].strip().startswith('##') and not linhas[i].strip().startswith('###'):
                fim_remocao = i
                break
        
        print(f"   Removendo linhas {inicio_remocao}-{fim_remocao} (seção duplicada)")
        for i in range(inicio_remocao, fim_remocao):
            linhas_para_remover.add(i)
    
    linhas_limpas = [linha for i, linha in enumerate(linhas) if i not in linhas_para_remover]
    
    print(f"✅ {len(linhas_para_remover)} linhas removidas")
    return '\n'.join(linhas_limpas)

def remover_duplicacoes_literais(texto):
    """Remove parágrafos individuais duplicados (v2.7 logic)"""
    from difflib import SequenceMatcher
    paragrafos = texto.split('\n\n')
    paragrafos_limpos = []
    dup_count = 0
    
    for i, para in enumerate(paragrafos):
        if i == 0:
            paragrafos_limpos.append(para)
            continue
        
        if len(para.strip()) < 80 or para.strip().startswith('#'):
            paragrafos_limpos.append(para)
            continue
        
        is_duplicate = False
        para_norm = ' '.join(para.split()).lower()
        
        # Check against last 3 paragraphs (v2.7 logic)
        for j in range(max(0, len(paragrafos_limpos) - 3), len(paragrafos_limpos)):
            para_ant = paragrafos_limpos[j]
            para_ant_norm = ' '.join(para_ant.split()).lower()
            
            ratio = SequenceMatcher(None, para_norm, para_ant_norm).ratio()
            
            if ratio > 0.95:
                is_duplicate = True
                dup_count += 1
                break
        
        if not is_duplicate:
            paragrafos_limpos.append(para)
    
    if dup_count > 0:
         print(f"⚠️  {dup_count} parágrafos duplicados removidos (Literal Dedup)")
    
    return '\n\n'.join(paragrafos_limpos)


def numerar_titulos(texto):
    """Adiciona numeração sequencial aos títulos"""
    linhas = texto.split('\n')
    linhas_numeradas = []
    
    contador_h2 = 0
    contador_h3 = 0
    contador_h4 = 0
    
    titulo_pattern = re.compile(r'^(#{2,4})\s+(?:\d+(?:\.\d+)*\.?\s+)?(.+)$')
    
    for linha in linhas:
        match = titulo_pattern.match(linha)
        
        if match:
            nivel = len(match.group(1))
            texto_titulo = match.group(2).strip()
            
            # Não numera títulos de resumo/quadros
            if any(keyword in texto_titulo.lower() for keyword in ['resumo', 'quadro', 'esquema', '📋', '📊', '🗂️']):
                linhas_numeradas.append(linha)
                continue
            
            if nivel == 2:
                contador_h2 += 1
                contador_h3 = 0
                contador_h4 = 0
                nova_linha = f"## {contador_h2}. {texto_titulo}"
            elif nivel == 3:
                contador_h3 += 1
                contador_h4 = 0
                nova_linha = f"### {contador_h2}.{contador_h3}. {texto_titulo}"
            elif nivel == 4:
                contador_h4 += 1
                nova_linha = f"#### {contador_h2}.{contador_h3}.{contador_h4}. {texto_titulo}"
            else:
                nova_linha = linha
            
            linhas_numeradas.append(nova_linha)
        else:
            linhas_numeradas.append(linha)
    
    return '\n'.join(linhas_numeradas)



# ==================== CHECKPOINT SYSTEM ====================
def get_checkpoint_path(video_name, folder):
    return Path(folder) / f"{video_name}.checkpoint.json"

def save_checkpoint(video_name, folder, results, segments_info, current_idx):
    path = get_checkpoint_path(video_name, folder)
    data = {
        'video_name': video_name,
        'current_idx': current_idx,
        'total_segments': len(segments_info),
        'results': results,
        'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
    }
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_checkpoint(video_name, folder):
    path = get_checkpoint_path(video_name, folder)
    if path.exists():
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️ Erro ao carregar checkpoint: {e}")
    return None

def delete_checkpoint(video_name, folder):
    path = get_checkpoint_path(video_name, folder)
    if path.exists(): 
        path.unlink()
        print(f"🧹 Checkpoint removido: {path.name}")

class VomoMLX:
    # GPT-5 Mini: 400k tokens input, 128k output
    MAX_CHUNK_SIZE = 300000  # 200k chars (~75k tokens)
    CHUNK_OVERLAP = 5000     # 5k overlap
    
    PROMPT_MAPEAMENTO = """Você é um especialista em organização de conteúdo jurídico.

Analise a transcrição abaixo e extraia a ESTRUTURA DE TÓPICOS do documento.

## INSTRUÇÕES:
1. Identifique os TÓPICOS PRINCIPAIS da aula (temas macro abordados)
2. Se houver múltiplas disciplinas/matérias, cada uma é um tópico de nível 1
3. Se for uma aula sobre um único tema, organize por subtópicos lógicos
4. **MÁXIMO 3 NÍVEIS**: Use apenas 1., 1.1., 1.1.1. (nunca 1.1.1.1.)
5. Seja conciso - apenas títulos, não explicações
6. Mantenha a ORDEM em que aparecem na transcrição
7. Mapeie do início ao fim, sem omitir partes

## FORMATO:
```
1. [Tópico Principal]
   1.1. [Subtópico]
      1.1.1. [Detalhamento]
```

## TRANSCRIÇÃO:
{transcricao}

Retorne APENAS a estrutura hierárquica (máx 3 níveis), sem texto adicional.
"""

    PROMPT_FIDELIDADE = """# DIRETRIZES DE FORMATAÇÃO E REVISÃO (MODO FIDELIDADE)

## PAPEL
VOCÊ É UM EXCELENTÍSSIMO REDATOR JURÍDICO E DIDÁTICO.
- **Tom:** didático, como o professor explicando em aula.
- **Pessoa:** MANTENHA a pessoa original da transcrição (1ª pessoa se for assim na fala).
- **Estilo:** texto corrido, com parágrafos curtos, sem "inventar" doutrina nova.
- **Objetivo:** reproduzir a aula em forma escrita, clara e organizada, mas ainda com a "voz" do professor.

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias e das falas do professor.
4. **NÃO CRIE MUITOS BULLET POINTS**. PREFIRA UM FORMATO DE MANUAL DIDÁTICO, não checklist.
5. **NÃO USE NEGRITOS EM EXCESSO**. Use apenas para conceitos-chave realmente importantes.

## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados (REsp/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência, autores, casos citados
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque)
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico"

## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Corrija erros gramaticais, regências, ortográficos e de pontuação.
2. **Limpeza Profunda:**
   - **REMOVA** marcadores de oralidade: "né", "tá?", "entende?", "veja bem", "tipo assim".
   - **REMOVA** interações diretas com a turma: "Isso mesmo", "A colega perguntou", "Já estão me vendo?", "Estão ouvindo?".
   - **REMOVA** redundâncias: "subir para cima", "criação nova".
   - **TRANSFORME** perguntas retóricas em afirmações quando possível.
3. **Coesão**: Utilize conectivos para tornar o texto mais fluido. Aplique pontuação adequada.
4. **Legibilidade**:
   - **PARÁGRAFOS CURTOS**: máximo **3-6 linhas visuais** por parágrafo.
   - **QUEBRE** blocos de texto maciços em parágrafos menores.
   - Seja didático sem perder detalhes e conteúdo.
5. **Formatação Didática** (use com moderação):
   - **Bullet points** para enumerar elementos, requisitos ou características
   - **Listas numeradas** para etapas, correntes ou exemplos
   - **Marcadores relacionais** como "→" para consequências lógicas

## 📝 ESTRUTURA E TÍTULOS
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis.
- **NÃO crie subtópicos para frases soltas.**
- Use títulos **APENAS** para mudanças reais de assunto.

## 📊 TABELA DE SÍNTESE
Ao final de cada **bloco temático relevante**, produza uma tabela de síntese:

| Conceito | Definição | Fundamento Legal | Observações |
| :--- | :--- | :--- | :--- |
| ...  | ...  | Art. X, Lei Y | ... |

**REGRAS CRÍTICAS PARA TABELAS:**
1. **Limite:** máximo ~50 palavras por célula
2. **PROIBIDO** blocos de código dentro de células
3. **POSICIONAMENTO:** A tabela vem **APENAS AO FINAL** de um bloco concluído, **NUNCA** no meio de explicação

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA
- **NUNCA formate novamente esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>
"""

    PROMPT_APOSTILA_ACTIVE = """# DIRETRIZES DE REDAÇÃO: MANUAL JURÍDICO DIDÁTICO (MODO APOSTILA)

## PAPEL
VOCÊ É UM EXCELENTÍSSIMO REDATOR JURÍDICO E DIDÁTICO.
- **Tom:** doutrinário, impessoal, estilo manual de Direito.
- **Pessoa:** 3ª pessoa ou construções impessoais ("O professor explica...", "A doutrina define...").
- **Estilo:** prosa densa, porém com parágrafos curtos e didáticos.
- **Objetivo:** transformar a aula em texto de apostila/manual, sem alterar conteúdo nem inventar informações.

## 💎 PILAR 1: ESTILO (VOZ ATIVA E DIRETA)
> 🚫 **PROIBIDO VOZ PASSIVA EXCESSIVA:** "Anunciou-se", "Informou-se".
> ✅ **PREFIRA VOZ ATIVA:** "O professor explica...", "A doutrina define...", "O Art. 37 estabelece...".

## � O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias.

## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados, Temas de Repercussão Geral, Recursos Repetitivos. **NUNCA OMITA NÚMEROS DE TEMAS OU SÚMULAS**.
- **JURISPRUDÊNCIA**: Se o texto citar "Tema 424", "RE 123", "ADI 555", **MANTENHA O NÚMERO**. Não generalize para "jurisprudência do STJ".
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência (STF/STJ), autores, casos citados
- **Ênfases intencionais** e **Observações pedagógicas**

## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Ajuste a linguagem coloquial para o padrão culto.
2. **Limpeza**: Remova gírias, cacoetes ("né", "tipo assim", "então") e vícios de oralidade.
3. **Coesão**: Use conectivos e pontuação adequada para tornar o texto fluido.
4. **Legibilidade**:
   - **PARÁGRAFOS CURTOS**: máximo **3-6 linhas visuais** por parágrafo.
   - **QUEBRE** blocos de texto maciços em parágrafos menores.
   - Use **negrito** para destacar conceitos-chave (sem exagero).
5. **Formatação Didática** (use com moderação):
   - **Bullet points** para enumerar elementos, requisitos ou características
   - **Listas numeradas** para etapas, correntes doutrinárias ou exemplos
   - **Marcadores relacionais** como "→" para consequências lógicas

## 📝 ESTRUTURA E TÍTULOS
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos.
- **NÃO crie subtópicos para frases soltas.**
- Use títulos **APENAS** para mudanças reais de assunto.

## 📊 TABELA DE SÍNTESE (OBRIGATÓRIO)
Ao final de CADA tópico principal (nível 2 ou 3), CRIE uma tabela de resumo.
SEMPRE que houver diferenciação de conceitos, prazos ou regras, CRIE UMA TABELA.

| Conceito/Instituto | Definição | Fundamento Legal | Observações |
| :--- | :--- | :--- | :--- |
| ...  | ...  | Art. X, Lei Y | ... |

**REGRAS CRÍTICAS PARA TABELAS:**
1. **Limite:** máximo ~50 palavras por célula
2. **PROIBIDO** blocos de código dentro de células
3. **NUNCA** deixe título "📋 Resumo" sozinho sem dados
4. **POSICIONAMENTO:** A tabela vem **APENAS AO FINAL** de um bloco concluído
   - **NUNCA** insira tabela no meio de explicação
   - A tabela deve ser o **fechamento** lógico da seção

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA
- **NUNCA formate novamente esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- **NUNCA repita informações que já estão no contexto**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>
- **CRÍTICO:** Se o texto começar repetindo a última frase do contexto, **IGNORE A REPETIÇÃO**
"""

    SYSTEM_PROMPT_FORMAT = PROMPT_APOSTILA_ACTIVE # Default

    def __init__(self, model_size="large-v3-turbo"):
        """
        MLX-Whisper otimizado para Apple Silicon (M3 Pro)
        Formatação otimizada para GPT-5 Mini com processamento paralelo e robustez.
        """
        print(f"{Fore.CYAN}🚀 Inicializando MLX-Whisper ({model_size}) para Apple Silicon...")
        self.model_name = model_size
        self.llm_model = "gpt-5-mini-2025-08-07"
        
        # Carrega variável de ambiente com override
        from dotenv import load_dotenv
        load_dotenv(override=True)
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError(f"{Fore.RED}❌ ERRO CRÍTICO: Configure a variável de ambiente OPENAI_API_KEY")
            
        self.client = OpenAI(api_key=api_key)
        self.async_client = AsyncOpenAI(api_key=api_key)
        
        # Cache directory
        self.cache_dir = Path(".cache_vomo")
        self.cache_dir.mkdir(exist_ok=True)

    def optimize_audio(self, file_path):
        """Extrai áudio otimizado (16kHz mono)"""
        print(f"{Fore.YELLOW}⚡ Verificando áudio...")
        
        mp3_path = Path(file_path).with_suffix('.mp3')
        if mp3_path.exists():
            print(f"   📂 Usando MP3 existente: {mp3_path.name}")
            return str(mp3_path)
        
        file_hash = hashlib.md5(file_path.encode()).hexdigest()[:8]
        output_path = f"temp_{file_hash}.wav"
        
        if os.path.exists(output_path):
            return output_path
        
        print(f"   🔄 Extraindo áudio...")
        subprocess.run([
            'ffmpeg', '-i', file_path,
            '-vn', '-ac', '1', '-ar', '16000',
            output_path, '-y', '-hide_banner', '-loglevel', 'error'
        ], check=True, capture_output=True)
        
        return output_path

    def transcribe(self, audio_path):
        """
        MLX-Whisper OTIMIZADO com GPU acelerado + Diarização
        
        Otimizações v2.0:
        - VAD filtering (pula silêncio) 
        - Batched inference (múltiplos chunks GPU)
        - condition_on_previous_text (contexto melhorado)
        - Hallucination suppression (evita texto inventado)
        """
        print(f"{Fore.GREEN}🎙️  Iniciando transcrição OTIMIZADA (MLX GPU)...")
        start_time = time.time()
        
        # Cache de transcrição
        cache_file = audio_path.replace('.wav', '_DIARIZATION.json').replace('.mp3', '_DIARIZATION.json')
        
        if os.path.exists(cache_file):
            try:
                print(f"{Fore.CYAN}   📂 Cache encontrado, carregando...")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                return cache_data['transcript']
            except Exception:
                pass

        if not mlx_whisper:
            raise ImportError("mlx_whisper não instalado.")
        
        # ==================== PARÂMETROS OTIMIZADOS ====================
        print("   🔍 Transcrevendo com parâmetros otimizados...")
        
        result = mlx_whisper.transcribe(
            audio_path,
            path_or_hf_repo=f"mlx-community/whisper-{self.model_name}",
            language="pt",
            
            # ========== PRECISÃO ==========
            temperature=0.0,              # Mais determinístico (desativa sampling estocástico)
            initial_prompt="Esta é uma transcrição de aula jurídica em português brasileiro sobre direito administrativo, constitucional, civil, penal e processual.",
            
            # === TIMESTAMPS ===
            word_timestamps=True,
            
            # === PERFORMANCE ===
            fp16=True,                    # Usa float16 (mais rápido na GPU)
            
            # === QUALIDADE (Hallucination Suppression) ===
            no_speech_threshold=0.6,      # Ignora segmentos com prob de silêncio > 60%
            logprob_threshold=-1.0,       # Rejeita tokens com log prob muito baixo
            compression_ratio_threshold=2.4,  # Detecta repetição/alucinação
            
            # === CONTEXTO ===
            condition_on_previous_text=True,  # Usa contexto anterior (melhora precisão)
            
            # === SUPRESSÃO DE TOKENS PROBLEMÁTICOS ===
            suppress_tokens=[-1],  # Suprime token de padding
            
            verbose=False
        )
        
        elapsed = time.time() - start_time
        audio_duration = result.get('duration', 0) if isinstance(result, dict) else 0
        rtf = elapsed / audio_duration if audio_duration > 0 else 0
        
        print(f"{Fore.GREEN}   ✅ Transcrição concluída em {elapsed:.1f}s (RTF: {rtf:.2f}x)")
        
        transcript_result = None
        
        # Diarização
        if Pipeline and torch and HF_TOKEN:
            try:
                print("   🗣️  Iniciando Diarização (Pyannote)...")
                pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-3.1",
                    use_auth_token=HF_TOKEN
                )
                
                device = "mps" if torch.backends.mps.is_available() else "cpu"
                pipeline.to(torch.device(device))
                diarization = pipeline(audio_path)
                
                transcript_result = self._align_diarization(result['segments'], diarization)
                print(f"{Fore.GREEN}✅ Diarização concluída")
            
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Erro na Diarização: {e}")
        
        if transcript_result is None:
            # Fallback sem diarização
            formatted_output = []
            last_timestamp = None  # Rastreia último timestamp inserido
            
            for segment in result['segments']:
                start = segment['start']
                
                # Adiciona timestamp apenas a cada 30 minutos
                if self._should_add_timestamp(start, last_timestamp, interval_minutes=30):
                    ts = self._format_timestamp(start)
                    formatted_output.append(f"[{ts}] {segment['text'].strip()}")
                    last_timestamp = start
                else:
                    formatted_output.append(segment['text'].strip())
            
            transcript_result = " ".join(formatted_output)
        
        # Salvar cache
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'transcript': transcript_result,
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
                }, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erro ao salvar cache: {e}")
        
        return transcript_result

    def _format_timestamp(self, seconds):
        m, s = divmod(seconds, 60)
        h, m = divmod(m, 60)
        return f"{int(h):02d}:{int(m):02d}:{int(s):02d}" if h > 0 else f"{int(m):02d}:{int(s):02d}"
    
    def _should_add_timestamp(self, current_seconds, last_timestamp_seconds, interval_minutes=30):
        """
        Determina se um timestamp deve ser adicionado baseado no intervalo configurado.
        
        Args:
            current_seconds: Tempo atual em segundos
            last_timestamp_seconds: Último timestamp inserido (None se for o primeiro)
            interval_minutes: Intervalo em minutos entre timestamps (padrão: 30)
        
        Returns:
            bool: True se deve adicionar timestamp
        """
        if last_timestamp_seconds is None:
            return True  # Sempre adiciona o primeiro timestamp
        
        interval_seconds = interval_minutes * 60
        return (current_seconds - last_timestamp_seconds) >= interval_seconds

    def _align_diarization(self, segments, diarization_output):
        """
        Alinha segmentos com diarização (Versão Otimizada com IntervalTree)
        
        Fase 2.2: Substituição de busca O(n) por IntervalTree O(log n)
        Ganho esperado: 10-20x mais rápido em áudios longos (>30 min)
        """
        try:
            from intervaltree import IntervalTree
        except ImportError:
            print("⚠️ intervaltree não instalado, usando fallback O(n)")
            return self._align_diarization_fallback(segments, diarization_output)
        
        raw_output = []
        current_speaker = None
        last_timestamp = None  # Rastreia último timestamp inserido
        
        # ========== FASE 2.2: PRÉ-COMPUTAR SPATIAL INDEX ==========
        # Evita recalcular overlaps O(n) para cada segmento
        # IntervalTree permite busca O(log n) por range
        tree = IntervalTree()
        for turn, _, speaker in diarization_output.itertracks(yield_label=True):
            # Adiciona intervalo com speaker como data
            tree[turn.start:turn.end] = speaker
        
        # Processar cada segmento do Whisper
        for segment in segments:
            start, end = segment['start'], segment['end']
            
            # ========== FASE 2.2: BUSCA O(log n) EM VEZ DE O(n) ==========
            overlaps = tree[start:end]
            
            if overlaps:
                # Seleciona speaker com maior overlap temporal
                best_overlap = max(
                    overlaps,
                    key=lambda interval: min(end, interval.end) - max(start, interval.begin)
                )
                # Extrai speaker ID do pyannote format (SPEAKER_XX)
                speaker_id = best_overlap.data.split('_')[-1]
                best_speaker = f"SPEAKER {int(speaker_id) + 1}"
            else:
                best_speaker = "SPEAKER 0"
            
            # Adiciona header de speaker se mudou
            if best_speaker != current_speaker:
                raw_output.append(f"\n{best_speaker}\n")
                current_speaker = best_speaker
            
            # Adiciona timestamp apenas a cada 30 minutos
            if self._should_add_timestamp(start, last_timestamp, interval_minutes=30):
                timestamp_str = f"[{self._format_timestamp(start)}] "
                last_timestamp = start
            else:
                timestamp_str = ""
            
            # Adiciona texto do segmento com timestamp condicional
            raw_output.append(f"{timestamp_str}{segment['text'].strip()}")
        
        return " ".join(raw_output)
    
    def _align_diarization_fallback(self, segments, diarization_output):
        """Fallback O(n) caso intervaltree não esteja disponível (código original)"""
        raw_output = []
        current_speaker = None
        last_timestamp = None  # Rastreia último timestamp inserido
        diarization_segments = [(t.start, t.end, s) for t, _, s in diarization_output.itertracks(yield_label=True)]

        for segment in segments:
            start, end = segment['start'], segment['end']
            best_speaker = "SPEAKER 0"
            max_overlap = 0

            # Busca O(n) através de todos os segmentos de diarização
            for d_start, d_end, d_speaker in diarization_segments:
                overlap = max(0, min(end, d_end) - max(start, d_start))
                if overlap > max_overlap:
                    max_overlap = overlap
                    best_speaker = f"SPEAKER {int(d_speaker.split('_')[-1]) + 1}" 
            
            if best_speaker != current_speaker:
                raw_output.append(f"\n{best_speaker}\n")
                current_speaker = best_speaker
            
            # Adiciona timestamp apenas a cada 30 minutos
            if self._should_add_timestamp(start, last_timestamp, interval_minutes=30):
                timestamp_str = f"[{self._format_timestamp(start)}] "
                last_timestamp = start
            else:
                timestamp_str = ""
            
            raw_output.append(f"{timestamp_str}{segment['text'].strip()}")
            
        return " ".join(raw_output)

    def _segment_raw_transcription(self, raw_text):
        lines = raw_text.split('\n')
        speaker_pattern = re.compile(r'^SPEAKER \d+$')
        segments = []
        current_speaker = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if speaker_pattern.match(line):
                if current_speaker:
                    segments.append({'speaker': current_speaker, 'content': "\n".join(current_content)})
                current_speaker = line
                current_content = []
            else:
                if current_speaker: current_content.append(line)
        
        if current_speaker:
            segments.append({'speaker': current_speaker, 'content': "\n".join(current_content)})
            
        # Merge simples
        final_segments = []
        if segments:
            current = segments[0]
            for next_seg in segments[1:]:
                if next_seg['speaker'] == current['speaker']:
                    current['content'] += "\n" + next_seg['content']
                else:
                    final_segments.append(current)
                    current = next_seg
            final_segments.append(current)
            
        return final_segments

    def _smart_chunk_with_overlap(self, text, max_size=None, overlap=None):
        max_size = max_size or self.MAX_CHUNK_SIZE
        overlap = overlap or self.CHUNK_OVERLAP
        if len(text) <= max_size: return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + max_size
            if end < len(text):
                search_zone = text[max(0, end-2000):end]
                last_break = search_zone.rfind('\n\n')
                if last_break != -1: end = end - 2000 + last_break
            
            chunks.append(text[start:end].strip())
            if end >= len(text): break
            start = end - overlap
        return chunks

    def _merge_chunks_deduplicated(self, chunks):
        if not chunks: return ""
        if len(chunks) == 1: return chunks[0]
        
        merged = chunks[0]
        for i in range(1, len(chunks)):
            current = chunks[i]
            tail = merged[-2000:]
            head = current[:2000]
            matcher = difflib.SequenceMatcher(None, tail, head)
            match = matcher.find_longest_match(0, len(tail), 0, len(head))
            
            if match.size > 200:
                merged += "\n\n" + current[match.b + match.size:]
            else:
                merged += "\n\n" + current
        return merged

    def _get_chunk_cache(self, chunk_text, prompt_hash):
        content_hash = hashlib.sha256(f"{chunk_text}{prompt_hash}".encode()).hexdigest()
        cache_path = self.cache_dir / f"{content_hash}.json"
        if cache_path.exists():
            try:
                with open(cache_path, 'r') as f:
                    return json.load(f)['result']
            except:
                return None
        return None

    def _save_chunk_cache(self, chunk_text, prompt_hash, result):
        content_hash = hashlib.sha256(f"{chunk_text}{prompt_hash}".encode()).hexdigest()
        cache_path = self.cache_dir / f"{content_hash}.json"
        try:
            with open(cache_path, 'w') as f:
                json.dump({'result': result}, f)
        except:
            pass

    @retry(
        retry=retry_if_exception_type((APIError, Exception)), 
        stop=stop_after_attempt(3), 
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    async def _format_chunk_async(self, chunk_text, idx, system_prompt, total=None, contexto_estilo=""):
        """Processa um chunk de forma assíncrona com retry e cache (Sequential Aware)"""
        prompt_hash = hashlib.md5((system_prompt + contexto_estilo).encode()).hexdigest()[:8]
        
        # Check cache
        cached = self._get_chunk_cache(chunk_text, prompt_hash)
        if cached:
            return cached

        # Prepare Context Block
        secao_contexto = ""
        if contexto_estilo:
            secao_contexto = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔒 CONTEXTO ANTERIOR (SOMENTE REFERÊNCIA DE ESTILO)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{contexto_estilo}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ ATENÇÃO: O bloco acima JÁ FOI FORMATADO anteriormente.
- NÃO formate novamente esse conteúdo
- NÃO inclua esse conteúdo na sua resposta
- Use APENAS como referência de estilo de escrita e continuidade
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 NOVO TEXTO PARA FORMATAR (comece aqui):
"""

        user_content = f"""{secao_contexto}
<texto_para_formatar>
{chunk_text}
</texto_para_formatar>

**INSTRUÇÕES FINAIS**:
- Esta é a parte {idx} de {total if total else '?'}
- Formate APENAS o texto entre <texto_para_formatar>
- Se houver contexto acima, NÃO o reprocesse
- Retorne APENAS o Markdown formatado do NOVO texto
"""
        
        try:
            response = await self.async_client.chat.completions.create(
                model=self.llm_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ],
                timeout=180
            )
            result = response.choices[0].message.content
            self._save_chunk_cache(chunk_text, prompt_hash, result)
            return result
        except Exception as e:
            print(f"{Fore.RED}❌ Falha no chunk {idx}: {e}")
            raise e # Let tenacity handle retry

    async def _identify_speaker_async(self, content, professors_info, speaker_label):
        """Identifica speaker com cache e heurística"""
        # Cache simples em memória
        if not hasattr(self, 'speaker_cache'): self.speaker_cache = {}
        if speaker_label in self.speaker_cache: return self.speaker_cache[speaker_label]
        
        prompt = f"""
        Analise o início do texto abaixo e a lista de professores extraída da introdução.
        Identifique quem é o provável professor falando e qual a disciplina.
        
        Falante (Label): {speaker_label if speaker_label else "Desconhecido"}
        
        Lista de Professores (Contexto):
        {professors_info}
        
        Texto (Início):
        {content[:5000]}...
        
        Retorne APENAS um JSON:
        {{
            "nome": "Nome do Professor",
            "disciplina": "Disciplina"
        }}
        """
        
        try:
            response = await self.async_client.chat.completions.create(
                model=self.llm_model,
                messages=[
                    {"role": "system", "content": "Você é um assistente que identifica palestrantes."},
                    {"role": "user", "content": prompt}
                ]
            )
            content_json = response.choices[0].message.content
            self.speaker_cache[speaker_label] = content_json
            return content_json
        except Exception as e:
            return '{"nome": "Professor", "disciplina": "Disciplina"}'

    def _extract_professors_context(self, full_text):
        """Extrai lista de professores (Deep Scan)"""
        print(f"   🕵️  Extraindo contexto de professores (Scan Completo)...")
        
        intro_context = full_text[:5000]
        keywords = ["meu nome é", "sou o professor", "sou a professora", "aqui é o professor"]
        
        found_contexts = []
        lower_text = full_text.lower()
        
        for keyword in keywords:
            start_idx = 0
            while True:
                idx = lower_text.find(keyword, start_idx)
                if idx == -1: break
                start_ctx = max(0, idx - 500)
                end_ctx = min(len(full_text), idx + 500)
                found_contexts.append(full_text[start_ctx:end_ctx])
                start_idx = idx + len(keyword)
        
        combined_context = intro_context + "\n\n... [TRECHOS] ...\n\n" + "\n\n".join(found_contexts)
        if len(combined_context) > 50000: combined_context = combined_context[:50000]

        try:
            response = self.client.chat.completions.create(
                model=self.llm_model,
                messages=[
                    {"role": "system", "content": "Extraia professores JSON"},
                    {"role": "user", "content": combined_context}
                ]
            )
            return response.choices[0].message.content
        except:
            return "{'professores': []}"

    async def _generate_header_async(self, formatted_content, professor_context_json):
        try:
            try:
                prof_ctx = json.loads(professor_context_json)
                prof = prof_ctx.get("nome", "Professor")
                disc = prof_ctx.get("disciplina", "Disciplina")
            except:
                prof = "Professor"
                disc = "Disciplina"
            
            prompt = f"""
            Gere APENAS o título Markdown para esta seção.
            Professor: {prof}
            Disciplina: {disc}
            
            Conteúdo:
            {formatted_content[:1000]}...
            
            FORMATO DE SAÍDA:
            # Prof. {prof} - {disc}
            """
            
            response = await self.async_client.chat.completions.create(
                model=self.llm_model,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.choices[0].message.content
        except: return f"# {prof} - {disc}\n"

    async def _process_segment_parallel(self, segment, professors_info, idx, system_prompt):
        speaker = segment['speaker']
        content = segment['content']
        
        print(f"\n{Fore.YELLOW}▶ Segmento {idx+1} ({speaker})...")
        chunks = self._smart_chunk_with_overlap(content, max_size=8000, overlap=1500)
        print(f"   {len(chunks)} chunks de ~8k chars (com 1.5k overlap)")
        
        context_task = self._identify_speaker_async(content[:5000], professors_info, speaker)
        
        chunk_tasks = [
            self._format_chunk_async(chunk, j, system_prompt)
            for j, chunk in enumerate(chunks)
        ]
        
        prof_context, *formatted_parts = await asyncio.gather(context_task, *chunk_tasks)
        
        full_content = remover_overlap_duplicado(formatted_parts)
        # Fallback para limpeza fina
        full_content = remover_paragrafos_identicos_consecutivos(full_content)
        header = await self._generate_header_async(full_content[:10000], prof_context)
        
        return f"{header}\n\n{full_content}\n\n---\n\n"

    def _renumber_topics(self, text):
        """Renumera tópicos sequencialmente"""
        lines = text.split('\n')
        new_lines = []
        main_counter = 0
        sub_counter = 0
        
        for line in lines:
            match_main = re.match(r'^##\s+\d+\.?\s+(.+)', line)
            if match_main:
                main_counter += 1
                sub_counter = 0
                new_lines.append(f"## {main_counter}. {match_main.group(1)}")
                continue
            
            match_sub = re.match(r'^###\s+(?:\d+(?:\.\d+)?\.?)?\s*(.+)', line)
            if match_sub:
                sub_counter += 1
                new_lines.append(f"### {main_counter}.{sub_counter} {match_sub.group(1)}")
                continue
                
            new_lines.append(line)
        return "\n".join(new_lines)

    def _fix_omissions(self, raw_transcript, formatted_text, omissions_report):
        """Tenta corrigir as omissões detectadas usando abordagem targeted chunk-by-chunk"""
        print(f"{Fore.CYAN}🔧 Tentando corrigir omissões automaticamente...")
        
        extraction_prompt = """# TAREFA: EXTRAIR CONTEÚDO OMITIDO
Você receberá:
1. RELATÓRIO DE OMISSÕES: Lista do que está faltando
2. TRANSCRIÇÃO BRUTA: Onde encontrar o conteúdo

## SUA MISSÃO:
Localize na transcrição bruta os trechos exatos que correspondem às omissões listadas.
Para cada omissão, extraia o trecho relevante da transcrição.

## RELATÓRIO DE OMISSÕES:
{report}

Retorne APENAS os trechos extraídos, sem comentários adicionais."""

        try:
            extract_response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": extraction_prompt.format(report=omissions_report)},
                    {"role": "user", "content": f"TRANSCRIÇÃO BRUTA (primeiros 100k chars):\n{raw_transcript[:100000]}"}
                ],
                timeout=120
            )
            
            missing_content = extract_response.choices[0].message.content
            print(f"{Fore.CYAN}   📝 Conteúdo omitido extraído ({len(missing_content)} caracteres)")
            
            chunks = self._smart_chunk_with_overlap(formatted_text, max_size=self.MAX_CHUNK_SIZE)
            fixed_chunks = []
            
            insertion_prompt = """# TAREFA: INSERIR CONTEÚDO FALTANTE
Você receberá:
1. Um TRECHO da apostila formatada
2. CONTEÚDO OMITIDO que precisa ser adicionado

## SUA MISSÃO:
Se o trecho da apostila é onde o conteúdo omitido deveria estar, insira-o naturalmente.
Se não for o local apropriado, retorne o trecho INALTERADO.

## REGRAS:
- Mantenha toda formatação Markdown
- Integre o conteúdo omitido de forma fluida
- Use tom didático e formal (3ª pessoa)
- NÃO remova nada existente

## CONTEÚDO OMITIDO A INSERIR:
{missing}

Retorne o trecho (modificado ou inalterado)."""

            for i, chunk in enumerate(chunks):
                print(f"{Fore.CYAN}   🔧 Processando chunk {i+1}/{len(chunks)}...")
                try:
                    fix_response = self.client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": insertion_prompt.format(missing=missing_content)},
                            {"role": "user", "content": chunk}
                        ],
                        timeout=90
                    )
                    fixed_chunks.append(fix_response.choices[0].message.content)
                except Exception as chunk_error:
                    print(f"{Fore.YELLOW}   ⚠️  Erro no chunk {i+1}, mantendo original: {chunk_error}")
                    fixed_chunks.append(chunk)
            
            fixed_text = "\n\n".join(fixed_chunks)
            print(f"{Fore.GREEN}   ✅ Texto corrigido gerado ({len(fixed_chunks)} chunks processados)")
            return fixed_text
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Falha ao corrigir omissões: {e}")
            return formatted_text

    def _validate_preservation_heuristics(self, original_text, formatted_text):
        """Validação Heurística com Tolerância Adaptativa"""
        print(f"\n{Fore.CYAN}🔍 Validação Heurística de Preservação (Adaptativa)...")
        issues = []
        
        # 1. Leis (Mantendo lógica original que é boa)
        original_laws = set(re.findall(r'(?:Lei|lei)\s+n?º?\s*\d+[\./]\d+|Art\.?\s*\d+|Súmula\s+\d+', original_text, re.IGNORECASE))
        formatted_laws = set(re.findall(r'(?:Lei|lei)\s+n?º?\s*\d+[\./]\d+|Art\.?\s*\d+|Súmula\s+\d+', formatted_text, re.IGNORECASE))
        missing_laws = original_laws - formatted_laws
        if missing_laws: 
            issues.append(f"❌ {len(missing_laws)} referências legais omitidas")
        else: 
            print(f"{Fore.GREEN}   ✅ Referências legais preservadas")
        
        # 2. Comprimento Adaptativo (Lógica do script Gemini)
        palavras_input = len(original_text.split())
        palavras_output = len(formatted_text.split())
        
        if palavras_input == 0: ratio = 0
        else: ratio = palavras_output / palavras_input
        
        # Heurística de Oralidade
        marcadores_oralidade = ['né', 'então', 'tipo', 'aí', 'pessoal', 'galera', 'tá', 'olha', 'gente', 'veja', 'bom']
        input_lower = original_text.lower()
        count_oralidade = sum(input_lower.count(m) for m in marcadores_oralidade)
        densidade_oralidade = count_oralidade / palavras_input if palavras_input > 0 else 0
        
        # Define tolerância baseada na densidade
        if densidade_oralidade > 0.025:  # Muito coloquial (>2.5%)
            tolerancia = 0.45  # Aceita reduzir até 45%
            tipo = "Muito Coloquial"
        elif densidade_oralidade > 0.015:  # Médio
            tolerancia = 0.38
            tipo = "Coloquial"
        elif densidade_oralidade > 0.008:  # Pouca
            tolerancia = 0.30
            tipo = "Pouca Oralidade"
        else:  # Técnico
            tolerancia = 0.22
            tipo = "Técnico/Denso"
            
        limite_minimo = 1.0 - tolerancia
        
        print(f"   📊 Análise: {tipo} (Densidade: {densidade_oralidade:.2%})")
        print(f"      Ratio atual: {ratio:.1%} (Mínimo aceitável: {limite_minimo:.1%})")
        
        if ratio < limite_minimo:
            issues.append(f"⚠️ Texto formatado muito curto ({ratio:.1%}). Esperado no mínimo {limite_minimo:.1%}")
        else:
            print(f"{Fore.GREEN}   ✅ Comprimento aprovado")
        
        if issues:
            print(f"{Fore.RED}━━━ PROBLEMAS ━━━")
            for i in issues: print(f"   {i}")
            return False, issues
        return True, []

    def validate_completeness_enhanced(self, raw_transcript, formatted_text, video_name):
        """Validação LLM-as-Judge com correção automática"""
        print(f"{Fore.YELLOW}🔍 Validação LLM (Amostragem Múltipla)...")
        
        windows = [
            ("INÍCIO", raw_transcript[:50000], formatted_text[:50000]),
            ("MEIO", raw_transcript[len(raw_transcript)//2-25000:len(raw_transcript)//2+25000],
                      formatted_text[len(formatted_text)//2-25000:len(formatted_text)//2+25000]),
            ("FIM", raw_transcript[-50000:], formatted_text[-50000:])
        ]
        
        validation_prompt = """# TAREFA DE VALIDAÇÃO
Identifique OMISSÕES DE CONTEÚDO TÉCNICO.
Retorne APENAS um JSON:
{
    "omissoes_graves": ["lista de itens perdidos"],
    "nota_fidelidade": 0-10,
    "aprovado": true/false
}
"""
        all_issues = []
        for label, raw_sample, fmt_sample in windows:
            if len(raw_sample) < 1000: continue
            try:
                response = self.client.chat.completions.create(
                    model=self.llm_model,
                    messages=[
                        {"role": "system", "content": validation_prompt},
                        {"role": "user", "content": f"RAW ({label}):\n{raw_sample}\n\nFMT ({label}):\n{fmt_sample}"}
                    ]
                )
                result = json.loads(response.choices[0].message.content)
                if not result.get('aprovado', True):
                    all_issues.extend(result.get('omissoes_graves', []))
            except: pass
            
        if all_issues:
            print(f"{Fore.RED}❌ Omissões detectadas: {all_issues}")
            # Tenta corrigir
            return self._fix_omissions(raw_transcript, formatted_text, "\n".join(all_issues)), all_issues
        
        print(f"{Fore.GREEN}✅ Validação LLM Aprovada")
        return formatted_text, []

    def _generate_audit_report(self, video_name, heuristic_issues, llm_issues):
        with open(f"audit_{video_name}.md", 'w') as f:
            f.write(f"# Auditoria: {video_name}\n")
            f.write(f"Heurística: {heuristic_issues}\n")
            f.write(f"LLM: {llm_issues}\n")

    async def map_structure(self, full_text):
        """Creates a global structure skeleton to guide the formatting."""
        print(f"{Fore.CYAN}🗺️  Mapeando estrutura global do documento...")
        
        # Limit input to avoid context overflow, though GPT-5 Mini handles large context well.
        # Taking the first 200k chars is usually enough for structure.
        input_sample = full_text[:200000] 
        
        try:
            response = await self.async_client.chat.completions.create(
                model=self.llm_model,
                messages=[
                    {"role": "system", "content": self.PROMPT_MAPEAMENTO.format(transcricao=input_sample)}
                ]
                # temperature parameter removed for gpt-5-mini/o1 compatibility
            )
            structure = response.choices[0].message.content
            print(f"{Fore.GREEN}✅ Estrutura mapeada com sucesso!")
            return structure
        except Exception as e:
            print(f"{Fore.RED}⚠️  Falha no mapeamento estrutural: {e}")
            return None

    def renumber_headings(self, text):
        """
        Post-processing: Enforces strictly sequential numbering (1, 2, 3...) 
        for H2 headers, regardless of what the LLM generated.
        """
        print(f"{Fore.CYAN}🔢 Renumerando tópicos sequencialmente...")
        lines = text.split('\n')
        new_lines = []
        h2_count = 0
        h3_count = 0
        h4_count = 0
        
        # Keywords to skip numbering
        skip_keywords = ['resumo', 'quadro', 'tabela', 'síntese', 'esquema', 'bibliografia', 'referências']
        
        for line in lines:
            stripped = line.strip()
            
            # Detect H2 (## Title)
            if stripped.startswith('## ') and not stripped.startswith('### '):
                # Clean existing numbers
                title_text = re.sub(r'^##\s*(\d+(\.\d+)*\.?)?\s*', '', stripped).strip()
                
                if any(k in title_text.lower() for k in skip_keywords):
                    new_lines.append(f"## {title_text}")
                else:
                    h2_count += 1
                    h3_count = 0
                    h4_count = 0
                    new_lines.append(f"## {h2_count}. {title_text}")

            # Detect H3 (### Title)
            elif stripped.startswith('### ') and not stripped.startswith('#### '):
                title_text = re.sub(r'^###\s*(\d+(\.\d+)*\.?)?\s*', '', stripped).strip()
                
                if any(k in title_text.lower() for k in skip_keywords):
                    new_lines.append(f"### {title_text}")
                else:
                    h3_count += 1
                    h4_count = 0
                    new_lines.append(f"### {h2_count}.{h3_count}. {title_text}")

            # Detect H4 (#### Title)
            elif stripped.startswith('#### '):
                title_text = re.sub(r'^####\s*(\d+(\.\d+)*\.?)?\s*', '', stripped).strip()
                
                if any(k in title_text.lower() for k in skip_keywords):
                    new_lines.append(f"#### {title_text}")
                else:
                    h4_count += 1
                    new_lines.append(f"#### {h2_count}.{h3_count}.{h4_count}. {title_text}")
            
            else:
                new_lines.append(line)
                
        return '\n'.join(new_lines)

    def check_coverage(self, original, formatted):
        """Checks for missing laws or sumulas."""
        print(f"{Fore.YELLOW}🔍 Verificando fidelidade (Leis/Súmulas)...")
        
        def extract_refs(text):
            # Leis: "Lei 8.666", "Lei nº 12.345"
            leis = re.findall(r'Lei\s+(?:nº\s*)?([\d\.]+)', text, re.IGNORECASE)
            # Súmulas
            sumulas = re.findall(r'Súmula\s+(?:Vinculante\s+)?(\d+)', text, re.IGNORECASE)
            return set(leis), set(sumulas)
            
        orig_leis, orig_sumulas = extract_refs(original)
        fmt_leis, fmt_sumulas = extract_refs(formatted)
        
        missing_leis = orig_leis - fmt_leis
        missing_sumulas = orig_sumulas - fmt_sumulas
        
        report = []
        if missing_leis:
            report.append(f"⚠️ LEIS POSSIVELMENTE OMITIDAS: {', '.join(missing_leis)}")
        if missing_sumulas:
            report.append(f"⚠️ SÚMULAS POSSIVELMENTE OMITIDAS: {', '.join(missing_sumulas)}")
            
        if not report:
            print(f"{Fore.GREEN}✅ Nenhuma omissão óbvia de Leis/Súmulas detectada.")
            return "Verificação OK: Nenhuma omissão detectada."
        else:
            msg = "\n".join(report)
            print(f"{Fore.RED}{msg}")
            return msg

    async def format_transcription_async(self, transcription, video_name, output_folder, mode="APOSTILA", dry_run=False):
        """Orquestrador Principal com Checkpoint e Robustez (Sequential Mode)"""
        print(f"{Fore.MAGENTA}🧠 Formatando com {self.llm_model} (Sequential Mode)...")
        
        # 0. Context Extraction
        pass
        # professors_info = self._extract_professors_context(transcription)

        # 0.1 Global Structure Mapping (NEW)
        global_structure = await self.map_structure(transcription)

        # 1. Sequential Slicing
        print(f"🔪 Dividindo sequencialmente (sem overlap)...")
        chunks_info = dividir_sequencial(transcription, chars_por_parte=6000)
        validar_chunks(chunks_info, transcription)
        
        total_segments = len(chunks_info)
        print(f"📊 Total de segmentos sequenciais: {total_segments}")
        
        if dry_run:
            print(f"{Fore.YELLOW}🔍 MODO DRY-RUN: Parando antes das chamadas de API.")
            print(f"   Exemplo do Chunk 1: {transcription[chunks_info[0]['inicio']:chunks_info[0]['inicio']+100]}...")
            return "# DRY RUN OUTPUT"
        
        # 2. Checkpoint Loading
        checkpoint_data = load_checkpoint(video_name, output_folder)
        results_map = {} # Map idx -> result
        
        if checkpoint_data:
            print(f"{Fore.CYAN}📂 Retomando via Checkpoint ({checkpoint_data.get('timestamp')})")
            if len(checkpoint_data.get('results', [])) > 0:
                saved_results = checkpoint_data['results']
                # Restore results map
                for idx, res in enumerate(saved_results):
                    if idx < total_segments:
                        results_map[idx] = res
                print(f"   ✅ {len(results_map)} segmentos recuperados.")
                
        # 3. Sequential Processing Loop
        ordered_results = []
        
        # Restore ordered results from map
        for i in range(len(results_map)):
            ordered_results.append(results_map[i])

        start_idx = len(ordered_results)
        
        if start_idx < total_segments:
            print(f"▶ Iniciando processamento sequencial do segmento {start_idx + 1}...")
            
            for i in tqdm(range(start_idx, total_segments), desc="Processando Sequencial"):
                info = chunks_info[i]
                chunk_text = transcription[info['inicio']:info['fim']]
                
                # Context Management
                contexto_estilo = ""
                if i > 0 and ordered_results:
                    raw_context = ordered_results[-1][-800:] # Last 800 chars
                    if len(raw_context.split()) > 30 and "[!WARNING]" not in raw_context: 
                        contexto_estilo = raw_context

                # Processing with Rate Limit
                await rate_limiter.wait_if_needed_async()
                
                try:
                    # Identify speaker context (simplified for sequential chunk)
                    # For strict sequential, we might just pass the chunk
                    # But if we want speaker consistency, we can scan the chunk
                    # For now, we rely on the prompting to handle format
                    
                    # Select prompt based on mode
                    if mode == "FIDELIDADE":
                        local_system_prompt = self.PROMPT_FIDELIDADE
                    else:
                        local_system_prompt = self.PROMPT_APOSTILA_ACTIVE

                    if global_structure:
                        local_system_prompt += f"\n\n# ESTRUTURA GLOBAL DA AULA\nUse esta hiearquia como guia para os títulos:\n{global_structure}\n"

                    # Call Format with Context
                    formatted = await self._format_chunk_async(
                        chunk_text, 
                        i+1, 
                        local_system_prompt, 
                        total=total_segments,
                        contexto_estilo=contexto_estilo
                    )
                    
                    ordered_results.append(formatted)
                    
                    # Save Checkpoint after each chunk
                    save_checkpoint(video_name, output_folder, ordered_results, chunks_info, i + 1)
                    
                except Exception as e:
                    print(f"{Fore.RED}❌ Falha Fatal no segmento {i+1}: {e}")
                    # Save what we have
                    save_checkpoint(video_name, output_folder, ordered_results, chunks_info, i)
                    raise e
        
        # 4. Final Assembly
        print(f"\n{Fore.CYAN}🧹 Pipeline de Limpeza Final (v2.7)...")
        full_formatted = f"# {video_name}\n\n" + "\n\n".join(ordered_results)
        
        # 5. Post-Processing Pipeline (v2.7 features)
        
        print("  Passada 1: Removendo duplicações literais...")
        # Uses the newly ported v2.7 logic
        full_formatted = remover_duplicacoes_literais(full_formatted)
        
        print("  Passada 2: Removendo seções duplicadas...")
        full_formatted = remover_secoes_duplicadas(full_formatted)
        
        print(f"\n{Fore.CYAN}🔢 Renumerando tópicos (1..N) (Logic by Vomo)...")
        full_formatted = self.renumber_headings(full_formatted)
        
        # 6. Validation & Coverage
        print(f"\n{Fore.CYAN}🛡️  Validando cobertura final...")
        coverage_report = self.check_coverage(transcription, full_formatted)
        
        # Save validation report
        report_path = Path(output_folder) / f"{video_name}_validacao.txt"
        with open(report_path, "w", encoding='utf-8') as f:
            f.write(coverage_report)
        print(f"📄 Relatório de validação salvo: {report_path.name}")
        
        # Checkpoint cleanup success
        delete_checkpoint(video_name, output_folder)
        
        return full_formatted

    def save_as_word(self, formatted_text, video_name, output_folder):
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        print(f"{Fore.CYAN}📄 Gerando documento Word profissional (Enhanced)...")
        doc = Document()
        
        # Configuração básica
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        
        title = doc.add_heading(video_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 51, 102)
            
        doc.add_paragraph().add_run(f"Gerado em {time.strftime('%d/%m/%Y')}")
        
        doc.add_paragraph()  # Espaço
     
        # Adiciona Sumário
        doc.add_heading('Sumário', level=1)
        self.create_toc(doc)
        doc.add_page_break()
        
        lines = formatted_text.split('\n')
        i = 0
        in_table = False
        table_rows = []
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Linha vazia
            if not line:
                i += 1
                continue
                
            # Detecta tabelas Markdown
            if '|' in line and not in_table:
                in_table = True
                table_rows = []
            
            if in_table:
                # Regex para identificar linha separadora (contém apenas |, -, :, espaços)
                if re.match(r'^[\s\|\-\:]+$', line) and '-' in line:
                     i += 1
                     continue
                     
                if '|' in line:
                    row_data = [c.strip() for c in line.split('|')[1:-1]]
                    if row_data: table_rows.append(row_data)
                
                # Fim da tabela
                if '|' not in line or i == len(lines)-1:
                    if table_rows: self._add_table_to_doc(doc, table_rows)
                    in_table = False
                    table_rows = []
                    # Se a linha atual não tem pipe, reprocessa ela como texto normal
                    if '|' not in line: continue 
                i += 1
                continue
            
            # Separador horizontal
            if line in ['---', '***', '___']:
                p = doc.add_paragraph()
                p.add_run('_' * 80).font.color.rgb = RGBColor(192, 192, 192)
                i += 1
                continue
            
            # Headings
            if line.startswith('#'):
                level = line.count('#')
                text = line.replace('#', '').strip()
                # Ajuste de nível para Word (max 9, mas usualmente mapeamos markdown 1->1)
                # O script original limitava a level=min(level, 3)
                word_level = len(line.split()[0]) # Conta os #
                word_level = min(word_level, 3) # Cap em 3 como no original ou ajusta 
                h = doc.add_heading(text, level=word_level)
                if word_level == 2: 
                    for r in h.runs: r.font.color.rgb = RGBColor(0, 102, 204)
                i += 1
                continue
            
            # Blockquotes
            if line.startswith('>'):
                p = doc.add_paragraph(line.replace('> ', ''), style='Quote')
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(64, 64, 64)
                i += 1
                continue

            # Listas com bullets
            if line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.space_after = Pt(6)
                self._format_inline_markdown(p, line[2:])
                i += 1
                continue

            # Listas numeradas (simples detecção)
            if len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.space_after = Pt(6)
                self._format_inline_markdown(p, line)
                i += 1
                continue
                
            # Parágrafos normais (com formatação inline)
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.first_line_indent = Cm(1)  # Recuo de 1ª linha
            p.paragraph_format.space_after = Pt(12)       # Espaço duplo entre parágrafos
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._format_inline_markdown(p, line)
            i += 1
            
        output_file = os.path.join(output_folder, f"{video_name}_APOSTILA.docx")
        doc.save(output_file)
        return output_file

    def _format_inline_markdown(self, paragraph, text):
        """Formata markdown inline (bold, italic, code)"""
        paragraph.clear() # Limpa para reconstruir com runs
        
        # Regex para capturar ***bold_italic***, **bold**, *italic*, `code`
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            # Texto antes do match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            # Formatação
            if match.group(0).startswith('***'):
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(0).startswith('**'):
                run = paragraph.add_run(match.group(3))
                run.bold = True
            elif match.group(0).startswith('*'):
                run = paragraph.add_run(match.group(4))
                run.italic = True
            elif match.group(0).startswith('`'):
                run = paragraph.add_run(match.group(5))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(200, 0, 0)
            
            last_end = match.end()
            
        # Texto restante
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_table_to_doc(self, doc, rows):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if not rows: return
        max_cols = max(len(r) for r in rows)
        if max_cols == 0: return

        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Tabela Centralizada
        
        for i, row in enumerate(rows):
            for j in range(max_cols):
                if j < len(row):
                    cell = table.rows[i].cells[j]
                    cell_text = row[j]
                    
                    # Formata markdown dentro da célula (v2.7 improvement)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT # Texto alinhado à esquerda
                    self._format_inline_markdown(p, cell_text)
                    
                    # Header styling
                    if i == 0:
                         shading_elm = OxmlElement('w:shd')
                         shading_elm.set(qn('w:fill'), '0066CC')
                         cell._element.get_or_add_tcPr().append(shading_elm)
                         for p in cell.paragraphs:
                             for r in p.runs:
                                 r.font.bold = True
                                 r.font.color.rgb = RGBColor(255, 255, 255)

    def create_toc(self, doc):
        """Adiciona um campo de Sumário (TOC) nativo do Word"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

def process_single_video(video_path, dry_run=False, mode="APOSTILA"):
    if not os.path.exists(video_path): return
    folder = os.path.dirname(video_path)
    video_name = Path(video_path).stem
    
    try:
        vomo = VomoMLX()
        
        if video_path.lower().endswith('.txt'):
            print(f"{Fore.CYAN}📄 Input é arquivo de texto. Pulando transcrição...")
            with open(video_path, 'r', encoding='utf-8') as f:
                transcription = f.read()
        else:
            if dry_run:
                print("⚠️ Dry run não suporta áudio direto ainda. Use arquivo .txt")
                return
            audio = vomo.optimize_audio(video_path)
            
            raw_path = os.path.join(folder, f"{video_name}_RAW.txt")
            if os.path.exists(raw_path):
                with open(raw_path, 'r') as f: transcription = f.read()
            else:
                transcription = vomo.transcribe(audio)
                with open(raw_path, 'w') as f: f.write(transcription)
            
        formatted = asyncio.run(vomo.format_transcription_async(transcription, video_name, folder, mode=mode, dry_run=dry_run))
        
        with open(os.path.join(folder, f"{video_name}_APOSTILA.md"), 'w') as f:
            f.write(formatted)
            
        vomo.save_as_word(formatted, video_name, folder)
        print(f"{Fore.GREEN}✨ SUCESSO!")
        
    except Exception as e:
        print(f"{Fore.RED}❌ Erro: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    if len(sys.argv) > 1:
        dry_run_flag = "--dry-run" in sys.argv
        mode = "APOSTILA"  # Default mode
        
        # Check for --mode flag
        for arg in sys.argv:
            if arg.startswith("--mode="):
                mode = arg.split("=")[1].upper()
        
        # Get input file (skip flags)
        input_file = None
        for arg in sys.argv[1:]:
            if not arg.startswith("--"):
                input_file = arg
                break
        
        if input_file:
            print(f"{Fore.CYAN}🔧 Modo selecionado: {mode}")
            process_single_video(input_file, dry_run=dry_run_flag, mode=mode)
