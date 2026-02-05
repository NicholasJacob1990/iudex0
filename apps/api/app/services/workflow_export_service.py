"""
Serviço de exportação de resultados de workflow runs.

Suporta exportação para Word (.docx), Excel (.xlsx) e PDF (.pdf).
"""

from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any, Dict, List


class WorkflowExportService:
    """Export workflow run outputs to various formats."""

    @staticmethod
    def _flatten_output(output_data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Flatten output_data into a list of {title, content} sections."""
        sections: List[Dict[str, str]] = []

        for key, value in output_data.items():
            if isinstance(value, dict):
                # Nested dict — flatten keys
                for sub_key, sub_val in value.items():
                    sections.append({
                        "title": f"{key} / {sub_key}",
                        "content": str(sub_val),
                    })
            elif isinstance(value, list):
                content_parts = []
                for i, item in enumerate(value):
                    if isinstance(item, dict):
                        content_parts.append(json.dumps(item, ensure_ascii=False, indent=2))
                    else:
                        content_parts.append(str(item))
                sections.append({
                    "title": str(key),
                    "content": "\n".join(content_parts),
                })
            else:
                sections.append({
                    "title": str(key),
                    "content": str(value),
                })

        return sections

    @staticmethod
    def _format_logs(logs: List[Dict[str, Any]]) -> str:
        """Format run logs into a readable string."""
        lines: List[str] = []
        for entry in logs:
            ts = entry.get("timestamp", "")
            msg = entry.get("message", entry.get("data", ""))
            if isinstance(msg, dict):
                msg = json.dumps(msg, ensure_ascii=False)
            lines.append(f"[{ts}] {msg}")
        return "\n".join(lines) if lines else "Sem logs disponíveis."

    async def export_to_docx(self, run_data: Dict[str, Any]) -> io.BytesIO:
        """Export run output as Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError(
                "python-docx não está instalado. Execute: pip install python-docx"
            )

        doc = Document()

        # Title
        workflow_name = run_data.get("workflow_name", "Workflow")
        title = doc.add_heading(workflow_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Status: {run_data.get('status', 'N/A')}")
        doc.add_paragraph(f"Exportado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")

        # Input data
        input_data = run_data.get("input_data", {})
        if input_data:
            doc.add_heading("Dados de Entrada", level=1)
            for key, value in input_data.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{key}: ")
                run.bold = True
                p.add_run(str(value))

        # Output data — main content
        output_data = run_data.get("output_data", {})
        if output_data:
            doc.add_heading("Resultado", level=1)
            sections = self._flatten_output(output_data)
            for section in sections:
                doc.add_heading(section["title"], level=2)
                # Split content into paragraphs to preserve formatting
                for line in section["content"].split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)

        # Logs
        logs = run_data.get("logs", [])
        if logs:
            doc.add_heading("Logs de Execução", level=1)
            log_text = self._format_logs(logs)
            p = doc.add_paragraph(log_text)
            for run in p.runs:
                run.font.size = Pt(8)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    async def export_to_xlsx(self, run_data: Dict[str, Any]) -> io.BytesIO:
        """Export run output as Excel spreadsheet."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill
        except ImportError:
            raise RuntimeError(
                "openpyxl não está instalado. Execute: pip install openpyxl"
            )

        wb = Workbook()
        header_font = Font(bold=True, size=12)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, size=11, color="FFFFFF")

        # --- Sheet 1: Resumo ---
        ws_summary = wb.active
        ws_summary.title = "Resumo"

        workflow_name = run_data.get("workflow_name", "Workflow")
        ws_summary.append(["Workflow", workflow_name])
        ws_summary.append(["Status", run_data.get("status", "N/A")])
        ws_summary.append(["Exportado em", datetime.now().strftime("%d/%m/%Y %H:%M")])
        ws_summary.append([])

        # Style header cells
        for row in ws_summary.iter_rows(min_row=1, max_row=3, min_col=1, max_col=1):
            for cell in row:
                cell.font = Font(bold=True)

        ws_summary.column_dimensions["A"].width = 20
        ws_summary.column_dimensions["B"].width = 60

        # Input data
        input_data = run_data.get("input_data", {})
        if input_data:
            ws_summary.append(["Dados de Entrada", ""])
            for key, value in input_data.items():
                ws_summary.append([str(key), str(value)])
            ws_summary.append([])

        # --- Sheet 2: Resultado ---
        output_data = run_data.get("output_data", {})
        if output_data:
            ws_result = wb.create_sheet("Resultado")
            ws_result.append(["Seção", "Conteúdo"])

            # Style headers
            for cell in ws_result[1]:
                cell.font = header_font_white
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")

            sections = self._flatten_output(output_data)
            for section in sections:
                ws_result.append([section["title"], section["content"]])

            ws_result.column_dimensions["A"].width = 30
            ws_result.column_dimensions["B"].width = 80

        # --- Sheet 3: Logs ---
        logs = run_data.get("logs", [])
        if logs:
            ws_logs = wb.create_sheet("Logs")
            ws_logs.append(["Timestamp", "Mensagem"])

            for cell in ws_logs[1]:
                cell.font = header_font_white
                cell.fill = header_fill

            for entry in logs:
                ts = str(entry.get("timestamp", ""))
                msg = entry.get("message", entry.get("data", ""))
                if isinstance(msg, dict):
                    msg = json.dumps(msg, ensure_ascii=False)
                ws_logs.append([ts, str(msg)])

            ws_logs.column_dimensions["A"].width = 25
            ws_logs.column_dimensions["B"].width = 80

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer

    async def export_to_pdf(self, run_data: Dict[str, Any]) -> io.BytesIO:
        """Export run output as PDF."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib.colors import HexColor
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                Table,
                TableStyle,
            )
        except ImportError:
            raise RuntimeError(
                "reportlab não está instalado. Execute: pip install reportlab"
            )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            "SectionTitle",
            parent=styles["Heading2"],
            spaceAfter=6,
            textColor=HexColor("#1e3a5f"),
        ))
        styles.add(ParagraphStyle(
            "ContentBody",
            parent=styles["BodyText"],
            spaceAfter=4,
            leading=14,
        ))

        elements: list = []

        # Title
        workflow_name = run_data.get("workflow_name", "Workflow")
        elements.append(Paragraph(workflow_name, styles["Title"]))
        elements.append(Spacer(1, 0.5 * cm))

        # Metadata
        elements.append(Paragraph(
            f"Status: {run_data.get('status', 'N/A')} | "
            f"Exportado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            styles["Normal"],
        ))
        elements.append(Spacer(1, 0.5 * cm))

        # Input data
        input_data = run_data.get("input_data", {})
        if input_data:
            elements.append(Paragraph("Dados de Entrada", styles["Heading1"]))
            elements.append(Spacer(1, 0.2 * cm))
            for key, value in input_data.items():
                text = f"<b>{_escape_xml(str(key))}:</b> {_escape_xml(str(value))}"
                elements.append(Paragraph(text, styles["ContentBody"]))
            elements.append(Spacer(1, 0.3 * cm))

        # Output data
        output_data = run_data.get("output_data", {})
        if output_data:
            elements.append(Paragraph("Resultado", styles["Heading1"]))
            elements.append(Spacer(1, 0.2 * cm))
            sections = self._flatten_output(output_data)
            for section in sections:
                elements.append(Paragraph(
                    _escape_xml(section["title"]),
                    styles["SectionTitle"],
                ))
                for line in section["content"].split("\n"):
                    if line.strip():
                        elements.append(Paragraph(
                            _escape_xml(line),
                            styles["ContentBody"],
                        ))
                elements.append(Spacer(1, 0.2 * cm))

        # Logs
        logs = run_data.get("logs", [])
        if logs:
            elements.append(Paragraph("Logs de Execução", styles["Heading1"]))
            elements.append(Spacer(1, 0.2 * cm))
            log_text = self._format_logs(logs)
            for line in log_text.split("\n")[:50]:  # Limit to 50 log lines
                elements.append(Paragraph(
                    _escape_xml(line),
                    styles["Code"] if "Code" in styles.byName else styles["Normal"],
                ))

        if not elements:
            elements.append(Paragraph("Nenhum dado para exportar.", styles["Normal"]))

        doc.build(elements)
        buffer.seek(0)
        return buffer


def _escape_xml(text: str) -> str:
    """Escape special XML characters for ReportLab paragraphs."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
