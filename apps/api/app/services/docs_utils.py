"""
Shared Document Utilities
Standardized functions for generating formatted legal documents (.docx)
"""

import os
import logging
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("DocsUtils")

def process_inline_formatting(paragraph, text: str):
    """Processa **negrito** e *itálico* no texto"""
    import re
    
    # Pattern para **negrito** e *itálico*
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # Adiciona texto antes do match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Verifica se é negrito ou itálico
        if match.group(2):  # **negrito**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *itálico*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        
        last_end = match.end()
    
    # Adiciona texto restante
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

def save_as_word_juridico(formatted_text: str, filename: str, output_folder: str, modo: str = "GENERICO") -> str:
    """Salva o documento com formatação jurídica (ABNT/Forense)"""
    logger.info(f"📄 Gerando DOCX Jurídico: {filename}...")
    
    os.makedirs(output_folder, exist_ok=True)
    
    doc = Document()
    
    # Configuração de Página (A4, Margens Padrão Peticionamento)
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    
    # Estilos Básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Cabeçalho Dinâmico baseado no modo
    cabecalhos = {
        "PETICAO_INICIAL": "EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA ___ VARA _____ DA COMARCA DE ______",
        "CONTESTACAO": "EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA ___ VARA _____ DA COMARCA DE ______",
        "RECURSO": "EGRÉGIO TRIBUNAL DE JUSTIÇA DO ESTADO DE ______",
        "PARECER": "PARECER JURÍDICO",
        "SENTENCA": "SENTENÇA",
        "CONTRATO": "INSTRUMENTO PARTICULAR DE CONTRATO",
        "AUDITORIA": "RELATÓRIO DE AUDITORIA JURÍDICA" 
    }
    
    header_text = cabecalhos.get(modo.upper(), "DOCUMENTO JURÍDICO")
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(header_text)
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Processamento do Markdown
    lines = formatted_text.split('\n')
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph()
            continue
            
        # Títulos (Outline)
        if line_stripped.startswith('#'):
            level = len(line_stripped) - len(line_stripped.lstrip('#'))
            text = line_stripped.lstrip('#').strip()
            heading_level = min(level, 3)  # Word suporta até Heading 9, mas usamos até 3
            h = doc.add_heading(text, level=heading_level)
            # Ajuste de cor para preto (petições não usam azul)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)
                r.font.name = 'Arial'
            continue
            
        # Citações (Blockquotes) - Jurisprudência
        if line_stripped.startswith('>'):
            quote_text = line_stripped.lstrip('>').strip()
            p = doc.add_paragraph(quote_text)
            p.paragraph_format.left_indent = Cm(4)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.italic = True
            continue
        
        # Tabelas Markdown (simplificado)
        if line_stripped.startswith('|') and '|' in line_stripped[1:]:
            # Skip table processing for now, treat as text
            p = doc.add_paragraph(line_stripped)
            p.paragraph_format.first_line_indent = Cm(0)
            continue
            
        # Parágrafo Normal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(2.0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        
        # Processar formatação inline (negrito/itálico)
        process_inline_formatting(p, line_stripped)

    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    logger.info(f"✅ Arquivo DOCX salvo: {output_path}")
    return output_path
