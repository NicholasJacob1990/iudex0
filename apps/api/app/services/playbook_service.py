"""
PlaybookService — Serviço de análise de contratos com base em Playbooks.

Inspirado no Harvey AI Playbook: aplica regras estruturadas a contratos,
classificando cláusulas como conformes, não-conformes ou ausentes,
e gerando sugestões de redline quando necessário.

Integração com a página /minuta:
    - Usuário abre um contrato em /minuta
    - Seleciona um Playbook no dropdown
    - PlaybookService.get_playbook_for_prompt() injeta as regras no system prompt do agente
    - O agente usa as regras para sinalizar problemas durante a revisão interativa
    - Para análise completa (batch), usa-se analyze_contract_with_playbook()
"""

from __future__ import annotations

import asyncio
import json
import logging
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document import Document
from app.models.playbook import Playbook, PlaybookRule
from app.schemas.playbook_analysis import (
    ClauseAnalysisResult,
    ClauseClassification,
    AnalysisSeverity,
    PlaybookAnalysisResult,
)
from app.services.playbook_prompts import (
    CLAUSE_EXTRACTION_PROMPT,
    CLAUSE_ANALYSIS_PROMPT,
    REDLINE_GENERATION_PROMPT,
    PLAYBOOK_GENERATION_PROMPT,
    PLAYBOOK_IMPORT_PROMPT,
    PLAYBOOK_SUMMARY_PROMPT,
    PLAYBOOK_FOR_AGENT_PROMPT,
    PLAYBOOK_RULE_TEMPLATE,
    WINNING_LANGUAGE_EXTRACTION_PROMPT,
)
from app.services.ai.agent_clients import (
    get_gemini_client,
    get_claude_client,
    call_vertex_gemini_async,
    call_anthropic_async,
)

logger = logging.getLogger("PlaybookService")

# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------

# Modelo padrão para análise (Gemini Flash — rápido e barato)
DEFAULT_ANALYSIS_MODEL = "gemini-3-flash"
# Modelo para geração de playbook e redlines (mais capaz)
DEFAULT_GENERATION_MODEL = "gemini-3-pro"
# Máximo de cláusulas para analisar em paralelo
MAX_CONCURRENT_ANALYSES = 5
# Timeout para chamadas de AI (segundos)
AI_TIMEOUT = 120

# Pesos de severidade para cálculo do risk_score
SEVERITY_WEIGHTS: Dict[str, float] = {
    "low": 1.0,
    "medium": 2.5,
    "high": 5.0,
    "critical": 10.0,
}

# Pesos de classificação para cálculo do risk_score
CLASSIFICATION_WEIGHTS: Dict[str, float] = {
    "compliant": 0.0,
    "needs_review": 0.4,
    "non_compliant": 1.0,
    "not_found": 0.6,  # cláusula ausente é risco moderado
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _safe_json_parse(text: str) -> Any:
    """Tenta fazer parse de JSON mesmo com markdown code fences."""
    if not text:
        return None
    # Remove markdown fences
    cleaned = text.strip()
    if cleaned.startswith("```"):
        # Remove primeira linha (```json) e última (```)
        lines = cleaned.split("\n")
        lines = [l for l in lines if not l.strip().startswith("```")]
        cleaned = "\n".join(lines)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        logger.warning("Falha ao parsear JSON da resposta da IA: %s...", cleaned[:200])
        return None


def _calculate_risk_score(clauses: List[ClauseAnalysisResult]) -> float:
    """
    Calcula score de risco (0-100) com base nas classificações e severidades.

    Fórmula: soma ponderada (severidade x classificação) / máximo teórico x 100
    """
    if not clauses:
        return 0.0

    total_weight = 0.0
    max_weight = 0.0

    for clause in clauses:
        sev_w = SEVERITY_WEIGHTS.get(clause.severity.value, 2.5)
        cls_w = CLASSIFICATION_WEIGHTS.get(clause.classification.value, 0.5)
        total_weight += sev_w * cls_w
        max_weight += sev_w * 1.0  # máximo é non_compliant (1.0)

    if max_weight == 0:
        return 0.0

    return round(min((total_weight / max_weight) * 100, 100.0), 1)


async def _call_ai(
    prompt: str,
    system_instruction: Optional[str] = None,
    model: str = DEFAULT_ANALYSIS_MODEL,
    max_tokens: int = 4000,
    temperature: float = 0.2,
) -> Optional[str]:
    """
    Chama o modelo de IA com fallback.
    Prioridade: Gemini -> Claude.
    """
    # Tenta Gemini primeiro (mais barato e rápido)
    gemini_client = get_gemini_client()
    if gemini_client:
        try:
            result = await call_vertex_gemini_async(
                client=gemini_client,
                prompt=prompt,
                model=model,
                max_tokens=max_tokens,
                temperature=temperature,
                timeout=AI_TIMEOUT,
                system_instruction=system_instruction,
            )
            if result:
                return result
        except Exception as e:
            logger.warning("Gemini falhou para playbook, tentando Claude: %s", e)

    # Fallback para Claude
    claude_client = get_claude_client()
    if claude_client:
        try:
            result = await call_anthropic_async(
                client=claude_client,
                prompt=prompt,
                model="claude-4.5-sonnet",
                max_tokens=max_tokens,
                temperature=temperature,
                timeout=AI_TIMEOUT,
                system_instruction=system_instruction,
            )
            if result:
                return result
        except Exception as e:
            logger.error("Claude também falhou para playbook: %s", e)

    return None


# ---------------------------------------------------------------------------
# PlaybookService
# ---------------------------------------------------------------------------


class PlaybookService:
    """Serviço para aplicar regras de Playbook na análise de contratos."""

    # -----------------------------------------------------------------------
    # analyze_contract_with_playbook
    # -----------------------------------------------------------------------

    async def analyze_contract_with_playbook(
        self,
        document_id: str,
        playbook_id: str,
        user_id: str,
        db: AsyncSession,
        contract_text_override: Optional[str] = None,
    ) -> PlaybookAnalysisResult:
        """
        Analisa um contrato contra as regras de um playbook.

        Fluxo:
        1. Carrega o playbook e suas regras
        2. Carrega/extrai o texto do contrato (do documento)
        3. Extrai cláusulas do contrato via IA
        4. Para cada regra, analisa a cláusula correspondente
        5. Gera redline para cláusulas não-conformes
        6. Calcula risk score e gera resumo executivo
        7. Retorna resultado estruturado

        Args:
            document_id: ID do documento de contrato
            playbook_id: ID do playbook a aplicar
            user_id: ID do usuário solicitante
            db: Sessão assíncrona do banco
            contract_text_override: Texto do contrato (opcional, sobrescreve o extraído)

        Returns:
            PlaybookAnalysisResult com análise completa

        Raises:
            ValueError: Se playbook, documento ou texto não encontrados
        """
        logger.info(
            "Iniciando análise de playbook=%s no documento=%s pelo user=%s",
            playbook_id, document_id, user_id,
        )

        # 1. Carregar playbook com regras
        playbook = await self._load_playbook(playbook_id, db)
        rules = await self._load_playbook_rules(playbook_id, db)

        if not rules:
            raise ValueError(f"Playbook {playbook_id} não possui regras ativas")

        # 2. Carregar texto do contrato
        contract_text = contract_text_override
        if not contract_text:
            contract_text = await self._load_document_text(document_id, user_id, db)

        if not contract_text or not contract_text.strip():
            raise ValueError(
                f"Documento {document_id} não possui texto extraído. "
                "Certifique-se de que o documento foi processado."
            )

        # 3. Extrair cláusulas do contrato
        extracted_clauses = await self._extract_clauses(contract_text)
        logger.info("Extraídas %d cláusulas do contrato", len(extracted_clauses))

        # 4. Analisar cada regra contra o contrato
        clause_results = await self._analyze_all_rules(
            rules=rules,
            extracted_clauses=extracted_clauses,
            contract_text=contract_text,
            party_perspective=getattr(playbook, "party_perspective", "neutro") or "neutro",
        )

        # 5. Gerar redlines para cláusulas não-conformes
        clause_results = await self._generate_redlines_for_non_compliant(
            clause_results=clause_results,
            rules_by_id={r.id: r for r in rules},
        )

        # 6. Calcular métricas
        compliant_count = sum(
            1 for c in clause_results
            if c.classification == ClauseClassification.COMPLIANT
        )
        needs_review_count = sum(
            1 for c in clause_results
            if c.classification == ClauseClassification.NEEDS_REVIEW
        )
        non_compliant_count = sum(
            1 for c in clause_results
            if c.classification == ClauseClassification.NON_COMPLIANT
        )
        not_found_count = sum(
            1 for c in clause_results
            if c.classification == ClauseClassification.NOT_FOUND
        )
        risk_score = _calculate_risk_score(clause_results)

        # 7. Gerar resumo executivo
        summary = await self._generate_summary(
            playbook_name=playbook.name,
            total_rules=len(rules),
            compliant=compliant_count,
            needs_review=needs_review_count,
            non_compliant=non_compliant_count,
            not_found=not_found_count,
            risk_score=risk_score,
            clause_results=clause_results,
        )

        result = PlaybookAnalysisResult(
            playbook_id=playbook_id,
            playbook_name=playbook.name,
            document_id=document_id,
            total_rules=len(rules),
            compliant=compliant_count,
            needs_review=needs_review_count,
            non_compliant=non_compliant_count,
            not_found=not_found_count,
            risk_score=risk_score,
            clauses=clause_results,
            summary=summary,
            analyzed_at=datetime.now(timezone.utc),
        )

        logger.info(
            "Análise concluída: risk_score=%.1f, compliant=%d, needs_review=%d, "
            "non_compliant=%d, not_found=%d",
            risk_score, compliant_count, needs_review_count,
            non_compliant_count, not_found_count,
        )

        return result

    # -----------------------------------------------------------------------
    # analyze_clause
    # -----------------------------------------------------------------------

    async def analyze_clause(
        self,
        clause_text: str,
        rule: PlaybookRule,
        contract_context: str,
        party_perspective: str = "neutro",
    ) -> ClauseAnalysisResult:
        """
        Analisa uma única cláusula contra uma regra do playbook.

        Args:
            clause_text: Texto da cláusula encontrada no contrato
            rule: Regra do playbook para comparar
            contract_context: Contexto adicional do contrato (primeiros 2000 chars)

        Returns:
            ClauseAnalysisResult com classificação e explicação
        """
        # Build party perspective section for the prompt
        perspective_labels = {
            "contratante": "Você está analisando na perspectiva do CONTRATANTE (parte que contrata/compra). Priorize proteções para o contratante.",
            "contratado": "Você está analisando na perspectiva do CONTRATADO (parte que presta o serviço/vende). Priorize proteções para o contratado.",
            "neutro": "Analise de forma neutra, considerando equilibrio entre as partes.",
        }
        perspective_text = perspective_labels.get(party_perspective, perspective_labels["neutro"])

        prompt = CLAUSE_ANALYSIS_PROMPT.format(
            rule_name=rule.rule_name,
            clause_type=rule.clause_type,
            rule_description=rule.description or "Sem descrição",
            preferred_position=rule.preferred_position,
            fallback_positions=json.dumps(rule.fallback_positions or [], ensure_ascii=False),
            rejected_positions=json.dumps(rule.rejected_positions or [], ensure_ascii=False),
            severity=rule.severity,
            guidance_notes=rule.guidance_notes or "Sem notas adicionais",
            clause_location="Não especificada",
            clause_text=clause_text,
            contract_context=contract_context[:2000],
            party_perspective_section=perspective_text,
        )

        response = await _call_ai(
            prompt=prompt,
            system_instruction="Você é um advogado revisor de contratos. Responda sempre em JSON válido.",
            model=DEFAULT_ANALYSIS_MODEL,
            temperature=0.1,
        )

        parsed = _safe_json_parse(response) if response else None

        if parsed and isinstance(parsed, dict):
            classification_str = parsed.get("classification", "needs_review")
            try:
                classification = ClauseClassification(classification_str)
            except ValueError:
                classification = ClauseClassification.NEEDS_REVIEW

            return ClauseAnalysisResult(
                rule_id=rule.id,
                rule_name=rule.rule_name,
                clause_type=rule.clause_type,
                found_in_contract=True,
                original_text=clause_text,
                location=None,
                classification=classification,
                severity=AnalysisSeverity(rule.severity),
                explanation=parsed.get("explanation", "Análise não disponível"),
                suggested_redline=None,  # Preenchido depois se non_compliant
                comment=parsed.get("comment"),
                confidence=min(max(float(parsed.get("confidence", 0.5)), 0.0), 1.0),
            )

        # Fallback: se a IA não retornou JSON válido
        logger.warning("IA não retornou JSON válido para regra %s", rule.rule_name)
        return ClauseAnalysisResult(
            rule_id=rule.id,
            rule_name=rule.rule_name,
            clause_type=rule.clause_type,
            found_in_contract=True,
            original_text=clause_text,
            location=None,
            classification=ClauseClassification.NEEDS_REVIEW,
            severity=AnalysisSeverity(rule.severity),
            explanation="Análise automática inconclusiva — revisão manual recomendada.",
            suggested_redline=None,
            comment=None,
            confidence=0.3,
        )

    # -----------------------------------------------------------------------
    # generate_redline
    # -----------------------------------------------------------------------

    async def generate_redline(
        self,
        original_clause: str,
        rule: PlaybookRule,
        analysis: ClauseAnalysisResult,
    ) -> str:
        """
        Gera texto sugerido de substituição para cláusula não-conforme.

        Args:
            original_clause: Texto original da cláusula
            rule: Regra do playbook
            analysis: Resultado da análise da cláusula

        Returns:
            Texto sugerido para redline
        """
        prompt = REDLINE_GENERATION_PROMPT.format(
            original_clause=original_clause,
            clause_type=rule.clause_type,
            preferred_position=rule.preferred_position,
            fallback_positions=json.dumps(rule.fallback_positions or [], ensure_ascii=False),
            severity=rule.severity,
            guidance_notes=rule.guidance_notes or "Sem notas adicionais",
            explanation=analysis.explanation,
        )

        response = await _call_ai(
            prompt=prompt,
            system_instruction="Você é um redator de contratos. Responda em JSON válido.",
            model=DEFAULT_GENERATION_MODEL,
            max_tokens=4000,
            temperature=0.3,
        )

        parsed = _safe_json_parse(response) if response else None

        if parsed and isinstance(parsed, dict):
            return parsed.get("suggested_text", original_clause)

        return original_clause  # fallback: retorna texto original

    # -----------------------------------------------------------------------
    # generate_playbook_from_contracts
    # -----------------------------------------------------------------------

    async def generate_playbook_from_contracts(
        self,
        document_ids: List[str],
        name: str,
        area: str,
        user_id: str,
        db: AsyncSession,
        description: Optional[str] = None,
    ) -> Playbook:
        """
        Gera um Playbook automaticamente a partir de 1-10 contratos.

        Inspirado na feature do Harvey AI: o usuário faz upload de contratos
        e a IA extrai padrões, tipos de cláusula e posições para criar
        um playbook reutilizável.

        Fluxo:
        1. Carrega os textos de todos os contratos
        2. Usa IA para identificar tipos de cláusula em todos os contratos
        3. Para cada tipo de cláusula, determina:
           - Posição mais comum -> preferida
           - Variações -> alternativas
           - Outliers -> potenciais rejeições
        4. Gera PlaybookRules
        5. Cria e persiste o Playbook no banco

        Args:
            document_ids: Lista de IDs de documentos (1-10)
            name: Nome do playbook
            area: Área jurídica
            user_id: ID do usuário criador
            db: Sessão do banco
            description: Descrição opcional

        Returns:
            Playbook criado com regras

        Raises:
            ValueError: Se documentos não encontrados ou textos vazios
        """
        if len(document_ids) > 10:
            raise ValueError("Máximo de 10 contratos para geração de playbook")

        logger.info(
            "Gerando playbook '%s' (área=%s) a partir de %d contratos",
            name, area, len(document_ids),
        )

        # 1. Carregar textos dos contratos
        contracts_text_parts: List[str] = []
        for idx, doc_id in enumerate(document_ids, 1):
            text = await self._load_document_text(doc_id, user_id, db)
            if text and text.strip():
                # Limitar cada contrato a ~15k chars para caber no prompt
                truncated = text[:15000]
                contracts_text_parts.append(
                    f"### CONTRATO {idx} (ID: {doc_id})\n{truncated}\n"
                )
            else:
                logger.warning("Documento %s sem texto extraído, ignorando", doc_id)

        if not contracts_text_parts:
            raise ValueError(
                "Nenhum dos documentos fornecidos possui texto extraído. "
                "Certifique-se de que os documentos foram processados."
            )

        contracts_text = "\n---\n".join(contracts_text_parts)

        # 2. Chamar IA para gerar regras
        prompt = PLAYBOOK_GENERATION_PROMPT.format(
            area=area,
            contracts_text=contracts_text,
        )

        response = await _call_ai(
            prompt=prompt,
            system_instruction=(
                "Você é um advogado sênior especialista em revisão contratual brasileira. "
                "Responda APENAS em JSON válido."
            ),
            model=DEFAULT_GENERATION_MODEL,
            max_tokens=8000,
            temperature=0.3,
        )

        rules_data = _safe_json_parse(response) if response else None

        if not rules_data or not isinstance(rules_data, list):
            raise ValueError(
                "Não foi possível gerar regras a partir dos contratos fornecidos. "
                "Tente novamente ou forneça contratos com mais conteúdo."
            )

        # 3. Criar Playbook no banco
        playbook_id = str(uuid.uuid4())
        playbook = Playbook(
            id=playbook_id,
            user_id=user_id,
            name=name,
            description=description or f"Playbook gerado automaticamente para área de {area}",
            area=area,
            rules=[],  # JSON inline (populado abaixo)
            is_active=True,
            is_template=False,
            scope="personal",
            version=1,
        )
        db.add(playbook)

        # 4. Criar PlaybookRules
        rules_json_inline: List[Dict[str, Any]] = []

        for idx, rule_data in enumerate(rules_data):
            if not isinstance(rule_data, dict):
                continue

            rule = PlaybookRule(
                id=str(uuid.uuid4()),
                playbook_id=playbook_id,
                clause_type=rule_data.get("clause_type", "geral"),
                rule_name=rule_data.get("rule_name", f"Regra {idx + 1}"),
                description=rule_data.get("description"),
                preferred_position=rule_data.get("preferred_position", ""),
                fallback_positions=rule_data.get("fallback_positions", []),
                rejected_positions=rule_data.get("rejected_positions", []),
                action_on_reject=rule_data.get("action_on_reject", "flag"),
                severity=rule_data.get("severity", "medium"),
                guidance_notes=rule_data.get("guidance_notes"),
                order=idx,
                is_active=True,
            )
            db.add(rule)
            rules_json_inline.append(rule_data)

        # Atualizar JSON inline no playbook
        playbook.rules = rules_json_inline

        await db.commit()
        await db.refresh(playbook)

        logger.info(
            "Playbook '%s' criado com %d regras (id=%s)",
            name, len(rules_json_inline), playbook_id,
        )

        return playbook

    # -----------------------------------------------------------------------
    # import_playbook_from_document
    # -----------------------------------------------------------------------

    async def import_playbook_from_document(
        self,
        document_id: str,
        name: str,
        area: str,
        user_id: str,
        db: AsyncSession,
        description: Optional[str] = None,
    ) -> Playbook:
        """
        Importa um playbook extraindo regras de um documento existente (PDF/DOCX).

        Fluxo:
        1. Carrega o texto do documento
        2. Chama IA com PLAYBOOK_IMPORT_PROMPT para extrair regras
        3. Parseia as regras extraidas
        4. Cria Playbook + PlaybookRules no banco
        5. Retorna o playbook criado

        Args:
            document_id: ID do documento fonte
            name: Nome do playbook
            area: Area juridica
            user_id: ID do usuario criador
            db: Sessao do banco
            description: Descricao opcional

        Returns:
            Playbook criado com regras extraidas

        Raises:
            ValueError: Se documento nao encontrado ou sem texto
        """
        logger.info(
            "Importando playbook '%s' (area=%s) do documento=%s pelo user=%s",
            name, area, document_id, user_id,
        )

        # 1. Carregar texto do documento
        document_text = await self._load_document_text(document_id, user_id, db)

        if not document_text or not document_text.strip():
            raise ValueError(
                f"Documento {document_id} nao possui texto extraido. "
                "Certifique-se de que o documento foi processado."
            )

        # Limitar a ~30k chars para caber no prompt
        truncated_text = document_text[:30000]

        # 2. Chamar IA para extrair regras
        prompt = PLAYBOOK_IMPORT_PROMPT.format(document_text=truncated_text)

        response = await _call_ai(
            prompt=prompt,
            system_instruction=(
                "Voce e um advogado senior especialista em revisao contratual brasileira. "
                "Extraia regras de playbook do documento. Responda APENAS em JSON valido."
            ),
            model=DEFAULT_GENERATION_MODEL,
            max_tokens=8000,
            temperature=0.3,
        )

        rules_data = _safe_json_parse(response) if response else None

        if not rules_data or not isinstance(rules_data, list):
            raise ValueError(
                "Nao foi possivel extrair regras do documento fornecido. "
                "Verifique se o documento contem um playbook ou manual de revisao."
            )

        # 3. Criar Playbook no banco
        playbook_id = str(uuid.uuid4())
        playbook = Playbook(
            id=playbook_id,
            user_id=user_id,
            name=name,
            description=description or f"Playbook importado de documento para area de {area}",
            area=area,
            rules=[],
            is_active=True,
            is_template=False,
            scope="personal",
            version=1,
        )
        db.add(playbook)

        # 4. Criar PlaybookRules
        rules_json_inline: List[Dict[str, Any]] = []

        for idx, rule_data in enumerate(rules_data):
            if not isinstance(rule_data, dict):
                continue

            rule = PlaybookRule(
                id=str(uuid.uuid4()),
                playbook_id=playbook_id,
                clause_type=rule_data.get("clause_type", "geral"),
                rule_name=rule_data.get("rule_name", f"Regra {idx + 1}"),
                description=rule_data.get("description"),
                preferred_position=rule_data.get("preferred_position", ""),
                fallback_positions=rule_data.get("fallback_positions", []),
                rejected_positions=rule_data.get("rejected_positions", []),
                action_on_reject=rule_data.get("action_on_reject", "flag"),
                severity=rule_data.get("severity", "medium"),
                guidance_notes=rule_data.get("guidance_notes"),
                order=idx,
                is_active=True,
            )
            db.add(rule)
            rules_json_inline.append(rule_data)

        # Atualizar JSON inline no playbook
        playbook.rules = rules_json_inline

        await db.commit()
        await db.refresh(playbook)

        logger.info(
            "Playbook '%s' importado com %d regras (id=%s)",
            name, len(rules_json_inline), playbook_id,
        )

        return playbook

    # -----------------------------------------------------------------------
    # extract_rules_from_text (for direct file upload — preview only)
    # -----------------------------------------------------------------------

    async def extract_rules_from_text(
        self,
        document_text: str,
        area: str,
    ) -> List[Dict[str, Any]]:
        """
        Extrai regras de playbook de um texto bruto (sem persistir no banco).

        Usado pelo endpoint de importacao direta de arquivo (PDF/DOCX upload).
        Retorna a lista de regras extraidas para preview no frontend,
        antes da confirmacao/criacao do playbook.

        Args:
            document_text: Texto extraido do documento
            area: Area juridica para contexto

        Returns:
            Lista de dicts com as regras extraidas

        Raises:
            ValueError: Se texto vazio ou IA nao conseguir extrair regras
        """
        if not document_text or not document_text.strip():
            raise ValueError("Texto do documento esta vazio.")

        # Limitar a ~30k chars para caber no prompt
        truncated_text = document_text[:30000]

        prompt = PLAYBOOK_IMPORT_PROMPT.format(document_text=truncated_text)

        response = await _call_ai(
            prompt=prompt,
            system_instruction=(
                "Voce e um advogado senior especialista em revisao contratual brasileira. "
                "Extraia regras de playbook do documento. Responda APENAS em JSON valido."
            ),
            model=DEFAULT_GENERATION_MODEL,
            max_tokens=8000,
            temperature=0.3,
        )

        rules_data = _safe_json_parse(response) if response else None

        if not rules_data or not isinstance(rules_data, list):
            raise ValueError(
                "Nao foi possivel extrair regras do documento fornecido. "
                "Verifique se o documento contem um playbook ou manual de revisao."
            )

        # Normalizar e validar cada regra
        valid_rules: List[Dict[str, Any]] = []
        for idx, rule_data in enumerate(rules_data):
            if not isinstance(rule_data, dict):
                continue

            valid_rules.append({
                "clause_type": rule_data.get("clause_type", "geral"),
                "rule_name": rule_data.get("rule_name", f"Regra {idx + 1}"),
                "description": rule_data.get("description"),
                "preferred_position": rule_data.get("preferred_position", ""),
                "fallback_positions": rule_data.get("fallback_positions", []),
                "rejected_positions": rule_data.get("rejected_positions", []),
                "action_on_reject": rule_data.get("action_on_reject", "flag"),
                "severity": rule_data.get("severity", "medium"),
                "guidance_notes": rule_data.get("guidance_notes"),
                "order": idx,
            })

        logger.info("Extraidas %d regras do texto enviado", len(valid_rules))
        return valid_rules

    # -----------------------------------------------------------------------
    # create_playbook_with_rules (batch create from pre-extracted rules)
    # -----------------------------------------------------------------------

    async def create_playbook_with_rules(
        self,
        name: str,
        area: str,
        user_id: str,
        rules_data: List[Dict[str, Any]],
        db: AsyncSession,
        description: Optional[str] = None,
    ) -> Playbook:
        """
        Cria um Playbook com regras pre-extraidas (confirmacao do preview).

        Args:
            name: Nome do playbook
            area: Area juridica
            user_id: ID do usuario
            rules_data: Lista de dicts com as regras
            db: Sessao do banco
            description: Descricao opcional

        Returns:
            Playbook criado com regras
        """
        playbook_id = str(uuid.uuid4())
        playbook = Playbook(
            id=playbook_id,
            user_id=user_id,
            name=name,
            description=description or f"Playbook importado de documento para area de {area}",
            area=area,
            rules=[],
            is_active=True,
            is_template=False,
            scope="personal",
            version=1,
        )
        db.add(playbook)

        rules_json_inline: List[Dict[str, Any]] = []

        for idx, rule_data in enumerate(rules_data):
            if not isinstance(rule_data, dict):
                continue

            rule = PlaybookRule(
                id=str(uuid.uuid4()),
                playbook_id=playbook_id,
                clause_type=rule_data.get("clause_type", "geral"),
                rule_name=rule_data.get("rule_name", f"Regra {idx + 1}"),
                description=rule_data.get("description"),
                preferred_position=rule_data.get("preferred_position", ""),
                fallback_positions=rule_data.get("fallback_positions", []),
                rejected_positions=rule_data.get("rejected_positions", []),
                action_on_reject=rule_data.get("action_on_reject", "flag"),
                severity=rule_data.get("severity", "medium"),
                guidance_notes=rule_data.get("guidance_notes"),
                order=rule_data.get("order", idx),
                is_active=True,
            )
            db.add(rule)
            rules_json_inline.append(rule_data)

        playbook.rules = rules_json_inline

        await db.commit()
        await db.refresh(playbook)

        logger.info(
            "Playbook '%s' criado com %d regras (id=%s)",
            name, len(rules_json_inline), playbook_id,
        )

        return playbook

    # -----------------------------------------------------------------------
    # export_playbook
    # -----------------------------------------------------------------------

    async def export_playbook(
        self,
        playbook_id: str,
        format: str,
        user_id: str,
        db: AsyncSession,
    ) -> Tuple[bytes, str, str]:
        """
        Exporta um playbook no formato especificado.

        Args:
            playbook_id: ID do playbook
            format: Formato de exportacao (json, pdf, docx)
            user_id: ID do usuario solicitante
            db: Sessao do banco

        Returns:
            Tupla (conteudo_bytes, nome_arquivo, content_type)

        Raises:
            ValueError: Se playbook nao encontrado ou formato invalido
        """
        if format not in ("json", "pdf", "docx"):
            raise ValueError(f"Formato '{format}' nao suportado. Use: json, pdf, docx")

        playbook = await self._load_playbook(playbook_id, db)
        rules = await self._load_playbook_rules(playbook_id, db)

        safe_name = playbook.name.replace(" ", "_").replace("/", "_")[:50]

        if format == "json":
            return self._export_as_json(playbook, rules, safe_name)
        elif format == "pdf":
            return self._export_as_pdf(playbook, rules, safe_name)
        elif format == "docx":
            return self._export_as_docx(playbook, rules, safe_name)
        else:
            raise ValueError(f"Formato '{format}' nao suportado")

    # -----------------------------------------------------------------------
    # Export helpers
    # -----------------------------------------------------------------------

    def _export_as_json(
        self,
        playbook: Playbook,
        rules: List[PlaybookRule],
        safe_name: str,
    ) -> Tuple[bytes, str, str]:
        """Exporta playbook como JSON formatado."""
        data = {
            "playbook": {
                "id": playbook.id,
                "name": playbook.name,
                "description": playbook.description,
                "area": playbook.area,
                "scope": playbook.scope,
                "version": playbook.version,
                "is_template": playbook.is_template,
                "created_at": playbook.created_at.isoformat() if playbook.created_at else None,
                "updated_at": playbook.updated_at.isoformat() if playbook.updated_at else None,
            },
            "rules": [
                {
                    "rule_name": r.rule_name,
                    "clause_type": r.clause_type,
                    "description": r.description,
                    "preferred_position": r.preferred_position,
                    "fallback_positions": r.fallback_positions or [],
                    "rejected_positions": r.rejected_positions or [],
                    "action_on_reject": r.action_on_reject,
                    "severity": r.severity,
                    "guidance_notes": r.guidance_notes,
                    "order": r.order,
                    "is_active": r.is_active,
                }
                for r in rules
            ],
            "meta": {
                "exported_at": datetime.now(timezone.utc).isoformat(),
                "total_rules": len(rules),
                "format_version": "1.0",
            },
        }
        content = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
        return content, f"{safe_name}.json", "application/json"

    def _export_as_pdf(
        self,
        playbook: Playbook,
        rules: List[PlaybookRule],
        safe_name: str,
    ) -> Tuple[bytes, str, str]:
        """Exporta playbook como PDF usando reportlab."""
        import io

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.lib.colors import HexColor
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                Table,
                TableStyle,
            )
        except ImportError:
            raise ValueError(
                "Biblioteca reportlab nao disponivel. "
                "Instale com: pip install reportlab"
            )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "PlaybookTitle",
            parent=styles["Title"],
            fontSize=18,
            spaceAfter=6,
            textColor=HexColor("#1e293b"),
        )
        heading_style = ParagraphStyle(
            "RuleHeading",
            parent=styles["Heading2"],
            fontSize=12,
            spaceAfter=4,
            textColor=HexColor("#334155"),
        )
        body_style = ParagraphStyle(
            "RuleBody",
            parent=styles["Normal"],
            fontSize=9,
            leading=12,
            spaceAfter=2,
        )
        note_style = ParagraphStyle(
            "NoteStyle",
            parent=styles["Normal"],
            fontSize=8,
            leading=10,
            textColor=HexColor("#64748b"),
            spaceAfter=4,
        )

        severity_colors = {
            "low": "#22c55e",
            "medium": "#eab308",
            "high": "#f97316",
            "critical": "#ef4444",
        }

        elements: list = []

        # Title
        elements.append(Paragraph(f"Playbook: {playbook.name}", title_style))
        if playbook.description:
            elements.append(Paragraph(playbook.description, body_style))
        elements.append(Paragraph(
            f"Area: {playbook.area or 'N/A'} | Regras: {len(rules)} | Versao: {playbook.version}",
            note_style,
        ))
        elements.append(Spacer(1, 8 * mm))

        # Rules
        for idx, rule in enumerate(rules, 1):
            sev_color = severity_colors.get(rule.severity, "#64748b")
            elements.append(Paragraph(
                f"{idx}. {rule.rule_name} "
                f'<font color="{sev_color}">[{rule.severity.upper()}]</font>',
                heading_style,
            ))

            elements.append(Paragraph(
                f"<b>Tipo:</b> {rule.clause_type}", body_style,
            ))
            if rule.description:
                elements.append(Paragraph(
                    f"<b>Descricao:</b> {rule.description}", body_style,
                ))
            elements.append(Paragraph(
                f"<b>Posicao Preferida:</b> {rule.preferred_position}", body_style,
            ))
            if rule.fallback_positions:
                fb = "; ".join(rule.fallback_positions)
                elements.append(Paragraph(
                    f"<b>Alternativas:</b> {fb}", body_style,
                ))
            if rule.rejected_positions:
                rp = "; ".join(rule.rejected_positions)
                elements.append(Paragraph(
                    f"<b>Posicoes Rejeitadas:</b> {rp}", body_style,
                ))
            elements.append(Paragraph(
                f"<b>Acao:</b> {rule.action_on_reject}", body_style,
            ))
            if rule.guidance_notes:
                elements.append(Paragraph(
                    f"<i>Notas: {rule.guidance_notes}</i>", note_style,
                ))
            elements.append(Spacer(1, 4 * mm))

        # Footer
        elements.append(Spacer(1, 10 * mm))
        elements.append(Paragraph(
            f"Exportado em {datetime.now(timezone.utc).strftime('%d/%m/%Y %H:%M UTC')} | Iudex",
            note_style,
        ))

        doc.build(elements)
        content = buffer.getvalue()
        buffer.close()

        return content, f"{safe_name}.pdf", "application/pdf"

    def _export_as_docx(
        self,
        playbook: Playbook,
        rules: List[PlaybookRule],
        safe_name: str,
    ) -> Tuple[bytes, str, str]:
        """Exporta playbook como DOCX usando python-docx."""
        import io

        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ValueError(
                "Biblioteca python-docx nao disponivel. "
                "Instale com: pip install python-docx"
            )

        doc = DocxDocument()

        # Title
        title = doc.add_heading(f"Playbook: {playbook.name}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if playbook.description:
            p = doc.add_paragraph(playbook.description)
            p.style.font.size = Pt(10)

        meta_text = (
            f"Area: {playbook.area or 'N/A'} | "
            f"Regras: {len(rules)} | "
            f"Versao: {playbook.version}"
        )
        meta_p = doc.add_paragraph(meta_text)
        meta_p.style.font.size = Pt(9)
        for run in meta_p.runs:
            run.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph()  # spacer

        severity_labels = {
            "low": "BAIXA",
            "medium": "MEDIA",
            "high": "ALTA",
            "critical": "CRITICA",
        }

        for idx, rule in enumerate(rules, 1):
            sev_label = severity_labels.get(rule.severity, rule.severity.upper())

            # Rule heading
            heading = doc.add_heading(
                f"{idx}. {rule.rule_name} [{sev_label}]", level=2
            )

            # Rule details
            doc.add_paragraph(f"Tipo de clausula: {rule.clause_type}").bold = True

            if rule.description:
                doc.add_paragraph(f"Descricao: {rule.description}")

            p_pref = doc.add_paragraph()
            run_bold = p_pref.add_run("Posicao Preferida: ")
            run_bold.bold = True
            p_pref.add_run(rule.preferred_position)

            if rule.fallback_positions:
                p_fb = doc.add_paragraph()
                run_bold = p_fb.add_run("Alternativas Aceitaveis: ")
                run_bold.bold = True
                p_fb.add_run("; ".join(rule.fallback_positions))

            if rule.rejected_positions:
                p_rj = doc.add_paragraph()
                run_bold = p_rj.add_run("Posicoes Rejeitadas: ")
                run_bold.bold = True
                p_rj.add_run("; ".join(rule.rejected_positions))

            p_action = doc.add_paragraph()
            run_bold = p_action.add_run("Acao: ")
            run_bold.bold = True
            p_action.add_run(rule.action_on_reject)

            if rule.guidance_notes:
                p_notes = doc.add_paragraph()
                run_notes = p_notes.add_run(f"Notas: {rule.guidance_notes}")
                run_notes.italic = True
                run_notes.font.color.rgb = RGBColor(100, 116, 139)

            doc.add_paragraph()  # spacer between rules

        # Footer
        footer_p = doc.add_paragraph(
            f"Exportado em {datetime.now(timezone.utc).strftime('%d/%m/%Y %H:%M UTC')} | Iudex"
        )
        for run in footer_p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(148, 163, 184)

        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()
        buffer.close()

        return (
            content,
            f"{safe_name}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # -----------------------------------------------------------------------
    # get_playbook_for_prompt
    # -----------------------------------------------------------------------

    async def get_playbook_for_prompt(
        self,
        playbook_id: str,
        db: AsyncSession,
    ) -> str:
        """
        Serializa as regras do playbook em um trecho de prompt estruturado
        para injeção no system prompt do agente de IA.

        Usado na página /minuta:
        1. Usuário abre contrato em /minuta
        2. Seleciona um Playbook no dropdown
        3. Este método gera o trecho de prompt
        4. O trecho é injetado no system prompt do agente
        5. O agente aplica as regras durante a revisão interativa

        Args:
            playbook_id: ID do playbook
            db: Sessão do banco

        Returns:
            String formatada para injeção no system prompt
        """
        playbook = await self._load_playbook(playbook_id, db)
        rules = await self._load_playbook_rules(playbook_id, db)

        if not rules:
            return ""

        rules_sections: List[str] = []
        for rule in rules:
            guidance_section = ""
            if rule.guidance_notes:
                guidance_section = f"- **Notas**: {rule.guidance_notes}"

            rule_text = PLAYBOOK_RULE_TEMPLATE.format(
                rule_name=rule.rule_name,
                severity=rule.severity.upper(),
                clause_type=rule.clause_type,
                preferred_position=rule.preferred_position,
                fallback_positions=", ".join(rule.fallback_positions or ["Nenhuma"]),
                rejected_positions=", ".join(rule.rejected_positions or ["Nenhuma"]),
                action_on_reject=rule.action_on_reject,
                guidance_notes_section=guidance_section,
            )
            rules_sections.append(rule_text)

        # Build party perspective section
        perspective = getattr(playbook, "party_perspective", "neutro") or "neutro"
        perspective_labels = {
            "contratante": "**Perspectiva**: CONTRATANTE — Priorize proteções para quem contrata/compra.",
            "contratado": "**Perspectiva**: CONTRATADO — Priorize proteções para quem presta o serviço/vende.",
            "neutro": "**Perspectiva**: Neutra — Análise equilibrada entre as partes.",
        }
        perspective_section = perspective_labels.get(perspective, perspective_labels["neutro"])

        return PLAYBOOK_FOR_AGENT_PROMPT.format(
            playbook_name=playbook.name,
            party_perspective_section=perspective_section,
            rules_section="\n".join(rules_sections),
        )

    # -----------------------------------------------------------------------
    # Métodos internos
    # -----------------------------------------------------------------------

    async def _load_playbook(self, playbook_id: str, db: AsyncSession) -> Playbook:
        """Carrega playbook do banco."""
        result = await db.execute(
            select(Playbook).where(
                Playbook.id == playbook_id,
                Playbook.is_active == True,  # noqa: E712
            )
        )
        playbook = result.scalar_one_or_none()
        if not playbook:
            raise ValueError(f"Playbook {playbook_id} não encontrado ou inativo")
        return playbook

    async def _load_playbook_rules(
        self, playbook_id: str, db: AsyncSession
    ) -> List[PlaybookRule]:
        """Carrega regras ativas de um playbook, ordenadas."""
        result = await db.execute(
            select(PlaybookRule)
            .where(
                PlaybookRule.playbook_id == playbook_id,
                PlaybookRule.is_active == True,  # noqa: E712
            )
            .order_by(PlaybookRule.order)
        )
        return list(result.scalars().all())

    async def _load_document_text(
        self, document_id: str, user_id: str, db: AsyncSession
    ) -> Optional[str]:
        """
        Carrega o texto de um documento.
        Prioriza extracted_text, depois content.
        """
        result = await db.execute(
            select(Document).where(
                Document.id == document_id,
                Document.user_id == user_id,
            )
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise ValueError(f"Documento {document_id} não encontrado para o usuário")

        return doc.extracted_text or doc.content

    async def _extract_clauses(
        self, contract_text: str
    ) -> List[Dict[str, str]]:
        """
        Extrai cláusulas do contrato via IA.

        Returns:
            Lista de dicts com keys: clause_type, title, location, text
        """
        prompt = CLAUSE_EXTRACTION_PROMPT.format(contract_text=contract_text[:30000])

        response = await _call_ai(
            prompt=prompt,
            system_instruction="Extraia cláusulas do contrato. Responda em JSON válido.",
            model=DEFAULT_ANALYSIS_MODEL,
            max_tokens=8000,
            temperature=0.1,
        )

        parsed = _safe_json_parse(response) if response else None

        if parsed and isinstance(parsed, list):
            return [c for c in parsed if isinstance(c, dict) and c.get("text")]

        logger.warning("Falha na extração de cláusulas, usando fallback")
        return []

    async def _analyze_all_rules(
        self,
        rules: List[PlaybookRule],
        extracted_clauses: List[Dict[str, str]],
        contract_text: str,
        party_perspective: str = "neutro",
    ) -> List[ClauseAnalysisResult]:
        """
        Analisa todas as regras contra as cláusulas extraídas.
        Usa semáforo para limitar concorrência.
        """
        # Criar mapa de tipo -> cláusulas
        clause_map: Dict[str, List[Dict[str, str]]] = {}
        for clause in extracted_clauses:
            ctype = clause.get("clause_type", "").lower()
            if ctype not in clause_map:
                clause_map[ctype] = []
            clause_map[ctype].append(clause)

        contract_context = contract_text[:2000]
        semaphore = asyncio.Semaphore(MAX_CONCURRENT_ANALYSES)

        async def _analyze_one(rule: PlaybookRule) -> ClauseAnalysisResult:
            async with semaphore:
                # Encontrar cláusula correspondente
                matching_clauses = clause_map.get(rule.clause_type.lower(), [])

                if not matching_clauses:
                    # Cláusula não encontrada no contrato
                    return ClauseAnalysisResult(
                        rule_id=rule.id,
                        rule_name=rule.rule_name,
                        clause_type=rule.clause_type,
                        found_in_contract=False,
                        original_text=None,
                        location=None,
                        classification=ClauseClassification.NOT_FOUND,
                        severity=AnalysisSeverity(rule.severity),
                        explanation=(
                            f"Cláusula do tipo '{rule.clause_type}' não foi encontrada "
                            f"no contrato. Verifique se o contrato aborda este tema."
                        ),
                        suggested_redline=None,
                        confidence=0.8,
                    )

                # Combinar texto de todas as cláusulas do mesmo tipo
                combined_text = "\n\n".join(
                    c.get("text", "") for c in matching_clauses
                )
                location = matching_clauses[0].get("location", "Não identificada")

                result = await self.analyze_clause(
                    clause_text=combined_text,
                    rule=rule,
                    contract_context=contract_context,
                    party_perspective=party_perspective,
                )
                # Atualizar location com o que foi extraído
                result.location = location
                return result

        # Executar análises em paralelo com limite de concorrência
        tasks = [_analyze_one(rule) for rule in rules]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        clause_results: List[ClauseAnalysisResult] = []
        for idx, res in enumerate(results):
            if isinstance(res, Exception):
                logger.error("Erro ao analisar regra %s: %s", rules[idx].rule_name, res)
                clause_results.append(
                    ClauseAnalysisResult(
                        rule_id=rules[idx].id,
                        rule_name=rules[idx].rule_name,
                        clause_type=rules[idx].clause_type,
                        found_in_contract=False,
                        original_text=None,
                        location=None,
                        classification=ClauseClassification.NEEDS_REVIEW,
                        severity=AnalysisSeverity(rules[idx].severity),
                        explanation=f"Erro na análise automática: {str(res)}",
                        suggested_redline=None,
                        confidence=0.0,
                    )
                )
            else:
                clause_results.append(res)

        return clause_results

    async def _generate_redlines_for_non_compliant(
        self,
        clause_results: List[ClauseAnalysisResult],
        rules_by_id: Dict[str, PlaybookRule],
    ) -> List[ClauseAnalysisResult]:
        """Gera redlines para cláusulas não-conformes."""
        semaphore = asyncio.Semaphore(MAX_CONCURRENT_ANALYSES)

        async def _gen_redline(result: ClauseAnalysisResult) -> ClauseAnalysisResult:
            if (
                result.classification != ClauseClassification.NON_COMPLIANT
                or not result.original_text
            ):
                return result

            rule = rules_by_id.get(result.rule_id)
            if not rule:
                return result

            # Só gera redline se a ação for redline ou suggest
            if rule.action_on_reject not in ("redline", "suggest"):
                return result

            async with semaphore:
                redline = await self.generate_redline(
                    original_clause=result.original_text,
                    rule=rule,
                    analysis=result,
                )
                result.suggested_redline = redline

            return result

        tasks = [_gen_redline(r) for r in clause_results]
        return await asyncio.gather(*tasks)

    async def _generate_summary(
        self,
        playbook_name: str,
        total_rules: int,
        compliant: int,
        needs_review: int,
        non_compliant: int,
        not_found: int,
        risk_score: float,
        clause_results: List[ClauseAnalysisResult],
    ) -> str:
        """Gera resumo executivo da análise."""
        clause_details = "\n".join(
            f"- **{c.rule_name}** ({c.clause_type}): {c.classification.value} "
            f"[{c.severity.value}] — {c.explanation[:150]}"
            for c in clause_results
        )

        prompt = PLAYBOOK_SUMMARY_PROMPT.format(
            playbook_name=playbook_name,
            total_rules=total_rules,
            compliant=compliant,
            needs_review=needs_review,
            non_compliant=non_compliant,
            not_found=not_found,
            risk_score=risk_score,
            clause_details=clause_details,
        )

        response = await _call_ai(
            prompt=prompt,
            system_instruction="Você é um advogado preparando um resumo executivo.",
            model=DEFAULT_ANALYSIS_MODEL,
            max_tokens=2000,
            temperature=0.3,
        )

        return response or (
            f"Análise do playbook '{playbook_name}' concluída. "
            f"Score de risco: {risk_score}/100. "
            f"Conformes: {compliant}, Revisão necessária: {needs_review}, "
            f"Não-conformes: {non_compliant}, Não encontradas: {not_found}."
        )

    # -----------------------------------------------------------------------
    # extract_winning_language
    # -----------------------------------------------------------------------

    async def extract_winning_language(
        self,
        document_ids: List[str],
        name: str,
        area: str,
        user_id: str,
        db: AsyncSession,
        description: Optional[str] = None,
    ) -> Playbook:
        """
        Extrai "linguagem vencedora" (winning language) de contratos já negociados.

        Diferente de generate_playbook_from_contracts (que identifica padrões gerais),
        este método foca em extrair a linguagem EXATA que foi aceita por ambas as partes
        em contratos já assinados, tratando-os como "precedentes de negociação".

        Fluxo:
        1. Carrega os textos de todos os contratos
        2. Usa IA com prompt especializado em winning language
        3. Para cada cláusula, extrai:
           - Linguagem vencedora (posição padrão)
           - Variações aceitas (posição de fallback)
           - Posições a evitar (inferidas do que NÃO foi aceito)
           - Importância baseada em recorrência entre contratos
        4. Cria Playbook com regras extraídas
        5. Persiste no banco

        Args:
            document_ids: Lista de IDs de contratos (1-10) já negociados
            name: Nome do playbook
            area: Área jurídica
            user_id: ID do usuário criador
            db: Sessão do banco
            description: Descrição opcional

        Returns:
            Playbook criado com regras de winning language

        Raises:
            ValueError: Se documentos não encontrados ou textos vazios
        """
        if len(document_ids) > 10:
            raise ValueError("Máximo de 10 contratos para extração de winning language")

        if not document_ids:
            raise ValueError("Informe ao menos 1 contrato para extração")

        logger.info(
            "Extraindo winning language para playbook '%s' (area=%s) de %d contratos",
            name, area, len(document_ids),
        )

        # 1. Carregar textos dos contratos
        contracts_text_parts: List[str] = []
        loaded_count = 0
        for idx, doc_id in enumerate(document_ids, 1):
            try:
                text = await self._load_document_text(doc_id, user_id, db)
                if text and text.strip():
                    truncated = text[:15000]
                    contracts_text_parts.append(
                        f"### CONTRATO {idx} (ID: {doc_id})\n{truncated}\n"
                    )
                    loaded_count += 1
                else:
                    logger.warning(
                        "Documento %s sem texto extraído, ignorando na extração de winning language",
                        doc_id,
                    )
            except ValueError as e:
                logger.warning("Documento %s não acessível: %s", doc_id, e)

        if not contracts_text_parts:
            raise ValueError(
                "Nenhum dos documentos fornecidos possui texto extraído. "
                "Certifique-se de que os documentos foram processados antes de extrair winning language."
            )

        contracts_text = "\n---\n".join(contracts_text_parts)

        # 2. Chamar IA com prompt de winning language
        prompt = WINNING_LANGUAGE_EXTRACTION_PROMPT.format(
            area=area,
            contracts_text=contracts_text,
        )

        response = await _call_ai(
            prompt=prompt,
            system_instruction=(
                "Você é um advogado sênior especialista em revisão contratual brasileira. "
                "Analise os contratos como 'linguagem vencedora' — cláusulas já negociadas e aceitas. "
                "Responda APENAS em JSON válido."
            ),
            model=DEFAULT_GENERATION_MODEL,
            max_tokens=8000,
            temperature=0.3,
        )

        rules_data = _safe_json_parse(response) if response else None

        if not rules_data or not isinstance(rules_data, list):
            raise ValueError(
                "Não foi possível extrair winning language dos contratos fornecidos. "
                "Tente novamente ou forneça contratos com mais conteúdo clausular."
            )

        # 3. Criar Playbook no banco
        playbook_id = str(uuid.uuid4())
        auto_description = description or (
            f"Playbook de winning language extraído de {loaded_count} contrato(s) "
            f"na área de {area}. Contém cláusulas já negociadas e aceitas."
        )
        playbook = Playbook(
            id=playbook_id,
            user_id=user_id,
            name=name,
            description=auto_description,
            area=area,
            rules=[],
            is_active=True,
            is_template=False,
            scope="personal",
            version=1,
            metadata_={
                "source": "winning_language_extraction",
                "source_document_ids": document_ids,
                "source_documents_loaded": loaded_count,
            },
        )
        db.add(playbook)

        # 4. Criar PlaybookRules
        rules_json_inline: List[Dict[str, Any]] = []

        for idx, rule_data in enumerate(rules_data):
            if not isinstance(rule_data, dict):
                continue

            rule = PlaybookRule(
                id=str(uuid.uuid4()),
                playbook_id=playbook_id,
                clause_type=rule_data.get("clause_type", "geral"),
                rule_name=rule_data.get("rule_name", f"Regra {idx + 1}"),
                description=rule_data.get("description"),
                preferred_position=rule_data.get("preferred_position", ""),
                fallback_positions=rule_data.get("fallback_positions", []),
                rejected_positions=rule_data.get("rejected_positions", []),
                action_on_reject=rule_data.get("action_on_reject", "flag"),
                severity=rule_data.get("severity", "medium"),
                guidance_notes=rule_data.get("guidance_notes"),
                order=idx,
                is_active=True,
                metadata_={"source": "winning_language"},
            )
            db.add(rule)
            rules_json_inline.append(rule_data)

        # Atualizar JSON inline no playbook
        playbook.rules = rules_json_inline

        await db.commit()
        await db.refresh(playbook)

        logger.info(
            "Winning language extraída: playbook '%s' criado com %d regras (id=%s)",
            name, len(rules_json_inline), playbook_id,
        )

        return playbook


# ---------------------------------------------------------------------------
# Singleton
# ---------------------------------------------------------------------------

playbook_service = PlaybookService()
