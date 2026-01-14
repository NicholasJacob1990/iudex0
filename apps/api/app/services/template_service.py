"""
Template Service
Serviço para aplicação de variáveis em templates DOCX
"""

from typing import Dict, List, Any, Optional
from loguru import logger
import re
from pathlib import Path
import shutil
import uuid
from docx import Document
from app.core.config import settings
from app.schemas.smart_template import SmartTemplate, TemplateRenderInput, BlockType

class TemplateService:
    """Serviço para manipulação de templates DOCX"""
    
    def __init__(self):
        self.storage_path = Path(settings.LOCAL_STORAGE_PATH)
        logger.info("TemplateService inicializado")
    
    async def apply_template(
        self,
        template_path: str,
        variables: Dict[str, str],
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Aplica variáveis em um template DOCX
        
        Args:
            template_path: Caminho para o arquivo template
            variables: Dicionário de variáveis {chave: valor}
            output_filename: Nome do arquivo de saída (opcional)
            
        Returns:
            Dict com caminho do arquivo gerado e metadados
        """
        logger.info(f"Aplicando template: {template_path}")
        
        try:
            # Verificar se arquivo existe
            if not Path(template_path).exists():
                raise FileNotFoundError(f"Template não encontrado: {template_path}")
            
            # Carregar documento
            doc = Document(template_path)
            
            # Substituir variáveis nos parágrafos
            replacements_count = 0
            for paragraph in doc.paragraphs:
                replacements_count += self._replace_in_paragraph(paragraph, variables)
            
            # Substituir variáveis nas tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replacements_count += self._replace_in_paragraph(paragraph, variables)
            
            # Gerar caminho de saída
            if not output_filename:
                output_filename = f"generated_{uuid.uuid4()}.docx"
            
            output_path = self.storage_path / "generated" / output_filename
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Salvar documento
            doc.save(output_path)
            logger.info(f"Documento gerado: {output_path} ({replacements_count} substituições)")
            
            return {
                "success": True,
                "file_path": str(output_path),
                "filename": output_filename,
                "replacements": replacements_count
            }
            
        except Exception as e:
            logger.error(f"Erro ao aplicar template: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _replace_in_paragraph(self, paragraph, variables: Dict[str, str]) -> int:
        """Substitui variáveis em um parágrafo preservando formatação"""
        count = 0
        if not paragraph.text:
            return 0
            
        # Identificar variáveis no texto: {{variavel}}
        text = paragraph.text
        matches = re.findall(r'\{\{([^}]+)\}\}', text)
        
        if not matches:
            return 0
            
        # Estratégia simples: substituir no texto completo do parágrafo
        # Nota: Isso pode perder formatação específica dentro da variável se ela estiver quebrada em runs
        # Para produção robusta, seria necessário iterar sobre runs, mas é complexo em python-docx
        
        new_text = text
        for var_name in matches:
            key = var_name.strip()
            if key in variables:
                new_text = new_text.replace(f"{{{{{var_name}}}}}", str(variables[key]))
                count += 1
        
        if new_text != text:
            paragraph.text = new_text
            
        return count
    
    async def extract_variables(self, file_path: str) -> List[str]:
        """Extrai lista de variáveis {{variavel}} de um DOCX"""
        try:
            doc = Document(file_path)
            variables = set()
            
            def find_vars(text):
                return re.findall(r'\{\{([^}]+)\}\}', text)
            
            for p in doc.paragraphs:
                for var in find_vars(p.text):
                    variables.add(var.strip())
                    
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for var in find_vars(p.text):
                                variables.add(var.strip())
                                
            return list(variables)
            
        except Exception as e:
            logger.error(f"Erro ao extrair variáveis: {e}")
            return []

    def assemble_smart_template(self, template: SmartTemplate, input_data: TemplateRenderInput) -> str:
        """
        Monta um documento final a partir de um SmartTemplate e inputs do usuário.
        Respeita as regras de bloqueio (lock) e permissões de edição.
        """
        doc_parts = []
        
        for block in template.blocks:
            # 1. Verificar Condição
            if block.condition:
                cond_val = input_data.variables.get(block.condition)
                if cond_val is False:
                    continue
            
            block_content = ""
            
            # 2. Resolver Conteúdo Base
            if block.type == BlockType.FIXED:
                block_content = block.content or ""
                
            elif block.type == BlockType.VARIABLE:
                var_name = block.variable_name
                val = input_data.variables.get(var_name)
                if val:
                    block_content = str(val)
                else:
                    block_content = f"[{var_name}]"
                    
            elif block.type == BlockType.AI:
                block_content = f"> [IA: {block.title}]"
            
            elif block.type == BlockType.CLAUSE:
                block_content = block.content or f"[Cláusula: {block.title}]"

            # 3. Aplicar Overrides (Edições do Usuário ou Conteúdo Gerado)
            if block.id in input_data.overrides:
                override_text = input_data.overrides[block.id]
                
                if block.type == BlockType.FIXED and not block.user_can_edit:
                    logger.warning(f"Tentativa de sobrescrever bloco fixo {block.id} ignorada.")
                else:
                    block_content = override_text

            doc_parts.append(block_content)
            
        return "\n\n".join(doc_parts)

# Instância global
template_service = TemplateService()
