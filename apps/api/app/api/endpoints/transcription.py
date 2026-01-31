from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks, Body, Request
from fastapi.responses import StreamingResponse, FileResponse
from typing import Optional, Any, Dict, List
from datetime import datetime
import asyncio
import shutil
import os
import uuid
import logging
import io
import json
import sys
from pathlib import Path
import hashlib
import subprocess
from app.schemas.transcription import TranscriptionRequest, HearingSpeakersUpdateRequest
from app.services.transcription_service import TranscriptionService
from app.services.job_manager import job_manager
from app.services.api_call_tracker import job_context, usage_context
from app.core.config import settings
from docx import Document
from pydantic import BaseModel
import re
from app.services.preventive_hil import build_preventive_hil_issues
from urllib.parse import urlparse
import ipaddress

class ExportRequest(BaseModel):
    content: str
    filename: str = "transcription.docx"
    document_theme: str = "classic"
    document_header: Optional[str] = None
    document_footer: Optional[str] = None
    document_margins: str = "normal"
    document_font_family: Optional[str] = None
    document_font_size: Optional[float] = None
    document_line_height: Optional[float] = None
    document_paragraph_spacing: Optional[float] = None


class JobQualityUpdateRequest(BaseModel):
    validation_report: Optional[Dict[str, Any]] = None
    analysis_result: Optional[Dict[str, Any]] = None
    selected_fix_ids: Optional[List[str]] = None
    applied_fixes: Optional[List[str]] = None
    suggestions: Optional[str] = None
    fixed_content: Optional[str] = None
    needs_revalidate: Optional[bool] = None
    applied_issue_ids: Optional[List[str]] = None


class JobContentUpdateRequest(BaseModel):
    content: Optional[str] = None
    rich_text_html: Optional[str] = None
    rich_text_json: Optional[Dict[str, Any]] = None
    rich_text_meta: Optional[Dict[str, Any]] = None
    needs_revalidate: Optional[bool] = None


class UrlVomoJobRequest(BaseModel):
    url: str
    mode: str = "APOSTILA"
    thinking_level: str = "medium"
    custom_prompt: Optional[str] = None
    document_theme: str = "classic"
    document_header: Optional[str] = None
    document_footer: Optional[str] = None
    document_margins: str = "normal"
    document_page_frame: bool = True
    document_show_header_footer: bool = True
    document_font_family: Optional[str] = None
    document_font_size: Optional[float] = None
    document_line_height: Optional[float] = None
    document_paragraph_spacing: Optional[float] = None
    model_selection: str = "gemini-3-flash-preview"
    high_accuracy: bool = False
    diarization: Optional[bool] = None
    diarization_strict: bool = False
    use_cache: bool = True
    auto_apply_fixes: bool = True
    auto_apply_content_fixes: bool = False
    skip_legal_audit: bool = False
    skip_audit: bool = False
    skip_fidelity_audit: bool = False
    skip_sources_audit: bool = False


class UrlHearingJobRequest(BaseModel):
    url: str
    case_id: str
    goal: str = "alegacoes_finais"
    thinking_level: str = "medium"
    model_selection: str = "gemini-3-flash-preview"
    high_accuracy: bool = False
    format_mode: str = "AUDIENCIA"
    custom_prompt: Optional[str] = None
    document_theme: str = "classic"
    document_header: Optional[str] = None
    document_footer: Optional[str] = None
    document_margins: str = "normal"
    document_page_frame: bool = True
    document_show_header_footer: bool = True
    document_font_family: Optional[str] = None
    document_font_size: Optional[float] = None
    document_line_height: Optional[float] = None
    document_paragraph_spacing: Optional[float] = None
    format_enabled: bool = True
    include_timestamps: bool = True
    allow_indirect: bool = False
    allow_summary: bool = False
    use_cache: bool = True
    auto_apply_fixes: bool = True
    auto_apply_content_fixes: bool = False
    skip_legal_audit: bool = False
    skip_fidelity_audit: bool = False
    skip_sources_audit: bool = False


def _url_is_public_and_allowed(url: str) -> tuple[bool, str]:
    """
    Política de segurança para importar URLs públicas (anti-SSRF).

    Default: permite apenas hosts de YouTube (configurável por env).
    """
    raw = (url or "").strip()
    if not raw:
        return False, "URL vazia."
    try:
        parsed = urlparse(raw)
    except Exception:
        return False, "URL inválida."
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return False, "URL deve começar com http:// ou https://"

    host = (parsed.hostname or "").strip().lower()
    if not host:
        return False, "Host inválido."

    allow_all = str(os.getenv("IUDEX_PUBLIC_URL_ALLOW_ALL", "0")).strip().lower() in {"1", "true", "yes", "on"}

    # Bloquear IPs privados/loopback/link-local (SSRF) quando allow_all estiver ligado também.
    try:
        ip = ipaddress.ip_address(host)
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast:
            return False, "Host/IP não permitido."
    except ValueError:
        # host não é IP; ok
        pass

    if allow_all:
        return True, ""

    allowlist = os.getenv("IUDEX_PUBLIC_URL_ALLOWLIST_HOSTS", "youtube.com,youtu.be")
    allowed = [h.strip().lower() for h in allowlist.split(",") if h.strip()]
    host_no_www = host[4:] if host.startswith("www.") else host
    for candidate in allowed:
        cand = candidate[4:] if candidate.startswith("www.") else candidate
        if host_no_www == cand or host_no_www.endswith(f".{cand}"):
            return True, ""
    return False, f"Host não permitido. Permitidos: {', '.join(allowed) or 'nenhum'}"


def _download_public_url_to_job_input(url: str, job_dir: Path, *, index: int = 1) -> tuple[str, str]:
    """
    Baixa mídia de URL pública usando `yt-dlp` para cache e copia para `job_dir/input`.

    Retorna (file_path, file_name).
    """
    ok, reason = _url_is_public_and_allowed(url)
    if not ok:
        raise HTTPException(status_code=400, detail=reason)

    ytdlp = (
        (os.getenv("IUDEX_YTDLP_PATH") or "").strip()
        or shutil.which("yt-dlp")
        or shutil.which("yt_dlp")
        or ("/opt/homebrew/bin/yt-dlp" if os.path.exists("/opt/homebrew/bin/yt-dlp") else None)
        or ("/usr/local/bin/yt-dlp" if os.path.exists("/usr/local/bin/yt-dlp") else None)
    )
    if not ytdlp:
        raise HTTPException(
            status_code=500,
            detail="Servidor sem `yt-dlp`. Instale `yt-dlp` para importar URLs (ex.: YouTube).",
        )

    storage_root = _get_storage_root()
    cache_dir = storage_root / "url_imports"
    cache_dir.mkdir(parents=True, exist_ok=True)

    url_norm = url.strip()
    url_hash = hashlib.sha256(url_norm.encode("utf-8")).hexdigest()[:12]
    base = f"url_{url_hash}"

    cached_mp3 = cache_dir / f"{base}.mp3"
    if not (cached_mp3.exists() and cached_mp3.stat().st_size > 1024):
        outtmpl = str(cache_dir / f"{base}.%(ext)s")
        cmd = [
            ytdlp,
            "--no-playlist",
            "--restrict-filenames",
            "-f",
            "bestaudio/best",
            "--extract-audio",
            "--audio-format",
            "mp3",
            "--audio-quality",
            "0",
            "-o",
            outtmpl,
            url_norm,
        ]
        max_mb = os.getenv("IUDEX_PUBLIC_URL_MAX_FILESIZE_MB")
        if max_mb:
            try:
                mb = int(str(max_mb).strip())
                if mb > 0:
                    cmd += ["--max-filesize", f"{mb}M"]
            except Exception:
                pass
        try:
            subprocess.run(cmd, check=True)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"Falha ao baixar URL: {exc}")

    input_dir = job_dir / "input"
    input_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"{base}.mp3"
    final_name = f"{index:02d}_{safe_name}"
    dest_path = input_dir / final_name
    shutil.copyfile(cached_mp3, dest_path)

    return str(dest_path), safe_name

router = APIRouter()
logger = logging.getLogger(__name__)

# Instância global do serviço (carrega modelos na init se necessário, aqui é leve)
service = TranscriptionService()

# Registro local de execução (permite cancelamento best-effort)
_transcription_tasks: Dict[str, asyncio.Task] = {}
_transcription_cancel_events: Dict[str, asyncio.Event] = {}


def _get_cancel_event(job_id: str) -> asyncio.Event:
    event = _transcription_cancel_events.get(job_id)
    if event is None:
        event = asyncio.Event()
        _transcription_cancel_events[job_id] = event
    return event


def _register_task(job_id: str, task: asyncio.Task) -> None:
    _transcription_tasks[job_id] = task


def _cleanup_task(job_id: str) -> None:
    _transcription_tasks.pop(job_id, None)
    _transcription_cancel_events.pop(job_id, None)


def _get_event_source_response():
    try:
        from sse_starlette.sse import EventSourceResponse
        return EventSourceResponse
    except Exception as exc:
        logger.error(f"SSE indisponível (sse_starlette): {exc}")
        raise HTTPException(
            status_code=500,
            detail="SSE indisponível no servidor. Instale 'sse-starlette' para habilitar streaming."
        )

def _get_transcription_jobs_dir() -> Path:
    storage_path = Path(settings.LOCAL_STORAGE_PATH) if settings else Path("./storage")
    if not storage_path.is_absolute():
        backend_root = Path(__file__).resolve().parents[3]
        storage_path = backend_root / storage_path
    base_dir = storage_path
    jobs_dir = base_dir / "transcription_jobs"
    jobs_dir.mkdir(parents=True, exist_ok=True)
    return jobs_dir

def _get_storage_root() -> Path:
    storage_path = Path(settings.LOCAL_STORAGE_PATH) if settings else Path("./storage")
    if not storage_path.is_absolute():
        backend_root = Path(__file__).resolve().parents[3]
        storage_path = backend_root / storage_path
    return storage_path.resolve()

def _get_allowed_report_roots() -> List[Path]:
    roots: List[Path] = []
    base_setting = Path(settings.LOCAL_STORAGE_PATH) if settings else Path("./storage")
    if base_setting.is_absolute():
        roots.append(base_setting.resolve())
    else:
        backend_root = Path(__file__).resolve().parents[3]
        roots.append((backend_root / base_setting).resolve())
        repo_root = Path(__file__).resolve().parents[5]
        roots.append((repo_root / base_setting).resolve())
    roots.append(_get_transcription_jobs_dir().resolve())
    unique: List[Path] = []
    for root in roots:
        if root not in unique:
            unique.append(root)
    return unique

def _safe_remove_path(path_value: str, allowed_roots: List[Path]) -> bool:
    if not path_value:
        return False
    resolved = _resolve_job_path(path_value).resolve()
    allowed = False
    for root in allowed_roots:
        try:
            resolved.relative_to(root)
            allowed = True
            break
        except ValueError:
            continue
    if not allowed or not resolved.exists():
        return False
    if resolved.is_dir():
        shutil.rmtree(resolved, ignore_errors=True)
    else:
        resolved.unlink(missing_ok=True)
    return True

def _get_job_dir(job_id: str) -> Path:
    job_dir = _get_transcription_jobs_dir() / job_id
    job_dir.mkdir(parents=True, exist_ok=True)
    return job_dir

def _resolve_job_path(path_value: str) -> Path:
    candidate = Path(path_value)
    if candidate.is_absolute():
        return candidate
    backend_root = Path(__file__).resolve().parents[3]
    alt_backend = backend_root / candidate
    if alt_backend.exists():
        return alt_backend
    repo_root = Path(__file__).resolve().parents[5]
    alt_repo = repo_root / candidate
    if alt_repo.exists():
        return alt_repo
    return candidate

def _infer_video_name_and_mode_suffix(reports: Dict[str, Any], fallback_mode: str = "APOSTILA") -> tuple[str, str]:
    """
    Infer `video_name` and `mode_suffix` from persisted report filenames.
    Expected patterns:
      - {video_name}_{MODE}.md
      - {video_name}_{MODE}_ANALISE.json
      - {video_name}_{MODE}_FIDELIDADE.json
    """
    mode_suffix = (fallback_mode or "APOSTILA").upper()
    candidates = [
        reports.get("md_path"),
        reports.get("analysis_path"),
        reports.get("validation_path"),
        reports.get("fidelity_path"),
        reports.get("audit_path"),
        reports.get("raw_path"),
    ]
    for value in candidates:
        if not value:
            continue
        name = Path(str(value)).name
        stem = Path(name).stem
        parts = stem.split("_")
        if len(parts) < 2:
            continue
        maybe_mode = parts[-1].upper()
        if maybe_mode in {"APOSTILA", "FIDELIDADE", "RAW"}:
            return "_".join(parts[:-1]), maybe_mode
        # Pattern: {video}_{mode}_{suffix}
        maybe_mode = parts[-2].upper()
        if maybe_mode in {"APOSTILA", "FIDELIDADE", "RAW"}:
            return "_".join(parts[:-2]), maybe_mode
    # Fallback: use job id-ish label
    return (Path(str(reports.get("output_dir") or "transcricao")).name or "transcricao"), mode_suffix

def _join_report_path(output_dir_value: str, filename: str) -> str:
    base = Path(output_dir_value)
    joined = base / filename
    # Keep the same “style” as the base path: absolute stays absolute, relative stays relative.
    return str(joined)

def _build_genai_client():
    from google import genai

    project = settings.GOOGLE_CLOUD_PROJECT or os.getenv("GOOGLE_CLOUD_PROJECT")
    location = settings.VERTEX_AI_LOCATION or os.getenv("VERTEX_AI_LOCATION") or "global"
    api_key = (
        settings.GOOGLE_API_KEY
        or os.getenv("GOOGLE_API_KEY")
        or os.getenv("GEMINI_API_KEY")
    )

    if project:
        return genai.Client(vertexai=True, project=project, location=location)
    if api_key:
        return genai.Client(api_key=api_key)
    raise RuntimeError("Google GenAI não configurado (defina GOOGLE_CLOUD_PROJECT ou GOOGLE_API_KEY/GEMINI_API_KEY).")

def _find_job_result_path(job: Dict[str, Any]) -> Path:
    result_path = job.get("result_path")
    if result_path:
        resolved = _resolve_job_path(result_path)
        if resolved.exists():
            return resolved

    job_id = job.get("job_id") or job.get("jobid")
    if job_id:
        candidates = [
            _get_transcription_jobs_dir() / job_id / "result.json",
            Path(__file__).resolve().parents[3] / "storage" / "transcription_jobs" / job_id / "result.json",
            Path(__file__).resolve().parents[5] / "storage" / "transcription_jobs" / job_id / "result.json",
        ]
        for candidate in candidates:
            if candidate.exists():
                return candidate

    raise FileNotFoundError("Result file not found")

def _save_uploaded_files(files: List[UploadFile], job_dir: Path) -> Dict[str, List[str]]:
    input_dir = job_dir / "input"
    input_dir.mkdir(parents=True, exist_ok=True)
    file_paths: List[str] = []
    file_names: List[str] = []
    for idx, f in enumerate(files):
        safe_name = Path(f.filename).name
        final_name = f"{idx + 1:02d}_{safe_name}"
        dest_path = input_dir / final_name
        with open(dest_path, "wb") as buffer:
            shutil.copyfileobj(f.file, buffer)
        file_paths.append(str(dest_path))
        file_names.append(safe_name)
    return {"file_paths": file_paths, "file_names": file_names}

def _write_vomo_job_result(
    job_dir: Path,
    result: Any,
    mode: str,
    file_names: List[str],
) -> str:
    if isinstance(result, dict):
        content = result.get("content") or ""
        raw_content = result.get("raw_content") or content
        reports = result.get("reports") or {}
        audit_issues = result.get("audit_issues") or []
        quality = result.get("quality")
    else:
        content = result or ""
        raw_content = content
        reports = {}
        audit_issues = []
        quality = None

    content_path = job_dir / "content.md"
    raw_path = job_dir / "raw.txt"
    content_path.write_text(content, encoding="utf-8")
    raw_path.write_text(raw_content, encoding="utf-8")

    reports_path = None
    if reports:
        reports_path = job_dir / "reports.json"
        reports_path.write_text(json.dumps(reports, ensure_ascii=False, indent=2), encoding="utf-8")

    audit_path = None
    if audit_issues:
        audit_path = job_dir / "audit_issues.json"
        audit_path.write_text(json.dumps(audit_issues, ensure_ascii=False, indent=2), encoding="utf-8")

    result_data = {
        "job_type": "vomo",
        "mode": mode,
        "file_names": file_names,
        "content_path": str(content_path),
        "raw_path": str(raw_path),
        "reports_path": str(reports_path) if reports_path else None,
        "audit_path": str(audit_path) if audit_path else None,
        "audit_issues": audit_issues,
        "quality": quality,
        "rich_text_html_path": None,
        "rich_text_json_path": None,
        "rich_text_meta_path": None,
    }
    result_path = job_dir / "result.json"
    result_path.write_text(json.dumps(result_data, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(result_path)

def _write_hearing_job_result(job_dir: Path, result: Dict[str, Any]) -> str:
    payload = result.get("hearing") if isinstance(result, dict) else {}
    paths = result.get("paths") if isinstance(result, dict) else {}

    payload_path = job_dir / "hearing_payload.json"
    payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    transcript_path = None
    if payload and payload.get("transcript_markdown"):
        transcript_path = job_dir / "hearing_transcript.md"
        transcript_path.write_text(payload.get("transcript_markdown"), encoding="utf-8")

    formatted_path = None
    if payload and payload.get("formatted_text"):
        formatted_path = job_dir / "hearing_formatted.md"
        formatted_path.write_text(payload.get("formatted_text"), encoding="utf-8")

    result_data = {
        "job_type": "hearing",
        "case_id": payload.get("case_id") if isinstance(payload, dict) else None,
        "payload_path": str(payload_path),
        "transcript_path": str(transcript_path) if transcript_path else None,
        "formatted_path": str(formatted_path) if formatted_path else None,
        "paths": paths or {},
        "rich_text_html_path": None,
        "rich_text_json_path": None,
        "rich_text_meta_path": None,
    }
    result_path = job_dir / "result.json"
    result_path.write_text(json.dumps(result_data, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(result_path)

def _load_job_result_payload(job: Dict[str, Any]) -> Dict[str, Any]:
    resolved_path = _find_job_result_path(job)
    if job.get("result_path") != str(resolved_path):
        job_manager.update_transcription_job(job.get("job_id"), result_path=str(resolved_path))
    with open(resolved_path, "r", encoding="utf-8") as f:
        result_data = json.load(f)

    job_type = result_data.get("job_type")
    if job_type == "hearing":
        payload_path = result_data.get("payload_path")
        if payload_path and os.path.exists(payload_path):
            with open(payload_path, "r", encoding="utf-8") as pf:
                payload = json.load(pf)
        else:
            payload = None
        paths = result_data.get("paths") or {}
        rich_text_html = None
        rich_text_json = None
        rich_text_meta = None
        rich_text_html_path = result_data.get("rich_text_html_path")
        rich_text_json_path = result_data.get("rich_text_json_path")
        rich_text_meta_path = result_data.get("rich_text_meta_path")
        if rich_text_html_path:
            candidate = _resolve_job_path(str(rich_text_html_path))
            if candidate.exists():
                rich_text_html = candidate.read_text(encoding="utf-8", errors="ignore")
        if rich_text_json_path:
            candidate = _resolve_job_path(str(rich_text_json_path))
            if candidate.exists():
                try:
                    rich_text_json = json.loads(candidate.read_text(encoding="utf-8", errors="ignore"))
                except Exception:
                    rich_text_json = None
        if rich_text_meta_path:
            candidate = _resolve_job_path(str(rich_text_meta_path))
            if candidate.exists():
                try:
                    rich_text_meta = json.loads(candidate.read_text(encoding="utf-8", errors="ignore"))
                except Exception:
                    rich_text_meta = None

        return {
            "job_type": "hearing",
            "payload": payload,
            "paths": paths,
            "reports": paths,
            "rich_text_html": rich_text_html,
            "rich_text_json": rich_text_json,
            "rich_text_meta": rich_text_meta,
        }

    content = ""
    raw_content = ""
    reports = None
    audit_issues = result_data.get("audit_issues") or []
    quality = result_data.get("quality")
    mode = result_data.get("mode")
    content_path = result_data.get("content_path")
    raw_path = result_data.get("raw_path")
    reports_path = result_data.get("reports_path")
    audit_path = result_data.get("audit_path")
    if content_path:
        content_candidate = _resolve_job_path(content_path)
        if content_candidate.exists():
            content = content_candidate.read_text(encoding="utf-8", errors="ignore")
    if raw_path:
        raw_candidate = _resolve_job_path(raw_path)
        if raw_candidate.exists():
            raw_content = raw_candidate.read_text(encoding="utf-8", errors="ignore")
    if reports_path:
        reports_candidate = _resolve_job_path(reports_path)
        if reports_candidate.exists():
            with open(reports_candidate, "r", encoding="utf-8") as rf:
                reports = json.load(rf)
    if not audit_issues and audit_path:
        audit_candidate = _resolve_job_path(audit_path)
        if audit_candidate.exists():
            with open(audit_candidate, "r", encoding="utf-8") as af:
                audit_issues = json.load(af)
    rich_text_html = None
    rich_text_json = None
    rich_text_meta = None
    rich_text_html_path = result_data.get("rich_text_html_path")
    rich_text_json_path = result_data.get("rich_text_json_path")
    rich_text_meta_path = result_data.get("rich_text_meta_path")
    if rich_text_html_path:
        candidate = _resolve_job_path(str(rich_text_html_path))
        if candidate.exists():
            rich_text_html = candidate.read_text(encoding="utf-8", errors="ignore")
    if rich_text_json_path:
        candidate = _resolve_job_path(str(rich_text_json_path))
        if candidate.exists():
            try:
                rich_text_json = json.loads(candidate.read_text(encoding="utf-8", errors="ignore"))
            except Exception:
                rich_text_json = None
    if rich_text_meta_path:
        candidate = _resolve_job_path(str(rich_text_meta_path))
        if candidate.exists():
            try:
                rich_text_meta = json.loads(candidate.read_text(encoding="utf-8", errors="ignore"))
            except Exception:
                rich_text_meta = None

    return {
        "job_type": "vomo",
        "mode": mode,
        "content": content,
        "raw_content": raw_content,
        "reports": reports,
        "audit_issues": audit_issues,
        "quality": quality,
        "rich_text_html": rich_text_html,
        "rich_text_json": rich_text_json,
        "rich_text_meta": rich_text_meta,
    }

@router.post("/export/docx")
async def export_docx(request: ExportRequest):
    """
    Converte texto/markdown para DOCX usando VomoMLX.save_as_word.
    Usa a mesma lógica de formatação premium do mlx_vomo.py.
    """
    import tempfile
    import sys
    import os
    
    # Add project root to path
    PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../../../../"))
    if PROJECT_ROOT not in sys.path:
        sys.path.append(PROJECT_ROOT)
    
    try:
        from mlx_vomo import VomoMLX
        
        # Create temp directory for output
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extract video name from filename
            video_name = request.filename.replace('.docx', '').replace('_', ' ')
            
            # Initialize VomoMLX (lightweight, just for save_as_word)
            vomo = VomoMLX(provider="gemini")
            
            # Call the real save_as_word method
            output_path = vomo.save_as_word(
                formatted_text=request.content,
                video_name=video_name,
                output_folder=temp_dir,
                mode="APOSTILA",
                document_theme=request.document_theme,
                document_header=request.document_header,
                document_footer=request.document_footer,
                document_margins=request.document_margins,
                document_font_family=request.document_font_family,
                document_font_size=request.document_font_size,
                document_line_height=request.document_line_height,
                document_paragraph_spacing=request.document_paragraph_spacing,
            )
            
            if output_path and os.path.exists(output_path):
                # Read the generated file
                with open(output_path, 'rb') as f:
                    docx_content = f.read()
                
                buffer = io.BytesIO(docx_content)
                buffer.seek(0)
                
                return StreamingResponse(
                    buffer,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={request.filename}"}
                )
            else:
                raise HTTPException(status_code=500, detail="Falha ao gerar documento Word")
                
    except ImportError as e:
        logger.error(f"Erro ao importar VomoMLX: {str(e)}")
        # Fallback to simple DOCX generation
        from docx import Document
        doc = Document()
        doc.add_heading(request.filename.replace('.docx', ''), 0)
        for para in request.content.split('\n'):
            if para.strip():
                doc.add_paragraph(para)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={request.filename}"}
        )
    except Exception as e:
        logger.error(f"Erro ao exportar DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/vomo/jobs")
async def create_vomo_job(
    files: list[UploadFile] = File(...),
    mode: str = Form("APOSTILA"),
    thinking_level: str = Form("medium"),
    custom_prompt: Optional[str] = Form(None),
    document_theme: str = Form("classic"),
    document_header: Optional[str] = Form(None),
    document_footer: Optional[str] = Form(None),
    document_margins: str = Form("normal"),
    document_page_frame: bool = Form(True),
    document_show_header_footer: bool = Form(True),
    document_font_family: Optional[str] = Form(None),
    document_font_size: Optional[float] = Form(None),
    document_line_height: Optional[float] = Form(None),
    document_paragraph_spacing: Optional[float] = Form(None),
    model_selection: str = Form("gemini-3-flash-preview"),
    high_accuracy: bool = Form(False),
    diarization: Optional[bool] = Form(None),
    diarization_strict: bool = Form(False),
    use_cache: bool = Form(True),
    auto_apply_fixes: bool = Form(True),
    auto_apply_content_fixes: bool = Form(False),
    skip_legal_audit: bool = Form(False),
    skip_audit: bool = Form(False),
    skip_fidelity_audit: bool = Form(False),
    skip_sources_audit: bool = Form(False),
):
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="No files provided")

    job_id = str(uuid.uuid4())
    job_dir = _get_job_dir(job_id)
    if custom_prompt is not None:
        custom_prompt = custom_prompt.strip() or None
    if document_header is not None:
        document_header = document_header.strip() or None
    if document_footer is not None:
        document_footer = document_footer.strip() or None
    if document_font_family is not None:
        document_font_family = document_font_family.strip() or None

    try:
        saved = _save_uploaded_files(files, job_dir)
    except Exception as exc:
        logger.error(f"Falha ao salvar arquivos do job {job_id}: {exc}")
        raise HTTPException(status_code=500, detail="Falha ao salvar arquivos do job.")

    file_paths = saved["file_paths"]
    file_names = saved["file_names"]
    mode = (mode or "APOSTILA").strip().upper()

    effective_skip_legal_audit = skip_legal_audit or skip_audit
    config = {
        "mode": mode,
        "thinking_level": thinking_level,
        "custom_prompt": custom_prompt,
        "document_theme": document_theme,
        "document_header": document_header,
        "document_footer": document_footer,
        "document_margins": document_margins,
        "document_page_frame": bool(document_page_frame),
        "document_show_header_footer": bool(document_show_header_footer),
        "document_font_family": document_font_family,
        "document_font_size": document_font_size,
        "document_line_height": document_line_height,
        "document_paragraph_spacing": document_paragraph_spacing,
        "model_selection": model_selection,
        "high_accuracy": high_accuracy,
        "diarization": diarization,
        "diarization_strict": diarization_strict,
        "use_cache": use_cache,
        "auto_apply_fixes": auto_apply_fixes,
        "auto_apply_content_fixes": auto_apply_content_fixes,
        "skip_legal_audit": effective_skip_legal_audit,
        "skip_fidelity_audit": skip_fidelity_audit,
        "skip_sources_audit": skip_sources_audit,
    }

    metadata_path = job_dir / "metadata.json"
    metadata_path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")

    job_manager.create_transcription_job(
        job_id=job_id,
        job_type="vomo",
        config=config,
        file_names=file_names,
        file_paths=file_paths,
        status="queued",
        progress=0,
        stage="queued",
        message="Aguardando início",
        result_path=str((job_dir / "result.json").resolve()),
    )

    cancel_event = _get_cancel_event(job_id)

    async def run_job():
        def ensure_not_cancelled():
            if cancel_event.is_set():
                raise asyncio.CancelledError()
            current = job_manager.get_transcription_job(job_id)
            if current and current.get("status") == "canceled":
                raise asyncio.CancelledError()

        async def on_progress(stage: str, progress: int, message: str):
            ensure_not_cancelled()
            job_manager.update_transcription_job(
                job_id,
                status="running",
                progress=progress,
                stage=stage,
                message=message,
            )

        try:
            ensure_not_cancelled()
            job_manager.update_transcription_job(
                job_id,
                status="running",
                progress=0,
                stage="starting",
                message="Iniciando processamento",
            )
            with job_context(job_id):
                if len(file_paths) == 1:
                    result = await service.process_file_with_progress(
                        file_path=file_paths[0],
                        mode=mode,
                        thinking_level=thinking_level,
                        custom_prompt=custom_prompt,
                        high_accuracy=high_accuracy,
                        diarization=diarization,
                        diarization_strict=diarization_strict,
                        on_progress=on_progress,
                        model_selection=model_selection,
                        use_cache=use_cache,
                        auto_apply_fixes=auto_apply_fixes,
                        auto_apply_content_fixes=auto_apply_content_fixes,
                        skip_legal_audit=effective_skip_legal_audit,
                        skip_audit=skip_audit,
                        skip_fidelity_audit=skip_fidelity_audit,
                        skip_sources_audit=skip_sources_audit,
                    )
                else:
                    result = await service.process_batch_with_progress(
                        file_paths=file_paths,
                        file_names=file_names,
                        mode=mode,
                        thinking_level=thinking_level,
                        custom_prompt=custom_prompt,
                        high_accuracy=high_accuracy,
                        diarization=diarization,
                        diarization_strict=diarization_strict,
                        model_selection=model_selection,
                        on_progress=on_progress,
                        use_cache=use_cache,
                        auto_apply_fixes=auto_apply_fixes,
                        auto_apply_content_fixes=auto_apply_content_fixes,
                        skip_legal_audit=effective_skip_legal_audit,
                        skip_audit=skip_audit,
                        skip_fidelity_audit=skip_fidelity_audit,
                        skip_sources_audit=skip_sources_audit,
                    )

            ensure_not_cancelled()
            result_path = _write_vomo_job_result(job_dir, result, mode, file_names)
            result_path = str(Path(result_path).resolve())
            job_manager.update_transcription_job(
                job_id,
                status="completed",
                progress=100,
                stage="complete",
                message="Concluído",
                result_path=result_path,
            )
        except asyncio.CancelledError:
            job_manager.update_transcription_job(
                job_id,
                status="canceled",
                progress=100,
                stage="canceled",
                message="Cancelado pelo usuário.",
            )
        except Exception as e:
            logger.error(f"Erro no job {job_id}: {e}")
            job_manager.update_transcription_job(
                job_id,
                status="error",
                progress=100,
                stage="error",
                message=str(e),
                error=str(e),
            )
        finally:
            _cleanup_task(job_id)

    task = asyncio.create_task(run_job())
    _register_task(job_id, task)

    return {"job_id": job_id, "status": "queued"}


@router.post("/vomo/jobs/url")
async def create_vomo_job_from_url(request: UrlVomoJobRequest = Body(...)):
    url = (request.url or "").strip()
    if request.custom_prompt is not None:
        request.custom_prompt = request.custom_prompt.strip() or None
    if request.document_header is not None:
        request.document_header = request.document_header.strip() or None
    if request.document_footer is not None:
        request.document_footer = request.document_footer.strip() or None
    if request.document_font_family is not None:
        request.document_font_family = request.document_font_family.strip() or None
    if request.document_font_family is not None:
        request.document_font_family = request.document_font_family.strip() or None

    job_id = str(uuid.uuid4())
    job_dir = _get_job_dir(job_id)

    file_path, file_name = _download_public_url_to_job_input(url, job_dir, index=1)

    file_paths = [file_path]
    file_names = [file_name]
    mode = (request.mode or "APOSTILA").strip().upper()

    effective_skip_legal_audit = bool(request.skip_legal_audit or request.skip_audit)
    config = {
        "mode": mode,
        "thinking_level": request.thinking_level,
        "custom_prompt": request.custom_prompt,
        "document_theme": request.document_theme,
        "document_header": request.document_header,
        "document_footer": request.document_footer,
        "document_margins": request.document_margins,
        "document_page_frame": bool(request.document_page_frame),
        "document_show_header_footer": bool(request.document_show_header_footer),
        "document_font_family": request.document_font_family,
        "document_font_size": request.document_font_size,
        "document_line_height": request.document_line_height,
        "document_paragraph_spacing": request.document_paragraph_spacing,
        "model_selection": request.model_selection,
        "high_accuracy": bool(request.high_accuracy),
        "diarization": request.diarization,
        "diarization_strict": bool(request.diarization_strict),
        "use_cache": bool(request.use_cache),
        "auto_apply_fixes": bool(request.auto_apply_fixes),
        "auto_apply_content_fixes": bool(request.auto_apply_content_fixes),
        "skip_legal_audit": effective_skip_legal_audit,
        "skip_fidelity_audit": bool(request.skip_fidelity_audit),
        "skip_sources_audit": bool(request.skip_sources_audit),
        "source_url": url,
    }

    metadata_path = job_dir / "metadata.json"
    metadata_path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")

    job_manager.create_transcription_job(
        job_id=job_id,
        job_type="vomo",
        config=config,
        file_names=file_names,
        file_paths=file_paths,
        status="queued",
        progress=0,
        stage="queued",
        message="Aguardando início",
        result_path=str((job_dir / "result.json").resolve()),
    )

    cancel_event = _get_cancel_event(job_id)

    async def run_job():
        def ensure_not_cancelled():
            if cancel_event.is_set():
                raise asyncio.CancelledError()
            current = job_manager.get_transcription_job(job_id)
            if current and current.get("status") == "canceled":
                raise asyncio.CancelledError()

        async def on_progress(stage: str, progress: int, message: str):
            ensure_not_cancelled()
            job_manager.update_transcription_job(
                job_id,
                status="running",
                progress=progress,
                stage=stage,
                message=message,
            )

        try:
            ensure_not_cancelled()
            job_manager.update_transcription_job(
                job_id,
                status="running",
                progress=0,
                stage="starting",
                message="Iniciando processamento",
            )
            with job_context(job_id):
                result = await service.process_file_with_progress(
                    file_path=file_paths[0],
                    mode=mode,
                    thinking_level=request.thinking_level,
                    custom_prompt=request.custom_prompt,
                    high_accuracy=bool(request.high_accuracy),
                    diarization=request.diarization,
                    diarization_strict=bool(request.diarization_strict),
                    on_progress=on_progress,
                    model_selection=request.model_selection,
                    use_cache=bool(request.use_cache),
                    auto_apply_fixes=bool(request.auto_apply_fixes),
                    auto_apply_content_fixes=bool(request.auto_apply_content_fixes),
                    skip_legal_audit=effective_skip_legal_audit,
                    skip_audit=bool(request.skip_audit),
                    skip_fidelity_audit=bool(request.skip_fidelity_audit),
                    skip_sources_audit=bool(request.skip_sources_audit),
                )

            ensure_not_cancelled()
            result_path = _write_vomo_job_result(job_dir, result, mode, file_names)
            result_path = str(Path(result_path).resolve())
            job_manager.update_transcription_job(
                job_id,
                status="completed",
                progress=100,
                stage="complete",
                message="Concluído",
                result_path=result_path,
            )
        except asyncio.CancelledError:
            job_manager.update_transcription_job(
                job_id,
                status="canceled",
                progress=100,
                stage="canceled",
                message="Cancelado pelo usuário.",
            )
        except Exception as e:
            logger.error(f"Erro no job {job_id}: {e}")
            job_manager.update_transcription_job(
                job_id,
                status="error",
                progress=100,
                stage="error",
                message=str(e),
                error=str(e),
            )
        finally:
            _cleanup_task(job_id)

    task = asyncio.create_task(run_job())
    _register_task(job_id, task)
    return {"job_id": job_id, "status": "queued"}

@router.post("/hearing/jobs")
async def create_hearing_job(
    file: UploadFile = File(...),
    case_id: str = Form(...),
    goal: str = Form("alegacoes_finais"),
    thinking_level: str = Form("medium"),
    model_selection: str = Form("gemini-3-flash-preview"),
    high_accuracy: bool = Form(False),
    format_mode: str = Form("AUDIENCIA"),
    custom_prompt: Optional[str] = Form(None),
    format_enabled: bool = Form(True),
    include_timestamps: bool = Form(True),
    document_theme: str = Form("classic"),
    document_header: Optional[str] = Form(None),
    document_footer: Optional[str] = Form(None),
    document_margins: str = Form("normal"),
    document_page_frame: bool = Form(True),
    document_show_header_footer: bool = Form(True),
    document_font_family: Optional[str] = Form(None),
    document_font_size: Optional[float] = Form(None),
    document_line_height: Optional[float] = Form(None),
    document_paragraph_spacing: Optional[float] = Form(None),
    allow_indirect: bool = Form(False),
    allow_summary: bool = Form(False),
    use_cache: bool = Form(True),
    auto_apply_fixes: bool = Form(True),
    auto_apply_content_fixes: bool = Form(False),
    skip_legal_audit: bool = Form(False),
    skip_fidelity_audit: bool = Form(False),
    skip_sources_audit: bool = Form(False),
):
    job_id = str(uuid.uuid4())
    job_dir = _get_job_dir(job_id)
    if custom_prompt is not None:
        custom_prompt = custom_prompt.strip() or None
    if document_header is not None:
        document_header = document_header.strip() or None
    if document_footer is not None:
        document_footer = document_footer.strip() or None
    if document_font_family is not None:
        document_font_family = document_font_family.strip() or None

    try:
        saved = _save_uploaded_files([file], job_dir)
    except Exception as exc:
        logger.error(f"Falha ao salvar arquivo do job {job_id}: {exc}")
        raise HTTPException(status_code=500, detail="Falha ao salvar arquivo do job.")

    file_paths = saved["file_paths"]
    file_names = saved["file_names"]

    config = {
        "case_id": case_id,
        "goal": goal,
        "thinking_level": thinking_level,
        "model_selection": model_selection,
        "high_accuracy": high_accuracy,
        "format_mode": format_mode,
        "custom_prompt": custom_prompt,
        "document_theme": document_theme,
        "document_header": document_header,
        "document_footer": document_footer,
        "document_margins": document_margins,
        "document_page_frame": bool(document_page_frame),
        "document_show_header_footer": bool(document_show_header_footer),
        "document_font_family": document_font_family,
        "document_font_size": document_font_size,
        "document_line_height": document_line_height,
        "document_paragraph_spacing": document_paragraph_spacing,
        "format_enabled": format_enabled,
        "include_timestamps": include_timestamps,
        "allow_indirect": allow_indirect,
        "allow_summary": allow_summary,
        "use_cache": use_cache,
        "auto_apply_fixes": auto_apply_fixes,
        "auto_apply_content_fixes": auto_apply_content_fixes,
        "skip_legal_audit": skip_legal_audit,
        "skip_fidelity_audit": skip_fidelity_audit,
        "skip_sources_audit": skip_sources_audit,
    }

    metadata_path = job_dir / "metadata.json"
    metadata_path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")

    job_manager.create_transcription_job(
        job_id=job_id,
        job_type="hearing",
        config=config,
        file_names=file_names,
        file_paths=file_paths,
        status="queued",
        progress=0,
        stage="queued",
        message="Aguardando início",
        result_path=str((job_dir / "result.json").resolve()),
    )

    cancel_event = _get_cancel_event(job_id)

    async def run_job():
        def ensure_not_cancelled():
            if cancel_event.is_set():
                raise asyncio.CancelledError()
            current = job_manager.get_transcription_job(job_id)
            if current and current.get("status") == "canceled":
                raise asyncio.CancelledError()

        async def on_progress(stage: str, progress: int, message: str):
            ensure_not_cancelled()
            job_manager.update_transcription_job(
                job_id,
                status="running",
                progress=progress,
                stage=stage,
                message=message,
            )

        try:
            ensure_not_cancelled()
            job_manager.update_transcription_job(
                job_id,
                status="running",
                progress=0,
                stage="starting",
                message="Iniciando processamento",
            )
            with job_context(job_id):
                result = await service.process_hearing_with_progress(
                    file_path=file_paths[0],
                    case_id=case_id,
                    goal=goal,
                    thinking_level=thinking_level,
                    model_selection=model_selection,
                    high_accuracy=high_accuracy,
                    format_mode=format_mode,
                    custom_prompt=custom_prompt,
                    format_enabled=format_enabled,
                    include_timestamps=include_timestamps,
                    allow_indirect=allow_indirect,
                    allow_summary=allow_summary,
                    use_cache=use_cache,
                    auto_apply_fixes=auto_apply_fixes,
                    auto_apply_content_fixes=auto_apply_content_fixes,
                    skip_legal_audit=skip_legal_audit,
                    skip_fidelity_audit=skip_fidelity_audit,
                    skip_sources_audit=skip_sources_audit,
                    on_progress=on_progress,
                )
            ensure_not_cancelled()
            result_path = _write_hearing_job_result(job_dir, result)
            result_path = str(Path(result_path).resolve())
            job_manager.update_transcription_job(
                job_id,
                status="completed",
                progress=100,
                stage="complete",
                message="Concluído",
                result_path=result_path,
            )
        except asyncio.CancelledError:
            job_manager.update_transcription_job(
                job_id,
                status="canceled",
                progress=100,
                stage="canceled",
                message="Cancelado pelo usuário.",
            )
        except Exception as e:
            logger.error(f"Erro no job {job_id}: {e}")
            job_manager.update_transcription_job(
                job_id,
                status="error",
                progress=100,
                stage="error",
                message=str(e),
                error=str(e),
            )
        finally:
            _cleanup_task(job_id)

    task = asyncio.create_task(run_job())
    _register_task(job_id, task)

    return {"job_id": job_id, "status": "queued"}


@router.post("/hearing/jobs/url")
async def create_hearing_job_from_url(request: UrlHearingJobRequest = Body(...)):
    url = (request.url or "").strip()
    if request.custom_prompt is not None:
        request.custom_prompt = request.custom_prompt.strip() or None
    if request.document_header is not None:
        request.document_header = request.document_header.strip() or None
    if request.document_footer is not None:
        request.document_footer = request.document_footer.strip() or None
    if not (request.case_id or "").strip():
        raise HTTPException(status_code=400, detail="case_id é obrigatório.")

    job_id = str(uuid.uuid4())
    job_dir = _get_job_dir(job_id)

    file_path, file_name = _download_public_url_to_job_input(url, job_dir, index=1)
    file_paths = [file_path]
    file_names = [file_name]

    config = {
        "case_id": request.case_id,
        "goal": request.goal,
        "thinking_level": request.thinking_level,
        "model_selection": request.model_selection,
        "high_accuracy": bool(request.high_accuracy),
        "format_mode": request.format_mode,
        "custom_prompt": request.custom_prompt,
        "document_theme": request.document_theme,
        "document_header": request.document_header,
        "document_footer": request.document_footer,
        "document_margins": request.document_margins,
        "document_page_frame": bool(request.document_page_frame),
        "document_show_header_footer": bool(request.document_show_header_footer),
        "document_font_family": request.document_font_family,
        "document_font_size": request.document_font_size,
        "document_line_height": request.document_line_height,
        "document_paragraph_spacing": request.document_paragraph_spacing,
        "format_enabled": bool(request.format_enabled),
        "include_timestamps": bool(getattr(request, "include_timestamps", True)),
        "allow_indirect": bool(request.allow_indirect),
        "allow_summary": bool(request.allow_summary),
        "use_cache": bool(request.use_cache),
        "auto_apply_fixes": bool(request.auto_apply_fixes),
        "auto_apply_content_fixes": bool(request.auto_apply_content_fixes),
        "skip_legal_audit": bool(request.skip_legal_audit),
        "skip_fidelity_audit": bool(request.skip_fidelity_audit),
        "skip_sources_audit": bool(request.skip_sources_audit),
        "source_url": url,
    }

    metadata_path = job_dir / "metadata.json"
    metadata_path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")

    job_manager.create_transcription_job(
        job_id=job_id,
        job_type="hearing",
        config=config,
        file_names=file_names,
        file_paths=file_paths,
        status="queued",
        progress=0,
        stage="queued",
        message="Aguardando início",
        result_path=str((job_dir / "result.json").resolve()),
    )

    cancel_event = _get_cancel_event(job_id)

    async def run_job():
        def ensure_not_cancelled():
            if cancel_event.is_set():
                raise asyncio.CancelledError()
            current = job_manager.get_transcription_job(job_id)
            if current and current.get("status") == "canceled":
                raise asyncio.CancelledError()

        async def on_progress(stage: str, progress: int, message: str):
            ensure_not_cancelled()
            job_manager.update_transcription_job(
                job_id,
                status="running",
                progress=progress,
                stage=stage,
                message=message,
            )

        try:
            ensure_not_cancelled()
            job_manager.update_transcription_job(
                job_id,
                status="running",
                progress=0,
                stage="starting",
                message="Iniciando processamento",
            )
            with job_context(job_id):
                result = await service.process_hearing_with_progress(
                    file_path=file_paths[0],
                    case_id=request.case_id,
                    goal=request.goal,
                    thinking_level=request.thinking_level,
                    model_selection=request.model_selection,
                    high_accuracy=bool(request.high_accuracy),
                    format_mode=request.format_mode,
                    custom_prompt=request.custom_prompt,
                    format_enabled=bool(request.format_enabled),
                    include_timestamps=bool(getattr(request, "include_timestamps", True)),
                    allow_indirect=bool(request.allow_indirect),
                    allow_summary=bool(request.allow_summary),
                    use_cache=bool(request.use_cache),
                    auto_apply_fixes=bool(request.auto_apply_fixes),
                    auto_apply_content_fixes=bool(request.auto_apply_content_fixes),
                    skip_legal_audit=bool(request.skip_legal_audit),
                    skip_fidelity_audit=bool(request.skip_fidelity_audit),
                    skip_sources_audit=bool(request.skip_sources_audit),
                    on_progress=on_progress,
                )

            ensure_not_cancelled()
            result_path = _write_hearing_job_result(job_dir, result)
            result_path = str(Path(result_path).resolve())
            job_manager.update_transcription_job(
                job_id,
                status="completed",
                progress=100,
                stage="complete",
                message="Concluído",
                result_path=result_path,
            )
        except asyncio.CancelledError:
            job_manager.update_transcription_job(
                job_id,
                status="canceled",
                progress=100,
                stage="canceled",
                message="Cancelado pelo usuário.",
            )
        except Exception as e:
            logger.error(f"Erro no job {job_id}: {e}")
            job_manager.update_transcription_job(
                job_id,
                status="error",
                progress=100,
                stage="error",
                message=str(e),
                error=str(e),
            )
        finally:
            _cleanup_task(job_id)

    task = asyncio.create_task(run_job())
    _register_task(job_id, task)
    return {"job_id": job_id, "status": "queued"}

@router.get("/jobs")
async def list_transcription_jobs(
    limit: int = 20,
    status: Optional[str] = None,
    job_type: Optional[str] = None,
):
    jobs = job_manager.list_transcription_jobs(limit=limit, status=status, job_type=job_type)
    return {"jobs": jobs}

@router.get("/jobs/{job_id}")
async def get_transcription_job(job_id: str):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@router.post("/jobs/{job_id}/cancel")
async def cancel_transcription_job(job_id: str):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    status = job.get("status")
    if status in {"completed", "error", "canceled"}:
        raise HTTPException(status_code=409, detail=f"Job already {status}")

    cancel_event = _get_cancel_event(job_id)
    cancel_event.set()

    task = _transcription_tasks.get(job_id)
    if task and not task.done():
        task.cancel()

    job_manager.update_transcription_job(
        job_id,
        status="canceled",
        progress=job.get("progress") or 0,
        stage="canceled",
        message="Cancelado pelo usuário.",
    )
    return {"success": True, "status": "canceled"}

@router.get("/jobs/{job_id}/result")
async def get_transcription_job_result(job_id: str):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Job not completed")
    try:
        payload = _load_job_result_payload(job)
        if payload.get("job_type") == "vomo" and not payload.get("audit_issues") and payload.get("content"):
            try:
                from app.services.quality_service import quality_service
                document_name = (job.get("file_names") or [job_id])[0]
                analysis_report = await quality_service.analyze_structural_issues(
                    content=payload.get("content") or "",
                    document_name=document_name,
                    raw_content=payload.get("raw_content") or ""
                )
                payload["audit_issues"] = service._build_audit_issues(
                    analysis_report,
                    document_name,
                    raw_content=payload.get("raw_content") or "",
                    formatted_content=payload.get("content") or "",
                )
            except Exception as audit_error:
                logger.warning(f"Falha ao recomputar auditoria HIL para job {job_id}: {audit_error}")
        return {"job_id": job_id, **payload}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to load result: {exc}")


class ConvertPreventiveToHilResponse(BaseModel):
    job_id: str
    added: int
    total: int
    audit_issues: List[Dict[str, Any]]


@router.post("/jobs/{job_id}/convert-preventive-to-hil", response_model=ConvertPreventiveToHilResponse)
async def convert_preventive_alerts_to_hil(job_id: str):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Job not completed")

    payload = _load_job_result_payload(job)
    if payload.get("job_type") != "vomo":
        raise HTTPException(status_code=409, detail="Only vomo jobs are supported for conversion")

    reports = payload.get("reports") or {}
    json_path = reports.get("preventive_fidelity_json_path")
    if not json_path:
        raise HTTPException(status_code=404, detail="Preventive fidelity report not found for this job")

    preventive_path = _resolve_job_path(str(json_path)).resolve()
    if not preventive_path.exists():
        raise HTTPException(status_code=404, detail="Preventive fidelity JSON file missing")

    try:
        with open(preventive_path, "r", encoding="utf-8") as f:
            preventive_audit = json.load(f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read preventive audit JSON: {e}")

    formatted_content = payload.get("content") or ""
    converted = build_preventive_hil_issues(preventive_audit, formatted_content=formatted_content)
    if not converted:
        return ConvertPreventiveToHilResponse(job_id=job_id, added=0, total=len(payload.get("audit_issues") or []), audit_issues=payload.get("audit_issues") or [])

    existing: List[Dict[str, Any]] = payload.get("audit_issues") or []
    existing_ids = {str(i.get("id")) for i in existing if isinstance(i, dict) and i.get("id")}
    existing_keys = {
        f"{i.get('type','')}:{i.get('reference','')}:{i.get('description','')}"
        for i in existing
        if isinstance(i, dict)
    }

    new_issues: List[Dict[str, Any]] = []
    for issue in converted:
        if not isinstance(issue, dict):
            continue
        issue_id = str(issue.get("id") or "")
        if issue_id and issue_id in existing_ids:
            continue
        key = f"{issue.get('type','')}:{issue.get('reference','')}:{issue.get('description','')}"
        if key in existing_keys:
            continue
        new_issues.append(issue)

    if not new_issues:
        return ConvertPreventiveToHilResponse(job_id=job_id, added=0, total=len(existing), audit_issues=existing)

    merged = existing + new_issues

    # Persist to job folder (audit_issues.json + result.json)
    job_dir = (_get_transcription_jobs_dir() / job_id).resolve()
    result_path = (job_dir / "result.json").resolve()
    if not result_path.exists():
        raise HTTPException(status_code=500, detail="Job result.json missing")

    try:
        with open(result_path, "r", encoding="utf-8") as f:
            result_data = json.load(f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read job result.json: {e}")

    audit_path = (job_dir / "audit_issues.json").resolve()
    try:
        audit_path.write_text(json.dumps(merged, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to write audit_issues.json: {e}")

    result_data["audit_path"] = str(audit_path)
    result_data["audit_issues"] = merged
    try:
        result_path.write_text(json.dumps(result_data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update job result.json: {e}")

    return ConvertPreventiveToHilResponse(job_id=job_id, added=len(new_issues), total=len(merged), audit_issues=merged)


class MergeAuditIssuesRequest(BaseModel):
    issues: List[Dict[str, Any]]


class MergeAuditIssuesResponse(BaseModel):
    job_id: str
    added: int
    total: int
    audit_issues: List[Dict[str, Any]]


@router.post("/jobs/{job_id}/audit-issues/merge", response_model=MergeAuditIssuesResponse)
async def merge_audit_issues(job_id: str, request: MergeAuditIssuesRequest):
    """
    Persistently merge extra audit issues into a job snapshot.

    Used to "send" alerts from the Quality tab to the HIL Corrections tab and
    keep them after reload (writes audit_issues.json + result.json).
    """
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Job not completed")

    payload = _load_job_result_payload(job)
    if payload.get("job_type") != "vomo":
        raise HTTPException(status_code=409, detail="Only vomo jobs are supported for merging issues")

    incoming = request.issues or []
    incoming = [i for i in incoming if isinstance(i, dict)]
    if not incoming:
        return MergeAuditIssuesResponse(job_id=job_id, added=0, total=len(payload.get("audit_issues") or []), audit_issues=payload.get("audit_issues") or [])

    existing: List[Dict[str, Any]] = payload.get("audit_issues") or []
    existing_ids = {str(i.get("id")) for i in existing if isinstance(i, dict) and i.get("id")}
    existing_keys = {
        f"{i.get('type','')}:{i.get('reference','')}:{i.get('description','')}"
        for i in existing
        if isinstance(i, dict)
    }

    new_issues: List[Dict[str, Any]] = []
    for issue in incoming:
        issue_id = str(issue.get("id") or "")
        if issue_id and issue_id in existing_ids:
            continue
        key = f"{issue.get('type','')}:{issue.get('reference','')}:{issue.get('description','')}"
        if key in existing_keys:
            continue
        new_issues.append(issue)

    if not new_issues:
        return MergeAuditIssuesResponse(job_id=job_id, added=0, total=len(existing), audit_issues=existing)

    merged = existing + new_issues

    job_dir = (_get_transcription_jobs_dir() / job_id).resolve()
    result_path = (job_dir / "result.json").resolve()
    if not result_path.exists():
        raise HTTPException(status_code=500, detail="Job result.json missing")

    try:
        result_data = json.loads(result_path.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read job result.json: {e}")

    audit_path = (job_dir / "audit_issues.json").resolve()
    try:
        audit_path.write_text(json.dumps(merged, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to write audit_issues.json: {e}")

    result_data["audit_path"] = str(audit_path)
    result_data["audit_issues"] = merged
    try:
        result_path.write_text(json.dumps(result_data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update job result.json: {e}")

    return MergeAuditIssuesResponse(job_id=job_id, added=len(new_issues), total=len(merged), audit_issues=merged)


@router.delete("/jobs/{job_id}")
async def delete_transcription_job(job_id: str, delete_outputs: bool = True):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") in {"queued", "running"}:
        raise HTTPException(status_code=409, detail="Job is still running")

    removed_paths: List[str] = []
    allowed_roots = _get_allowed_report_roots()

    if delete_outputs:
        try:
            payload = _load_job_result_payload(job)
            reports = payload.get("reports") or {}
            output_dir = reports.get("output_dir")
            if output_dir and _safe_remove_path(output_dir, allowed_roots):
                removed_paths.append(output_dir)
        except Exception:
            pass

    job_dir = (_get_transcription_jobs_dir() / job_id).resolve()
    if _safe_remove_path(str(job_dir), allowed_roots):
        removed_paths.append(str(job_dir))

    job_manager.clear_events(job_id)
    removed = job_manager.delete_transcription_job(job_id)
    if not removed:
        raise HTTPException(status_code=500, detail="Failed to delete job")

    return {"success": True, "removed_paths": removed_paths}

@router.get("/jobs/{job_id}/reports/{report_key}")
async def download_transcription_report(job_id: str, report_key: str):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Job not completed")
    try:
        payload = _load_job_result_payload(job)
        reports = payload.get("reports") or {}
        report_path = reports.get(report_key)
        if not report_path:
            raise HTTPException(status_code=404, detail="Report not found")
        resolved = _resolve_job_path(report_path).resolve()
        allowed_roots = _get_allowed_report_roots()
        allowed = False
        for root in allowed_roots:
            try:
                resolved.relative_to(root)
                allowed = True
                break
            except ValueError:
                continue
        if not allowed:
            raise HTTPException(status_code=403, detail="Report path not allowed")
        if not resolved.exists():
            raise HTTPException(status_code=404, detail="Report file missing")
        suffix = resolved.suffix.lower()
        media_type = "application/octet-stream"
        if suffix in [".md", ".markdown"]:
            media_type = "text/markdown"
        elif suffix == ".txt":
            media_type = "text/plain"
        elif suffix == ".json":
            media_type = "application/json"
        elif suffix == ".docx":
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return FileResponse(resolved, media_type=media_type, filename=resolved.name)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to download report: {exc}")

@router.get("/jobs/{job_id}/media")
async def get_job_media(job_id: str, index: int = 0):
    """
    Serve the original audio/video file from a transcription job.
    Used by the frontend player to replay the source media.
    """
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    job_dir = _get_transcription_jobs_dir() / job_id
    input_dir = job_dir / "input"

    if not input_dir.exists():
        raise HTTPException(status_code=404, detail="Input directory not found")

    # List input files sorted by name (01_file.mp3, 02_file.mp3, etc.)
    input_files = sorted(input_dir.iterdir())
    if not input_files:
        raise HTTPException(status_code=404, detail="No input files found")

    if index < 0 or index >= len(input_files):
        raise HTTPException(status_code=404, detail=f"File index {index} out of range (0-{len(input_files)-1})")

    file_path = input_files[index]
    suffix = file_path.suffix.lower()

    # Determine media type
    media_types = {
        ".mp3": "audio/mpeg",
        ".wav": "audio/wav",
        ".m4a": "audio/mp4",
        ".aac": "audio/aac",
        ".ogg": "audio/ogg",
        ".mp4": "video/mp4",
        ".mov": "video/quicktime",
        ".mkv": "video/x-matroska",
        ".webm": "video/webm",
    }
    media_type = media_types.get(suffix, "application/octet-stream")

    return FileResponse(
        file_path,
        media_type=media_type,
        filename=file_path.name,
        headers={"Accept-Ranges": "bytes"}  # Enable seeking in audio/video
    )

@router.get("/jobs/{job_id}/media/list")
async def list_job_media(job_id: str):
    """
    List all media files available for a transcription job.
    """
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    job_dir = _get_transcription_jobs_dir() / job_id
    input_dir = job_dir / "input"

    if not input_dir.exists():
        return {"files": []}

    input_files = sorted(input_dir.iterdir())
    files = []
    for idx, f in enumerate(input_files):
        files.append({
            "index": idx,
            "name": f.name,
            "size": f.stat().st_size,
            "url": f"/api/transcription/jobs/{job_id}/media?index={idx}"
        })

    return {"files": files}

@router.post("/jobs/{job_id}/preventive-audit/recompute")
async def recompute_preventive_audit(job_id: str):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Job not completed")

    try:
        payload = _load_job_result_payload(job)
        reports = payload.get("reports") or {}
        if not isinstance(reports, dict):
            reports = {}

        raw_content = (payload.get("raw_content") or "").strip()
        formatted_content = (payload.get("content") or "").strip()
        if not raw_content or not formatted_content:
            raise HTTPException(status_code=409, detail="Job sem RAW ou conteúdo formatado para auditoria preventiva.")

        job_dir = (_get_transcription_jobs_dir() / job_id).resolve()
        metadata_path = job_dir / "metadata.json"
        metadata: Dict[str, Any] = {}
        if metadata_path.exists():
            try:
                metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
            except Exception:
                metadata = {}

        fallback_mode = str(metadata.get("mode") or payload.get("mode") or "APOSTILA")
        video_name, mode_suffix = _infer_video_name_and_mode_suffix(reports, fallback_mode=fallback_mode)

        output_dir_value = reports.get("output_dir")
        if output_dir_value:
            output_dir_path = _resolve_job_path(str(output_dir_value)).resolve()
        else:
            # Fallback: persist under job dir to keep it downloadable.
            output_dir_path = job_dir
            output_dir_value = str(output_dir_path)
            reports["output_dir"] = output_dir_value

        output_dir_path.mkdir(parents=True, exist_ok=True)

        preventive_json_key = "preventive_fidelity_json_path"
        preventive_md_key = "preventive_fidelity_md_path"

        json_filename = f"{video_name}_{mode_suffix}_AUDITORIA_FIDELIDADE.json"
        md_filename = f"{video_name}_{mode_suffix}_AUDITORIA_FIDELIDADE.md"

        json_path_value = _join_report_path(str(output_dir_value), json_filename)
        md_path_value = _join_report_path(str(output_dir_value), md_filename)
        json_path = _resolve_job_path(json_path_value).resolve()
        md_path = _resolve_job_path(md_path_value).resolve()

        include_sources = not bool(metadata.get("skip_sources_audit"))
        modo = str(metadata.get("mode") or mode_suffix or "APOSTILA").upper()

        project_root = str(Path(__file__).resolve().parents[5])
        if project_root not in sys.path:
            sys.path.append(project_root)

        from audit_fidelity_preventive import auditar_fidelidade_preventiva, gerar_relatorio_markdown_completo

        try:
            client = _build_genai_client()
            result = auditar_fidelidade_preventiva(
                client,
                raw_content,
                formatted_content,
                video_name,
                output_path=str(json_path),
                modo=modo,
                include_sources=include_sources,
            )
        except Exception as exc:
            palavras_raw = len(raw_content.split()) if raw_content else 0
            palavras_fmt = len(formatted_content.split()) if formatted_content else 0
            taxa_retencao = (palavras_fmt / palavras_raw) if palavras_raw > 0 else 0
            result = {
                "aprovado": True,
                "nota_fidelidade": 0,
                "gravidade_geral": "INFO",
                "erro": str(exc),
                "omissoes_criticas": [],
                "distorcoes": [],
                "problemas_estruturais": [],
                "problemas_contexto": [],
                "alucinacoes": [],
                "metricas": {
                    "palavras_raw": palavras_raw,
                    "palavras_formatado": palavras_fmt,
                    "taxa_retencao": round(taxa_retencao, 4),
                    "dispositivos_legais_raw": 0,
                    "dispositivos_legais_formatado": 0,
                    "taxa_preservacao_dispositivos": 0,
                },
                "observacoes_gerais": f"Auditoria preventiva não pôde ser recalculada: {exc}",
                "recomendacao_hil": {"pausar_para_revisao": False, "motivo": "", "areas_criticas": []},
                "source": "preventive_audit_recompute_fallback",
            }
            try:
                json_path.parent.mkdir(parents=True, exist_ok=True)
                json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
            except Exception:
                pass

        # Garantir persistência do JSON mesmo quando a função retornou sem escrever (defensivo).
        try:
            if not json_path.exists() or json_path.stat().st_size == 0:
                json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
        except Exception:
            pass

        try:
            gerar_relatorio_markdown_completo(result, str(md_path), video_name)
        except Exception as exc:
            # Ainda assim, tente deixar um MD mínimo para o usuário.
            try:
                md_path.write_text(
                    f"# Auditoria Preventiva de Fidelidade\n\nFalha ao gerar relatório markdown: {exc}\n",
                    encoding="utf-8",
                )
            except Exception:
                pass

        reports[preventive_json_key] = json_path_value
        reports[preventive_md_key] = md_path_value

        reports_path_value = payload.get("reports_path")
        if reports_path_value:
            reports_path = _resolve_job_path(str(reports_path_value))
        else:
            reports_path = job_dir / "reports.json"
            # Best effort: update result.json pointer if missing
            try:
                result_path = _find_job_result_path(job)
                result_data = json.loads(result_path.read_text(encoding="utf-8"))
                if not result_data.get("reports_path"):
                    result_data["reports_path"] = str(reports_path)
                    result_path.write_text(json.dumps(result_data, ensure_ascii=False, indent=2), encoding="utf-8")
            except Exception:
                pass

        reports_path.write_text(json.dumps(reports, ensure_ascii=False, indent=2), encoding="utf-8")
        return {"success": True, "reports": reports}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to recompute preventive audit: {exc}")


@router.post("/jobs/{job_id}/quality")
async def update_transcription_job_quality(job_id: str, request: JobQualityUpdateRequest):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Job not completed")
    try:
        result_path = _find_job_result_path(job)
        with open(result_path, "r", encoding="utf-8") as f:
            result_data = json.load(f)

        quality = result_data.get("quality") or {}
        update = request.model_dump(exclude_none=True)
        fixed_content = update.pop("fixed_content", None)
        applied_issue_ids = update.pop("applied_issue_ids", None)
        needs_revalidate = update.get("needs_revalidate")

        if update:
            update["updated_at"] = datetime.utcnow().isoformat()
            quality.update(update)
            result_data["quality"] = quality

        if fixed_content is not None:
            content_path = result_data.get("content_path")
            if not content_path:
                raise HTTPException(status_code=404, detail="Content path not found for job")
            resolved_content = _resolve_job_path(content_path)
            resolved_content.write_text(fixed_content, encoding="utf-8")
            quality["content_updated_at"] = datetime.utcnow().isoformat()
            if needs_revalidate is None:
                quality["needs_revalidate"] = True
            result_data["quality"] = quality

        # Persist applied issues removal (avoid reappearing after reload)
        if applied_issue_ids:
            applied_set = {str(x) for x in applied_issue_ids if x}
            existing_issues: List[Dict[str, Any]] = []
            current = result_data.get("audit_issues")
            if isinstance(current, list):
                existing_issues = [i for i in current if isinstance(i, dict)]
            if not existing_issues:
                audit_path = result_data.get("audit_path")
                if audit_path:
                    audit_candidate = _resolve_job_path(str(audit_path))
                    if audit_candidate.exists():
                        try:
                            with open(audit_candidate, "r", encoding="utf-8") as af:
                                loaded = json.load(af)
                                if isinstance(loaded, list):
                                    existing_issues = [i for i in loaded if isinstance(i, dict)]
                        except Exception:
                            existing_issues = []

            def _should_keep(issue: Dict[str, Any]) -> bool:
                issue_id = str(issue.get("id") or "")
                return not issue_id or issue_id not in applied_set

            remaining_issues = [i for i in existing_issues if _should_keep(i)]
            result_data["audit_issues"] = remaining_issues

            # Ensure we have a job-local audit_issues.json to persist edits
            job_dir = (_get_transcription_jobs_dir() / job_id).resolve()
            audit_path_value = result_data.get("audit_path")
            audit_file = _resolve_job_path(str(audit_path_value)) if audit_path_value else (job_dir / "audit_issues.json")
            try:
                Path(audit_file).write_text(json.dumps(remaining_issues, ensure_ascii=False, indent=2), encoding="utf-8")
                result_data["audit_path"] = str(Path(audit_file).resolve())
            except Exception:
                pass

        # Recompute deterministic analysis_result + auto audit issues after content changes
        if fixed_content is not None:
            try:
                from app.services.quality_service import quality_service

                raw_path = result_data.get("raw_path")
                raw_content = ""
                if raw_path:
                    raw_candidate = _resolve_job_path(str(raw_path))
                    if raw_candidate.exists():
                        raw_content = raw_candidate.read_text(encoding="utf-8", errors="ignore")

                document_name = ((job.get("file_names") or [job_id])[0]) if isinstance(job, dict) else job_id
                analysis_report = await quality_service.analyze_structural_issues(
                    content=fixed_content,
                    document_name=document_name,
                    raw_content=raw_content or None,
                )

                quality = result_data.get("quality") or {}
                quality["analysis_result"] = analysis_report
                if needs_revalidate is None:
                    quality["needs_revalidate"] = True
                result_data["quality"] = quality

                # Best-effort: refresh validation_report so the Quality tab stays consistent after applying fixes.
                if raw_content:
                    try:
                        validation_timeout = int(
                            os.getenv("IUDEX_HIL_VALIDATION_TIMEOUT_SECONDS")
                            or os.getenv("IUDEX_HIL_CONTENT_TIMEOUT_SECONDS", "900")
                        )
                    except Exception:
                        validation_timeout = 900
                    modo = str(result_data.get("mode") or "APOSTILA").upper()
                    try:
                        validation_report = await asyncio.wait_for(
                            quality_service.validate_document(
                                raw_content=raw_content,
                                formatted_content=fixed_content,
                                document_name=document_name,
                                mode=modo,
                            ),
                            timeout=validation_timeout,
                        )
                        quality["validation_report"] = validation_report
                        quality["needs_revalidate"] = False
                        result_data["quality"] = quality
                    except asyncio.TimeoutError:
                        logger.warning("Timeout ao revalidar documento após correções")
                        quality["needs_revalidate"] = True
                        result_data["quality"] = quality
                    except Exception as validation_error:
                        logger.warning(f"Falha ao revalidar documento após correções: {validation_error}")
                        quality["needs_revalidate"] = True
                        result_data["quality"] = quality

                # Replace auto-derived audit issues with freshly computed ones, keeping manual/preventive issues.
                derived = service._build_audit_issues(
                    analysis_report,
                    document_name,
                    raw_content=raw_content or "",
                    formatted_content=fixed_content or "",
                )
                derived_types = {
                    "duplicate_section",
                    "duplicate_paragraph",
                    "heading_numbering",
                    "missing_law",
                    "missing_sumula",
                    "missing_decreto",
                    "missing_julgado",
                    "compression_warning",
                    "legal_audit",
                }

                existing_list = result_data.get("audit_issues")
                if isinstance(existing_list, list):
                    existing_issues = [i for i in existing_list if isinstance(i, dict)]
                else:
                    existing_issues = []
                preserved = [i for i in existing_issues if str(i.get("type") or "") not in derived_types]

                merged: List[Dict[str, Any]] = []
                seen_ids = set()
                seen_keys = set()
                for issue in preserved + (derived or []):
                    if not isinstance(issue, dict):
                        continue
                    issue_id = str(issue.get("id") or "")
                    if issue_id and issue_id in seen_ids:
                        continue
                    key = f"{issue.get('type','')}:{issue.get('reference','')}:{issue.get('description','')}"
                    if key in seen_keys:
                        continue
                    if issue_id:
                        seen_ids.add(issue_id)
                    seen_keys.add(key)
                    merged.append(issue)
                result_data["audit_issues"] = merged

                job_dir = (_get_transcription_jobs_dir() / job_id).resolve()
                audit_file = job_dir / "audit_issues.json"
                try:
                    audit_file.write_text(json.dumps(merged, ensure_ascii=False, indent=2), encoding="utf-8")
                    result_data["audit_path"] = str(audit_file.resolve())
                except Exception:
                    pass
            except Exception as e:
                logger.warning(f"Falha ao recomputar analysis/audit_issues após correção: {e}")

        with open(result_path, "w", encoding="utf-8") as f:
            json.dump(result_data, f, ensure_ascii=False, indent=2)

        return {"success": True, "quality": result_data.get("quality") or {}}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to update quality: {exc}")

@router.post("/jobs/{job_id}/content")
async def update_transcription_job_content(job_id: str, request: JobContentUpdateRequest):
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Job not completed")
    try:
        result_path = _find_job_result_path(job)
        with open(result_path, "r", encoding="utf-8") as f:
            result_data = json.load(f)

        job_dir = (_get_transcription_jobs_dir() / job_id).resolve()
        job_dir.mkdir(parents=True, exist_ok=True)

        job_type = result_data.get("job_type") or "vomo"
        content_changed = False

        if request.content is not None:
            if job_type == "hearing":
                payload_path = result_data.get("payload_path")
                payload: Dict[str, Any] = {}
                if payload_path and os.path.exists(payload_path):
                    try:
                        with open(payload_path, "r", encoding="utf-8") as pf:
                            payload = json.load(pf)
                    except Exception:
                        payload = {}
                payload = payload or {}
                payload["formatted_text"] = request.content
                if payload_path:
                    Path(payload_path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
                formatted_path = result_data.get("formatted_path") or str(job_dir / "hearing_formatted.md")
                Path(formatted_path).write_text(request.content, encoding="utf-8")
                result_data["formatted_path"] = formatted_path
            else:
                content_path = result_data.get("content_path") or str(job_dir / "content.md")
                resolved_content = _resolve_job_path(content_path)
                resolved_content.write_text(request.content, encoding="utf-8")
                result_data["content_path"] = str(resolved_content)
            content_changed = True

        if request.rich_text_html is not None:
            html_path = result_data.get("rich_text_html_path") or str(job_dir / "rich_text.html")
            resolved_html = _resolve_job_path(html_path)
            resolved_html.write_text(request.rich_text_html, encoding="utf-8")
            result_data["rich_text_html_path"] = str(resolved_html)

        if request.rich_text_json is not None:
            json_path = result_data.get("rich_text_json_path") or str(job_dir / "rich_text.json")
            resolved_json = _resolve_job_path(json_path)
            resolved_json.write_text(json.dumps(request.rich_text_json, ensure_ascii=False, indent=2), encoding="utf-8")
            result_data["rich_text_json_path"] = str(resolved_json)

        if request.rich_text_meta is not None:
            meta_path = result_data.get("rich_text_meta_path") or str(job_dir / "rich_text_meta.json")
            resolved_meta = _resolve_job_path(meta_path)
            resolved_meta.write_text(json.dumps(request.rich_text_meta, ensure_ascii=False, indent=2), encoding="utf-8")
            result_data["rich_text_meta_path"] = str(resolved_meta)

        if content_changed:
            quality = result_data.get("quality") or {}
            quality["content_updated_at"] = datetime.utcnow().isoformat()
            if request.needs_revalidate is None:
                quality["needs_revalidate"] = True
            else:
                quality["needs_revalidate"] = bool(request.needs_revalidate)
            result_data["quality"] = quality
        elif request.needs_revalidate is not None:
            quality = result_data.get("quality") or {}
            quality["needs_revalidate"] = bool(request.needs_revalidate)
            result_data["quality"] = quality

        result_path.write_text(json.dumps(result_data, ensure_ascii=False, indent=2), encoding="utf-8")
        return {
            "success": True,
            "content_updated": content_changed,
            "rich_text_html_path": result_data.get("rich_text_html_path"),
            "rich_text_json_path": result_data.get("rich_text_json_path"),
            "rich_text_meta_path": result_data.get("rich_text_meta_path"),
            "quality": result_data.get("quality"),
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to update job content: {exc}")

@router.get("/jobs/{job_id}/stream")
async def stream_transcription_job(job_id: str):
    EventSourceResponse = _get_event_source_response()

    async def event_generator():
        last_snapshot = None
        while True:
            job = job_manager.get_transcription_job(job_id)
            if not job:
                yield {"event": "error", "data": json.dumps({"error": "Job not found"})}
                break

            status = job.get("status")
            if status == "error":
                yield {"event": "error", "data": json.dumps({"error": job.get("error") or "Job failed"})}
                break
            if status == "canceled":
                yield {"event": "error", "data": json.dumps({"error": job.get("message") or "Job canceled"})}
                break
            if status == "completed":
                try:
                    payload = _load_job_result_payload(job)
                    yield {
                        "event": "complete",
                        "data": json.dumps({"status": "success", "job_id": job_id, **payload})
                    }
                except Exception as exc:
                    yield {"event": "error", "data": json.dumps({"error": f"Failed to load result: {exc}"})}
                break

            snapshot = {
                "stage": job.get("stage") or "queued",
                "progress": job.get("progress") or 0,
                "message": job.get("message") or "",
            }
            if snapshot != last_snapshot:
                last_snapshot = snapshot
                yield {"event": "progress", "data": json.dumps(snapshot)}

            await asyncio.sleep(1.5)

    return EventSourceResponse(
        event_generator(),
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Content-Encoding": "identity",
        }
    )

@router.post("/vomo", response_model=dict)
async def transcribe_vomo(
    file: UploadFile = File(...),
    mode: str = Form("APOSTILA"),
    thinking_level: str = Form("medium"),
    custom_prompt: Optional[str] = Form(None),
    model_selection: str = Form("gemini-3-flash-preview"),
    high_accuracy: bool = Form(False),
    diarization: Optional[bool] = Form(None),
    diarization_strict: bool = Form(False),
    use_cache: bool = Form(True),
    auto_apply_fixes: bool = Form(True),
    auto_apply_content_fixes: bool = Form(False),
    skip_legal_audit: bool = Form(False),
    skip_audit: bool = Form(False),
    skip_fidelity_audit: bool = Form(False),
    skip_sources_audit: bool = Form(False),
):
    """
    Endpoint para transcrição e formatação usando MLX Vomo.
    Suporta arquivos de áudio e vídeo.
    Retorna o texto transcrito/formatado.
    
    WARNING: Processamento síncrono/longo por enquanto (MVP).
    Idealmente mover para BackgroundTasks e retornar JobID.
    """
    temp_file_path = f"/tmp/{uuid.uuid4()}_{file.filename}"
    run_id = str(uuid.uuid4())
    
    try:
        if custom_prompt is not None:
            custom_prompt = custom_prompt.strip() or None
        # Salvar arquivo temporário
        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        logger.info(f"📁 Arquivo recebido: {file.filename} ({mode})")
        
        # Processar
        run_id = str(uuid.uuid4())
        with usage_context("transcription", run_id):
            final_text = await service.process_file(
                file_path=temp_file_path,
                mode=mode,
                thinking_level=thinking_level,
                custom_prompt=custom_prompt,
                high_accuracy=high_accuracy,
                diarization=diarization,
                diarization_strict=diarization_strict,
                model_selection=model_selection,
                use_cache=use_cache,
                auto_apply_fixes=auto_apply_fixes,
                auto_apply_content_fixes=auto_apply_content_fixes,
                skip_legal_audit=skip_legal_audit,
                skip_audit=skip_audit,
                skip_fidelity_audit=skip_fidelity_audit,
                skip_sources_audit=skip_sources_audit,
            )
        
        return {
            "status": "success",
            "filename": file.filename,
            "content": final_text
        }
        
    except Exception as e:
        logger.error(f"Erro na transcrição: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
        
    finally:
        # Cleanup
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)


@router.post("/vomo/stream")
async def transcribe_vomo_stream(
    file: UploadFile = File(...),
    mode: str = Form("APOSTILA"),
    thinking_level: str = Form("medium"),
    custom_prompt: Optional[str] = Form(None),
    model_selection: str = Form("gemini-3-flash-preview"),
    high_accuracy: bool = Form(False),
    diarization: Optional[bool] = Form(None),
    diarization_strict: bool = Form(False),
    use_cache: bool = Form(True),
    auto_apply_fixes: bool = Form(True),
    auto_apply_content_fixes: bool = Form(False),
    skip_legal_audit: bool = Form(False),
    skip_audit: bool = Form(False),
    skip_fidelity_audit: bool = Form(False),
    skip_sources_audit: bool = Form(False),
):
    """
    SSE endpoint that streams transcription progress in real-time.
    
    Events:
    - progress: { stage, progress, message }
    - complete: { status, filename, content }
    - error: { error }
    """
    import json
    import asyncio

    EventSourceResponse = _get_event_source_response()
    
    temp_file_path = f"/tmp/{uuid.uuid4()}_{file.filename}"
    run_id = str(uuid.uuid4())
    
    # Save uploaded file first (outside generator)
    try:
        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as exc:
        logger.error(f"Falha ao salvar arquivo temporário: {exc}")
        raise HTTPException(status_code=500, detail="Falha ao salvar arquivo temporário para transcrição.")
    
    logger.info(f"📁 SSE: Arquivo recebido: {file.filename} ({mode})")
    
    async def event_generator():
        progress_queue = asyncio.Queue()
        final_result = {"content": None, "raw_content": None, "reports": None, "error": None}
        
        async def on_progress(stage: str, progress: int, message: str):
            """Callback chamado pelo service para reportar progresso."""
            await progress_queue.put({
                "event": "progress",
                "data": {"stage": stage, "progress": progress, "message": message}
            })
        
        async def process_task():
            """Task que executa o processamento."""
            try:
                with usage_context("transcription", run_id):
                    result = await service.process_file_with_progress(
                        file_path=temp_file_path,
                        mode=mode,
                        thinking_level=thinking_level,
                        custom_prompt=custom_prompt,
                        high_accuracy=high_accuracy,
                        diarization=diarization,
                        diarization_strict=diarization_strict,
                        model_selection=model_selection,
                        on_progress=on_progress,
                        use_cache=use_cache,
                        auto_apply_fixes=auto_apply_fixes,
                        auto_apply_content_fixes=auto_apply_content_fixes,
                        skip_legal_audit=skip_legal_audit,
                        skip_audit=skip_audit,
                        skip_fidelity_audit=skip_fidelity_audit,
                        skip_sources_audit=skip_sources_audit,
                    )
                if isinstance(result, dict):
                    final_result["content"] = result.get("content")
                    final_result["raw_content"] = result.get("raw_content")
                    final_result["reports"] = result.get("reports")
                else:
                    final_result["content"] = result
                    final_result["raw_content"] = result
            except Exception as e:
                final_result["error"] = str(e)
            finally:
                # Signal done
                await progress_queue.put(None)
        
        # Start processing in background
        task = asyncio.create_task(process_task())
        
        try:
            # Yield progress events as they come
            while True:
                item = await progress_queue.get()
                if item is None:
                    break
                yield {
                    "event": item["event"],
                    "data": json.dumps(item["data"])
                }
            
            # Wait for task to complete
            await task
            
            # Yield final result
            if final_result["error"]:
                yield {
                    "event": "error",
                    "data": json.dumps({"error": final_result["error"]})
                }
            else:
                yield {
                    "event": "complete",
                    "data": json.dumps({
                        "status": "success",
                        "filename": file.filename,
                        "content": final_result["content"],
                        "raw_content": final_result.get("raw_content"),
                        "reports": final_result.get("reports")
                    })
                }
        finally:
            # Cleanup temp file
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
    
    return EventSourceResponse(
        event_generator(),
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Content-Encoding": "identity",
        }
    )


@router.post("/vomo/batch/stream")
async def transcribe_batch_stream(
    files: list[UploadFile] = File(...),
    mode: str = Form("APOSTILA"),
    thinking_level: str = Form("medium"),
    custom_prompt: Optional[str] = Form(None),
    model_selection: str = Form("gemini-3-flash-preview"),
    high_accuracy: bool = Form(False),
    diarization: Optional[bool] = Form(None),
    diarization_strict: bool = Form(False),
    use_cache: bool = Form(True),
    auto_apply_fixes: bool = Form(True),
    auto_apply_content_fixes: bool = Form(False),
    skip_legal_audit: bool = Form(False),
    skip_audit: bool = Form(False),
    skip_fidelity_audit: bool = Form(False),
    skip_sources_audit: bool = Form(False),
):
    """
    SSE endpoint for batch transcription of multiple files.
    
    Processes files in order and unifies output into a single document.
    Files are processed sequentially to maintain order (Aula 1, 2, 3...).
    
    Events:
    - progress: { stage, progress, message }
    - complete: { status, filenames, content }
    - error: { error }
    """
    import json
    import asyncio

    EventSourceResponse = _get_event_source_response()
    
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="No files provided")
    
    # Save all uploaded files first
    temp_paths = []
    file_names = []
    
    for f in files:
        path = f"/tmp/{uuid.uuid4()}_{f.filename}"
        try:
            with open(path, "wb") as buffer:
                shutil.copyfileobj(f.file, buffer)
        except Exception as exc:
            logger.error(f"Falha ao salvar arquivo temporário ({f.filename}): {exc}")
            raise HTTPException(status_code=500, detail="Falha ao salvar arquivos temporários para transcrição em lote.")
        temp_paths.append(path)
        file_names.append(f.filename)
    
    logger.info(f"📁 BATCH SSE: {len(files)} arquivos recebidos: {', '.join(file_names)}")
    run_id = str(uuid.uuid4())
    
    async def event_generator():
        progress_queue = asyncio.Queue()
        final_result = {"content": None, "raw_content": None, "reports": None, "error": None}
        
        async def on_progress(stage: str, progress: int, message: str):
            await progress_queue.put({
                "event": "progress",
                "data": {"stage": stage, "progress": progress, "message": message}
            })
        
        async def process_task():
            try:
                with usage_context("transcription", run_id):
                    result = await service.process_batch_with_progress(
                        file_paths=temp_paths,
                        file_names=file_names,
                        mode=mode,
                        thinking_level=thinking_level,
                        custom_prompt=custom_prompt,
                        high_accuracy=high_accuracy,
                        diarization=diarization,
                        diarization_strict=diarization_strict,
                        model_selection=model_selection,
                        on_progress=on_progress,
                        use_cache=use_cache,
                        auto_apply_fixes=auto_apply_fixes,
                        auto_apply_content_fixes=auto_apply_content_fixes,
                        skip_legal_audit=skip_legal_audit,
                        skip_audit=skip_audit,
                        skip_fidelity_audit=skip_fidelity_audit,
                        skip_sources_audit=skip_sources_audit,
                    )
                if isinstance(result, dict):
                    final_result["content"] = result.get("content")
                    final_result["raw_content"] = result.get("raw_content")
                    final_result["reports"] = result.get("reports")
                else:
                    final_result["content"] = result
                    final_result["raw_content"] = result
            except Exception as e:
                final_result["error"] = str(e)
            finally:
                await progress_queue.put(None)
        
        task = asyncio.create_task(process_task())
        
        try:
            while True:
                item = await progress_queue.get()
                if item is None:
                    break
                yield {
                    "event": item["event"],
                    "data": json.dumps(item["data"])
                }
            
            await task
            
            if final_result["error"]:
                yield {
                    "event": "error",
                    "data": json.dumps({"error": final_result["error"]})
                }
            else:
                yield {
                    "event": "complete",
                    "data": json.dumps({
                        "status": "success",
                        "filenames": file_names,
                        "total_files": len(file_names),
                        "content": final_result["content"],
                        "raw_content": final_result.get("raw_content"),
                        "reports": final_result.get("reports")
                    })
                }
        finally:
            # Cleanup all temp files
            for path in temp_paths:
                if os.path.exists(path):
                    os.remove(path)
    
    return EventSourceResponse(
        event_generator(),
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Content-Encoding": "identity",
        }
    )


@router.post("/apply-revisions")
async def apply_revisions(request: Request):
    """
    Apply revisions to transcription based on user-approved issues.
    
    Applies both:
    - Structural fixes (duplicate removal, heading renumbering) - local operations
    - Content fixes (missing laws, omissions) - via LLM (Gemini or OpenAI)
    
    Uses raw Request to avoid Pydantic validation conflicts.
    """
    payload_summary = None
    try:
        # Parse JSON manually to bypass all Pydantic validation
        try:
            payload = await request.json()
        except Exception as parse_error:
            logger.error(f"Erro ao parsear JSON: {parse_error}")
            raise HTTPException(status_code=400, detail="JSON inválido no corpo da requisição")
        
        # Manual extraction - no validation
        job_id = payload.get("job_id") or payload.get("jobId") or payload.get("jobID")
        content = payload.get("content") or ""
        raw_content = payload.get("raw_content") or ""
        approved_issues = payload.get("approved_issues") or []
        model_selection = payload.get("model_selection") or "gemini-2.0-flash"
        mode = payload.get("mode") or payload.get("format_mode") or payload.get("formatMode")

        if (not content or not raw_content) and job_id:
            try:
                job = job_manager.get_transcription_job(str(job_id))
            except Exception:
                job = None
            if not job:
                raise HTTPException(status_code=404, detail="Job not found")
            if job.get("status") != "completed":
                raise HTTPException(status_code=409, detail="Job not completed")
            job_payload = _load_job_result_payload(job)
            if job_payload.get("job_type") == "hearing":
                raise HTTPException(status_code=409, detail="Hearing jobs are not supported for apply-revisions")
            if not content:
                content = job_payload.get("content") or ""
            if not raw_content:
                raw_content = job_payload.get("raw_content") or ""
            if not mode:
                mode = job_payload.get("mode")

        if not approved_issues:
            return {"revised_content": content, "changes_made": 0}
            
        from app.services.quality_service import quality_service

        structural_issues = [i for i in approved_issues if i.get("fix_type") == "structural"]
        content_issues = [i for i in approved_issues if i.get("fix_type") != "structural"]

        payload_summary = {
            "job_id": str(job_id) if job_id else None,
            "content_len": len(content) if content else 0,
            "raw_content_len": len(raw_content) if raw_content else 0,
            "approved_issues": len(approved_issues),
            "structural_issues": len(structural_issues),
            "content_issues": len(content_issues),
            "model_selection": model_selection,
            "mode": mode,
        }
        logger.info(f"HIL apply-revisions payload summary: {payload_summary}")

        revised_content = content
        applied_structural = []
        structural_error_msg = None

        # 1. Apply structural fixes (local, no LLM needed)
        if structural_issues:
            try:
                structural_timeout = int(os.getenv("IUDEX_HIL_STRUCTURAL_TIMEOUT_SECONDS", "60"))
                structural_result = await asyncio.wait_for(
                    quality_service.apply_structural_fixes_from_issues(
                        content=revised_content,
                        approved_issues=structural_issues,
                    ),
                    timeout=structural_timeout,
                )
                revised_content = structural_result.get("content", revised_content)
                applied_structural = structural_result.get("fixes", []) or []
                if structural_result.get("error") and not structural_error_msg:
                    structural_error_msg = structural_result.get("error")
            except asyncio.TimeoutError:
                logger.error("Timeout ao aplicar correções estruturais (HIL)")
                structural_error_msg = "Timeout ao aplicar correções estruturais."
            except Exception as structural_error:
                logger.error(f"Erro ao aplicar correções estruturais: {structural_error}")
                structural_error_msg = str(structural_error)

        # 2. Apply content fixes via LLM (Gemini or OpenAI based on model_selection)
        applied_content = []
        content_error_msg = None
        skipped_issue_ids: list[str] = []
        skipped_reason: str | None = None
        
        if content_issues and not raw_content:
            skipped_issue_ids = [issue.get("id") for issue in content_issues if issue.get("id")]
            skipped_reason = "raw_content ausente (necessário para correções de conteúdo)"
            content_error_msg = "raw_content não fornecido para correções de conteúdo"
        if content_issues and raw_content:
            try:
                # Content fixes may involve large documents + LLM latency; keep aligned with the web proxy long timeout.
                # Env is in seconds.
                content_timeout = int(os.getenv("IUDEX_HIL_CONTENT_TIMEOUT_SECONDS", "900"))
                content_result = await asyncio.wait_for(
                    quality_service.fix_content_issues_with_llm(
                        content=revised_content,
                        raw_content=raw_content or "",
                        issues=content_issues,
                        model_selection=model_selection,
                        mode=mode,
                    ),
                    timeout=content_timeout,
                )
                if content_result.get("fixes"):
                    revised_content = content_result.get("content", revised_content)
                    applied_content = content_result.get("fixes", [])
                if content_result.get("error"):
                    content_error_msg = content_result.get("error")
            except asyncio.TimeoutError:
                logger.error("Timeout ao aplicar correções de conteúdo via LLM (HIL)")
                content_error_msg = "Timeout ao aplicar correções de conteúdo via IA."
            except Exception as content_error:
                logger.error(f"Erro ao aplicar correções de conteúdo via LLM: {content_error}")
                content_error_msg = str(content_error)

        def _normalize_title(value: str) -> str:
            cleaned = re.sub(r"^#+\s*\d*\.?\s*", "", value or "").strip().lower()
            return cleaned

        def _extract_structural_applied_ids(
            issues: List[Dict[str, Any]],
            applied_fixes: List[Any],
        ) -> List[str]:
            if not issues or not applied_fixes:
                return []

            fix_strings = [str(fix) for fix in applied_fixes]
            applied_fps = set()
            applied_titles = set()
            has_heading_fix = any("renumbered h2 headings" in fix.lower() for fix in fix_strings)

            for fix in fix_strings:
                fp_match = re.search(r"fingerprint:\s*([0-9a-f]+)", fix, re.IGNORECASE)
                if fp_match:
                    applied_fps.add(fp_match.group(1))
                if "removed duplicate section" in fix.lower():
                    title = fix.split(":", 1)[-1].strip()
                    if title:
                        applied_titles.add(_normalize_title(title))

            applied_ids: List[str] = []
            for issue in issues:
                issue_id = issue.get("id")
                if not issue_id:
                    continue
                issue_type = issue.get("type")
                if issue_type == "duplicate_paragraph":
                    fingerprint = issue.get("fingerprint")
                    if fingerprint and fingerprint in applied_fps:
                        applied_ids.append(issue_id)
                elif issue_type == "duplicate_section":
                    title = issue.get("title")
                    if title and _normalize_title(title) in applied_titles:
                        applied_ids.append(issue_id)
                elif issue_type == "heading_numbering":
                    if has_heading_fix:
                        applied_ids.append(issue_id)
            return applied_ids

        applied_issue_ids = []
        applied_issue_ids.extend(_extract_structural_applied_ids(structural_issues, applied_structural))
        applied_issue_ids.extend([issue_id for issue_id in applied_content if issue_id])

        # Calculate total changes
        all_applied = applied_structural + applied_content
        changes_made = len(all_applied)
        
        # CRITICAL: Never return empty content - fallback to original
        final_content = revised_content
        if not final_content or not final_content.strip():
            logger.warning("HIL apply-revisions: revised_content is empty, falling back to original")
            final_content = content
            if not structural_error_msg:
                structural_error_msg = "Conteúdo revisado ficou vazio - original preservado"

        # Detectar se houve mudança real no conteúdo
        try:
            before_hash = hashlib.md5((content or "").encode("utf-8")).hexdigest()
            after_hash = hashlib.md5((final_content or "").encode("utf-8")).hexdigest()
        except Exception:
            before_hash = None
            after_hash = None
        content_changed = bool(before_hash and after_hash and before_hash != after_hash)
        content_change = {
            "before_chars": len(content or ""),
            "after_chars": len(final_content or ""),
            "delta_chars": len((final_content or "")) - len((content or "")),
        }
        
        # Sanity check: if content is drastically smaller, warn but still return it
        if len(final_content) < len(content) * 0.3 and len(content) > 100:
            logger.warning(f"HIL apply-revisions: content shrank significantly ({len(content)} -> {len(final_content)})")
        
        return {
            "revised_content": final_content,
            "changes_made": changes_made,
            "issues_applied": applied_issue_ids,
            "applied_issue_ids": applied_issue_ids,
            "structural_fixes_applied": applied_structural,
            "content_fixes_applied": applied_content,
            "structural_error": structural_error_msg,
            "content_error": content_error_msg,
            "skipped_issue_ids": skipped_issue_ids,
            "skipped_reason": skipped_reason,
            "content_changed": content_changed,
            "content_change": content_change,
            "model_used": model_selection if content_issues else None,
        }
        
    except HTTPException:
        raise
    except Exception as e:
        if payload_summary:
            logger.error(f"Erro ao aplicar revisões HIL (summary={payload_summary}): {e}", exc_info=True)
        else:
            logger.error(f"Erro ao aplicar revisões HIL: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Falha na revisão: {str(e)}")


class HearingApplyRevisionsRequest(BaseModel):
    approved_issues: List[Dict[str, Any]] = []
    model_selection: Optional[str] = None
    regenerate_formatted: bool = True


@router.post("/jobs/{job_id}/hearing/apply-revisions")
async def apply_hearing_revisions(job_id: str, request: HearingApplyRevisionsRequest):
    """
    Apply AI-assisted revisions to a hearing/meeting job snapshot.

    - Patches affected segments (preserving ordering/timestamps)
    - Re-renders transcript_markdown from segments
    - Optionally patches formatted_text using the transcript_markdown as evidence
    - Persists updated hearing_payload.json (+ hearing_transcript.md / hearing_formatted.md)
    """
    job = job_manager.get_transcription_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Job not completed")
    if job.get("job_type") != "hearing":
        raise HTTPException(status_code=409, detail="Only hearing jobs are supported for hearing apply-revisions")

    job_dir = (_get_transcription_jobs_dir() / job_id).resolve()
    result_path = (job_dir / "result.json").resolve()
    if not result_path.exists():
        raise HTTPException(status_code=500, detail="Job result.json missing")

    payload = _load_job_result_payload(job)
    hearing_payload = payload.get("payload") if isinstance(payload, dict) else None
    if not isinstance(hearing_payload, dict):
        raise HTTPException(status_code=500, detail="Hearing payload missing")

    approved = [i for i in (request.approved_issues or []) if isinstance(i, dict)]
    if not approved:
        return {"success": True, "changes_made": 0, "issues_applied": [], "payload": hearing_payload}

    # Read config for defaults
    metadata_path = job_dir / "metadata.json"
    metadata: Dict[str, Any] = {}
    if metadata_path.exists():
        try:
            metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
        except Exception:
            metadata = {}

    mode = str(metadata.get("format_mode") or hearing_payload.get("formatted_mode") or "AUDIENCIA").upper()
    model_selection = (request.model_selection or metadata.get("model_selection") or "gemini-3-flash-preview")

    from app.services.quality_service import quality_service

    # 1) Patch segments
    segments = hearing_payload.get("segments") or []
    speakers = hearing_payload.get("speakers") or []
    seg_result = await quality_service.fix_hearing_segments_with_llm(
        segments=segments if isinstance(segments, list) else [],
        speakers=speakers if isinstance(speakers, list) else [],
        issues=approved,
        model_selection=model_selection,
        mode=mode,
    )
    updated_segments = seg_result.get("segments") if isinstance(seg_result, dict) else None
    if isinstance(updated_segments, list) and updated_segments:
        hearing_payload["segments"] = updated_segments

    # 2) Re-render transcript_markdown
    try:
        transcript_markdown = service._render_hearing_markdown(hearing_payload)
        hearing_payload["transcript_markdown"] = transcript_markdown
    except Exception:
        transcript_markdown = hearing_payload.get("transcript_markdown") or ""

    # 3) Patch formatted_text using transcript as evidence (optional)
    applied_content_ids: List[str] = []
    content_error: Optional[str] = None
    if request.regenerate_formatted and isinstance(hearing_payload.get("formatted_text"), str) and transcript_markdown:
        mapped_issues: List[Dict[str, Any]] = []
        for issue in approved:
            mapped = dict(issue)
            mapped.setdefault("fix_type", "content")
            mapped.setdefault("source", "hearing_validation")
            mapped.setdefault("raw_evidence", [{"snippet": transcript_markdown[:4000]}])
            mapped_issues.append(mapped)
        try:
            content_result = await quality_service.fix_content_issues_with_llm(
                content=hearing_payload.get("formatted_text") or "",
                raw_content=transcript_markdown,
                issues=mapped_issues,
                model_selection=model_selection,
                mode=mode,
            )
            if isinstance(content_result, dict):
                if content_result.get("content"):
                    hearing_payload["formatted_text"] = content_result.get("content")
                applied_content_ids = [x for x in (content_result.get("fixes") or []) if x]
                content_error = content_result.get("error")
        except Exception as e:
            content_error = str(e)

    # 4) Persist updated payload (with backup)
    payload_path_value = None
    try:
        result_data = json.loads(result_path.read_text(encoding="utf-8"))
        payload_path_value = result_data.get("payload_path")
    except Exception:
        payload_path_value = None

    payload_path = Path(payload_path_value) if payload_path_value else (job_dir / "hearing_payload.json")
    payload_path = payload_path.resolve()

    try:
        backup_path = job_dir / f"hearing_payload.backup_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json"
        if payload_path.exists():
            backup_path.write_text(payload_path.read_text(encoding="utf-8", errors="ignore"), encoding="utf-8")
    except Exception:
        pass

    payload_path.write_text(json.dumps(hearing_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    transcript_path = job_dir / "hearing_transcript.md"
    if hearing_payload.get("transcript_markdown"):
        transcript_path.write_text(hearing_payload.get("transcript_markdown") or "", encoding="utf-8")

    formatted_path = job_dir / "hearing_formatted.md"
    if hearing_payload.get("formatted_text"):
        formatted_path.write_text(hearing_payload.get("formatted_text") or "", encoding="utf-8")

    # Keep result.json pointers fresh
    try:
        result_data = json.loads(result_path.read_text(encoding="utf-8"))
        result_data["payload_path"] = str(payload_path)
        result_data["transcript_path"] = str(transcript_path) if transcript_path.exists() else None
        result_data["formatted_path"] = str(formatted_path) if formatted_path.exists() else None
        result_data["updated_at"] = datetime.utcnow().isoformat()
        result_path.write_text(json.dumps(result_data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass

    applied_ids = [x for x in (seg_result.get("fixes") if isinstance(seg_result, dict) else []) or [] if x]
    applied_ids.extend(applied_content_ids)
    return {
        "success": True,
        "changes_made": len(set(applied_ids)),
        "issues_applied": list(dict.fromkeys(applied_ids)),
        "segment_error": (seg_result.get("error") if isinstance(seg_result, dict) else None),
        "content_error": content_error,
        "model_used": model_selection,
        "mode": mode,
        "payload": hearing_payload,
    }


@router.post("/hearing/stream")
async def transcribe_hearing_stream(
    file: UploadFile = File(...),
    case_id: str = Form(...),
    goal: str = Form("alegacoes_finais"),
    thinking_level: str = Form("medium"),
    model_selection: str = Form("gemini-3-flash-preview"),
    high_accuracy: bool = Form(False),
    format_mode: str = Form("AUDIENCIA"),
    custom_prompt: Optional[str] = Form(None),
    format_enabled: bool = Form(True),
    include_timestamps: bool = Form(True),
    allow_indirect: bool = Form(False),
    allow_summary: bool = Form(False),
    use_cache: bool = Form(True),
    auto_apply_fixes: bool = Form(True),
    auto_apply_content_fixes: bool = Form(False),
    skip_legal_audit: bool = Form(False),
    skip_fidelity_audit: bool = Form(False),
    skip_sources_audit: bool = Form(False),
):
    """
    SSE endpoint for hearings/reunions transcription (structured JSON + evidence).
    """
    import json
    import asyncio

    EventSourceResponse = _get_event_source_response()

    temp_file_path = f"/tmp/{uuid.uuid4()}_{file.filename}"

    try:
        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as exc:
        logger.error(f"Falha ao salvar arquivo temporário (hearing): {exc}")
        raise HTTPException(status_code=500, detail="Falha ao salvar arquivo temporário da audiência.")

    logger.info(f"📁 HEARING SSE: Arquivo recebido: {file.filename} (case_id={case_id})")
    if custom_prompt is not None:
        custom_prompt = custom_prompt.strip() or None

    async def event_generator():
        progress_queue = asyncio.Queue()
        final_result = {"payload": None, "error": None}

        async def on_progress(stage: str, progress: int, message: str):
            await progress_queue.put({
                "event": "progress",
                "data": {"stage": stage, "progress": progress, "message": message}
            })

        async def process_task():
            try:
                with usage_context("transcription", run_id):
                    result = await service.process_hearing_with_progress(
                        file_path=temp_file_path,
                        case_id=case_id,
                        goal=goal,
                        thinking_level=thinking_level,
                        model_selection=model_selection,
                        high_accuracy=high_accuracy,
                        format_mode=format_mode,
                        custom_prompt=custom_prompt,
                        format_enabled=format_enabled,
                        include_timestamps=include_timestamps,
                        allow_indirect=allow_indirect,
                        allow_summary=allow_summary,
                        use_cache=use_cache,
                        auto_apply_fixes=auto_apply_fixes,
                        auto_apply_content_fixes=auto_apply_content_fixes,
                        skip_legal_audit=skip_legal_audit,
                        skip_fidelity_audit=skip_fidelity_audit,
                        skip_sources_audit=skip_sources_audit,
                        on_progress=on_progress,
                    )
                final_result["payload"] = result
            except Exception as e:
                final_result["error"] = str(e)
            finally:
                await progress_queue.put(None)

        task = asyncio.create_task(process_task())

        try:
            while True:
                item = await progress_queue.get()
                if item is None:
                    break
                yield {
                    "event": item["event"],
                    "data": json.dumps(item["data"])
                }

            await task

            if final_result["error"]:
                yield {
                    "event": "error",
                    "data": json.dumps({"error": final_result["error"]})
                }
            else:
                yield {
                    "event": "complete",
                    "data": json.dumps({
                        "status": "success",
                        "filename": file.filename,
                        "payload": final_result["payload"],
                    })
                }
        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)

    return EventSourceResponse(
        event_generator(),
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Content-Encoding": "identity",
        }
    )


@router.post("/hearing/speakers")
async def update_hearing_speakers(request: HearingSpeakersUpdateRequest):
    """
    Update hearing speaker registry (manual edits).
    """
    if not request.speakers:
        raise HTTPException(status_code=400, detail="Nenhum falante informado")
    speakers = service.update_hearing_speakers(request.case_id, [s.model_dump() for s in request.speakers])
    return {"status": "success", "speakers": speakers}


@router.post("/hearing/enroll")
async def enroll_hearing_speaker(
    file: UploadFile = File(...),
    case_id: str = Form(...),
    name: str = Form(...),
    role: str = Form("outro"),
):
    """
    Enroll speaker audio for a case (voice profile seed).
    """
    temp_file_path = f"/tmp/{uuid.uuid4()}_{file.filename}"
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    try:
        speaker = service.enroll_hearing_speaker(case_id=case_id, name=name, role=role, file_path=temp_file_path)
        return {"status": "success", "speaker": speaker}
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
