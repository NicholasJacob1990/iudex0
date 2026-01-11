#!/usr/bin/env python3
"""
Script para processar transcrições de aulas usando Gemini via OpenRouter
Formata conforme diretrizes específicas para concursos de procuradorias
"""

import os
import sys
import time
import re
import requests
from pathlib import Path
from docx import Document
from typing import List, Optional

# Configuração
OPENROUTER_API_KEY = "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"
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL_NAME = "google/gemini-2.0-flash-001"
MAX_OUTPUT_TOKENS = 65_536
MAX_TOKENS_PER_CHUNK = 800_000  # Deixando margem de segurança
OUTPUT_DIR = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ')

PROMPT_FORMATACAO = """Você é um especialista em revisão de transcrições para concursos de procuradorias. Sua tarefa é revisar a transcrição fornecida de uma aula, corrigindo erros de português, erros gramaticais e de pontuação, melhorando a formatação para facilitar a leitura, e mantendo o conteúdo original. Não resuma, não parafraseie e não adicione informações que não estejam na transcrição original. Siga estas diretrizes:

- Mantenha o modo em primeira pessoa
- Corrija erros gramaticais, ortográficos e de pontuação, tornando o texto gramaticalmente correto e claro
- Mantenha todo o conteúdo original, incluindo ideias, exemplos, explicações, pausas, hesitações e ideias incompletas, fazendo o uso apropriado de aspas, parênteses. Não resuma, não omita informações nem altere o significado
- Melhore a formatação para facilitar a leitura
- Mantenha todo o conteúdo original, mas corrija erros da linguagem coloquial para torná-la mais clara, didática e legível
- Ajuste a linguagem coloquial para um português padrão, mantendo o significado original
- Elimine vícios da oralidade e gírias
- Preserve a sequência exata das falas e ideias apresentadas
- Utilize formatação e estrutura com parágrafos bem definidos, facilitando a leitura e compreensão, para melhorar a legibilidade, seguindo o fluxo natural do discurso. Evite parágrafos longos
- Reproduza fielmente as informações, apenas melhorando a clareza e a legibilidade
- Utilize conectivos necessários para tornar o texto mais fluido. Aplique a pontuação devida para deixar o texto coeso e coerente
- Corrija vícios de linguagem, como repetições desnecessárias, uso excessivo de advérbios, linguagem vaga ou imprecisa, gírias, expressões redundantes, e outros erros que afetem a clareza e a eficácia da comunicação, sem alterar o significado do texto
- Identifique e rotule os diferentes falantes, se existentes, organizando suas falas de forma clara
- Divida a aula em tópicos e subtópicos para melhor organização e visualização do conteúdo
- Enumere os tópicos e subtópicos, use negrito quando mais apropriado
- Intercale parágrafos curtos e longos conforme a ideia neles contida, tornando a leitura menos cansativa
- Seja didático sem perder detalhes e conteúdo
- USE TEXTO CORRIDO NA MEDIDA DO POSSÍVEL
- **Ao final de cada tópico/capítulo, sintetize/resume o assunto de forma esquematizada, preferencialmente por tabelas

**Ao final de CADA tópico principal, crie uma tabela de síntese (EXEMPLIFICATIVO, PODENDO SER COMPOSTO POR OUTROS ELEMENTOS NAS COLUNAS):**

| Conceito/Instituto | Definição | Fundamento Legal | Observações |
|-------------------|-----------|-----------------|-------------|
| [Preencher] | [Resumo] | [Art. X, Lei Y] | [Dicas/Exceções] |

Por favor, forneça a versão revisada da transcrição, seguindo estritamente as diretrizes acima. Lembre-se: o objetivo é manter o conteúdo fiel ao original, melhorando apenas a clareza e legibilidade.

<transcrição>
{texto}
</transcrição>
"""


def extrair_texto_txt(caminho_arquivo: str) -> str:
    """Extrai texto de um arquivo .txt"""
    with open(caminho_arquivo, 'r', encoding='utf-8') as f:
        return f.read()


def extrair_texto_docx(caminho_arquivo: str) -> str:
    """Extrai texto de um arquivo .docx"""
    try:
        doc = Document(caminho_arquivo)
        texto_completo = []
        
        for paragrafo in doc.paragraphs:
            if paragrafo.text.strip():
                texto_completo.append(paragrafo.text)
        
        return '\n\n'.join(texto_completo)
    except Exception as e:
        raise Exception(f"Erro ao ler arquivo DOCX: {e}")


def dividir_texto_em_chunks(texto: str, max_chars: int = 3_200_000) -> List[str]:
    """
    Divide o texto em chunks para não exceder o limite de tokens do modelo.
    Tenta dividir em pontos naturais (parágrafos).
    Com ~4 chars por token, 800k tokens = ~3.2M chars
    """
    paragrafos = texto.split('\n\n')
    chunks = []
    chunk_atual = []
    tamanho_atual = 0
    
    for paragrafo in paragrafos:
        tamanho_paragrafo = len(paragrafo)
        
        if tamanho_atual + tamanho_paragrafo > max_chars and chunk_atual:
            chunks.append('\n\n'.join(chunk_atual))
            chunk_atual = [paragrafo]
            tamanho_atual = tamanho_paragrafo
        else:
            chunk_atual.append(paragrafo)
            tamanho_atual += tamanho_paragrafo
    
    if chunk_atual:
        chunks.append('\n\n'.join(chunk_atual))
    
    return chunks


def processar_com_openrouter(texto: str, parte_num: int = 1, total_partes: int = 1) -> str:
    """Processa um chunk de texto com Gemini via OpenRouter"""
    
    # Preparar prompt
    if total_partes > 1:
        contexto_adicional = f"\n\n**IMPORTANTE: Esta é a PARTE {parte_num} de {total_partes} do documento. Mantenha a numeração de tópicos contínua e consistente.**\n\n"
    else:
        contexto_adicional = ""
    
    prompt_completo = PROMPT_FORMATACAO.format(texto=texto) + contexto_adicional
    
    print(f"\n{'='*60}")
    print(f"Processando parte {parte_num}/{total_partes}")
    print(f"Tamanho do texto: {len(texto):,} caracteres")
    print(f"{'='*60}\n")
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://iudex.app",
        "X-Title": "Iudex Transcription Formatter"
    }
    
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {
                "role": "user",
                "content": prompt_completo
            }
        ],
        "temperature": 0.1,
        "max_tokens": MAX_OUTPUT_TOKENS,
        "top_p": 0.95
    }
    
    try:
        response = requests.post(
            OPENROUTER_BASE_URL,
            headers=headers,
            json=payload,
            timeout=600  # 10 minutos de timeout
        )
        response.raise_for_status()
        
        result = response.json()
        return result['choices'][0]['message']['content']
        
    except requests.exceptions.RequestException as e:
        print(f"⚠️  Erro na requisição: {e}")
        if hasattr(e, 'response') and e.response:
            print(f"   Resposta: {e.response.text}")
        raise


def salvar_resultado(texto: str, nome_arquivo: str) -> Path:
    """Salva o resultado em um arquivo .docx"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    
    for paragrafo in texto.split('\n'):
        if paragrafo.strip():
            p = doc.add_paragraph(paragrafo)
    
    caminho_saida = OUTPUT_DIR / nome_arquivo
    doc.save(str(caminho_saida))
    
    return caminho_saida


def salvar_resultado_markdown(texto: str, nome_arquivo: str) -> Path:
    """Salva o resultado em um arquivo .md"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    caminho_saida = OUTPUT_DIR / nome_arquivo
    with open(caminho_saida, 'w', encoding='utf-8') as f:
        f.write(texto)
    
    return caminho_saida


def main():
    """Função principal"""
    
    # Caminho do arquivo de entrada
    if len(sys.argv) > 1:
        arquivo_entrada = Path(sys.argv[1])
    else:
        arquivo_entrada = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/Administrativo.txt')
    
    if not arquivo_entrada.exists():
        print(f"❌ Arquivo não encontrado: {arquivo_entrada}")
        return
    
    print("📄 Extraindo texto do documento...")
    
    if arquivo_entrada.suffix == '.docx':
        texto_original = extrair_texto_docx(str(arquivo_entrada))
    else:
        texto_original = extrair_texto_txt(str(arquivo_entrada))
    
    print(f"✅ Texto extraído: {len(texto_original):,} caracteres")
    print(f"   Estimativa: ~{len(texto_original)//4:,} tokens")
    
    # Dividir em chunks se necessário
    chunks = dividir_texto_em_chunks(texto_original)
    print(f"\n📊 Documento dividido em {len(chunks)} parte(s)")
    
    # Processar cada chunk
    resultados = []
    
    for i, chunk in enumerate(chunks, 1):
        print(f"\n🔄 Processando parte {i}/{len(chunks)}...")
        
        try:
            resultado = processar_com_openrouter(
                texto=chunk,
                parte_num=i,
                total_partes=len(chunks)
            )
            resultados.append(resultado)
            
            print(f"✅ Parte {i} processada com sucesso")
            
            if i < len(chunks):
                print("⏳ Aguardando 2 segundos...")
                time.sleep(2)
                
        except Exception as e:
            print(f"❌ Erro ao processar parte {i}: {e}")
            print(f"   Continuando com as partes restantes...")
            continue
    
    # Combinar resultados
    texto_final = '\n\n---\n\n'.join(resultados)
    
    # Salvar resultados
    print("\n💾 Salvando resultados...")
    
    nome_base = arquivo_entrada.stem
    
    # Salvar em Markdown
    arquivo_md = salvar_resultado_markdown(
        texto_final,
        f'{nome_base}_GEMINI_Formatado.md'
    )
    print(f"✅ Markdown salvo: {arquivo_md}")
    
    # Salvar em DOCX
    try:
        arquivo_docx = salvar_resultado(
            texto_final,
            f'{nome_base}_GEMINI_Formatado.docx'
        )
        print(f"✅ DOCX salvo: {arquivo_docx}")
    except Exception as e:
        print(f"⚠️  Erro ao salvar DOCX: {e}")
    
    print("\n" + "="*60)
    print("🎉 PROCESSAMENTO CONCLUÍDO!")
    print(f"   Total de partes processadas: {len(resultados)}/{len(chunks)}")
    print(f"   Tamanho final: {len(texto_final):,} caracteres")
    print("="*60)


if __name__ == '__main__':
    main()
