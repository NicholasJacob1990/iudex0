#!/usr/bin/env python3
"""
Script v2.7 - Formatação de Transcrições com Gemini 2.5 Flash
CORREÇÃO CRÍTICA: Anti-duplicação de seções

Mudanças v2.7 vs v2.6:
- Detecção agressiva de duplicações de seções inteiras
- Instruções de prompt muito mais rígidas contra reprocessamento
- Validação de sobreposição de chunks (modo debug)
- Delimitadores mais visíveis para o contexto
- Post-processing em múltiplas passadas

Uso: python format_transcription_gemini.py <entrada.txt> [saida]
"""

import os
import sys
import time
import random
import json
import re
import threading
from pathlib import Path
from time import sleep
from difflib import SequenceMatcher
import hashlib

try:
    from google import genai
    from google.genai import types
except ImportError:
    print("❌ Erro: Biblioteca google-genai não instalada.")
    print("   Instale com: pip install google-genai")
    sys.exit(1)

try:
    from tqdm import tqdm
except ImportError:
    print("⚠️ Aviso: tqdm não instalado. Progress bar desabilitada.")
    tqdm = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ Aviso: python-docx não disponível. Saída em Word desabilitada.")
    DOCX_AVAILABLE = False

import logging

# =============================================================================
# CONFIGURAÇÕES v2.7
# =============================================================================

CHARS_POR_PARTE = 10000
CONTEXTO_ESTILO = 750
OUTPUT_TOKEN_LIMIT = 12000
CACHE_TTL = '7200s'
MIN_CHARS_PARA_CACHE = 16000
MAX_RETRIES = 5
MAX_RPM = 3

# v2.7: FORÇAR delimitadores visíveis para evitar confusão
USE_FANCY_DELIMITERS = True

# Preços API Gemini 2.5 Flash (Dezembro 2025)
PRECO_INPUT_SEM_CACHE = 0.30
PRECO_INPUT_COM_CACHE = 0.03
PRECO_OUTPUT = 2.50

# =============================================================================
# LOGGING
# =============================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S',
    handlers=[
        logging.FileHandler('formatacao.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# =============================================================================
# PROMPTS v2.7 - INSTRUÇÕES ANTI-DUPLICAÇÃO REFORÇADAS
# =============================================================================

PROMPT_FIDELIDADE = """# DIRETRIZES DE FORMATAÇÃO E REVISÃO

## PAPEL
VOCÊ É UM EXCELENTISSIMO REDATOR JURÍDICO E DIDÁTICO


## OBJETIVO
Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, MANTENDO A FIDELIDADE TOTAL ao conteúdo original.

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias e das falas do professor.
4. **NÃO CRIE MUITOS BULLET POINTS** ou frases curtas demasiadamente. PREFIRA UM FORMATO DE MANUAL DIDÁTICO, não checklist.
5. **NÃO USE NEGRITOS EM EXCESSO**. Use apenas para conceitos-chave realmente importantes.
6. **NÃO PARAFRASEIE**. Mantenha todo o conteúdo original, incluindo ideias, exemplos, explicações, pausas e hesitações.


## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados (REDI/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência, autores, casos citados
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque)
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico"

## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Corrija erros gramaticais, regências, ortográficos e de pontuação, tornando o texto gramaticalmente correto e claro.
2. **Limpeza**: Elimine vícios de oralidade, gírias, cacoetes ("né", "tipo assim", "então"), repetições desnecessárias, uso excessivo de advérbios, linguagem vaga ou imprecisa.
3. **Coesão**: Utilize conectivos necessários para tornar o texto mais fluido. Aplique a pontuação devida para deixar o texto coeso e coerente.
4. **Legibilidade**:
   - **USE TEXTO CORRIDO NA MEDIDA DO POSSÍVEL**
   - Utilize formatação e estrutura com parágrafos bem definidos, facilitando a leitura e compreensão
   - Evite parágrafos longos (máximo 3-4 linhas visuais)
   - Evite blocos de texto maciços
   - Seja didático sem perder detalhes e conteúdo
5. **Linguagem**: Ajuste a linguagem coloquial para um português padrão, mantendo o significado original.
6. **Citações**: Use itálico para citações curtas e recuo em itálico para citações longas.
7. -Use **negrito** para destacar conceitos-chave (sem exagero).

## 📝 ESTRUTURA
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis a partir do contexto.

## 📊 TABELAS DE RESUMO
Ao final de cada **tópico principal** (seção com título ##), crie uma tabela de resumo no formato:

```
### 📋 Resumo do Tópico

| Conceito/Ponto | Explicação Resumida |
|----------------|---------------------|
| **Conceito 1** | Definição ou ideia principal em 1-2 frases |
| **Conceito 2** | Definição ou ideia principal em 1-2 frases |
```

**Critérios para incluir tabela:**
- Use apenas ao final de seções substantivas (com conteúdo técnico/conceitual)
- Não use em seções introdutórias ou de transição curtas
- Limite a 5-7 linhas por tabela

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA
- **NUNCA formate novamente esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- **NUNCA repita informações que já estão no contexto**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>
- Se o texto_para_formatar começar com algo similar ao fim do contexto, NÃO duplique, apenas continue naturalmente
"""

PROMPT_APOSTILA = """# DIRETRIZES DE REDAÇÃO: MANUAL JURÍDICO DIDÁTICO (MODO APOSTILA)
## PAPEL
VOCÊ É UM EXCELENTISSIMO REDATOR JURÍDICO E DIDÁTICO

## OBJETIVO
Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, em formato de manual didático

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias 


❌ PRESERVE obrigatoriamente:
- **NÚMEROS EXATOS**: Artigos, Leis, Artigos Súmulas, Julgados (REDI/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência, autores, casos citados
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque)
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico"

## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Ajuste a linguagem coloquial para o padrão culto.
2. **Limpeza**: Remova gírias, cacoetes ("né", "tipo assim", "então") e vícios de oralidade.
3. **Coesão**: Use conectivos e pontuação adequada para tornar o texto fluido.
4. **Legibilidade**:
   - Use parágrafos bem definidos e curtos (máximo 3-4 linhas visuais).
   - Evite blocos de texto maciços.
   - Use **negrito** para destacar conceitos-chave (sem exagero).

## 📝 ESTRUTURA
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis a partir do contexto.

## 📊 TABELAS DE RESUMO
Ao final de cada **tópico principal** (seção com título ##), crie uma tabela de resumo no formato:

```
### 📋 Resumo do Tópico

| Conceito/Ponto | Explicação Resumida |
|----------------|---------------------|
| **Conceito 1** | Definição ou ideia principal em 1-2 frases |
| **Conceito 2** | Definição ou ideia principal em 1-2 frases |
```

**Critérios para incluir tabela:**
- Use apenas ao final de seções substantivas (com conteúdo técnico/conceitual)
- Não use em seções introdutórias ou de transição curtas
- Limite a 5-7 linhas por tabela

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA
- **NUNCA formate novamente esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- **NUNCA repita informações que já estão no contexto**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>
- Se o texto_para_formatar começar com algo similar ao fim do contexto, NÃO duplique, apenas continue naturalmente
"""

PPROMPT_APOSTILA = """# DIRETRIZES DE REDAÇÃO: MANUAL JURÍDICO DIDÁTICO (MODO APOSTILA)
## PAPEL
VOCÊ É UM EXCELENTISSIMO REDATOR JURÍDICO E DIDÁTICO

## OBJETIVO
Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, em formato de manual didático

## 1. CORREÇÕES LINGUÍSTICAS
✅ Corrija:
- Erros gramaticais, ortográficos e de pontuação
- Concordâncias verbais e nominais incorretas
- Linguagem coloquial excessiva → português padrão culto
  Exemplo: "a gente viu" → "vimos" | "tipo assim" → [remover]

✅ Remova vícios de linguagem que não agregam conteúdo:
- Palavras de preenchimento: "né", "tipo assim", "sabe"
- Repetições acidentais: "é, é, é importante" → "é importante"
- Falsos inícios: "O artigo... quer dizer, a norma..." → "A norma..."


❌ PRESERVE obrigatoriamente:
- **NÚMEROS EXATOS**: Artigos, Leis, Artigos Súmulas, Julgados (REDI/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência, autores, casos citados
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque)
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico"

# 📝 ESTILO DE ESCRITA
1. **Tom Doutrinário**: Escreva como um livro de Direito.
2. **Texto Corrido**: Transforme as explicações do professor em parágrafos expositivos sólidos.
3. **Impessoalidade**: Use impessoalidade acadêmica ("Entende-se", "Observa-se", "A doutrina aponta").
4. **Conceitos**: Destaque termos técnicos e princípios jurídicos em **negrito**.
5. **Densidade**: Condense quando adequado, mas preserve toda a substância jurídica.

## 🚫 O QUE EVITAR
- Não mantenha cacoetes de oralidade.
- Não invente informações.
- Não crie exemplos não mencionados pelo professor.


## 2. ESTRUTURA E FORMATAÇÃO (MARKDOWN)
### Hierarquia de Tópicos:
1. Tópico Principal (mudança de assunto)
   1.1 Subtópico (divisão interna)
       1.1.1 Sub-subtópico (detalhamento)

### Organização do Texto:
- Parágrafos curtos (máximo 5-6 linhas)
- Divisão lógica por temas/subtemas
- Identificação de falantes se houver: Professor:, Aluno:
- Conectivos apropriados para fluidez entre ideias

### Elementos de Destaque:
- **Negrito** para institutos jurídicos, princípios e conceitos-chave
- *Itálico* para ênfases ou termos estrangeiros
- > Blockquote para transcrição literal de lei ou jurisprudência mencionada
- Listas numeradas (1. 2. 3.) para requisitos/etapas/correntes doutrinárias
- Listas com marcadores (-) para enumerações simples

## 3. TABELAS COMPARATIVAS (USE SEMPRE QUE APLICÁVEL):
- Comparação entre institutos (ex: Nulidade vs. Anulabilidade)
- Divergências doutrinárias (1ª Corrente | 2ª Corrente | STF)
- Requisitos temporais (Antes da Lei X | Depois da Lei X)

### Sintaxe:
| Aspecto | Posição A | Posição B |
|---------|-----------|-----------|
| ...     | ...       | ...       |

## 4. SÍNTESE DE SEÇÕES COMPLEXAS
Ao final de tópicos com múltiplos conceitos ou comparações, crie:
### 📋 RESUMO DO TÓPICO:
- Conceito
- Regra Geral
- Exceções/Observações
- [extrair do conteúdo da aula]

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter coerência de estilo
- **NUNCA reprocesse esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- **NUNCA repita tópicos/seções que já apareceram no contexto**
- Formate APENAS o texto entre <texto_para_formatar>
- Se o texto_para_formatar repetir algo do contexto, PULE a repetição
"""

# =============================================================================
# ESCOLHA O MODO
# =============================================================================

PROMPT_FORMATACAO = PROMPT_FIDELIDADE
# PROMPT_FORMATACAO = PROMPT_APOSTILA

# =============================================================================
# DETECÇÃO DE MODO
# =============================================================================

FIDELIDADE_MODE = "NÃO RESUMA" in PROMPT_FORMATACAO
APOSTILA_MODE = "MANUAL JURÍDICO" in PROMPT_FORMATACAO

if APOSTILA_MODE:
    THRESHOLD_MINIMO = 0.50
    THRESHOLD_CRITICO = 0.40
    MODO_NOME = "APOSTILA"
elif FIDELIDADE_MODE:
    THRESHOLD_MINIMO = 0.80
    THRESHOLD_CRITICO = 0.75
    MODO_NOME = "FIDELIDADE"
else:
    THRESHOLD_MINIMO = 0.60
    THRESHOLD_CRITICO = 0.50
    MODO_NOME = "PADRÃO"

logger.info(f"🎯 Modo: {MODO_NOME} (threshold={THRESHOLD_MINIMO:.0%})")
logger.info(f"🛡️  Anti-duplicação: ATIVADA (v2.7)")

# =============================================================================
# RATE LIMITER
# =============================================================================

class RateLimiter:
    def __init__(self, max_requests_per_minute=MAX_RPM):
        self.max_rpm = max_requests_per_minute
        self.requests = []
        self.lock = threading.Lock()
    
    def wait_if_needed(self):
        with self.lock:
            now = time.time()
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
                logger.info(f"⏱️  Rate limit: aguardando {wait_time:.1f}s...")
                sleep(wait_time)
                self.requests = [t for t in self.requests if time.time() - t < 60]
            
            self.requests.append(time.time())

rate_limiter = RateLimiter()

# =============================================================================
# CHECKPOINT/RESUME
# =============================================================================

def get_checkpoint_path(input_file):
    return Path(input_file).with_suffix('.checkpoint.json')

def save_checkpoint(input_file, resultados, chunks_info, secao_atual):
    checkpoint_path = get_checkpoint_path(input_file)
    checkpoint_data = {
        'input_file': str(input_file),
        'secao_atual': secao_atual,
        'total_secoes': len(chunks_info),
        'chunks_info': chunks_info,
        'resultados': resultados,
        'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
        'version': '2.7',
        'modo': MODO_NOME
    }
    with open(checkpoint_path, 'w', encoding='utf-8') as f:
        json.dump(checkpoint_data, f, ensure_ascii=False, indent=2)

def load_checkpoint(input_file):
    checkpoint_path = get_checkpoint_path(input_file)
    if checkpoint_path.exists():
        try:
            with open(checkpoint_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if data.get('version') not in ('2.6', '2.7'):
                    logger.warning("Checkpoint de versão anterior. Reiniciando...")
                    return None
                return data
        except Exception as e:
            logger.error(f"Erro ao carregar checkpoint: {e}")
    return None

def delete_checkpoint(input_file):
    checkpoint_path = get_checkpoint_path(input_file)
    if checkpoint_path.exists():
        checkpoint_path.unlink()
        logger.info("🧹 Checkpoint removido")

# =============================================================================
# DIVISÃO SEQUENCIAL
# =============================================================================

def dividir_sequencial(transcricao_completa):
    """Divide documento em chunks SEQUENCIAIS sem sobreposição"""
    chunks = []
    tamanho_total = len(transcricao_completa)
    inicio = 0
    
    while inicio < tamanho_total:
        fim = min(inicio + CHARS_POR_PARTE, tamanho_total)
        
        if fim < tamanho_total:
            janela = transcricao_completa[max(0, fim - 500):min(tamanho_total, fim + 500)]
            titulo_match = re.search(r'\n(#{2,4}\s+.+)\n', janela)
            
            if titulo_match:
                pos_titulo = janela.find(titulo_match.group(0))
                fim = max(0, fim - 500) + pos_titulo
            else:
                quebra = transcricao_completa.rfind('\n\n', fim - 300, fim + 300)
                if quebra != -1 and quebra > inicio:
                    fim = quebra
                else:
                    quebra = transcricao_completa.rfind('. ', fim - 150, fim + 150)
                    if quebra != -1 and quebra > inicio:
                        fim = quebra + 1
        
        chunks.append({'inicio': inicio, 'fim': fim})
        inicio = fim
    
    return chunks

def validar_chunks(chunks, transcricao_completa):
    """v2.7: Validação rigorosa de chunks sequenciais"""
    logger.info("🔍 Validando chunks sequenciais...")
    
    for i in range(len(chunks)):
        chunk = chunks[i]
        
        # Verifica se início == fim do anterior
        if i > 0:
            anterior = chunks[i-1]
            if chunk['inicio'] != anterior['fim']:
                logger.error(f"❌ Gap/Overlap no chunk {i+1}!")
                logger.error(f"   Anterior termina em: {anterior['fim']}")
                logger.error(f"   Atual começa em: {chunk['inicio']}")
                logger.error(f"   Diferença: {chunk['inicio'] - anterior['fim']} chars")
                
                # Mostra preview
                if chunk['inicio'] < anterior['fim']:
                    overlap_text = transcricao_completa[chunk['inicio']:anterior['fim']]
                    logger.error(f"   OVERLAP: '{overlap_text[:100]}...'")
                
                return False
    
    logger.info(f"✅ {len(chunks)} chunks validados (sequenciais, sem overlap)")
    return True

# =============================================================================
# FUNÇÕES AUXILIARES
# =============================================================================

def limpar_tags_xml(texto):
    texto = re.sub(r'</?[a-z_][\w\-]*>', '', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<[a-z_][\w\-]*\s+[^>]+>', '', texto, flags=re.IGNORECASE)
    return texto

def carregar_transcricao(arquivo):
    try:
        with open(arquivo, 'r', encoding='utf-8-sig') as f:
            conteudo = f.read()
        
        if not conteudo.strip():
            logger.error("Arquivo está vazio.")
            sys.exit(1)
            
        return conteudo
    except FileNotFoundError:
        logger.error(f"Arquivo '{arquivo}' não encontrado.")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Erro ao ler arquivo: {e}")
        sys.exit(1)

def estimar_custo(transcricao, usar_cache, num_chunks=1):
    tokens_in = len(transcricao) // 4
    
    if APOSTILA_MODE:
        tokens_out_estimado = int(tokens_in * 0.65)
    elif FIDELIDADE_MODE:
        tokens_out_estimado = int(tokens_in * 1.00)
    else:
        tokens_out_estimado = int(tokens_in * 0.85)
    
    tokens_prompt = len(PROMPT_FORMATACAO) // 4
    tokens_in_total = tokens_in + (tokens_prompt * num_chunks)
    
    if usar_cache:
        custo_input = (tokens_in * PRECO_INPUT_COM_CACHE + tokens_prompt * num_chunks * PRECO_INPUT_SEM_CACHE) / 1_000_000
        custo_output = (tokens_out_estimado * PRECO_OUTPUT) / 1_000_000
        custo = custo_input + custo_output
    else:
        custo = (tokens_in_total * PRECO_INPUT_SEM_CACHE + tokens_out_estimado * PRECO_OUTPUT) / 1_000_000
    
    logger.info(f"💰 Custo estimado: ${custo:.4f} USD (modo {MODO_NOME})")

# v2.7: DESABILITAR cache para evitar "vazamento" de conteúdo
def criar_cache_contexto(client, transcricao_completa):
    """v2.7: Cache DESABILITADO temporariamente para debug de duplicações"""
    logger.warning("⚠️  Cache desabilitado na v2.7 (debug anti-duplicação)")
    return None

# =============================================================================
# PROCESSAMENTO
# =============================================================================

def processar_simples(client, transcricao_bruta):
    logger.info("📄 Documento pequeno - processando em requisição única...")
    
    prompt = f"""{PROMPT_FORMATACAO}

<texto_para_formatar>
{transcricao_bruta}
</texto_para_formatar>

Retorne APENAS o Markdown formatado."""
    
    for tentativa in range(MAX_RETRIES):
        try:
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=OUTPUT_TOKEN_LIMIT,
                    safety_settings=[
                        types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                    ]
                )
            )
            resultado = response.text
            return limpar_tags_xml(resultado)
        except Exception as e:
            if tentativa < MAX_RETRIES - 1:
                wait = (2 ** tentativa) + random.uniform(0, 1)
                logger.warning(f"Erro, retry em {wait:.1f}s...")
                sleep(wait)
            else:
                raise

def processar_chunk(client, cache, texto_chunk, numero, total, contexto_estilo=""):
    rate_limiter.wait_if_needed()
    
    # v2.7: Delimitadores MUITO visíveis e instruções reforçadas
    secao_contexto = ""
    if contexto_estilo:
        secao_contexto = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔒 CONTEXTO ANTERIOR (SOMENTE REFERÊNCIA DE ESTILO)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{contexto_estilo}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ ATENÇÃO: O bloco acima JÁ FOI FORMATADO anteriormente.
- NÃO formate novamente esse conteúdo
- NÃO inclua esse conteúdo na sua resposta
- Use APENAS como referência de estilo de escrita
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 NOVO TEXTO PARA FORMATAR (comece aqui):
"""
    
    prompt = f"""{PROMPT_FORMATACAO}
{secao_contexto}

<texto_para_formatar>
{texto_chunk}
</texto_para_formatar>

**INSTRUÇÕES FINAIS**:
- Esta é a parte {numero} de {total}
- Formate APENAS o texto entre <texto_para_formatar>
- Se houver contexto acima, NÃO o reprocesse
- Retorne APENAS o Markdown formatado do NOVO texto
"""

    for tentativa in range(MAX_RETRIES):
        try:
            safety_config = [
                types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
            ]
            
            # v2.7: Sem cache por enquanto
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=OUTPUT_TOKEN_LIMIT,
                    safety_settings=safety_config
                )
            )
            
            resultado = response.text
            
            if not resultado or not resultado.strip():
                raise Exception("Resposta vazia da API")
            
            resultado = limpar_tags_xml(resultado)
            
            # Validação de tamanho
            palavras_in = len(texto_chunk.split())
            palavras_out = len(resultado.split())
            razao = palavras_out / palavras_in if palavras_in > 0 else 1.0
            
            if razao < THRESHOLD_MINIMO:
                logger.warning(f"Seção {numero}: Output curto ({razao:.0%}). Reprocessando...")
                if tentativa < MAX_RETRIES - 1:
                    continue
                else:
                    logger.error(f"Seção {numero}: Falha após {MAX_RETRIES} tentativas (razão={razao:.0%})")
            
            return resultado
            
        except Exception as e:
            erro_msg = str(e)
            is_recoverable = any(code in erro_msg for code in ['503', '429', '500', 'RESOURCE_EXHAUSTED']) or "Resposta vazia" in erro_msg
            
            if tentativa < MAX_RETRIES - 1 and is_recoverable:
                wait = (2 ** tentativa) + random.uniform(0, 1)
                logger.warning(f"Erro seção {numero}, retry {tentativa+2}/{MAX_RETRIES} em {wait:.1f}s")
                sleep(wait)
            else:
                logger.error(f"Falha seção {numero}: {erro_msg}")
                return f"\n\n> [!WARNING]\n> Falha ao processar seção {numero}. Texto original:\n\n{texto_chunk}"
    
    return texto_chunk

# =============================================================================
# PÓS-PROCESSAMENTO v2.7 - ANTI-DUPLICAÇÃO AGRESSIVA
# =============================================================================

def extrair_titulos_h2(texto):
    """Extrai todos os títulos de nível 2 (##) do texto"""
    titulos = []
    for linha in texto.split('\n'):
        if linha.strip().startswith('##') and not linha.strip().startswith('###'):
            titulo_limpo = re.sub(r'^##\s*\d+\.\s*', '', linha.strip())
            titulos.append(titulo_limpo.lower())
    return titulos

def detectar_secoes_duplicadas(texto):
    """v2.7: Detecta seções inteiras duplicadas por títulos"""
    logger.info("🔍 Detectando seções duplicadas...")
    
    linhas = texto.split('\n')
    titulos_vistos = {}
    secoes_duplicadas = []
    
    for i, linha in enumerate(linhas):
        if linha.strip().startswith('##') and not linha.strip().startswith('###'):
            # Normaliza título (remove numeração)
            titulo_normalizado = re.sub(r'^##\s*\d+\.?\s*', '', linha.strip()).lower()
            titulo_normalizado = re.sub(r'[📋📊🗂️]', '', titulo_normalizado).strip()
            
            if titulo_normalizado in titulos_vistos:
                logger.warning(f"⚠️  Título duplicado encontrado: '{linha.strip()}'")
                logger.warning(f"   Primeira ocorrência: linha {titulos_vistos[titulo_normalizado]}")
                logger.warning(f"   Segunda ocorrência: linha {i}")
                secoes_duplicadas.append({
                    'titulo': titulo_normalizado,
                    'primeira_linha': titulos_vistos[titulo_normalizado],
                    'duplicada_linha': i
                })
            else:
                titulos_vistos[titulo_normalizado] = i
    
    if secoes_duplicadas:
        logger.error(f"❌ {len(secoes_duplicadas)} seções duplicadas detectadas!")
        return secoes_duplicadas
    else:
        logger.info("✅ Nenhuma seção duplicada detectada")
        return []

def remover_secoes_duplicadas(texto):
    """v2.7: Remove seções duplicadas mantendo apenas a primeira ocorrência"""
    secoes_dup = detectar_secoes_duplicadas(texto)
    
    if not secoes_dup:
        return texto
    
    logger.info("🧹 Removendo seções duplicadas...")
    
    linhas = texto.split('\n')
    linhas_para_remover = set()
    
    for dup in secoes_dup:
        # Marca para remoção todas as linhas da seção duplicada
        inicio_remocao = dup['duplicada_linha']
        
        # Encontra onde a seção duplicada termina (próximo ## ou fim do arquivo)
        fim_remocao = len(linhas)
        for i in range(inicio_remocao + 1, len(linhas)):
            if linhas[i].strip().startswith('##') and not linhas[i].strip().startswith('###'):
                fim_remocao = i
                break
        
        logger.info(f"   Removendo linhas {inicio_remocao}-{fim_remocao} (seção duplicada)")
        for i in range(inicio_remocao, fim_remocao):
            linhas_para_remover.add(i)
    
    # Reconstrói texto sem as linhas duplicadas
    linhas_limpas = [linha for i, linha in enumerate(linhas) if i not in linhas_para_remover]
    
    logger.info(f"✅ {len(linhas_para_remover)} linhas removidas")
    return '\n'.join(linhas_limpas)

def remover_duplicacoes_literais(texto):
    """Remove parágrafos individuais duplicados"""
    paragrafos = texto.split('\n\n')
    paragrafos_limpos = []
    dup_count = 0
    
    for i, para in enumerate(paragrafos):
        if i == 0:
            paragrafos_limpos.append(para)
            continue
        
        if len(para.strip()) < 80 or para.strip().startswith('#'):
            paragrafos_limpos.append(para)
            continue
        
        is_duplicate = False
        para_norm = ' '.join(para.split()).lower()
        
        for j in range(max(0, len(paragrafos_limpos) - 3), len(paragrafos_limpos)):
            para_ant = paragrafos_limpos[j]
            para_ant_norm = ' '.join(para_ant.split()).lower()
            
            ratio = SequenceMatcher(None, para_norm, para_ant_norm).ratio()
            
            if ratio > 0.95:
                is_duplicate = True
                dup_count += 1
                break
        
        if not is_duplicate:
            paragrafos_limpos.append(para)
    
    if dup_count > 5:
        logger.warning(f"⚠️  {dup_count} parágrafos duplicados removidos")
    
    return '\n\n'.join(paragrafos_limpos)

def numerar_titulos(texto):
    """Adiciona numeração sequencial aos títulos"""
    linhas = texto.split('\n')
    linhas_numeradas = []
    
    contador_h2 = 0
    contador_h3 = 0
    contador_h4 = 0
    
    titulo_pattern = re.compile(r'^(#{2,4})\s+(?:\d+(?:\.\d+)*\.?\s+)?(.+)$')
    
    for linha in linhas:
        match = titulo_pattern.match(linha)
        
        if match:
            nivel = len(match.group(1))
            texto_titulo = match.group(2).strip()
            
            # Não numera títulos de resumo/quadros
            if any(keyword in texto_titulo.lower() for keyword in ['resumo', 'quadro', 'esquema', '📋', '📊', '🗂️']):
                linhas_numeradas.append(linha)
                continue
            
            if nivel == 2:
                contador_h2 += 1
                contador_h3 = 0
                contador_h4 = 0
                nova_linha = f"## {contador_h2}. {texto_titulo}"
            elif nivel == 3:
                contador_h3 += 1
                contador_h4 = 0
                nova_linha = f"### {contador_h2}.{contador_h3}. {texto_titulo}"
            elif nivel == 4:
                contador_h4 += 1
                nova_linha = f"#### {contador_h2}.{contador_h3}.{contador_h4}. {texto_titulo}"
            else:
                nova_linha = linha
            
            linhas_numeradas.append(nova_linha)
        else:
            linhas_numeradas.append(linha)
    
    return '\n'.join(linhas_numeradas)

# =============================================================================
# FLUXO PRINCIPAL
# =============================================================================

def formatar_transcricao(transcricao_completa, usar_cache=True, input_file=None):
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        logger.error("GEMINI_API_KEY não configurada.")
        logger.error("Configure com: export GEMINI_API_KEY='sua-chave'")
        sys.exit(1)
    
    client = genai.Client(api_key=api_key)
    
    # Dry-run mode
    if '--dry-run' in sys.argv:
        logger.info("🔍 MODO DRY-RUN: Validando divisão de chunks")
        chunks = dividir_sequencial(transcricao_completa)
        validar_chunks(chunks, transcricao_completa)
        for i, c in enumerate(chunks):
            print(f"\n  Chunk {i+1}/{len(chunks)}:")
            print(f"    Posição: {c['inicio']:,} → {c['fim']:,} ({c['fim']-c['inicio']:,} chars)")
            inicio_preview = transcricao_completa[c['inicio']:c['inicio']+80].replace('\n', '↵')
            fim_preview = transcricao_completa[max(0, c['fim']-80):c['fim']].replace('\n', '↵')
            print(f"    Início: {inicio_preview}...")
            print(f"    Fim: ...{fim_preview}")
        sys.exit(0)
    
    tamanho_total = len(transcricao_completa)
    logger.info(f"📊 Tamanho: {tamanho_total:,} caracteres")
    
    chunks = dividir_sequencial(transcricao_completa)
    
    # v2.7: Validação rigorosa
    if not validar_chunks(chunks, transcricao_completa):
        logger.error("❌ Chunks inválidos! Abortando.")
        sys.exit(1)
    
    num_partes = len(chunks)
    chunks_info = [{'inicio': c['inicio'], 'fim': c['fim']} for c in chunks]
    
    estimar_custo(transcricao_completa, usar_cache, num_partes)
    
    if num_partes == 1:
        return processar_simples(client, transcricao_completa), None
    
    checkpoint = None
    resultados = []
    inicio_secao = 0
    
    if input_file:
        checkpoint = load_checkpoint(input_file)
        if checkpoint:
            logger.info(f"📁 Checkpoint: seção {checkpoint['secao_atual']}/{checkpoint['total_secoes']}")
            resposta = input("   Continuar? (s/n): ").strip().lower()
            if resposta == 's':
                resultados = checkpoint['resultados']
                inicio_secao = checkpoint['secao_atual']
                chunks_info = checkpoint['chunks_info']
            else:
                delete_checkpoint(input_file)
    
    # v2.7: Cache desabilitado temporariamente
    cache = None
    
    try:
        iterator = range(inicio_secao, num_partes)
        if tqdm:
            iterator = tqdm(iterator, desc="Formatando", initial=inicio_secao, total=num_partes)
        
        for i in iterator:
            chunk = chunks_info[i]
            texto_chunk = transcricao_completa[chunk['inicio']:chunk['fim']]
            
            # Contexto com validação
            contexto_estilo = ""
            if i > 0 and resultados:
                raw_context = resultados[-1][-CONTEXTO_ESTILO:]
                if len(raw_context.split()) < 50 or "[!WARNING]" in raw_context:
                    logger.warning(f"Contexto chunk {i+1} descartado")
                else:
                    contexto_estilo = raw_context
            
            resultado = processar_chunk(
                client, cache, texto_chunk,
                i + 1, num_partes,
                contexto_estilo=contexto_estilo
            )
            
            resultados.append(resultado)
            
            if input_file:
                save_checkpoint(input_file, resultados, chunks_info, i + 1)
            
            if not tqdm:
                logger.info(f"✅ Seção {i+1}/{num_partes}")
        
        # v2.7: Post-processing em múltiplas passadas
        logger.info("🧹 Iniciando limpeza (v2.7)...")
        
        texto_final = '\n\n'.join(resultados)
        
        logger.info("  Passada 1: Removendo duplicações literais...")
        texto_final = remover_duplicacoes_literais(texto_final)
        
        logger.info("  Passada 2: Detectando seções duplicadas...")
        texto_final = remover_secoes_duplicadas(texto_final)
        
        logger.info("  Passada 3: Numerando títulos...")
        texto_final = numerar_titulos(texto_final)
        
        # Validação final
        palavras_in = len(transcricao_completa.split())
        palavras_out = len(texto_final.split())
        razao = palavras_out / palavras_in if palavras_in > 0 else 1.0
        
        logger.info(f"✅ Validação: {razao:.0%} do original ({palavras_out:,}/{palavras_in:,} palavras)")
        
        if razao < THRESHOLD_CRITICO:
            if FIDELIDADE_MODE:
                logger.error(f"❌ POSSÍVEL PERDA DE CONTEÚDO ({razao:.0%})")
                logger.error(f"   Esperado: >{THRESHOLD_CRITICO:.0%} | Obtido: {razao:.0%}")
            elif APOSTILA_MODE:
                logger.warning(f"⚠️  Texto condensado: {razao:.0%}")
                logger.info(f"   ✅ Esperado no modo {MODO_NOME}")
            
            if razao < 0.30 or (FIDELIDADE_MODE and razao < THRESHOLD_CRITICO):
                resposta = input("\n   Continuar? (s/n): ").strip().lower()
                if resposta != 's':
                    logger.info("Cancelado.")
                    sys.exit(1)
        
        if input_file:
            delete_checkpoint(input_file)
        
        return texto_final, cache
    
    except KeyboardInterrupt:
        logger.warning("\n⚠️  Interrompido. Checkpoint salvo.")
        sys.exit(1)

# =============================================================================
# EXPORTAÇÃO WORD
# =============================================================================

def create_toc(doc):
    """Adiciona Sumário (Table of Contents) nativo do Word"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def _format_inline_markdown(paragraph, text):
    """Formata markdown inline (negrito, itálico, código)"""
    paragraph.clear()
    
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        if match.group(0).startswith('***'):
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(0).startswith('**'):
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(0).startswith('*'):
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(0).startswith('`'):
            run = paragraph.add_run(match.group(5))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(200, 0, 0)
        
        last_end = match.end()
    
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

def _add_table_to_doc(doc, rows):
    """Adiciona tabela formatada ao documento Word"""
    if len(rows) < 2:
        return
    
    max_cols = max(len(row) for row in rows)
    if max_cols == 0:
        return
    
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(rows):
        for j in range(max_cols):
            cell = table.rows[i].cells[j]
            cell_text = row_data[j] if j < len(row_data) else ""
            
            # Formata markdown dentro da célula
            _format_inline_markdown(cell.paragraphs[0], cell_text)
            
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '0066CC')
                cell._element.get_or_add_tcPr().append(shading_elm)

def save_as_word(formatted_text, video_name, output_file):
    """Salva markdown formatado como documento Word (.docx)"""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx não disponível. Salvando apenas Markdown.")
        return None
    
    logger.info("📄 Gerando documento Word...")
    
    doc = Document()
    
    # Margens
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    
    # Título principal
    title = doc.add_heading(video_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Data de geração
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')} - Modo: {MODO_NOME}")
    date_run.italic = True
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Sumário
    doc.add_heading('Sumário', level=1)
    create_toc(doc)
    doc.add_page_break()
    
    # Processa conteúdo markdown
    lines = formatted_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Tabelas
        if '|' in line and not in_table:
            in_table = True
            table_rows = []
        
        if in_table:
            if '|' in line and not line.startswith('|--'):
                table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
            
            if '|' not in line or i == len(lines) - 1:
                if len(table_rows) > 0:
                    _add_table_to_doc(doc, table_rows)
                in_table = False
                table_rows = []
                if '|' not in line:
                    continue
            i += 1
            continue
        
        # Headings
        if line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# ') and line != f"# {video_name}":
            doc.add_heading(line[2:], level=1)
        # Separadores
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.add_run('_' * 80).font.color.rgb = RGBColor(192, 192, 192)
        # Quotes
        elif line.startswith('>'):
            p = doc.add_paragraph(line[1:].strip(), style='Quote')
            p.paragraph_format.left_indent = Inches(0.5)
        # Listas não-ordenadas
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            _format_inline_markdown(p, line[2:])
        # Listas numeradas
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-0.63)
            _format_inline_markdown(p, line)
        # Parágrafo normal
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _format_inline_markdown(p, line)
        
        i += 1
    
    doc.save(output_file)
    logger.info(f"✅ Word salvo: {output_file}")
    return output_file

def salvar_resultado(conteudo, arquivo):
    with open(arquivo, 'w', encoding='utf-8') as f:
        f.write(conteudo)
    logger.info(f"✅ Markdown salvo: {arquivo}")

# =============================================================================
# MAIN
# =============================================================================

def main():
    if len(sys.argv) < 2 or '--help' in sys.argv:
        print("=" * 70)
        print("FORMATADOR v2.7 - ANTI-DUPLICAÇÃO")
        print("=" * 70)
        print("\nUso: python format_transcription_gemini.py <entrada.txt> [saida]")
        print("\nOpções:")
        print("  --dry-run    Valida chunks e mostra preview")
        print("  --help       Mostra esta mensagem")
        print("\n🛡️  CORREÇÕES v2.7:")
        print("  • Detecção agressiva de seções duplicadas")
        print("  • Validação rigorosa de chunks sequenciais")
        print("  • Delimitadores de contexto mais visíveis")
        print("  • Post-processing em múltiplas passadas")
        print("  • Cache desabilitado (debug)")
        print("\nDependências:")
        print("  pip install google-genai python-docx tqdm")
        sys.exit(1)
    
    arquivo_entrada = sys.argv[1]
    
    if len(sys.argv) > 2 and not sys.argv[2].startswith('--'):
        arquivo_saida = sys.argv[2]
    else:
        base = arquivo_entrada.replace('.txt', '_formatada_v2.7')
        arquivo_saida = f"{base}.md"
    
    video_name = Path(arquivo_entrada).stem
    
    logger.info("=" * 60)
    logger.info(f"FORMATADOR v2.7 ANTI-DUPLICAÇÃO - Modo {MODO_NOME}")
    logger.info("=" * 60)
    logger.info(f"📂 Entrada: {arquivo_entrada}")
    logger.info(f"📝 Saída: {arquivo_saida}")
    
    transcricao = carregar_transcricao(arquivo_entrada)
    
    try:
        resultado, cache = formatar_transcricao(transcricao, input_file=arquivo_entrada)
    except Exception as e:
        logger.error(f"\n❌ Falha: {e}", exc_info=True)
        sys.exit(1)
    
    salvar_resultado(resultado, arquivo_saida)
    
    tokens_in = len(transcricao) // 4
    tokens_out = len(resultado) // 4
    custo = (tokens_in * PRECO_INPUT_SEM_CACHE + tokens_out * PRECO_OUTPUT) / 1_000_000
    
    logger.info("=" * 60)
    logger.info(f"💰 Custo: ${custo:.4f} USD")
    logger.info(f"✨ Concluído! (v2.7 anti-duplicação)")
    logger.info("=" * 60)

if __name__ == "__main__":
    main()
