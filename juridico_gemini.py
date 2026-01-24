import os
import sys
import time
import json
import re
import logging
import traceback
from pathlib import Path
from typing import List, Dict, Optional, Generator, Any
import datetime
from app.core.config import settings
try:
    from audit_juridico import audit_document_text
    AUDIT_AVAILABLE = True
except ImportError:
    print(f"⚠️ Módulo de auditoria não encontrado (audit_juridico.py)")
    AUDIT_AVAILABLE = False

# Third-party imports
try:
    import anthropic
except ImportError:
    anthropic = None
    print("⚠️ Anthropic SDK não instalado. Fallback desativado.")

try:
    import google.genai as genai
    from google.genai import types
    from jinja2 import Template
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from difflib import SequenceMatcher
except ImportError as e:
    print(f"❌ Erro de importação: {e}")
    print("Instale as dependências: pip install -r requirements.txt")
    # sys.exit(1) # Removido para permitir testes/diagnóstico parcial

try:
    from colorama import Fore, Style, init
    init(autoreset=True)
except ImportError:
    # Fallback for colorama if not available
    class Fore:
        CYAN = YELLOW = GREEN = RED = MAGENTA = ""
    class Style:
        RESET_ALL = ""

try:
    from rag_module import get_deduplicator
    HAS_RAG_DEDUPE = True
except ImportError:
    HAS_RAG_DEDUPE = False
    get_deduplicator = None

# Configuration - Vertex AI (v3.0)
# Credentials: Uses GOOGLE_APPLICATION_CREDENTIALS or ADC
from dotenv import load_dotenv
load_dotenv(override=True)

PROJECT_ID = os.getenv("GOOGLE_CLOUD_PROJECT")
# Credenciais devem vir via variável de ambiente GOOGLE_APPLICATION_CREDENTIALS
if os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
    print(f"{Fore.CYAN}🔑 Credenciais Vertex detectadas via ambiente.")
else:
    print(f"{Fore.YELLOW}⚠️  Variável GOOGLE_APPLICATION_CREDENTIALS não definida. Tentando ADC padrão...")

# Gemini 3 Flash Preview (v3.0 - melhor custo-benefício)
DEFAULT_MODEL = "gemini-3-flash-preview"

# Logging Setup
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger("JuridicoAI")


# =============================================================================
# PROMPTS & TEMPLATES (Jinja2)
# =============================================================================

# =============================================================================
# PROMPTS & TEMPLATES (Jinja2)
# =============================================================================

PROMPT_MAP = {
    "PETICAO_INICIAL": """
    Você é um Advogado de Autores focado em procedência.
    Crie um SUMÁRIO para PETIÇÃO INICIAL de {{ resumo_caso }}.
    Estrutura Obrigatória (Art. 319 CPC):
    1. Endereçamento e Qualificação das Partes
    2. Dos Fatos (Narrativa cronológica e persuasiva)
    3. Do Direito (Dividido por Teses/Fundamentos - cite súmulas e jurisprudência)
    4. Das Provas a Produzir (Art. 319, VI CPC)
    5. Dos Pedidos (Exaustivos - principais e subsidiários)
    6. Do Valor da Causa
    7. Opção por Audiência de Conciliação (Art. 319, VII CPC)
    """,
    
    "CONTESTACAO": """
    Você é um Advogado de Defesa focado em improcedência.
    Crie um SUMÁRIO para CONTESTAÇÃO de {{ resumo_caso }}.
    Estrutura Obrigatória (Art. 336-337 CPC):
    1. Tempestividade
    2. Síntese da Inicial
    3. Preliminares (Art. 337 CPC - Inépcia, Ilegitimidade, Litispendência, etc)
    4. Prejudicial de Mérito (Prescrição, Decadência)
    5. Mérito (Impugnação ESPECÍFICA de cada fato - Art. 341 CPC)
    6. Reconvenção (se cabível e conexa - Art. 343 CPC)
    7. Provas a Produzir
    8. Requerimentos Finais
    """,
    
    "PARECER": """
    Você é um Consultor Jurídico Sênior (neutro e analítico).
    Crie um SUMÁRIO para PARECER JURÍDICO sobre: {{ resumo_caso }}.
    Estrutura:
    1. Objeto da Consulta
    2. Relatório dos Fatos
    3. Fundamentação (Doutrina majoritária + Jurisprudência dominante)
    4. Análise de Riscos (Tabela: Tese vs Probabilidade de Êxito)
    5. Conclusão (Opinativo final com recomendação)
    """,

    "SENTENCA": """
    Você é um Juiz de Direito imparcial e técnico.
    Crie um SUMÁRIO para SENTENÇA judicial sobre: {{ resumo_caso }}.
    Estrutura Rígida (Art. 489 CPC - sob pena de nulidade):
    1. Relatório (Resumo do processo sem juízo de valor)
    2. Fundamentação (Análise de CADA argumento das partes - Art. 489, §1º)
    3. Dispositivo (Decisão clara: procedente/improcedente, condenação, custas e honorários)
    """,

    "RECURSO": """
    Você é um Advogado Recorrente.
    Crie um SUMÁRIO para RECURSO DE APELAÇÃO contra a decisão de {{ resumo_caso }}.
    Estrutura:
    1. Tempestividade e Preparo (Art. 1.007 CPC)
    2. Síntese do Processo e da Decisão Recorrida
    3. Do Cabimento (Art. 1.009 CPC)
    4. Razões de Reforma (Error in judicando/procedendo)
    5. Prequestionamento (Art. 1.025 CPC - para eventual REsp/RE)
    6. Pedidos (Reforma ou Anulação com retorno dos autos)
    """,
    
    "CONTRATO": """
    Você é um Advogado Especialista em Contratos.
    Crie um SUMÁRIO para CONTRATO de {{ resumo_caso }}.
    Estrutura:
    1. Qualificação das Partes (CNPJ/CPF, Endereço, Representante Legal)
    2. Objeto do Contrato (Definição clara e delimitada)
    3. Obrigações da Parte A (CONTRATANTE)
    4. Obrigações da Parte B (CONTRATADA)
    5. Preço e Condições de Pagamento (Tabela se necessário)
    6. Prazo de Vigência e Rescisão
    7. Penalidades e Multas (Cláusula penal compensatória)
    8. Confidencialidade (se aplicável)
    9. Foro de Eleição
    """,
    
    "DESPACHO": """
    Você é um Juiz ou Servidor Judiciário.
    Crie um SUMÁRIO para DESPACHO de mero expediente sobre {{ resumo_caso }}.
    Estrutura (breve e imperativo):
    1. Autos/Processo nº
    2. Determinação (ex: "Cite-se", "Intime-se", "Junte-se", "Dê-se vista")
    3. Prazo (se aplicável)
    4. Providências complementares
    """,
    
    "DECISAO_INTERLOCUTORIA": """
    Você é um Juiz de Direito decidindo questão incidental.
    Crie um SUMÁRIO para DECISÃO INTERLOCUTÓRIA sobre {{ resumo_caso }}.
    Estrutura (Art. 203, §2º CPC):
    1. Autos/Processo nº
    2. Relatório breve da questão incidental
    3. Fundamentação (Art. 11 CPC - motivação obrigatória)
    4. Dispositivo (ex: "Defiro a tutela provisória...", "Indefiro o pedido de...")
    """,
    
    "VOTO": """
    Você é um Desembargador/Ministro relator.
    Crie um SUMÁRIO para VOTO em acórdão sobre {{ resumo_caso }}.
    Estrutura:
    1. Ementa (resumo da tese fixada)
    2. Relatório (síntese do recurso e das contrarrazões)
    3. Voto (Fundamentação detalhada conduzindo à conclusão)
    4. Conclusão (ex: "Dou PROVIMENTO ao recurso para...")
    """,
    
    "NOTA_TECNICA": """
    Você é um Analista Técnico de órgão público.
    Crie um SUMÁRIO para NOTA TÉCNICA sobre {{ resumo_caso }}.
    Estrutura:
    1. Identificação (Nº, Órgão, Data, Interessado)
    2. Assunto (Objeto da análise)
    3. Contextualização (Fatos/Normas relevantes)
    4. Análise Técnica (Parecer objetivo)
    5. Conclusão e Recomendações
    """,
    
    "OFICIO": """
    Você é um servidor público redigindo comunicação oficial.
    Crie um SUMÁRIO para OFÍCIO sobre {{ resumo_caso }}.
    Estrutura (Manual de Redação Oficial):
    1. Endereçamento (A quem, Cargo, Órgão)
    2. Referência (Processo/Assunto)
    3. Corpo (Apresentação, Solicitação/Comunicação, Justificativa)
    4. Fecho e Assinatura
    """,
    
    "CI": """
    Você é um servidor redigindo comunicação interna (memorando/CI).
    Crie um SUMÁRIO para COMUNICAÇÃO INTERNA sobre {{ resumo_caso }}.
    Estrutura (formato simples e objetivo):
    1. De/Para (Destinatário e Remetente com cargos)
    2. Assunto (Linha de assunto clara)
    3. Corpo (Informação ou solicitação, justificativa breve)
    4. Prazo para resposta (se aplicável)
    """,
    
    "NOTA_JURIDICA": """
    Você é um Procurador ou Advogado Público redigindo nota jurídica opinativa.
    Crie um SUMÁRIO para NOTA JURÍDICA sobre {{ resumo_caso }}.
    Estrutura:
    1. Processo/Referência (SEI, PAD, etc.)
    2. Consulente (Unidade que solicitou)
    3. Síntese da Consulta
    4. Fundamentação Jurídica (Doutrina, Lei, Jurisprudência)
    5. Conclusão (Opinativo vinculante ou não vinculante)
    """,
    
    "ESCRITURA": """
    Você é um Tabelião de Notas ou Escrevente Autorizado.
    Crie um SUMÁRIO para ESCRITURA PÚBLICA de {{ resumo_caso }}.
    Estrutura (Ato Notarial Solene):
    1. Cabeçalho (Data, Local, Tabelionato, Livro/Folha)
    2. Qualificação das Partes (Outorgantes e Outorgados completíssima)
    3. Objeto e Declarações (Venda e Compra, Doação, Inventário, etc.)
    4. Da Transmissão e do Preço (se oneroso)
    5. Das Certidões e Documentos Apresentados (DOI, CNDs)
    6. Notas do Tabelião (Fé pública, leitura do ato, emissão de custas)
    7. Assinaturas e Encerramento
    """
}

# v1.1: Mapping from args.mode to tipo_peca field in pecas_modelo metadata
MODE_TO_TIPO_PECA = {
    "PETICAO_INICIAL": "peticao_inicial",
    "PETICAO": "peticao_inicial",
    "CONTESTACAO": "contestacao",
    "PARECER": "parecer",
    "SENTENCA": "sentenca",
    "RECURSO": "recurso",
    "APELACAO": "recurso",
    "CONTRATO": "contrato",
    "DESPACHO": "despacho",
    "DECISAO_INTERLOCUTORIA": "decisao_interlocutoria",
    "VOTO": "voto",
    "NOTA_TECNICA": "nota_tecnica",
    "OFICIO": "oficio",
    "CI": "ci",
    "COMUNICACAO_INTERNA": "ci",
    "NOTA_JURIDICA": "nota_juridica",
    "ESCRITURA": "escritura_notas",
}

TEMPLATE_OUTLINE_BASE = """
{{ role_instruction }}

## Contexto do Caso
{{ resumo_caso }}

## Tese / Instrução do Usuário
{{ tese_usuario }}

## Formato de Saída (JSON)
Responda APENAS com um JSON válido contendo uma lista de strings (títulos das seções).
Exemplo de formato esperado:
["1. Endereçamento e Qualificação", "2. Dos Fatos", "3. Do Direito", "4. Dos Pedidos"]

NÃO inclua explicações, apenas o JSON puro.
"""

TEMPLATE_SECAO = """
<role>
Você é um Especialista Jurídico redigindo a seção "{{ titulo_secao }}".
Tipo de Peça: {{ tipo_peca }}.
</role>

<style>
{% if tipo_peca == "PETICAO_INICIAL" or tipo_peca == "CONTESTACAO" or tipo_peca == "RECURSO" %}
Tom: PERSUASIVO e combativo. Você é advogado defendendo seu cliente.
Use linguagem que convença o juiz. Cite súmulas e jurisprudência dominante.
{% elif tipo_peca == "SENTENCA" %}
Tom: TÉCNICO e imparcial. Você é juiz analisando o caso.
Use linguagem decisória. Fundamente CADA ponto. Evite ambiguidades.
{% elif tipo_peca == "PARECER" %}
Tom: ANALÍTICO e neutro. Você é consultor avaliando riscos.
Apresente prós e contras. Use tabelas para síntese.
{% elif tipo_peca == "CONTRATO" %}
Tom: PRECISO e formal. Use linguagem contratual técnica.
Evite ambiguidades. Defina termos quando necessário.
{% else %}
Tom: Forense e técnico. Formatação: Markdown.
{% endif %}
</style>

<anti_hallucination_rules>
⚠️ REGRAS CRÍTICAS ANTI-ALUCINAÇÃO:
1. NUNCA invente fatos, leis, súmulas ou jurisprudência.
2. Se houver fontes RAG anexas abaixo, cite APENAS essas fontes.
3. Se não houver fonte RAG para uma afirmação, use "conforme documentos anexos" ou "a conferir".
4. Prefira citar legislação consolidada (CF, CC, CPC) a inventar precedentes.
5. Em caso de dúvida sobre existência de um julgado, NÃO cite.
</anti_hallucination_rules>

<context>
Estamos na seção {{ indice }} de {{ total }}.
Resumo do que já foi escrito (para manter coesão):
{{ contexto_anterior }}
</context>

{% if fontes_rag %}
<rag_sources>
## FONTES RECUPERADAS (Use APENAS estas para citações):
{{ fontes_rag }}
</rag_sources>
{% endif %}

<citation_format>
Ao citar, use EXATAMENTE estes formatos:
- Súmulas: "Súmula 123 do STJ" (número + tribunal)
- Jurisprudência: "REsp 1.234.567/SP, Rel. Min. Fulano, DJe 01/01/2024"
- Leis: "Art. 123, §1º, inciso II, da Lei 8.666/93"
- Docs dos Autos: "[TIPO - Doc. X, p. Y]"
</citation_format>

<task>
Escreva APENAS o conteúdo da seção "{{ titulo_secao }}".
DESENVOLVIMENTO COMPLETO: Cubra todos os aspectos relevantes.
Se houver dados numéricos ou comparações, USE TABELAS.
Tamanho: adequado à complexidade do tópico (nem curto demais, nem prolixo).
</task>
"""


# =============================================================================
# METRICS & SAFETY (v2.2 - Ported from mlx_vomo.py)
# =============================================================================

import threading
import difflib
import asyncio

class RateLimiter:
    """Controla requisições por minuto para não estourar rate limit da API (Fixed for Async)"""
    def __init__(self, max_requests_per_minute=60):
        self.max_rpm = max_requests_per_minute
        self.requests = []
        self.lock = threading.Lock()
        self._async_lock = None # Lazy init to avoid loop issues
    
    @property
    def async_lock(self):
        if self._async_lock is None:
            self._async_lock = asyncio.Lock()
        return self._async_lock

    def wait_if_needed(self):
        with self.lock:
            now = time.time()
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
                if wait_time > 0:
                    print(f"{Fore.YELLOW}⏱️  Rate limit: aguardando {wait_time:.1f}s...")
                    time.sleep(wait_time)
                self.requests = [t for t in self.requests if time.time() - t < 60]
            
            self.requests.append(time.time())

    async def wait_if_needed_async(self):
        """Versão async do rate limiter com asyncio.Lock"""
        wait_time = 0
        
        # Use async lock to protect shared state access in async context
        async with self.async_lock:
            now = time.time()
            # Clean old requests
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
            
            # If we don't need to wait, register now
            if wait_time <= 0:
                 self.requests.append(time.time())
        
        # If we need to wait, release lock, wait, then re-acquire to register
        if wait_time > 0:
            print(f"{Fore.YELLOW}⏱️  Rate limit (Async): aguardando {wait_time:.1f}s...")
            await asyncio.sleep(wait_time)
            async with self.async_lock:
                self.requests.append(time.time())

# Instância global
rate_limiter = RateLimiter(max_requests_per_minute=60)

# =============================================================================
# AGENT HELPERS (v3.1 - Key Documents Selection)
# =============================================================================

def _safe_parse_date(date_str: str) -> str:
    # Mantém string (YYYY-MM-DD) ou vazio; ordenação lexicográfica funciona p/ ISO.
    return date_str or ""

def select_key_documents(
    local_index,
    max_docs: int = 8,
    prefer_types: Optional[List[str]] = None
) -> List[Dict[str, Any]]:
    """
    Seleciona PDFs-chave de forma determinística a partir do índice local.
    Retorna lista de dicts: {doc_id, tipo_doc, data_doc, file_path, filename}.
    """
    if not local_index:
        return []

    prefer_types = prefer_types or [
        "sentenca", "acordao", "decisao", "despacho",
        "laudo", "peticao_inicial", "contestacao", "recurso", "parecer"
    ]

    # Map doc_id -> melhor file_path (primeiro encontrado)
    docid_to_path = {}
    docid_to_filename = {}
    for meta in getattr(local_index, "_metadatas", []):
        doc_id = meta.get("doc_id", "")
        fp = meta.get("file_path", "")
        fn = meta.get("filename", "")
        if doc_id and fp and doc_id not in docid_to_path:
            docid_to_path[doc_id] = fp
            docid_to_filename[doc_id] = fn

    # Cronologia agregada por doc_id (já ordenada por data asc no rag_local)
    docs = local_index.cronologia()

    # Ordena por (prioridade tipo) e depois por data desc (mais recente primeiro)
    type_rank = {t: i for i, t in enumerate(prefer_types)}
    def sort_key(d):
        t = d.get("tipo_doc", "")
        rank = type_rank.get(t, 999)
        data = _safe_parse_date(d.get("data_doc", ""))
        # data desc => usa negativo via tuple invertida (ou simplesmente ordena depois)
        return (rank, data)

    # primeiro por rank e data asc, depois inverte por data dentro do mesmo rank
    docs_sorted = sorted(docs, key=sort_key)

    selected = []
    used_paths = set()

    # Para cada tipo preferido, pega o mais recente disponível (varrendo do fim)
    for tipo in prefer_types:
        candidates = [d for d in docs_sorted if d.get("tipo_doc") == tipo]
        if not candidates:
            continue
        # mais recente = maior data_doc
        candidates = sorted(candidates, key=lambda x: _safe_parse_date(x.get("data_doc","")), reverse=True)
        for c in candidates:
            doc_id = c.get("doc_id", "")
            fp = docid_to_path.get(doc_id, "")
            if fp and fp.lower().endswith(".pdf") and fp not in used_paths:
                selected.append({
                    "doc_id": doc_id,
                    "tipo_doc": c.get("tipo_doc", ""),
                    "data_doc": c.get("data_doc", ""),
                    "file_path": fp,
                    "filename": docid_to_filename.get(doc_id, c.get("filename",""))
                })
                used_paths.add(fp)
                break
        if len(selected) >= max_docs:
            break

    return selected[:max_docs]


def build_key_docs_text_pack(
    local_index,
    selected_docs: List[Dict[str, Any]],
    max_chars: int = 50000,
    snippet_chars: int = 700
) -> str:
    """
    Monta um text_pack estável com 1–2 trechos por doc_id usando os próprios chunks do índice.
    """
    if not local_index or not selected_docs:
        return ""

    # Junta chunks por doc_id e pega menor página (primeiros trechos)
    chunks_by_doc = {}
    for doc_text, meta in zip(getattr(local_index, "_documents", []), getattr(local_index, "_metadatas", [])):
        doc_id = meta.get("doc_id", "")
        if not doc_id:
            continue
        chunks_by_doc.setdefault(doc_id, []).append((meta.get("pagina", 0), doc_text, meta))

    lines = []
    for d in selected_docs:
        doc_id = d["doc_id"]
        tipo = (d.get("tipo_doc") or "DOC").upper()
        items = chunks_by_doc.get(doc_id, [])
        if not items:
            continue
        items.sort(key=lambda x: x[0])  # menor página primeiro
        # pega 1–2 trechos iniciais
        for (pagina, txt, meta) in items[:2]:
            r = {"text": txt, "metadata": meta, "final_score": 1.0}
            cit = local_index.format_citation(r)
            snippet = (txt or "").strip().replace("\n", " ")
            snippet = snippet[:snippet_chars] + ("..." if len(snippet) > snippet_chars else "")
            lines.append(f"- {cit}: \"{snippet}\"")

    out = "### 📌 DOCUMENTOS-CHAVE (TRECHOS) \n" + "\n".join(lines) + "\n"
    return out if len(out) <= max_chars else out[:max_chars] + "\n\n[... text_pack truncado ...]\n"

# =============================================================================
# TEXT PROCESSING HELPERS (Ported from mlx_vomo.py)
# =============================================================================

def remover_eco_do_contexto(resposta_api: str, contexto_enviado: str) -> str:
    """Remove o início da resposta se for apenas um 'eco' do final do contexto."""
    if not contexto_enviado or not resposta_api:
        return resposta_api

    final_contexto = contexto_enviado.strip()[-300:]
    inicio_resposta = resposta_api.strip()[:300]

    matcher = difflib.SequenceMatcher(None, final_contexto, inicio_resposta)
    match = matcher.find_longest_match(0, len(final_contexto), 0, len(inicio_resposta))

    if match.size > 50:
        print(f"{Fore.YELLOW}✂️ Eco detectado! Removendo {match.size} chars repetidos.{Style.RESET_ALL}")
        return resposta_api.strip()[match.size:].strip()
    
    return resposta_api

def titulos_sao_similares(t1: str, t2: str, threshold: float = 0.90) -> bool:
    """Verifica se dois títulos são semanticamente iguais (fuzzy matching)."""
    def normalizar(t):
        return re.sub(r'[^a-z0-9 ]', '', t.lower())
    
    nt1 = normalizar(t1)
    nt2 = normalizar(t2)
    
    if not nt1 or not nt2: return False
    
    # Proteção: Diferença de tamanho
    nt1_compact = nt1.replace(' ', '')
    nt2_compact = nt2.replace(' ', '')
    len_ratio = min(len(nt1_compact), len(nt2_compact)) / max(len(nt1_compact), len(nt2_compact)) if max(len(nt1_compact), len(nt2_compact)) > 0 else 0
    if len_ratio < 0.8: return False
    
    # Proteção: Palavras exclusivas
    palavras1 = set(nt1.split())
    palavras2 = set(nt2.split())
    diferenca = palavras1.symmetric_difference(palavras2)
    if any(len(w) > 3 for w in diferenca): return False
        
    return difflib.SequenceMatcher(None, nt1_compact, nt2_compact).ratio() > threshold

def limpar_inicio_redundante(texto_novo: str, texto_acumulado: str) -> str:
    """Remove título no início do novo chunk se similar ao do acumulado."""
    if not texto_acumulado.strip(): return texto_novo
    
    # Pega último título do acumulado
    match_acum = re.search(r'^(#{1,4}\s+.+?)$', texto_acumulado.strip().split('\n')[-5:][0] if texto_acumulado.strip() else '', re.MULTILINE)
    # Pega primeiro título do novo
    match_novo = re.search(r'^(#{1,4}\s+.+?)$', texto_novo.strip()[:500], re.MULTILINE)
    
    if match_acum and match_novo:
        titulo_acum = match_acum.group(1)
        titulo_novo = match_novo.group(1)
        if titulos_sao_similares(titulo_acum, titulo_novo):
            # Remove o título duplicado do início
            print(f"{Fore.YELLOW}✂️ Título duplicado removido: {titulo_novo[:50]}...{Style.RESET_ALL}")
            return texto_novo.replace(titulo_novo, '', 1).strip()
    
    return texto_novo

# =============================================================================
# DEDUPLICATION HELPERS (v3.0 - Ported from mlx_vomo.py)
# =============================================================================

def detectar_secoes_duplicadas(texto):
    """v3.0: Detecta seções duplicadas por títulos (Fuzzy Matching)"""
    linhas = texto.split('\n')
    titulos_vistos = []
    secoes_duplicadas = []
    
    for i, linha in enumerate(linhas):
        linha_strip = linha.strip()
        if linha_strip.startswith('##'):
            titulo_normalizado = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', linha_strip)
            titulo_normalizado = re.sub(r'[📋📊🗂]', '', titulo_normalizado).strip()
            titulo_para_comparar = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_normalizado, flags=re.IGNORECASE).strip()
            
            duplicado = False
            for t_visto, linha_visto in titulos_vistos:
                if titulos_sao_similares(titulo_para_comparar, t_visto):
                    secoes_duplicadas.append({
                        'titulo': titulo_normalizado,
                        'primeira_linha': linha_visto,
                        'duplicada_linha': i
                    })
                    duplicado = True
                    break
            
            if not duplicado:
                titulos_vistos.append((titulo_para_comparar, i))
    
    return secoes_duplicadas

def remover_secoes_duplicadas(texto, limiar=0.60):
    """v3.0: Remove seções duplicadas com comparação de conteúdo"""
    secoes_dup = detectar_secoes_duplicadas(texto)
    if not secoes_dup: return texto
    
    print(f"{Fore.YELLOW}🧹 Removendo {len(secoes_dup)} seções duplicadas...{Style.RESET_ALL}")
    linhas = texto.split('\n')
    linhas_para_remover = set()
    
    for dup in secoes_dup:
        idx_dup = dup['duplicada_linha']
        header_dup = linhas[idx_dup].strip()
        match_dup = re.match(r'^(#+)', header_dup)
        nivel_dup = len(match_dup.group(1)) if match_dup else 2
        
        # Find end of section
        fim_dup_idx = len(linhas)
        for i in range(idx_dup + 1, len(linhas)):
            line = linhas[i].strip()
            if line.startswith('#'):
                match_now = re.match(r'^(#+)', line)
                if match_now and len(match_now.group(1)) <= nivel_dup:
                    fim_dup_idx = i
                    break
        
        # Mark lines for removal
        for i in range(idx_dup, fim_dup_idx):
            linhas_para_remover.add(i)
    
    resultado = [linha for i, linha in enumerate(linhas) if i not in linhas_para_remover]
    print(f"{Fore.GREEN}   ✅ {len(linhas_para_remover)} linhas removidas{Style.RESET_ALL}")
    return '\n'.join(resultado)

def remover_paragrafos_duplicados(texto: str, min_chars: int = 80, threshold: float = 0.90) -> str:
    """v3.1: Remove parágrafos duplicados (Híbrido: Hash + Semântico)"""
    import unicodedata
    
    blocos = re.split(r'\n\s*\n', texto)
    if not blocos:
        return texto
        
    print(f"{Fore.CYAN}🧹 Iniciando deduplicação de {len(blocos)} parágrafos...{Style.RESET_ALL}")
    
    # FASE 1: Deduplicação exata (Hash) - Rápido
    blocos_unicos = []
    vistos_hash = set()
    
    for bloco in blocos:
        if len(bloco.strip()) < min_chars:
            blocos_unicos.append(bloco)
            continue
            
        # Normalize for hash
        norm = unicodedata.normalize('NFKD', bloco.lower()).encode('ASCII', 'ignore').decode('ASCII')
        norm = re.sub(r'\s+', ' ', norm).strip()
        h = hash(norm)
        
        if h not in vistos_hash:
            vistos_hash.add(h)
            blocos_unicos.append(bloco)
    
    # FASE 2: Deduplicação Semântica (Embeddings) - Se disponível
    if HAS_RAG_DEDUPE and get_deduplicator():
        print(f"{Fore.CYAN}   🧠 Executando deduplicação semântica (Embeddings)...{Style.RESET_ALL}")
        
        candidatos_idx = [i for i, b in enumerate(blocos_unicos) if len(b) >= min_chars]
        candidatos_texto = [blocos_unicos[i] for i in candidatos_idx]
        
        if candidatos_texto:
            deduplicator = get_deduplicator()
            indices_validos = deduplicator.deduplicate(candidatos_texto, threshold=threshold)
            
            # Map back
            indices_preservados = set([candidatos_idx[k] for k in indices_validos])
            
            blocos_finais = []
            removidos_sem = 0
            
            for i, b in enumerate(blocos_unicos):
                if i in candidatos_idx:
                    if i in indices_preservados:
                        blocos_finais.append(b)
                    else:
                        removidos_sem += 1
                else:
                    blocos_finais.append(b)
                
            if removidos_sem > 0:
                print(f"{Fore.GREEN}   ✅ {removidos_sem} duplicatas semânticas removidas{Style.RESET_ALL}")
            
            return '\n\n'.join(blocos_finais)

    return '\n\n'.join(blocos_unicos)

# =============================================================================
# ADAPTIVE RAG - ROUTING (v4.0)
# =============================================================================

# Strategy constants
STRATEGY_LOCAL_ONLY = "LOCAL_ONLY"       # Search only in local process documents
STRATEGY_GLOBAL_SINGLE = "GLOBAL_SINGLE"  # Search only in global bases (lei, juris, etc.)
STRATEGY_HYBRID = "HYBRID"                # Search both local and global
STRATEGY_NO_RETRIEVAL = "NO_RETRIEVAL"    # Skip RAG entirely (simple/template sections)
# v5.0: Advanced strategies
STRATEGY_HYDE = "HYDE"                    # Use HyDE for complex/abstract queries
STRATEGY_GRAPH = "GRAPH"                  # Use GraphRAG for multi-hop reasoning

def default_route_for_section(section_title: str, tipo_peca: str = "") -> Dict[str, Any]:
    """
    Heuristic-based router that decides RAG strategy based on section title.
    
    Returns a dict with:
        - strategy: one of STRATEGY_* constants
        - sources: list of RAG sources to query (e.g., ['lei', 'juris'])
        - top_k: number of results to fetch
        - bm25_weight: weight for keyword matching (0.0-1.0)
        - semantic_weight: weight for semantic search (0.0-1.0)
        - needs_query_expansion: bool, whether to expand query with LLM
    """
    title_lower = section_title.lower()
    
    # Default config
    config = {
        "strategy": STRATEGY_HYBRID,
        "sources": ["lei", "juris", "pecas_modelo"],
        "top_k": 8,
        "bm25_weight": 0.4,
        "semantic_weight": 0.6,
        "needs_query_expansion": True,
        "reason": "default hybrid"
    }
    
    # === PATTERN MATCHING FOR LEGAL SECTIONS ===
    
    # 1. LOCAL_ONLY: Factual sections that require process-specific info
    local_patterns = [
        r"(dos?\s+)?fatos?",
        r"(da\s+)?narrativa",
        r"síntese\s*(da|do)?\s*(inicial|proceso|fatos)?",
        r"relatório",
        r"qualificação\s*(das?\s+partes)?",
        r"histórico",
        r"(do\s+)?caso",
    ]
    for pattern in local_patterns:
        if re.search(pattern, title_lower):
            return {
                "strategy": STRATEGY_LOCAL_ONLY,
                "sources": [],  # Only local process docs
                "top_k": 5,
                "bm25_weight": 0.6,
                "semantic_weight": 0.4,
                "needs_query_expansion": False,
                "reason": f"matched local pattern: {pattern}"
            }
    
    # 2. HYDE: Complex fundamentação (abstract concepts - HyDE excels here)
    hyde_patterns = [
        r"fundament(o|ação)\s*(jurídica)?",
        r"tese(s)?\s+(central|principal|jurídica)",
        r"doutrina",
        r"teoria\s*(geral|do)?",
    ]
    for pattern in hyde_patterns:
        if re.search(pattern, title_lower):
            return {
                "strategy": STRATEGY_HYDE,
                "sources": ["lei", "juris", "pecas_modelo"],
                "top_k": 10,
                "bm25_weight": 0.3,
                "semantic_weight": 0.7,
                "needs_query_expansion": False,  # HyDE replaces query expansion
                "reason": f"matched HyDE pattern: {pattern}"
            }
    
    # 3. GRAPH: Jurisprudência sections (multi-hop reasoning: precedentes → súmulas → teses)
    graph_patterns = [
        r"juris(prudência)?",
        r"súmula(s)?",
        r"precedent(e|es)",
        r"entendimento\s+(do|da)\s+(st[fj]|tst|tribunal)",
    ]
    for pattern in graph_patterns:
        if re.search(pattern, title_lower):
            return {
                "strategy": STRATEGY_GRAPH,
                "sources": ["juris"],
                "top_k": 8,
                "bm25_weight": 0.4,
                "semantic_weight": 0.6,
                "needs_query_expansion": True,
                "graph_hops": 2,  # How deep to traverse relationships
                "reason": f"matched Graph pattern: {pattern}"
            }
    
    # 4. GLOBAL_SINGLE: Simple legal doctrine (lei, direito)
    global_patterns = [
        r"(do\s+)?direito",
        r"legislação",
        r"mérito",
    ]
    for pattern in global_patterns:
        if re.search(pattern, title_lower):
            return {
                "strategy": STRATEGY_GLOBAL_SINGLE,
                "sources": ["lei", "juris"],
                "top_k": 10,
                "bm25_weight": 0.5,
                "semantic_weight": 0.5,
                "needs_query_expansion": True,
                "reason": f"matched global pattern: {pattern}"
            }
    
    # 3. NO_RETRIEVAL: Procedural/template sections
    no_rag_patterns = [
        r"(dos?\s+)?pedidos?",
        r"(do\s+)?valor\s*(da\s+causa)?",
        r"(opção\s*(por)?\s*)?(audiência|conciliação)",
        r"endereçamento",
        r"conclusão",
        r"fecho",
        r"requerimentos?\s*finais?",
        r"tempestividade",
        r"preparo",
    ]
    for pattern in no_rag_patterns:
        if re.search(pattern, title_lower):
            return {
                "strategy": STRATEGY_NO_RETRIEVAL,
                "sources": [],
                "top_k": 0,
                "bm25_weight": 0,
                "semantic_weight": 0,
                "needs_query_expansion": False,
                "reason": f"matched no-rag pattern: {pattern}"
            }
    
    # 4. HYBRID for "Provas" sections (needs local facts + procedural law)
    if re.search(r"provas?\s*(a\s+produzir)?", title_lower):
        return {
            "strategy": STRATEGY_HYBRID,
            "sources": ["lei"],  # Minimal law refs for procedure
            "top_k": 5,
            "bm25_weight": 0.5,
            "semantic_weight": 0.5,
            "needs_query_expansion": False,
            "reason": "provas section - hybrid minimal"
        }
    
    # 5. Return default for unmatched sections
    return config


def crag_gate_retrieve(
    query: str,
    rag_manager,
    sources: List[str],
    top_k: int = 8,
    bm25_weight: float = 0.4,
    semantic_weight: float = 0.6,
    tenant_id: str = "default",
    user_id: str = None,
    tipo_peca_filter: str = None,
    min_best_score: float = 0.45,
    min_avg_top3_score: float = 0.35,
    max_retries: int = 2,
    verbose: bool = False
) -> Dict[str, Any]:
    """
    CRAG (Corrective RAG) Gate: Retrieves results and validates quality before passing to generation.
    
    If quality is below thresholds, attempts corrective actions:
    1. Retry with higher top_k and BM25 weight
    2. If still failing, return with safe_mode=True to trigger conservative generation
    
    Returns:
        {
            "results": List[Dict],  # RAG results
            "gate_passed": bool,    # True if quality thresholds met
            "safe_mode": bool,      # True if generation should be extra conservative
            "best_score": float,    # Best chunk score
            "avg_top3": float,      # Avg of top 3 scores
            "attempts": int,        # Number of retrieval attempts
            "reason": str           # Explanation
        }
    """
    gate_info = {
        "results": [],
        "gate_passed": False,
        "safe_mode": False,
        "best_score": 0.0,
        "avg_top3": 0.0,
        "attempts": 0,
        "reason": ""
    }
    
    if not sources or not rag_manager:
        gate_info["reason"] = "No sources or RAG manager, skipping gate"
        gate_info["gate_passed"] = True  # Pass-through if no RAG
        return gate_info
    
    def evaluate_quality(results: List[Dict]) -> tuple:
        """Calculate quality metrics from results."""
        if not results:
            return 0.0, 0.0
        scores = [r.get('final_score', 0) for r in results]
        best = max(scores) if scores else 0.0
        top3 = scores[:3]
        avg_top3 = sum(top3) / len(top3) if top3 else 0.0
        return best, avg_top3
    
    current_top_k = top_k
    current_bm25 = bm25_weight
    
    for attempt in range(max_retries + 1):
        gate_info["attempts"] = attempt + 1
        
        try:
            results = rag_manager.hybrid_search(
                query=query,
                sources=sources,
                top_k=current_top_k,
                bm25_weight=current_bm25,
                semantic_weight=max(0.0, 1.0 - current_bm25),
                user_id=user_id,
                tenant_id=tenant_id,
                tipo_peca_filter=tipo_peca_filter,
            )
        except Exception as e:
            logger.warning(f"CRAG Gate: Retrieval error attempt {attempt+1}: {e}")
            results = []
        
        gate_info["results"] = results
        best, avg_top3 = evaluate_quality(results)
        gate_info["best_score"] = best
        gate_info["avg_top3"] = avg_top3
        
        if verbose:
            logger.info(f"  CRAG Gate (attempt {attempt+1}): best={best:.2f}, avg_top3={avg_top3:.2f}, threshold_best={min_best_score}, threshold_avg={min_avg_top3_score}")
        
        # Check if quality is acceptable
        if best >= min_best_score and avg_top3 >= min_avg_top3_score:
            gate_info["gate_passed"] = True
            gate_info["reason"] = f"Quality OK: best={best:.2f} >= {min_best_score}, avg={avg_top3:.2f} >= {min_avg_top3_score}"
            return gate_info
        
        # Corrective action: increase top_k and BM25 weight
        if attempt < max_retries:
            current_top_k = min(current_top_k + 4, 20)  # Cap at 20
            current_bm25 = min(current_bm25 + 0.15, 0.8)  # Boost keyword matching
            if verbose:
                logger.info(f"  CRAG Gate: Quality low, retrying with top_k={current_top_k}, bm25_weight={current_bm25}")
    
    # All retries exhausted
    gate_info["safe_mode"] = True
    gate_info["reason"] = f"Quality below threshold after {gate_info['attempts']} attempts: best={best:.2f} < {min_best_score} or avg={avg_top3:.2f} < {min_avg_top3_score}"
    
    if verbose:
        logger.warning(f"  ⚠️ CRAG Gate: Entering SAFE MODE - {gate_info['reason']}")
    
    return gate_info


# Safe Mode prompt injection for CRAG
SAFE_MODE_INSTRUCTION = """
⚠️ MODO SEGURO ATIVADO: A base de conhecimento retornou resultados de baixa confiança para esta seção.

REGRAS IMPERATIVAS:
1. NÃO invente citações de leis, súmulas ou jurisprudência.
2. Se usar uma referência legal, MARQUE com [VERIFICAR]: "Art. X da Lei Y [VERIFICAR]"
3. Prefira argumentação principiológica e doutrinária genérica.
4. Use expressões como "conforme entendimento majoritário" ao invés de citar decisões específicas.
5. Ao final da seção, adicione: "[Nota: Esta seção requer revisão manual das citações]"
"""

# =============================================================================
# METRICS COLLECTOR (v2.2)
# =============================================================================

class MetricsCollector:
    """Coleta e reporta métricas de execução para otimização de custos."""
    def __init__(self):
        self.reset()
    
    def reset(self):
        self.api_calls = 0
        self.total_prompt_tokens = 0
        self.total_cached_tokens = 0
        self.total_completion_tokens = 0
        self.total_time_seconds = 0.0
        self.start_time = None
    
    def start_timer(self):
        self.start_time = time.time()
    
    def stop_timer(self):
        if self.start_time:
            self.total_time_seconds += time.time() - self.start_time
            self.start_time = None
    
    def record_usage(self, usage_metadata):
        """Record usage from Gemini SDK response"""
        if not usage_metadata: return
        self.api_calls += 1
        self.total_prompt_tokens += (usage_metadata.prompt_token_count or 0) - (usage_metadata.cached_content_token_count or 0)
        self.total_cached_tokens += (usage_metadata.cached_content_token_count or 0)
        self.total_completion_tokens += (usage_metadata.candidates_token_count or 0)
    
    def record_call(self, provider: str, prompt_tokens: int, completion_tokens: int, duration: float):
        """Record a call with explicit values (compatible with mlx_vomo v2.2)"""
        self.api_calls += 1
        self.total_prompt_tokens += prompt_tokens
        self.total_completion_tokens += completion_tokens
        self.total_time_seconds += duration

    def get_report(self):
        # Preços (Estimativa Gemini 1.5 Pro)
        input_price = 1.25   # / 1M
        cached_price = 0.30 
        output_price = 5.00
        
        cost = (
            (self.total_prompt_tokens * input_price) + 
            (self.total_cached_tokens * cached_price) +
            (self.total_completion_tokens * output_price)
        ) / 1_000_000
        
        return f"""
        📊 MÉTRICAS JURIDICO AI
        -----------------------
        Chamadas API: {self.api_calls}
        Tokens Input: {self.total_prompt_tokens:,}
        Tokens Cache: {self.total_cached_tokens:,}
        Tokens Output: {self.total_completion_tokens:,}
        Tempo Total:  {self.total_time_seconds:.1f}s
        Custo Est.:   ${cost:.4f}
        """

metrics = MetricsCollector()

# Safety Settings (Permissive for Legal Context - types-based for new SDK)
SAFETY_SETTINGS = [
    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
]

# =============================================================================
# CLASS: LEGAL DRAFTER (v3.0 - Vertex AI + Gemini 3 Flash)
# =============================================================================

class LegalDrafter:
    def __init__(self, model_name=None, cache_ttl_mins=60, section_template=TEMPLATE_SECAO):
        self.model_name = model_name or DEFAULT_MODEL
        
        self.cache_ttl_mins = cache_ttl_mins
        
        # Initialize Anthropic Client (Fallback)
        self.anthropic_key = os.getenv("ANTHROPIC_API_KEY")
        self.anthropic_client = anthropic.Anthropic(api_key=self.anthropic_key) if (anthropic and self.anthropic_key) else None
        if self.anthropic_client:
            print(f"{Fore.GREEN}   ✅ Anthropic Fallback configurado.")

        self.cache = None
        self.files = []
        self.section_template = section_template
        
        # Initialize Vertex AI Client
        print(f"{Fore.CYAN}☁️  Conectando via Vertex AI (Global Endpoint)...")
        self.client = genai.Client(
            vertexai=True,
            project=PROJECT_ID,
            location="global"
        )
        # Quick test
        try:
            self.client.models.count_tokens(model=self.model_name, contents="test")
            print(f"{Fore.GREEN}   ✅ Conectado ao Vertex AI com sucesso ({self.model_name}).")
        except Exception as e:
            print(f"{Fore.RED}   ❌ Falha ao conectar: {e}")
            raise e

    def _generate_with_retry(self, prompt, config=None, stream=False, thinking_level="LOW"):
        """Wrapper for API calls with robust retry logic (v3.0 - Vertex AI + ThinkingConfig)"""
        max_retries = 3
        base_delay = 2
        
        # Build generation config
        gen_config_kwargs = {
            "max_output_tokens": config.get("max_output_tokens", 8192) if config else 8192,
            "temperature": config.get("temperature", 0.2) if config else 0.2,
            "safety_settings": SAFETY_SETTINGS,
            "cached_content": self.cache.name if self.cache else None, # Injetar cache se houver
        }

        # thinking_config is optional and may not be supported in all SDK versions
        thinking_config = None
        try:
            thinking_config = types.ThinkingConfig(
                include_thoughts=False,
                thinking_level=thinking_level
            )
        except Exception:
            try:
                thinking_config = types.ThinkingConfig(include_thoughts=False)
            except Exception:
                thinking_config = None

        if thinking_config:
            gen_config_kwargs["thinking_config"] = thinking_config

        gen_config = types.GenerateContentConfig(**gen_config_kwargs)
        
        # Override for JSON response
        if config and config.get("response_mime_type"):
            gen_config.response_mime_type = config.get("response_mime_type")
        
        for attempt in range(max_retries):
            try:
                # Rate limiting antes de cada chamada
                rate_limiter.wait_if_needed()
                
                metrics.start_timer()
                if stream:
                    return self.client.models.generate_content_stream(
                        model=self.model_name,
                        contents=prompt,
                        config=gen_config
                    )
                else:
                    response = self.client.models.generate_content(
                        model=self.model_name,
                        contents=prompt,
                        config=gen_config
                    )
                    metrics.stop_timer()
                    
                    if hasattr(response, 'usage_metadata'):
                        metrics.record_usage(response.usage_metadata)
                    
                    return response
                    
            except Exception as e:
                metrics.stop_timer()
                print(f"{Fore.YELLOW}⚠️ Erro API (Tentativa {attempt+1}/{max_retries}): {e}")
                
                # Check for Fallback to Anthropic
                if attempt == max_retries - 1 and self.anthropic_client:
                    print(f"{Fore.MAGENTA}🔄 Acionando FALLBACK Anthropic (Claude 3.5 Sonnet)...")
                    try:
                        metrics.start_timer()
                        claude_response = self.anthropic_client.messages.create(
                            model="claude-3-5-sonnet-20241022",
                            max_tokens=config.get("max_output_tokens", 8192) if config else 8192,
                            temperature=config.get("temperature", 0.2) if config else 0.2,
                            system="Você é um Assistente Jurídico de elite. Responda com precisão técnica e formalidade.",
                            messages=[{"role": "user", "content": prompt}]
                        )
                        metrics.stop_timer()
                        
                        # Mock the response object to look like Gemini response for compatibility
                        class MockResponse:
                            def __init__(self, text):
                                self.text = text
                        
                        return MockResponse(claude_response.content[0].text)
                    except Exception as ae:
                        print(f"{Fore.RED}❌ Falha no Fallback Anthropic: {ae}")

                if "429" in str(e) or "503" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                    wait_time = base_delay * (2 ** attempt)
                    print(f"{Fore.YELLOW}   ⏱️ Aguardando {wait_time}s (backoff exponencial)...")
                    time.sleep(wait_time)
                else:
                    if attempt == max_retries - 1: raise e
                    time.sleep(base_delay)
        return None


    def upload_files(self, file_paths: List[str]):
        """Faz upload de arquivos locais para o Gemini (v3.0 - Vertex AI)"""
        print(f"{Fore.YELLOW}📂 Enviando {len(file_paths)} arquivos para o Gemini...")
        uploaded_files = []
        for path in file_paths:
            if not os.path.exists(path):
                print(f"{Fore.RED}⚠️ Arquivo não encontrado: {path}")
                continue
            
            print(f"   ⬆️ Uploading: {Path(path).name}...")
            try:
                # New SDK: client.files.upload
                f = self.client.files.upload(file=path)
                
                # Wait for processing (new SDK uses different state check)
                while hasattr(f, 'state') and f.state.name == "PROCESSING":
                    time.sleep(1)
                    f = self.client.files.get(name=f.name)
                
                uploaded_files.append(f)
                print(f"{Fore.GREEN}   ✅ Pronto: {Path(path).name}")
            except Exception as e:
                print(f"{Fore.RED}❌ Falha no upload de {path}: {e}")
        
        self.files.extend(uploaded_files)

    def create_context_cache(self, system_instruction: str = "Você é um assistente jurídico. Use os arquivos como verdade."):
        """Cria o cache de contexto com os arquivos carregados (v3.0 - Vertex AI)"""
        if not self.files:
            print(f"{Fore.YELLOW}⚠️ Nenhum arquivo para cache. Usando modelo padrão.")
            return

        print(f"{Fore.YELLOW}🧠 Criando Cache de Contexto Jurídico (TTL={self.cache_ttl_mins}m)...")
        try:
            # New SDK: client.caches.create
            ttl_seconds = f"{self.cache_ttl_mins * 60}s"
            self.cache = self.client.caches.create(
                model=self.model_name,
                display_name="juridico_cache_" + datetime.datetime.now().strftime("%Y%m%d_%H%M"),
                system_instruction=system_instruction,
                contents=[types.Content(parts=[types.Part.from_uri(file_uri=f.uri, mime_type=f.mime_type) for f in self.files])],
                ttl=ttl_seconds
            )
            print(f"{Fore.GREEN}✅ Cache ATIVO: {self.cache.name}")
            if hasattr(self.cache, 'usage_metadata') and self.cache.usage_metadata:
                print(f"{Fore.CYAN}   ℹ️ Tokens no cache: {self.cache.usage_metadata.total_token_count}")
        except Exception as e:
            print(f"{Fore.RED}❌ Erro ao criar cache: {e}")
            print(f"{Fore.YELLOW}   Continuando sem cache...")

    def generate_rag_query(self, section_title: str, resumo_caso: str, max_chars: int = 800) -> str:
        """
        Gera uma consulta expandida (ou hipotetica) para recuperar fontes RAG.

        Nota: Quando usado com HyDE, este texto serve como "documento hipotetico"
        para alinhar embeddings e melhorar a busca semantica.
        """
        base_query = f"{section_title}. Contexto: {resumo_caso[:500]}"
        prompt = f"""
Você esta ajudando um motor de busca juridico (RAG). Gere um texto curto e rico em termos
para recuperar fontes relevantes para a secao abaixo. Regras:
- Responda com um unico paragrafo, sem listas ou titulos.
- Inclua termos juridicos, sinonimos e possiveis fundamentos (sem inventar fatos).
- Nao cite numeros de artigos/precedentes se nao estiverem no contexto.
- Limite a resposta a ~{max_chars} caracteres.

SECAO: {section_title}
CONTEXTO DO CASO: {resumo_caso}
"""
        try:
            response = self._generate_with_retry(
                prompt,
                config={"max_output_tokens": 300, "temperature": 0.2},
                thinking_level="LOW"
            )
            if not response or not getattr(response, "text", None):
                return base_query

            text = response.text.strip()
            if "```" in text:
                import re
                match = re.search(r"```(?:[a-zA-Z0-9_-]+)?\n([\s\S]*?)```", text)
                if match:
                    text = match.group(1).strip()

            if not text:
                return base_query

            return text[:max_chars].strip()
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Falha ao gerar query RAG: {e}")
            return base_query


    def generate_outline(self, tipo_peca: str, resumo_caso: str, tese_usuario: str) -> List[str]:
        """Gera o sumário estruturado da peça"""
        print(f"{Fore.CYAN}📑 Planejando estrutura (Outline)...")
        
        # 1. Select Role Instruction from Map
        role_instruction = PROMPT_MAP.get(tipo_peca, PROMPT_MAP["PETICAO_INICIAL"])
        if tipo_peca not in PROMPT_MAP:
             print(f"{Fore.YELLOW}⚠️ Modo '{tipo_peca}' desconhecido. Usando PETICAO_INICIAL como base.")

        # 2. Render Prompt
        t = Template(TEMPLATE_OUTLINE_BASE)
        prompt = t.render(
            role_instruction=role_instruction,
            resumo_caso=resumo_caso, 
            tese_usuario=tese_usuario
        )
        
        try:
            # v3.0: Use HIGH thinking for outline planning (complex reasoning)
            response = self._generate_with_retry(
                prompt,
                config={"response_mime_type": "application/json"},
                thinking_level="HIGH"
            )
            if not response: return []
            
            # Robust JSON parsing with fallback
            try:
                outline = json.loads(response.text)
            except json.JSONDecodeError:
                # Tenta extrair JSON de dentro do texto
                import re
                json_match = re.search(r'\[.*\]', response.text, re.DOTALL)
                if json_match:
                    outline = json.loads(json_match.group())
                else:
                    print(f"{Fore.RED}❌ Resposta não contém JSON válido.")
                    print(f"   Texto recebido: {response.text[:200]}...")
                    return []
            
            if not isinstance(outline, list):
                print(f"{Fore.RED}❌ Outline não é uma lista.")
                return []
                
            print(f"{Fore.GREEN}✅ Outline Gerado com {len(outline)} seções.")
            for item in outline:
                print(f"   - {item}")
            return outline
        except Exception as e:
            print(f"{Fore.RED}❌ Erro ao gerar outline: {e}")
            return []

    def stream_section_updates(self, prompt: str, target_words: int = 1000) -> Generator[Dict, None, str]:
        """Gera texto da seção com streaming e updates parciais (20/40/60/80%)"""
        
        # Simula thresholds baseados em contagem de palavras (estimativa)
        thresholds = [int(target_words * 0.2), int(target_words * 0.4), int(target_words * 0.6), int(target_words * 0.8)]
        next_th_idx = 0
        
        full_text = ""
        last_sent_len = 0
        
        try:
            # Use retry wrapper but force stream=True
            response_stream = self._generate_with_retry(prompt, stream=True)
            if not response_stream: return ""

            for chunk in response_stream:
                if not chunk.text: continue
                
                full_text += chunk.text
                word_count = len(full_text.split())
                
                # Check for thresholds
                if next_th_idx < len(thresholds) and word_count >= thresholds[next_th_idx]:
                    perc = (next_th_idx + 1) * 20
                    yield {
                        "type": "section_progress",
                        "percent": perc,
                        "delta": full_text[last_sent_len:],
                        "word_count": word_count
                    }
                    last_sent_len = len(full_text)
                    next_th_idx += 1
            
            # Record final usage manually since stream usage is in the last chunk usually
            # But the SDK handling of stream usage metadata varies. Simplified here.
            
            # Final update
            yield {
                "type": "section_done",
                "markdown": full_text,
                "word_count": len(full_text.split())
            }
            return full_text
            
        except Exception as e:
            print(f"{Fore.RED}❌ Erro no streaming: {e}")
            return full_text

    def cleanup(self):
        """Limpa o cache para economizar"""
        print(metrics.get_report()) # Print metrics on cleanup
        if self.cache:
            try:
                self.cache.delete()
                print(f"{Fore.YELLOW}🗑️ Cache deletado para economizar custos.")
            except:
                pass


# =============================================================================
# WORD GENERATION (Legal Formatting)
# =============================================================================

def save_as_word_juridico(formatted_text: str, filename: str, output_folder: str, modo: str):
    """Salva o documento com formatação jurídica (ABNT/Forense)"""
    print(f"{Fore.CYAN}📄 Gerando DOCX Jurídico: {filename}...")
    
    doc = Document()
    
    # Configuração de Página (A4, Margens Padrão Peticionamento)
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    
    # Estilos Básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Cabeçalho Dinâmico baseado no modo
    cabecalhos = {
        "PETICAO_INICIAL": "EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA ___ VARA _____ DA COMARCA DE ______",
        "CONTESTACAO": "EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA ___ VARA _____ DA COMARCA DE ______",
        "RECURSO": "EGRÉGIO TRIBUNAL DE JUSTIÇA DO ESTADO DE ______",
        "PARECER": "PARECER JURÍDICO",
        "SENTENCA": "SENTENÇA",
        "CONTRATO": "INSTRUMENTO PARTICULAR DE CONTRATO"
    }
    
    header_text = cabecalhos.get(modo, "DOCUMENTO JURÍDICO")
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(header_text)
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Processamento do Markdown
    lines = formatted_text.split('\n')
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph()
            continue
            
        # Títulos (Outline)
        if line_stripped.startswith('#'):
            level = len(line_stripped) - len(line_stripped.lstrip('#'))
            text = line_stripped.lstrip('#').strip()
            heading_level = min(level, 3)  # Word suporta até Heading 9, mas usamos até 3
            h = doc.add_heading(text, level=heading_level)
            # Ajuste de cor para preto (petições não usam azul)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)
                r.font.name = 'Arial'
            continue
            
        # Citações (Blockquotes) - Jurisprudência
        if line_stripped.startswith('>'):
            quote_text = line_stripped.lstrip('>').strip()
            p = doc.add_paragraph(quote_text)
            p.paragraph_format.left_indent = Cm(4)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.italic = True
            continue
        
        # Tabelas Markdown (simplificado)
        if line_stripped.startswith('|') and '|' in line_stripped[1:]:
            # Skip table processing for now, treat as text
            p = doc.add_paragraph(line_stripped)
            p.paragraph_format.first_line_indent = Cm(0)
            continue
            
        # Parágrafo Normal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(2.0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        
        # Processar formatação inline (negrito/itálico)
        process_inline_formatting(p, line_stripped)

    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    print(f"{Fore.GREEN}✅ Arquivo DOCX salvo: {output_path}")
    return output_path


def process_inline_formatting(paragraph, text: str):
    """Processa **negrito** e *itálico* no texto"""
    import re
    
    # Pattern para **negrito** e *itálico*
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # Adiciona texto antes do match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Verifica se é negrito ou itálico
        if match.group(2):  # **negrito**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *itálico*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        
        last_end = match.end()
    
    # Adiciona texto restante
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

# =============================================================================
# CHAT MODE (v3.1 - Conversação com Documentos)
# =============================================================================

SYSTEM_PROMPT_CHAT = """Você é um Assistente Jurídico Especializado.
Você tem acesso aos documentos do processo/caso que foram carregados no contexto.

Regras:
1. Responda SEMPRE em Português do Brasil.
2. Use linguagem técnica jurídica quando apropriado.
3. Cite trechos dos documentos quando relevante.
4. Se não souber algo, diga claramente.
5. Formato: Markdown para melhor leitura.

Você pode:
- Resumir documentos
- Analisar argumentos das partes
- Identificar pontos fortes e fracos
- Sugerir teses e estratégias
- Comparar jurisprudência
- Responder perguntas específicas sobre o caso
"""

def chat_with_documents(drafter, document_content: str, single_prompt: str = None):
    """
    Modo de chat interativo com os documentos jurídicos.
    Se single_prompt for fornecido, executa apenas uma consulta.
    Caso contrário, entra em loop interativo.
    """
    print(f"\n{Fore.CYAN}{'='*60}")
    print(f"💬 MODO CHAT JURÍDICO (v3.1)")
    print(f"{'='*60}{Style.RESET_ALL}")
    print(f"{Fore.YELLOW}Documentos carregados no contexto. Digite sua pergunta.")
    print(f"Comandos: 'sair' ou 'exit' para encerrar.{Style.RESET_ALL}\n")
    
    # Histórico de conversa para contexto
    conversation_history = []
    
    def send_message(user_message: str) -> str:
        """Envia mensagem e retorna resposta"""
        # Construir prompt com contexto do documento
        full_prompt = f"""{SYSTEM_PROMPT_CHAT}

## Documentos do Caso:
{document_content[:50000]}  # Limita para não estourar contexto

## Histórico da Conversa:
{chr(10).join([f"Usuário: {h['user']}{chr(10)}Assistente: {h['assistant']}" for h in conversation_history[-5:]])}

## Nova Pergunta do Usuário:
{user_message}

## Sua Resposta:
"""
        try:
            response = drafter._generate_with_retry(
                full_prompt,
                config={"max_output_tokens": 8192, "temperature": 0.3},
                thinking_level="HIGH"  # Raciocínio profundo para análise jurídica
            )
            if response and response.text:
                return response.text
            return "⚠️ Não foi possível gerar uma resposta."
        except Exception as e:
            return f"❌ Erro: {e}"
    
    # Modo single prompt
    if single_prompt:
        print(f"{Fore.GREEN}📤 Prompt: {single_prompt}{Style.RESET_ALL}\n")
        response = send_message(single_prompt)
        print(f"{Fore.WHITE}{response}{Style.RESET_ALL}")
        return response
    
    # Modo interativo (loop)
    while True:
        try:
            user_input = input(f"{Fore.GREEN}Você: {Style.RESET_ALL}").strip()
            
            if not user_input:
                continue
            
            if user_input.lower() in ['sair', 'exit', 'quit', 'q']:
                print(f"{Fore.YELLOW}\n👋 Encerrando chat. Até logo!{Style.RESET_ALL}")
                break
            
            print(f"{Fore.CYAN}🤔 Pensando...{Style.RESET_ALL}")
            response = send_message(user_input)
            
            # Adicionar ao histórico
            conversation_history.append({
                "user": user_input,
                "assistant": response[:500] + "..." if len(response) > 500 else response
            })
            
            print(f"\n{Fore.WHITE}🤖 Assistente:{Style.RESET_ALL}")
            print(f"{response}\n")
            
        except KeyboardInterrupt:
            print(f"\n{Fore.YELLOW}👋 Chat interrompido.{Style.RESET_ALL}")
            break
        except EOFError:
            break
    
    return conversation_history

# =============================================================================
# PROGRAMMATIC API (for library import by backend)
# =============================================================================

def generate_document_programmatic(
    input_text: str,
    mode: str = "PETICAO_INICIAL",
    tese: str = "A favor do cliente",
    model: str = None,
    target_pages: int = 0,
    min_pages: int = 0,
    max_pages: int = 0,
    local_files: List[str] = None,
    rag_enabled: bool = True,
    rag_sources: List[str] = None,
    rag_manager = None,
    rag_top_k: int = 8,
    processo_local_path: str = None,
    processo_id: str = None,
    sistema: str = "SEI",
    tenant_id: str = "default",
    agent_mode: bool = False,
    gpt_model: str = "gpt-5.2-chat-latest",
    claude_model: str = "claude-sonnet-4-5@20250929",
    run_audit: bool = True,
    include_toc: bool = False,
    include_section_summaries: bool = False,
    section_template: str = None,
    context_files: List[str] = None, # v3.4: Context Caching
    cache_ttl: int = 60,
    # v4.0: Adaptive RAG & CRAG Gate
    adaptive_routing: bool = False,
    crag_gate: bool = False,
    crag_min_best: float = 0.45,
    crag_min_avg: float = 0.35,
    verbose_rag: bool = False,
    # v4.1: Unification Flags
    approve_outline: bool = False,
    deep_research: bool = False,
    web_search: bool = False,
    reasoning_level: str = "medium",
    # v5.0: GraphRAG & HyDE
    hyde_enabled: bool = False,
    graph_rag_enabled: bool = False,
    graph_hops: int = 2
) -> Dict[str, Any]:
    """
    Programmatic entry point for document generation.
    
    This function allows the backend API to use juridico_gemini.py as a library
    instead of calling it as a subprocess.
    
    Args:
        input_text: The input facts/notes/case description
        mode: Document type (PETICAO_INICIAL, CONTESTACAO, PARECER, etc.)
        tese: Central thesis instruction
        model: Gemini model to use (defaults to gemini-3-flash-preview)
        target_pages: Target page count (0 = no limit)
        min_pages: Minimum page count (0 = auto)
        max_pages: Maximum page count (0 = auto)
        local_files: List of local files to index as RAG Local
        rag_enabled: Enable RAG (retrieval augmented generation)
        rag_sources: List of RAG sources ["lei", "juris", "pecas_modelo"]
        rag_manager: Pre-initialized RAGManager instance (optional)
        rag_top_k: Number of RAG results per section
        processo_local_path: Path to local process folder for RAG Local
        processo_id: Process ID for RAG Local
        sistema: Source system (SEI, PJe, eproc)
        tenant_id: Tenant ID for RBAC
        agent_mode: Enable multi-agent mode (GPT + Claude + Gemini)
        gpt_model: GPT model for agent mode
        claude_model: Claude model for agent mode
        run_audit: Run legal audit after generation
        include_toc: Include table of contents
        include_section_summaries: Include section summaries
        section_template: Custom Jinja2 template for sections
    
    Returns:
        Dict with keys:
            - markdown: str (the generated document in markdown)
            - docx_bytes: bytes (the DOCX file as bytes, if generated)
            - audit: Dict (audit result if run_audit=True)
            - citations_log: List (citation log for explainability)
            - metrics: Dict (cost/token metrics)
            - outline: List[str] (section titles)
    """
    import io
    from datetime import datetime
    
    # Defaults
    model = model or DEFAULT_MODEL
    reasoning_level = (reasoning_level or "medium").lower()
    rag_sources = rag_sources or ["lei", "juris", "pecas_modelo"]
    section_template = section_template or TEMPLATE_SECAO
    
    # Calculate word budget
    WORDS_PER_PAGE = 350
    min_pages = int(min_pages or 0)
    max_pages = int(max_pages or 0)
    if min_pages < 0:
        min_pages = 0
    if max_pages < 0:
        max_pages = 0
    if min_pages and max_pages and max_pages < min_pages:
        max_pages = min_pages
    if max_pages and not min_pages:
        min_pages = 1
    if min_pages and not max_pages:
        max_pages = min_pages

    min_word_budget = 0
    max_word_budget = 0
    if min_pages or max_pages:
        target_pages = (min_pages + max_pages) // 2
        min_word_budget = min_pages * WORDS_PER_PAGE
        max_word_budget = max_pages * WORDS_PER_PAGE

    total_word_budget = target_pages * WORDS_PER_PAGE if target_pages > 0 else 0
    
    # Initialize drafter
    drafter = LegalDrafter(model_name=model, cache_ttl_mins=cache_ttl, section_template=section_template)
    
    # v3.4: Context Caching Upload
    if context_files:
        valid_files = [f for f in context_files if os.path.exists(f)]
        if valid_files:
            drafter.upload_files(valid_files)
            drafter.create_context_cache(
                system_instruction="Você é um assistente jurídico especializado. " +
                                   "Use os arquivos anexos como FONTE DA VERDADE para fatos e provas. " +
                                   "NUNCA invente fatos ou leis."
            )
    
    result = {
        "markdown": "",
        "docx_bytes": None,
        "audit": None,
        "citations_log": [],
        "metrics": {},
        "outline": [],
        "divergencias": []
    }
    
    try:
        # Initialize RAG Manager if enabled and not provided
        if rag_enabled and not rag_manager:
            try:
                from rag_module import create_rag_manager
                rag_manager = create_rag_manager()
                logger.info(f"RAG Manager inicializado: {rag_manager.get_stats()}")
            except ImportError:
                logger.warning("RAG Module não encontrado")
                rag_manager = None
        
        # Initialize Local RAG if path provided
        local_index = None
        key_pdf_paths = []
        key_text_pack = ""
        if processo_local_path:
            try:
                from rag_local import LocalProcessIndex
                pid = processo_id or Path(processo_local_path).name
                local_index = LocalProcessIndex(
                    processo_id=pid,
                    sistema=sistema,
                    tenant_id=tenant_id
                )
                local_index.index_pasta(processo_local_path)
                
                # Select key documents
                key_docs = select_key_documents(local_index, max_docs=8)
                key_pdf_paths = [d["file_path"] for d in key_docs]
                key_text_pack = build_key_docs_text_pack(local_index, key_docs, max_chars=30000)
                logger.info(f"RAG Local inicializado: {len(key_pdf_paths)} PDFs selecionados")
            except ImportError:
                logger.warning("RAG Local não encontrado")
            except Exception as e:
                logger.error(f"Erro no RAG Local: {e}")
        elif local_files:
            try:
                from rag_local import LocalProcessIndex
                pid = processo_id or "upload"
                local_index = LocalProcessIndex(
                    processo_id=pid,
                    sistema="UPLOAD",
                    tenant_id=tenant_id
                )
                for path in local_files:
                    if path and os.path.exists(path):
                        local_index.index_documento(path)

                key_docs = select_key_documents(local_index, max_docs=8)
                key_pdf_paths = [d["file_path"] for d in key_docs]
                key_text_pack = build_key_docs_text_pack(local_index, key_docs, max_chars=30000)
                logger.info(f"RAG Local (arquivos soltos) inicializado: {len(key_pdf_paths)} PDFs selecionados")
            except ImportError:
                logger.warning("RAG Local não encontrado")
            except Exception as e:
                logger.error(f"Erro no RAG Local (arquivos soltos): {e}")
        
        # v5.0: Initialize Knowledge Graph for GraphRAG
        knowledge_graph = None
        if graph_rag_enabled:
            try:
                from rag_graph import LegalKnowledgeGraph
                knowledge_graph = LegalKnowledgeGraph()
                logger.info(f"GraphRAG inicializado: {knowledge_graph.get_stats()}")
            except ImportError:
                logger.warning("GraphRAG Module não encontrado")
                graph_rag_enabled = False
            except Exception as e:
                logger.error(f"Erro ao inicializar GraphRAG: {e}")
                graph_rag_enabled = False
        
        # Generate Outline
        outline = drafter.generate_outline(mode, resumo_caso=input_text, tese_usuario=tese)
        if not outline:
            raise ValueError("Falha ao gerar outline")
        # v4.1: Unification - Outline Approval
        if approve_outline:
            print(f"\n===== OUTLINE GERADO ({len(outline)} seções) =====")
            for i, section in enumerate(outline):
                print(f"  {i+1}. {section}")
            print("===============================================\n")
            
            # Simple interactive loop (blocking)
            # Only useful if running in terminal/script
            try:
                user_input = input("Aprovar? [s]im / [e]ditar / [c]ancelar: ").strip().lower()
                if user_input in ['e', 'editar']:
                    print("Digite os novos títulos (linha vazia p/ terminar):")
                    new_outline = []
                    while True:
                        line = input("  > ").strip()
                        if not line: break
                        new_outline.append(line)
                    if new_outline:
                        outline = new_outline
                        result["outline"] = outline
                        print(f"Outline atualizado: {len(outline)} seções.")
                elif user_input in ['c', 'cancelar']:
                    raise ValueError("Geração cancelada pelo usuário.")
            except EOFError:
                pass # Can't read input, assume approval
        
        # v4.1: Unification - Deep Research
        deep_research_context = ""
        if deep_research:
            logger.info("🔬 DEEP RESEARCH (Programmatic)...")
            try:
                from app.services.ai.deep_research_service import deep_research_service
                query = f"{mode}: {tese}. Seções: {', '.join(outline[:3])}"
                loop = asyncio.get_event_loop()
                if loop.is_running():
                     # We are likely in an async context already if called from FastAPI
                     # But this function is sync. Use run_in_executor or just run_until_complete if loop not running.
                     # Since this is a sync function, we assume it might be running in a thread.
                     # Safe way: use a new loop if none, or ensure we have one.
                     # Actually, for simplicity in sync function:
                    import nest_asyncio
                    nest_asyncio.apply()
                
                dr_res = asyncio.run(deep_research_service.run_research_task(query))
                if dr_res and dr_res.get("summary"):
                    deep_research_context = dr_res["summary"]
                    logger.info(f"✅ Deep Research OK: {len(deep_research_context)} chars")
            except Exception as e:
                logger.warning(f"⚠️ Deep Research failed: {e}")

        # v4.1: Unification - Web Search
        web_search_context = ""
        if web_search:
            logger.info("🌐 WEB SEARCH (Programmatic)...")
            try:
                from app.services.web_search_service import web_search_service
                wquery = f"{mode} {tese[:100]}"
                # Handle async call in sync function
                import nest_asyncio
                nest_asyncio.apply()
                
                web_res = asyncio.run(web_search_service.search(wquery, max_results=5))
                if web_res:
                    web_search_context = "\n".join([f"- {r.get('title')}: {r.get('body','')[:200]}" for r in web_res])
                    logger.info(f"✅ Web Search OK: {len(web_res)} results")
            except Exception as e:
                logger.warning(f"⚠️ Web Search failed: {e}")
                
        result["outline"] = outline
        
        # Calculate word budget per section
        num_sections = len(outline)
        word_budget_per_section = total_word_budget // num_sections if total_word_budget > 0 else 1000
        range_min_words_per_section = 0
        range_max_words_per_section = 0
        if min_word_budget > 0 and max_word_budget > 0 and num_sections > 0:
            range_min_words_per_section = max(1, min_word_budget // num_sections)
            range_max_words_per_section = max(1, max_word_budget // num_sections)

        min_words_per_section_map = {
            "low": 300,
            "medium": 450,
            "high": 600,
        }
        min_words_per_section = min_words_per_section_map.get(reasoning_level, 450)
        adjusted_target_pages = target_pages
        if total_word_budget > 0 and num_sections > 0 and word_budget_per_section < min_words_per_section:
            total_word_budget = min_words_per_section * num_sections
            word_budget_per_section = min_words_per_section
            import math
            adjusted_target_pages = int(math.ceil(total_word_budget / WORDS_PER_PAGE))
            logger.info(
                f"📏 Ajuste de completude: {num_sections} seções, "
                f"mínimo {min_words_per_section} palavras/seção → "
                f"{adjusted_target_pages} páginas (estimado)"
            )
        if min_pages or max_pages:
            if min_pages and adjusted_target_pages and adjusted_target_pages < min_pages:
                adjusted_target_pages = min_pages
            if max_pages and adjusted_target_pages and adjusted_target_pages > max_pages:
                adjusted_target_pages = max_pages
        
        # Initialize agent clients if agent mode
        gpt_client = None
        claude_client = None
        if agent_mode:
            try:
                from agent_clients import init_openai_client, init_anthropic_client
                gpt_client = init_openai_client()
                claude_client = init_anthropic_client()
                logger.info(f"Agentes inicializados: GPT={bool(gpt_client)}, Claude={bool(claude_client)}")
            except ImportError as e:
                logger.warning(f"Agent clients não disponíveis: {e}")
                agent_mode = False

        agent_mode_used = bool(agent_mode and gpt_client and claude_client)
        agents_used = ["gemini"]
        if agent_mode_used:
            agents_used = ["gpt", "claude", "gemini"]
        result["agents_used"] = agents_used
        result["agent_mode_used"] = agent_mode_used
        
        # Serial Generation
        full_document = ""
        citation_log = []
        all_divergencias = []
        contexto_anterior = "Início do documento."
        
        for i, section_title in enumerate(outline):
            logger.info(f"Gerando seção {i+1}/{len(outline)}: {section_title}")
            
            # v4.0: Adaptive Routing - Decide strategy for this section
            route_config = {"strategy": STRATEGY_HYBRID, "sources": rag_sources, "top_k": rag_top_k}
            safe_mode_active = False
            
            if adaptive_routing:
                route_config = default_route_for_section(section_title, mode)
                logger.info(f"  Routing: {route_config['strategy']}")
            
            # RAG Global search (skip if NO_RETRIEVAL or LOCAL_ONLY)
            rag_context = ""
            rag_results = []
            graph_context = ""  # v5.0: GraphRAG context
            
            if rag_manager and route_config["strategy"] not in [STRATEGY_NO_RETRIEVAL, STRATEGY_LOCAL_ONLY]:
                effective_sources = route_config.get("sources") or rag_sources
                effective_top_k = route_config.get("top_k", rag_top_k)
                
                expanded_query = drafter.generate_rag_query(section_title, input_text)
                
                # v5.0: Route to appropriate search strategy
                use_hyde = (hyde_enabled and route_config["strategy"] == STRATEGY_HYDE)
                use_graph = (graph_rag_enabled and route_config["strategy"] == STRATEGY_GRAPH)
                
                if use_hyde:
                    # HyDE: Generate hypothetical document, then search
                    logger.info(f"  🔮 Using HyDE for semantic-rich search")
                    logger.info("  🔮 HyDE: usando query hipotetica gerada pelo LegalDrafter")
                    logger.info(f"  🔮 HyDE query chars={len(expanded_query)} preview={expanded_query[:160]!r}")
                    rag_results = rag_manager.hyde_search(
                        query=expanded_query,
                        sources=effective_sources,
                        top_k=effective_top_k,
                        tenant_id=tenant_id,
                        tipo_peca_filter=MODE_TO_TIPO_PECA.get(mode.upper()),
                        verbose=verbose_rag
                    )
                elif use_graph:
                    # GraphRAG: Primary graph query with vector fallback
                    logger.info(f"  📊 Using GraphRAG for multi-hop reasoning")
                    if knowledge_graph:
                        hops = route_config.get("graph_hops", graph_hops)
                        graph_context, _ = knowledge_graph.query_context_from_text(
                            expanded_query,
                            hops=hops
                        )
                    if not graph_context:
                        rag_results = rag_manager.hybrid_search(
                            query=expanded_query,
                            sources=effective_sources,
                            top_k=effective_top_k,
                            tenant_id=tenant_id,
                            tipo_peca_filter=MODE_TO_TIPO_PECA.get(mode.upper())
                        )
                        if knowledge_graph and rag_results:
                            hops = route_config.get("graph_hops", graph_hops)
                            graph_context = knowledge_graph.enrich_context(rag_results, hops=hops)
                elif crag_gate:
                    # v4.0: Use CRAG Gate if enabled
                    gate_result = crag_gate_retrieve(
                        query=expanded_query,
                        rag_manager=rag_manager,
                        sources=effective_sources,
                        top_k=effective_top_k,
                        tenant_id=tenant_id,
                        tipo_peca_filter=MODE_TO_TIPO_PECA.get(mode.upper()),
                        min_best_score=crag_min_best,
                        min_avg_top3_score=crag_min_avg,
                        verbose=verbose_rag
                    )
                    rag_results = gate_result["results"]
                    safe_mode_active = gate_result["safe_mode"]
                else:
                    # Standard hybrid search
                    rag_results = rag_manager.hybrid_search(
                        query=expanded_query,
                        sources=effective_sources,
                        top_k=effective_top_k,
                        tenant_id=tenant_id,
                        tipo_peca_filter=MODE_TO_TIPO_PECA.get(mode.upper())
                    )
                
                if rag_results:
                    rag_context = rag_manager.format_sources_for_prompt(rag_results)
            
            # RAG Local search
            rag_local_context = ""
            local_citations_used = []
            if local_index:
                local_query = f"{section_title}. Contexto: {input_text[:500]}"
                local_results = local_index.search(local_query, top_k=5)
                if local_results:
                    rag_local_context = "\n\n### 📁 FATOS DO PROCESSO:\n"
                    for r in local_results:
                        rag_local_context += f"- {r['citacao']}: \"{r['text'][:300]}...\"\n"
                        local_citations_used.append({
                            "citacao": r['citacao'],
                            "doc_id": r.get('metadata', {}).get('doc_id', ''),
                            "tipo_doc": r.get('metadata', {}).get('tipo_doc', ''),
                            "score": r.get('final_score', 0)
                        })
            
            # Render section prompt
            t = Template(drafter.section_template)
            budget_instruction = ""
            if total_word_budget > 0:
                if range_min_words_per_section and range_max_words_per_section:
                    budget_instruction = (
                        f"\nLIMITE: Esta seção deve ficar entre {range_min_words_per_section}-{range_max_words_per_section} palavras "
                        f"(meta total {min_pages}-{max_pages} páginas)."
                    )
                else:
                    budget_instruction = f"\nLIMITE: Escreva aproximadamente {word_budget_per_section} palavras."
            
            prompt = t.render(
                titulo_secao=section_title,
                tipo_peca=mode,
                indice=i+1,
                total=len(outline),
                contexto_anterior=contexto_anterior[-2000:]
            ) + budget_instruction
            
            # Add RAG contexts
            if rag_local_context:
                prompt += rag_local_context
                prompt += "\n⚠️ Para fatos do processo acima: CITE SEMPRE a origem."
            
            if rag_context:
                prompt += "\n\n⚠️ INSTRUÇÕES DE USO DE FONTES RAG:"
                prompt += "\n- **Leis/Jurisprudência**: CITE EXPLICITAMENTE."
                prompt += "\n- **Modelos de Peças**: Use como REFERÊNCIA DE ESTILO."
            
            # v5.0: Add GraphRAG knowledge graph context
            if graph_context:
                prompt += "\n\n" + graph_context
                prompt += "\n⚠️ O grafo acima mostra RELAÇÕES entre entidades jurídicas. Use para encadear argumentação."
            
            # v4.0: Inject SAFE MODE instruction if CRAG Gate failed
            if safe_mode_active:
                prompt += "\n\n" + SAFE_MODE_INSTRUCTION
            
            # Generate section
            section_text = ""
            
            if agent_mode and (gpt_client or claude_client):
                try:
                    from agent_clients import CaseBundle, generate_section_agent_mode_async
                    import asyncio
                    
                    section_text_pack = key_text_pack + "\n\n" + rag_local_context
                    case_bundle = CaseBundle(
                        processo_id=processo_id or "N/A",
                        text_pack=section_text_pack,
                        pdf_paths=key_pdf_paths
                    )
                    
                    section_text, divergencias, drafts = asyncio.run(generate_section_agent_mode_async(
                        section_title=section_title,
                        prompt_base=prompt,
                        case_bundle=case_bundle,
                        rag_local_context=rag_local_context,
                        drafter=drafter,
                        gpt_client=gpt_client,
                        claude_client=claude_client,
                        gpt_model=gpt_model,
                        claude_model=claude_model,
                        thesis=tese,
                        mode=mode,
                        previous_sections=outline[:i],
                        reasoning_level=reasoning_level,
                        web_search=web_search
                    ))
                    if divergencias:
                        all_divergencias.append({"secao": section_title, "divergencias": divergencias})
                except Exception as e:
                    logger.error(f"Erro no modo agente: {e}")
            
            # Fallback: Single model generation
            if not section_text:
                for update in drafter.stream_section_updates(prompt, target_words=word_budget_per_section):
                    if update['type'] == 'section_done':
                        section_text = update['markdown']
            
            # Clean echo
            section_text = remover_eco_do_contexto(section_text, contexto_anterior)
            
            full_document += f"\n\n# {section_title}\n\n{section_text}"
            contexto_anterior += f"\nRESUMO SEÇÃO {section_title}: {section_text[:500]}..."
            
            citation_log.append({
                "secao": section_title,
                "rag_local_usadas": local_citations_used,
                "rag_global_count": len(rag_results) if rag_results else 0
            })
        
        # Add TOC if requested
        if include_toc:
            toc_md = "# Sumário\n\n"
            for i, title in enumerate(outline):
                toc_md += f"{i+1}. [{title}](#{title.lower().replace(' ', '-')})\n"
            toc_md += "\n---\n"
            full_document = toc_md + full_document
        
        # Deduplication
        full_document = remover_secoes_duplicadas(full_document)
        full_document = remover_paragrafos_duplicados(full_document)
        
        result["markdown"] = full_document
        result["citations_log"] = citation_log
        result["divergencias"] = all_divergencias
        
        # Generate DOCX bytes
        try:
            docx_buffer = io.BytesIO()
            # Use existing save function but capture to buffer
            doc = create_docx_from_markdown(full_document, mode)
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            result["docx_bytes"] = docx_buffer.read()
        except Exception as e:
            logger.warning(f"Não foi possível gerar DOCX: {e}")
        
        # Run Audit
        if run_audit and AUDIT_AVAILABLE:
            try:
                audit_result = audit_document_text(
                    client=drafter.client,
                    model_name="gemini-1.5-pro-002",
                    text=full_document,
                    rag_manager=rag_manager
                )
                result["audit"] = audit_result
            except Exception as e:
                logger.error(f"Erro na auditoria: {e}")
        
        # Collect metrics
        result["metrics"] = {
            "api_calls": metrics.api_calls,
            "total_prompt_tokens": metrics.total_prompt_tokens,
            "total_cached_tokens": metrics.total_cached_tokens,
            "total_completion_tokens": metrics.total_completion_tokens,
            "total_time_seconds": metrics.total_time_seconds,
            "sections": len(outline),
            "target_pages_requested": target_pages,
            "target_pages_adjusted": adjusted_target_pages,
            "target_pages_min": min_pages,
            "target_pages_max": max_pages,
            "word_budget_per_section": word_budget_per_section
        }
        
    except Exception as e:
        logger.error(f"Erro na geração: {e}")
        raise
    finally:
        drafter.cleanup()
    
    return result


def create_docx_from_markdown(markdown_text: str, mode: str) -> 'Document':
    """Helper to create DOCX from markdown (uses save_as_word_juridico logic)"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Parse markdown and add to doc
    lines = markdown_text.split('\n')
    for line in lines:
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.strip():
            p = doc.add_paragraph(line)
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(12)
    
    return doc


# =============================================================================
# MAIN WORKFLOW
# =============================================================================

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Juridico AI - Gerador de Peças Jurídicas")
    parser.add_argument("--mode", default="PETICAO_INICIAL", 
                        choices=list(PROMPT_MAP.keys()) + ["CHAT"],
                        help="Tipo de peça: PETICAO_INICIAL, CONTESTACAO, PARECER, SENTENCA, RECURSO, CONTRATO, CHAT")
    parser.add_argument("--input", required=True, help="Arquivo de texto com Fatos/Notas ou Pasta com arquivos")
    parser.add_argument("--processo", help="Caminho para PDF/TXT do Processo (Opcional)")
    parser.add_argument("--tese", default="A favor do cliente", help="Instrução central da tese")
    parser.add_argument("--template", help="Caminho para arquivo de Template Jinja2 personalizado (Opcional)")
    parser.add_argument("--model", default=DEFAULT_MODEL, 
                        help="Modelo Gemini: gemini-3-flash-preview, gemini-1.5-pro-002, etc")
    
    # Novas flags de formatação (spec v2)
    parser.add_argument("--include-toc", action="store_true", 
                        help="Incluir Sumário automático no início")
    parser.add_argument("--include-section-summaries", action="store_true",
                        help="Gerar quadro-síntese ao final de cada seção")
    parser.add_argument("--include-summary-table", action="store_true",
                        help="Gerar tabela de síntese final do documento")
    parser.add_argument("--target-pages", type=int, default=0,
                        help="Número alvo de páginas (0=sem limite). Calibração: ~350 palavras/página A4.")
    parser.add_argument("--min-pages", type=int, default=0,
                        help="Número MÍNIMO de páginas (0=sem mínimo). Usado junto com --max-pages para intervalo.")
    parser.add_argument("--max-pages", type=int, default=0,
                        help="Número MÁXIMO de páginas (0=sem máximo). Usado junto com --min-pages para intervalo.")
    
    # v3.1: Modo Chat/Prompt Customizado
    parser.add_argument("--chat", action="store_true",
                        help="Modo interativo: conversa com os documentos como no ChatGPT")
    parser.add_argument("--prompt", type=str, default=None,
                        help="Prompt customizado único (ex: 'Resuma os principais argumentos do réu')")
    
    # v3.2: RAG - Retrieval Augmented Generation
    parser.add_argument("--rag", action="store_true",
                        help="Ativar RAG (busca em bases de lei, juris, SEI, modelos)")
    parser.add_argument("--sources", type=str, default="lei,juris,pecas_modelo",
                        help="Bases RAG a consultar (ex: 'lei,juris,sei,pecas_modelo')")
    parser.add_argument("--rag-top-k", type=int, default=8,
                        help="Número de trechos RAG por seção (default: 8)")
    parser.add_argument("--tenant-id", type=str, default="default",
                        help="Tenant ID para documentos SEI (multi-tenancy)")
    parser.add_argument("--user-id", type=str, default=None,
                        help="User ID para RBAC em documentos SEI sigilosos")
    parser.add_argument("--fast-rag", action="store_true",
                        help="Modo rápido: pula Query Expansion (usa título + resumo direto)")
    parser.add_argument("--verbose-rag", action="store_true",
                        help="Modo debug: imprime scores de cada fonte RAG encontrada")
    parser.add_argument("--native-pdf", action="store_true",
                        help="Suporte nativo a PDF: envia o PDF original para o Gemini e usa Context Caching")
    
    # v2.0: RAG Local - Dual Retrieval
    parser.add_argument("--processo-local", type=str, default=None,
                        help="Pasta com autos do processo para RAG Local (dual retrieval)")
    parser.add_argument("--processo-id", type=str, default=None,
                        help="ID do processo para RAG Local (ex: 'SEI-12345/2024')")
    parser.add_argument("--sistema-processo", type=str, default="SEI",
                        choices=["SEI", "PJe", "eproc", "SAPIENS"],
                        help="Sistema de origem do processo")
    
    # v2.1: Agent Mode - Multi-Model Committee
    parser.add_argument("--agent-mode", action="store_true",
                        help="Modo Agente: gera seções com GPT + Claude, consolida com Gemini")
    parser.add_argument("--save-drafts", action="store_true", help="Salva rascunhos individuais dos agentes (se agent-mode ativado)")
    
    # Audit Flag
    parser.add_argument(
        "--audit",
        action="store_true",
        help="Rodar auditoria jurídica ao final e gerar relatório *_AUDITORIA.md"
    )
    
    # Context Caching & Files (v3.4)
    parser.add_argument("--context-files", type=str, default=None,
                        help="Arquivos PDF para upload ao Cache do Gemini (separados por vírgula)")
    parser.add_argument("--cache-ttl", type=int, default=60,
                        help="Tempo de vida do cache em minutos (default: 60)")
    
    parser.add_argument("--gpt-model", type=str, default="gpt-4o",
                        help="Modelo GPT a usar (default: gpt-4o)")
    parser.add_argument("--claude-model", type=str, default="claude-sonnet-4-20250514",
                        help="Modelo Claude a usar (default: claude-sonnet-4)")
    
    # v4.0: Adaptive RAG & CRAG Gate
    parser.add_argument("--adaptive-routing", action="store_true",
                        help="Ativar roteamento adaptativo: escolhe estratégia RAG por seção")
    parser.add_argument("--crag-gate", action="store_true",
                        help="Ativar CRAG Gate: valida qualidade do RAG antes de gerar")
    parser.add_argument("--crag-min-best", type=float, default=0.45,
                        help="CRAG: score mínimo do melhor chunk (default: 0.45)")
    parser.add_argument("--crag-min-avg", type=float, default=0.35,
                        help="CRAG: score médio mínimo dos top 3 chunks (default: 0.35)")
    
    # v4.1: HIL & Research Flags (Unification with LangGraph API)
    parser.add_argument("--approve-outline", action="store_true",
                        help="Pausar para aprovar/editar outline antes de gerar seções (interativo)")
    parser.add_argument("--deep-research", action="store_true",
                        help="Ativar Deep Research: pesquisa aprofundada antes de gerar")
    parser.add_argument("--web-search", action="store_true",
                        help="Ativar Web Search: busca na web por fatos atuais")
    
    # v5.0: GraphRAG & HyDE
    parser.add_argument("--hyde", action="store_true",
                        help="Ativar HyDE: gera documento hipotético para busca semântica aprimorada")
    parser.add_argument("--graph-rag", action="store_true",
                        help="Ativar GraphRAG: enriquece resultados com grafo de conhecimento jurídico")
    parser.add_argument("--graph-hops", type=int, default=2,
                        help="GraphRAG: profundidade de relacionamentos a explorar (default: 2)")
    
    args = parser.parse_args()
    
    # Calcular word_budget baseado em target_pages ou intervalo min/max
    WORDS_PER_PAGE = 350  # Calibração para A4, Arial 12pt, margens ABNT
    
    # Prioridade: Se --min-pages/--max-pages definidos, usar intervalo
    if args.min_pages > 0 or args.max_pages > 0:
        min_pages = args.min_pages if args.min_pages > 0 else 1
        max_pages = args.max_pages if args.max_pages > 0 else min_pages * 2
        # Target é o ponto médio do intervalo
        target_pages_effective = (min_pages + max_pages) // 2
        total_word_budget = target_pages_effective * WORDS_PER_PAGE
        min_word_budget = min_pages * WORDS_PER_PAGE
        max_word_budget = max_pages * WORDS_PER_PAGE
        print(f"{Fore.CYAN}📏 Intervalo de Páginas: {min_pages}-{max_pages} páginas ({min_word_budget}-{max_word_budget} palavras)")
    elif args.target_pages > 0:
        total_word_budget = args.target_pages * WORDS_PER_PAGE
        min_word_budget = 0
        max_word_budget = 0
    else:
        total_word_budget = 0
        min_word_budget = 0
        max_word_budget = 0
    
    input_path = args.input
    files_to_cache = []
    
    # 0. Carregar Template Personalizado (se houver)
    section_template = TEMPLATE_SECAO
    if args.template:
        try:
            with open(args.template, 'r', encoding='utf-8') as f:
                section_template = f.read()
            print(f"{Fore.CYAN}🎨 Usando Template Personalizado: {args.template}")
        except Exception as e:
            print(f"{Fore.RED}❌ Erro ao ler template personalizado: {e}")
            return

    # 1. Carregar Inputs
    resumo_input = ""
    if os.path.isdir(input_path):
        files_to_cache.extend([str(p) for p in Path(input_path).glob("*") if p.is_file()])
        resumo_input = "Veja os arquivos anexos no contexto."
    else:
        ext = Path(input_path).suffix.lower()
        if ext == '.pdf':
            if args.native_pdf:
                print(f"{Fore.CYAN}📄 Modo Nativo PDF: Arquivo será enviado via File API.")
                files_to_cache.append(input_path)
                resumo_input = "Os fatos estão contidos no arquivo PDF anexo."
            else:
                print(f"{Fore.YELLOW}📑 Extraindo texto de PDF (modo legacy)...")
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(input_path)
                    for page in reader.pages:
                        resumo_input += (page.extract_text() or "") + "\n"
                except ImportError:
                    print(f"{Fore.RED}❌ PyPDF2 não instalado. Use --native-pdf ou instale: pip install PyPDF2")
                    return
                except Exception as e:
                    print(f"{Fore.RED}❌ Erro ao extrair PDF: {e}")
                    return
        else:
            with open(input_path, 'r', encoding='utf-8') as f:
                resumo_input = f.read()
            # Adicionar arquivo de input ao cache também
            files_to_cache.append(input_path)
            
    if args.processo:
        files_to_cache.append(args.processo)

    # v3.4: Add explicit context files from CLI
    if args.context_files:
        extra_files = [f.strip() for f in args.context_files.split(",") if f.strip()]
        for f in extra_files:
            if os.path.exists(f):
                files_to_cache.append(f)
            else:
                print(f"{Fore.YELLOW}⚠️ Arquivo de contexto não encontrado: {f}")
        
    # 2. Inicializar Drafter com Template e Modelo
    drafter = LegalDrafter(model_name=args.model, cache_ttl_mins=args.cache_ttl, section_template=section_template)
    
    try:
        # 3. Upload & Cache
        if files_to_cache:
            drafter.upload_files(files_to_cache)
            drafter.create_context_cache(
                system_instruction="Você é um assistente jurídico especializado. " +
                                   "Use os arquivos anexos como FONTE DA VERDADE para fatos e provas. " +
                                   "NUNCA invente fatos ou leis."
            )
        
        # 3.1 MODO CHAT/PROMPT CUSTOMIZADO (v3.1)
        if args.chat or args.prompt or args.mode == "CHAT":
            chat_with_documents(drafter, resumo_input, single_prompt=args.prompt)
            drafter.cleanup()
            return
            
        # 4. Planejamento (Outline) - modo padrão de geração de peças
        outline = drafter.generate_outline(args.mode, resumo_caso=resumo_input, tese_usuario=args.tese)
        
        if not outline:
            print("❌ Falha no Outline.")
            return

        # v4.1: Outline Approval HIL (interactive)
        if args.approve_outline:
            print(f"\n{Fore.CYAN}{'='*60}")
            print(f"{Fore.CYAN}📋 OUTLINE GERADO - APROVAÇÃO NECESSÁRIA")
            print(f"{Fore.CYAN}{'='*60}")
            for i, section in enumerate(outline):
                print(f"  {i+1}. {section}")
            print(f"{Fore.CYAN}{'='*60}\n")
            
            while True:
                user_input = input(f"{Fore.YELLOW}Aprovar? [s]im / [e]ditar / [c]ancelar: {Style.RESET_ALL}").strip().lower()
                if user_input in ['s', 'sim', 'y', 'yes', '']:
                    print(f"{Fore.GREEN}✅ Outline aprovado!")
                    break
                elif user_input in ['e', 'editar', 'edit']:
                    print(f"{Fore.YELLOW}Digite os títulos das seções (um por linha, linha vazia para terminar):")
                    new_outline = []
                    while True:
                        line = input("  > ").strip()
                        if not line:
                            break
                        new_outline.append(line)
                    if new_outline:
                        outline = new_outline
                        print(f"{Fore.GREEN}✅ Outline atualizado com {len(outline)} seções.")
                    break
                elif user_input in ['c', 'cancelar', 'cancel']:
                    print(f"{Fore.RED}❌ Geração cancelada pelo usuário.")
                    return
                else:
                    print(f"{Fore.RED}Opção inválida. Use 's', 'e' ou 'c'.")

        # v4.1: Deep Research (antes de gerar seções)
        deep_research_context = ""
        if args.deep_research:
            print(f"\n{Fore.MAGENTA}🔬 DEEP RESEARCH: Pesquisando contexto aprofundado...")
            try:
                # Import deep research service
                from app.services.ai.deep_research_service import deep_research_service
                
                research_query = f"{args.mode}: {args.tese}. Seções: {', '.join(outline[:3])}"
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                research_result = loop.run_until_complete(
                    deep_research_service.run_research_task(research_query)
                )
                loop.close()
                
                if research_result and research_result.get("summary"):
                    deep_research_context = research_result["summary"]
                    print(f"{Fore.GREEN}✅ Deep Research concluído: {len(deep_research_context)} chars")
                else:
                    print(f"{Fore.YELLOW}⚠️ Deep Research não retornou resultados.")
            except ImportError:
                print(f"{Fore.YELLOW}⚠️ Deep Research indisponível (serviço não encontrado).")
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Deep Research falhou: {e}")
        
        # v4.1: Web Search (busca fatos atuais)
        web_search_context = ""
        if args.web_search:
            print(f"\n{Fore.BLUE}🌐 WEB SEARCH: Buscando informações atuais...")
            try:
                from app.services.web_search_service import web_search_service
                
                search_query = f"{args.mode} {args.tese[:100]}"
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                web_results = loop.run_until_complete(
                    web_search_service.search(search_query, max_results=5)
                )
                loop.close()
                
                if web_results:
                    web_search_context = "\n\n".join([
                        f"[{r.get('title', 'Sem título')}]: {r.get('snippet', '')}"
                        for r in web_results
                    ])
                    print(f"{Fore.GREEN}✅ Web Search: {len(web_results)} resultados encontrados")
                else:
                    print(f"{Fore.YELLOW}⚠️ Web Search não retornou resultados.")
            except ImportError:
                print(f"{Fore.YELLOW}⚠️ Web Search indisponível (serviço não encontrado).")
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Web Search falhou: {e}")

        # 4.1 Calcular word_budget por seção (distribuição proporcional)
        num_sections = len(outline)
        if total_word_budget > 0:
            word_budget_per_section = total_word_budget // num_sections
            min_words_per_section = 450
            if word_budget_per_section < min_words_per_section and num_sections > 0:
                total_word_budget = min_words_per_section * num_sections
                word_budget_per_section = min_words_per_section
                import math
                adjusted_pages = int(math.ceil(total_word_budget / WORDS_PER_PAGE))
                print(
                    f"{Fore.CYAN}📏 Ajuste: {num_sections} seções → "
                    f"{adjusted_pages} páginas (min {min_words_per_section} palavras/seção)"
                )
            else:
                print(f"{Fore.CYAN}📏 Target: {args.target_pages} páginas → {total_word_budget} palavras")
                print(f"   ~{word_budget_per_section} palavras/seção")
        else:
            word_budget_per_section = 1000  # Default sem limite

        # 4.2 Inicializar RAG (se ativado)
        rag_manager = None
        if args.rag:
            try:
                from rag_module import create_rag_manager
                print(f"\n{Fore.CYAN}🔍 Inicializando RAG (sources: {args.sources})...")
                rag_manager = create_rag_manager()
                stats = rag_manager.get_stats()
                print(f"   📊 Documentos indexados: {stats}")
            except ImportError:
                print(f"{Fore.YELLOW}⚠️ RAG Module não encontrado. Instale: pip install chromadb sentence-transformers rank_bm25")
                rag_manager = None
            except Exception as e:
                print(f"{Fore.RED}❌ Erro ao inicializar RAG: {e}")
                rag_manager = None
        
        # 4.3 Inicializar RAG Local (v2.0 - Dual Retrieval)
        local_index = None
        key_pdf_paths = []  # v3.1: Default empty to avoid NameError
        key_text_pack = ""  # v3.1: Default empty to avoid NameError
        if args.processo_local:
            try:
                from rag_local import LocalProcessIndex
                processo_id = args.processo_id or Path(args.processo_local).name
                print(f"\n{Fore.CYAN}📁 Inicializando RAG Local para processo: {processo_id}")
                local_index = LocalProcessIndex(
                    processo_id=processo_id,
                    sistema=args.sistema_processo
                )
                local_index.index_pasta(args.processo_local)
                
                # v3.1: Determinstic Key Doc Selection
                print(f"{Fore.CYAN}🔑 Selecionando documentos-chave para contexto dos Agentes...")
                key_docs = select_key_documents(local_index, max_docs=8)
                key_pdf_paths = [d["file_path"] for d in key_docs]
                key_text_pack = build_key_docs_text_pack(local_index, key_docs, max_chars=30000) # Reduced to 30k for safety
                
                print(f"   Selecionados {len(key_pdf_paths)} PDFs principais.")
                if key_pdf_paths:
                    print(f"   Contexto inicial gerado: {len(key_text_pack)} chars.")
            except ImportError:
                print(f"{Fore.YELLOW}⚠️ RAG Local não encontrado (rag_local.py)")
                local_index = None
            except Exception as e:
                print(f"{Fore.RED}❌ Erro ao inicializar RAG Local: {e}")
                local_index = None

        # 4.4 Inicializar GraphRAG (v5.0)
        knowledge_graph = None
        if args.graph_rag:
            try:
                from rag_graph import LegalKnowledgeGraph
                print(f"\n{Fore.CYAN}🕸️  Inicializando GraphRAG...")
                knowledge_graph = LegalKnowledgeGraph()
                print(f"   📊 Grafo: {knowledge_graph.get_stats()}")
            except ImportError:
                print(f"{Fore.YELLOW}⚠️ GraphRAG Module não encontrado.")
                args.graph_rag = False
            except Exception as e:
                print(f"{Fore.RED}❌ Erro ao inicializar GraphRAG: {e}")
                args.graph_rag = False

        # 5. Execução Serial
        full_document = ""
        section_summaries = []  # Para tabela final
        citation_log = []  # v2.1: Track all citations for explainability
        all_divergencias = []  # v2.1: Agent mode divergencias log
        contexto_anterior = "Início do documento."
        rag_sources_list = args.sources.split(",") if args.rag else []
        
        # v2.1: Initialize agent clients if agent mode is enabled
        gpt_client = None
        claude_client = None
        if args.agent_mode:
            print(f"\n{Fore.CYAN}🤖 Modo Agente ativado. Inicializando clientes...")
            try:
                from agent_clients import init_openai_client, init_anthropic_client, generate_section_agent_mode
                gpt_client = init_openai_client()
                claude_client = init_anthropic_client()
                print(f"   GPT: {'✅' if gpt_client else '❌'} | Claude: {'✅' if claude_client else '❌'}")
            except ImportError as e:
                print(f"{Fore.RED}❌ Erro ao importar agent_clients: {e}")
                args.agent_mode = False
        
        print(f"\n{Fore.MAGENTA}🚀 Iniciando Redação Serial ({len(outline)} seções)...{' [AGENT MODE]' if args.agent_mode else ''}\n")
        
        for i, section_title in enumerate(outline):
            print(f"{Fore.BLUE}✍️  Escrevendo Seção {i+1}/{len(outline)}: {section_title}...")
            
            # v4.0: Adaptive Routing - Decide strategy for this section
            route_config = {"strategy": STRATEGY_HYBRID, "sources": rag_sources_list, "top_k": args.rag_top_k}
            safe_mode_active = False
            
            if args.adaptive_routing:
                route_config = default_route_for_section(section_title, args.mode)
                print(f"   🧭 Routing: {route_config['strategy']} ({route_config['reason'][:40]}...)")
            
            # 5.1 RAG Search por Seção (se ativado e strategy != NO_RETRIEVAL)
            rag_context = ""
            rag_results = []
            graph_context = ""  # v5.0
            
            if rag_manager and route_config["strategy"] != STRATEGY_NO_RETRIEVAL:
                effective_sources = route_config.get("sources") or rag_sources_list
                effective_top_k = route_config.get("top_k", args.rag_top_k)
                
                # Skip global RAG if LOCAL_ONLY
                if route_config["strategy"] == STRATEGY_LOCAL_ONLY:
                    print(f"   ℹ️  Skipping global RAG (LOCAL_ONLY strategy)")
                else:
                    # Query Expansion (v3.3)
                    if args.fast_rag or not route_config.get("needs_query_expansion", True):
                        expanded_query = f"{section_title}. {resumo_input[:500]}"
                    else:
                        expanded_query = drafter.generate_rag_query(section_title, resumo_input)
                    
                    # v5.0: Route to appropriate search strategy
                    use_hyde = (args.hyde and route_config["strategy"] == STRATEGY_HYDE)
                    use_graph = (args.graph_rag and route_config["strategy"] == STRATEGY_GRAPH)
                    
                    if use_hyde:
                        print(f"   🔮 HyDE: Gerando documento hipotético e buscando...")
                        print("   🔮 HyDE: usando query hipotetica gerada pelo LegalDrafter")
                        print(f"   🔮 HyDE query chars={len(expanded_query)} preview={expanded_query[:160]!r}")
                        rag_results = rag_manager.hyde_search(
                            query=expanded_query,
                            sources=effective_sources,
                            top_k=effective_top_k,
                            tenant_id=args.tenant_id,
                            tipo_peca_filter=MODE_TO_TIPO_PECA.get(args.mode.upper()),
                            verbose=args.verbose_rag
                        )
                    elif use_graph:
                        print(f"   📊 GraphRAG: Buscando com enriquecimento de grafo...")
                        rag_results = rag_manager.hybrid_search(
                            query=expanded_query,
                            sources=effective_sources,
                            top_k=effective_top_k,
                            tenant_id=args.tenant_id,
                            tipo_peca_filter=MODE_TO_TIPO_PECA.get(args.mode.upper())
                        )
                        if knowledge_graph and rag_results:
                            hops = route_config.get("graph_hops", args.graph_hops)
                            graph_context = knowledge_graph.enrich_context(rag_results, hops=hops)
                            print(f"   🕸️  Grafo: Contexto enriquecido ({len(graph_context)} chars)")
                    elif args.crag_gate:
                        # v4.0: Use CRAG Gate if enabled
                        gate_result = crag_gate_retrieve(
                            query=expanded_query,
                            rag_manager=rag_manager,
                            sources=effective_sources,
                            top_k=effective_top_k,
                            tenant_id=args.tenant_id,
                            user_id=args.user_id,
                            tipo_peca_filter=MODE_TO_TIPO_PECA.get(args.mode.upper()),
                            min_best_score=args.crag_min_best,
                            min_avg_top3_score=args.crag_min_avg,
                            verbose=args.verbose_rag
                        )
                        rag_results = gate_result["results"]
                        safe_mode_active = gate_result["safe_mode"]
                        
                        if gate_result["gate_passed"]:
                            print(f"   ✅ CRAG Gate: PASSED (best={gate_result['best_score']:.2f})")
                        elif safe_mode_active:
                            print(f"   ⚠️ CRAG Gate: SAFE MODE (best={gate_result['best_score']:.2f} < {args.crag_min_best})")
                    else:
                        print(f"   🔍 Hybrid RAG Normal...")
                        rag_results = rag_manager.hybrid_search(
                            query=expanded_query,
                            sources=effective_sources,
                            top_k=effective_top_k,
                            tenant_id=args.tenant_id,
                            tipo_peca_filter=MODE_TO_TIPO_PECA.get(args.mode.upper())
                        )

                    
                    if rag_results:
                        rag_context = rag_manager.format_sources_for_prompt(rag_results)
                        print(f"   ✅ {len(rag_results)} fontes globais encontradas")
                        
                        # Verbose: imprimir detalhes de cada fonte
                        if args.verbose_rag:
                            for j, r in enumerate(rag_results):
                                print(f"      [{j+1}] {r['source']}: score={r['final_score']:.2f} bm25={r.get('bm25_score', 0):.2f}")
            
            # v2.0: RAG Local - Buscar fatos do processo específico
            rag_local_context = ""
            local_citations_used = []  # v2.1: Track citations for explainability
            if local_index:
                print(f"   🔍 Buscando nos autos do processo...")
                # v2.1: Query orientada à seção (section_title + contexto do caso)
                local_query = f"{section_title}. Contexto: {resumo_input[:500]}"
                local_results = local_index.search(local_query, top_k=5)
                if local_results:
                    rag_local_context = "\n\n### 📁 FATOS DO PROCESSO:\n"
                    for r in local_results:
                        rag_local_context += f"- {r['citacao']}: \"{r['text'][:300]}...\"\n"
                        local_citations_used.append({
                            "citacao": r['citacao'],
                            "doc_id": r.get('metadata', {}).get('doc_id', ''),
                            "tipo_doc": r.get('metadata', {}).get('tipo_doc', ''),
                            "score": r.get('final_score', 0)
                        })
                    print(f"   ✅ {len(local_results)} trechos do processo encontrados")
            
            # Render Prompt com word_budget
            t = Template(drafter.section_template)
            budget_instruction = f"\nLIMITE: Escreva aproximadamente {word_budget_per_section} palavras." if total_word_budget > 0 else ""
            
            prompt = t.render(
                titulo_secao=section_title,
                tipo_peca=args.mode,
                indice=i+1,
                total=len(outline),
                contexto_anterior=contexto_anterior[-2000:]
            ) + budget_instruction
            
            # Injetar fontes RAG no prompt (v2.0: Dual retrieval with differentiated instructions)
            if rag_local_context:
                prompt += rag_local_context
                prompt += "\n⚠️ Para fatos do processo acima: CITE SEMPRE a origem (ex: [LAUDO - Doc. SEI nº 12345, p. 3])."
            
            if rag_context:
                prompt += "\n\n⚠️ INSTRUÇÕES DE USO DE FONTES RAG:"
                prompt += "\n- **Leis/Jurisprudência**: CITE EXPLICITAMENTE (ex: 'conforme art. X da Lei Y', 'nos termos da Súmula Z do STJ')."
                prompt += "\n- **Modelos de Peças**: Use como REFERÊNCIA DE ESTILO e ESTRUTURA. Imite o tom, vocabulário e formatação. NÃO copie literalmente."
            
            # v5.0: Add GraphRAG knowledge graph context
            if graph_context:
                prompt += "\n\n" + graph_context
                prompt += "\n⚠️ O grafo acima mostra RELAÇÕES entre entidades jurídicas. Use para encadear argumentação."
            
            # v4.0: Inject SAFE MODE instruction if CRAG Gate failed
            if safe_mode_active:
                prompt += "\n\n" + SAFE_MODE_INSTRUCTION
            
            # v2.1: Agent Mode - Multi-model generation with judge
            # v2.1: Agent Mode - Multi-model generation with judge
            section_text = ""
            if args.agent_mode and (gpt_client or claude_client):
                print(f"   🤖 Gerando via comitê de agentes (Async)...")
                
                # Create CaseBundle for this section context
                from agent_clients import CaseBundle, generate_section_agent_mode_async
                import asyncio
                
                # Stable context pack from key docs + relevant search results
                section_text_pack = key_text_pack + "\n\n" + rag_local_context

                case_bundle = CaseBundle(
                    processo_id=f"Processo {sistema_processo or 'N/A'}",
                    text_pack=section_text_pack,
                    pdf_paths=key_pdf_paths # Use stable key PDF paths
                )
                
                section_text, divergencias, drafts = asyncio.run(generate_section_agent_mode_async(
                    section_title=section_title,
                    prompt_base=prompt,
                    case_bundle=case_bundle,
                    rag_local_context=rag_local_context,
                    drafter=drafter,
                    gpt_client=gpt_client,
                    claude_client=claude_client,
                    gpt_model=args.gpt_model,
                    claude_model=args.claude_model
                ))
                if divergencias:
                    all_divergencias.append({"secao": section_title, "divergencias": divergencias})
                
                # Save drafts if requested
                if args.save_drafts:
                    drafts_dir = os.path.join(os.path.dirname(input_path) if os.path.isfile(input_path) else input_path, "_drafts")
                    os.makedirs(drafts_dir, exist_ok=True)
                    for agent_name, draft_text in drafts.items():
                        draft_file = os.path.join(drafts_dir, f"section_{i+1}_{agent_name}.md")
                        with open(draft_file, 'w', encoding='utf-8') as f:
                            f.write(draft_text or "")
            
            # Fallback: Se agente falhou (ou não foi ativado), usa geração padrão (Gemini)
            if not section_text:
                if args.agent_mode and (gpt_client or claude_client):
                     logger.warning(f"⚠️ Falha no comitê de agentes para seção '{section_title}'. Usando fallback (Gemini Solo).")
                     
                # Stream Generation (default single-model mode)
                for update in drafter.stream_section_updates(prompt, target_words=word_budget_per_section):
                    if update['type'] == 'section_progress':
                        sys.stdout.write(f"\r   ⏳ Progresso: {update['percent']}% ({update['word_count']} palavras)")
                        sys.stdout.flush()
                    elif update['type'] == 'section_done':
                        section_text = update['markdown']
            
            # v3.0: Apply echo removal to clean context repetition
            section_text = remover_eco_do_contexto(section_text, contexto_anterior)
            
            word_count = len(section_text.split())
            print(f"\n   ✅ Concluído. ({word_count} palavras)")    
            
            full_document += f"\n\n# {section_title}\n\n{section_text}"
            
            # 5.1 Gerar quadro-síntese da seção (se flag ativa)
            if getattr(args, 'include_section_summaries', False):
                print(f"   📊 Gerando quadro-síntese...")
                summary_prompt = f"""
                Leia a seção "{section_title}" abaixo e gere um QUADRO-SÍNTESE em formato de lista (bullets):
                - 3 a 5 pontos principais
                - Máximo 1 linha por ponto
                
                SEÇÃO:
                {section_text[:3000]}
                """
                try:
                    summary_response = drafter._generate_with_retry(summary_prompt)
                    if summary_response and summary_response.text:
                        section_summary = summary_response.text.strip()
                        full_document += f"\n\n> **Síntese da Seção:**\n{section_summary}\n"
                        section_summaries.append({"title": section_title, "summary": section_summary})
                except:
                    pass
            
            # Update Link Context (Summary for next section)
            contexto_anterior += f"\nRESUMO SECÃO {section_title}: {section_text[:500]}..."
            
            # v2.1: Log citations used for explainability
            citation_log.append({
                "secao": section_title,
                "rag_local_usadas": local_citations_used,
                "rag_global_count": len(rag_results) if rag_results else 0
            })
        
        # 5.2 Gerar tabela de síntese final (se flag ativa)
        if getattr(args, 'include_summary_table', False) and section_summaries:
            print(f"\n{Fore.CYAN}📊 Gerando tabela de síntese final...")
            table_md = "\n\n# Tabela de Síntese\n\n| Seção | Pontos Principais |\n|-------|-------------------|\n"
            for item in section_summaries:
                # Limpa bullets e quebras de linha para tabela
                clean_summary = item['summary'].replace('\n', ' ').replace('- ', '• ')[:200]
                table_md += f"| {item['title']} | {clean_summary} |\n"
            full_document += table_md
        
        # 5.3 Inserir TOC no início (se flag ativa)
        if getattr(args, 'include_toc', False):
            toc_md = "# Sumário\n\n"
            for i, title in enumerate(outline):
                toc_md += f"{i+1}. [{title}](#{title.lower().replace(' ', '-')})\n"
            toc_md += "\n---\n"
            full_document = toc_md + full_document
        
        # 5.4 DEDUPLICAÇÃO (v3.0 - Portado de mlx_vomo.py)
        print(f"\n{Fore.CYAN}🧹 Pipeline de Deduplicação (v3.0)...")
        full_document = remover_secoes_duplicadas(full_document)
        full_document = remover_paragrafos_duplicados(full_document)
        print(f"{Fore.GREEN}   ✅ Deduplicação concluída.{Style.RESET_ALL}")
            
        # 6. Salvar outputs
        output_folder = os.path.dirname(input_path) if os.path.isfile(input_path) else input_path
        if os.path.isfile(output_folder):
             output_folder = os.path.dirname(output_folder)
             
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        base_name = f"JuridicoAI_{args.mode}_{timestamp}"
        
        md_file = os.path.join(output_folder, f"{base_name}.md")
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(full_document)
        print(f"\n💾 Markdown salvo: {md_file}")
        
        # v2.1: Salvar log de citações para explainability
        if citation_log:
            citation_log_file = os.path.join(output_folder, f"{base_name}_CITACOES.json")
            with open(citation_log_file, 'w', encoding='utf-8') as f:
                json.dump(citation_log, f, indent=2, ensure_ascii=False)
            print(f"📋 Log de citações salvo: {citation_log_file}")
        
            div_file = os.path.join(output_folder, f"{base_name}_DIVERGENCIAS.json")
            with open(div_file, 'w', encoding='utf-8') as f:
                json.dump(all_divergencias, f, indent=2, ensure_ascii=False)
            print(f"📋 Log de divergências salvo: {div_file}")

        # ============================================================
        # 6. AUDITORIA JURÍDICA (opcional)
        # ============================================================

            
            # v2.2: Salvar métricas de custo e gerar dashboard
            try:
                from agent_clients import agent_metrics, generate_divergence_dashboard
                
                # Save metrics JSON
                metrics_file = os.path.join(output_folder, f"{base_name}_METRICAS.json")
                agent_metrics.save(metrics_file)
                print(f"📊 Métricas de custo salvas: {metrics_file}")
                print(f"   💵 Custo estimado: ${agent_metrics.estimated_cost_usd():.4f} USD")
                
                # Generate HTML dashboard
                dashboard_dir = os.path.join(output_folder, "_dashboard")
                dashboard_file = os.path.join(dashboard_dir, f"{base_name}_divergencias.html")
                
                # Prepare sections data for dashboard
                dashboard_secoes = []
                for i, div_data in enumerate(all_divergencias):
                    dashboard_secoes.append({
                        "titulo": div_data.get("secao", f"Seção {i+1}"),
                        "status": "warning" if div_data.get("divergencias") else "success",
                        "status_label": "Com Divergências" if div_data.get("divergencias") else "Consolidado",
                        "versao_gpt": div_data.get("drafts", {}).get("gpt_v2", "N/A"),
                        "versao_claude": div_data.get("drafts", {}).get("claude_v2", "N/A"),
                        "divergencias": []  # Parse from div_data if needed
                    })
                
                generate_divergence_dashboard(
                    titulo_documento=f"{args.mode} - {base_name}",
                    secoes=dashboard_secoes,
                    metrics=agent_metrics,
                    output_path=dashboard_file
                )
            except ImportError:
                print("⚠️ Métricas do modo agente não disponíveis.")
        


        # 8. DOCX (passa flags para formatação)
        docx_file = f"{base_name}.docx"
        save_as_word_juridico(full_document, docx_file, output_folder, args.mode)
        
        # 9. Auditoria Jurídica (se habilitado)
        if args.audit and AUDIT_AVAILABLE:
            print(f"\n{Fore.CYAN}{'='*60}")
            print(f"⚖️  AUDITORIA JURÍDICA")
            print(f"{'='*60}{Style.RESET_ALL}\n")
            
            try:
                audit_result = audit_document_text(
                    client=drafter.client,
                    model_name="gemini-1.5-pro-002",  # Modelo de auditoria
                    text=full_document,
                    rag_manager=rag_manager if args.rag else None
                )
                
                # Salvar relatório MD
                audit_md_path = os.path.splitext(md_file)[0] + "_AUDITORIA.md"
                with open(audit_md_path, 'w', encoding='utf-8') as f:
                    f.write(f"# ⚖️ Auditoria Jurídica\n\n")
                    f.write(f"**Documento Auditado:** `{os.path.basename(md_file)}`\n")
                    f.write(f"**Data:** {audit_result['audit_date']}\n\n")
                    f.write("---\n\n")
                    f.write(audit_result["audit_report_markdown"])
                    
                    # Adicionar tabela de citações se houver
                    if audit_result.get("citations"):
                        f.write("\n\n---\n\n## 📊 Análise Detalhada de Citações\n\n")
                        f.write("| Citação | Status | Score | Mensagem |\n")
                        f.write("|---------|--------|-------|----------|\n")
                        
                        icon_map = {
                            "valid": "🟢",
                            "suspicious": "🟠",
                            "hallucination": "🔴",
                            "warning": "⚠️",
                            "not_found": "❓"
                        }
                        
                        for cit in audit_result["citations"]:
                            icon = icon_map.get(cit.get("status", "not_found"), "❓")
                            f.write(f"| {cit['citation']} | {icon} {cit.get('status', 'N/A')} | "
                                   f"{cit.get('score', 0):.2f} | {cit.get('message', '')} |\n")
                
                print(f"{Fore.GREEN}✅ Relatório de Auditoria: {audit_md_path}")
                
                # Salvar JSON estruturado (opcional - útil para API/UI)
                audit_json_path = os.path.splitext(md_file)[0] + "_AUDITORIA.json"
                with open(audit_json_path, 'w', encoding='utf-8') as f:
                    json.dump(audit_result, f, indent=2, ensure_ascii=False)
                
                print(f"{Fore.GREEN}✅ Dados Estruturados: {audit_json_path}")
                
                # Mostrar resumo rápido no terminal
                citations = audit_result.get("citations", [])
                if citations:
                    suspicious = [c for c in citations if c.get("status") in ["suspicious", "hallucination"]]
                    valid = [c for c in citations if c.get("status") == "valid"]
                    
                    print(f"\n{Fore.CYAN}📊 Resumo de Citações:")
                    print(f"   🟢 Válidas: {len(valid)}")
                    print(f"   🔴 Suspeitas/Alucinações: {len(suspicious)}")
                    
                    if suspicious:
                        print(f"\n{Fore.YELLOW}⚠️  Atenção às seguintes citações:{Style.RESET_ALL}")
                        for s in suspicious[:5]:  # Mostrar no máx 5
                            print(f"   - {s['citation']}: {s.get('message', '')}")
            
            except Exception as e:
                logger.error(f"❌ Erro na auditoria: {e}")
                print(f"{Fore.RED}❌ Auditoria falhou: {e}{Style.RESET_ALL}")
        
        elif args.audit and not AUDIT_AVAILABLE:
            print(f"{Fore.YELLOW}⚠️ Flag --audit ignorada (módulo não disponível){Style.RESET_ALL}")
        
    except KeyboardInterrupt:
        print("\n🛑 Cancelado pelo usuário.")
    except Exception as e:
        print(f"\n❌ Erro Geral: {e}")
        traceback.print_exc()
    finally:
        drafter.cleanup()



def chat_programmatic(
    message: str,
    history: List[Dict[str, str]] = None,
    context_files: List[str] = None,
    cache_ttl: int = 60,
    model: str = None,
    rag_config: Dict[str, Any] = None,
    tenant_id: str = "default",
    custom_prompt: str = None
) -> Dict[str, Any]:
    """
    Programmatic entry point for Chat with Documents.
    """
    history = history or []
    model = model or DEFAULT_MODEL
    
    # Initialize drafter
    drafter = LegalDrafter(model_name=model, cache_ttl_mins=cache_ttl)
    
    # Handle Context Caching
    doc_context = ""
    if context_files:
        valid_files = [f for f in context_files if os.path.exists(f)]
        if valid_files:
            try:
                drafter.upload_files(valid_files)
                drafter.create_context_cache(
                    system_instruction="Você é um Assistente Jurídico. Use os arquivos carregados como verdade absoluta."
                )
                doc_context += f"Contexto carregado: {len(valid_files)} arquivos.\n"
            except Exception as e:
                logger.error(f"Erro ao criar cache de contexto no chat: {e}")
                doc_context += f"Erro ao carregar arquivos: {e}\n"

    # Handle RAG Local Search
    rag_context = ""
    sources_used = []
    rag_paths = []
    if rag_config:
        raw_paths = rag_config.get("paths") or rag_config.get("context_files")
        if isinstance(raw_paths, list):
            rag_paths.extend([str(p) for p in raw_paths if str(p).strip()])
        raw_path = rag_config.get("path")
        if raw_path:
            rag_paths.append(str(raw_path))
    if not rag_paths and context_files:
        rag_paths = [str(p) for p in context_files if str(p).strip()]

    if rag_paths:
        try:
            from pathlib import Path
            from rag_local import LocalProcessIndex

            allowed_exts = {".pdf", ".txt", ".md"}
            max_files = max(1, int(settings.ATTACHMENT_RAG_LOCAL_MAX_FILES))
            files = []
            for raw_path in rag_paths:
                path = Path(raw_path)
                if path.is_dir():
                    for ext in allowed_exts:
                        files.extend(path.rglob(f"*{ext}"))
                elif path.is_file() and path.suffix.lower() in allowed_exts:
                    files.append(path)

            files = [str(p) for p in files[:max_files]]
            if files:
                local_index = LocalProcessIndex(
                    processo_id=f"upload-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}",
                    sistema="UPLOAD",
                    tenant_id=tenant_id
                )
                try:
                    for file_path in files:
                        local_index.index_documento(file_path)
                    top_k = max(1, int(settings.ATTACHMENT_RAG_LOCAL_TOP_K))
                    if rag_config:
                        top_k = int(rag_config.get("rag_top_k") or rag_config.get("top_k") or top_k)
                    results = local_index.search(message, top_k=top_k)
                finally:
                    local_index.cleanup()

                if results:
                    rag_context = "\n### 📂 Trechos Relevantes (RAG Local):\n"
                    for r in results:
                        rag_context += f"- **{r.get('citacao', 'Doc')}**: {r['text'][:400]}...\n"
                        sources_used.append({
                            "doc_id": r.get("metadata", {}).get("doc_id"),
                            "citation": r.get('citacao'),
                            "score": r.get('final_score'),
                            "text_snippet": r['text'][:100]
                        })
        except ImportError:
            logger.warning("RAG Local module could not be imported.")
        except Exception as e:
            logger.error(f"Error in RAG Local search during chat: {e}")

    # Construct Prompt
    base_instruction = custom_prompt if custom_prompt else SYSTEM_PROMPT_CHAT
    
    formatted_history = "\n".join([
        f"{'Usuário' if h['role'] == 'user' else 'Assistente'}: {h['content']}"
        for h in history
    ])
    
    full_prompt = f"""{base_instruction}

## Contexto de Arquivos (Upload):
{doc_context}

{rag_context}

## Histórico da Conversa:
{formatted_history}

## Nova Pergunta do Usuário:
{message}

## Sua Resposta:
"""

    try:
        response = drafter._generate_with_retry(
            full_prompt,
            config={"max_output_tokens": 4096, "temperature": 0.3},
            thinking_level="HIGH"
        )
        
        reply_text = response.text if response else "Não foi possível gerar uma resposta."
        
        return {
            "reply": reply_text,
            "sources": sources_used
        }
        
    except Exception as e:
        logger.error(f"Chat generation failed: {e}")
        return {
            "reply": f"Erro na geração: {e}",
            "sources": []
        }

if __name__ == "__main__":
    main()
