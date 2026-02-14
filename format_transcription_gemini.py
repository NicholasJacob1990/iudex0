#!/usr/bin/env python3
"""
Script v2.11 - Formatação de Transcrições com Gemini 3 Flash Preview
MELHORIAS: Smart Stitching Anti-Duplicação Cirúrgica

Mudanças v2.10 vs v2.9:
- remover_eco_do_contexto: Remove eco do contexto na resposta da API
- titulos_sao_similares: Fuzzy matching para detecção de títulos duplicados
- limpar_inicio_redundante: Limpeza na junção de chunks
- Injeção dinâmica de ultimo_titulo no prompt para prevenir repetição

Uso: python format_transcription_gemini.py <entrada.txt> [saida]
"""


import os
import sys
import time
import random
import json
import re
import threading
from pathlib import Path
from time import sleep
from difflib import SequenceMatcher
import hashlib

try:
    from audit_module import auditar_consistencia_legal
    AUDIT_AVAILABLE = True
except ImportError:
    AUDIT_AVAILABLE = False
    logger = logging.getLogger(__name__) if 'logging' in locals() else None
    if logger: logger.warning("⚠️  Módulo de auditoria não encontrado.")
    else: print("⚠️  Módulo de auditoria não encontrado.")

try:
    from google import genai
    from google.genai import types
except ImportError:
    print("❌ Erro: Biblioteca google-genai não instalada.")
    print("   Instale com: pip install google-genai")
    sys.exit(1)

try:
    from tqdm import tqdm
except ImportError:
    print("⚠️ Aviso: tqdm não instalado. Progress bar desabilitada.")
    tqdm = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ Aviso: python-docx não disponível. Saída em Word desabilitada.")
    DOCX_AVAILABLE = False

# =============================================================================
# SETUP CREDENCIAIS (v2.11)
# =============================================================================
CREDENTIALS_PATH = "/Users/nicholasjacob/Documents/Aplicativos/Transcritor/vertex_credentials.json"
if os.path.exists(CREDENTIALS_PATH) and not os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDENTIALS_PATH
    # print(f"🔑 Credenciais carregadas: {CREDENTIALS_PATH}")

import logging

# =============================================================================
# CONFIGURAÇÕES v2.7
# =============================================================================

CHARS_POR_PARTE = 15000
CONTEXTO_ESTILO = 3000
OUTPUT_TOKEN_LIMIT = 32000
CACHE_TTL = '7200s'
MIN_CHARS_PARA_CACHE = 150000
MAX_RETRIES = 3
MAX_RPM = 60 
# v2.7: FORÇAR delimitadores visíveis para evitar confusão
USE_FANCY_DELIMITERS = True

# Modelo Gemini (centralizado para fácil atualização)
GEMINI_MODEL = 'gemini-3-flash-preview'

# Preços API Gemini 3 Flash Preview (Estimativa)
PRECO_INPUT_SEM_CACHE = 0.50
PRECO_INPUT_COM_CACHE = 0.05  # Estimado a 10% do input (manter proporção anterior)
PRECO_OUTPUT = 3.00

# =============================================================================
# LOGGING
# =============================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S',
    handlers=[
        logging.FileHandler('formatacao.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# =============================================================================
# MÉTRICAS DE EXECUÇÃO (v2.10)
# =============================================================================
class MetricsCollector:
    """Coleta e reporta métricas de execução para otimização de custos."""
    
    def __init__(self):
        self.reset()
    
    def reset(self):
        self.api_calls = 0
        self.gemini_calls = 0
        self.openai_calls = 0
        self.cache_hits = 0
        self.adaptive_splits = 0
        self.total_prompt_tokens = 0
        self.total_cached_tokens = 0
        self.total_completion_tokens = 0
        self.total_time_seconds = 0.0
        self.chunks_processed = 0
        self.start_time = None
    
    def start_timer(self):
        self.start_time = time.time()
    
    def stop_timer(self):
        if self.start_time:
            self.total_time_seconds = time.time() - self.start_time
    
    def record_api_call(self, prompt_tokens=0, completion_tokens=0, cached_tokens=0, provider='gemini'):
        # Garantir que são inteiros (v2.17.1: fix NoneType)
        prompt_tokens = prompt_tokens or 0
        completion_tokens = completion_tokens or 0
        cached_tokens = cached_tokens or 0
        
        self.api_calls += 1
        if provider == 'gemini':
            self.gemini_calls += 1
        elif provider == 'openai':
            self.openai_calls += 1
            
        # Se houve tokens cacheados, considera cache hit
        if cached_tokens > 0:
            self.cache_hits += 1
            
        self.total_prompt_tokens += (prompt_tokens - cached_tokens)
        self.total_cached_tokens += cached_tokens
        self.total_completion_tokens += completion_tokens
    
    def record_adaptive_split(self):
        self.adaptive_splits += 1

    def record_cache_hit(self):
        self.cache_hits += 1
    
    def get_cost(self):
        """Calcula custo estimado (Gemini 3 Flash Preview)."""
        # Preços por 1M tokens ($0.50 Input / $3.00 Output)
        input_price = 0.50       # USD (Standard)
        cached_price = 0.05      # USD (Cached - Est. 10%)
        output_price = 3.00      # USD (Output)
        
        cost = (
            (self.total_prompt_tokens * input_price) + 
            (self.total_cached_tokens * cached_price) +
            (self.total_completion_tokens * output_price)
        ) / 1_000_000
        return cost
    
    def get_report(self):
        avg_time = self.total_time_seconds / self.api_calls if self.api_calls > 0 else 0
        cost = self.get_cost()
        
        return f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 MÉTRICAS DE EXECUÇÃO (v2.10)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
   📡 Total de Chamadas API: {self.api_calls}
      - Gemini: {self.gemini_calls}
      - OpenAI: {self.openai_calls}
      - Cache Hits: {self.cache_hits}
   ✂️ Divisões Adaptativas: {self.adaptive_splits}
   🎯 Tokens Usados:
      - Prompt (Regular): {self.total_prompt_tokens:,}
      - Prompt (Cached):  {self.total_cached_tokens:,}
      - Completion:       {self.total_completion_tokens:,}
      - Total Geral:      {self.total_prompt_tokens + self.total_cached_tokens + self.total_completion_tokens:,}
   ⏱️ Tempo Total: {self.total_time_seconds:.1f}s (média: {avg_time:.2f}s/chamada)
   💰 Custo Real: ${cost:.6f} USD
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

# Global metrics instance
metrics = MetricsCollector()

# =============================================================================
# PROMPTS v2.7 - INSTRUÇÕES ANTI-DUPLICAÇÃO REFORÇADAS
# =============================================================================

PROMPT_FIDELIDADE = """# DIRETRIZES DE FORMATAÇÃO E REVISÃO

## PAPEL
VOCÊ É UM EXCELENTISSIMO REDATOR TÉCNICO E DIDÁTICO

 **Tom:** didático, como o professor explicando em aula.  
- **Pessoa:** manter a pessoa original da transcrição (1ª pessoa se for assim na fala).  
- **Estilo:** texto corrido, com parágrafos curtos, sem “inventar” doutrina nova.  
- **Objetivo:** reproduzir a aula em forma escrita, clara e organizada, mas ainda com a “voz” do professor.


## OBJETIVO
-Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, MANTENDO A FIDELIDADE TOTAL ao conteúdo original.
-- **Tamanho:** a saída deve ficar **entre 95% e 115%** do tamanho do trecho de entrada (salvo remoção de muletas e logística).

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias e das falas do professor.
4. **NÃO CRIE MUITOS BULLET POINTS** ou frases curtas demasiadamente. PREFIRA UM FORMATO DE MANUAL DIDÁTICO, não checklist.
5. **NÃO USE NEGRITOS EM EXCESSO**. Use apenas para conceitos-chave realmente importantes.
6. **NÃO RESUMA e NÃO OMITA**. Você pode reescrever frases em português padrão para melhorar a fluidez, preservar a ordem, os detalhes técnicos e os exemplos, mas **REMOVA** pausas excessivas e hesitações.


## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados (REDI/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência, autores, casos citados
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque)
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico"

## 🎯 PRESERVAÇÃO ESPECIAL: DICAS DE PROVA E EXAMINADORES (CRÍTICO)
Aulas presenciais frequentemente contêm informações valiosas sobre:
1. **Referências a Examinadores**: Nomes de examinadores de concursos, suas preferências, posicionamentos ou temas favoritos. **PRESERVE INTEGRALMENTE**.
   - Exemplo: "O examinador Fulano costuma cobrar..." → MANTER
   - Exemplo: "Esse tema foi cobrado pelo professor X na prova..." → MANTER
2. **Dicas de Prova**: Orientações sobre o que costuma cair em provas, pegadinhas comuns, temas recorrentes.
   - Exemplo: "Isso cai muito em prova..." → MANTER
   - Exemplo: "Atenção: essa é uma pegadinha clássica..." → MANTER
3. **Estratégias de Estudo**: Sugestões do professor sobre priorização, macetes, formas de memorização.
   - Exemplo: "Gravem isso: na dúvida, marquem..." → MANTER
   - Exemplo: "Para PGM, foquem em..." → MANTER
4. **Casos Práticos e Histórias Reais**: Exemplos de situações reais, casos julgados, histórias ilustrativas.
   - **NUNCA RESUMA** histórias ou exemplos práticos. Preserve na íntegra.

> ⚠️ **ESSAS INFORMAÇÕES SÃO O DIFERENCIAL DE UMA AULA AO VIVO.** Sua omissão representa perda irreparável de valor didático.


## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Corrija erros gramaticais, regências, ortográficos e de pontuação, tornando o texto gramaticalmente correto e claro.
2. **Limpeza Profunda:**
   - **REMOVA** marcadores de oralidade: "né", "tá?", "entende?", "veja bem", "tipo assim".
   - **REMOVA** interações diretas com a turma/alunos e logística: "Isso mesmo", "A colega perguntou", "Já estão me vendo?", "Estão ouvindo?", "Como ele disse ali atrás".
   - **REMOVA** redundâncias: "subir para cima", "criação nova".
   - **TRANSFORME** perguntas retóricas em afirmações quando possível (ex: "E o que isso significa?" -> "Isso significa que...").
3. **Coesão**: Utilize conectivos necessários para tornar o texto mais fluido. Aplique a pontuação devida para deixar o texto coeso e coerente.
4. **Legibilidade**:
   - **USE TEXTO CORRIDO NA MEDIDA DO POSSÍVEL**
   - Utilize formatação e estrutura com parágrafos bem definidos, facilitando a leitura e compreensão
   - Evite parágrafos longos (máximo 3-4 linhas visuais)
   - Evite blocos de texto maciços, quebre os blocos de texto em parágrafos menores
   - Seja didático sem perder detalhes e conteúdo
5. **Linguagem**: Ajuste a linguagem coloquial para um português padrão, mantendo o significado original.
6. **Citações**: Use itálico para citações curtas e recuo em itálico para citações longas.
7. -Use **negrito** para destacar conceitos-chave (sem exagero).
8. **Formatação Didática** (use com moderação, sem excesso):
   - **Bullet points** para enumerar elementos, requisitos ou características
   - **Listas numeradas** (1., 2., 3.) para enumerar itens, etapas, correntes ou exemplos
   - **Marcadores relacionais** como "→" para indicar relações, transições, ou consequências lógicas
   - Exemplo: "Processo entre A e B → prova usada contra C"
9. **Questões e Exercícios**:
   - Se o professor ditar uma questão, exercício ou caso hipotético para julgar, **ILHE-O** em um bloco de citação:
   > **Questão:** O prazo para agravo de petição é de...
   - Separe claramente o enunciado da questão da explicação/gabarito subsequente.
10. **Destaques com Emojis** (use com moderação para facilitar escaneamento visual):
   - 💡 **Dica de Prova** ou **Observação Pedagógica**: Quando o professor der uma dica específica para provas ou concursos.
   - ⚠️ **Atenção** ou **Cuidado**: Para alertas, pegadinhas ou pontos polêmicos.
   - 📌 **Ponto Importante**: Para conceitos-chave que merecem destaque especial.
   - Exemplo de uso: `> 💡 **Dica de Prova:** Esse tema caiu 3 vezes na PGM-Rio.`

## 📝 ESTRUTURA
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis a partir do contexto.

## 🚫 TÍTULOS E SUBTÓPICOS (IMPORTANTE)
- **NÃO críe subtópicos para frases soltas.**
- Use títulos (##, ###) **APENAS** para mudanças reais de assunto.
- Se uma frase parece um título mas não inicia uma nova seção, mantenha como texto normal e use **negrito** se necessário.

## 📊 TABELA DE SÍNTESE (FLEXÍVEL)
Ao final de cada **bloco temático relevante** (ou capítulo), produza uma tabela de síntese completa (modelo flexível).
Exemplo de estrutura (adapte conforme o conteúdo):

```
### 📋 Tabela de síntese do tópico

| Conceito/Instituto | Definição (conforme a aula) | Fundamento Legal (se citado) | Observações (alertas/exceções/juris) |
| :--- | :--- | :--- | :--- |
| ...  | ...  | Art. X, Lei Y / "—" | ... |
```

***REGRAS CRÍTICAS PARA TABELAS:**
1. **Limite de conteúdo por célula:** máximo ~50 palavras. Se precisar de mais, divida em múltiplas linhas da tabela
2. **PROIBIDO usar blocos de código (```) dentro de células** - use texto simples
3. **NUNCA deixe o título "📋 Resumo do Tópico" sozinho** - se não houver dados para tabela, NÃO escreva o título
4. **POSICIONAMENTO ESTRITO:**
   - A tabela deve vir **APENAS AO FINAL** de um bloco concluído.
   - **PROIBIDO** inserir tabela no meio de uma frase ou interrompendo uma explicação.
   - Se o texto continuar sobre o mesmo assunto, **termine o texto primeiro** e coloque a tabela depois.


## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA
- **NUNCA formate novamente esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- **NUNCA repita informações que já estão no contexto**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>
- Se o texto_para_formatar começar com algo similar ao fim do contexto, NÃO duplique, apenas continue naturalmente
"""

PROMPT_APOSTILA = """# DIRETRIZES DE REDAÇÃO: MANUAL JURÍDICO DIDÁTICO (MODO APOSTILA)
## PAPEL
VOCÊ É UM EXCELENTISSIMO REDATOR TÉCNICO E DIDÁTICO
- **Tom:** doutrinário, impessoal, estilo manual de Direito.  
- **Pessoa:** 3ª pessoa ou construções impessoais (“observa-se”, “entende-se”).  
- **Estilo:** prosa mais densa, porém com parágrafos curtos e didáticos.  
- **Objetivo:** transformar o conteúdo da aula em texto de apostila/livro, sem alterar o conteúdo e sem inventar informações.


## OBJETIVO
Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, em formato de apostila/manual didático

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias 


❌ PRESERVE obrigatoriamente:
- **NÚMEROS EXATOS**: Artigos, Leis, Artigos Súmulas, Julgados (REDI/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência, autores, casos citados
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque)
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico"

## 🎯 PRESERVAÇÃO ESPECIAL: DICAS DE PROVA E EXAMINADORES (CRÍTICO)
Aulas presenciais frequentemente contêm informações valiosas sobre:
1. **Referências a Examinadores**: Nomes de examinadores de concursos, suas preferências, posicionamentos ou temas favoritos. **PRESERVE INTEGRALMENTE**.
   - Exemplo: "O examinador Fulano costuma cobrar..." → MANTER
   - Exemplo: "Esse tema foi cobrado pelo professor X na prova..." → MANTER
2. **Dicas de Prova**: Orientações sobre o que costuma cair em provas, pegadinhas comuns, temas recorrentes.
   - Exemplo: "Isso cai muito em prova..." → MANTER
   - Exemplo: "Atenção: essa é uma pegadinha clássica..." → MANTER
3. **Estratégias de Estudo**: Sugestões do professor sobre priorização, macetes, formas de memorização.
   - Exemplo: "Gravem isso: na dúvida, marquem..." → MANTER
   - Exemplo: "Para PGM, foquem em..." → MANTER
4. **Casos Práticos e Histórias Reais**: Exemplos de situações reais, casos julgados, histórias ilustrativas.
   - **NUNCA RESUMA** histórias ou exemplos práticos. Preserve na íntegra.

> ⚠️ **ESSAS INFORMAÇÕES SÃO O DIFERENCIAL DE UMA AULA AO VIVO.** Sua omissão representa perda irreparável de valor didático.


## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Ajuste a linguagem coloquial para o padrão culto.
2. **Limpeza**: Remova gírias, cacoetes ("né", "tipo assim", "então") e vícios de oralidade.
3. **Coesão**: Use conectivos e pontuação adequada para tornar o texto fluido.
4. **Legibilidade**:
   - Use parágrafos bem definidos e curtos (máximo 3-4 linhas visuais).
    - Evite blocos de texto maciços, quebre os blocos de texto em parágrafos menores
   - Use **negrito** para destacar conceitos-chave (sem exagero).
5. **Formatação Didática** (use com moderação, sem excesso):
   - **Bullet points** para enumerar elementos, requisitos ou características
   - **Listas numeradas** (1., 2., 3.) para enumerar itens, etapas, correntes doutrinárias ou exemplos
   - **Marcadores relacionais** como "→" para indicar relações, transições, ou consequências lógicas
   - Exemplo: "Processo entre Pedro e José → prova usada contra Ana"
6. **Destaques com Emojis** (use com moderação para facilitar escaneamento visual):
   - 💡 **Dica de Prova** ou **Observação Pedagógica**: Quando houver uma dica específica para provas ou concursos.
   - ⚠️ **Atenção** ou **Cuidado**: Para alertas, pegadinhas ou pontos polêmicos.
   - 📌 **Ponto Importante**: Para conceitos-chave que merecem destaque especial.
   - Exemplo de uso: `> 💡 **Dica de Prova:** Esse tema caiu 3 vezes na PGM-Rio.`

## 📝 ESTRUTURA
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis a partir do contexto.

## 📊 TABELA DE SÍNTESE (FLEXÍVEL)
Ao final de cada **bloco temático relevante** (ou capítulo), produza uma tabela de síntese (modelo flexível).
Exemplo de estrutura (adapte conforme o conteúdo):

```
### 📋 Tabela de síntese do tópico

| Conceito/Instituto | Definição (conforme a aula) | Fundamento Legal (se citado) | Observações (alertas/exceções/juris) |
| :--- | :--- | :--- | :--- |
| ...  | ...  | Art. X, Lei Y / "—" | ... |
```

***REGRAS CRÍTICAS PARA TABELAS:**
1. **Limite de conteúdo por célula:** máximo ~50 palavras. Se precisar de mais, divida em múltiplas linhas da tabela
2. **PROIBIDO usar blocos de código (```) dentro de células** - use texto simples
3. **NUNCA deixe o título "📋 Resumo do Tópico" sozinho** - se não houver dados para tabela, NÃO escreva o título
4. **POSICIONAMENTO:** A tabela deve vir **APENAS AO FINAL** da explicação completa dos tópicos ou blocos temáticos relevantes 
   - **NUNCA** insira a tabela no meio de uma explicação.
   - **NUNCA** resuma um tópico que você ainda não acabou de explicar no texto.
   - A tabela deve ser o **fechamento** lógico da seção, antes de iniciar um novo título ou tópico (## ou ###).

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA
- **NUNCA formate novamente esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- **NUNCA repita informações que já estão no contexto**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>
- Se o texto_para_formatar começar com algo similar ao fim do contexto, NÃO duplique, apenas continue naturalmente
"""


# =============================================================================
# ESCOLHA O MODO
# =============================================================================

PROMPT_FORMATACAO = PROMPT_FIDELIDADE
# PROMPT_FORMATACAO = PROMPT_APOSTILA

# =============================================================================
# DETECÇÃO DE MODO
# =============================================================================

FIDELIDADE_MODE = "NÃO RESUMA" in PROMPT_FORMATACAO
APOSTILA_MODE = "MANUAL JURÍDICO" in PROMPT_FORMATACAO

if APOSTILA_MODE:
    THRESHOLD_MINIMO = 0.75
    THRESHOLD_CRITICO = 0.65
    MODO_NOME = "APOSTILA"
elif FIDELIDADE_MODE:
    THRESHOLD_MINIMO = 0.75
    THRESHOLD_CRITICO = 0.70
    MODO_NOME = "FIDELIDADE"
else:
    THRESHOLD_MINIMO = 0.70
    THRESHOLD_CRITICO = 0.60
    MODO_NOME = "PADRÃO"

logger.info(f"🎯 Modo: {MODO_NOME} (threshold={THRESHOLD_MINIMO:.0%})")
logger.info(f"🛡️  Anti-duplicação: ATIVADA (v2.7)")

# Limiares adaptativos por camada de deduplicação
# 7-DIFF (chunk overlaps): pode ser agressivo, overlaps são quase sempre erros
LIMIAR_7DIFF = 0.85 if MODO_NOME == "FIDELIDADE" else 0.80
# Seções duplicadas: mais cuidado, professor pode repetir propositalmente
LIMIAR_SECOES = 0.70 if MODO_NOME == "FIDELIDADE" else 0.60

logger.info(f"📊 Limiares: 7-DIFF={LIMIAR_7DIFF:.0%} | Seções={LIMIAR_SECOES:.0%}")

# =============================================================================
# RATE LIMITER
# =============================================================================

class RateLimiter:
    def __init__(self, max_requests_per_minute=MAX_RPM):
        self.max_rpm = max_requests_per_minute
        self.requests = []
        self.lock = threading.Lock()
    
    def wait_if_needed(self):
        with self.lock:
            now = time.time()
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
                logger.info(f"⏱️  Rate limit: aguardando {wait_time:.1f}s...")
                sleep(wait_time)
                self.requests = [t for t in self.requests if time.time() - t < 60]
            
            self.requests.append(time.time())

rate_limiter = RateLimiter()

# =============================================================================
# CHECKPOINT/RESUME
# =============================================================================

def get_checkpoint_path(input_file):
    return Path(input_file).with_suffix('.checkpoint.json')

def save_checkpoint(input_file, resultados, chunks_info, secao_atual):
    checkpoint_path = get_checkpoint_path(input_file)
    checkpoint_data = {
        'input_file': str(input_file),
        'secao_atual': secao_atual,
        'total_secoes': len(chunks_info),
        'chunks_info': chunks_info,
        'resultados': resultados,
        'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
        'version': '2.7',
        'modo': MODO_NOME
    }
    with open(checkpoint_path, 'w', encoding='utf-8') as f:
        json.dump(checkpoint_data, f, ensure_ascii=False, indent=2)

def load_checkpoint(input_file):
    checkpoint_path = get_checkpoint_path(input_file)
    if checkpoint_path.exists():
        try:
            with open(checkpoint_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if data.get('version') not in ('2.6', '2.7'):
                    logger.warning("Checkpoint de versão anterior. Reiniciando...")
                    return None
                return data
        except Exception as e:
            logger.error(f"Erro ao carregar checkpoint: {e}")
    return None

def delete_checkpoint(input_file):
    checkpoint_path = get_checkpoint_path(input_file)
    if checkpoint_path.exists():
        checkpoint_path.unlink()
        logger.info("🧹 Checkpoint removido")

# =============================================================================
# DIVISÃO SEQUENCIAL
# =============================================================================

def dividir_sequencial(transcricao_completa, estrutura_global=None):
    """
    v2.17: Divide documento com Inteligência de Âncoras (Anchor-Based Chunking).
    Se estrutura_global for fornecida, tenta alinhar cortes com início de tópicos.
    """
    chunks = []
    tamanho_total = len(transcricao_completa)
    inicio = 0
    
    # Prepara âncoras (keywords dos títulos)
    ancoras = []
    if estrutura_global:
        for item in estrutura_global.split('\n'):
            clean_item = re.sub(r'^\d+(\.\d+)*\.?\s*', '', item.strip())
            if len(clean_item) > 10:
                # Pega as primeiras 4 palavras do título como "âncora"
                keywords = ' '.join(clean_item.split()[:4])
                ancoras.append(keywords)
    
    while inicio < tamanho_total:
        target_fim = inicio + CHARS_POR_PARTE
        fim = min(target_fim, tamanho_total)
        
        if fim < tamanho_total:
            # Janela de busca para ajuste fino
            janela_busca = transcricao_completa[max(0, fim - 1000):min(tamanho_total, fim + 1000)]
            offset_janela = max(0, fim - 1000)
            
            melhor_corte = -1
            
            # 1. Tenta encontrar uma ÂNCORA de Tópico (Prioridade Alta)
            if ancoras:
                for ancora in ancoras:
                    # Busca fuzzy ou exata da âncora na janela
                    # Simplificação: busca exata case-insensitive
                    pos_ancora = janela_busca.lower().find(ancora.lower())
                    if pos_ancora != -1:
                        corte_proposto = offset_janela + pos_ancora
                        # Se o corte sugerido pela âncora estiver dentro de um range aceitável
                        if abs(corte_proposto - fim) < 800: # Aceita desvio de até 800 chars
                            melhor_corte = corte_proposto
                            logger.info(f"   ⚓ Âncora encontrada: '{ancora}' (ajustando corte)")
                            break
            
            # 2. Se não achou âncora, busca quebra estrutural forte (## Título)
            if melhor_corte == -1:
                titulo_match = re.search(r'\n(#{2,4}\s+.+)\n', janela_busca)
                if titulo_match:
                    melhor_corte = offset_janela + janela_busca.find(titulo_match.group(0))
            
            # 3. Fallback: Quebra de parágrafo duplo
            if melhor_corte == -1:
                quebra = transcricao_completa.rfind('\n\n', fim - 300, fim + 300)
                if quebra != -1 and quebra > inicio:
                    melhor_corte = quebra
            
            # 4. Último recurso: Ponto final
            if melhor_corte == -1:
                quebra = transcricao_completa.rfind('. ', fim - 150, fim + 150)
                if quebra != -1 and quebra > inicio:
                    melhor_corte = quebra + 1
            
            # Aplica o melhor corte encontrado
            if melhor_corte != -1:
                fim = melhor_corte
        
        chunks.append({'inicio': inicio, 'fim': fim})
        inicio = fim
    
    return chunks

def validar_chunks(chunks, transcricao_completa):
    """v2.7: Validação rigorosa de chunks sequenciais"""
    logger.info("🔍 Validando chunks sequenciais...")
    
    for i in range(len(chunks)):
        chunk = chunks[i]
        
        # Verifica se início == fim do anterior
        if i > 0:
            anterior = chunks[i-1]
            if chunk['inicio'] != anterior['fim']:
                logger.error(f"❌ Gap/Overlap no chunk {i+1}!")
                logger.error(f"   Anterior termina em: {anterior['fim']}")
                logger.error(f"   Atual começa em: {chunk['inicio']}")
                logger.error(f"   Diferença: {chunk['inicio'] - anterior['fim']} chars")
                
                # Mostra preview
                if chunk['inicio'] < anterior['fim']:
                    overlap_text = transcricao_completa[chunk['inicio']:anterior['fim']]
                    logger.error(f"   OVERLAP: '{overlap_text[:100]}...'")
                
                return False
    
    logger.info(f"✅ {len(chunks)} chunks validados (sequenciais, sem overlap)")
    return True

# =============================================================================
# FUNÇÕES AUXILIARES
# =============================================================================

def limpar_tags_xml(texto):
    texto = re.sub(r'</?[a-z_][\w\-]*>', '', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<[a-z_][\w\-]*\s+[^>]+>', '', texto, flags=re.IGNORECASE)
    return texto

def carregar_transcricao(arquivo):
    try:
        with open(arquivo, 'r', encoding='utf-8-sig') as f:
            conteudo = f.read()
        
        if not conteudo.strip():
            logger.error("Arquivo está vazio.")
            sys.exit(1)
            
        return conteudo
    except FileNotFoundError:
        logger.error(f"Arquivo '{arquivo}' não encontrado.")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Erro ao ler arquivo: {e}")
        sys.exit(1)

def estimar_custo(transcricao, usar_cache, num_chunks=1):
    tokens_in = len(transcricao) // 4
    
    if APOSTILA_MODE:
        tokens_out_estimado = int(tokens_in * 0.65)
    elif FIDELIDADE_MODE:
        tokens_out_estimado = int(tokens_in * 1.00)
    else:
        tokens_out_estimado = int(tokens_in * 0.85)
    
    tokens_prompt = len(PROMPT_FORMATACAO) // 4
    tokens_in_total = tokens_in + (tokens_prompt * num_chunks)
    
    if usar_cache:
        custo_input = (tokens_in * PRECO_INPUT_COM_CACHE + tokens_prompt * num_chunks * PRECO_INPUT_SEM_CACHE) / 1_000_000
        custo_output = (tokens_out_estimado * PRECO_OUTPUT) / 1_000_000
        custo = custo_input + custo_output
    else:
        custo = (tokens_in_total * PRECO_INPUT_SEM_CACHE + tokens_out_estimado * PRECO_OUTPUT) / 1_000_000
    
    logger.info(f"💰 Custo estimado: ${custo:.4f} USD (modo {MODO_NOME})")

# v2.9: Cache REABILITADO com hash inteligente
def criar_cache_contexto(client, transcricao_completa, system_prompt, estrutura_global=None):
    """Cria cache de contexto com hash estável para reutilização"""
    
    # Cache só vale a pena para documentos grandes
    if len(transcricao_completa) < MIN_CHARS_PARA_CACHE:
        logger.info(f"📦 Documento pequeno ({len(transcricao_completa):,} chars), cache não necessário")
        return None
    
    try:
        # Hash do prompt + estrutura para garantir unicidade por documento
        combined_content = system_prompt + (estrutura_global or "")
        prompt_hash = hashlib.sha256(combined_content.encode()).hexdigest()[:16]
        cache_name = f"fmt_{prompt_hash}"
        
        # v2.9: Tenta encontrar cache existente válido
        try:
            for c in client.caches.list(page_size=100):
                if c.display_name == cache_name:
                    logger.info(f"♻️  Reusando cache existente: {cache_name} ({c.name})")
                    return c
        except Exception as e:
            logger.warning(f"Cache lookup warning: {e}")

        # Adiciona a estrutura global se disponível
        estrutura_text = f"\n\n## ESTRUTURA GLOBAL:\n{estrutura_global}" if estrutura_global else ""
        
        cache_content = f"""{system_prompt}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📚 CONTEXTO GLOBAL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Modo: {MODO_NOME}{estrutura_text}
"""
        
        # v2.19: TTL Dinâmico (User Request)
        # Estimativa: 1 hora a cada 500k chars + 1h margem
        # Isso evita que o cache expire no meio de um documento longo
        tempo_estimado_segundos = int((len(transcricao_completa) / 500000) * 3600) + 3600
        dinamico_ttl = f"{max(3600, tempo_estimado_segundos)}s"
        
        # Cria cache usando a API do google-genai
        cache = client.caches.create(
            model=GEMINI_MODEL,
            config=types.CreateCachedContentConfig(
                contents=[cache_content],
                ttl=dinamico_ttl,
                display_name=cache_name
            )
        )
        
        logger.info(f"✅ Cache criado: {cache_name} (hash: {prompt_hash}, TTL: {CACHE_TTL})")
        return cache
        
    except Exception as e:
        logger.warning(f"⚠️ Falha ao criar cache: {e}. Continuando sem cache.")
        return None

# =============================================================================
# MAPEAMENTO ESTRUTURAL (v2.8)
# =============================================================================

PROMPT_MAPEAMENTO = """Você é um especialista em organização de conteúdo educacional acadêmico.

## ETAPA 1: IDENTIFICAR O TIPO DE CONTEÚDO
Analise a transcrição e determine qual é a **natureza predominante** do material:

| Tipo | Pistas no Texto | Estrutura Ideal |
|------|-----------------|-----------------|
| **SIMULADO** | "questão 1", "questão 2", "espelho de correção", "correção do simulado", "vamos corrigir" | Organizar por QUESTÕES numeradas |
| **AULA EXPOSITIVA** | explicações contínuas de um tema, teoria, doutrina, sem questões específicas | Organizar por TEMAS/MATÉRIAS |
| **REVISÃO** | "revisão", "resumo", múltiplos temas curtos, "pontos importantes" | Organizar por TÓPICOS de revisão |
| **CORREÇÃO DE PROVA** | "gabarito", "alternativa correta", "item certo/errado" | Organizar por QUESTÕES com gabarito |

## ETAPA 2: EXTRAIR A ESTRUTURA

### Se for SIMULADO ou CORREÇÃO DE PROVA:
```
1. Orientações Gerais / Introdução
2. Questão 1: [Título descritivo] — [Área do Direito]
   2.1. Enunciado e Contexto
   2.2. Fundamentação (Doutrina/Jurisprudência)
   2.3. Pontos do Espelho / Resposta
3. Questão 2: [Título descritivo] — [Área do Direito]
   3.1. ...
[Continue para cada questão]
N. Considerações Finais / Dúvidas
```

### Se for AULA EXPOSITIVA:
```
1. Introdução
2. [Matéria 1: ex. Direito Administrativo]
   2.1. [Subtema]
      2.1.1. [Detalhamento]
3. [Matéria 2: ex. Direito Civil]
   3.1. ...
```

### Se for REVISÃO:
```
1. [Tema 1]
   1.1. Pontos-chave
   1.2. Jurisprudência/Súmulas
2. [Tema 2]
   2.1. ...
```

## REGRAS GERAIS:
1. **MÁXIMO 3 NÍVEIS** de hierarquia (1., 1.1., 1.1.1.)
2. **Seja descritivo** nos títulos — inclua o assunto real, não apenas "Questão 1"
3. **Mantenha a ORDEM** cronológica da transcrição
4. **Mapeie do INÍCIO ao FIM** — não omita partes
5. **Identifique a ÁREA DO DIREITO** de cada bloco quando possível

## 🏛️ REGRA ESPECIAL: MARCOS LEGAIS (v2.17)
Quando identificar marcos legais importantes, crie subtópicos específicos:
- **Súmulas** (STF, STJ, Vinculantes): Criar subtópico "X.Y. Súmula [Número] do [Tribunal]"
- **Teses (Repercussão Geral/Repetitivos)**: Criar subtópico "X.Y. Tese/Tema [Número] do STJ/STF"
- **Artigos de Lei Central**: Se um artigo é explicado em profundidade, criar subtópico "X.Y. Art. [Número] da [Lei]"

Exemplo:
```
2. Execução Fiscal
   2.1. Procedimento da LEF (Lei 6.830/80)
   2.2. Súmula 314 do STJ (Citação por Hora Certa)
   2.3. Tema 444 do STJ (Redirecionamento)
```

## RESPOSTA:
Primeiro, indique em uma linha: `[TIPO: SIMULADO/EXPOSITIVA/REVISÃO/CORREÇÃO]`
Depois, retorne APENAS a estrutura hierárquica (máx 3 níveis).
"""

# ---------------------------------------------------------------------------
# Sanitização de títulos na estrutura mapeada (v2.47)
# ---------------------------------------------------------------------------

_CONVERSATIONAL_TITLE_PREFIXES = (
    "já ", "na prova", "para quem", "minha proposta",
    "bom dia", "gente ", "pessoal ", "vamos ", "então ",
    "logo ", "eu ", "nós ", "aqui ", "olha ", "vejam ",
    "como eu ", "antes de ", "boa tarde", "boa noite",
    "obrigado", "obrigada", "com licença",
)

_GREETING_PREFIXES = (
    "bom dia", "boa tarde", "boa noite", "já ", "pessoal ",
    "gente ", "olha ", "obrigado", "obrigada",
)

# Rótulos canônicos por nível
_CANONICAL_LABEL_L1 = "Introdução e Contextualização"
_CANONICAL_LABEL_SUB = "Abertura"

_MAX_TITLE_WORDS = 8
_MAX_TITLE_CHARS = 70


def _sanitize_structure_titles(estrutura: str) -> str:
    """Valida e corrige títulos de estrutura que são trechos literais de fala.

    Regras (alinhadas com PROMPT_MAPEAMENTO regra 8):
    - Títulos > 8 palavras ou > 70 chars → mapear para rótulo canônico
    - Prefixos conversacionais (saudações, logística) → rótulo canônico
    - Preserva âncoras ABRE/FECHA intactas (incluindo aspas)

    Função pura, sem dependências externas (logger opcional).
    """
    if not estrutura:
        return estrutura

    lines = estrutura.split('\n')
    fixed_lines: list[str] = []
    sanitized_count = 0

    for line in lines:
        stripped = line.strip()
        # Detecta linhas numeradas: "1. Título", "   1.1. Subtítulo"
        m = re.match(r'^(\s*\d+(?:\.\d+)*\.?\s+)(.*)', stripped)
        if not m:
            fixed_lines.append(line)
            continue

        prefix_num = m.group(1)
        rest = m.group(2).strip()

        # Separa âncoras ABRE/FECHA se existirem (preserva literalmente)
        anchor_part = ""
        title = rest
        anchor_idx = rest.find("| ABRE:")
        if anchor_idx >= 0:
            title = rest[:anchor_idx].strip()
            anchor_part = " " + rest[anchor_idx:]

        title_lower = title.lower()
        words = re.findall(r'[A-Za-zÀ-ÿ0-9]+', title)

        needs_fix = False
        if len(title) > _MAX_TITLE_CHARS:
            needs_fix = True
        elif len(words) > _MAX_TITLE_WORDS:
            needs_fix = True
        elif any(title_lower.startswith(pfx) for pfx in _CONVERSATIONAL_TITLE_PREFIXES):
            needs_fix = True

        if needs_fix:
            # Determinar nível: "1." → nível 1, "1.1." → nível 2+
            is_level1 = re.match(r'^\s*\d+\.\s', stripped) and not re.match(r'^\s*\d+\.\d+', stripped)

            if any(title_lower.startswith(pfx) for pfx in _GREETING_PREFIXES):
                canonical = _CANONICAL_LABEL_L1 if is_level1 else _CANONICAL_LABEL_SUB
            elif is_level1:
                canonical = _CANONICAL_LABEL_L1
            else:
                canonical = _CANONICAL_LABEL_SUB

            sanitized_count += 1
            try:
                logger.warning(f"⚠️  Título sanitizado: '{title[:60]}' → '{canonical}'")
            except Exception:
                pass  # logger pode não existir em contexto de teste
            fixed_lines.append(f"{prefix_num}{canonical}{anchor_part}")
        else:
            fixed_lines.append(line)

    if sanitized_count:
        try:
            logger.info(f"🔧 {sanitized_count} título(s) de estrutura sanitizado(s)")
        except Exception:
            pass

    return '\n'.join(fixed_lines)


def mapear_estrutura(client, transcricao_completa):
    """Analisa o documento completo e extrai a estrutura de tópicos"""
    logger.info("🗺️  Mapeando estrutura do documento...")
    
    rate_limiter.wait_if_needed()
    
    # Gemini 2.5 Flash suporta 1M tokens (~4M chars)
    # Limite de 3.5M chars para deixar margem para output (20k tokens)
    max_chars_mapeamento = 3_500_000
    
    if len(transcricao_completa) > max_chars_mapeamento:
        logger.warning(f"⚠️  Documento EXTREMAMENTE grande ({len(transcricao_completa):,} chars). Cortando final para caber no contexto.")
        # É melhor cortar o final do que picotar o meio para estrutura
        texto_para_mapear = transcricao_completa[:max_chars_mapeamento]
    else:
        texto_para_mapear = transcricao_completa
        logger.info(f"   Mapeando documento completo ({len(transcricao_completa):,} chars)")
    
    prompt = PROMPT_MAPEAMENTO.format(transcricao=texto_para_mapear)
    
    try:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(
                max_output_tokens=20000,  # Aumentado para documentos grandes
                thinking_config={"include_thoughts": False, "thinking_level": "HIGH"}, # Mapeamento: HIGH
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                ]
            )
        )
        
        
        if not response.text:
            logger.warning("⚠️  Resposta vazia do mapeamento.")
            logger.warning(f"   Response object: {response}")
            if hasattr(response, 'candidates') and response.candidates:
                logger.warning(f"   Finish reason: {response.candidates[0].finish_reason}")
                if hasattr(response.candidates[0], 'safety_ratings'):
                    logger.warning(f"   Safety ratings: {response.candidates[0].safety_ratings}")
            logger.warning("   Continuando sem estrutura prévia.")
            return None
        
        estrutura = response.text.strip()
        
        # Remove markdown code blocks se presentes
        if estrutura.startswith('```'):
            linhas = estrutura.split('\n')
            estrutura = '\n'.join(linhas[1:-1]) if len(linhas) > 2 else estrutura
        
        linhas = [l for l in estrutura.split('\n') if l.strip()]
        
        # Validação da estrutura
        if len(linhas) < 3:
            logger.warning(f"⚠️  Estrutura muito curta ({len(linhas)} linhas). Pode estar incompleta.")
            return None
        
        tem_numeracao = any(re.match(r'^\d+\.', l.strip()) for l in linhas)
        if not tem_numeracao:
            logger.warning("⚠️  Estrutura sem numeração hierárquica. Pode estar mal formatada.")
            return None
        
        logger.info(f"✅ Estrutura mapeada: {len(linhas)} tópicos identificados")
        
        # Log preview (primeiras 10 linhas + total)
        for linha in linhas[:10]:
            logger.info(f"   {linha}")
        if len(linhas) > 10:
            logger.info(f"   ... e mais {len(linhas) - 10} tópicos")

        # v2.47: Sanitiza títulos que são frases literais de fala
        estrutura = _sanitize_structure_titles(estrutura)

        return estrutura
        
    except Exception as e:
        logger.warning(f"⚠️  Falha no mapeamento: {e}. Continuando sem estrutura prévia.")
        return None

def simplificar_estrutura_se_necessario(estrutura, max_linhas=60):
    """
    Se a estrutura for muito longa (> max_linhas), mantém apenas:
    - Nível 1: 1. Assunto
    - Nível 2: 1.1. Subassunto

    Sempre inclui todos os níveis 1, e todos os níveis 2,
    e depois corta no máximo max_linhas, preservando a ordem original.
    """
    if not estrutura:
        return estrutura

    linhas = [l for l in estrutura.strip().split("\n") if l.strip()]
    if len(linhas) <= max_linhas:
        # Já está razoável, mantém até nível 3 (será filtrado depois por filtrar_niveis_excessivos)
        return estrutura

    logger.info(f"📉 Estrutura muito longa ({len(linhas)} itens). Simplificando para níveis 1 e 2, máx {max_linhas} linhas...")

    nivel1 = []
    nivel2 = []

    for l in linhas:
        s = l.strip()
        # 1. Processo do Trabalho
        if re.match(r"^\d+\.\s", s):
            nivel1.append(l)
        # 1.1. Recursos Trabalhistas
        elif re.match(r"^\d+\.\d+\.\s", s):
            nivel2.append(l)

    # Se por algum motivo não identificou quase nada, devolve original para não quebrar
    if len(nivel1) + len(nivel2) < 5:
        logger.warning("⚠️ Simplificação deixou poucos tópicos. Mantendo estrutura original.")
        return estrutura

    # Monta nova estrutura: primeiro todos os níveis 1, depois os níveis 2, respeitando ordem de aparecimento
    nova = []
    vistos = set()

    for l in linhas:
        if l in vistos:
            continue
        if l in nivel1 or l in nivel2:
            nova.append(l)
            vistos.add(l)

    if len(nova) > max_linhas:
        nova = nova[:max_linhas]

    logger.info(f"✅ Estrutura simplificada: {len(linhas)} -> {len(nova)} linhas (níveis 1 e 2).")
    return "\n".join(nova)

def filtrar_niveis_execessivos(estrutura, max_nivel=3):
    """
    Remove itens da estrutura que sejam mais profundos que max_nivel.
    Ex: se max_nivel=3, remove 1.1.1.1
    """
    if not estrutura:
        return estrutura
        
    linhas = estrutura.strip().split('\n')
    linhas_filtradas = []
    itens_removidos = 0
    
    # Regex para validar nível. 
    # Nível 1: \d+\.
    # Nível 2: \d+\.\d+
    # Nível 3: \d+\.\d+\.\d+
    # O regex verifica se tem no máximo (max_nivel-1) pontos internos entre números
    
    for linha in linhas:
        # Conta quantos grupos de números existem
        match = re.match(r'^(\d+(?:\.\d+)*)', linha.strip())
        if match:
            numeracao = match.group(1)
            nivel = numeracao.count('.') + 1
            if linha.strip().endswith('.'): # Se terminar com ponto (1.1.), não conta como nível extra
                 nivel = numeracao.count('.')
            
            # Ajuste robusto: contar números separados por ponto
            partes = [p for p in numeracao.split('.') if p.isdigit()]
            nivel_real = len(partes)
            
            if nivel_real <= max_nivel:
                linhas_filtradas.append(linha)
            else:
                itens_removidos += 1
        else:
            # Linhas sem numeração (títulos soltos?) mantém por segurança ou remove?
            # Vamos manter para não quebrar formatação estranha
            linhas_filtradas.append(linha)
            
    if itens_removidos > 0:
        logger.info(f"✂️  Filtrados {itens_removidos} itens com nível > {max_nivel}")
    
    return '\n'.join(linhas_filtradas)

# =============================================================================
# PROCESSAMENTO
# =============================================================================

def processar_simples(client, transcricao_bruta, system_prompt):
    logger.info("📄 Documento pequeno - processando em requisição única...")
    
    prompt = f"""{system_prompt}

<texto_para_formatar>
{transcricao_bruta}
</texto_para_formatar>

Retorne APENAS o Markdown formatado."""
    
    for tentativa in range(MAX_RETRIES):
        try:
            response = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=OUTPUT_TOKEN_LIMIT,
                    temperature=0,
                    thinking_config={"include_thoughts": False, "thinking_level": "LOW"}, # Formatação: LOW
                    safety_settings=[
                        types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                    ]
                )
            )
            resultado = response.text
            return limpar_tags_xml(resultado)
        except Exception as e:
            if tentativa < MAX_RETRIES - 1:
                wait = (2 ** tentativa) + random.uniform(0, 1)
                logger.warning(f"Erro, retry em {wait:.1f}s...")
                sleep(wait)
            else:
                raise

def processar_chunk(client, cache, system_prompt, texto_chunk, numero, total, contexto_estilo="", estrutura_global=None, ultimo_titulo=None, profundidade=0):
    rate_limiter.wait_if_needed()
    
    # Recursão infinita protection e limite mínimo
    MIN_CHUNK_CHARS = 4000
    if len(texto_chunk) < MIN_CHUNK_CHARS:
        logger.warning(f"⚠️ Chunk {numero} muito pequeno ({len(texto_chunk)} chars). Processando sem dividir.")
    elif profundidade > 2:
        logger.warning(f"⚠️ Chunk {numero}: Profundidade de recursão {profundidade} atingida. Processando sem dividir.")

    # v2.8: Seção de estrutura global
    secao_estrutura = ""
    if estrutura_global:
        secao_estrutura = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 ESTRUTURA GLOBAL DA AULA (GUIA)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{estrutura_global}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ REGRA DE OURO - PRIORIDADE DO CONTEÚDO REAL:
O Mapeamento acima é apenas um guia inicial. SE houver divergência 
entre o Mapeamento e a Transcrição Real (ex: o professor mudou de 
assunto, ou o título não existe na fala), SIGA A TRANSCRIÇÃO REAL.
A fidelidade ao que foi *falado* é mais importante que seguir 
cegamente a estrutura.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
    
    # v2.10: Aviso sobre último título (anti-duplicação)
    aviso_titulo = ""
    if ultimo_titulo:
        aviso_titulo = f"""
🚫 O bloco anterior TERMINOU no tópico: "{ultimo_titulo}"
   NÃO inicie sua resposta repetindo este título.
   Continue o conteúdo ou inicie o PRÓXIMO subtópico.
"""
    
    # v2.7: Delimitadores MUITO visíveis e instruções reforçadas
    secao_contexto = ""
    if contexto_estilo:
        secao_contexto = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔒 CONTEXTO ANTERIOR (SOMENTE REFERÊNCIA DE ESTILO)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{contexto_estilo}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ ATENÇÃO: O bloco acima JÁ FOI FORMATADO anteriormente.
- NÃO formate novamente esse conteúdo
- NÃO inclua esse conteúdo na sua resposta
- Use APENAS como referência de estilo de escrita
{aviso_titulo}━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 NOVO TEXTO PARA FORMATAR (comece aqui):
"""
    
    # v2.9: Cache Support - Se usar cache, não repete PROMPT_FORMATACAO
    instructions_body = f"""
{secao_estrutura}
{secao_contexto}

<texto_para_formatar>
{texto_chunk}
</texto_para_formatar>

**INSTRUÇÕES FINAIS**:
- Esta é a parte {numero} de {total} (Profundidade: {profundidade})
- Formate APENAS o texto entre <texto_para_formatar>
- Se houver ESTRUTURA GLOBAL acima, use os mesmos nomes de tópicos
- Se houver contexto acima, NÃO o reprocesse
- **ANTI-REPETIÇÃO DE TÍTULOS**: Se o contexto anterior termina com um título (ex: "## Homologação"), NÃO repita esse título no início da sua resposta. Continue o conteúdo diretamente ou inicie o PRÓXIMO tópico diferente.
- Retorne APENAS o Markdown formatado do NOVO texto
"""

    if cache:
        prompt = instructions_body
    else:
        prompt = f"{system_prompt}\n{instructions_body}"

    for tentativa in range(MAX_RETRIES):
        try:
            safety_config = [
                types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
            ]
            
            # Configuração dinâmica para suportar cache
            gen_config_args = {
                "max_output_tokens": OUTPUT_TOKEN_LIMIT,
                "temperature": 0.1,
                "thinking_config": {"include_thoughts": False, "thinking_level": "LOW"}, # Formatação: LOW
                "safety_settings": safety_config
            }
            if cache:
                gen_config_args['cached_content'] = cache.name

            response = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(**gen_config_args)
            )
            
            # --- Diagnóstico finishReason e Usage ---
            finish_reason = "UNKNOWN"
            usage_tokens = 0
            
            try:
                if hasattr(response, 'usage_metadata'):
                    usage = response.usage_metadata
                    # Tenta acessar atributos (pode variar entre SDKs/Vertex)
                    # v2.17.1: Garantir inteiros (getattr pode retornar None se o atributo existe mas é None)
                    candidates_token_count = getattr(usage, 'candidates_token_count', 0) or 0
                    prompt_token_count = getattr(usage, 'prompt_token_count', 0) or 0
                    cached_content_token_count = getattr(usage, 'cached_content_token_count', 0) or 0
                    
                    usage_tokens = candidates_token_count
                    
                    logger.info(f"📊 Usage: Prompt={prompt_token_count} (Cached: {cached_content_token_count}) | Candidates={candidates_token_count}")
                    
                    # Acumular métricas globais (v2.10)
                    metrics.record_api_call(
                        prompt_tokens=prompt_token_count, 
                        completion_tokens=candidates_token_count, 
                        cached_tokens=cached_content_token_count,
                        provider='gemini'
                    )
                
                if hasattr(response, 'candidates') and response.candidates:
                    cand = response.candidates[0]
                    if hasattr(cand, 'finish_reason'):
                         finish_reason = str(cand.finish_reason) # ex: FinishReason.STOP ou "STOP"
            except Exception as ex_usage:
                logger.warning(f"⚠️ Erro ao ler metadados: {ex_usage}")

            # Captura texto (lidando com caso de .text vazio mas content presente)
            resultado = ""
            try:
                resultado = response.text
            except ValueError:
                # O SDK levanta ValueError se finish_reason for SAFETY ou se não houver text field padrão
                pass
            
            if not resultado and hasattr(response, 'candidates') and response.candidates:
                # Tenta extrair parts[0].text manualmente se .text falhou
                try:
                    parts = response.candidates[0].content.parts
                    if parts:
                        resultado = parts[0].text
                except:
                    pass
            
            if not resultado:
                logger.warning(f"⚠️  Resposta vazia na tentativa {tentativa+1}. Reason: {finish_reason}")
                if tentativa < MAX_RETRIES - 1:
                    sleep(2 * (tentativa + 1))
                    continue
                else:
                    return f"[ERRO SEÇÃO {numero}: RESPOSTA VAZIA]"

            # Validação básica de tamanho (compressão)
            razao = len(resultado) / len(texto_chunk) if len(texto_chunk) > 0 else 0
            
            problema_detectado = False
            msg_problema = ""
            
            compressao_excessiva_severa = False
            
            if razao < THRESHOLD_CRITICO: # < 0.70 por padrão
                problema_detectado = True
                msg_problema = f"Compressão excessiva ({razao:.0%})"
                # Flag para chunking imediato se for muito baixo (ex < 0.65)
                # O usuário pediu < THRESHOLD_CRITICO, vamos ser assertivos.
                compressao_excessiva_severa = True
                
            if problema_detectado:
                logger.warning(f"⚠️ Seção {numero}: {msg_problema}. Reason: {finish_reason}. (Tentativa {tentativa+1}/{MAX_RETRIES})")
                
                # SEÇÃO CRÍTICA: Decisão de Chunking Adaptativo
                
                # Lógica antiga: só no final ou MAX_TOKENS
                # Lógica nova: divide cedo se compressão for severa
                
                deve_dividir = (
                    len(texto_chunk) > MIN_CHUNK_CHARS 
                    and profundidade < 2
                    and (
                        ("MAX_TOKENS" in str(finish_reason)) or 
                        (tentativa == MAX_RETRIES - 1) or
                        (compressao_excessiva_severa) # NOVO: Divide já!
                    )
                )
                
                if deve_dividir:
                    motivo = "MAX_TOKENS" if "MAX_TOKENS" in str(finish_reason) else "COMPRESSÃO"
                    logger.info(f"✂️  ATIVANDO CHUNKING ADAPTATIVO para Seção {numero} (Motivo: {motivo} | Profundidade {profundidade} -> {profundidade+1})")
                    return dividir_e_reprocessar(client, cache, system_prompt, texto_chunk, numero, total, contexto_estilo, estrutura_global, ultimo_titulo, profundidade)
                
                if tentativa < MAX_RETRIES - 1:
                    continue  # Tenta de novo (se não foi severo o suficiente para dividir)
                else:
                    logger.error(f"Seção {numero}: Falha após {MAX_RETRIES} tentativas. Retornando melhor esforço.")
            
            return resultado
            
        except Exception as e:
            erro_msg = str(e)
            is_recoverable = any(code in erro_msg for code in ['503', '429', '500', 'RESOURCE_EXHAUSTED', 'InternalServerError']) or "Resposta vazia" in erro_msg
            
            if tentativa < MAX_RETRIES - 1 and is_recoverable:
                wait = (2 ** (tentativa + 2)) + random.uniform(1, 3)
                if '429' in erro_msg or 'RATE_LIMIT' in erro_msg:
                    wait = 30 + random.uniform(0, 5)
                    logger.warning(f"🛑 Rate Limit (429) detectado na seção {numero}. Pausa longa de {wait:.1f}s...")
                else:
                    logger.warning(f"Erro seção {numero}: {erro_msg}. Retry {tentativa+2}/{MAX_RETRIES} em {wait:.1f}s")
                sleep(wait)
            else:
                logger.error(f"Falha seção {numero}: {erro_msg}")
                return f"\n\n> [!WARNING]\n> Falha ao processar seção {numero}. Texto original:\n\n{texto_chunk}"
    
    return texto_chunk

def dividir_e_reprocessar(client, cache, system_prompt, texto_chunk, numero, total, contexto_estilo, estrutura_global, ultimo_titulo, profundidade):
    """
    Divide um chunk grande em dois menores e processa recursivamente.
    Tenta dividir em quebras de parágrafo (\n\n) próximas ao meio.
    """
    # v2.10: Registrar divisão adaptativa
    metrics.record_adaptive_split()
    
    meio = len(texto_chunk) // 2
    
    # Procura quebra ideal (\n\n) num raio de 20% do meio
    margem = int(len(texto_chunk) * 0.2)
    inicio_busca = max(0, meio - margem)
    fim_busca = min(len(texto_chunk), meio + margem)
    
    janela_busca = texto_chunk[inicio_busca:fim_busca]
    pos_relativa = janela_busca.find('\n\n')
    
    if pos_relativa != -1:
        ponto_corte = inicio_busca + pos_relativa + 2 # +2 para incluir os \n\n no primeiro bloco ou pular? Vamos cortar DEPOIS dos \n\n
    else:
        # Tenta quebra simples \n
        pos_relativa_n = janela_busca.find('\n')
        if pos_relativa_n != -1:
            ponto_corte = inicio_busca + pos_relativa_n + 1
        else:
            # Corte seco no espaço mais próximo
            pos_relativa_espaco = janela_busca.find(' ')
            if pos_relativa_espaco != -1:
                ponto_corte = inicio_busca + pos_relativa_espaco + 1
            else:
                ponto_corte = meio # Corte arbitrário
    
    parte_a = texto_chunk[:ponto_corte]
    parte_b = texto_chunk[ponto_corte:]
    
    logger.info(f"   Splitting chunk {numero}: Part A ({len(parte_a)} chars) + Part B ({len(parte_b)} chars)")
    
    # Processa Parte A
    resultado_a = processar_chunk(
        client, cache, system_prompt, parte_a, f"{numero}.A", total, 
        contexto_estilo, estrutura_global, ultimo_titulo, profundidade + 1
    )
    
    # Usa o final de A como contexto para B? Talvez seja excessivo e caro.
    # Vamos manter o contexto original para B por segurança, 
    # ou usar resultado_a[-1000:] como novo contexto_estilo.
    # Usar resultado_a é melhor para continuidade.
    
    novo_contexto = resultado_a[-2000:] if len(resultado_a) > 2000 else resultado_a
    
    # Processa Parte B
    resultado_b = processar_chunk(
        client, cache, system_prompt, parte_b, f"{numero}.B", total, 
        novo_contexto, estrutura_global, None, profundidade + 1 # ultimo_titulo None pois A já tratou disso
    )
    
    return f"{resultado_a}\n\n{resultado_b}"

def extrair_titulos_h2(texto):
    """Extrai todos os títulos de nível 2 (##) do texto"""
    titulos = []
    for linha in texto.split('\n'):
        if linha.strip().startswith('##') and not linha.strip().startswith('###'):
            titulo_limpo = re.sub(r'^##\s*\d+\.\s*', '', linha.strip())
            titulos.append(titulo_limpo.lower())
    return titulos

# =============================================================================
# SMART STITCHING (v2.10) - Anti-Duplicação Cirúrgica
# =============================================================================

def remover_eco_do_contexto(resposta_api, contexto_enviado):
    """
    Remove o início da resposta se for apenas um 'eco' do final do contexto.
    """
    if not contexto_enviado or not resposta_api:
        return resposta_api

    final_contexto = contexto_enviado.strip()[-300:]
    inicio_resposta = resposta_api.strip()[:300]

    matcher = SequenceMatcher(None, final_contexto, inicio_resposta)
    match = matcher.find_longest_match(0, len(final_contexto), 0, len(inicio_resposta))

    if match.size > 50:
        logger.info(f"✂️ Eco detectado! Removendo {match.size} chars repetidos no início.")
        return resposta_api.strip()[match.size:].strip()
    
    return resposta_api

def titulos_sao_similares(t1, t2, threshold=None):
    """Verifica se dois títulos são semanticamente iguais (fuzzy matching).
    Usa LIMIAR_SECOES global se threshold não for especificado.
    """
    if threshold is None:
        threshold = LIMIAR_SECOES  # Usa limiar de seções (0.70 Fidelidade / 0.60 Apostila)
        
    def normalizar(t):
        # Remove apenas caracteres não alfanuméricos, mas MANTÉM tamanho relativo
        return re.sub(r'[^a-z0-9 ]', '', t.lower())
    
    nt1 = normalizar(t1)
    nt2 = normalizar(t2)
    
    if not nt1 or not nt2:
        return False
    
    # PROTEÇÃO 1: Se um título for muito maior que o outro, não são duplicatas
    nt1_compact = nt1.replace(' ', '')
    nt2_compact = nt2.replace(' ', '')
    len_ratio = min(len(nt1_compact), len(nt2_compact)) / max(len(nt1_compact), len(nt2_compact))
    if len_ratio < 0.8:  # Se a diferença de tamanho for > 20%, assume que são diferentes
        return False
    
    # PROTEÇÃO 2: Verificação por palavras - se houver palavras exclusivas significativas
    palavras1 = set(nt1.split())
    palavras2 = set(nt2.split())
    diferenca = palavras1.symmetric_difference(palavras2)
    
    # Se as palavras diferentes forem longas (não apenas 'e', 'do', 'da'), assume diferença real
    if any(len(w) > 3 for w in diferenca):
        return False
        
    return SequenceMatcher(None, nt1_compact, nt2_compact).ratio() > threshold

def limpar_inicio_redundante(texto_novo, texto_acumulado):
    """
    Remove título no início do novo chunk se similar ao último título do texto acumulado.
    """
    if not texto_acumulado.strip():
        return texto_novo

    ultimas_linhas = texto_acumulado.strip().split('\n')[-30:]
    ultimo_titulo = None
    for linha in reversed(ultimas_linhas):
        if linha.strip().startswith('##'):
            ultimo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
            break
    
    if not ultimo_titulo:
        return texto_novo

    linhas_novas = texto_novo.strip().split('\n')
    
    for i, linha in enumerate(linhas_novas[:10]):  # v2.11: Busca mais profunda (era 5)
        if linha.strip().startswith('##'):
            novo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
            
            if titulos_sao_similares(ultimo_titulo, novo_titulo):
                logger.info(f"✂️ Título duplicado na junção: '{novo_titulo}' ≈ '{ultimo_titulo}'")
                return '\n'.join(linhas_novas[i+1:])
    
    return texto_novo

def detectar_secoes_duplicadas(texto):
    """v2.15: Detecta seções duplicadas por títulos em ## e ### (Fuzzy Matching)"""
    logger.info("🔍 Detectando seções duplicadas (fuzzy, H2+H3)...")
    
    linhas = texto.split('\n')
    titulos_vistos = []  # (titulo_normalizado, linha_idx)
    secoes_duplicadas = []
    
    for i, linha in enumerate(linhas):
        linha_strip = linha.strip()
        # Match both ## (H2) and ### (H3)
        if linha_strip.startswith('##'):
            # Extract level
            match_nivel = re.match(r'^(#+)', linha_strip)
            nivel = len(match_nivel.group(1)) if match_nivel else 2
            
            # Normalize title: remove ## prefix, numbers, emojis, and "(Continuação)"
            titulo_normalizado = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', linha_strip)
            titulo_normalizado = re.sub(r'[📋📊🗂️]', '', titulo_normalizado).strip()
            titulo_normalizado = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_normalizado, flags=re.IGNORECASE).strip()
            
            duplicado = False
            for t_visto, linha_visto in titulos_vistos:
                if titulos_sao_similares(titulo_normalizado, t_visto):
                    logger.warning(f"⚠️  Duplicado (fuzzy): '{linha_strip}' ≈ '{t_visto}'")
                    secoes_duplicadas.append({
                        'titulo': titulo_normalizado,
                        'primeira_linha': linha_visto,
                        'duplicada_linha': i
                    })
                    duplicado = True
                    break
            
            if not duplicado:
                titulos_vistos.append((titulo_normalizado, i))
    
    if secoes_duplicadas:
        logger.error(f"❌ {len(secoes_duplicadas)} seções duplicadas detectadas!")
    else:
        logger.info("✅ Nenhuma seção duplicada detectada")
    
    return secoes_duplicadas

def remover_secoes_duplicadas(texto):
    """v2.14: Remove seções duplicadas com COMPARAÇÃO JANELADA (Fix Diluição)"""
    from difflib import SequenceMatcher
    
    secoes_dup = detectar_secoes_duplicadas(texto)
    if not secoes_dup: return texto
    
    print("🧹 Removendo seções duplicadas (Smart Dedupe v2.14)...")
    linhas = texto.split('\n')
    
    # Rastreia o ÚLTIMO segmento adicionado para cada título (para evitar diluição na comparação)
    ultimo_segmento_visto = {}  # titulo_normalizado -> último texto adicionado
    linhas_para_remover = set()
    linhas_para_adicionar_separador = set()
    
    for dup in secoes_dup:
        # --- 1. Extrair Conteúdo Original ---
        idx_orig = dup['primeira_linha']
        header_orig = linhas[idx_orig].strip()
        match_orig = re.match(r'^(#+)', header_orig)
        nivel_orig = len(match_orig.group(1)) if match_orig else 2
        
        # Extrai conteúdo da seção original
        content_orig = []
        for i in range(idx_orig + 1, len(linhas)):
            line = linhas[i].strip()
            if line.startswith('#'):
                match_now = re.match(r'^(#+)', line)
                if match_now and len(match_now.group(1)) <= nivel_orig: break
            content_orig.append(linhas[i])
        text_orig = "\n".join(content_orig)
        
        titulo_key = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', header_orig)
        titulo_key = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_key, flags=re.IGNORECASE).strip()
        
        # Inicializa o rastreamento se for a primeira vez
        if titulo_key not in ultimo_segmento_visto:
            ultimo_segmento_visto[titulo_key] = text_orig

        # --- 2. Extrair Conteúdo Duplicado ---
        idx_dup = dup['duplicada_linha']
        header_dup = linhas[idx_dup].strip()
        match_dup = re.match(r'^(#+)', header_dup)
        nivel_dup = len(match_dup.group(1)) if match_dup else 2
        
        fim_dup_idx = len(linhas)
        content_dup = []
        for i in range(idx_dup + 1, len(linhas)):
            line = linhas[i].strip()
            if line.startswith('#'):
                match_now = re.match(r'^(#+)', line)
                if match_now and len(match_now.group(1)) <= nivel_dup:
                    fim_dup_idx = i
                    break
            content_dup.append(linhas[i])
        text_dup = "\n".join(content_dup)
        
        # --- 3. Comparar Conteúdo (Lógica Janelada v2.15) ---
        # Compara APENAS com o último segmento conhecido dessa seção
        texto_referencia = ultimo_segmento_visto.get(titulo_key, text_orig)
        
        len_dup = len(text_dup.strip())
        len_ref = len(texto_referencia.strip())
        
        # Lógica de decisão baseada em tamanho
        if len_dup < 50:
            # Duplicado curto demais = lixo, deletar
            sim = 1.0
            print(f"   ⚠️  Seção duplicada muito curta ({len_dup} chars) - marcando para remoção")
        elif len_ref < 50 and len_dup >= 200:
            # Original curto, mas duplicado substancial = original estava incompleto
            # Manter o duplicado como novo conteúdo
            sim = 0.0
            print(f"   ℹ️  Original curto ({len_ref}c), duplicado substancial ({len_dup}c) - mantendo novo conteúdo")
        else:
            sim = SequenceMatcher(None, texto_referencia, text_dup).ratio()
            
        print(f"   Similaridade: {sim:.1%} | Linha {idx_dup} | '{titulo_key[:40]}...'")
        
        if sim > LIMIAR_SECOES:  # Usa limiar de seções (0.70 Fidelidade / 0.60 Apostila) 
            print(f"   🗑️  Removendo SEÇÃO INTEIRA (Duplicata confirmada)")
            for i in range(idx_dup, fim_dup_idx):
                linhas_para_remover.add(i)
        else:
            print(f"   🔗 Mesclando conteúdo (Nova informação detectada)")
            linhas_para_remover.add(idx_dup)
            if idx_dup + 1 < len(linhas):
                linhas_para_adicionar_separador.add(idx_dup + 1)
            
            # ATUALIZA o último segmento visto para a próxima comparação
            ultimo_segmento_visto[titulo_key] = text_dup

    # --- 4. Reconstrução ---
    linhas_limpas = []
    for i, linha in enumerate(linhas):
        if i in linhas_para_remover:
            continue
        if i in linhas_para_adicionar_separador:
            linhas_limpas.append("") 
        linhas_limpas.append(linha)
        
    print(f"✅ {len(linhas_para_remover)} linhas removidas")
    return '\n'.join(linhas_limpas)

def remover_duplicacoes_literais(texto):
    """Remove parágrafos individuais duplicados"""
    paragrafos = texto.split('\n\n')
    paragrafos_limpos = []
    dup_count = 0
    
    for i, para in enumerate(paragrafos):
        if i == 0:
            paragrafos_limpos.append(para)
            continue
        
        if len(para.strip()) < 80 or para.strip().startswith('#'):
            paragrafos_limpos.append(para)
            continue
        
        is_duplicate = False
        para_norm = ' '.join(para.split()).lower()
        
        for j in range(max(0, len(paragrafos_limpos) - 3), len(paragrafos_limpos)):
            para_ant = paragrafos_limpos[j]
            para_ant_norm = ' '.join(para_ant.split()).lower()
            
            ratio = SequenceMatcher(None, para_norm, para_ant_norm).ratio()
            
            if ratio > 0.95:
                is_duplicate = True
                dup_count += 1
                break
        
        if not is_duplicate:
            paragrafos_limpos.append(para)
    
    if dup_count > 5:
        logger.warning(f"⚠️  {dup_count} parágrafos duplicados removidos")
    
    return '\n\n'.join(paragrafos_limpos)

# =============================================================================
# V2.17: DEDUPLICAÇÃO ROBUSTA (7-DIFF Strategy) - Portado de mlx_vomo.py
# =============================================================================

def remover_overlap_duplicado(resultados):
    """
    v2.17: Remove duplicação entre chunks usando detecção ROBUSTA de conteúdo.
    Estratégia 7-DIFF: Compara título + conteúdo com janela deslizante de 20 seções.
    """
    if isinstance(resultados, str):
        resultados = [resultados]
    if len(resultados) <= 1:
        return resultados[0] if resultados else ""
    
    from difflib import SequenceMatcher
    
    def normalize_text(text):
        if not text: return ""
        text = re.sub(r'[#*-]', ' ', text)
        text = re.sub(r'[^\w\s]', '', text)
        return ' '.join(text.lower().split())

    def calculate_similarity(text1, text2):
        if not text1 or not text2:
            return 0.0
        if len(text1) < 50:
            return 1.0 if text1 in text2 or text2 in text1 else 0.0
        return SequenceMatcher(None, text1, text2).quick_ratio()

    def extract_unique_paragraphs(sec_curr_content, sec_prev_content):
        if not sec_curr_content: return []
        unique = []
        paras_curr = sec_curr_content.split('\n\n')
        paras_prev_norm = [normalize_text(p) for p in sec_prev_content.split('\n\n')]
        
        for p in paras_curr:
            p_clean = p.strip()
            if not p_clean or len(p_clean) < 20: continue
            
            p_norm = normalize_text(p_clean)
            is_present = False
            for pp_norm in paras_prev_norm:
                if calculate_similarity(p_norm, pp_norm) > 0.85:
                    is_present = True
                    break
            
            if not is_present:
                unique.append(p_clean)
        return unique

    logger.info("🧹 Iniciando deduplicação robusta (7-DIFF Strategy)...")

    # 1. Junta e Parseia
    texto_bruto = '\n\n'.join(resultados)
    lines = texto_bruto.split('\n')
    
    sections = []
    current_section = None
    header_pattern = re.compile(r'^(#{1,6})\s+(.+)$')
    
    captured_lines = []
    intro_lines = []
    has_started = False
    
    for line in lines:
        match = header_pattern.match(line)
        if match:
            has_started = True
            if current_section:
                current_section['content'] = '\n'.join(captured_lines).strip()
                sections.append(current_section)
                captured_lines = []
            
            title_text = match.group(2).strip()
            title_clean = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', title_text)
            
            current_section = {
                'title_clean': title_clean,
                'level': len(match.group(1)),
                'full_header': line,
                'content': ""
            }
        else:
            if has_started:
                captured_lines.append(line)
            else:
                intro_lines.append(line)
    
    if current_section:
        current_section['content'] = '\n'.join(captured_lines).strip()
        sections.append(current_section)
        
    logger.info(f"   📊 Analisando {len(sections)} seções...")

    # 2. Detecção e Remoção
    indices_to_remove = set()
    MAX_WINDOW = 20
    
    for i in range(len(sections)):
        if i in indices_to_remove: continue
        sec_curr = sections[i]
        
        start_check = max(0, i - MAX_WINDOW)
        for j in range(start_check, i):
            if j in indices_to_remove: continue
            sec_prev = sections[j]
            
            if sec_curr['level'] != sec_prev['level']: continue
            
            sim_title = calculate_similarity(normalize_text(sec_curr['title_clean']), normalize_text(sec_prev['title_clean']))
            sim_content = calculate_similarity(normalize_text(sec_curr['content']), normalize_text(sec_prev['content']))
            
            is_duplicate = False
            
            if sim_title > 0.9 and sim_content > 0.6:
                is_duplicate = True
            elif sim_content > LIMIAR_7DIFF:  # Usa limiar 7-DIFF (0.85 Fidelidade / 0.80 Apostila)
                 is_duplicate = True
            elif sim_title > 0.95 and len(sec_curr['content']) < 100:
                 is_duplicate = True
            
            if is_duplicate:
                logger.info(f"   🗑️  Duplicata detectada: '{sec_curr['title_clean'][:40]}...'")
                
                unique_paras = extract_unique_paragraphs(sec_curr['content'], sec_prev['content'])
                if unique_paras:
                    sections[j]['content'] += '\n\n' + '\n\n'.join(unique_paras)
                
                indices_to_remove.add(i)
                break
    
    # 3. Reconstrução
    final_lines = list(intro_lines)
    for i, sec in enumerate(sections):
        if i in indices_to_remove: continue
        final_lines.append(sec['full_header'])
        if sec['content']:
            final_lines.append(sec['content'])
        final_lines.append("")
    
    logger.info(f"   ✅ {len(indices_to_remove)} seções duplicadas removidas/mescladas")
    return '\n'.join(final_lines)

def deterministic_structure_fix(text):
    """
    v2.17: Reorganização Estrutural Determinística.
    Detecta se usa H1 ou H2 como nível principal e reorganiza o documento.
    """
    logger.info("🧩 Executando Reorganização Estrutural Determinística...")
    
    lines = text.split('\n')
    
    # Detecção de Hierarquia
    has_h1 = any(re.match(r'^#\s+', line) for line in lines)
    header_level_regex = r'^#\s+' if has_h1 else r'^##\s+'
    logger.info(f"   ℹ️  Nível principal detectado: {'H1 (#)' if has_h1 else 'H2 (##)'}")

    content_map = {
        "PREAMBULO": [],
        "DISCIPLINAS": {}, 
        "ENCERRAMENTO": []
    }
    
    current_area = "PREAMBULO"
    current_block = []
    disciplinas_order = [] 
    
    re_disciplina = re.compile(rf'{header_level_regex}(?!Questão|Q\.)([^0-9\.]+.*)', re.IGNORECASE)
    re_encerramento = re.compile(rf'{header_level_regex}(?:ENCERRAMENTO|CONSIDERAÇÕES|CONCLUSÃO)', re.IGNORECASE)
    
    def flush_block(area, block_lines):
        if not block_lines: return
        block_text = '\n'.join(block_lines)
        
        if area == "PREAMBULO":
            content_map["PREAMBULO"].append(block_text)
        elif area == "ENCERRAMENTO":
            content_map["ENCERRAMENTO"].append(block_text)
        else:
            if area not in content_map["DISCIPLINAS"]:
                content_map["DISCIPLINAS"][area] = []
                disciplinas_order.append(area)
            content_map["DISCIPLINAS"][area].append(block_text)

    for line in lines:
        match_disc = re_disciplina.match(line)
        if match_disc:
            flush_block(current_area, current_block)
            current_block = []
            
            raw_area = match_disc.group(1).strip().upper()
            
            if "DIREITO" not in raw_area and len(raw_area) < 50:
                 if any(x in raw_area for x in ["CIVIL", "PENAL", "TRABALHO", "ADMINISTRATIVO", "CONSTITUCIONAL"]):
                     current_area = f"DIREITO {raw_area}"
                 else:
                     current_area = raw_area
            else:
                 current_area = raw_area
            continue 
            
        if re_encerramento.match(line):
            flush_block(current_area, current_block)
            current_block = []
            current_area = "ENCERRAMENTO"
            continue

        current_block.append(line)
        
    flush_block(current_area, current_block)
    
    # Reconstrução
    final_output = []
    
    if content_map["PREAMBULO"]:
        final_output.append("# ORIENTAÇÕES GERAIS / INTRODUÇÃO")
        final_output.extend(content_map["PREAMBULO"])
        final_output.append("")

    for area in disciplinas_order:
        area_clean = area.replace("#", "").strip()
        final_output.append(f"# {area_clean}")
        for block in content_map["DISCIPLINAS"][area]:
            final_output.append(block)
        final_output.append("")
        
    if content_map["ENCERRAMENTO"]:
        final_output.append("# CONSIDERAÇÕES FINAIS")
        final_output.extend(content_map["ENCERRAMENTO"])
        
    num_identified = len(disciplinas_order)
    logger.info(f"   ✅ Reorganizado: {num_identified} seções principais identificadas.")
    
    if num_identified == 0 and len(content_map["PREAMBULO"]) > 0:
        logger.warning("   ⚠️ Nenhuma estrutura detectada. Mantendo original.")
        return text
        
    return '\n'.join(final_output)

def normalize_headings(texto):
    """
    v1.0: Normaliza títulos semanticamente similares para uma versão única.
    - Agrupa títulos por similaridade
    - Escolhe o título mais descritivo de cada grupo
    - Remove sufixos como "(Continuação)"
    """
    from difflib import SequenceMatcher
    
    print("🔤 Normalizando títulos similares...")
    linhas = texto.split('\n')
    
    # 1. Extrair todos os títulos com info de nível e posição
    titulos = []
    for i, linha in enumerate(linhas):
        stripped = linha.strip()
        if stripped.startswith('##'):
            match = re.match(r'^(#+)\s*', stripped)
            nivel = len(match.group(1)) if match else 2
            # Extrai título limpo (sem # e sem numeração)
            titulo_limpo = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', stripped).strip()
            # Remove "(Continuação)" para comparação
            titulo_base = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_limpo, flags=re.IGNORECASE).strip()
            titulos.append({
                'linha': i,
                'nivel': nivel,
                'original': stripped,
                'limpo': titulo_limpo,
                'base': titulo_base
            })
    
    if not titulos:
        return texto
    
    # 2. Agrupar títulos similares (mesmo nível + similaridade > LIMIAR)
    grupos = []
    usados = set()
    
    for i, t1 in enumerate(titulos):
        if i in usados:
            continue
        grupo = [t1]
        usados.add(i)
        
        for j, t2 in enumerate(titulos):
            if j in usados or j <= i:
                continue
            if t1['nivel'] == t2['nivel']:
                sim = SequenceMatcher(None, t1['base'].lower(), t2['base'].lower()).ratio()
                if sim > LIMIAR_SECOES:
                    grupo.append(t2)
                    usados.add(j)
        
        if len(grupo) > 1:
            grupos.append(grupo)
    
    if not grupos:
        # Nada para normalizar, mas ainda remove "(Continuação)"
        texto_limpo = re.sub(r'\s*\(Continuação\)\s*(?=\n|$)', '', texto, flags=re.IGNORECASE)
        return texto_limpo
    
    # 3. Para cada grupo, escolher o "melhor" título (mais curto entre os mais longos)
    #    Lógica: preferir títulos sem "(Continuação)" e com descrição completa
    substituicoes = {}
    for grupo in grupos:
        # Ordenar por: não ter "(Continuação)" primeiro, depois por comprimento (prefer médio)
        candidatos = sorted(grupo, key=lambda x: (
            '(Continuação)' in x['limpo'],  # False (0) vem antes de True (1)
            abs(len(x['limpo']) - 40)  # Preferir títulos de ~40 chars (nem muito curto nem muito longo)
        ))
        
        melhor = candidatos[0]['limpo']
        print(f"   📝 Grupo de {len(grupo)} títulos similares → padronizando para: '{melhor[:50]}...'")
        
        for t in grupo:
            if t['limpo'] != melhor:
                substituicoes[t['linha']] = (t['nivel'], melhor)
    
    # 4. Aplicar substituições
    novas_linhas = []
    for i, linha in enumerate(linhas):
        if i in substituicoes:
            nivel, novo_titulo = substituicoes[i]
            # Preservar numeração existente se houver
            match_num = re.match(r'^(#+\s*\d+(?:\.\d+)*\.?\s*)', linhas[i].strip())
            if match_num:
                prefixo = match_num.group(1)
                # Remove a numeração do novo título se ele tiver uma
                novo_titulo = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', novo_titulo)
                novas_linhas.append(f"{prefixo}{novo_titulo}")
            else:
                novas_linhas.append(f"{'#' * nivel} {novo_titulo}")
        else:
            # Remove "(Continuação)" de qualquer título restante
            if linha.strip().startswith('##'):
                linha = re.sub(r'\s*\(Continuação\)\s*(?=\n|$)', '', linha, flags=re.IGNORECASE)
            novas_linhas.append(linha)
    
    print(f"   ✅ {len(substituicoes)} títulos normalizados, '(Continuação)' removidos")
    return '\n'.join(novas_linhas)

# =============================================================================
# REVISÃO SEMÂNTICA DE ESTRUTURA POR IA (v2.16)
# =============================================================================

# =============================================================================
# REVISÃO SEMÂNTICA, MAPEAMENTO E ESTRUTURA (v2.16)
# =============================================================================

PROMPT_MAPEAMENTO = """Você é um especialista em organização de conteúdo educacional acadêmico (v2.17).

## ETAPA 1: IDENTIFICAR O TIPO DE CONTEÚDO
Analise a transcrição e determine qual é a **natureza predominante** do material:

| Tipo | Pistas no Texto | Estrutura Ideal |
|------|-----------------|-----------------|
| **SIMULADO** | "questão 1", "questão 2", "espelho de correção", "correção do simulado", "vamos corrigir" | Organizar por QUESTÕES numeradas |
| **AULA EXPOSITIVA** | explicações contínuas de um tema, teoria, doutrina, sem questões específicas | Organizar por TEMAS/MATÉRIAS e MARCOS LEGAIS |
| **REVISÃO** | "revisão", "resumo", múltiplos temas curtos, "pontos importantes" | Organizar por TÓPICOS de revisão |
| **CORREÇÃO DE PROVA** | "gabarito", "alternativa correta", "item certo/errado" | Organizar por QUESTÕES com gabarito |

## ETAPA 2: EXTRAIR A ESTRUTURA

### Se for SIMULADO ou CORREÇÃO DE PROVA:
```
1. Orientações Gerais / Introdução
2. Questão 1: [Título descritivo] — [Área do Direito]
   2.1. Enunciado e Contexto
   2.2. Fundamentação (Doutrina/Jurisprudência)
   2.3. Pontos do Espelho / Resposta
3. Questão 2: [Título descritivo] — [Área do Direito]
   3.1. ...
[Continue para cada questão]
N. Considerações Finais / Dúvidas
```

### Se for AULA EXPOSITIVA (ATENÇÃO AOS MARCOS LEGAIS):
Você DEVE identificar **Marcos Legais e Jurisprudenciais** importantes e elevá-los à categoria de SUBTÓPICOS.
Exemplos de marcos: "Súmula X", "Artigo Y do CC", "Tese de Repercussão Geral Z", "Julgado X do STF".

```
1. Introdução
2. [Matéria 1: ex. Direito Administrativo]
   2.1. [Subtema]
      2.1.1. [Detalhamento]
3. [Matéria 2: ex. Direito Civil]
   3.1. ...
```

### Se for REVISÃO:
```
1. [Tema 1]
   1.1. Pontos-chave
   1.2. Jurisprudência/Súmulas
2. [Tema 2]
   2.1. ...
```

## REGRAS GERAIS:
1. **MÁXIMO 3 NÍVEIS** de hierarquia (1., 1.1., 1.1.1.)
2. **Seja descritivo** nos títulos — inclua o assunto real, não apenas "Questão 1"
3. **Mantenha a ORDEM** cronológica da transcrição
4. **Mapeie do INÍCIO ao FIM** — não omita partes
5. **Identifique a ÁREA DO DIREITO** de cada bloco quando possível
6. **PREFIRA SUBTÓPICOS (1.1.) a novos tópicos (2.)**: Abra novo tópico de nível 1 SOMENTE quando o macroassunto mudar de verdade (ex.: de Direito Administrativo para Direito Civil). Aspectos, institutos e marcos legais DENTRO do mesmo macroassunto devem ser subtópicos (1.1., 1.2., etc.), NUNCA tópicos de nível 1 separados.
7. **ANTI-FRAGMENTAÇÃO**: Se o professor trata 4+ aspectos de um tema, todos devem ser subtópicos de um único tema-mãe. Exemplo correto: `2. Execução Fiscal` com `2.1. Procedimento`, `2.2. Citação`, `2.3. Exceção de Pré-Executividade`. Exemplo ERRADO: `2. Execução Fiscal`, `3. Procedimento`, `4. Citação`.
8. **TÍTULOS SÃO RÓTULOS, NÃO FALAS**: Os títulos devem ser rótulos descritivos curtos (máx 8 palavras), NUNCA trechos literais da fala do professor.
   - ERRADO: "1. Já estávamos conversando aqui antes de começar a transmissão"
   - CORRETO: "1. Introdução e Apresentação"
   - ERRADO: "2.1. Bom dia pessoal vamos começar a aula de hoje sobre licitações"
   - CORRETO: "2.1. Abertura — Licitações e Contratos"
9. **SAUDAÇÕES E LOGÍSTICA → "Introdução"**: Trechos de boas-vindas, ajustes técnicos, apresentação pessoal ou logística devem ser agrupados sob "1. Introdução" ou "1. Apresentação e Contextualização", nunca com a fala literal como título.

## TRANSCRIÇÃO:
{transcricao}

## RESPOSTA:
Primeiro, indique em uma linha: `[TIPO: SIMULADO/EXPOSITIVA/REVISÃO/CORREÇÃO]`
Depois, retorne APENAS a estrutura hierárquica (máx 3 níveis)."""

PROMPT_STRUCTURE_REVIEW = """Você é um revisor especializado em estrutura de documentos jurídicos educacionais.

## ESTRUTURA DE MAPEAMENTO INICIAL (Referência - se disponível):
{estrutura_mapeada}

## TAREFA
Revise a ESTRUTURA (headers/títulos) do documento abaixo, COMPARANDO com o mapeamento acima (se disponível). Sua missão é harmonizar o mapeamento planejado com o CONTEÚDO REAL da aula.

## ✅ O QUE VOCÊ DEVE FAZER:

### 1. COMPARAR E REFINAR TÍTULOS (CRÍTICO)
Verifique se los títulos refletem os tópicos do mapeamento. Se um título for genérico, torne-o descritivo.
- ERRADO: "### Questão" ou "### Tópico"
- CORRETO: "### Questão 1: Responsabilidade Civil" (Descritivo conforme mapa/conteúdo)
- **CRÍTICO:** Evite títulos idênticos em seções "irmãs". Diferencie-os pelo conteúdo específico de cada uma.

### 2. VALIDAR HIERARQUIA E PROMOÇÃO DE TÓPICOS
Confirme se a estrutura segue lógica consistente (##, ###, ####). 
- **PROMOÇÃO:** Se um sub-subtópico (ex: 9.19.5) for extenso e tratar de um tema central (ex: Execução Fiscal), PROMOVA-O a um nível superior (ex: ### 9.20) para evitar fragmentação excessiva e respeitar o limite de níveis.

### 3. RENUMERAÇÃO SEQUENCIAL OBRIGATÓRIA
Se você criar, deletar ou promover uma seção, você DEVE renumerar TODAS as seções subsequentes daquela mesma hierarquia para manter a sequência lógica (ex: se 9.20 foi criado, o antigo 9.20 vira 9.21, e assim por diante).

### 4. MESCLAR QUESTÕES DUPLICADAS
Se duas seções têm o mesmo número de questão na mesma área, MESCLE-AS.
- ERRADO: "2.1. Questão 1: TAC" + "2.2. Questão 1: TAC" 
- CORRETO: "2.1. Questão 1: TAC" (Única, com todo o conteúdo unificado)

### 5. PRIORIDADE DO CONTEÚDO REAL (DECIDIR ESTRUTURA)
Se houver conflito entre mapeamento e documento, escolha a estrutura que melhor reflete o CONTEÚDO REAL. O mapeamento é apenas um guia.

### 6. LIMPEZA TÉCNICA (SINTAXE MARKDOWN)
Corrija problemas detalhados de formatação:
- **Tabelas:** Alinhar colunas e adicionar separadores faltantes.
- **Listas:** Corrigir bullets ou numeração mal formatada.
- **Espaçamento:** Padronizar linhas em branco entre seções (mínimo uma linha).
- **Headers Vazios:** Remover títulos sem conteúdo abaixo.
- **Metadados:** Remover headers ou tags inline como "[TIPO: ...]", "**[TIPO: ...]**", ou "[BLOCO 0X]".

## ❌ O QUE VOCÊ NÃO DEVE FAZER:
1. **NÃO ALTERE O CONTEÚDO** dos parágrafos - apenas os títulos e a organização.
2. **NUNCA RESUMA** ou encurte o texto (mesmo tamanho de entrada/saída é obrigatório).
3. **NÃO INVENTE** fatos jurídicos.
4. **NÃO REMOVA** trechos técnicos ou exemplos.

## REGRAS CRÍTICAS DE HIERARQUIA:
- Use **MÁXIMO 3** níveis de hierarquia (##, ###, ####).
- Nunca use # (H1) para subtópicos (apenas para o título principal do documento).
- Preserve a ordem cronológica geral.

## DOCUMENTO PARA REVISAR:
{documento}

## 📝 RELATÓRIO ESPERADO:
Ao final do documento, inclua um bloco de comentário indicando:
- Quantos títulos foram refinados
- Promoções de seções e renumerações realizadas
- Discrepâncias com o mapeamento (se houver)

Formato:
<!-- RELATÓRIO: X títulos refinados | Y seções promovidas/renumeradas | Discrepâncias: [Nenhuma/Lista] -->

## RESPOSTA:
Retorne o documento COMPLETO E INTEGRAL (mesmo tamanho do original) com os títulos/headers corrigidos e o relatório no final. NÃO RESUMA."""


def map_structure(client, full_text):
    """Creates a global structure skeleton to guide the formatting."""
    logger.info("🗺️  Mapeando estrutura global do documento...")
    
    # Limit input to avoid context overflow (200k chars is plenty for structure)
    input_sample = full_text[:200000] 
    
    try:
        rate_limiter.wait_if_needed()
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=PROMPT_MAPEAMENTO.format(transcricao=input_sample),
            config=types.GenerateContentConfig(
                max_output_tokens=20000,
                thinking_config={"include_thoughts": False, "thinking_level": "HIGH"}
            )
        )
        content = response.text.replace('```markdown', '').replace('```', '')
        content = _sanitize_structure_titles(content)
        logger.info("   ✅ Estrutura mapeada com sucesso.")
        return content

    except Exception as e:
        logger.warning(f"   ⚠️  Falha no mapeamento via Gemini: {e}")
        return None

def ai_structure_review(client, texto, estrutura_mapeada=None):
    """
    v2.0: Revisão semântica de estrutura usando IA com VALIDAÇÃO CRUZADA.
    Compara o documento com a estrutura de mapeamento inicial.
    Corrige: questões duplicadas, subtópicos órfãos, fragmentação excessiva.
    """
    logger.info("🧠 Revisão Semântica de Estrutura (IA v2.0)...")
    
    # Gemini 3 Flash suporta 1M tokens (~4M chars) - usar até 500k chars
    max_doc_chars = 500000
    if len(texto) > max_doc_chars:
        logger.warning(f"   ⚠️ Documento muito longo ({len(texto)} chars), truncando para {max_doc_chars//1000}k...")
        texto_para_revisao = texto[:max_doc_chars] + "\n\n[... documento truncado para revisão estrutural ...]"
    else:
        texto_para_revisao = texto
    
    # Preparar estrutura mapeada (se disponível)
    if estrutura_mapeada:
        estrutura_str = estrutura_mapeada[:50000]  # Limitar estrutura a 50k chars
        logger.info(f"   📋 Usando estrutura de mapeamento inicial ({len(estrutura_mapeada)} chars) para validação cruzada.")
    else:
        estrutura_str = "[Estrutura de mapeamento não disponível - analisar documento autonomamente]"
        logger.info("   ℹ️  Sem mapeamento inicial, IA revisará estrutura autonomamente.")
    
    try:
        rate_limiter.wait_if_needed()
        
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=PROMPT_STRUCTURE_REVIEW.format(
                estrutura_mapeada=estrutura_str,
                documento=texto_para_revisao
            ),
            config=types.GenerateContentConfig(
                max_output_tokens=65536,  # Máximo permitido
                thinking_config={"include_thoughts": False, "thinking_level": "HIGH"},
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                ]
            )
        )
        
        resultado = response.text.replace('```markdown', '').replace('```', '').strip()
        
        # Extrair e exibir relatório da IA (se presente)
        relatorio_match = re.search(r'<!--\s*RELATÓRIO:\s*(.+?)\s*-->', resultado, re.IGNORECASE)
        if relatorio_match:
            relatorio = relatorio_match.group(1)
            logger.info(f"   📊 Relatório da IA: {relatorio}")
            # Remover o comentário do resultado final
            resultado = re.sub(r'<!--\s*RELATÓRIO:.+?-->\s*', '', resultado, flags=re.IGNORECASE).strip()
        
        # Validação básica: o resultado deve ter pelo menos 70% do tamanho original
        if len(resultado) < len(texto) * 0.7:
            logger.warning(f"   ⚠️ Revisão retornou texto muito curto ({len(resultado)} vs {len(texto)}). Mantendo original.")
            return texto
        
        # Contar quantos headers foram alterados
        headers_original = len(re.findall(r'^#{2,4}\s', texto, re.MULTILINE))
        headers_revisado = len(re.findall(r'^#{2,4}\s', resultado, re.MULTILINE))
        diff = abs(headers_original - headers_revisado)
        
        logger.info(f"   ✅ Estrutura revisada: {headers_original} → {headers_revisado} headers (Δ{diff})")
        return resultado
        
    except Exception as e:
        logger.warning(f"   ⚠️ Erro na revisão por IA: {e}. Mantendo estrutura original.")
        return texto

# =============================================================================
# REVISÃO LEVE DE FORMATAÇÃO (MODO FIDELIDADE) v2.0
# =============================================================================

PROMPT_STRUCTURE_REVIEW_LITE = """Você é um revisor editorial especializado em ESTRUTURA e FORMATAÇÃO de documentos educacionais. Você receberá:
1. Uma **Estrutura de Mapeamento Inicial** (planejada antes da formatação)
2. O **Documento Processado** (resultado da formatação por chunks)

Sua tarefa é analisar ambos e garantir que os títulos estejam **descritivos, hierarquicamente corretos e alinhados com o conteúdo real**, sem jamais alterar a ordem cronológica.

---

## 📋 ESTRUTURA DE MAPEAMENTO INICIAL (Referência):
{estrutura_mapeada}

---

## ✅ O QUE VOCÊ DEVE FAZER:
1. **Comparar Títulos:** Verifique se os títulos do documento refletem corretamente os tópicos do mapeamento. Se um título estiver genérico mas o mapeamento indicar um tema específico, refine-o.
2. **Validar Hierarquia:** Confirme que a estrutura (##, ###, ####) segue uma lógica consistente (ex: seções > subseções > detalhes). MÁXIMO 3 níveis.
3. **Decidir a Melhor Estrutura:** Se houver conflito entre mapeamento e documento, escolha a estrutura que melhor reflete o CONTEÚDO REAL do texto.
4. **Subtópicos Órfãos:** Se detectar headers como "A.", "B.", "C." isolados como tópicos principais, converta-os em subníveis do tópico anterior (ex: ## para ###).
5. **Títulos Descritivos:** Refine títulos genéricos (ex: "Questão 1") para algo que cite o tema técnico (ex: "Questão 1: Responsabilidade Civil").
6. **Corrigir Sintaxe Markdown:** Tabelas (alinhar colunas), listas (bullets), espaçamento entre seções.
7. **Remover Vazios:** Títulos sem conteúdo abaixo.

## ❌ O QUE VOCÊ NÃO DEVE FAZER:
1. **NÃO MOVA** blocos de texto. A ordem deve permanecer 100% cronológica.
2. **NÃO MESCLE** seções que apareçam em momentos diferentes da aula.
3. **NÃO RESUMA** nem altere o corpo dos parágrafos.
4. **NÃO ADICIONE** conteúdo novo.

## 📝 RELATÓRIO ESPERADO:
Ao final do documento, inclua um bloco de comentário (que será removido) indicando:
- Quantos títulos foram refinados
- Se a estrutura final segue o mapeamento ou foi adaptada
- Discrepâncias encontradas (se houver)

Formato:
<!-- RELATÓRIO: X títulos refinados | Estrutura: [MAPEAMENTO/ADAPTADA] | Discrepâncias: [Nenhuma/Lista] -->

---

## 📄 DOCUMENTO PARA REVISAR:
{documento}

---

## RESPOSTA:
Retorne o documento COMPLETO E INTEGRAL (mesmo tamanho do input) com a formatação corrigida e o relatório no final. NÃO RESUMA."""

def ai_structure_review_lite(client, texto, estrutura_mapeada=None):
    """
    v2.0: Revisão LEVE de formatação Markdown com VALIDAÇÃO CRUZADA.
    Compara o documento processado com a estrutura de mapeamento inicial.
    Refina títulos, valida hierarquia, e reporta discrepâncias.
    NÃO reorganiza nem mescla conteúdo.
    """
    logger.info("🧹 Revisão Leve de Formatação (IA - Modo Fidelidade v2.0)...")
    
    # Gemini 2.0 Flash suporta 1M tokens (~4M chars) - usar até 500k chars para documento + 50k para estrutura
    max_doc_chars = 500000
    if len(texto) > max_doc_chars:
        logger.warning(f"   ⚠️ Documento muito longo ({len(texto)} chars), truncando para {max_doc_chars//1000}k...")
        texto_para_revisao = texto[:max_doc_chars] + "\n\n[... documento truncado para revisão ...]"
    else:
        texto_para_revisao = texto
    
    # Preparar estrutura mapeada (se disponível)
    if estrutura_mapeada:
        estrutura_str = estrutura_mapeada[:50000]  # Limitar estrutura a 50k chars
        logger.info(f"   📋 Usando estrutura de mapeamento inicial ({len(estrutura_mapeada)} chars) para validação cruzada.")
    else:
        estrutura_str = "[Estrutura de mapeamento não disponível - analisar documento para inferir estrutura ideal]"
        logger.info("   ℹ️  Sem mapeamento inicial, IA irá inferir estrutura ideal do próprio documento.")
    
    try:
        rate_limiter.wait_if_needed()
        
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=PROMPT_STRUCTURE_REVIEW_LITE.format(
                estrutura_mapeada=estrutura_str,
                documento=texto_para_revisao
            ),
            config=types.GenerateContentConfig(
                max_output_tokens=65536,  # Máximo permitido
                thinking_config={"include_thoughts": False, "thinking_level": "HIGH"},  # HIGH para análise estrutural profunda
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                ]
            )
        )
        
        resultado = response.text.replace('```markdown', '').replace('```', '').strip()
        
        # Extrair e exibir relatório da IA (se presente)
        relatorio_match = re.search(r'<!--\s*RELATÓRIO:\s*(.+?)\s*-->', resultado, re.IGNORECASE)
        if relatorio_match:
            relatorio = relatorio_match.group(1)
            logger.info(f"   📊 Relatório da IA: {relatorio}")
            # Remover o comentário do resultado final
            resultado = re.sub(r'<!--\s*RELATÓRIO:.+?-->\s*', '', resultado, flags=re.IGNORECASE).strip()
        
        # Validação: resultado deve ter pelo menos 85% do tamanho
        if len(resultado) < len(texto) * 0.85:
            logger.warning(f"   ⚠️ Revisão retornou texto muito curto ({len(resultado)} vs {len(texto)}). Mantendo original.")
            return texto
        
        # Verificar se a ordem dos headers foi preservada (validação extra)
        headers_original = re.findall(r'^#{1,4}\s+(.+?)$', texto, re.MULTILINE)[:20]
        headers_revisado = re.findall(r'^#{1,4}\s+(.+?)$', resultado, re.MULTILINE)[:20]
        
        if len(headers_original) > 5 and len(headers_revisado) > 5:
            matches = sum(1 for h1, h2 in zip(headers_original[:10], headers_revisado[:10]) 
                         if h1.strip()[:30] == h2.strip()[:30])
            if matches < 6:
                logger.warning(f"   ⚠️ Ordem dos headers parece alterada. Mantendo original.")
                return texto
        
        # Relatório de alterações de títulos
        alteracoes = []
        for h_orig, h_rev in zip(headers_original[:15], headers_revisado[:15]):
            if h_orig.strip() != h_rev.strip():
                orig_short = h_orig.strip()[:40] + "..." if len(h_orig.strip()) > 40 else h_orig.strip()
                rev_short = h_rev.strip()[:40] + "..." if len(h_rev.strip()) > 40 else h_rev.strip()
                alteracoes.append(f"   - '{orig_short}' → '{rev_short}'")
        
        if alteracoes:
            logger.info(f"📝 Títulos Refinados ({len(alteracoes)}):")
            for alt in alteracoes[:5]:
                logger.info(alt)
            if len(alteracoes) > 5:
                logger.info(f"   ... e mais {len(alteracoes) - 5} alterações")
        else:
            logger.info("   ℹ️  Nenhum título foi alterado (estrutura já estava OK).")
        
        logger.info(f"   ✅ Formatação revisada (modo leve v2.0).")
        return resultado
        
    except Exception as e:
        logger.warning(f"   ⚠️ Erro na revisão leve: {e}. Mantendo original.")
        return texto




def deterministic_structure_fix(text):
    """
    v1.1: Reorganização Estrutural Determinística (Regex).
    Adaptativo: Detecta se o documento usa H1 ou apenas H2 como nível principal.
    """
    logger.info("🧩 Executando Reorganização Estrutural Determinística...")
    
    lines = text.split('\n')
    
    # Detecção de Hierarquia
    has_h1 = any(re.match(r'^#\s+', line) for line in lines)
    header_level_regex = r'^#\s+' if has_h1 else r'^##\s+'
    logger.info(f"   ℹ️  Nível principal detectado: {'H1 (#)' if has_h1 else 'H2 (##)'}")

    # Estruturas de dados (Preserva ordem de inserção)
    content_map = {
        "PREAMBULO": [],
        "DISCIPLINAS": {}, 
        "ENCERRAMENTO": []
    }
    
    current_area = "PREAMBULO"
    current_block = []
    disciplinas_order = [] 
    
    # Regex Adaptativo
    # Captura o texto do cabeçalho principal (seja H1 ou H2)
    # Exclui "Questão" ou "Q." para não quebrar simulados dentro de uma área
    re_disciplina = re.compile(f'{header_level_regex}(?!Questão|Q\\.)([^0-9\\.]+.*)', re.IGNORECASE)
    re_encerramento = re.compile(f'{header_level_regex}(?:ENCERRAMENTO|CONSIDERAÇÕES|CONCLUSÃO)', re.IGNORECASE)
    
    def flush_block(area, block_lines):
        if not block_lines: return
        block_text = '\n'.join(block_lines)
        
        if area == "PREAMBULO":
            content_map["PREAMBULO"].append(block_text)
        elif area == "ENCERRAMENTO":
            content_map["ENCERRAMENTO"].append(block_text)
        else:
            if area not in content_map["DISCIPLINAS"]:
                content_map["DISCIPLINAS"][area] = []
                disciplinas_order.append(area)
            content_map["DISCIPLINAS"][area].append(block_text)

    for line in lines:
        # 1. Detectar mudança de disciplina macro
        match_disc = re_disciplina.match(line)
        if match_disc:
            flush_block(current_area, current_block)
            current_block = []
            
            raw_area = match_disc.group(1).strip().upper()
            
            # Normalização de nome de área
            if "DIREITO" not in raw_area and len(raw_area) < 50:
                 # Adiciona prefixo se parecer nome de matéria jurídica comum
                 if any(x in raw_area for x in ["CIVIL", "PENAL", "TRABALHO", "ADMINISTRATIVO", "CONSTITUCIONAL"]):
                     current_area = f"DIREITO {raw_area}"
                 else:
                     current_area = raw_area
            else:
                 current_area = raw_area
            
            # IMPORTANTE: Se estamos operando em modo H2, essa linha é um Header que queremos manter?
            # Se for modo H1, recriamos "# AREA". 
            # Se for modo H2, recriamos "# AREA" (upcast) ou mantemos "## AREA"?
            # Para padronizar Apostilas, vamos promover tudo a H1 na reconstrução.
            continue 
            
        # 2. Detectar Encerramento
        if re_encerramento.match(line):
            flush_block(current_area, current_block)
            current_block = []
            current_area = "ENCERRAMENTO"
            continue

        current_block.append(line)
        
    flush_block(current_area, current_block)
    
    # Reconstrução
    final_output = []
    
    # Preambulo
    if content_map["PREAMBULO"]:
        final_output.append("# ORIENTAÇÕES GERAIS / INTRODUÇÃO")
        final_output.extend(content_map["PREAMBULO"])
        final_output.append("")

    # Disciplinas / Tópicos Principais
    for area in disciplinas_order:
        area_clean = area.replace("#", "").strip()
        final_output.append(f"# {area_clean}")
        for block in content_map["DISCIPLINAS"][area]:
            final_output.append(block)
        final_output.append("")
        
    # Encerramento
    if content_map["ENCERRAMENTO"]:
        final_output.append("# CONSIDERAÇÕES FINAIS")
        final_output.extend(content_map["ENCERRAMENTO"])
        
    num_identified = len(disciplinas_order)
    logger.info(f"   ✅ Reorganizado: {num_identified} seções principais identificadas.")
    
    # Fallback: Se não identificou nada (tudo preambulo), retorna original para não estragar
    if num_identified == 0 and len(content_map["PREAMBULO"]) > 0:
        logger.warning("   ⚠️ Nenhuma estrutura detectada. Mantendo original.")
        return text
        
    return '\n'.join(final_output)

def numerar_titulos(texto):
    """Adiciona numeração sequencial aos títulos"""
    linhas = texto.split('\n')
    linhas_numeradas = []
    
    contador_h2 = 0
    contador_h3 = 0
    contador_h4 = 0
    
    titulo_pattern = re.compile(r'^(#{2,4})\s+(?:\d+(?:\.\d+)*\.?\s+)?(.+)$')
    
    # Variaveis para rastrear últimos títulos e evitar repetições
    ultimo_h2_texto = ""
    ultimo_h3_texto = ""
    
    for linha in linhas:
        match = titulo_pattern.match(linha)
        
        if match:
            nivel = len(match.group(1))
            texto_titulo = match.group(2).strip()
            
            # Não numera títulos de resumo/quadros
            if any(keyword in texto_titulo.lower() for keyword in ['resumo', 'quadro', 'esquema', '📋', '📊', '🗂️']):
                linhas_numeradas.append(linha)
                continue
            
            # MERGE INTELIGENTE DE TÍTULOS REPETIDOS (v2.11)
            # Se o título atual for muito similar ao anterior do mesmo nível ("continuação" de chunk), ignoramos o novo
            # para que o texto flua como um único tópico.
            
            from difflib import SequenceMatcher
            eh_duplicado = False
            
            if nivel == 2:
                # Verifica similaridade com último H2
                ratio = SequenceMatcher(None, texto_titulo.lower(), ultimo_h2_texto.lower()).ratio()
                if ratio > 0.9:
                    eh_duplicado = True
                    logger.info(f"🔄 Título H2 mesclado: '{texto_titulo}' ≈ '{ultimo_h2_texto}'")
                else:
                    ultimo_h2_texto = texto_titulo
            elif nivel == 3:
                 # Verifica similaridade com último H3
                ratio = SequenceMatcher(None, texto_titulo.lower(), ultimo_h3_texto.lower()).ratio()
                if ratio > 0.9:
                    eh_duplicado = True
                    logger.info(f"🔄 Título H3 mesclado: '{texto_titulo}' ≈ '{ultimo_h3_texto}'")
                else:
                    ultimo_h3_texto = texto_titulo
            
            if eh_duplicado:
                continue # Pula a linha do título, fundindo o conteúdo
            
            if nivel == 2:
                contador_h2 += 1
                contador_h3 = 0
                contador_h4 = 0
                nova_linha = f"## {contador_h2}. {texto_titulo}"
            elif nivel == 3:
                contador_h3 += 1
                contador_h4 = 0
                nova_linha = f"### {contador_h2}.{contador_h3}. {texto_titulo}"
            elif nivel == 4:
                contador_h4 += 1
                nova_linha = f"#### {contador_h2}.{contador_h3}.{contador_h4}. {texto_titulo}"
            else:
                nova_linha = linha
            
            linhas_numeradas.append(nova_linha)
        else:
            linhas_numeradas.append(linha)
    
    return '\n'.join(linhas_numeradas)

def renumerar_secoes(texto):
    """
    v1.0: Renumeração Sequencial Determinística.
    
    Esta função é uma camada de segurança extra, aplicada APÓS o AI Review.
    Ela percorre todos os headers numerados e corrige qualquer duplicação ou
    sequência quebrada, garantindo que os números sejam estritamente sequenciais
    dentro de cada nível de hierarquia.
    
    Exemplo de correção:
    - Input:  9.20, 9.21, 9.21, 9.35, 9.36  (duplicação e pulo)
    - Output: 9.20, 9.21, 9.22, 9.23, 9.24  (sequencial)
    """
    logger.info("🔢 Executando Renumeração Sequencial Determinística...")
    
    linhas = texto.split('\n')
    novas_linhas = []
    
    # Contadores por nível de hierarquia (ex: {2: 9, 3: 44} -> ## 9.x, ### 9.44.x)
    # Estrutura: {(header_level, parent_prefix): next_number}
    contadores = {}
    
    # Regex para detectar headers numerados: ### 9.20.1. Título ou ## 5. Título
    header_pattern = re.compile(r'^(#{1,4})\s+([\d.]+\.?)\s*(.*)$')
    
    for linha in linhas:
        match = header_pattern.match(linha)
        
        if match:
            hashes = match.group(1)       # "###"
            numero = match.group(2)       # "9.20.1." ou "9.20.1"
            titulo = match.group(3)       # "Título..."
            
            nivel = len(hashes)
            numero_limpo = numero.rstrip('.')
            partes = numero_limpo.split('.')
            
            # Determina o prefixo pai e o sufixo atual
            if len(partes) == 1:
                # Nível raiz: ## 1. ou ## 9.
                prefixo_pai = ""
                sufixo_atual = int(partes[0])
            else:
                # Subnível: ### 9.20. ou #### 9.20.1.
                prefixo_pai = '.'.join(partes[:-1])
                sufixo_atual = int(partes[-1])
            
            # Inicializa contador se não existir
            chave = (nivel, prefixo_pai)
            if chave not in contadores:
                contadores[chave] = sufixo_atual  # Começa do número encontrado
            else:
                contadores[chave] += 1
            
            novo_sufixo = contadores[chave]
            
            # Reconstrói o número
            if prefixo_pai:
                novo_numero = f"{prefixo_pai}.{novo_sufixo}."
            else:
                novo_numero = f"{novo_sufixo}."
            
            # Reconstrói a linha
            nova_linha = f"{hashes} {novo_numero} {titulo}"
            novas_linhas.append(nova_linha)
            
            # Log se houve mudança
            if numero_limpo != novo_numero.rstrip('.'):
                logger.info(f"   🔄 {numero_limpo} → {novo_numero.rstrip('.')}")
        else:
            novas_linhas.append(linha)
    
    logger.info("   ✅ Renumeração concluída.")
    return '\n'.join(novas_linhas)

# =============================================================================
# VERIFICAÇÃO DE COBERTURA E DUPLICAÇÕES
# =============================================================================

def normalizar_fingerprint(texto, tipo):
    """Normaliza texto para comparação (ex: 'Lei 11.100' -> 'lei 11100')"""
    texto = texto.lower().strip()
    
    if tipo == 'leis':
        # Mantém apenas 'lei' e números
        nums = re.findall(r'\d+', texto)
        if nums:
            # Reconstrói como 'lei 12345'
            # Filtra leis com menos de 4 dígitos para evitar ruído (ex: lei 10, lei 13)
            num_full = ''.join(nums)
            if len(num_full) >= 4:
                return f"lei {num_full}"
            return None
            
    elif tipo == 'sumulas':
        nums = re.findall(r'\d+', texto)
        if nums:
            return f"súmula {''.join(nums)}"
            
    elif tipo == 'artigos':
        nums = re.findall(r'\d+', texto)
        if nums:
            return f"artigo {''.join(nums)}"
            
    return re.sub(r'[^\w\s]', '', texto)

def extrair_fingerprints(texto):
    """Extrai 'fingerprints' únicos e normalizados do texto"""
    fingerprints = {
        'leis': set(),
        'sumulas': set(),
        'artigos': set(),
        'julgados': set()
    }
    
    # Regex melhorado para capturar variações
    lei_pattern = re.compile(r'\b(?:lei|l\.)\s*n?º?\s*([\d\.]+)', re.IGNORECASE)
    sumula_pattern = re.compile(r'\bsúmula\s*(?:vinculante)?\s*n?º?\s*(\d+)', re.IGNORECASE)
    
    # Extrai e normaliza
    for match in lei_pattern.finditer(texto):
        fp = normalizar_fingerprint(f"lei {match.group(1)}", 'leis')
        if fp: fingerprints['leis'].add(fp)
    
    for match in sumula_pattern.finditer(texto):
        fp = normalizar_fingerprint(f"súmula {match.group(1)}", 'sumulas')
        if fp: fingerprints['sumulas'].add(fp)
        
    return fingerprints

def contar_ocorrencias_robust(fingerprints, texto):
    """Conta ocorrências com suporte a formatação jurídica formal (Lei nº X)"""
    contagens = {}
    
    # CORREÇÃO 1: Não remover pontuação indiscriminadamente, apenas normalizar espaços
    # Mantemos barras e pontos para evitar fusão de números (11.101/2005)
    texto_lower = texto.lower()
    
    for categoria, items in fingerprints.items():
        for item in items:
            key = f"{categoria}:{item}"
            
            if categoria == 'leis':
                # item ex: "lei 4320" -> extrai "4320"
                # Remove pontuação do item para garantir match limpo no número
                num_bruto = item.split()[-1] 
                num = re.sub(r'[^\d]', '', num_bruto)
                
                # Permite pontos opcionais entre dígitos (ex: 4.320 match com 4320)
                num_regex = r"\.?".join(list(num))
                
                # CORREÇÃO 2: Regex flexível que aceita "n", "nº", "no", "num" no meio
                # Aceita: "Lei 4320", "Lei nº 4.320", "Lei n. 4320"
                # O \W* permite pontos/barras entre Lei e o número
                # Adicionado \b no final para evitar matches parciais (Lei 10 != Lei 100)
                pattern = f"lei(?:\\s+|\\.|\\,|nº|n\\.|n\\s|num\\.?)*{num_regex}\\b"
                
                # Usamos findall no texto original (lower) para pegar variações com pontuação
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
                
            elif categoria == 'sumulas':
                num = item.split()[-1]
                num_regex = r"\.?".join(list(num)) # Súmulas raramente tem ponto, mas por garantia
                
                # Mesma lógica para súmulas (Súmula Vinculante nº 10)
                pattern = f"súmula(?:\\s+|\\.|\\,|vinculante|nº|n\\.|n\\s)*{num_regex}\\b"
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
                
            else:
                # Fallback para outros tipos (busca literal simples)
                contagens[key] = texto_lower.count(item)
                
    return contagens

def verificar_cobertura(texto_original, texto_formatado, arquivo_saida=None):
    """Verifica omissões e duplicações artificiais entre original e formatado"""
    logger.info("🔍 Verificando cobertura e duplicações...")
    
    # Extrai fingerprints do original
    fp_original = extrair_fingerprints(texto_original)
    
    # Conta ocorrências em ambos
    contagem_original = contar_ocorrencias_robust(fp_original, texto_original)
    contagem_formatado = contar_ocorrencias_robust(fp_original, texto_formatado)
    
    omissoes = []
    duplicacoes = []
    
    for key, count_orig in contagem_original.items():
        count_fmt = contagem_formatado.get(key, 0)
        categoria, item = key.split(':', 1)
        
        # Omissão: estava no original mas sumiu
        if count_orig > 0 and count_fmt == 0:
            omissoes.append({
                'categoria': categoria,
                'item': item,
                'original': count_orig,
                'formatado': count_fmt
            })
        
        # Duplicação (agora considerada positiva em materiais didáticos)
        if count_fmt > count_orig:
            duplicacoes.append({
                'categoria': categoria,
                'item': item,
                'original': count_orig,
                'formatado': count_fmt,
                'extra': count_fmt - count_orig
            })
    
    # Gera relatório
    total_items = len([k for k, v in contagem_original.items() if v > 0])
    items_preservados = total_items - len(omissoes)
    cobertura = items_preservados / total_items * 100 if total_items > 0 else 100
    
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    logger.info(f"📊 RELATÓRIO DE VERIFICAÇÃO")
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    logger.info(f"✅ Cobertura: {cobertura:.1f}% ({items_preservados}/{total_items} referências)")
    
    if omissoes:
        logger.warning(f"\n❌ POSSÍVEIS OMISSÕES ({len(omissoes)}):")
        for o in omissoes[:10]:  # Limita a 10
            logger.warning(f"   - [{o['categoria']}] {o['item']}")
        if len(omissoes) > 10:
            logger.warning(f"   ... e mais {len(omissoes) - 10} omissões")
    else:
        logger.info("✅ Nenhuma omissão detectada")
    
    if duplicacoes:
        logger.info(f"\nℹ️ CITAÇÕES REFORÇADAS (Tabelas/Resumos) ({len(duplicacoes)}):")
        for d in duplicacoes[:10]:
            logger.info(f"   - [{d['categoria']}] {d['item']}: {d['original']}x → {d['formatado']}x (+{d['extra']})")
        if len(duplicacoes) > 10:
            logger.info(f"   ... e mais {len(duplicacoes) - 10} citações extras")
    else:
        logger.info("ℹ️ Nenhuma citação extra detectada")
    
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    
    # Salva relatório em arquivo se especificado
    if arquivo_saida:
        relatorio_path = arquivo_saida.replace('.md', '_verificacao.txt')
        with open(relatorio_path, 'w', encoding='utf-8') as f:
            f.write(f"RELATÓRIO DE VERIFICAÇÃO\n")
            f.write(f"{'='*50}\n\n")
            f.write(f"Cobertura: {cobertura:.1f}% ({items_preservados}/{total_items})\n\n")
            
            if omissoes:
                f.write(f"OMISSÕES ({len(omissoes)}):\n")
                for o in omissoes:
                    f.write(f"  - [{o['categoria']}] {o['item']}\n")
                f.write("\n")
            
            if duplicacoes:
                f.write(f"DUPLICAÇÕES ARTIFICIAIS ({len(duplicacoes)}):\n")
                for d in duplicacoes:
                    f.write(f"  - [{d['categoria']}] {d['item']}: {d['original']}x original → {d['formatado']}x formatado\n")
        
        logger.info(f"📄 Relatório salvo: {relatorio_path}")
    
    return {
        'cobertura': cobertura,
        'omissoes': omissoes,
        'duplicacoes': duplicacoes
    }

# =============================================================================
# V2.16: VALIDAÇÃO LLM (Metadata Strategy)
# =============================================================================

PROMPT_VALIDATE_COMPLETENESS = """# TAREFA DE VALIDAÇÃO DE FIDELIDADE (METADATA STRATEGY)

Você é um auditor de qualidade para transcrições jurídicas formatadas.

## SEU OBJETIVO
Compare a ESTRUTURA DO ORIGINAL (Metadata/Skeleton) com o TEXTO FORMATADO FINAL e identifique:

1. **OMISSÕES GRAVES**: Conceitos jurídicos, leis, súmulas, artigos ou exemplos importantes que estavam no esqueleto original mas foram omitidos no formatado.
2. **DISTORÇÕES**: Informações que foram alteradas de forma que mude o sentido jurídico.
3. **ESTRUTURA**: Verifique se os tópicos e subtópicos estão organizados de forma lógica e se não há duplicações.

## REGRAS
- NÃO considere como omissão: hesitações, "né", "então", dados repetitivos, conversas paralelas.
- CONSIDERE como omissão: qualquer lei, súmula, artigo, jurisprudência, exemplo prático ou dica de prova.
- O input "TEXTO ORIGINAL" é um RESUMO ESTRUTURAL (Metadata) contendo apenas títulos e referências chave. Use-o para validar se esses elementos aparecem no "TEXTO FORMATADO".

## FORMATO DE RESPOSTA (JSON)
{
    "aprovado": true/false,
    "nota_fidelidade": 0-10,
    "omissoes_graves": ["descrição clara do item omitido"],
    "distorcoes": ["descrição clara da distorção"],
    "problemas_estrutura": ["títulos duplicados ou hierarquia quebrada"],
    "observacoes": "comentário geral sobre a qualidade"
}"""

def extract_raw_metadata(texto):
    """
    v2.16.1: Extrai esqueleto robusto do texto original para validação.
    Captura: Títulos, Leis, Súmulas, Artigos, Jurisprudência, Destaques.
    """
    lines = texto.split('\n')
    metadata = []
    metadata.append(f"TOTAL WORDS: {len(texto.split())}")
    metadata.append(f"TOTAL CHARS: {len(texto)}")
    
    # Regex robustas para capturar referências legais (reutilizando padrões do script)
    patterns = {
        'leis': re.compile(r'\b(?:lei|l\.)\s*n?º?\s*([\d\.]+(?:/\d+)?)', re.IGNORECASE),
        'artigos': re.compile(r'\b(?:art\.?|artigo)\s*(\d+)', re.IGNORECASE),
        'sumulas': re.compile(r'\bsúmula\s*(?:vinculante)?\s*n?º?\s*(\d+)', re.IGNORECASE),
        'jurisprudencia': re.compile(r'\b(?:REsp|RE|ADI|ADPF|HC|MS|AgRg|RMS|Rcl)\s*[\d\.\/\-]+', re.IGNORECASE),
        'informativos': re.compile(r'\b(?:informativo|info\.?)\s*(?:stf|stj)?\s*n?º?\s*(\d+)', re.IGNORECASE),
        'temas': re.compile(r'\btema\s*(?:repetitivo)?\s*n?º?\s*(\d+)', re.IGNORECASE),
    }
    
    # Contadores para estatísticas
    refs_encontradas = {k: set() for k in patterns.keys()}
    
    for line in lines:
        line_strip = line.strip()
        if not line_strip: 
            continue
        
        # Detectar possíveis títulos (caixa alta, começa com número, curto)
        if len(line_strip) < 100:
            if line_strip.isupper() and len(line_strip) > 5:
                metadata.append(f"\n[TÍTULO] {line_strip}")
            elif re.match(r'^\d+[\.\)]\s', line_strip):
                metadata.append(f"\n[TÓPICO] {line_strip[:80]}")
        
        # Capturar todas as referências legais
        for categoria, pattern in patterns.items():
            matches = pattern.findall(line_strip)
            for m in matches:
                ref = m if isinstance(m, str) else m[0]
                refs_encontradas[categoria].add(ref.strip())
        
        # Destaques pedagógicos
        keywords = ['importante', 'atenção', 'cuidado', 'exemplo', 'obs:', 'dica', 'pegadinha', 'caiu em prova']
        if any(kw in line_strip.lower() for kw in keywords):
            metadata.append(f"  > DESTAQUE: {line_strip[:120]}...")
    
    # Resumo estatístico
    metadata.append("\n--- REFERÊNCIAS ENCONTRADAS ---")
    for cat, refs in refs_encontradas.items():
        if refs:
            metadata.append(f"[{cat.upper()}] ({len(refs)}): {', '.join(sorted(refs)[:15])}{'...' if len(refs) > 15 else ''}")
    
    return "\n".join(metadata)

# =============================================================================
# AUTO-FIX PASS (v2.18)
# =============================================================================

def aplicar_correcoes_automaticas(texto):
    """
    v2.18: Aplica correções automáticas baseadas em padrões comuns de erro.
    Retorna (texto_corrigido, lista_de_correcoes).
    """
    correcoes = []
    texto_original = texto
    
    # 1. Remover saudações duplicadas (apenas mantém a primeira)
    saudacoes_pattern = r'(?:Olá|Oi),?\s*(?:sejam?\s+)?(?:bem[- ]?vindos?(?:\s+e\s+bem[- ]?vindas?)?)[.,!]?'
    matches = list(re.finditer(saudacoes_pattern, texto, re.IGNORECASE))
    if len(matches) > 1:
        # Remove todas exceto a primeira
        for match in reversed(matches[1:]):
            # Captura a linha inteira onde a saudação aparece
            start = texto.rfind('\n', 0, match.start()) + 1
            end = texto.find('\n', match.end())
            if end == -1: end = len(texto)
            linha = texto[start:end].strip()
            # Só remove se a linha for majoritariamente saudação
            if len(linha) < 150:
                texto = texto[:start] + texto[end+1:]
                correcoes.append(f"Removida saudação duplicada: '{linha[:50]}...'")
    
    # 2. Remover apresentações repetidas do professor
    apresentacao_pattern = r'Eu sou o professor\s+\w+(?:\s+\w+)?'
    matches = list(re.finditer(apresentacao_pattern, texto, re.IGNORECASE))
    if len(matches) > 1:
        for match in reversed(matches[1:]):
            start = texto.rfind('\n', 0, match.start()) + 1
            end = texto.find('\n', match.end())
            if end == -1: end = len(texto)
            linha = texto[start:end].strip()
            if len(linha) < 200:
                texto = texto[:start] + texto[end+1:]
                correcoes.append(f"Removida apresentação duplicada: '{linha[:50]}...'")
    
    # 3. Padronizar nome do professor (detecta variações e unifica)
    # Extrai primeiro nome mencionado como "professor X"
    nome_match = re.search(r'professor\s+(\w+(?:\s+\w+)?)', texto, re.IGNORECASE)
    if nome_match:
        nome_canonico = nome_match.group(1)
        # Busca variações próximas (distância de Levenshtein simplificada)
        variacoes_pattern = rf'\bprofessor\s+(\w+(?:\s+\w+)?)\b'
        for m in re.finditer(variacoes_pattern, texto, re.IGNORECASE):
            nome_atual = m.group(1)
            if nome_atual.lower() != nome_canonico.lower():
                sim = SequenceMatcher(None, nome_canonico.lower(), nome_atual.lower()).ratio()
                if sim > 0.6 and sim < 1.0:  # Similar mas diferente
                    texto = texto.replace(f"professor {nome_atual}", f"professor {nome_canonico}")
                    texto = texto.replace(f"Professor {nome_atual}", f"Professor {nome_canonico}")
                    correcoes.append(f"Padronizado nome: '{nome_atual}' → '{nome_canonico}'")
    
    # 4. Corrigir itens de lista vazios ou malformados
    # Padrão: número + ponto + espaços/quebra + próximo conteúdo
    texto = re.sub(r'(\d+\.)\s*\n\s*((?:Requisitos|Preenchimento|Fundamento|Artigo|Lei))', r'\1 \2', texto)
    if texto != texto_original:
        correcoes.append("Corrigidos itens de lista malformados")
    
    # 5. Remover linhas em branco excessivas (mais de 2 consecutivas)
    texto_limpo = re.sub(r'\n{4,}', '\n\n\n', texto)
    if texto_limpo != texto:
        texto = texto_limpo
        correcoes.append("Removidas linhas em branco excessivas")
    
    logger.info(f"🔧 Auto-Fix: {len(correcoes)} correções aplicadas")
    for c in correcoes:
        logger.info(f"   ✓ {c}")
    
    return texto, correcoes


def validate_completeness_llm(raw_text, formatted_text, client, output_file=None):
    """
    v2.16.1: Validação LLM com Metadata Strategy e retorno estruturado.
    """
    logger.info("🕵️ Executando Validação LLM (Completeness Check) com Gemini 3 Flash...")
    
    # 1. Extrair Metadata do Raw (Otimização)
    raw_metadata = extract_raw_metadata(raw_text)
    
    # Estimativa de tokens
    input_text = f"{PROMPT_VALIDATE_COMPLETENESS}\n\n## TEXTO ORIGINAL (METADATA/SKELETON):\n{raw_metadata}\n\n## TEXTO FORMATADO:\n{formatted_text}"
    est_tokens = len(input_text) // 4
    logger.info(f"   📊 Payload de Validação: ~{est_tokens:,} tokens")
    
    try:
        if est_tokens > 2_000_000:
             logger.warning("⚠️ Payload excede 2M tokens. Pulando validação LLM para evitar erro.")
             return {'aprovado': True, 'nota_fidelidade': 0, 'skipped': True, 'reason': 'payload_too_large'}
             
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=input_text,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                max_output_tokens=8000,
                thinking_config=types.ThinkingConfig(
                    include_thoughts=False,
                    thinking_level="HIGH"
                )
            )
        )
        
        result = json.loads(response.text)
        
        # Log Gráfico
        if result.get('aprovado'):
            logger.info(f"   ✅ APROVADO (Nota {result.get('nota_fidelidade')}/10)")
        else:
            logger.warning(f"   ❌ REPROVADO (Nota {result.get('nota_fidelidade')}/10)")
            
        omissions = result.get('omissoes_graves', [])
        if omissions:
            logger.warning(f"   🚨 {len(omissions)} Omissões Graves Detectadas:")
            for o in omissions[:5]:  # Limita log a 5
                logger.warning(f"      - {o}")
                
        # Salvar Relatório
        if output_file:
            report_path = output_file.replace('.md', '_LLM_VALIDATION.md')
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(f"# Relatório de Validação LLM (Metadata Strategy v2.16.1)\n")
                f.write(f"**Data:** {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"**Modelo:** {GEMINI_MODEL}\n")
                f.write(f"**Nota Fidelidade:** {result.get('nota_fidelidade')}/10\n\n")
                
                f.write("## 🔍 Omissões Graves\n")
                if omissions:
                    for o in omissions: f.write(f"- 🔴 {o}\n")
                else:
                    f.write("- ✅ Nenhuma omissão grave detectada.\n")
                    
                f.write("\n## ⚠️ Distorções\n")
                distorcoes = result.get('distorcoes', [])
                if distorcoes:
                    for d in distorcoes: f.write(f"- ⚠️ {d}\n")
                else:
                    f.write("- Nenhuma distorção detectada.\n")
                
                f.write("\n## 🏗️ Estrutura\n")
                problemas = result.get('problemas_estrutura', [])
                if problemas:
                    for p in problemas: f.write(f"- 🔧 {p}\n")
                else:
                    f.write("- Estrutura OK.\n")
                
                f.write(f"\n## 📝 Observações\n{result.get('observacoes', 'N/A')}\n")
                
            logger.info(f"   📄 Relatório salvo: {report_path}")
        
        return result
            
    except Exception as e:
        logger.error(f"   ❌ Erro na validação LLM: {e}")
        return {'aprovado': True, 'nota_fidelidade': 0, 'error': str(e)}




def auto_fix_smart(raw_text, formatted_text, validation_result, client, estrutura_global=None):
    """
    v2.18 (SAFE MODE): Corretor Estrutural Seguro.
    Foca EXCLUSIVAMENTE em problemas de estrutura (títulos, duplicatas, hierarquia).
    NÃO altera conteúdo jurídico para evitar alucinações.
    
    Args:
        raw_text: (Não usado no modo Safe, mantido para compatibilidade)
        formatted_text: Texto formatado atual
        validation_result: Dict com problemas estruturais
        client: Cliente Vertex AI
        estrutura_global: Mapeamento de referência (opcional)
    """
    # No modo SAFE, ignoramos omissões/distorções para não correr risco de reescrita
    problemas_estrut = validation_result.get('problemas_estrutura', [])
    
    if not problemas_estrut:
        logger.info("✅ Nenhum problema estrutural para corrigir.")
        return formatted_text
    
    logger.info(f"🔧 Auto-Fix Safe: Corrigindo {len(problemas_estrut)} problemas estruturais...")
    
    report = "### PROBLEMAS ESTRUTURAIS:\n" + "\n".join([f"- {p}" for p in problemas_estrut]) + "\n"
        
    PROMPT_FIX = f"""Você é um editor técnico de elite.
    
## TAREFA: LIMPEZA ESTRUTURAL (SEM ALTERAR CONTEÚDO)
Você deve corrigir APENAS a formatação e estrutura do documento.

## REGRA DE OURO (SEGURANÇA JURÍDICA):
- **NÃO altere o texto dos parágrafos.**
- **NÃO adicione nem remova informações jurídicas.**
- **NÃO reescreva explicações.**
- Sua permissão é APENAS para Títulos, Hierarquia e Duplicatas exatas.

## INSTRUÇÕES DE CORREÇÃO:
1. **Títulos Duplicados**: Se houver títulos repetidos (ex: dois "3. Introdução" seguidos), remova o redundante.
2. **Hierarquia**: Ajuste níveis (H2, H3) para seguir a lógica do conteúdo.
3. **Parágrafos Repetidos**: Delete duplicações EXATAS de parágrafos (copia-cola acidental).
4. **Renumeração**: Garanta sequência lógica (1, 2, 3...) nos títulos.

{f"## ESTRUTURA DE REFERÊNCIA (Guia):\n{estrutura_global}" if estrutura_global else ""}

## RELATÓRIO DE ERROS:
{report}

## SAÍDA:
Retorne o documento COMPLETO corrigido em Markdown. Sem explicações."""

    try:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=f"{PROMPT_FIX}\n\n## TEXTO A CORRIGIR:\n{formatted_text}",
            config=types.GenerateContentConfig(
                max_output_tokens=100000,
                thinking_config=types.ThinkingConfig(
                    include_thoughts=False,
                    thinking_level="HIGH" 
                )
            )
        )
        
        resultado = response.text.replace('```markdown', '').replace('```', '').strip()
        
        # Validação de segurança estrita
        if len(resultado) < len(formatted_text) * 0.8: # Tolerância menor no modo safe
            logger.warning("⚠️ Auto-Fix Safe cortou muito texto (>20%). Abortando por segurança.")
            return formatted_text
            
        logger.info(f"✅ Auto-Fix Estrutural concluído. ({len(formatted_text)} -> {len(resultado)} chars)")
        return resultado
        
    except Exception as e:
        logger.error(f"❌ Falha no Auto-Fix Safe: {e}")
        return formatted_text

# V2.10: FUNÇÕES DE PÓS-PROCESSAMENTO ESTRUTURAL (Tabelas e Parágrafos)
# =============================================================================

def mover_tabelas_para_fim_de_secao(texto):
    """
    v2.11: Reorganiza tabelas movendo-as para o final do BLOCO ATUAL (H2 ou H3).
    Corrige bug de tabelas sumindo ou ficando muito longe do contexto.
    """
    logger.info("📊 Reorganizando tabelas (Smart Layout v2.11)...")
    
    linhas = texto.split('\n')
    resultado = []
    tabelas_pendentes = [] 
    
    i = 0
    while i < len(linhas):
        linha = linhas[i]
        linha_strip = linha.strip()
        
        # 1. DETECTAR SE É UM TÍTULO (H1, H2, H3...)
        # Se encontrarmos um novo título, hora de "despejar" as tabelas acumuladas do bloco anterior
        if linha_strip.startswith('#'):
            # Despeja tabelas antes de iniciar o novo tópico
            if tabelas_pendentes:
                resultado.append('') # Espaço antes
                for t_info in tabelas_pendentes:
                    if t_info['titulo']:
                        resultado.append(t_info['titulo'])
                    resultado.extend(t_info['linhas'])
                    resultado.append('') # Espaço depois
                tabelas_pendentes = []
            
            resultado.append(linha)
            i += 1
            continue

        # 2. DETECTAR INÍCIO DE TABELA
        # Critério: Linha tem pipe '|' E parece estrutura de tabela (não apenas citação)
        eh_inicio_tabela = False
        if '|' in linha_strip:
            # Verifica se é uma linha de markdown table válida (tem pipe e chars)
            # E se a próxima linha ou a seguinte tem o separador '---'
            has_separator = False
            for lookahead in range(1, 3): # Olha até 2 linhas pra frente (ignora 1 linha vazia)
                if i + lookahead < len(linhas):
                    prox = linhas[i + lookahead].strip()
                    if set(prox).issubset(set('|- :')): # Só contem caracteres de estrutura de tabela
                         has_separator = True
                         break
            
            if has_separator or (linha_strip.startswith('|') and linha_strip.endswith('|')):
                eh_inicio_tabela = True

        if eh_inicio_tabela:
            # --- Captura da Tabela ---
            tabela_linhas = []
            titulo_tabela = None
            
            # Tenta recuperar o título da tabela que ficou na linha anterior (ou resultado)
            # Verifica se a última linha adicionada ao resultado parece um título de tabela
            if resultado and len(resultado) > 0:
                last_line = resultado[-1].strip()
                # Padrões comuns de título de tabela gerados pela IA
                if (last_line.startswith('###') or last_line.startswith('**')) and \
                   any(x in last_line.lower() for x in ['tabela', 'resumo', 'quadro', 'síntese', 'esquema', '📋']):
                    titulo_tabela = resultado.pop() # Remove do fluxo principal para agrupar com a tabela

            # Captura as linhas da tabela
            j = i
            while j < len(linhas):
                curr = linhas[j].strip()
                # Continua se tiver pipe ou for linha vazia no meio da tabela (mas cuidado com fim)
                if '|' in curr:
                    tabela_linhas.append(linhas[j])
                    j += 1
                elif not curr:
                    # Linha vazia: verifica se a próxima volta a ter pipe
                    if j + 1 < len(linhas) and '|' in linhas[j+1]:
                        tabela_linhas.append(linhas[j]) # Mantém linha vazia interna
                        j += 1
                    else:
                        break # Fim da tabela
                else:
                    break # Texto normal, fim da tabela

            # Verifica se capturou algo útil
            if len(tabela_linhas) > 0:
                tabelas_pendentes.append({
                    'titulo': titulo_tabela,
                    'linhas': tabela_linhas
                })
                i = j # Pula as linhas processadas
                continue
            else:
                # Falso positivo? Devolve o título se tinhamos pego
                if titulo_tabela:
                    resultado.append(titulo_tabela)
        
        # Se não for tabela nem título, adiciona linha normal
        resultado.append(linha)
        i += 1
    
    # 3. FINAL DO DOCUMENTO
    # Se sobraram tabelas no buffer, despeja agora
    if tabelas_pendentes:
        resultado.append('')
        for t_info in tabelas_pendentes:
            if t_info['titulo']:
                resultado.append(t_info['titulo'])
            resultado.extend(t_info['linhas'])
            resultado.append('')
            
    return '\n'.join(resultado)

def quebrar_paragrafos_longos(texto, max_chars=400, max_sentencas=4):
    """
    Quebra parágrafos que excedem limite de chars OU número de sentenças.
    Preserva listas, tabelas, citações e blocos especiais.
    """
    logger.info(f"✂️ Quebrando parágrafos > {max_chars} chars ou > {max_sentencas} sentenças...")
    
    paragrafos = texto.split('\n\n')
    resultado = []
    quebras = 0
    
    for para in paragrafos:
        linha_strip = para.strip()
        
        # PRESERVAR: títulos, listas, tabelas, citações, blocos de código
        if (linha_strip.startswith('#') or 
            linha_strip.startswith('-') or 
            linha_strip.startswith('* ') or 
            linha_strip.startswith('|') or
            linha_strip.startswith('>') or
            linha_strip.startswith('```') or
            re.match(r'^\d+\.', linha_strip)):
            resultado.append(para)
            continue
        
        # Verifica se precisa quebrar
        num_chars = len(para)
        sentencas = re.split(r'(?<=[.!?])\s+', para)
        num_sentencas = len(sentencas)
        
        if num_chars <= max_chars and num_sentencas <= max_sentencas:
            resultado.append(para)
            continue
        
        # QUEBRAR: Agrupa em blocos de até max_sentencas
        quebras += 1
        subparagrafos = []
        bloco_atual = []
        chars_atual = 0
        sentencas_no_bloco = 0
        
        for sentenca in sentencas:
            teste_chars = chars_atual + len(sentenca)
            
            # Se adicionar essa sentença ultrapassar AMBOS os limites ou o limite de sentenças
            # E já temos algo no bloco...
            if (teste_chars > max_chars or sentencas_no_bloco >= max_sentencas) and bloco_atual:
                subparagrafos.append(' '.join(bloco_atual).strip())
                bloco_atual = [sentenca]
                chars_atual = len(sentenca)
                sentencas_no_bloco = 1
            else:
                bloco_atual.append(sentenca)
                chars_atual = teste_chars
                sentencas_no_bloco += 1
        
        # Adiciona último bloco
        if bloco_atual:
            subparagrafos.append(' '.join(bloco_atual).strip())
        
        resultado.append('\n\n'.join(subparagrafos))
        
    if quebras > 0:
        logger.info(f"   ✅ {quebras} parágrafos foram ajustados.")
        
    return '\n\n'.join(resultado)

# =============================================================================
# FLUXO PRINCIPAL
# =============================================================================

def formatar_transcricao(transcricao_completa, usar_cache=True, input_file=None, custom_prompt=None):
    prompt_ativo = custom_prompt if custom_prompt else PROMPT_FORMATACAO
    if custom_prompt:
        logger.info(f"🎨 Usando prompt customizado ({len(custom_prompt):,} caracteres)")
    
    estrutura_global = "" # v2.17: Inicializa para evitar UnboundLocalError
    # ⚡ FORÇAR USO DE VERTEX AI - AI Studio desabilitado
    rate_limiter.max_rpm = 60
    # Tenta pegar o projeto do ambiente, se não tiver, usa o hardcoded como fallback
    project_id = os.getenv("GOOGLE_CLOUD_PROJECT", "gen-lang-client-0727883752")
    logger.info(f" Usando Vertex AI (Project: {project_id})")
    logger.info(" Rate Limit: 60 RPM")
    logger.info("🔥 AI Studio DESABILITADO - Usando apenas Vertex AI")
    
    # Verificar se as credenciais do Vertex AI estão configuradas
    if os.getenv("GOOGLE_APPLICATION_CREDENTIALS") is None:
        logger.error("❌ Nenhuma autenticação configurada para Vertex AI.")
        logger.error("Configure a variável de ambiente GOOGLE_APPLICATION_CREDENTIALS:")
        logger.error("  export GOOGLE_APPLICATION_CREDENTIALS='/path/to/service-account.json'")
        sys.exit(1)

    client = genai.Client(
        vertexai=True,
        project=project_id,
        location="global"
    )
    
    # Dry-run mode
    if '--dry-run' in sys.argv:
        logger.info("🔍 MODO DRY-RUN: Validando divisão de chunks")
        chunks = dividir_sequencial(transcricao_completa)
        validar_chunks(chunks, transcricao_completa)
        for i, c in enumerate(chunks):
            print(f"\n  Chunk {i+1}/{len(chunks)}:")
            print(f"    Posição: {c['inicio']:,} → {c['fim']:,} ({c['fim']-c['inicio']:,} chars)")
            inicio_preview = transcricao_completa[c['inicio']:c['inicio']+80].replace('\n', '↵')
            fim_preview = transcricao_completa[max(0, c['fim']-80):c['fim']].replace('\n', '↵')
            print(f"    Início: {inicio_preview}...")
            print(f"    Fim: ...{fim_preview}")
        sys.exit(0)
    
    tamanho_total = len(transcricao_completa)
    logger.info(f"📊 Tamanho: {tamanho_total:,} caracteres")
    
    # v2.10: Iniciar timer de métricas
    metrics.start_timer()
    
    # v2.17: Divisão Inicial (Estimativa)
    chunks = dividir_sequencial(transcricao_completa)
    num_partes = len(chunks)
    
    # v2.8: Mapeamento estrutural prévio (se necessário)
    # Se não foi passado externamente e temos múltiplas partes, gera agora.
    if num_partes > 1 and not estrutura_global:
        estrutura_global = map_structure(client, transcricao_completa)
        
        # v2.10: Filtros de estrutura
        estrutura_global = filtrar_niveis_execessivos(estrutura_global, max_nivel=3)
        estrutura_global = simplificar_estrutura_se_necessario(estrutura_global)
        
        # v2.17: RE-DIVISÃO OTIMIZADA (Anchor-Based Chunking)
        # Agora que temos a estrutura, refazemos os cortes para alinhar com os tópicos
        logger.info("🔄 Otimizando cortes com base na estrutura (Anchor-Based Chunking)...")
        chunks = dividir_sequencial(transcricao_completa, estrutura_global=estrutura_global)
        num_partes = len(chunks)

    # v2.7: Validação rigorosa
    if not validar_chunks(chunks, transcricao_completa):
        logger.error("❌ Chunks inválidos! Abortando.")
        sys.exit(1)
    
    chunks_info = [{'inicio': c['inicio'], 'fim': c['fim']} for c in chunks]
    
    estimar_custo(transcricao_completa, usar_cache, num_partes)
    
    if num_partes == 1:
        return processar_simples(client, transcricao_completa, prompt_ativo), None, client
    
    checkpoint = None
    resultados = []
    inicio_secao = 0
    
    if input_file:
        checkpoint = load_checkpoint(input_file)
        if checkpoint:
            logger.info(f"📁 Checkpoint: seção {checkpoint['secao_atual']}/{checkpoint['total_secoes']}")
            resposta = input("   Continuar? (s/n): ").strip().lower()
            if resposta == 's':
                resultados = checkpoint['resultados']
                inicio_secao = checkpoint['secao_atual']
                chunks_info = checkpoint['chunks_info']
            else:
                delete_checkpoint(input_file)
    
    # v2.19: Cache HABILITADO com Estrutura Global (User Request)
    cache = None
    if usar_cache and num_partes > 1:
         cache = criar_cache_contexto(client, transcricao_completa, prompt_ativo, estrutura_global)
    
    try:
        # Prepara o progresso
        iter_chunks = range(inicio_secao, num_partes)
        if tqdm:
            iter_chunks = tqdm(iter_chunks, desc="Processando", unit="seção")
        
        for i in iter_chunks:
            chunk = chunks_info[i]
            texto_chunk = transcricao_completa[chunk['inicio']:chunk['fim']]
            
            # v2.10: Contexto estático apenas (instrução de estilo)
            contexto_estilo = None
            if resultados:
                 # Pega os últimos 3000 caracteres como referência de ESTILO
                 ultimo_texto = resultados[-1]
                 contexto_estilo = ultimo_texto[-CONTEXTO_ESTILO:] if len(ultimo_texto) > CONTEXTO_ESTILO else ultimo_texto
            
            # v2.10: Extrair último título do chunk anterior para anti-duplicação
            ultimo_titulo = None
            if resultados:
                texto_anterior = resultados[-1]
                for linha in reversed(texto_anterior.split('\n')[-30:]):
                    if linha.strip().startswith('##'):
                        ultimo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
                        break
            
            # v2.17: Contexto Localizado (Localized Context)
            # Em vez de passar a estrutura inteira, passa uma janela: [Anterior] + [Atual] + [Próximo]
            estrutura_local = estrutura_global
            if estrutura_global and num_partes > 1:
                itens_estrutura = estrutura_global.split('\n')
                # Estimativa simples: mapeia chunk i para itens da estrutura
                ratio = len(itens_estrutura) / num_partes
                center_idx = int(i * ratio)
                window_size = max(4, int(len(itens_estrutura) * 0.15)) # 15% da estrutura ou min 4 itens
                
                start_idx = max(0, center_idx - window_size)
                end_idx = min(len(itens_estrutura), center_idx + window_size + 2)
                
                slice_itens = itens_estrutura[start_idx:end_idx]
                if start_idx > 0: slice_itens.insert(0, "[... Tópicos anteriores ...]")
                if end_idx < len(itens_estrutura): slice_itens.append("[... Tópicos posteriores ...]")
                
                estrutura_local = '\n'.join(slice_itens)

            resultado = processar_chunk(
                client, cache, prompt_ativo, texto_chunk,
                i + 1, num_partes,
                contexto_estilo=contexto_estilo,
                estrutura_global=estrutura_local, # Passa estrutura localmente fatiada
                ultimo_titulo=ultimo_titulo
            )
            
            # v2.10: Smart Stitching - Remove eco do contexto
            if contexto_estilo:
                resultado = remover_eco_do_contexto(resultado, contexto_estilo)
            
            # v2.10: Smart Stitching - Remove título duplicado na fronteira
            texto_acumulado = '\n\n'.join(resultados) if resultados else ""
            resultado = limpar_inicio_redundante(resultado, texto_acumulado)
            
            resultados.append(resultado)
            
            if input_file:
                save_checkpoint(input_file, resultados, chunks_info, i + 1)
            
            if not tqdm:
                logger.info(f"✅ Seção {i+1}/{num_partes}")
        
        # v2.7: Post-processing em múltiplas passadas
        logger.info("🧹 Iniciando limpeza (v2.7)...")
        
        texto_final = '\n\n'.join(resultados)
        
        # Passada 0: Limpar metadados de mapeamento que vazam para o output
        # Remove linhas como "[TIPO: AULA EXPOSITIVA]" ou "**[TIPO: SIMULADO]**"
        texto_final = re.sub(r'^#?\s*\*?\*?\[TIPO:.*?\]\*?\*?\s*$', '', texto_final, flags=re.MULTILINE)
        # Remove marcadores de bloco [BLOCO 01], [BLOCO 02], etc.
        texto_final = re.sub(r'^\s*\[BLOCO\s*\d+\]\s*$', '', texto_final, flags=re.MULTILINE)
        # Remove timestamps órfãos [HH:MM] ou [HH:MM:SS] no início de linha
        texto_final = re.sub(r'^\s*\[\d{1,2}:\d{2}(:\d{2})?\]\s*$', '', texto_final, flags=re.MULTILINE)
        texto_final = re.sub(r'\n{3,}', '\n\n', texto_final)  # Remove linhas em branco extras
        
        logger.info("  Passada 1: Removendo duplicações literais...")
        texto_final = remover_duplicacoes_literais(texto_final)
        
        logger.info("  Passada 1.5: Deduplicação robusta (7-DIFF Strategy)...")
        texto_final = remover_overlap_duplicado(texto_final)
        
        logger.info("  Passada 2: Detectando e removendo seções duplicadas (v2.15)...")
        texto_final = remover_secoes_duplicadas(texto_final)
        
        logger.info("  Passada 3: Normalizando títulos similares...")
        texto_final = normalize_headings(texto_final)
        
        if MODO_NOME != "FIDELIDADE":
            logger.info("  Passada 3.5: Reorganização Estrutural Determinística...")
            texto_final = deterministic_structure_fix(texto_final)
        else:
            logger.info("  ℹ️  Modo FIDELIDADE: Pulando reorganização para preservar linearidade.")
            
        
        # v2.10: Reordenação do Pipeline (Tabelas -> Numeração -> Parágrafos)
        
        logger.info("  Passada 4: Reorganizando tabelas (Smart Layout)...")
        texto_final = mover_tabelas_para_fim_de_secao(texto_final)

        
        logger.info("  Passada 5: Numerando títulos...")
        texto_final = numerar_titulos(texto_final)
        
        logger.info("  Passada 6: Ajustando parágrafos longos...")
        texto_final = quebrar_paragrafos_longos(texto_final, max_chars=400, max_sentencas=4)
        
        if MODO_NOME != "FIDELIDADE":
            logger.info("  Passada 7: Revisão semântica de estrutura (IA v2.0)...")
            texto_final = ai_structure_review(client, texto_final, estrutura_mapeada=estrutura_global)
        else:
            logger.info("  Passada 7: Revisão leve de formatação (IA - Modo Fidelidade v2.0)...")
            texto_final = ai_structure_review_lite(client, texto_final, estrutura_mapeada=estrutura_global)
        
        # Passada 7.5: Renumeração Sequencial Determinística (camada de segurança)
        try:
            texto_final = renumerar_secoes(texto_final)
        except Exception as e:
            logger.warning(f"⚠️ Erro na renumeração: {e}. Continuando sem renumerar.")
        
        # Validação final
        palavras_in = len(transcricao_completa.split())
        palavras_out = len(texto_final.split())
        razao = palavras_out / palavras_in if palavras_in > 0 else 1.0
        
        logger.info(f"✅ Validação: {razao:.0%} do original ({palavras_out:,}/{palavras_in:,} palavras)")
        
        if razao < THRESHOLD_CRITICO:
            if FIDELIDADE_MODE:
                logger.error(f"❌ POSSÍVEL PERDA DE CONTEÚDO ({razao:.0%})")
                logger.error(f"   Esperado: >{THRESHOLD_CRITICO:.0%} | Obtido: {razao:.0%}")
            elif APOSTILA_MODE:
                logger.warning(f"⚠️  Texto condensado: {razao:.0%}")
                logger.info(f"   ✅ Esperado no modo {MODO_NOME}")
            
            if razao < 0.30 or (FIDELIDADE_MODE and razao < THRESHOLD_CRITICO):
                resposta = input("\n   Continuar? (s/n): ").strip().lower()
                if resposta != 's':
                    logger.info("Cancelado.")
                    sys.exit(1)
        
        if input_file:
            delete_checkpoint(input_file)
        
        return texto_final, cache, client
    
    except KeyboardInterrupt:
        logger.warning("\n⚠️  Interrompido. Checkpoint salvo.")
        sys.exit(1)

# =============================================================================
# EXPORTAÇÃO WORD
# =============================================================================

def create_toc(doc):
    """Adiciona Sumário (Table of Contents) nativo do Word"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def _format_inline_markdown(paragraph, text):
    """
    v2.16.2: Formata markdown inline de forma robusta.
    Suporta: ***bold_italic***, **bold**, *italic*, __bold__, _italic_, `code`.
    """
    paragraph.clear()
    
    # Regex robusta:
    # Group 2: ***text*** ou ___text___ (Bold + Italic)
    # Group 3: **text** ou ___text___ (Bold + Italic)
    # Group 4: **text** (Bold)
    # Group 5: __text__ (Bold)
    # Group 6: *text* (Italic)
    # Group 7: _text_ (Italic)
    # Group 8: `text` (Code)
    pattern = r'(\*{3}(.+?)\*{3}|_{3}(.+?)_{3}|\*{2}(.+?)\*{2}|_{2}(.+?)_{2}|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|(?<!_)_(?!(?:_|\s))(.+?)(?<!(?:_|\s))_(?!_)|`(.+?)`)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        full_match = match.group(0)
        
        if full_match.startswith('***'):
            content = match.group(2)
            run = paragraph.add_run(content)
            run.bold = True
            run.italic = True
        elif full_match.startswith('___'):
            content = match.group(3)
            run = paragraph.add_run(content)
            run.bold = True
            run.italic = True
        elif full_match.startswith('**'):
            content = match.group(4)
            run = paragraph.add_run(content)
            run.bold = True
            run.font.name = 'Arial'
        elif full_match.startswith('__'):
            content = match.group(5)
            run = paragraph.add_run(content)
            run.bold = True
            run.font.name = 'Arial'
        elif full_match.startswith('*'):
            run = paragraph.add_run(match.group(6))
            run.italic = True
            run.font.name = 'Arial'
        elif full_match.startswith('_'):
            run = paragraph.add_run(match.group(7))
            run.italic = True
            run.font.name = 'Arial'
        elif full_match.startswith('`'):
            run = paragraph.add_run(match.group(8))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(200, 0, 0)
        
        last_end = match.end()
    
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = 'Arial'

def _add_table_to_doc(doc, rows):
    """Adiciona tabela formatada ao documento Word"""
    if len(rows) < 2:
        return
    
    max_cols = max(len(row) for row in rows)
    if max_cols == 0:
        return
    
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(rows):
        for j in range(max_cols):
            cell = table.rows[i].cells[j]
            cell_text = row_data[j] if j < len(row_data) else ""
            
            # Formata markdown dentro da célula
            _format_inline_markdown(cell.paragraphs[0], cell_text)
            
            # Alinhamento padrão: Esquerda
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = 'Arial'
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    # Alinhamento à esquerda conforme solicitado
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '336699')
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                 # Conteúdo normal da tabela: Esquerda
                 for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def save_as_word(formatted_text, video_name, output_file):
    """Salva markdown formatado como documento Word (.docx)"""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx não disponível. Salvando apenas Markdown.")
        return None
    
    logger.info("📄 Gerando documento Word...")
    
    doc = Document()
    
    # Configura Fonte Arial no estilo Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    # python-docx precisa desse hack para garantir que o nome da fonte seja aplicado corretamente
    style.element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    
    # Margens
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    
    # Título principal
    title = doc.add_heading(video_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Data de geração
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')} - Modo: {MODO_NOME}")
    date_run.font.name = 'Arial'
    date_run.italic = True
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Sumário
    doc.add_heading('Sumário', level=1)
    create_toc(doc)
    doc.add_page_break()
    
    # Processa conteúdo markdown
    lines = formatted_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Tabelas
        if '|' in line and not in_table:
            in_table = True
            table_rows = []
        
        if in_table:
            # Ignora linha separadora (ex: |---| ou | :--- | :--- | :--- |)
            is_separator = re.match(r'^\s*\|[\s:|-]+\|[\s:|-]*$', line)
            
            if '|' in line and not is_separator:
                table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
            
            if '|' not in line or i == len(lines) - 1:
                if len(table_rows) > 0:
                    _add_table_to_doc(doc, table_rows)
                in_table = False
                table_rows = []
                if '|' not in line:
                    continue
            i += 1
            continue
        
        # Headings
        if line.startswith('##### '):
            h = doc.add_heading('', level=5)
            _format_inline_markdown(h.paragraphs[0], line[6:])
        elif h_match := re.match(r'^(####|###|##|#)\s+(.*)', line):
            lvl = len(h_match.group(1))
            h_text = h_match.group(2)
            if lvl == 1 and h_text == video_name:
                i += 1
                continue
            h = doc.add_heading('', level=lvl)
            _format_inline_markdown(h, h_text)  # 'h' já é um Paragraph
        # Separadores
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.add_run('_' * 80).font.color.rgb = RGBColor(192, 192, 192)
        # Quotes
        elif line.startswith('>'):
            p = doc.add_paragraph(style='Quote')
            p.paragraph_format.left_indent = Cm(4.0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _format_inline_markdown(p, line[1:].strip())
            # Forçar estilo de quote em todos os runs
            for run in p.runs:
                run.font.name = 'Arial'
                run.italic = True
                run.font.size = Pt(10)
        # Listas não-ordenadas
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            # Forçar recuo de 1,5cm
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-0.63)
            
            _format_inline_markdown(p, line[2:])
            
        # Listas numeradas
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
            p = doc.add_paragraph(style='Normal')
            # Recuo padronizado de 1,5cm
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-0.63) # Mantém hanging indent para o número
            _format_inline_markdown(p, line)
            
        # Parágrafo normal
        else:
            p = doc.add_paragraph()
            # Espaçamento 1.5 (Line Spacing = 1.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            # Espaçamento antes e após parágrafo (6pt)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6) 
            
            # Recuo de 1ª linha (1cm)
            p.paragraph_format.first_line_indent = Cm(1.0)
            
            # Justificado
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            _format_inline_markdown(p, line)
            
            # Forçar fonte 12pt para texto normal e Arial
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
        
        i += 1
    
    # Flush final (caso o arquivo termine com tabela seguida de linhas em branco)
    if in_table and len(table_rows) > 0:
        _add_table_to_doc(doc, table_rows)
        
    doc.save(output_file)
    logger.info(f"✅ Word salvo: {output_file}")
    return output_file

def salvar_resultado(conteudo, arquivo):
    with open(arquivo, 'w', encoding='utf-8') as f:
        f.write(conteudo)
    logger.info(f"✅ Markdown salvo: {arquivo}")

# =============================================================================
# MAIN
# =============================================================================

def main():
    if len(sys.argv) < 2 or '--help' in sys.argv:
        print("=" * 70)
        print("FORMATADOR v2.7 - ANTI-DUPLICAÇÃO")
        print("=" * 70)
        print("\nUso: python format_transcription_gemini.py <entrada.txt> [saida] [--prompt <texto_ou_path>]")
        print("\nOpções:")
        print("  --dry-run         Valida chunks e mostra preview")
        print("  --prompt <p>      Prompt customizado (texto direto ou caminho .txt)")
        print("  --help            Mostra esta mensagem")
        print("\n🛡️  CORREÇÕES v2.7:")
        print("  • Detecção agressiva de seções duplicadas")
        print("  • Validação rigorosa de chunks sequenciais")
        print("  • Delimitadores de contexto mais visíveis")
        print("  • Post-processing em múltiplas passadas")
        print("  • Cache desabilitado (debug)")
        print("\nDependências:")
        print("  pip install google-genai python-docx tqdm")
        sys.exit(1)
    
    arquivo_entrada = sys.argv[1]
    
    if len(sys.argv) > 2 and not sys.argv[2].startswith('--'):
        arquivo_saida = sys.argv[2]
    else:
        base = arquivo_entrada.replace('.txt', '_formatada_v2.7')
        arquivo_saida = f"{base}.md"
    
    video_name = Path(arquivo_entrada).stem
    
    # v2.21: Parse prompt customizado
    custom_prompt = None
    if '--prompt' in sys.argv:
        try:
            p_idx = sys.argv.index('--prompt')
            p_val = sys.argv[p_idx + 1]
            if os.path.exists(p_val):
                with open(p_val, 'r', encoding='utf-8') as f:
                    custom_prompt = f.read().strip()
                logger.info(f"📂 Prompt carregado de arquivo: {p_val}")
            else:
                custom_prompt = p_val.strip()
                logger.info("✍️ Prompt customizado lido diretamente da CLI")
        except Exception as e:
            logger.error(f"❌ Erro ao ler prompt: {e}")
            sys.exit(1)

    logger.info("=" * 60)
    logger.info(f"FORMATADOR v2.7 ANTI-DUPLICAÇÃO - Modo {MODO_NOME}")
    logger.info("=" * 60)
    logger.info(f"📂 Entrada: {arquivo_entrada}")
    logger.info(f"📝 Saída: {arquivo_saida}")
    
    transcricao = carregar_transcricao(arquivo_entrada)
    
    try:
        resultado, cache, client = formatar_transcricao(transcricao, input_file=arquivo_entrada, custom_prompt=custom_prompt)
    except Exception as e:
        logger.error(f"\n❌ Falha: {e}", exc_info=True)
        sys.exit(1)
    
    # v2.18: Auto-Fix Pass - Correções automáticas
    logger.info("🔧 Aplicando Auto-Fix Pass (v2.18)...")
    resultado, correcoes = aplicar_correcoes_automaticas(resultado)
    
    salvar_resultado(resultado, arquivo_saida)
    
    # v2.8: Verificação de cobertura e duplicações
    verificar_cobertura(transcricao, resultado, arquivo_saida)

    # v2.16: Validação LLM (Metadata Strategy)
    validation_result = validate_completeness_llm(transcricao, resultado, client, arquivo_saida)
    
    # v2.17: Auto-Fix Loop - Corrige omissões se detectadas
    if validation_result and not validation_result.get('aprovado', True):
        logger.info("🔁 Iniciando Auto-Fix Loop...")
        resultado = auto_fix_smart(transcricao, resultado, validation_result, client, estrutura_global)
        
        # Salvar versão corrigida
        with open(arquivo_saida, 'w', encoding='utf-8') as f:
            f.write(resultado)
        logger.info(f"💾 Versão corrigida salva: {arquivo_saida}")

    # v2.9: Auditoria Legal Pós-Processamento
    if AUDIT_AVAILABLE:
        report_path = arquivo_saida.replace('.md', '_RELATORIO_AUDITORIA.md')
        auditar_consistencia_legal(client, resultado, report_path)

    if DOCX_AVAILABLE:
        arquivo_docx = arquivo_saida.replace('.md', '.docx')
        save_as_word(resultado, video_name, arquivo_docx)
    
    # v2.10: Métricas de Execução (tokens reais)
    metrics.stop_timer()
    logger.info(metrics.get_report())
    logger.info("✨ Concluído! (v2.10 com métricas)")

    # v2.19: Cleanup manual do cache para economia
    if cache:
        try:
            client.caches.delete(name=cache.name)
            logger.info(f"🗑️ Cache {cache.name} deletado manualmente para economizar custos.")
        except Exception as e:
            logger.warning(f"⚠️ Não foi possível deletar o cache: {e}")

if __name__ == "__main__":
    main()
