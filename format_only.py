import os
import time
import asyncio
from pathlib import Path
from openai import OpenAI, AsyncOpenAI, APIError
from colorama import Fore, init
from tqdm import tqdm
import re
import traceback
import json
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

init(autoreset=True)

class VomoFormatter:
    def __init__(self):
        """Formatador de transcrições Vomo (apenas texto) - VERSÃO OTIMIZADA"""
        
        # Cache para evitar re-identificação de speakers
        self.speaker_cache = {}
        print(f"{Fore.CYAN}🚀 Inicializando Vomo Formatter (OpenAI GPT-5-mini)...")
        
        # Carrega variável de ambiente com override
        from dotenv import load_dotenv
        load_dotenv(override=True)
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError(f"{Fore.RED}❌ Configure: export OPENAI_API_KEY='sk-...'")
        
        # Cliente OpenAI direto (sem OpenRouter)
        self.async_client = AsyncOpenAI(
            api_key=api_key
        )
        
        # Modelo GPT-5-mini (OpenAI)
        self.model = "gpt-5-mini-2025-08-07"
        
        # Cliente síncrono (se necessário)
        self.client = OpenAI(
            api_key=api_key
        )
        self.llm_client = self.client
    
    def _get_cache_path(self, file_path, cache_type="FORMATTED"):
        """Gera caminho do arquivo de cache baseado no arquivo original"""
        base_path = Path(file_path)
        cache_file = base_path.parent / f"{base_path.stem}_CACHE_{cache_type}.json"
        return str(cache_file)
    
    def _load_cache(self, cache_file):
        """Carrega resultado processado do cache se disponível e válido"""
        if not os.path.exists(cache_file):
            return None
        
        try:
            with open(cache_file, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)
            
            # Validação básica
            required_fields = ['formatted_text', 'timestamp', 'version']
            if not all(field in cache_data for field in required_fields):
                print(f"{Fore.YELLOW}   ⚠️ Cache inválido (campos faltando), reprocessando...")
                return None
            
            # Verifica versão (compatibilidade)
            if cache_data.get('version') != '0.3':
                print(f"{Fore.YELLOW}   ⚠️ Cache de versão antiga ({cache_data.get('version')}), reprocessando...")
                return None
            
            print(f"{Fore.GREEN}   📂 Cache válido encontrado ({cache_data['timestamp']})")
            return cache_data
        
        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ Erro ao carregar cache ({e}), reprocessando...")
            return None
    
    def _save_cache(self, cache_file, formatted_text, original_file):
        """Salva resultado processado em cache JSON"""
        try:
            cache_data = {
                'formatted_text': formatted_text,
                'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
                'original_file': os.path.basename(original_file),
                'version': '0.3',
                'model': self.model
            }
            
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(cache_data, f, ensure_ascii=False, indent=2)
            
            print(f"{Fore.CYAN}   💾 Cache salvo: {os.path.basename(cache_file)}")
        
        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ Erro ao salvar cache ({e})")
    


    SYSTEM_PROMPT_FORMAT = """# SEU PAPEL
    Você é um **revisor sênior de material didático jurídico** com 20 anos de experiência em preparatórios para concursos de procuradoria (PGM/PGE). Seu trabalho é transformar transcrições brutas de aulas em apostilas de alta qualidade, mantendo 100% do conteúdo técnico.

    ---

    # MISSÃO

    Transformar a transcrição de aula em uma **apostila didática** com estilo formal, impessoal e expositivo, **PRESERVANDO TODO O CONTEÚDO ORIGINAL**.

    ⚠️ **REGRA ABSOLUTA: PRESERVE 100% DO CONTEÚDO - NÃO RESUMA, NÃO OMITA NADA.**

    ---

    # EXEMPLOS DE TRANSFORMAÇÃO (Few-Shot)

    ## Exemplo 1: Transformação de Linguagem
    **INPUT (Transcrição bruta):**
    > "Então pessoal, a Lei 8.666, né, ela fala lá no artigo 37 que a gente não pode fazer contratação direta, tá? Exceto em alguns casos, tipo... quando tem emergência, sabe?"

    **OUTPUT (Apostila formatada):**
    > Observa-se que o *Art. 37 da Lei nº 8.666/93* estabelece a vedação à contratação direta como regra geral. Contudo, o legislador previu hipóteses excepcionais, notadamente os casos de emergência, nos termos do *Art. 24, inciso IV* do referido diploma legal.

    ## Exemplo 2: Preservação de Dicas de Prova
    **INPUT:**
    > "Isso aqui cai muito em prova, hein! Atenção total nesse artigo 37."

    **OUTPUT:**
    > Reveste-se de especial importância este dispositivo para fins de certames públicos, sendo recorrentemente objeto de cobrança em provas de procuradorias.

    ## Exemplo 3: Preservação de Exemplos Completos
    **INPUT:**
    > "Por exemplo, teve um caso no Rio que o prefeito tentou fazer uma contratação emergencial pra comprar uns equipamentos, mas o TCE glosou porque não tinha a urgência real."

    **OUTPUT:**
    > A título ilustrativo, cumpre mencionar caso ocorrido no Estado do Rio de Janeiro, no qual determinado prefeito municipal tentou realizar contratação emergencial para aquisição de equipamentos. Contudo, o Tribunal de Contas do Estado glosou o procedimento, fundamentando sua decisão na ausência de urgência real que justificasse a dispensa de licitação.

    ---

    # ❌ ERROS A EVITAR

    **ERRO 1 - Resumir exemplos:**
    - ❌ "O professor citou casos de jurisprudência"
    - ✅ Transcreva cada caso completo com nome, ano e decisão

    **ERRO 2 - Generalizar referências:**
    - ❌ "Conforme a legislação de licitações"
    - ✅ "Conforme o *Art. 37, caput, da Lei nº 8.666/93*"

    **ERRO 3 - Omitir dicas de prova:**
    - ❌ [Ignorar comentário do professor sobre a prova]  
    - ✅ Inclua: "Reveste-se de especial importância para certames..."

    **ERRO 4 - Perder contexto histórico:**
    - ❌ "Houve mudanças na lei"
    - ✅ "A reforma introduzida pela Lei nº 14.133/2021 alterou significativamente..."

    ---

    # DIRETRIZES DE ESTILO (APOSTILA FORMAL)

    ## Tom e Linguagem
    ✅ **USE SEMPRE:**
    - **Terceira pessoa impessoal**: "observa-se", "constata-se", "verifica-se", "cumpre destacar"
    - **Conectivos acadêmicos**: "nesse sentido", "ademais", "cumpre ressaltar", "destarte"
    - **Verbos formais**: "configura", "caracteriza", "evidencia", "compreende"

    ❌ **NUNCA USE:**
    - Primeira/segunda pessoa: "eu", "você", "nós", "a gente"
    - Vícios de fala: "né", "tá", "aí", "então", "tipo", "beleza"
    - Expressões informais: "pô", "cara", "galera", "pessoal"

    ---

    # ESTRUTURA OBRIGATÓRIA DO OUTPUT

    ```markdown
    # [Título da Aula/Disciplina]

    ## 1. [Primeiro Tópico Principal]

    [Conteúdo em prosa acadêmica fluida, preservando 100% das informações]

    **SÍNTESE DO TÓPICO:**

    | Conceito/Instituto | Definição | Fundamento Legal | Observações |
    |-------------------|-----------|-----------------|-------------|
    | [Nome] | [Descrição] | [Art. X, Lei Y] | [Dicas/Exceções] |

    ## 2. [Segundo Tópico Principal]
    ...
    ```

    ---

    # RESTRIÇÕES OBRIGATÓRIAS

    1. **Tamanho:** Output deve ter entre **90% e 120%** do tamanho do input
    2. **Linguagem:** APENAS português brasileiro formal
    3. **Formato:** APENAS Markdown válido
    4. **Numeração:** Tópicos DEVEM ser numerados sequencialmente (1., 1.1, 1.2, 2., etc.)
    5. **Tabelas:** OBRIGATÓRIAS ao final de CADA tópico principal

    ## PROIBIDO:
    - Comentários meta ("Continuação...", "Parte X...", "Chunk...")
    - Emojis (exceto em alertas de dica de prova)
    - Links externos inventados
    - Conteúdo não presente na transcrição original
    - Reorganizar a ordem cronológica da aula

    ---

    # AUTO-VERIFICAÇÃO (EXECUTE ANTES DE ENVIAR)

    □ Mantive TODOS os artigos de lei com números corretos?
    □ Mantive TODOS os exemplos e casos completos (não resumidos)?
    □ Mantive TODAS as dicas de prova e observações do professor?
    □ Criei tabelas de síntese ao final de CADA tópico principal?
    □ Usei terceira pessoa impessoal em TODO o texto?
    □ Output tem pelo menos 90% do tamanho do input?
    □ Numeração dos tópicos está sequencial e correta?

    **Se QUALQUER item for NÃO → REVISE antes de enviar.**

    ---

    # FORMATO DE RESPOSTA

    Retorne **APENAS** a apostila em Markdown, sem meta-comentários ou explicações sobre o processo.
    """

    async def _summarize_chunk_async(self, text_chunk, idx, total):
        """Resume um bloco de texto formatado (Async)"""
        try:
            response = await self.async_client.chat.completions.create(
                model=self.model,  # Gemini 2.5 Flash via OpenRouter
                messages=[
                    {"role": "system", "content": self.SUMMARY_PROMPT},
                    {"role": "user", "content": f"Resuma a parte {idx+1}/{total}:\n\n{text_chunk}"}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"{Fore.RED}⚠️ Erro ao resumir chunk {idx}: {e}")
            return text_chunk # Retorna original em caso de erro para não perder

    def generate_summary_version(self, full_formatted_text):
        """Gera a versão sintética a partir da versão completa"""
        print(f"\n{Fore.CYAN}📉 Gerando versão RESUMIDA (Sintética)...")
        
        # Divide o texto completo em pedaços maiores (o resumo aceita mais contexto)
        # Usamos 25k caracteres pois o texto já está limpo
        chunks = self._smart_chunk(full_formatted_text, 25000)
        
        async def process_summary():
            results = []
            for i, chunk in enumerate(chunks):
                print(f"   Resumindo parte {i+1}/{len(chunks)}...")
                res = await self._summarize_chunk_async(chunk, i, len(chunks))
                results.append(res)
            return "\n\n---\n\n".join(results)
            
        summary_text = asyncio.run(process_summary())
        
        # Adiciona um cabeçalho distintivo
        header = "# RESUMO ESQUEMÁTICO (Revisão Rápida)\n> Baseado na aula completa.\n\n"
        return header + summary_text

    def _smart_chunk(self, text, max_size):
        """Divide texto respeitando parágrafos (ORIGINAL)"""
        if len(text) <= max_size:
            return [text]
        
        paragraphs = text.split('\n\n')
        chunks = []
        current = ""
        
        for para in paragraphs:
            if len(current) + len(para) + 2 <= max_size:
                current += para + "\n\n"
            else:
                if current:
                    chunks.append(current.strip())
                
                if len(para) > max_size:
                    sentences = para.replace('? ', '?|').replace('! ', '!|').replace('. ', '.|').split('|')
                    temp = ""
                    for s in sentences:
                        if len(temp) + len(s) <= max_size:
                            temp += s
                        else:
                            chunks.append(temp)
                            temp = s
                    if temp:
                        current = temp + "\n\n"
                    else:
                        current = ""
                else:
                    current = para + "\n\n"
        
        if current:
            chunks.append(current.strip())
            
        return chunks

    def _smart_chunk_overlapping(self, text, max_size=8000, overlap=1000):
        """
        Divide texto com SOBREPOSIÇÃO para evitar perda nas bordas
        """
        if len(text) <= max_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_size
            
            # Se não for o último chunk, encontra quebra natural
            if end < len(text):
                # Procura por quebra de parágrafo nos últimos 1000 chars
                search_zone = text[end-1000:end]
                last_break = search_zone.rfind('\n\n')
                
                if last_break != -1:
                    end = end - 1000 + last_break
            
            chunk = text[start:end]
            chunks.append(chunk)
            
            # Próximo chunk começa com sobreposição (exceto no final)
            if end < len(text):
                start = end - overlap  # 500 chars de overlap
            else:
                break
        
        return chunks

    def _validate_preservation_heuristics(self, original_text, formatted_text):
        """
        Validação automática baseada em contagem de elementos críticos
        """
        import re
        
        print(f"\n{Fore.CYAN}🔍 Validação Heurística de Preservação...")
        
        issues = []
        
        # 1. Contagem de Referências Legais
        original_laws = set(re.findall(
            r'(?:Lei|lei)\s+n?º?\s*\d+[\./]\d+|Art\.?\s*\d+|Súmula\s+\d+',
            original_text,
            re.IGNORECASE
        ))
        formatted_laws = set(re.findall(
            r'(?:Lei|lei)\s+n?º?\s*\d+[\./]\d+|Art\.?\s*\d+|Súmula\s+\d+',
            formatted_text,
            re.IGNORECASE
        ))
        
        missing_laws = original_laws - formatted_laws
        if missing_laws:
            issues.append(f"❌ {len(missing_laws)} referências legais omitidas: {list(missing_laws)[:3]}")
        else:
            print(f"{Fore.GREEN}   ✅ Leis/Artigos: {len(formatted_laws)}/{len(original_laws)} preservados")
        
        # 2. Contagem de Autores
        original_authors = set(re.findall(
            r'(?:Prof|professor|Doutor|Dr\.?)\s+([A-ZÇÁÀÃÉÊ][a-zçáàãéê]+(?:\s+[A-ZÇÁÀÃÉÊ][a-zçáàãéê]+)+)',
            original_text
        ))
        formatted_authors = set(re.findall(
            r'(?:Prof|professor|Doutor|Dr\.?)\s+([A-ZÇÁÀÃÉÊ][a-zçáàãéê]+(?:\s+[A-ZÇÁÀÃÉÊ][a-zçáàãéê]+)+)',
            formatted_text
        ))
        
        missing_authors = original_authors - formatted_authors
        if missing_authors:
            issues.append(f"❌ {len(missing_authors)} autores omitidos: {list(missing_authors)[:3]}")
        else:
            print(f"{Fore.GREEN}   ✅ Autores: {len(formatted_authors)}/{len(original_authors)} preservados")
        
        # 3. Contagem de Palavras-Chave de Dicas de Prova
        tip_keywords = [
            r'cai muito', r'atenção', r'pegadinha', r'cuidado',
            r'importante', r'não confund', r'diferença entre'
        ]
        
        original_tips = sum([len(re.findall(kw, original_text, re.IGNORECASE)) for kw in tip_keywords])
        formatted_tips = sum([len(re.findall(kw, formatted_text, re.IGNORECASE)) for kw in tip_keywords])
        
        if formatted_tips < original_tips * 0.8:  # Tolerância de 20%
            issues.append(f"⚠️ Possível perda de dicas: {formatted_tips}/{original_tips} ocorrências")
        else:
            print(f"{Fore.GREEN}   ✅ Dicas de prova: {formatted_tips}/{original_tips} preservadas")
        
        # 4. Análise de Comprimento Relativo
        original_words = len(original_text.split())
        formatted_words = len(formatted_text.split())
        ratio = formatted_words / original_words if original_words > 0 else 0
        
        if ratio < 0.70:  # Se o texto formatado for <70% do original
            issues.append(f"⚠️ ALERTA: Texto formatado é {ratio:.1%} do original (esperado >70%)")
        else:
            print(f"{Fore.GREEN}   ✅ Comprimento relativo: {ratio:.1%} (adequado)")
        
        # 5. Detecção de Frases Cortadas
        truncation_patterns = [
            r'\.\.\.$',  # Texto terminando em ...
            r'\w+\s*$(?![.!?])',  # Palavra sem pontuação final
            r'exemplo:|por exemplo:(?!\s*\w)',  # "exemplo:" sem continuação
        ]
        
        for pattern in truncation_patterns:
            if re.search(pattern, formatted_text, re.MULTILINE):
                issues.append(f"⚠️ Possível truncamento detectado (padrão: {pattern})")
        
        # RESULTADO
        if issues:
            print(f"\n{Fore.RED}━━━ PROBLEMAS DETECTADOS ━━━")
            for issue in issues:
                print(f"{Fore.YELLOW}   {issue}")
            print(f"{Fore.RED}━━━━━━━━━━━━━━━━━━━━━━━━━━━")
            return False, issues
        else:
            print(f"\n{Fore.GREEN}✅ Validação heurística: APROVADA (nenhuma omissão detectada)")
            return True, []

    def _generate_audit_report(self, video_name, heuristic_issues, llm_issues):
        """Gera relatório de auditoria da formatação"""
        
        report_path = f"audit_{video_name.replace(' ', '_')}.md"
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(f"# Relatório de Auditoria - {video_name}\n\n")
            f.write(f"**Data:** {time.strftime('%d/%m/%Y %H:%M')}\n\n")
            
            f.write("## Validação Heurística\n")
            if not heuristic_issues:
                f.write("✅ Nenhuma omissão detectada\n\n")
            else:
                for issue in heuristic_issues:
                    f.write(f"- {issue}\n")
                f.write("\n")
            
            f.write("## Validação LLM (3 janelas)\n")
            if not llm_issues:
                f.write("✅ Nenhuma omissão detectada\n\n")
            else:
                for issue in llm_issues:
                    f.write(f"- {issue}\n")
                f.write("\n")
            
            if not heuristic_issues and not llm_issues:
                f.write("## ✅ CONCLUSÃO: Apostila APROVADA (preservação completa)\n")
            else:
                f.write("## ⚠️ CONCLUSÃO: Revisar seções indicadas acima\n")
        
        print(f"{Fore.CYAN}📄 Relatório de auditoria salvo: {report_path}")
    
    def _segment_raw_transcription(self, raw_text):
        """Segmenta por SPEAKER (ORIGINAL - sem mudanças)"""
        lines = raw_text.split('\n')
        speaker_pattern = re.compile(r'^SPEAKER \d+$')
        
        segments = []
        current_speaker = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if speaker_pattern.match(line):
                if current_speaker:
                    segments.append({
                        'speaker': current_speaker,
                        'content': "\n".join(current_content)
                    })
                current_speaker = line
                current_content = []
            else:
                if current_speaker:
                    current_content.append(line)
        
        if current_speaker:
            segments.append({
                'speaker': current_speaker,
                'content': "\n".join(current_content)
            })
        
        # Mesmas 3 etapas de merge/filtering do original
        merged_same_speaker = []
        if segments:
            current_seg = segments[0]
            for next_seg in segments[1:]:
                if next_seg['speaker'] == current_seg['speaker']:
                    current_seg['content'] += "\n" + next_seg['content']
                else:
                    merged_same_speaker.append(current_seg)
                    current_seg = next_seg
            merged_same_speaker.append(current_seg)
        
        filtered_segments = []
        if merged_same_speaker:
            current_seg = merged_same_speaker[0]
            for next_seg in merged_same_speaker[1:]:
                if len(next_seg['content']) < 100:
                    current_seg['content'] += "\n" + next_seg['content']
                else:
                    filtered_segments.append(current_seg)
                    current_seg = next_seg
            filtered_segments.append(current_seg)
        
        final_segments = []
        if filtered_segments:
            current_seg = filtered_segments[0]
            for next_seg in filtered_segments[1:]:
                if next_seg['speaker'] == current_seg['speaker']:
                    current_seg['content'] += "\n" + next_seg['content']
                else:
                    final_segments.append(current_seg)
                    current_seg = next_seg
            final_segments.append(current_seg)
        
        return final_segments
    
    # VERSÕES ASYNC DOS MÉTODOS CRÍTICOS (mantendo o SYSTEM_PROMPT original)
    # Retry para robustez contra falhas de API
    @retry(
        retry=retry_if_exception_type((APIError, Exception)), 
        stop=stop_after_attempt(3), 
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    async def _format_chunk_async(self, chunk_text, chunk_idx, system_prompt):
        """VERSÃO ASYNC de _format_chunk_content_only - PRESERVA FORMATAÇÃO"""
        word_count = len(chunk_text.split())
        
        user_content = f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ MODO LITERAL ABSOLUTO - PRESERVAÇÃO DE CONTEÚDO ⚠️
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TAREFA: Formatar o texto abaixo mantendo ~{word_count} palavras originais.

📏 REQUISITOS:
1. Mantenha TODOS os exemplos, datas, nomes e críticas.
2. Mantenha TODAS as dicas de prova ("Atenção", "Cai muito").
3. NÃO resuma histórias ou anedotas.
4. Apenas limpe vícios ("né", "tipo") e formate leis.

[TEXTO ORIGINAL - CHUNK {chunk_idx + 1}]
{chunk_text}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
        
        try:
            response = await self.async_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},  # PROMPT COMPLETO ORIGINAL
                    {"role": "user", "content": user_content}
                ],

                presence_penalty=0.0,  # Sem penalidade de repetição
                frequency_penalty=0.0   # Permite repetir termos técnicos
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"{Fore.RED}⚠️ Erro ao formatar chunk {chunk_idx}: {e}")
            return chunk_text
    
    async def _identify_speaker_async(self, content, professors_info, speaker_label):
        """VERSÃO ASYNC de _identify_speaker_context com cache"""
        
        # Cache: Se já identificamos este speaker, reutiliza
        if speaker_label in self.speaker_cache:
            return self.speaker_cache[speaker_label]
        
        transition = self._detect_discipline_transition(content)
        
        prompt = f"""
        Analise o início do texto abaixo e a lista de professores extraída da introdução.
        Identifique quem é o provável professor falando e qual a disciplina.
        
        IMPORTANTE:
        1. Se o texto indicar um professor diferente da lista (ex: gênero diferente, nome citado, autoapresentação), PREFIRA a informação do texto.
        2. Se o 'Falante (Label)' for "SPEAKER 2" (ou maior) e a lista tiver apenas um professor, ASSUMA QUE É UM NOVO PROFESSOR.
        3. Se não houver nome claro, mas houver indicação de gênero (ex: "sou a professora"), use "Professora Desconhecida" ou tente inferir o nome pelo contexto.
        4. Identifique a disciplina pelo conteúdo TÉCNICO-JURÍDICO.
           - IGNORE avisos sobre editais, datas de prova, conselhos de estudo ou "conversa fiada".
           - Foque no tema substantivo (ex: se falar de "Poder Constituinte", a disciplina é "Direito Constitucional").
           - SEJA ESPECÍFICO (ex: "Direito Constitucional", "Direito Administrativo", "Processo Civil").

        Falante (Label): {speaker_label if speaker_label else "Desconhecido"}
        
        Lista de Professores (Contexto):
        {professors_info}
        
        Texto (Início do segmento - Amostra Ampliada):
        {content[:5000]}...
        
        Retorne APENAS um JSON no seguinte formato:
        {{
            "nome": "Nome do Professor",
            "disciplina": "Disciplina Específica"
        }}
        """
        
        try:
            response = await self.async_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "Você é um assistente que identifica palestrantes."},
                    {"role": "user", "content": prompt}
                ]
            )
            content_json = response.choices[0].message.content
            
            # --- HEURÍSTICA DE CORREÇÃO ---
            # Se o LLM insistir no nome do primeiro professor para um SPEAKER diferente, forçamos "Professor(a) Convidado(a)"
            try:
                import json
                data = json.loads(content_json)
                detected_name = data.get("nome", "")
                
                if speaker_label and "SPEAKER" in speaker_label and not speaker_label.endswith("1"):
                    # Verifica se o nome detectado é o mesmo do único professor conhecido
                    try:
                        prof_ctx = json.loads(professors_info)
                        prof_list = prof_ctx.get("professores", [])
                        if len(prof_list) == 1:
                            first_prof_name = prof_list[0].get("nome", "")
                            if detected_name == first_prof_name:
                                print(f"{Fore.YELLOW}⚠️ Detectado mesmo nome ({detected_name}) para {speaker_label}. Aplicando override.")
                                data["nome"] = "Professor(a) Convidado(a)"
                                content_json = json.dumps(data)
                    except Exception as e:
                        print(f"Erro ao aplicar heurística de professor: {e}")
            except:
                pass
            # ------------------------------
            
            # Armazena no cache para reutilizar\n            self.speaker_cache[speaker_label] = content_json\n            return content_json
        except Exception as e:
            print(f"Erro ao identificar speaker: {e}")
            return '{"nome": "Professor", "disciplina": "Disciplina"}'
    
    async def _generate_header_async(self, formatted_content, professor_context_json):
        """VERSÃO ASYNC de _generate_discipline_header"""
        try:
            import json
            try:
                prof_ctx = json.loads(professor_context_json)
                prof_name = prof_ctx.get("nome", "Professor")
                discipline = prof_ctx.get("disciplina", "Disciplina")
            except:
                prof_name = "Professor"
                discipline = "Disciplina"
            
            prompt = f"""
            Gere APENAS o título Markdown para esta seção.
            Professor: {prof_name}
            Disciplina: {discipline}
            
            Conteúdo:
            {formatted_content[:1000]}...
            
            FORMATO DE SAÍDA:
            # Prof. {prof_name} - {discipline}
            """
            
            response = await self.async_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "Você é um assistente que gera cabeçalhos de estudo."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"Erro ao gerar header: {e}")
            return f"# {prof_name} - {discipline}\n"
    
    def _extract_professors_context(self, full_text):
        """
        Extrai a lista de professores e disciplinas analisando todo o texto.
        Procura por padrões de introdução em todo o arquivo.
        """
        print(f"   🕵️  Extraindo contexto de professores (Scan Completo)...")
        
        # 1. Pega o início (padrão)
        intro_context = full_text[:5000]
        
        # 2. Busca por padrões de introdução no restante do texto
        keywords = [
            "meu nome é", "sou o professor", "sou a professora", 
            "aqui é o professor", "aqui é a professora", 
            "comigo", "recebam", "passo a palavra", "agora com",
            "boa tarde", "bom dia", "boa noite", "olá pessoal", "olá a todos",
            "começar a aula", "iniciar a aula", "mudando de assunto",
            "próximo professor", "próxima professora"
        ]
        
        found_contexts = []
        lower_text = full_text.lower()
        
        for keyword in keywords:
            start_idx = 0
            while True:
                idx = lower_text.find(keyword, start_idx)
                if idx == -1:
                    break
                
                # Se encontrou, pega um contexto ao redor (500 chars antes e depois)
                start_ctx = max(0, idx - 500)
                end_ctx = min(len(full_text), idx + 500)
                found_contexts.append(full_text[start_ctx:end_ctx])
                
                start_idx = idx + len(keyword)
        
        # Combina tudo
        combined_context = intro_context + "\n\n... [TRECHOS RELEVANTES ENCONTRADOS] ...\n\n" + "\n\n".join(found_contexts)
        
        # Limita tamanho total para não estourar contexto
        if len(combined_context) > 50000:
            combined_context = combined_context[:50000]

        system_prompt = """
        Você é um assistente especializado em extrair informações de introduções de aulas.
        Analise o texto fornecido (que contém o início da aula e trechos onde professores podem ter se apresentado) e extraia a lista de TODOS os professores e suas respectivas disciplinas.
        
        Retorne APENAS um JSON no seguinte formato:
        {
            "professores": [
                {"nome": "Nome", "disciplina": "Disciplina", "ordem": "primeiro/segundo/último"}
            ]
        }
        Se não encontrar informações, retorne {"professores": []}.
        """
        
        try:
            response = self.llm_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": combined_context}
                ],

                presence_penalty=0.0,
                frequency_penalty=0.0
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"Erro ao extrair professores: {e}")
            return "{'professores': []}"
    
    def _check_transition(self, text_window):
        """MANTÉM SÍNCRONO - chamado internamente"""
        prompt = """
        Analise o texto abaixo e verifique se há uma transição explícita de disciplina ou professor.
        Exemplos: "Agora vamos falar de Direito Penal", "Passando a palavra para o professor X".
        
        Texto:
        {text}
        
        Retorne APENAS a transição encontrada ou "None".
        """
        try:
            response = self.llm_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt.format(text=text_window)}]
            )
            content = response.choices[0].message.content.strip()
            return content if content != "None" else None
        except:
            return None
    
    def _detect_discipline_transition(self, chunk_text):
        """MANTÉM SÍNCRONO"""
        windows = [
            chunk_text[:10000],
            chunk_text[len(chunk_text)//2-5000:len(chunk_text)//2+5000],
            chunk_text[-10000:]
        ]
        
        for window in windows:
            if len(window) < 100: continue
            result = self._check_transition(window)
            if result:
                return result
        return None
    
    async def _extract_critical_entities(self, text):
        """Extrai entidades críticas que DEVEM constar no output"""
        prompt = """
        Analise o texto jurídico abaixo e liste TODAS as entidades críticas que NÃO podem ser omitidas.
        
        Categorias obrigatórias:
        1. LEGISLAÇÃO: Leis, Artigos, Parágrafos, Incisos (ex: "Art. 37 da CF", "Lei 8.666")
        2. JURISPRUDÊNCIA: Súmulas, Temas de Repercussão Geral, Nomes de Casos (ex: "Súmula Vinculante 13", "Tema 105", "Caso Sindispreve")
        3. PESSOAS/INSTITUIÇÕES: Nomes de autores, políticos, órgãos (ex: "Hely Lopes", "STF", "Fazenda Pública")
        4. CONCEITOS/TERMOS-CHAVE: Termos técnicos específicos definidos ou explicados.
        4. CONCEITOS/TERMOS-CHAVE: Termos técnicos específicos definidos ou explicados.
        5. EXEMPLOS/HISTÓRIAS: Resumo de 3-5 palavras de exemplos, casos hipotéticos ou histórias pessoais.
        
        Texto:
        {text}
        
        Retorne APENAS um JSON:
        {{
            "legislacao": ["item1", "item2"],
            "jurisprudencia": ["item1", "item2"],
            "entidades": ["item1", "item2"],
            "exemplos": ["item1", "item2"]
        }}
        """
        try:
            response = await self.async_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt.format(text=text)}],
                response_format={"type": "json_object"}
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            print(f"{Fore.RED}⚠️ Erro ao extrair entidades: {e}")
            return {}

    def _verify_entities_presence(self, formatted_text, entities):
        """Verifica se as entidades extraídas estão presentes no texto formatado"""
        missing = []
        formatted_lower = formatted_text.lower()
        
        # Verifica Legislação e Jurisprudência (Prioridade Alta)
        for category in ["legislacao", "jurisprudencia", "entidades", "exemplos"]:
            if category in entities:
                for item in entities[category]:
                    # Busca flexível (case insensitive e parcial)
                    item_clean = item.lower().replace("art.", "").replace("lei", "").strip()
                    if len(item_clean) < 4: continue # Pula itens muito curtos
                    
                    # Verifica se o item (ou parte significativa dele) está no texto
                    if item.lower() not in formatted_lower:
                        # Tenta busca mais flexível (palavras-chave)
                        keywords = [w for w in item.lower().split() if len(w) > 3]
                        matches = sum(1 for k in keywords if k in formatted_lower)
                        
                        # Threshold de 60% para flexibilidade em exemplos (narrativas variam mais)
                        threshold = 0.6 if category == "exemplos" else 0.7
                        
                        if matches < len(keywords) * threshold:
                            missing.append(f"[{category.upper()}] {item}")
                            
        return missing

    async def _repair_omissions(self, formatted_text, missing_items, original_text):
        """Repara cirurgicamente o texto inserindo itens omitidos"""
        print(f"{Fore.MAGENTA}   🔧 Iniciando REPARO CIRÚRGICO para {len(missing_items)} itens...")
        
        prompt = f"""
        CRITICAL RESTORATION TASK.
        The following items were OMITTED from the formatted text but are MANDATORY:
        {json.dumps(missing_items, ensure_ascii=False)}
        
        Source Context (where these items appear in original text):
        {original_text}
        
        Current Formatted Text:
        {formatted_text}
        
        INSTRUCTION:
        Rewrite the Formatted Text to include ALL the missing items naturally.
        - Maintain the current structure and flow.
        - Just weave the missing details back in where they belong.
        - DO NOT summarize. DO NOT remove anything else.
        - RETURN ONLY THE REPAIRED TEXT.
        """
        
        try:
            response = await self.async_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert legal text editor. Your goal is to restore missing information without altering the rest of the text."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro no reparo: {e}")
            return formatted_text

    async def _format_with_validation(self, chunk_text, idx, system_prompt, max_retries=0):
        """Formatação RÁPIDA sem validação (1 chamada LLM por chunk)"""
        try:
            response = await self.async_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"FORMAT THIS TEXT:\n\n{chunk_text}"}
                ]
            )
            formatted = response.choices[0].message.content
            print(f"   ✅ Chunk {idx+1} processado")
            return formatted
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro: {e}")
            return chunk_text

    
    def _merge_chunks_with_deduplication(self, formatted_chunks):
        """
        Remove duplicações causadas por overlap entre chunks formatados.
        Usa difflib para encontrar sobreposição exata de subsequências.
        
        Args:
            formatted_chunks (list): Lista de chunks já formatados pelo LLM
            
        Returns:
            str: Texto unificado sem redundâncias
        """
        import difflib
        
        if not formatted_chunks:
            return ""
        
        if len(formatted_chunks) == 1:
            return formatted_chunks[0]
        
        print(f"{Fore.CYAN}   🔗 Deduplicando overlap entre {len(formatted_chunks)} chunks...")
        
        merged = formatted_chunks[0]
        duplications_found = 0
        total_chars_removed = 0
        
        for i in range(1, len(formatted_chunks)):
            current_chunk = formatted_chunks[i]
            
            # Pega os últimos 2000 chars do anterior e os primeiros 2000 do atual
            lookback = 2000
            tail = merged[-lookback:] if len(merged) > lookback else merged
            head = current_chunk[:lookback] if len(current_chunk) > lookback else current_chunk
            
            # Encontra a maior substring comum sequencial (LCS)
            matcher = difflib.SequenceMatcher(None, tail, head)
            match = matcher.find_longest_match(0, len(tail), 0, len(head))
            
            # Se a sobreposição for significativa (> 100 chars), faz o merge cirúrgico
            if match.size > 100:
                # ANÁLISE DE PREFIXO (Evitar cortes indevidos)
                # Verifica se há conteúdo no chunk atual ANTES do match que foi ignorado
                if match.b > 0:
                    head_prefix = current_chunk[:match.b]
                    tail_prefix = tail[:match.a]
                    
                    # Se o prefixo ignorado for muito diferente do final do anterior,
                    # provavelmente é conteúdo novo que foi omitido no anterior (ou alucinação).
                    # Na dúvida, PRESERVA (melhor duplicar que omitir).
                    similarity = difflib.SequenceMatcher(None, tail_prefix, head_prefix).ratio()
                    
                    if similarity < 0.6: # 60% de similaridade
                        print(f"      ⚠️ Recuperando {len(head_prefix)} chars antes do overlap (Similaridade: {similarity:.2f})")
                        merged += head_prefix
                    else:
                        print(f"      ℹ️ Ignorando prefixo redundante ({len(head_prefix)} chars, Sim: {similarity:.2f})")

                # Corta apenas o que já estava no chunk anterior (parte coincidente)
                unique_part = current_chunk[match.b + match.size:]
                merged += unique_part
                duplications_found += 1
                total_chars_removed += match.size
                print(f"      Chunk {i+1}: Overlap exato de {match.size} chars removido")
            else:
                # Fallback: Concatena com quebra de parágrafo se não achar overlap claro
                merged += "\n\n" + current_chunk
        
        if duplications_found > 0:
            print(f"{Fore.GREEN}   ✅ {duplications_found} overlaps removidos ({total_chars_removed} chars)")
        else:
            print(f"{Fore.GREEN}   ✅ Nenhum overlap detectado")
        
        return merged

    # NOVO MÉTODO: Processa 1 segmento inteiro em PARALELO
    async def _process_segment_parallel(self, segment, professors_info, idx, system_prompt):
        """
        Processa 1 segmento completo (chunks em paralelo) - PRESERVA FORMATAÇÃO MARKDOWN
        """
        speaker = segment['speaker']
        content = segment['content']
        
        print(f"\n{Fore.YELLOW}▶ Segmento {idx+1} ({speaker})...")
        
        # Chunks otimizados (8k chars = ~2k tokens) com overlap para evitar perda
        chunks = self._smart_chunk_overlapping(content, max_size=8000, overlap=1500)
        print(f"   {len(chunks)} chunks de ~8k chars (com 1.5k overlap)")
        
        # 1. Identifica contexto (paralelo)
        context_task = self._identify_speaker_async(content[:5000], professors_info, speaker)
        
        # 2. Formata chunks SEQUENCIALMENTE com VALIDAÇÃO
        # Cada chunk é verificado quanto à preservação de conteúdo (>80%)
        formatted_parts = []
        import gc
        
        for j, chunk in enumerate(chunks):
            print(f"   Processando chunk {j+1}/{len(chunks)}...")
            part = await self._format_with_validation(chunk, j, system_prompt, max_retries=0)
            formatted_parts.append(part)
            
            # Limpeza forçada de memória após cada chunk
            gc.collect()
        
        # Aguarda contexto (se ainda não terminou)
        prof_context = await context_task
        
        # Pass 1: Formatar Conteúdo (Chunking com overlap para preservar contexto)
        # 🎯 TAMANHO REDUZIDO: 8k chars (~2k tokens) com 1k overlap
        # Isso evita perda de contexto nas bordas e previne compressão do output
        chunks = self._smart_chunk_overlapping(content, max_size=8000)
        formatted_segment_parts = []
        

        
        print(f"   {len(chunks)} chunks de ~8k chars (com 1k overlap)")
        
        # 1. Identifica contexto (paralelo)
        context_task = self._identify_speaker_async(content[:5000], professors_info, speaker)
        
        # 2. Formata TODOS os chunks em paralelo (MANTÉM O SYSTEM_PROMPT COMPLETO)
        chunk_tasks = [
            self._format_chunk_async(chunk, j, system_prompt)
            for j, chunk in enumerate(chunks)
        ]
        
        # Aguarda tudo em paralelo
        prof_context, *formatted_parts = await asyncio.gather(context_task, *chunk_tasks)
        
        # 3. Gera header
        # Concatena com deduplicação de overlaps
        full_content = self._merge_chunks_with_deduplication(formatted_parts)
        header = await self._generate_header_async(full_content[:10000], prof_context)
        
        return f"{header}\n\n{full_content}\n\n---\n\n", prof_context
    
    def format_transcription(self, transcript_text, video_name, output_folder):
        """
        VERSÃO OTIMIZADA: Processa segmentos em PARALELO mantendo formatação original
        """
        print(f"{Fore.MAGENTA}🧠 Formatando com GPT-5-mini (Async Otimizado)...")
        
        # 1. Extrai contexto da intro (1 chamada síncrona)
        # USANDO NOVO MÉTODO DE SCAN COMPLETO
        professors_info = self._extract_professors_context(transcript_text)
        print(f"   📋 Contexto: {professors_info}")
        
        # 2. Segmenta por speaker
        segments = []
        if any(re.match(r'^SPEAKER \d+$', line.strip()) for line in transcript_text.split('\n')):
            print(f"   🗣️  Usando segmentos de falante...")
            segments = self._segment_raw_transcription(transcript_text)
        else:
            print(f"   ⚠️  Sem tags de falante. Processando como bloco único.")
            segments = [{'speaker': 'SPEAKER 1', 'content': transcript_text}]
        
        # 3. Processa segmentos em PARALELO (máximo 5 concorrentes)
        async def process_all():
            semaphore = asyncio.Semaphore(5)  # Limite de rate
            
            async def process_with_limit(seg, i):
                async with semaphore:
                    return await self._process_segment_parallel(seg, professors_info, i, self.SYSTEM_PROMPT_FORMAT)
            
            results = await asyncio.gather(*[
                process_with_limit(seg, i) for i, seg in enumerate(segments)
            ])
            
            return results
        
        # Executa tudo
        results = asyncio.run(process_all())
        
        # 4. Monta resultado final
        formatted_chunks = [result[0] for result in results]
        full_formatted = f"# {video_name}\n\n" + "\n\n".join(formatted_chunks)
        
        # 🆕 RENUMERAÇÃO AUTOMÁTICA (garante sequência correta)
        print(f"\n{Fore.CYAN}🔢 Renumerando tópicos sequencialmente...")
        full_formatted = self._renumber_topics(full_formatted)
        
        # VALIDAÇÃO EM DUAS CAMADAS
        print(f"\n{Fore.MAGENTA}🔒 Iniciando validação de preservação...")
        
        # Camada 1: Heurística (rápida)
        h_passed, h_issues = self._validate_preservation_heuristics(transcript_text, full_formatted)
        
        # Camada 2: LLM (precisa, mas lenta)
        full_formatted, llm_issues = self.validate_completeness_enhanced(transcript_text, full_formatted, video_name)
        
        # Gera relatório
        self._generate_audit_report(video_name, h_issues, llm_issues)
        
        return full_formatted
    
    def _renumber_topics(self, markdown_text):
        """
        Renumera TODOS os tópicos e subtópicos sequencialmente,
        garantindo numeração correta independente do LLM.
        """
        import re
        
        lines = markdown_text.split('\n')
        output_lines = []
        
        # Contadores para cada nível
        # counters[0] = ## (nível 2)
        # counters[1] = ### (nível 3)
        # counters[2] = #### (nível 4)
        counters = [0, 0, 0]
        
        # Regex para detectar headings numerados
        # Captura: ## 1.2.3 Título ou ## 1. Título ou ## Título
        heading_pattern = re.compile(r'^(#{2,4})\s*(?:[\d\.]+\s+)?(.+)$')
        
        for line in lines:
            match = heading_pattern.match(line)
            
            if match:
                hashes = match.group(1)  # ##, ###, ou ####
                title = match.group(2).strip()  # Texto do título
                
                level = len(hashes) - 2  # ## = 0, ### = 1, #### = 2
                
                if level > 2:
                    # Ignora níveis muito profundos (##### ou mais)
                    output_lines.append(line)
                    continue
                
                # Incrementa contador do nível atual
                counters[level] += 1
                
                # Reseta contadores dos níveis inferiores
                for i in range(level + 1, 3):
                    counters[i] = 0
                
                # Constrói número sequencial
                if level == 0:  # ##
                    number = f"{counters[0]}"
                elif level == 1:  # ###
                    number = f"{counters[0]}.{counters[1]}"
                elif level == 2:  # ####
                    number = f"{counters[0]}.{counters[1]}.{counters[2]}"
                
                # Reconstrói linha com número correto
                new_line = f"{hashes} {number}. {title}"
                output_lines.append(new_line)
                
            else:
                # Linha normal (não é heading)
                output_lines.append(line)
        
        renumbered_text = '\n'.join(output_lines)
        
        # Conta quantos tópicos foram renumerados
        topic_count = counters[0]
        print(f"{Fore.GREEN}   ✅ {topic_count} tópicos principais renumerados sequencialmente")
        
        return renumbered_text
    
    def validate_completeness_enhanced(self, raw_transcript, formatted_text, video_name):
        """
        Validação LLM-as-Judge com amostragem estratégica
        """
        print(f"{Fore.YELLOW}🔍 Validação LLM (Amostragem Múltipla)...")
        
        # Em vez de apenas os primeiros 50k, valida 3 JANELAS diferentes
        windows = [
            ("INÍCIO", raw_transcript[:50000], formatted_text[:50000]),
            ("MEIO", raw_transcript[len(raw_transcript)//2-25000:len(raw_transcript)//2+25000],
                      formatted_text[len(formatted_text)//2-25000:len(formatted_text)//2+25000]),
            ("FIM", raw_transcript[-50000:], formatted_text[-50000:])
        ]
        
        validation_prompt = """# TAREFA DE VALIDAÇÃO
Você receberá:
1. TRANSCRIÇÃO BRUTA (amostra)
2. APOSTILA FORMATADA (amostra correspondente)

Identifique conteúdo OMITIDO na apostila.

## O QUE PROCURAR (PRIORIDADE ALTA):
✅ Exemplos práticos, casos, histórias
✅ Críticas do professor
✅ Especulações, "apostas"
✅ Argumentos estratégicos
✅ Contexto histórico/político
✅ Dicas de prova
✅ Nuances técnicas

## FORMATO DE RESPOSTA:
Se TUDO preservado:
```
VALIDAÇÃO: COMPLETA
```

Se houver omissões:
```
OMISSÕES DETECTADAS:
1. [Tipo]: [Descrição concisa]
2. [Tipo]: [Descrição]
```

Seja RIGOROSO mas CONCISO (máx 5 itens)."""
        
        all_issues = []
        
        for window_name, raw_window, formatted_window in windows:
            print(f"   Validando janela: {window_name}...")
            
            try:
                response = self.llm_client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": validation_prompt},
                        {"role": "user", "content": f"RAW:\n{raw_window}\n\nFORMATTED:\n{formatted_window}"}
                    ],
                    timeout=120
                )
                
                result = response.choices[0].message.content
                
                if "VALIDAÇÃO: COMPLETA" not in result:
                    all_issues.append(f"[{window_name}] {result}")
            
            except Exception as e:
                print(f"{Fore.RED}   ❌ Erro na validação {window_name}: {e}")
        
        # RESULTADO CONSOLIDADO
        if not all_issues:
            print(f"{Fore.GREEN}   ✅ Todas as janelas validadas: Nenhuma omissão detectada")
        else:
            print(f"\n{Fore.YELLOW}━━━ OMISSÕES DETECTADAS (LLM Validation) ━━━")
            for issue in all_issues:
                print(f"{Fore.YELLOW}   {issue}")
            print(f"{Fore.YELLOW}━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
            print(f"{Fore.CYAN}   💡 Revise manualmente o arquivo .md")
        
        return formatted_text, all_issues
    
    def save_as_word(self, formatted_text, video_name, output_folder):
        """Salva a apostila formatada em Word (.docx) com formatação profissional"""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        print(f"{Fore.CYAN}📄 Gerando documento Word profissional...")
        
        doc = Document()
        
        # Configurações do documento
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        
        # Título Principal
        title = doc.add_heading(video_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Data
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')}")
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espaço
        
        # Processa o texto linha por linha
        lines = formatted_text.split('\n')
        i = 0
        in_table = False
        table_rows = []
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Linha vazia ou código markdown - PULA
            if not line or line.startswith('```'):
                i += 1
                continue
            
            # Detecta tabelas Markdown
            if '|' in line and not in_table:
                in_table = True
                table_rows = []
            
            if in_table:
                if '|' in line and not line.startswith('|--'):  # Ignora separador
                    table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
                
                # Fim da tabela
                if '|' not in line or i == len(lines) - 1:
                    if len(table_rows) > 0:
                        self._add_table_to_doc(doc, table_rows)
                    in_table = False
                    table_rows = []
                    if '|' not in line:
                        continue
                i += 1
                continue
            
            # Seções especiais (Summary, Key Takeaways, Action Items)
            if line.startswith('## Summary') or line.startswith('## Key Takeaways') or line.startswith('## Action Items'):
                heading = doc.add_heading(line.replace('## ', ''), level=2)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 102, 204)
                i += 1
                continue
            
            # Headings
            if line.startswith('#### '):
                doc.add_heading(line.replace('#### ', ''), level=4)
                i += 1
                continue
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
                i += 1
                continue
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
                i += 1
                continue
            elif line.startswith('# ') and line != f"# {video_name}":  # Evita duplicar título
                doc.add_heading(line.replace('# ', ''), level=1)
                i += 1
                continue
            
            # Separador horizontal
            if line.startswith('---'):
                p = doc.add_paragraph()
                p.add_run('_' * 80).font.color.rgb = RGBColor(192, 192, 192)
                i += 1
                continue
            
            # Blockquote
            if line.startswith('>'):
                p = doc.add_paragraph(line.replace('> ', ''), style='Quote')
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                for run in p.runs:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(64, 64, 64)
                i += 1
                continue
            
            # Listas com bullets
            if line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(6)
                self._format_inline_markdown(p, line[2:])
                i += 1
                continue
            
            # Listas numeradas
            if len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
                text = line.split('. ', 1)[1] if '. ' in line else line.split(') ', 1)[1]
                p = doc.add_paragraph(text, style='List Number')
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(6)
                self._format_inline_markdown(p, text)
                i += 1
                continue
            
            # Texto normal (Parágrafo de conteúdo)
            p = doc.add_paragraph()
            
            # Formatação solicitada:
            # - Espaçamento 1.5
            # - 6pt antes e depois
            # - Recuo na primeira linha (1.25 cm)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificado fica melhor em apostila
            
            self._format_inline_markdown(p, line)
            i += 1
        
        # Salva o documento
        docx_file = os.path.join(output_folder, f"{video_name}_APOSTILA.docx")
        doc.save(docx_file)
        
        return docx_file
    
    def _add_table_to_doc(self, doc, rows):
        """Adiciona tabela formatada ao documento"""
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        if len(rows) < 2:
            return
        
        # Determina o número máximo de colunas
        max_cols = max(len(row) for row in rows)
        
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                
                # Formata cabeçalho
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Cor de fundo do cabeçalho
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '0066CC')
                    cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _format_inline_markdown(self, paragraph, text):
        """Formata markdown inline (negrito, itálico) em um parágrafo existente usando Regex"""
        paragraph.clear()
        
        # Regex que captura **bold**, __bold__, *italic*, _italic_
        # A ordem importa: verifica bold primeiro
        tokens = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)', text)
        
        for token in tokens:
            if not token: continue
            
            run = paragraph.add_run()
            if (token.startswith('**') and token.endswith('**')) or \
               (token.startswith('__') and token.endswith('__')):
                run.bold = True
                run.text = token[2:-2] # Remove marcadores
            elif (token.startswith('*') and token.endswith('*')) or \
                 (token.startswith('_') and token.endswith('_')):
                run.italic = True
                run.text = token[1:-1] # Remove marcadores
            else:
                run.text = token


def process_text_file(file_path):
    """Processa um arquivo de texto específico gerando versões completa e resumida"""
    if not os.path.exists(file_path):
        print(f"{Fore.RED}❌ Arquivo não encontrado: {file_path}")
        return
    
    folder = os.path.dirname(file_path)
    video_name = Path(file_path).stem.replace('_RAW', '').replace('_TRANSCRICAO', '')
    
    print(f"\n{Fore.CYAN}{'='*60}")
    print(f"{Fore.CYAN}📄 Processando: {video_name}")
    print(f"{Fore.CYAN}{'='*60}\n")
    
    try:
        formatter = VomoFormatter()
        
        with open(file_path, "r", encoding="utf-8") as f:
            transcription = f.read()
        
        # =============================================================
        # ETAPA 1: VERSÃO COMPLETA (com cache)
        # =============================================================
        print(f"\n{Fore.BLUE}{'='*60}")
        print(f"{Fore.BLUE}ETAPA 1: VERSÃO COMPLETA")
        print(f"{Fore.BLUE}{'='*60}")
        
        # Verifica cache primeiro
        cache_file = formatter._get_cache_path(file_path, cache_type="FORMATTED")
        cached_data = formatter._load_cache(cache_file)
        
        if cached_data:
            # Usa versão em cache
            print(f"{Fore.GREEN}⚡ Usando versão em cache (processamento instantâneo)")
            formatted_text = cached_data['formatted_text']
            validation_issues = []  # Cache já foi validado anteriormente
        else:
            # Processa do zero
            print(f"{Fore.CYAN}🔄 Processando transcrição (pode levar alguns minutos)...")
            formatted_text = formatter.format_transcription(transcription, video_name, folder)
            
            # Valida completude comparando com transcrição bruta
            formatted_text, validation_issues = formatter.validate_completeness_enhanced(transcription, formatted_text, video_name)
            
            # Salva cache para próxima execução
            formatter._save_cache(cache_file, formatted_text, file_path)
        
        # Salva MD Completo
        output_md = os.path.join(folder, f"{video_name}_APOSTILA_COMPLETA.md")
        with open(output_md, 'w', encoding='utf-8') as f:
            f.write(f"# {video_name}\n\n{formatted_text}")
        print(f"{Fore.GREEN}📝 MD Completo salvo: {output_md}")
        
        # Salva Word Completo
        docx_completo = formatter.save_as_word(formatted_text, f"{video_name}_COMPLETA", folder)
        print(f"{Fore.GREEN}📄 Word Completo salvo: {docx_completo}")

        
        # =============================================================
        # ETAPA 2: VERSÃO RESUMIDA (com cache)
        # =============================================================
        print(f"\n{Fore.BLUE}{'='*60}")
        print(f"{Fore.BLUE}ETAPA 2: VERSÃO RESUMIDA")
        print(f"{Fore.BLUE}{'='*60}")
        
        # Verifica cache do resumo
        summary_cache_file = formatter._get_cache_path(file_path, cache_type="SUMMARY")
        cached_summary = formatter._load_cache(summary_cache_file)
        
        if cached_summary:
            # Usa resumo em cache
            print(f"{Fore.GREEN}⚡ Usando resumo em cache (processamento instantâneo)")
            summary_text = cached_summary['formatted_text']
        else:
            # Gera resumo do zero
            print(f"{Fore.CYAN}🔄 Gerando versão resumida...")
            summary_text = formatter.generate_summary_version(formatted_text)
            
            # Salva cache do resumo
            formatter._save_cache(summary_cache_file, summary_text, file_path)
        
        # Salva MD Resumido
        output_md_summary = os.path.join(folder, f"{video_name}_RESUMO.md")
        with open(output_md_summary, 'w', encoding='utf-8') as f:
            f.write(summary_text)
        print(f"{Fore.GREEN}📝 MD Resumo salvo: {output_md_summary}")
        
        # Salva Word Resumido
        docx_resumo = formatter.save_as_word(summary_text, f"{video_name}_RESUMO", folder)
        print(f"{Fore.GREEN}📄 Word Resumo salvo: {docx_resumo}")

        
        # =============================================================
        # SUCESSO FINAL
        # =============================================================
        print(f"\n{Fore.GREEN}{'='*60}")
        print(f"{Fore.GREEN}✨ SUCESSO! Apostila Completa + Resumo gerados.")
        print(f"{Fore.GREEN}📄 Completa: {docx_completo}")
        print(f"{Fore.GREEN}📄 Resumo: {docx_resumo}")
        print(f"{Fore.GREEN}{'='*60}")
        
    except KeyboardInterrupt:
        print(f"\n{Fore.YELLOW}⚠️  Interrompido pelo usuário")
        
    except Exception as e:
        print(f"{Fore.RED}❌ Erro: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        process_text_file(file_path)
    else:
        print(f"{Fore.YELLOW}⚠️  Uso: python format_only.py <caminho_do_arquivo.txt>")
        print(f"{Fore.CYAN}💡 O script gerará automaticamente duas versões:")
        print(f"{Fore.CYAN}   1. COMPLETA: Apostila detalhada com todo o conteúdo")
        print(f"{Fore.CYAN}   2. RESUMO: Versão esquemática para revisão rápida")
