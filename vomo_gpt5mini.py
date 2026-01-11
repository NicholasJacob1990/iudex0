#!/usr/bin/env python3
"""
VomoGPT5Mini - Transcrição (MLX-Whisper) + Formatação (GPT-5 Mini)
Versão Híbrida com DOCX Export e Cost Estimation
GPT-5 Mini: 400k contexto, 128k saída máxima
"""

import os
import sys
import time
import subprocess
import hashlib
import json
import re
from pathlib import Path

from colorama import Fore, init
from tqdm import tqdm

from openai import OpenAI, APIError

init(autoreset=True)

# ============== OPCIONAIS ==============
try:
    import mlx_whisper
except ImportError:
    mlx_whisper = None
    print(f"{Fore.YELLOW}⚠️ mlx_whisper não instalado. Use: pip install mlx-whisper")

try:
    from pyannote.audio import Pipeline
    import torch
    HF_TOKEN = os.getenv("HUGGING_FACE_TOKEN")
except ImportError:
    Pipeline = None
    torch = None
    HF_TOKEN = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print(f"{Fore.YELLOW}⚠️ python-docx não disponível. Apenas Markdown será gerado.")

# ============ PREÇOS GPT-5 MINI ============
GPT5_PRICE_INPUT = 0.25   # USD por 1M tokens (Dezembro 2024)
GPT5_PRICE_OUTPUT = 2.00  # USD por 1M tokens

# ============ SYSTEM PROMPT ============

SYSTEM_PROMPT = """# SEU PAPEL
Você é um **revisor sênior de material didático jurídico** com 20 anos de experiência em preparatórios para concursos de procuradoria (PGM/PGE). Seu trabalho é transformar transcrições brutas de aulas em apostilas de alta qualidade, mantendo 100% do conteúdo técnico.

---

# MISSÃO

Transformar a transcrição de aula em uma **apostila didática** com estilo formal, impessoal e expositivo, **PRESERVANDO TODO O CONTEÚDO ORIGINAL**.

⚠️ **REGRA ABSOLUTA: PRESERVE 100% DO CONTEÚDO - NÃO RESUMA, NÃO OMITA NADA.**

---

# EXEMPLOS DE TRANSFORMAÇÃO (Few-Shot)

## Exemplo 1: Transformação de Linguagem
**INPUT (Transcrição bruta):**
> "Então pessoal, a Lei 8.666, né, ela fala lá no artigo 37 que a gente não pode fazer contratação direta, tá? Exceto em alguns casos, tipo... quando tem emergência, sabe?"

**OUTPUT (Apostila formatada):**
> Observa-se que o *Art. 37 da Lei nº 8.666/93* estabelece a vedação à contratação direta como regra geral. Contudo, o legislador previu hipóteses excepcionais, notadamente os casos de emergência, nos termos do *Art. 24, inciso IV* do referido diploma legal.

## Exemplo 2: Preservação de Dicas de Prova
**INPUT:**
> "Isso aqui cai muito em prova, hein! Atenção total nesse artigo 37."

**OUTPUT:**
> Reveste-se de especial importância este dispositivo para fins de certames públicos, sendo recorrentemente objeto de cobrança em provas de procuradorias.

## Exemplo 3: Preservação de Exemplos Completos
**INPUT:**
> "Por exemplo, teve um caso no Rio que o prefeito tentou fazer uma contratação emergencial pra comprar uns equipamentos, mas o TCE glosou porque não tinha a urgência real."

**OUTPUT:**
> A título ilustrativo, cumpre mencionar caso ocorrido no Estado do Rio de Janeiro, no qual determinado prefeito municipal tentou realizar contratação emergencial para aquisição de equipamentos. Contudo, o Tribunal de Contas do Estado glosou o procedimento, fundamentando sua decisão na ausência de urgência real que justificasse a dispensa de licitação.

---

# DIRETRIZES DE FORMATAÇÃO

1. **Estrutura Hierárquica**: Use títulos Markdown (##, ###, ####) para organizar tópicos
2. **Linguagem**: Terceira pessoa impessoal, tempo presente
3. **Citações**: Itálico para leis (*Lei nº 8.666/93*), negrito para conceitos-chave
4. **Tabelas**: Crie tabelas comparativas ao final de tópicos complexos
5. **Preservação Total**: Mantenha TODOS os exemplos, casos, jurisprudência e dicas
6. **Numeração**: Numere os tópicos principais (1., 2., 3.) e subtópicos (1.1, 1.2)

---

# ❌ ERROS A EVITAR

- ❌ Resumir exemplos ou casos práticos
- ❌ Omitir artigos de lei ou números
- ❌ Inventar conteúdo não presente na transcrição
- ❌ Usar primeira pessoa ("eu", "nós")
- ❌ Reorganizar cronologicamente (mantenha ordem da aula)

---

# AUTO-VERIFICAÇÃO (EXECUTE ANTES DE ENVIAR)

□ Mantive TODOS os artigos de lei com números corretos?
□ Mantive TODOS os exemplos e casos completos?
□ Mantive TODAS as dicas de prova?
□ Criei tabelas ao final de tópicos principais?
□ Usei terceira pessoa em TODO o texto?
□ Output tem pelo menos 80% do tamanho do input?

---

# FORMATO DE RESPOSTA

Retorne **APENAS** a apostila em Markdown, começando com # Título da Apostila, sem meta-comentários.
"""


# ============= CLASSE PRINCIPAL =============

class VomoGPT5Mini:
    def __init__(self, whisper_model="large-v3-turbo", gpt_model="gpt-5-mini-2025-08-07"):
        """
        Transcrição com MLX-Whisper + Formatação com GPT-5 Mini.
        GPT-5 Mini: 400k contexto, 128k saída máx.
        """
        print(f"{Fore.CYAN}🚀 Inicializando VomoGPT5Mini...")
        print(f"   🎙️ Whisper: {whisper_model}")
        print(f"   🧠 LLM: {gpt_model}")

        self.whisper_model = whisper_model
        self.gpt_model = gpt_model

        from dotenv import load_dotenv
        load_dotenv(override=True)

        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError(f"{Fore.RED}❌ Configure OPENAI_API_KEY no .env")

        self.client = OpenAI(api_key=api_key)

        self.cache_dir = Path(".cache_vomo_gpt5")
        self.cache_dir.mkdir(exist_ok=True)

        print(f"{Fore.GREEN}✅ Inicialização concluída!\n")

    # ============== TRANSCRIÇÃO ==============

    def optimize_audio(self, file_path: str) -> str:
        """Extrai/converte áudio para 16kHz mono com ffmpeg."""
        print(f"{Fore.YELLOW}⚡ Verificando áudio...{Fore.RESET}")
        src = Path(file_path)

        mp3_path = src.with_suffix(".mp3")
        if mp3_path.exists():
            print(f"   📂 Usando MP3 existente: {mp3_path.name}")
            return str(mp3_path)

        file_hash = hashlib.md5(str(src).encode()).hexdigest()[:8]
        out_wav = src.parent / f"temp_{file_hash}.wav"

        if out_wav.exists():
            print(f"   📂 WAV já existe: {out_wav.name}")
            return str(out_wav)

        print("   🔄 Extraindo áudio otimizado (16kHz mono)...")
        try:
            subprocess.run(
                [
                    "ffmpeg", "-i", str(src),
                    "-vn", "-ac", "1", "-ar", "16000",
                    str(out_wav), "-y",
                    "-hide_banner", "-loglevel", "error",
                ],
                check=True,
                capture_output=True,
            )
        except subprocess.CalledProcessError as e:
            print(f"{Fore.RED}❌ Erro ffmpeg: {e}")
            sys.exit(1)

        print(f"   ✅ Áudio pronto: {out_wav.name}")
        return str(out_wav)

    def transcribe(self, audio_path: str) -> str:
        """Transcreve com MLX‑Whisper; usa Pyannote se disponível."""
        if not mlx_whisper:
            raise ImportError("mlx_whisper não instalado. pip install mlx-whisper")

        print(f"{Fore.GREEN}🎙️  Iniciando transcrição (MLX GPU)...{Fore.RESET}")
        start = time.time()

        cache_file = Path(audio_path).with_suffix(".gpt5_transcription.json")
        if cache_file.exists():
            try:
                print(f"{Fore.CYAN}   📂 Cache de transcrição encontrado...")
                with open(cache_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                print(f"{Fore.GREEN}   ✅ Transcrição carregada do cache.")
                return data["transcript"]
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Erro ao ler cache: {e}")

        print("   🔍 Transcrevendo com parâmetros otimizados...")
        result = mlx_whisper.transcribe(
            audio_path,
            path_or_hf_repo=f"mlx-community/whisper-{self.whisper_model}",
            language="pt",
            word_timestamps=True,
            fp16=True,
            no_speech_threshold=0.6,
            logprob_threshold=-1.0,
            compression_ratio_threshold=2.4,
            condition_on_previous_text=True,
            verbose=False,
        )

        elapsed = time.time() - start
        duration = result.get("duration", 0) if isinstance(result, dict) else 0
        rtf = elapsed / duration if duration > 0 else 0
        print(f"{Fore.GREEN}   ✅ Transcrição em {elapsed:.1f}s (RTF: {rtf:.2f}x)")

        transcript = self._with_optional_diarization(audio_path, result)

        try:
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(
                    {
                        "transcript": transcript,
                        "duration": duration,
                        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
                    },
                    f,
                    ensure_ascii=False,
                    indent=2,
                )
            print(f"   💾 Cache salvo: {cache_file.name}")
        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ Erro ao salvar cache: {e}")

        return transcript

    def _with_optional_diarization(self, audio_path, whisper_result):
        if not (Pipeline and torch and HF_TOKEN):
            return self._format_no_diar(whisper_result)

        try:
            print("   🗣️  Iniciando diarização (Pyannote)...")
            pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1", use_auth_token=HF_TOKEN
            )
            device = "mps" if torch.backends.mps.is_available() else "cpu"
            pipeline.to(torch.device(device))
            diar = pipeline(audio_path)
            text = self._align_diarization(whisper_result["segments"], diar)
            print(f"{Fore.GREEN}   ✅ Diarização concluída.")
            return text
        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ Falha na diarização: {e}")
            return self._format_no_diar(whisper_result)

    def _format_no_diar(self, result):
        parts = []
        for seg in result["segments"]:
            ts = self._fmt_ts(seg["start"])
            parts.append(f"[{ts}] {seg['text'].strip()}")
        return " ".join(parts)

    def _fmt_ts(self, sec):
        m, s = divmod(sec, 60)
        h, m = divmod(m, 60)
        return f"{int(h):02d}:{int(m):02d}:{int(s):02d}" if h > 0 else f"{int(m):02d}:{int(s):02d}"

    def _align_diarization(self, segments, diarization_output):
        diar_segs = [
            (t.start, t.end, s) for t, _, s in diarization_output.itertracks(yield_label=True)
        ]
        out = []
        current_speaker = None

        for seg in segments:
            start, end = seg["start"], seg["end"]
            best_spk = "SPEAKER 0"
            max_ov = 0
            for ds, de, spk in diar_segs:
                ov = max(0, min(end, de) - max(start, ds))
                if ov > max_ov:
                    max_ov = ov
                    best_spk = f"SPEAKER {int(spk.split('_')[-1]) + 1}"
            if best_spk != current_speaker:
                out.append(f"\n\n{best_spk}\n")
                current_speaker = best_spk
            out.append(f"[{self._fmt_ts(start)}] {seg['text'].strip()}")
        return " ".join(out)

    # ============== FORMATAÇÃO GPT‑5 MINI ==============

    def format_with_gpt5(self, transcription: str, video_name: str) -> str:
        """
        Usa GPT‑5 Mini para formatar a transcrição inteira em UMA chamada.
        Contexto 400k e saída teórica até 128k tokens.
        """
        print(f"\n{Fore.CYAN}{'='*70}")
        print(f"{Fore.CYAN}🧠 FORMATAÇÃO COM GPT‑5 MINI")
        print(f"{Fore.CYAN}{'='*70}{Fore.RESET}")

        n_chars = len(transcription)
        est_tokens_in = n_chars // 4
        print(f"   📊 Transcrição: {n_chars:,} chars (~{est_tokens_in:,} tokens estimados)")

        if est_tokens_in > 272_000:
            print(f"{Fore.YELLOW}   ⚠️ AVISO: pode ultrapassar o input recomendado (~272k).")
            print("      Considere dividir a aula ou cortar silêncio/introdução.\n")

        cache_hash = hashlib.sha256(transcription.encode()).hexdigest()[:16]
        cache_file = self.cache_dir / f"formatted_{cache_hash}.md"
        
        if cache_file.exists():
            try:
                print(f"{Fore.CYAN}   📂 Apostila encontrada em cache!")
                with open(cache_file, "r", encoding="utf-8") as f:
                    cached = f.read()
                est_tokens_out = len(cached) // 4
                cost = (est_tokens_in * GPT5_PRICE_INPUT + est_tokens_out * GPT5_PRICE_OUTPUT) / 1_000_000
                print(f"   💰 Custo estimado (se fosse chamada real): ${cost:.4f} USD")
                return cached
            except Exception:
                pass

        user_prompt = f"""# AULA PARA FORMATAR

Título da aula: {video_name}

TRANSCRIÇÃO COMPLETA (inclui marcações de tempo e speakers):

{transcription}

---

Gere a apostila completa conforme o system prompt, em Markdown.
"""

        print("   🚀 Chamando GPT‑5 Mini...")
        start = time.time()

        try:
            resp = self.client.chat.completions.create(
                model=self.gpt_model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.2,
                max_tokens=110_000,  # margem de segurança < 128k
            )
        except APIError as e:
            print(f"{Fore.RED}   ❌ Erro API OpenAI: {e}")
            raise
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro inesperado: {e}")
            raise

        elapsed = time.time() - start
        content = resp.choices[0].message.content or ""
        content = content.strip()

        est_tokens_out = len(content) // 4
        cost = (est_tokens_in * GPT5_PRICE_INPUT + est_tokens_out * GPT5_PRICE_OUTPUT) / 1_000_000
        
        print(f"{Fore.GREEN}   ✅ Formatação em {elapsed:.1f}s (~{est_tokens_out:,} tokens saída)")
        print(f"   💰 Custo estimado: ${cost:.4f} USD")

        try:
            with open(cache_file, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"   💾 Apostila em cache: {cache_file.name}")
        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ Erro ao salvar cache: {e}")

        return content

    # ============== EXPORTAÇÃO DOCX ==============

    def save_docx(self, markdown_text, video_name, output_path):
        """Converte Markdown para DOCX profissional"""
        if not DOCX_AVAILABLE:
            print(f"{Fore.YELLOW}   ⚠️ python-docx não disponível, pulando DOCX")
            return None

        print(f"{Fore.CYAN}📄 Gerando DOCX...")

        doc = Document()

        # Margens
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

        # Título
        title = doc.add_heading(video_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 51, 102)

        # Data
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')} - GPT-5 Mini")
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Processa Markdown
        lines = markdown_text.split('\n')
        in_table = False
        table_rows = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Tabelas
            if '|' in line and not line.startswith('|--'):
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append([c.strip() for c in line.split('|')[1:-1]])
                continue
            elif in_table:
                self._add_table(doc, table_rows)
                in_table = False
                table_rows = []

            # Headers
            if line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:], style='Quote')
                p.paragraph_format.left_indent = Inches(0.5)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                self._format_inline(p, line[2:])
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                self._format_inline(p, line)

        # Tabela final
        if table_rows:
            self._add_table(doc, table_rows)

        doc.save(output_path)
        print(f"{Fore.GREEN}   ✅ DOCX salvo: {output_path}")
        return output_path

    def _format_inline(self, paragraph, text):
        """Formata bold/italic inline"""
        paragraph.clear()
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'

        last_end = 0
        for match in re.finditer(pattern, text):
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            if match.group(0).startswith('***'):
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(0).startswith('**'):
                run = paragraph.add_run(match.group(3))
                run.bold = True
            elif match.group(0).startswith('*'):
                run = paragraph.add_run(match.group(4))
                run.italic = True
            elif match.group(0).startswith('`'):
                run = paragraph.add_run(match.group(5))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(200, 0, 0)

            last_end = match.end()

        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_table(self, doc, rows):
        """Adiciona tabela formatada"""
        if len(rows) < 2:
            return

        max_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'

        for i, row in enumerate(rows):
            for j in range(max_cols):
                if j < len(row):
                    cell = table.rows[i].cells[j]
                    self._format_inline(cell.paragraphs[0], row[j])

                    if i == 0:
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '0066CC')
                        cell._element.get_or_add_tcPr().append(shading_elm)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(255, 255, 255)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============== PIPELINE COMPLETO ==============

    def process_video(self, video_path: str, output_folder: str = "apostilas_gpt5", export_docx=True):
        print(f"\n{Fore.CYAN}{'='*70}")
        print(f"{Fore.CYAN}🎬 PROCESSANDO VÍDEO (GPT‑5 Mini)")
        print(f"{Fore.CYAN}{'='*70}{Fore.RESET}\n")

        video = Path(video_path)
        if not video.exists():
            print(f"{Fore.RED}❌ Arquivo não encontrado: {video}")
            sys.exit(1)

        video_name = video.stem
        out_dir = Path(output_folder)
        out_dir.mkdir(exist_ok=True)
        
        md_file = out_dir / f"{video_name}.md"
        docx_file = out_dir / f"{video_name}.docx"

        print(f"📹 Vídeo: {video.name}")
        print(f"📁 Saída: {out_dir}\n")

        audio = self.optimize_audio(str(video))
        transcription = self.transcribe(audio)
        
        # Salva transcrição bruta
        raw_file = out_dir / f"{video_name}_RAW.txt"
        with open(raw_file, 'w', encoding='utf-8') as f:
            f.write(transcription)
        print(f"   💾 Transcrição bruta salva: {raw_file.name}")
        
        apostila = self.format_with_gpt5(transcription, video_name)

        print(f"\n{Fore.CYAN}💾 Salvando apostila...{Fore.RESET}")
        try:
            with open(md_file, "w", encoding="utf-8") as f:
                f.write(apostila)
            size_kb = md_file.stat().st_size / 1024
            print(f"{Fore.GREEN}   ✅ Markdown salvo: {md_file.name} ({size_kb:.1f} KB)")
        except Exception as e:
            print(f"{Fore.RED}❌ Erro ao salvar: {e}")
            sys.exit(1)

        # Exportar DOCX
        if export_docx and DOCX_AVAILABLE:
            try:
                self.save_docx(apostila, video_name, docx_file)
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Erro ao gerar DOCX: {e}")

        print(f"\n{Fore.GREEN}{'='*70}")
        print(f"{Fore.GREEN}✅ PIPELINE CONCLUÍDO (GPT‑5 Mini)")
        print(f"{Fore.GREEN}{'='*70}{Fore.RESET}")
        print(f"📄 Markdown: {md_file.absolute()}")
        if export_docx and DOCX_AVAILABLE:
            print(f"📄 DOCX: {docx_file.absolute()}")
        
        return str(md_file)


# ============== CLI ==============

def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="VomoGPT5Mini - Transcrição + Apostila com GPT‑5 Mini",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos de uso:
  python vomo_gpt5mini.py video.mp4
  python vomo_gpt5mini.py aula.mp3 -o apostilas --no-docx
  python vomo_gpt5mini.py video.mov --whisper large-v3
        """
    )
    parser.add_argument("video", help="Arquivo de vídeo/áudio (mp4, mp3, etc.)")
    parser.add_argument(
        "-o", "--output", default="apostilas_gpt5", help="Pasta de saída (default: apostilas_gpt5)"
    )
    parser.add_argument(
        "--whisper", default="large-v3-turbo", help="Modelo Whisper (default: large-v3-turbo)"
    )
    parser.add_argument(
        "--gpt", default="gpt-5-mini-2025-08-07", help="Modelo GPT‑5 Mini (default: gpt-5-mini-2025-08-07)"
    )
    parser.add_argument('--no-docx', action='store_true', help='Não gerar arquivo DOCX')

    args = parser.parse_args()

    try:
        app = VomoGPT5Mini(whisper_model=args.whisper, gpt_model=args.gpt)
        app.process_video(args.video, args.output, export_docx=not args.no_docx)
    except KeyboardInterrupt:
        print(f"\n{Fore.YELLOW}⚠️ Interrompido pelo usuário")
        sys.exit(0)
    except Exception as e:
        print(f"\n{Fore.RED}❌ Erro fatal: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
