#!/usr/bin/env python3
"""
Script para formatar transcrições usando um prompt personalizado com GPT-5 mini.
Uso: python format_with_custom_prompt.py <arquivo_transcricao.txt>
"""

import os
import sys
import asyncio
from pathlib import Path
from openai import AsyncOpenAI
from colorama import Fore, init
from tqdm import tqdm
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

init(autoreset=True)

# ============================================
# PROMPT PERSONALIZADO (FORNECIDO PELO USUÁRIO)
# ============================================
CUSTOM_SYSTEM_PROMPT = """Você é um especialista em revisão de transcrições para concursos de procuradorias. Sua tarefa é revisar a transcrição fornecida de uma aula, corrigindo erros de português, erros gramaticais e de pontuação, melhorando a formatação para facilitar a leitura, e mantendo o conteúdo original. Não resuma, não parafraseie e não adicione informações que não estejam na transcrição original. Siga estas diretrizes:
 
-mantenha o modo em primeira pessoa
-Corrija erros gramaticais, ortográficos e de pontuação, tornando o texto gramaticalmente correto e claro,
-Mantenha todo o conteúdo original, incluindo ideias, exemplos, explicações, pausas, hesitações e ideias incompletas, fazendo o uso apropriado de aspas, parênteses. Não resuma, não omita informações nem altere o significado.
-Melhore a formatação para facilitar a leitura.
-Mantenha todo o conteúdo original, mas corrija erros da linguagem coloquial para torná-la mais clara, didática e legível
-Ajuste a linguagem coloquial para um português padrão, mantendo o significado original.
-Elimine vicios da oralidade e gírias
-Preserve a sequência exata das falas e ideias apresentadas
-Utilize formatação e estrutura com parágrafos bem definidos, facilitando a leitura e compreensão, para melhorar a legibilidade, seguindo o fluxo natural do discurso. Evite parágrafos longos.
-Reproduza fielmente as informações, apenas melhorando a clareza e a legibilidade.
-Utilize conectivos necessários para tornar o texto mais fluido. Aplique a pontuação devida para deixar o texto coeso e coerente.
-Corrija vícios de linguagem, como repetições desnecessárias, uso excessivo de advérbios, linguagem vaga ou imprecisa, gírias, expressões redundantes, e outros erros que afetem a clareza e a eficácia da comunicação, sem alterar o significado do texto.
-Identifique e rotule os diferentes falantes, se existentes, organizando suas falas de forma clara.
-Divida a aula em tópicos e subtópicos para melhor organização e visualização do conteúdo
-Enumere os tópicos e subtópicos, use negrito quando mais apropriado
-Intercale parágrafos curtos e longos conforme a idéia neles contida, tornando a leitura menos cansativa
-Seja didático sem perder detalhes e conteúdo
-USE TEXTO CORRIDO NA MEDIDA DO POSSÍVEL
-**Ao final de cada tópico/capítulo, sintetize/resume o assunto de forma esquematizada, preferencialmente por tabelas

**Ao final de CADA tópico principal, crie uma tabela de síntese (EXEMPLIFICATIVO, PODENDO SER COMPOSTO POR OUTROS ELEMENTOS NAS COLUNAS:**

| Conceito/Instituto | Definição | Fundamento Legal | Observações |
|-------------------|-----------|-----------------|-------------|
| [Preencher] | [Resumo] | [Art. X, Lei Y] | [Dicas/Exceções] |

Por favor, forneça a versão revisada da transcrição, seguindo estritamente as diretrizes acima. Lembre-se: o objetivo é manter o conteúdo fiel ao original, melhorando apenas a clareza e legibilidade.

IMPORTANTE:
- Retorne APENAS o texto formatado em Markdown
- NÃO adicione comentários como "Continuação...", "Parte X...", etc.
- Mantenha a numeração sequencial dos tópicos ao longo de todo o documento
"""


class CustomPromptFormatter:
    def __init__(self):
        """Inicializa o formatador com GPT-5 mini"""
        print(f"{Fore.CYAN}🚀 Inicializando Formatador com Prompt Personalizado...")
        
        from dotenv import load_dotenv
        load_dotenv(override=True)
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError(f"{Fore.RED}❌ Configure: export OPENAI_API_KEY='sk-...'")
        
        self.async_client = AsyncOpenAI(api_key=api_key)
        self.model = "gpt-5-mini-2025-08-07"
        print(f"{Fore.GREEN}✅ Usando modelo: {self.model}")

    def _smart_chunk_overlapping(self, text, max_size=12000, overlap=1500):
        """
        Divide texto com SOBREPOSIÇÃO para evitar perda nas bordas
        """
        if len(text) <= max_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_size
            
            # Se não for o último chunk, encontra quebra natural
            if end < len(text):
                search_zone = text[max(0, end-2000):end]
                
                # Procura por quebra de parágrafo ou ponto final
                last_para = search_zone.rfind('\n\n')
                last_period = search_zone.rfind('. ')
                
                if last_para != -1:
                    end = end - 2000 + last_para
                elif last_period != -1:
                    end = end - 2000 + last_period + 1
            
            chunk = text[start:end]
            chunks.append(chunk)
            
            if end < len(text):
                start = end - overlap
            else:
                break
        
        return chunks

    async def _format_chunk_async(self, chunk_text, chunk_idx, total_chunks):
        """Formata um chunk usando o prompt personalizado"""
        word_count = len(chunk_text.split())
        
        user_content = f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TAREFA: Revisar e formatar a transcrição abaixo.
Tamanho original: ~{word_count} palavras.
Chunk {chunk_idx + 1} de {total_chunks}.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

<transcrição>
{chunk_text}
</transcrição>
"""
        
        try:
            response = await self.async_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": CUSTOM_SYSTEM_PROMPT},
                    {"role": "user", "content": user_content}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"{Fore.RED}⚠️ Erro ao formatar chunk {chunk_idx}: {e}")
            return chunk_text

    def _merge_chunks_with_deduplication(self, formatted_chunks):
        """Remove duplicações causadas por overlap entre chunks"""
        if len(formatted_chunks) <= 1:
            return formatted_chunks[0] if formatted_chunks else ""
        
        from difflib import SequenceMatcher
        
        final_text = formatted_chunks[0]
        
        for i in range(1, len(formatted_chunks)):
            current_chunk = formatted_chunks[i]
            
            # Pega as últimas 500 palavras do texto acumulado
            tail_words = final_text.split()[-500:]
            tail_text = " ".join(tail_words)
            
            # Pega as primeiras 500 palavras do chunk atual
            head_words = current_chunk.split()[:500]
            head_text = " ".join(head_words)
            
            # Encontra a sobreposição
            matcher = SequenceMatcher(None, tail_text, head_text)
            match = matcher.find_longest_match(0, len(tail_text), 0, len(head_text))
            
            if match.size > 50:  # Sobreposição significativa
                # Encontra onde começa a parte nova
                overlap_end_in_chunk = match.b + match.size
                
                # Calcula a posição no chunk original
                words_to_skip = len(head_text[:overlap_end_in_chunk].split())
                chunk_words = current_chunk.split()
                
                if words_to_skip < len(chunk_words):
                    new_content = " ".join(chunk_words[words_to_skip:])
                    final_text = final_text + "\n\n" + new_content
                else:
                    # Chunk está totalmente contido
                    pass
            else:
                # Sem sobreposição: adiciona tudo
                final_text = final_text + "\n\n" + current_chunk
        
        return final_text

    def save_as_word(self, markdown_text, output_path, title="Apostila de Direito Administrativo"):
        """Salva o texto formatado como documento Word PROFISSIONAL (Apostila)"""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re
        
        doc = Document()
        
        # ============================================
        # CONFIGURAÇÃO DE ESTILOS PROFISSIONAIS
        # ============================================
        
        # Configuração de margens do documento
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1)
            section.header_distance = Inches(0.4)
            section.footer_distance = Inches(0.4)
        
        # Estilo do corpo do texto
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Georgia'
        font.size = Pt(11)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.first_line_indent = Cm(1.25)
        
        # Configurar estilos de Heading
        for i in range(1, 5):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Arial'
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)  # Azul escuro
            if i == 1:
                heading_style.font.size = Pt(18)
                heading_style.paragraph_format.space_before = Pt(24)
                heading_style.paragraph_format.space_after = Pt(12)
            elif i == 2:
                heading_style.font.size = Pt(14)
                heading_style.paragraph_format.space_before = Pt(18)
                heading_style.paragraph_format.space_after = Pt(8)
            elif i == 3:
                heading_style.font.size = Pt(12)
                heading_style.paragraph_format.space_before = Pt(14)
                heading_style.paragraph_format.space_after = Pt(6)
            else:
                heading_style.font.size = Pt(11)
                heading_style.paragraph_format.space_before = Pt(10)
                heading_style.paragraph_format.space_after = Pt(4)
        
        # ============================================
        # CAPA PROFISSIONAL
        # ============================================
        
        # Espaço superior
        for _ in range(6):
            doc.add_paragraph()
        
        # Título principal
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.upper())
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        
        # Subtítulo
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("Material de Estudo para Concursos")
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        
        # Linha decorativa
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("━" * 40)
        line_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        
        # Espaço
        for _ in range(4):
            doc.add_paragraph()
        
        # Área de concurso
        concurso_para = doc.add_paragraph()
        concurso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        concurso_run = concurso_para.add_run("PGM/PGE - Procuradorias")
        concurso_run.font.name = 'Arial'
        concurso_run.font.size = Pt(16)
        concurso_run.font.bold = True
        
        # Data
        from datetime import datetime
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime("%B de %Y").title())
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(12)
        date_run.font.italic = True
        
        # Quebra de página após capa
        doc.add_page_break()
        
        # ============================================
        # SUMÁRIO (placeholder - preenchido dinamicamente)
        # ============================================
        
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run("SUMÁRIO")
        toc_run.font.name = 'Arial'
        toc_run.font.size = Pt(16)
        toc_run.font.bold = True
        toc_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        
        doc.add_paragraph()  # Espaço
        
        # Nota sobre sumário
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note.add_run("(Índice gerado automaticamente - atualize o documento no Word)")
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        doc.add_page_break()
        
        # ============================================
        # CABEÇALHO E RODAPÉ
        # ============================================
        
        section = doc.sections[0]
        
        # Cabeçalho
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run(title)
        header_run.font.name = 'Arial'
        header_run.font.size = Pt(9)
        header_run.font.italic = True
        header_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        # Rodapé com número de página
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adiciona número de página
        self._add_page_number(footer_para)
        
        # ============================================
        # CONTEÚDO PRINCIPAL
        # ============================================
        
        lines = markdown_text.split('\n')
        current_table_data = []
        in_table = False
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                if in_table and current_table_data:
                    self._add_professional_table(doc, current_table_data)
                    current_table_data = []
                    in_table = False
                continue
            
            # Detecta tabelas markdown
            if line_stripped.startswith('|') and line_stripped.endswith('|'):
                in_table = True
                # Ignora linha de separação (|---|---|)
                if not re.match(r'^[\|\-\:\s]+$', line_stripped):
                    cells = [cell.strip() for cell in line_stripped.split('|')[1:-1]]
                    if cells and any(c for c in cells):  # Não adiciona linhas vazias
                        current_table_data.append(cells)
                continue
            
            if in_table and current_table_data:
                self._add_professional_table(doc, current_table_data)
                current_table_data = []
                in_table = False
            
            # Headings
            if line_stripped.startswith('# '):
                p = doc.add_heading(line_stripped[2:], level=1)
            elif line_stripped.startswith('## '):
                p = doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith('### '):
                p = doc.add_heading(line_stripped[4:], level=3)
            elif line_stripped.startswith('#### '):
                p = doc.add_heading(line_stripped[5:], level=4)
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, line_stripped[2:])
            elif re.match(r'^\d+\.\s', line_stripped):
                p = doc.add_paragraph(style='List Number')
                # Extrai o texto após o número
                text = re.sub(r'^\d+\.\s*', '', line_stripped)
                self._add_formatted_text(p, text)
            elif line_stripped.startswith('>'):
                # Citação/bloco de destaque
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.right_indent = Cm(0.5)
                self._add_formatted_text(p, line_stripped[1:].strip())
                for run in p.runs:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
            else:
                # Texto normal com formatação inline
                p = doc.add_paragraph()
                self._add_formatted_text(p, line_stripped)
        
        # Finaliza última tabela se existir
        if in_table and current_table_data:
            self._add_professional_table(doc, current_table_data)
        
        doc.save(output_path)
        print(f"{Fore.GREEN}📄 Documento Word PROFISSIONAL salvo: {output_path}")

    def _add_page_number(self, paragraph):
        """Adiciona número de página ao rodapé"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        instr_text.text = "PAGE"
        
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

    def _add_professional_table(self, doc, table_data):
        """Adiciona uma tabela profissional ao documento"""
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        if not table_data or not table_data[0]:
            return
        
        num_cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Estilo de tabela com bordas
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            # Altura mínima da linha
            row.height = Pt(20)
            
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = cell_text
                    
                    # Formatação do texto na célula
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(4)
                        paragraph.paragraph_format.space_after = Pt(4)
                        
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                            
                            # Cabeçalho em negrito com fundo colorido
                            if i == 0:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # Branco
                    
                    # Cor de fundo do cabeçalho
                    if i == 0:
                        self._set_cell_shading(cell, "1A3C6E")  # Azul escuro
                    elif i % 2 == 0:
                        self._set_cell_shading(cell, "F0F4F8")  # Cinza claro alternado
        
        # Espaço após tabela
        doc.add_paragraph()

    def _set_cell_shading(self, cell, color_hex):
        """Define a cor de fundo de uma célula"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        tcPr.append(shading)

    def _add_formatted_text(self, paragraph, text):
        """Adiciona texto com formatação inline (negrito, itálico)"""
        from docx.shared import RGBColor
        import re
        
        # Padrão para **negrito**, *itálico* e `código`
        pattern = r'(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # Vermelho escuro
            else:
                paragraph.add_run(part)

    async def process_file(self, input_path):
        """Processa o arquivo de transcrição completo"""
        print(f"\n{Fore.CYAN}📂 Lendo arquivo: {input_path}")
        
        with open(input_path, 'r', encoding='utf-8') as f:
            raw_text = f.read()
        
        print(f"{Fore.GREEN}   ✅ {len(raw_text):,} caracteres lidos")
        print(f"{Fore.GREEN}   ✅ ~{len(raw_text.split()):,} palavras")
        
        # Divide em chunks com overlap
        chunks = self._smart_chunk_overlapping(raw_text, max_size=12000, overlap=1500)
        print(f"\n{Fore.CYAN}📊 Dividido em {len(chunks)} chunks para processamento")
        
        # Processa chunks em paralelo (com limite de concorrência)
        formatted_chunks = []
        semaphore = asyncio.Semaphore(3)  # Máximo 3 requisições simultâneas
        
        async def process_with_semaphore(chunk, idx):
            async with semaphore:
                print(f"   Processando chunk {idx + 1}/{len(chunks)}...")
                result = await self._format_chunk_async(chunk, idx, len(chunks))
                return result
        
        tasks = [process_with_semaphore(chunk, i) for i, chunk in enumerate(chunks)]
        formatted_chunks = await asyncio.gather(*tasks)
        
        print(f"\n{Fore.CYAN}🔗 Mesclando chunks com deduplicação...")
        final_text = self._merge_chunks_with_deduplication(formatted_chunks)
        
        # Gera caminhos de saída
        input_path = Path(input_path)
        output_md = input_path.parent / f"{input_path.stem}_FORMATADO.md"
        output_docx = input_path.parent / f"{input_path.stem}_FORMATADO.docx"
        
        # Salva Markdown
        with open(output_md, 'w', encoding='utf-8') as f:
            f.write(final_text)
        print(f"{Fore.GREEN}📝 Markdown salvo: {output_md}")
        
        # Salva Word
        self.save_as_word(final_text, str(output_docx))
        
        print(f"\n{Fore.GREEN}{'='*50}")
        print(f"{Fore.GREEN}✅ PROCESSAMENTO CONCLUÍDO!")
        print(f"{Fore.GREEN}{'='*50}")
        print(f"   📝 Markdown: {output_md}")
        print(f"   📄 Word: {output_docx}")
        
        return str(output_md), str(output_docx)


async def main():
    if len(sys.argv) < 2:
        print(f"{Fore.RED}❌ Uso: python format_with_custom_prompt.py <arquivo_transcricao.txt>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    if not os.path.exists(input_file):
        print(f"{Fore.RED}❌ Arquivo não encontrado: {input_file}")
        sys.exit(1)
    
    formatter = CustomPromptFormatter()
    await formatter.process_file(input_file)


if __name__ == "__main__":
    asyncio.run(main())
