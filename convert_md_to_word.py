from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import sys

def convert_md_to_docx(md_file, docx_file):
    """Converte arquivo Markdown para Word (.docx)"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Configurações do documento
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Extrai título (primeira linha começando com #)
    lines = content.split('\n')
    title_text = "Apostila"
    for line in lines[:5]:
        if line.startswith('# '):
            title_text = line.replace('# ', '').strip()
            break
    
    # Título
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data
    date_para = doc.add_paragraph()
    date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')}").italic = True
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Espaço
    
    # Processa o conteúdo
    skip_first_title = True
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
        
        # Pula primeira linha de título (já usada)
        if skip_first_title and line.startswith('# '):
            skip_first_title = False
            continue
        
        # Pula metadata do fim
        if line.startswith('---') or line.startswith('*Gerado em'):
            continue
        
        # Detecta headings
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        
        # Detecta listas
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            text = line.split('. ', 1)[1] if '. ' in line else line
            doc.add_paragraph(text, style='List Number')
        
        # Texto normal
        else:
            p = doc.add_paragraph()
            # Processa negritos
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    p.add_run(part)
                else:
                    p.add_run(part).bold = True
    
    doc.save(docx_file)
    print(f"✅ Convertido: {docx_file}")

if __name__ == "__main__":
    md_file = "/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/03_Direito_Financeiro_Luiz_Oliveira_APOSTILA.md"
    docx_file = md_file.replace('.md', '.docx')
    convert_md_to_docx(md_file, docx_file)
