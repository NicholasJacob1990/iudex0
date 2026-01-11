#!/usr/bin/env python3
"""
VOMO Gemini 2.5 Flash - Script Simplificado para Transcrição e Formatação
===========================================================================

Este script aproveita o limite de 65k tokens de output do Gemini 2.5 Flash
para processar aulas COMPLETAS em uma única chamada, eliminando:
- Chunking/Divisão sequencial
- Deduplicação de overlap
- Context management
- Rate limiting complexo

Requisitos:
    pip install google-generativeai mlx-whisper python-docx colorama tqdm

Uso:
    python vomo_gemini_flash.py <arquivo_audio_ou_texto>
    python vomo_gemini_flash.py aula.mp3
    python vomo_gemini_flash.py transcricao.txt

Variáveis de Ambiente:
    GEMINI_API_KEY - Chave da API do Google AI Studio
"""

import os
import sys
import time
import subprocess
import hashlib
import json
from pathlib import Path

try:
    from colorama import Fore, init
    init(autoreset=True)
except ImportError:
    class Fore:
        RED = GREEN = YELLOW = CYAN = MAGENTA = ""

try:
    import google.generativeai as genai
    from google.generativeai.types import GenerationConfig
except ImportError:
    print("❌ Instale: pip install google-generativeai")
    sys.exit(1)

try:
    import mlx_whisper
except ImportError:
    mlx_whisper = None
    print("⚠️ mlx_whisper não disponível. Apenas arquivos .txt serão aceitos.")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx não disponível.")

# =============================================================================
# PROMPT OTIMIZADO PARA CHAMADA ÚNICA
# =============================================================================

SYSTEM_PROMPT = """# SEU PAPEL
Você é um EXCELENTÍSSIMO REDATOR JURÍDICO E DIDÁTICO.

# MISSÃO
Transformar a transcrição COMPLETA de uma aula em uma APOSTILA DIDÁTICA perfeita.

# DIRETRIZES ABSOLUTAS

## ❌ O QUE NUNCA FAZER
1. **NÃO RESUMA** - O output deve ser tão detalhado quanto a entrada
2. **NÃO OMITA** - Exemplos, casos, jurisprudência, artigos de lei
3. **NÃO INVENTE** - Apenas formate o que está na transcrição

## ✅ O QUE FAZER
1. **Correção Linguística**
   - Ajuste linguagem coloquial → padrão culto
   - Remova vícios: "né", "tipo assim", "então", "tá"
   - Corrija concordância e pontuação

2. **Estruturação (Markdown)**
   - Use ## para tópicos principais
   - Use ### para subtópicos
   - Use **negrito** para conceitos-chave
   - Use > blockquote para citações de lei

3. **Tabelas de Síntese**
   Ao final de CADA tópico principal (##), crie uma tabela:
   
   | Conceito | Definição | Fundamento Legal |
   |----------|-----------|------------------|
   | [Nome]   | [Resumo]  | [Art. X, Lei Y]  |

## PRESERVE OBRIGATORIAMENTE
- **TODOS os números**: Artigos, Leis, Súmulas, REsp, RE
- **TODOS os exemplos** citados pelo professor
- **TODAS as referências** a autores e julgados
- **Ênfases pedagógicas**: "isso cai muito!", "atenção aqui"

# FORMATO DE SAÍDA
Retorne APENAS o Markdown formatado da apostila completa.
Não inclua comentários, explicações ou meta-texto.
"""

# =============================================================================
# CLASSE PRINCIPAL
# =============================================================================

class VomoGeminiFlash:
    def __init__(self):
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError(f"{Fore.RED}❌ Configure GEMINI_API_KEY")
        
        genai.configure(api_key=api_key)
        
        # Modelo Gemini 2.5 Flash
        self.model = genai.GenerativeModel(
            model_name="gemini-2.5-flash-preview-05-20",
            system_instruction=SYSTEM_PROMPT
        )
        
        # Whisper model
        self.whisper_model = "large-v3-turbo"
        
        print(f"{Fore.GREEN}✅ Gemini 2.5 Flash inicializado")

    def optimize_audio(self, file_path):
        """Extrai áudio otimizado (16kHz mono)"""
        print(f"{Fore.YELLOW}⚡ Processando áudio...")
        
        # Verifica se já existe MP3
        mp3_path = Path(file_path).with_suffix('.mp3')
        if mp3_path.exists():
            return str(mp3_path)
        
        # Cria WAV otimizado
        file_hash = hashlib.md5(file_path.encode()).hexdigest()[:8]
        output_path = f"temp_{file_hash}.wav"
        
        if os.path.exists(output_path):
            return output_path
        
        subprocess.run([
            'ffmpeg', '-i', file_path,
            '-vn', '-ac', '1', '-ar', '16000',
            output_path, '-y', '-hide_banner', '-loglevel', 'error'
        ], check=True, capture_output=True)
        
        return output_path

    def transcribe(self, audio_path):
        """Transcrição com MLX-Whisper otimizado"""
        if not mlx_whisper:
            raise ImportError("mlx_whisper não disponível")
        
        print(f"{Fore.GREEN}🎙️ Transcrevendo com MLX-Whisper...")
        start = time.time()
        
        # Cache
        cache_file = Path(audio_path).with_suffix('.transcript.json')
        if cache_file.exists():
            print(f"{Fore.CYAN}   📂 Cache encontrado")
            with open(cache_file, 'r') as f:
                return json.load(f)['text']
        
        result = mlx_whisper.transcribe(
            audio_path,
            path_or_hf_repo=f"mlx-community/whisper-{self.whisper_model}",
            language="pt",
            word_timestamps=True,
            fp16=True,
            no_speech_threshold=0.6,
            condition_on_previous_text=True,
            verbose=False
        )
        
        # Extrai texto
        text = result.get('text', '')
        if not text and 'segments' in result:
            text = ' '.join(s['text'] for s in result['segments'])
        
        # Salva cache
        with open(cache_file, 'w') as f:
            json.dump({'text': text}, f, ensure_ascii=False)
        
        elapsed = time.time() - start
        print(f"{Fore.GREEN}   ✅ Transcrito em {elapsed:.1f}s ({len(text)} chars)")
        
        return text

    def generate_apostila(self, transcription, video_name):
        """
        Gera apostila em CHAMADA ÚNICA usando Gemini 2.5 Flash
        
        Sem chunking, sem overlap, sem deduplicação.
        O modelo processa até ~1M tokens de input e gera até 65k de output.
        """
        print(f"\n{Fore.MAGENTA}⚡ Gerando apostila com Gemini 2.5 Flash...")
        print(f"   📄 Input: {len(transcription):,} caracteres (~{len(transcription)//4:,} tokens)")
        
        start = time.time()
        
        # Config para output máximo
        config = GenerationConfig(
            temperature=0.3,
            max_output_tokens=65536,  # 65k tokens de output!
            top_p=0.95,
        )
        
        prompt = f"""# Título: {video_name}

## TRANSCRIÇÃO COMPLETA DA AULA:

{transcription}

---

Gere a apostila formatada seguindo TODAS as diretrizes do sistema.
Preserve 100% do conteúdo. Use Markdown."""

        try:
            response = self.model.generate_content(
                prompt,
                generation_config=config
            )
            
            result = response.text
            elapsed = time.time() - start
            
            print(f"{Fore.GREEN}   ✅ Apostila gerada em {elapsed:.1f}s")
            print(f"   📄 Output: {len(result):,} caracteres (~{len(result)//4:,} tokens)")
            
            return result
            
        except Exception as e:
            print(f"{Fore.RED}❌ Erro: {e}")
            raise

    def save_markdown(self, content, output_path):
        """Salva como Markdown"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"{Fore.GREEN}✅ Markdown salvo: {output_path}")

    def save_docx(self, content, video_name, output_path):
        """Converte Markdown para DOCX profissional"""
        if not DOCX_AVAILABLE:
            return
        
        print(f"{Fore.CYAN}📄 Gerando DOCX...")
        
        doc = Document()
        
        # Margens
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        
        # Título
        title = doc.add_heading(video_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        doc.add_paragraph().add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')} - Gemini 2.5 Flash")
        doc.add_page_break()
        
        # Processa Markdown
        lines = content.split('\n')
        in_table = False
        table_rows = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Tabelas
            if '|' in line and not line.startswith('|--'):
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append([c.strip() for c in line.split('|')[1:-1]])
                continue
            elif in_table:
                self._add_table(doc, table_rows)
                in_table = False
                table_rows = []
            
            # Headers
            if line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:], style='Quote')
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._format_inline(p, line)
        
        # Tabela final
        if table_rows:
            self._add_table(doc, table_rows)
        
        doc.save(output_path)
        print(f"{Fore.GREEN}✅ DOCX salvo: {output_path}")

    def _format_inline(self, paragraph, text):
        """Formata bold/italic inline"""
        import re
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'
        last_end = 0
        
        for match in re.finditer(pattern, text):
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            if match.group(0).startswith('**'):
                run = paragraph.add_run(match.group(2))
                run.bold = True
            elif match.group(0).startswith('*'):
                run = paragraph.add_run(match.group(3))
                run.italic = True
            
            last_end = match.end()
        
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_table(self, doc, rows):
        """Adiciona tabela formatada"""
        if len(rows) < 2:
            return
        
        max_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < max_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text
                    if i == 0:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True


def main():
    if len(sys.argv) < 2:
        print("=" * 60)
        print("VOMO GEMINI 2.5 FLASH - Script Simplificado")
        print("=" * 60)
        print("\nUso: python vomo_gemini_flash.py <arquivo>")
        print("\nExemplos:")
        print("  python vomo_gemini_flash.py aula.mp3")
        print("  python vomo_gemini_flash.py transcricao.txt")
        print("\nVariável de Ambiente:")
        print("  GEMINI_API_KEY - Chave do Google AI Studio")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    if not os.path.exists(input_file):
        print(f"{Fore.RED}❌ Arquivo não encontrado: {input_file}")
        sys.exit(1)
    
    video_name = Path(input_file).stem
    output_folder = Path(input_file).parent
    
    print("=" * 60)
    print(f"VOMO GEMINI 2.5 FLASH")
    print("=" * 60)
    print(f"📂 Input: {input_file}")
    
    try:
        vomo = VomoGeminiFlash()
        
        # Determina se é áudio ou texto
        if input_file.lower().endswith('.txt'):
            print(f"{Fore.CYAN}📄 Lendo arquivo de texto...")
            with open(input_file, 'r', encoding='utf-8') as f:
                transcription = f.read()
        else:
            # Áudio
            audio = vomo.optimize_audio(input_file)
            transcription = vomo.transcribe(audio)
            
            # Salva transcrição bruta
            raw_path = output_folder / f"{video_name}_RAW.txt"
            with open(raw_path, 'w', encoding='utf-8') as f:
                f.write(transcription)
            print(f"   💾 Transcrição salva: {raw_path}")
        
        # Gera apostila
        apostila = vomo.generate_apostila(transcription, video_name)
        
        # Salva outputs
        md_path = output_folder / f"{video_name}_APOSTILA.md"
        vomo.save_markdown(apostila, md_path)
        
        if DOCX_AVAILABLE:
            docx_path = output_folder / f"{video_name}_APOSTILA.docx"
            vomo.save_docx(apostila, video_name, docx_path)
        
        print("\n" + "=" * 60)
        print(f"{Fore.GREEN}✨ SUCESSO!")
        print("=" * 60)
        
    except Exception as e:
        print(f"{Fore.RED}❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
