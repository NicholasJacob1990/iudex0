#!/usr/bin/env python3
"""
Script usando Gemini DIRETO via Google API (não OpenRouter)
"""

import os
import sys
import time
import google.generativeai as genai
from pathlib import Path
from docx import Document
from typing import List

# Configuração
GEMINI_API_KEY = "AIzaSyBZoYa8UeboO54a5xiUqkt_Qc5gKqFVxqk"
MODEL_NAME = "gemini-2.5-flash"
MAX_CHARS_PER_CHUNK = 5_000
OUTPUT_DIR = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ')

PROMPT_FORMATACAO = """REVISOR DE TEXTO - NÃO RESUMIR!

⛔ REGRA: O texto de saída deve ter tamanho IGUAL OU MAIOR que a entrada!

FAÇA:
✅ Corrija gramática e pontuação
✅ Transforme linguagem coloquial em formal
✅ Organize em parágrafos claros
✅ MANTENHA TODO o conteúdo (exemplos, casos, dicas)
✅ Use **negrito** para termos importantes

NÃO FAÇA:
❌ Não resuma
❌ Não omita informações
❌ Não corte exemplos

[Parte {parte}/{total}]

TEXTO ORIGINAL:
{texto}

TEXTO FORMATADO:"""


def extrair_texto(caminho: str) -> str:
    path = Path(caminho)
    if path.suffix.lower() == '.docx':
        doc = Document(caminho)
        return '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    with open(caminho, 'r', encoding='utf-8') as f:
        return f.read()


def dividir_em_chunks(texto: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> List[str]:
    """Divide texto em chunks, forçando split em parágrafos grandes"""
    chunks = []
    current_chunk = ""
    
    # Split por quebras de linha simples para granularidade
    lines = texto.split('\n')
    
    for line in lines:
        # Se a linha sozinha é maior que max_chars, divide por frases
        if len(line) > max_chars:
            # Salva chunk atual primeiro
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            # Divide linha grande por frases (. ou ; ou :)
            import re
            sentences = re.split(r'(?<=[.;:])\s+', line)
            temp_chunk = ""
            for sentence in sentences:
                if len(temp_chunk) + len(sentence) > max_chars:
                    if temp_chunk:
                        chunks.append(temp_chunk.strip())
                    temp_chunk = sentence
                else:
                    temp_chunk += " " + sentence if temp_chunk else sentence
            if temp_chunk:
                current_chunk = temp_chunk
        elif len(current_chunk) + len(line) + 1 > max_chars:
            chunks.append(current_chunk.strip())
            current_chunk = line
        else:
            current_chunk += "\n" + line if current_chunk else line
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return [c for c in chunks if c]  # Remove chunks vazios


def processar_chunk(model, texto: str, parte: int, total: int) -> str:
    """Processa via Gemini API"""
    
    prompt = PROMPT_FORMATACAO.format(texto=texto, parte=parte, total=total)
    
    print(f"   📝 {parte}/{total}: {len(texto):,} chars", end="", flush=True)
    
    response = model.generate_content(prompt)
    resultado = response.text
    
    ratio = len(resultado) / len(texto) * 100
    status = "✅" if ratio >= 90 else "⚠️"
    print(f" → {len(resultado):,} ({ratio:.0f}%) {status}")
    
    return resultado


def salvar_arquivos(texto: str, nome_base: str):
    """Salva MD e DOCX"""
    md_path = OUTPUT_DIR / f'{nome_base}_GEMINI_FINAL.md'
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(texto)
    print(f"✅ MD: {md_path}")
    
    docx_path = OUTPUT_DIR / f'{nome_base}_GEMINI_FINAL.docx'
    doc = Document()
    for linha in texto.split('\n'):
        if linha.strip():
            doc.add_paragraph(linha)
    doc.save(str(docx_path))
    print(f"✅ DOCX: {docx_path}")


def main():
    arquivo = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/Administrativo.txt')
    
    if not arquivo.exists():
        print(f"❌ Não encontrado: {arquivo}")
        return
    
    print(f"\n📄 Lendo: {arquivo.name}")
    texto = extrair_texto(str(arquivo))
    print(f"   {len(texto):,} caracteres\n")
    
    print(f"🔧 Configurando Gemini API...")
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel(
        MODEL_NAME,
        generation_config={
            "temperature": 0.1,
            "top_p": 0.95,
            "max_output_tokens": 8192,
        }
    )
    
    print(f"📊 Dividindo em chunks de {MAX_CHARS_PER_CHUNK:,} chars")
    chunks = dividir_em_chunks(texto)
    print(f"   {len(chunks)} partes\n")
    
    resultados = []
    
    for i, chunk in enumerate(chunks, 1):
        try:
            resultado = processar_chunk(model, chunk, i, len(chunks))
            resultados.append(resultado)
            
            if i < len(chunks):
                time.sleep(1)  # Rate limiting
                
        except Exception as e:
            print(f" ❌ Erro: {e}")
            resultados.append(chunk)  # Mantém original se falhar
    
    texto_final = '\n\n'.join(resultados)
    
    print(f"\n💾 Salvando...")
    salvar_arquivos(texto_final, arquivo.stem)
    
    ratio = len(texto_final) / len(texto) * 100
    print(f"\n{'='*60}")
    print(f"🎉 CONCLUÍDO!")
    print(f"   Entrada:  {len(texto):,} chars")
    print(f"   Saída:    {len(texto_final):,} chars ({ratio:.0f}%)")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
