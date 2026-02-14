import os
import sys
import time
import argparse
import subprocess
import traceback
import hashlib
import shutil
from urllib.parse import urlparse
from pathlib import Path
from typing import Optional
try:
    from colorama import init, Fore, Style
except ImportError:
    class MockColor:
        def __getattr__(self, name): return ""
    init = lambda *a, **k: None
    Fore = MockColor()
    Style = MockColor()

try:
    from tqdm import tqdm
except ImportError:
    def tqdm(iterable, *args, **kwargs):
        return iterable
import re
import difflib
import json
import asyncio
from google import genai
from google.genai import types
from openai import OpenAI, AsyncOpenAI
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import threading
import random
from collections import deque
from time import sleep # Added for RateLimiter fallback if needed
import logging

# Carrega .env no início do módulo para garantir variáveis disponíveis
try:
    from dotenv import load_dotenv
    load_dotenv(override=False)
except ImportError:
    pass

init(autoreset=True)

try:
    from app.services.api_call_tracker import record_api_call as _record_api_call
except Exception:
    _record_api_call = None


def _safe_int(value):
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _safe_float(value):
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _env_truthy(name: str, default: Optional[bool] = None) -> Optional[bool]:
    value = os.getenv(name)
    if value is None:
        return default
    value_norm = str(value).strip().lower()
    if value_norm in ("1", "true", "yes", "y", "on", "enable", "enabled"):
        return True
    if value_norm in ("0", "false", "no", "n", "off", "disable", "disabled"):
        return False
    return default


def _record_llm_usage(
    *,
    provider: str,
    model: Optional[str],
    tokens_in: Optional[int] = None,
    tokens_out: Optional[int] = None,
    cached_tokens_in: Optional[int] = None,
    seconds_audio: Optional[float] = None,
    seconds_video: Optional[float] = None,
):
    if not _record_api_call or not model:
        return
    meta = {}
    if tokens_in is not None:
        meta["tokens_in"] = int(tokens_in)
        meta["context_tokens"] = int(tokens_in)
    if tokens_out is not None:
        meta["tokens_out"] = int(tokens_out)
    if cached_tokens_in is not None:
        meta["cached_tokens_in"] = int(cached_tokens_in)
    if seconds_audio is not None:
        meta["seconds_audio"] = float(seconds_audio)
    if seconds_video is not None:
        meta["seconds_video"] = float(seconds_video)
    try:
        _record_api_call(kind="llm", provider=provider, model=model, success=True, meta=meta)
    except Exception:
        pass


def _record_openai_usage(response, *, model: Optional[str], provider: str = "openai"):
    usage = getattr(response, "usage", None)
    tokens_in = _safe_int(getattr(usage, "prompt_tokens", None) or getattr(usage, "input_tokens", None))
    tokens_out = _safe_int(getattr(usage, "completion_tokens", None) or getattr(usage, "output_tokens", None))
    details = getattr(usage, "prompt_tokens_details", None) or getattr(usage, "input_tokens_details", None)
    cached_tokens = _safe_int(getattr(details, "cached_tokens", None) or getattr(details, "cached", None))
    _record_llm_usage(
        provider=provider,
        model=model,
        tokens_in=tokens_in,
        tokens_out=tokens_out,
        cached_tokens_in=cached_tokens,
    )


def _record_genai_usage(response, *, model: Optional[str], provider: str = "gemini"):
    usage = getattr(response, "usage_metadata", None)
    tokens_in = _safe_int(
        getattr(usage, "prompt_token_count", None) or getattr(usage, "input_tokens", None)
    )
    tokens_out = _safe_int(
        getattr(usage, "candidates_token_count", None) or getattr(usage, "output_tokens", None)
    )
    cached_tokens = _safe_int(
        getattr(usage, "cached_content_token_count", None)
        or getattr(usage, "cached_token_count", None)
        or getattr(usage, "cached_tokens", None)
    )
    seconds_audio = _safe_float(
        getattr(usage, "prompt_audio_duration_seconds", None)
        or getattr(usage, "audio_duration_seconds", None)
        or getattr(usage, "audio_seconds", None)
    )
    seconds_video = _safe_float(
        getattr(usage, "prompt_video_duration_seconds", None)
        or getattr(usage, "video_duration_seconds", None)
        or getattr(usage, "video_seconds", None)
    )
    _record_llm_usage(
        provider=provider,
        model=model,
        tokens_in=tokens_in,
        tokens_out=tokens_out,
        cached_tokens_in=cached_tokens,
        seconds_audio=seconds_audio,
        seconds_video=seconds_video,
    )

class HILCheckpointException(Exception):
    """Interrompe o pipeline quando a revisão humana é obrigatória."""
    pass

# v2.18: Import do módulo de auditoria jurídica (desativado por padrão)
LEGAL_AUDIT_ENABLED = os.getenv("ENABLE_LEGAL_AUDIT", "").lower() in ("1", "true", "yes", "on")
if LEGAL_AUDIT_ENABLED:
    try:
        from audit_module import auditar_consistencia_legal
        AUDIT_AVAILABLE = True
    except ImportError:
        AUDIT_AVAILABLE = False
        print("⚠️ audit_module não encontrado. Auditoria jurídica desabilitada.")
else:
    AUDIT_AVAILABLE = False

# v2.27: Auditoria preventiva de fidelidade (não limitada à autoria)
FIDELITY_AUDIT_ENABLED = os.getenv("ENABLE_FIDELITY_AUDIT", "1").lower() in ("1", "true", "yes", "on")
# Backup opcional da validação full-context
FIDELITY_BACKUP_ENABLED = os.getenv("ENABLE_FIDELITY_BACKUP", "1").lower() in ("1", "true", "yes", "on")
try:
    from audit_fidelity_preventive import (
        auditar_fidelidade_preventiva,
        gerar_relatorio_markdown_completo,
    )
    FIDELITY_AUDIT_AVAILABLE = True
except ImportError as e:
    FIDELITY_AUDIT_AVAILABLE = False
    print(f"⚠️ audit_fidelity_preventive não encontrado ou erro de importação: {e}. Auditoria preventiva desabilitada.")

# v2.27: Auditoria de fontes integrada (controla inclusão no relatório preventivo)
SOURCES_AUDIT_ENABLED = os.getenv("ENABLE_SOURCES_AUDIT", "1").lower() in ("1", "true", "yes", "on")

# v2.24: Import auto_fix_apostilas for post-processing
try:
    from auto_fix_apostilas import analyze_structural_issues, apply_structural_fixes_to_file
    AUTO_FIX_AVAILABLE = True
except ImportError:
    AUTO_FIX_AVAILABLE = False

# v3.0: Relatório unificado (cross-referencing entre camadas)
try:
    from audit_unified import UnifiedAuditEngine, generate_unified_markdown, UnifiedReport, compare_reports
    UNIFIED_AUDIT_AVAILABLE = True
except ImportError:
    UNIFIED_AUDIT_AVAILABLE = False

try:
    import mlx_whisper
except ImportError:
    mlx_whisper = None

# Faster-Whisper (Beam Search Backend)
try:
    from faster_whisper import WhisperModel
    FASTER_WHISPER_AVAILABLE = True
except ImportError:
    FASTER_WHISPER_AVAILABLE = False

# Imports Pyannote e Torch
try:
    from pyannote.audio import Pipeline
    import torch
    
    HF_TOKEN = os.getenv("HUGGING_FACE_TOKEN")
    if not HF_TOKEN:
        pass
    
except ImportError:
    Pipeline = None
    torch = None

# ==================== CONTEXT CACHING (v2.2) ====================
# Só usa cache para documentos grandes (economia de tokens)
MIN_CHARS_PARA_CACHE = 150000  # 150k chars (~37k tokens)
CACHE_TTL = '7200s'            # 2 horas (padrão)

def criar_cache_contexto(client, transcricao_completa, system_prompt, estrutura_global=None, model_name="gemini-3-flash-preview"):
    """
    v2.2: Cria cache de contexto com hash estável para reutilização.
    
    Args:
        client: Cliente Gemini
        transcricao_completa: Texto completo (usado para calcular TTL dinâmico)
        system_prompt: Prompt de formatação (PROMPT_APOSTILA ou PROMPT_FIDELIDADE)
        estrutura_global: Estrutura mapeada (opcional)
        model_name: Nome do modelo Gemini
    
    Returns:
        Cache object ou None se falhar/não necessário
    """
    # Cache só vale a pena para documentos grandes
    if len(transcricao_completa) < MIN_CHARS_PARA_CACHE:
        print(f"{Fore.YELLOW}📦 Documento pequeno ({len(transcricao_completa):,} chars), cache não necessário{Style.RESET_ALL}")
        return None
    
    try:
        # Hash do prompt + estrutura para garantir unicidade por documento
        combined_content = system_prompt + (estrutura_global or "")
        prompt_hash = hashlib.sha256(combined_content.encode()).hexdigest()[:16]
        cache_name = f"vomo_{prompt_hash}"
        
        # Tenta encontrar cache existente válido
        try:
            for c in client.caches.list(page_size=100):
                if c.display_name == cache_name:
                    print(f"{Fore.GREEN}♻️  Reusando cache existente: {cache_name}{Style.RESET_ALL}")
                    return c
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Cache lookup warning: {e}{Style.RESET_ALL}")

        # Adiciona a estrutura global se disponível
        estrutura_text = f"\n\n## ESTRUTURA GLOBAL (GUIA):\n{estrutura_global}" if estrutura_global else ""
        
        cache_content = f"""{system_prompt}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📚 CONTEXTO GLOBAL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{estrutura_text}
"""
        
        # TTL Dinâmico: 1 hora a cada 500k chars + 1h margem
        tempo_estimado_segundos = int((len(transcricao_completa) / 500000) * 3600) + 3600
        dinamico_ttl = f"{max(3600, tempo_estimado_segundos)}s"
        
        # Cria cache usando a API do google-genai
        cache = client.caches.create(
            model=model_name,
            config=types.CreateCachedContentConfig(
                contents=[cache_content],
                ttl=dinamico_ttl,
                display_name=cache_name
            )
        )
        
        print(f"{Fore.GREEN}✅ Cache criado: {cache_name} (hash: {prompt_hash}, TTL: {dinamico_ttl}){Style.RESET_ALL}")
        return cache
        
    except Exception as e:
        print(f"{Fore.YELLOW}⚠️ Falha ao criar cache: {e}. Continuando sem cache.{Style.RESET_ALL}")
        return None

# ==================== RATE LIMITER ====================
class RateLimiter:
    """Controla requisições por minuto para não estourar rate limit da API"""
    def __init__(self, max_requests_per_minute=60): # Vertex AI Limit
        self.max_rpm = max_requests_per_minute
        self._window_seconds = 60.0
        self._requests_sync = deque()
        self._requests_async = deque()
        self._lock = threading.Lock()
        self._async_lock = asyncio.Lock()

    def _prune_requests(self, requests, now):
        cutoff = now - self._window_seconds
        while requests and requests[0] <= cutoff:
            requests.popleft()
    
    def wait_if_needed(self):
        while True:
            with self._lock:
                now = time.monotonic()
                self._prune_requests(self._requests_sync, now)

                if len(self._requests_sync) < self.max_rpm:
                    self._requests_sync.append(now)
                    return

                oldest = self._requests_sync[0]
                wait_time = self._window_seconds - (now - oldest) + 0.5

            if wait_time > 0:
                print(f"{Fore.YELLOW}⏱️  Rate limit: aguardando {wait_time:.1f}s...")
                sleep(wait_time)

    async def wait_if_needed_async(self):
        """Versão async do rate limiter para não bloquear o event loop"""
        while True:
            async with self._async_lock:
                now = time.monotonic()
                self._prune_requests(self._requests_async, now)

                if len(self._requests_async) < self.max_rpm:
                    self._requests_async.append(now)
                    return

                oldest = self._requests_async[0]
                wait_time = self._window_seconds - (now - oldest) + 0.5

            if wait_time > 0:
                print(f"{Fore.YELLOW}⏱️  Rate limit (Async): aguardando {wait_time:.1f}s...")
                await asyncio.sleep(wait_time)

# Instância global
rate_limiter = RateLimiter(max_requests_per_minute=60)

def remover_overlap_duplicado(resultados, mode="APOSTILA"):
    """Remove duplicação causada pelo overlap entre chunks usando detecção ROBUSTA de conteúdo
    
    v2.17: Usa LIMIAR_7DIFF diferenciado por modo.
    - FIDELIDADE: 0.85 (mais conservador)
    - APOSTILA: 0.80 (mais agressivo - overlaps são quase sempre erros)
    """
    # Limiares adaptativos por camada de deduplicação
    LIMIAR_7DIFF = 0.85 if mode == "FIDELIDADE" else 0.80
    if len(resultados) <= 1:
        return resultados[0] if resultados else ""
    
    import re
    from difflib import SequenceMatcher
    
    # === FUNÇÕES AUXILIARES DA ESTRATÉGIA ROBUSTA (Portadas de clean_redundancy.py) ===
    
    def normalize_text(text):
        if not text: return ""
        text = re.sub(r'[#*-]', ' ', text)
        text = re.sub(r'[^\w\s]', '', text)
        return ' '.join(text.lower().split())

    def calculate_similarity(text1, text2):
        if not text1 or not text2:
            return 0.0
        if len(text1) < 50:
            return 1.0 if text1 in text2 or text2 in text1 else 0.0
        return SequenceMatcher(None, text1, text2).quick_ratio()

    def extract_unique_paragraphs(sec_curr_content, sec_prev_content):
        if not sec_curr_content: return []
        unique = []
        paras_curr = sec_curr_content.split('\n\n')
        paras_prev_norm = [normalize_text(p) for p in sec_prev_content.split('\n\n')]
        
        for p in paras_curr:
            p_clean = p.strip()
            if not p_clean or len(p_clean) < 20: continue
            
            p_norm = normalize_text(p_clean)
            is_present = False
            for pp_norm in paras_prev_norm:
                if calculate_similarity(p_norm, pp_norm) > 0.85:
                    is_present = True
                    break
            
            if not is_present:
                unique.append(p_clean)
        return unique
    
    # ==================================================================================

    print("🧹 Iniciando deduplicação robusta (7-DIFF Strategy)...")

    # 1. Junta e Parseia
    texto_bruto = '\n\n'.join(resultados)
    lines = texto_bruto.split('\n')
    
    sections = []
    current_section = None
    header_pattern = re.compile(r'^(#{1,6})\s+(.+)$')
    
    captured_lines = []
    intro_lines = []
    has_started = False
    
    for i, line in enumerate(lines):
        match = header_pattern.match(line)
        if match:
            has_started = True
            if current_section:
                current_section['content'] = '\n'.join(captured_lines).strip()
                sections.append(current_section)
                captured_lines = []
            
            title_text = match.group(2).strip()
            # Remove numeração original para comparação agnóstica
            title_clean = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', title_text)
            
            current_section = {
                'title_clean': title_clean,
                'level': len(match.group(1)),
                'full_header': line,
                'content': ""
            }
        else:
            if has_started:
                captured_lines.append(line)
            else:
                intro_lines.append(line)
    
    if current_section:
        current_section['content'] = '\n'.join(captured_lines).strip()
        sections.append(current_section)
        
    print(f"   📊 Analisando {len(sections)} seções...")

    # 2. Detecção e Remoção
    indices_to_remove = set()
    MAX_WINDOW = 20 # Olha até 20 seções para trás (cobre overlaps grandes)
    
    for i in range(len(sections)):
        if i in indices_to_remove: continue
        sec_curr = sections[i]
        
        # Janela deslizante
        start_check = max(0, i - MAX_WINDOW)
        for j in range(start_check, i):
            if j in indices_to_remove: continue
            sec_prev = sections[j]
            
            if sec_curr['level'] != sec_prev['level']: continue
            
            sim_title = calculate_similarity(normalize_text(sec_curr['title_clean']), normalize_text(sec_prev['title_clean']))
            sim_content = calculate_similarity(normalize_text(sec_curr['content']), normalize_text(sec_prev['content']))
            
            is_duplicate = False
            
            if sim_title > 0.9 and sim_content > 0.6:
                is_duplicate = True
            elif sim_content > LIMIAR_7DIFF:  # Usa limiar 7-DIFF (0.85 Fidelidade / 0.80 Apostila)
                 is_duplicate = True
            elif sim_title > 0.95 and len(sec_curr['content']) < 100:
                 is_duplicate = True
            
            if is_duplicate:
                print(f"   🗑️  Detectado: '{sec_curr['title_clean'][:30]}...' duplica seção anterior")
                
                # Mescla conteúdo único antes de excluir
                unique_paras = extract_unique_paragraphs(sec_curr['content'], sec_prev['content'])
                if unique_paras:
                    sections[j]['content'] += '\n\n' + '\n\n'.join(unique_paras)
                
                indices_to_remove.add(i)
                break
    
    # 3. Reconstrução
    final_lines = list(intro_lines)
    for i, sec in enumerate(sections):
        if i in indices_to_remove: continue
        final_lines.append(sec['full_header'])
        if sec['content']:
            final_lines.append(sec['content'])
            final_lines.append("")
            
    texto_limpo = '\n'.join(final_lines)
    print(f"   ✅ Removidas {len(indices_to_remove)} seções duplicadas.")

    return texto_limpo


# ==================== v2.28: VALIDAÇÃO E SANITIZAÇÃO DE MARKDOWN ====================

class TruncamentoError(Exception):
    """Exceção levantada quando truncamento crítico é detectado."""
    pass


def corrigir_headings_duplicados(texto: str) -> str:
    """
    v2.28: Corrige headings duplicados como '#### #### Título' → '#### Título'

    Também normaliza variações como '## ## #' → '##'
    """
    # Padrão: múltiplos grupos de # separados por espaços (ex: "#### #### Título")
    # Importante: exige pelo menos 1 espaço entre os grupos para não degradar headings normais ("#### Título").
    pattern = r'^(#{1,6})(?:\s+#{1,6})+\s*(.*)$'

    def fix_heading(match):
        level = match.group(1)  # Primeiro conjunto de #
        title = (match.group(2) or "").strip()
        return f"{level} {title}" if title else level

    linhas = texto.split('\n')
    linhas_corrigidas = []
    correcoes = 0

    for linha in linhas:
        if re.match(r'^#{1,6}\s+#{1,6}', linha):
            linha_corrigida = re.sub(pattern, fix_heading, linha)
            if linha_corrigida != linha:
                correcoes += 1
                print(f"{Fore.YELLOW}   🔧 Heading corrigido: '{linha[:50]}...' → '{linha_corrigida[:50]}...'")
            linhas_corrigidas.append(linha_corrigida)
        else:
            linhas_corrigidas.append(linha)

    if correcoes > 0:
        print(f"{Fore.GREEN}   ✅ Corrigidos {correcoes} headings duplicados")

    return '\n'.join(linhas_corrigidas)


def padronizar_separadores(texto: str, estilo: str = "remover") -> str:
    """
    v2.28: Padroniza separadores horizontais (---, ***, ___).

    Args:
        texto: Texto markdown
        estilo: 'remover' (remove todos), 'padronizar' (usa --- apenas), 'manter' (não altera)

    Returns:
        Texto com separadores padronizados
    """
    if estilo == "manter":
        return texto

    # Padrão: linha contendo apenas hífens, asteriscos ou underscores (3+)
    pattern = r'^[\s]*[-*_]{3,}[\s]*$'

    linhas = texto.split('\n')
    linhas_novas = []
    removidos = 0

    for linha in linhas:
        if re.match(pattern, linha):
            if estilo == "remover":
                removidos += 1
                continue  # Pula a linha
            elif estilo == "padronizar":
                linhas_novas.append("---")
                continue
        linhas_novas.append(linha)

    if removidos > 0:
        print(f"{Fore.CYAN}   🔧 Removidos {removidos} separadores horizontais")

    return '\n'.join(linhas_novas)


def detectar_tabelas_em_par(texto: str) -> list:
    """
    v2.28: Detecta pares de tabelas (Quadro-síntese + Pegadinhas).

    Padrão esperado (flexível):
    - #### 📋 [título contextual]
    - Tabela 5 colunas
    - #### 🎯 [título contextual]
    - Tabela 3 colunas

    Returns:
        Lista de dicts com informações sobre cada par de tabelas
    """
    pares = []
    linhas = texto.split('\n')

    i = 0
    while i < len(linhas):
        linha = linhas[i].strip()

        # Detectar início de quadro-síntese / tabela principal (título contextual com 📋)
        if re.match(r'^#{3,5}\s*📋', linha):
            par = {
                'quadro_titulo': linha,
                'quadro_linha': i,
                'quadro_tabela_inicio': None,
                'quadro_tabela_linhas': 0,
                'pegadinha_titulo': None,
                'pegadinha_linha': None,
                'pegadinha_tabela_inicio': None,
                'pegadinha_tabela_linhas': 0,
                'completo': False
            }

            # Procurar tabela do quadro
            j = i + 1
            while j < len(linhas) and j < i + 20:
                if linhas[j].strip().startswith('|'):
                    par['quadro_tabela_inicio'] = j
                    # Contar linhas da tabela
                    k = j
                    while k < len(linhas) and linhas[k].strip().startswith('|'):
                        par['quadro_tabela_linhas'] += 1
                        k += 1
                    break
                elif re.match(r'^#{1,5}\s', linhas[j]):
                    break  # Novo heading, tabela ausente
                j += 1

            # Procurar tabela de pegadinhas
            j = (
                par['quadro_tabela_inicio'] + par['quadro_tabela_linhas']
                if par['quadro_tabela_inicio'] is not None
                else i + 1
            )
            while j < len(linhas) and j < i + 50:
                if re.match(r'^#{3,5}\s*🎯', linhas[j]):
                    par['pegadinha_titulo'] = linhas[j].strip()
                    par['pegadinha_linha'] = j

                    # Procurar tabela de pegadinhas
                    k = j + 1
                    while k < len(linhas) and k < j + 15:
                        if linhas[k].strip().startswith('|'):
                            par['pegadinha_tabela_inicio'] = k
                            m = k
                            while m < len(linhas) and linhas[m].strip().startswith('|'):
                                par['pegadinha_tabela_linhas'] += 1
                                m += 1
                            break
                        k += 1
                    break
                elif re.match(r'^#{1,2}\s', linhas[j]):
                    break  # Novo bloco temático
                j += 1

            # Verificar se par está completo
            par['completo'] = (
                par['quadro_tabela_linhas'] >= 3 and  # Header + separador + pelo menos 1 dado
                par['pegadinha_tabela_linhas'] >= 3
            )

            pares.append(par)
            i = j if j > i else i + 1
        else:
            i += 1

    # Log de diagnóstico
    completos = sum(1 for p in pares if p['completo'])
    print(f"{Fore.CYAN}   📊 Pares de tabelas detectados: {len(pares)} ({completos} completos)")

    for p in pares:
        if not p['completo']:
            print(f"{Fore.YELLOW}   ⚠️  Par incompleto: {p['quadro_titulo'][:40]}... "
                  f"(Quadro: {p['quadro_tabela_linhas']} linhas, Pegadinha: {p['pegadinha_tabela_linhas']} linhas)")

    return pares


def validar_celulas_tabela(texto: str) -> tuple:
    """
    v2.28: Valida integridade das células de tabela.

    Detecta:
    1. Células truncadas (texto cortado no meio de palavra)
    2. Headers incompletos (ex: 'Comcobra' em vez de 'Como a banca cobra')
    3. Linhas de tabela sem fechamento de pipe

    Returns:
        Tuple (is_valid, list of issues)
    """
    issues = []
    linhas = texto.split('\n')

    # Padrões conhecidos de truncamento
    TRUNCAMENTOS_CONHECIDOS = [
        (r'\bonto\b', 'truncamento de "o território" ou similar'),
        (r'Comcobra', 'header truncado: "Como a banca cobra"'),
        (r'urbanístamos', 'palavra cortada: "urbanístico. Vamos"'),
        (r'\bsitua\s+competência', 'frase cortada'),
        (r'\bEls\s+sobre', 'início de frase cortada'),
        (r'\bou\s+[a-z]{1,3}\s+[A-Z]', 'possível corte no meio de frase'),
    ]

    for i, linha in enumerate(linhas):
        # Verificar padrões de truncamento
        for pattern, desc in TRUNCAMENTOS_CONHECIDOS:
            if re.search(pattern, linha):
                issues.append({
                    'tipo': 'truncamento',
                    'linha': i + 1,
                    'descricao': desc,
                    'texto': linha[:100] + '...' if len(linha) > 100 else linha
                })

        # Verificar células de tabela
        if linha.strip().startswith('|'):
            # Linha de tabela deve terminar com |
            if not linha.strip().endswith('|'):
                issues.append({
                    'tipo': 'tabela_aberta',
                    'linha': i + 1,
                    'descricao': 'Linha de tabela não fechada com |',
                    'texto': linha[-50:] if len(linha) > 50 else linha
                })

            # Verificar células muito curtas (possível truncamento)
            celulas = linha.split('|')[1:-1]  # Remove primeiro e último vazio
            for j, celula in enumerate(celulas):
                celula_limpa = celula.strip()
                # Célula com menos de 3 chars e não é separador pode ser truncamento
                if len(celula_limpa) < 3 and not re.match(r'^[-:]+$', celula_limpa) and celula_limpa != '—':
                    issues.append({
                        'tipo': 'celula_suspeita',
                        'linha': i + 1,
                        'descricao': f'Célula {j+1} muito curta: "{celula_limpa}"',
                        'texto': linha[:80]
                    })

    is_valid = len(issues) == 0

    if not is_valid:
        print(f"{Fore.RED}   ⚠️  Encontrados {len(issues)} problemas de integridade:")
        for issue in issues[:5]:  # Mostrar no máximo 5
            print(f"{Fore.YELLOW}      L{issue['linha']}: {issue['descricao']}")
        if len(issues) > 5:
            print(f"{Fore.YELLOW}      ... e mais {len(issues) - 5} problemas")

    return is_valid, issues


def chunk_texto_seguro(texto: str, max_chars: int = 25000, overlap_chars: int = 2000) -> list:
    """
    v2.28: Chunking inteligente que respeita limites naturais do texto.

    Prioridades de corte (em ordem):
    1. Antes de heading ## ou ### (novo bloco temático)
    2. Após tabela completa (#### 🎯 + tabela)
    3. Parágrafo duplo (\\n\\n)
    4. Final de frase (. seguido de espaço ou newline)
    5. Qualquer newline

    Nunca corta:
    - No meio de uma tabela
    - No meio de uma palavra
    - Imediatamente após heading (deixa pelo menos 500 chars)

    Args:
        texto: Texto completo
        max_chars: Tamanho máximo de cada chunk
        overlap_chars: Caracteres de overlap entre chunks

    Returns:
        Lista de chunks com integridade preservada
    """
    if len(texto) <= max_chars:
        return [texto]

    chunks = []
    inicio = 0

    # Pré-processar: identificar zonas "proibidas" para corte
    zonas_proibidas = []  # Lista de (inicio, fim) onde não cortar

    # Encontrar todas as tabelas
    linhas = texto.split('\n')
    pos = 0
    em_tabela = False
    tabela_inicio = 0

    for i, linha in enumerate(linhas):
        if linha.strip().startswith('|') and not em_tabela:
            em_tabela = True
            tabela_inicio = pos
        elif not linha.strip().startswith('|') and em_tabela:
            em_tabela = False
            zonas_proibidas.append((tabela_inicio, pos))
        pos += len(linha) + 1  # +1 pelo \n
    if em_tabela:
        zonas_proibidas.append((tabela_inicio, pos))

    def esta_em_zona_proibida(posicao):
        for inicio_z, fim_z in zonas_proibidas:
            if inicio_z <= posicao <= fim_z:
                return True
        return False

    def encontrar_ponto_corte_seguro(texto_slice, pos_inicio_global):
        """Encontra o melhor ponto de corte dentro do slice."""

        # Zona de busca: últimos 30% do chunk
        zona_busca_inicio = int(len(texto_slice) * 0.7)
        zona_busca = texto_slice[zona_busca_inicio:]

        # Prioridade 1: Antes de heading ## ou ###
        headings = list(re.finditer(r'(?m)^#{2,3}\s+', zona_busca))
        if headings:
            pos = headings[-1].start()
            pos_global = pos_inicio_global + zona_busca_inicio + pos
            if not esta_em_zona_proibida(pos_global):
                return zona_busca_inicio + pos

        # Prioridade 2: Após tabela de pegadinhas (🎯)
        match = re.search(r'\n(?=####?\s*🎯)', zona_busca)
        if match:
            # Encontrar fim da tabela após o heading
            after_heading = zona_busca[match.end():]
            # Procurar fim da tabela (linha que não começa com |)
            lines_after = after_heading.split('\n')
            pos_apos_tabela = match.end()
            for j, line in enumerate(lines_after):
                if j > 2 and not line.strip().startswith('|'):  # Passou da tabela
                    pos_apos_tabela += sum(len(l) + 1 for l in lines_after[:j])
                    break
            pos_global = pos_inicio_global + zona_busca_inicio + pos_apos_tabela
            if not esta_em_zona_proibida(pos_global) and pos_apos_tabela < len(zona_busca):
                return zona_busca_inicio + pos_apos_tabela

        # Prioridade 3: Parágrafo duplo
        pos = zona_busca.rfind('\n\n')
        if pos != -1:
            pos_global = pos_inicio_global + zona_busca_inicio + pos
            if not esta_em_zona_proibida(pos_global):
                return zona_busca_inicio + pos + 2  # +2 para incluir os \n\n

        # Prioridade 4: Final de frase
        finais_frase = list(re.finditer(r'[.!?][\s\n]+', zona_busca))
        if finais_frase:
            pos = finais_frase[-1].end()
            pos_global = pos_inicio_global + zona_busca_inicio + pos
            if not esta_em_zona_proibida(pos_global):
                return zona_busca_inicio + pos

        # Prioridade 5: Qualquer newline
        pos = zona_busca.rfind('\n')
        if pos != -1:
            pos_global = pos_inicio_global + zona_busca_inicio + pos
            if not esta_em_zona_proibida(pos_global):
                return zona_busca_inicio + pos + 1

        # Fallback: cortar no max_chars mesmo
        return len(texto_slice)

    print(f"{Fore.CYAN}   🔪 Iniciando chunking seguro (max: {max_chars} chars, overlap: {overlap_chars})...")

    while inicio < len(texto):
        fim_ideal = min(inicio + max_chars, len(texto))

        if fim_ideal >= len(texto):
            # Último chunk
            chunks.append(texto[inicio:])
            break

        texto_slice = texto[inicio:fim_ideal]
        ponto_corte_relativo = encontrar_ponto_corte_seguro(texto_slice, inicio)
        fim_real = inicio + ponto_corte_relativo

        chunk = texto[inicio:fim_real].strip()
        chunks.append(chunk)

        print(f"{Fore.GREEN}   ✂️  Chunk {len(chunks)}: {inicio} → {fim_real} ({len(chunk)} chars)")

        # Próximo início com overlap
        inicio = fim_real - overlap_chars if fim_real > overlap_chars else fim_real

    print(f"{Fore.GREEN}   ✅ Criados {len(chunks)} chunks com integridade preservada")

    return chunks


SEGMENT_BOUNDARY_RE = re.compile(
    r'(?m)(?=^\s*(?:'
    r'\[\d{1,2}:\d{2}(?::\d{2})?\]\s+'
    r'|\*\*[^*]{1,40}\*\*:\s+'
    r'|(?:SPEAKER|FALANTE)\s*\d{1,3}\s*[:\-]'
    r'))'
)


def _segmentar_texto_para_mapeamento(texto: str) -> list[str]:
    """Segmenta o texto em blocos naturais (timestamps/speaker labels) para mapeamento."""
    if not texto:
        return []
    if not SEGMENT_BOUNDARY_RE.search(texto):
        return []
    partes = [p for p in re.split(SEGMENT_BOUNDARY_RE, texto) if p and p.strip()]
    return partes


def chunk_texto_por_segmentos(
    texto: str,
    *,
    max_chars: int = 25000,
    overlap_chars: int = 2000,
    min_segments: int = 3,
) -> Optional[list[str]]:
    """
    Chunking baseado em segmentos (timestamps/speaker labels), evitando cortes artificiais.
    Retorna None quando não há segmentos suficientes.
    """
    segmentos = _segmentar_texto_para_mapeamento(texto)
    if not segmentos or len(segmentos) < min_segments:
        return None

    chunks: list[str] = []
    cur: list[str] = []
    cur_len = 0

    def _flush_current():
        if not cur:
            return
        chunks.append("\n\n".join(cur).strip())

    def _build_overlap(prev_segments: list[str]) -> list[str]:
        if not prev_segments or overlap_chars <= 0:
            return []
        overlap_list: list[str] = []
        acc = 0
        for seg in reversed(prev_segments):
            seg_len = len(seg) + (2 if overlap_list else 0)
            acc += seg_len
            overlap_list.insert(0, seg)
            if acc >= overlap_chars:
                break
        return overlap_list

    for seg in segmentos:
        seg = seg.strip()
        if not seg:
            continue
        seg_len = len(seg) + (2 if cur else 0)
        if cur and (cur_len + seg_len) > max_chars:
            _flush_current()
            cur = _build_overlap(cur)
            cur_len = sum(len(s) for s in cur) + max(0, len(cur) - 1) * 2
        if not cur:
            cur = [seg]
            cur_len = len(seg)
        else:
            cur.append(seg)
            cur_len += seg_len

    _flush_current()
    return [c for c in chunks if c]


def validar_integridade_pos_merge(texto: str, raise_on_error: bool = False) -> tuple:
    """
    v2.28: Validação completa de integridade após merge de chunks.

    Verifica:
    1. Palavras cortadas no meio
    2. Headers incompletos
    3. Tabelas abertas (sem fechamento)
    4. Padrões conhecidos de truncamento

    Args:
        texto: Texto merged
        raise_on_error: Se True, levanta TruncamentoError em problemas críticos

    Returns:
        Tuple (is_valid, issues, texto_corrigido)
    """
    issues = []
    texto_corrigido = texto

    # 1. Validar células de tabela
    is_table_valid, table_issues = validar_celulas_tabela(texto)
    issues.extend(table_issues)

    # 2. Verificar headings duplicados
    if re.search(r'^#{1,6}\s*#{1,6}', texto, re.MULTILINE):
        texto_corrigido = corrigir_headings_duplicados(texto_corrigido)
        issues.append({
            'tipo': 'heading_duplicado',
            'descricao': 'Headings duplicados encontrados e corrigidos',
            'linha': 0
        })

    # 3. Detectar pares de tabelas incompletos
    pares = detectar_tabelas_em_par(texto)
    incompletos = [p for p in pares if not p['completo']]
    for p in incompletos:
        issues.append({
            'tipo': 'par_tabela_incompleto',
            'descricao': f"Par incompleto: {p['quadro_titulo'][:40]}",
            'linha': p['quadro_linha']
        })

    # 4. Padrões de texto truncado (mais genéricos)
    pattern_truncado = r'\b([a-záéíóúàâãêîôûç]{2,})\s{2,}([a-záéíóúàâãêîôûç]{2,})\b'
    matches = list(re.finditer(pattern_truncado, texto, re.IGNORECASE))
    for match in matches[:5]:  # Limitar a 5
        # Verificar se parece truncamento (palavras soltas)
        antes = match.group(1)
        depois = match.group(2)
        if len(antes) < 6 and len(depois) < 6:
            issues.append({
                'tipo': 'possivel_truncamento',
                'descricao': f'Possível corte: "{antes} {depois}"',
                'linha': texto[:match.start()].count('\n') + 1
            })

    is_valid = len([i for i in issues if i['tipo'] in ['truncamento', 'tabela_aberta']]) == 0

    if raise_on_error and not is_valid:
        criticos = [i for i in issues if i['tipo'] in ['truncamento', 'tabela_aberta']]
        raise TruncamentoError(f"Detectados {len(criticos)} problemas críticos de truncamento")

    return is_valid, issues, texto_corrigido


def remover_marcadores_continua(texto: str) -> str:
    """
    Remove marcadores artificiais de continuação inseridos pelo LLM.

    Exemplos removidos:
    - [continua], [continuação], [continuacao]
    - (continua), (continuação), (continuacao)
    """
    if not texto:
        return texto

    # Linha isolada com marcador
    out = re.sub(
        r"(?im)^[ \t]*(?:\[\s*(?:continua|continuação|continuacao)\s*\]|\(\s*(?:continua|continuação|continuacao)\s*\))[ \t]*\n",
        "",
        texto,
    )

    # Marcador inline (substitui por um espaço)
    out = re.sub(
        r"(?i)\s*(?:\[\s*(?:continua|continuação|continuacao)\s*\]|\(\s*(?:continua|continuação|continuacao)\s*\))\s*",
        " ",
        out,
    )

    # Normalizar espaços múltiplos gerados pela remoção
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out


def sanitizar_markdown_final(texto: str) -> str:
    """
    v2.28: Sanitização final do markdown antes de salvar.

    Aplica todas as correções em sequência:
    1. Corrige headings duplicados
    2. Padroniza separadores
    3. Remove linhas em branco excessivas
    4. Valida integridade (sem raise)

    Returns:
        Texto sanitizado
    """
    print(f"{Fore.CYAN}🧹 Sanitizando markdown final...")

    # 1. Headings duplicados
    texto = corrigir_headings_duplicados(texto)

    # 2. Separadores
    texto = padronizar_separadores(texto, estilo="remover")

    # 3. Linhas em branco excessivas (mais de 2 consecutivas → 2)
    texto = re.sub(r'\n{4,}', '\n\n\n', texto)

    # 3.5 Remover vocativos/gírias em forma de chamada ("Meu irmão,", "cara,", ...)
    texto = remover_vocativos_girias(texto)

    # 3.6 Normalizar referências "Tema" frequentemente erradas por ASR (ex.: 234→1.234, 1933→1.033)
    texto = normalizar_temas_markdown(texto)

    # 3.7 Remover marcadores artificiais de continuação (ex.: "[continua]")
    texto = remover_marcadores_continua(texto)

    # 4. Validação (sem raise, apenas log)
    is_valid, issues, texto = validar_integridade_pos_merge(texto, raise_on_error=False)

    if is_valid:
        print(f"{Fore.GREEN}   ✅ Markdown validado sem problemas críticos")
    else:
        print(f"{Fore.YELLOW}   ⚠️  {len(issues)} issues encontradas (não-críticas mantidas)")

    return texto


def normalizar_temas_markdown(texto: str) -> str:
    """
    Normaliza variações comuns de "Tema" geradas por ASR/edição que criam referências inexistentes.

    Regras (conservadoras):
    - Só corrige quando o documento já contém a forma canônica.
    - Remove parênteses/apostos que preservam variantes erradas (ex.: "(ou 1933)").
    """
    if not texto:
        return texto

    out = texto

    def _has(pattern: str) -> bool:
        try:
            return re.search(pattern, out, flags=re.IGNORECASE) is not None
        except re.error:
            return False

    # 234 -> 1.234 (apenas se "Tema 1.234" já aparece no documento)
    if _has(r"\b[Tt]ema\s+1\.234\b"):
        out = re.sub(r"\b([Tt]ema)\s+234\b", r"\1 1.234", out)
        # Remove explicações que mantêm a variante errada (somente casos típicos: "(ou 234)", "(tema 234)")
        out = re.sub(r"\s*\(\s*(?:ou\s+)?(?:tema\s+)?234\s*\)", "", out, flags=re.IGNORECASE)

    # 1933 / 1.933 -> 1.033 (apenas se "Tema 1.033" já aparece no documento)
    if _has(r"\b[Tt]ema\s+1\.033\b"):
        # Normalize numeric variants
        out = re.sub(r"\b([Tt]ema)\s+1933\b", r"\1 1.033", out)
        out = re.sub(r"\b([Tt]ema)\s+1\.933\b", r"\1 1.033", out)
        out = re.sub(r"\b([Tt]ema)\s+1\s*933\b", r"\1 1.033", out)
        # Remove parenthetical that keeps wrong aliases (somente casos típicos: "(ou 1933)", "(tema 1933)")
        out = re.sub(r"\s*\(\s*(?:ou\s+)?(?:tema\s+)?(?:1933|1\.933)\s*\)", "", out, flags=re.IGNORECASE)
        # Normalize table-style combos: "Tema 1.033 / 1933" etc.
        out = re.sub(r"\b(Tema\s+1\.033)\s*/\s*(?:1933|1\.933)\b", r"\1", out, flags=re.IGNORECASE)
        out = re.sub(r"\b(Tema\s+1\.033)\s*\(\s*ou\s*(?:1933|1\.933)\s*\)", r"\1", out, flags=re.IGNORECASE)

    return out


# ==================== HELPERS PORTED FROM GPT SCRIPT ====================

def limpar_tags_xml(texto):
    texto = re.sub(r'</?[a-z_][\w\-]*>', '', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<[a-z_][\w\-]*\s+[^>]+>', '', texto, flags=re.IGNORECASE)
    return texto


def remover_vocativos_girias(texto: str) -> str:
    """
    Remove vocativos/gírias comuns que não agregam conteúdo e não devem constar no formatado.
    Ex.: "Meu irmão,", "cara,", "mano!", "minha gente:" etc.

    Observação: aplica apenas em texto fora de code fences e com pontuação típica de vocativo,
    para evitar apagar informação factual ("meu irmão" como parentesco) quando não estiver em forma de vocativo.
    """
    if not texto:
        return texto

    vocativos = [
        r"meu\s+irm[aã]o",
        r"mano",
        r"cara",
        r"minha\s+gente",
        r"galera",
        r"meu\s+velho",
    ]
    voc = "|".join(vocativos)
    # Start-of-line vocative: "Meu irmão, ..."
    re_start = re.compile(rf"^(\s*)(?:{voc})\s*[,!?:;\-–—]+\s*", flags=re.IGNORECASE)
    # Mid-line after sentence boundary: ". Meu irmão, ..."
    re_mid = re.compile(rf"([.!?:;])\s+(?:{voc})\s*[,!?:;\-–—]+\s*", flags=re.IGNORECASE)

    out_lines = []
    in_fence = False
    for line in texto.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_fence = not in_fence
            out_lines.append(line)
            continue
        if in_fence:
            out_lines.append(line)
            continue
        new_line = re_start.sub(r"\1", line)
        new_line = re_mid.sub(r"\1 ", new_line)
        out_lines.append(new_line)
    return "\n".join(out_lines)

# ==================== LOGGER SHIM ====================
class Logger:
    def info(self, msg):
        print(f"{Fore.CYAN}{msg}")
    def warning(self, msg):
        print(f"{Fore.YELLOW}{msg}")
    def error(self, msg):
        print(f"{Fore.RED}{msg}")

logger = Logger()

# ==================== FUNCIONALIDADES PORTADAS DO GEMINI SCRIPT ====================

# ==================== SMART STITCHING & FIDELIDADE (V2.10 PROTADAS) ====================

def remover_eco_do_contexto(resposta_api, contexto_enviado):
    """
    Remove o início da resposta se for apenas um 'eco' do final do contexto.
    """
    if not contexto_enviado or not resposta_api:
        return resposta_api

    final_contexto = contexto_enviado.strip()[-300:]
    inicio_resposta = resposta_api.strip()[:300]

    matcher = difflib.SequenceMatcher(None, final_contexto, inicio_resposta)
    match = matcher.find_longest_match(0, len(final_contexto), 0, len(inicio_resposta))

    if match.size > 50:
        print(f"Scissors Eco detectado! Removendo {match.size} chars repetidos no início.")
        return resposta_api.strip()[match.size:].strip()
    
    return resposta_api


def _extract_style_context(text: str, max_chars: int = 2500) -> str:
    """
    v2.27: Extrai um contexto de estilo maior, INCLUINDO tabelas recentes para continuidade.
    
    Melhorias v2.27:
    - Se houver tabela nas últimas 100 linhas, inclui a tabela COMPLETA no contexto
    - Permite contexto maior (até 1.5x max_chars) quando há tabela
    - Garante que o LLM veja a estrutura da tabela do chunk anterior
    """
    if not text:
        return ""

    lines = text.splitlines()
    
    # v2.27: Detectar se há tabela recente que deve ser incluída
    last_table_start = None
    last_table_end = None
    search_range = min(100, len(lines))
    
    for i in range(len(lines) - 1, max(0, len(lines) - search_range), -1):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|') and '---' not in line:
            if last_table_end is None:
                last_table_end = i
            last_table_start = i
        elif last_table_start is not None and not line.startswith('|'):
            # Encontramos o início da tabela (linha antes não é tabela)
            break
    
    # Se há tabela recente, incluir ela completa no contexto
    if last_table_start is not None and last_table_end is not None:
        # Incluir algumas linhas antes da tabela (título/contexto)
        table_context_start = max(0, last_table_start - 5)
        table_with_context = '\n'.join(lines[table_context_start:])
        
        # Permitir contexto maior se tiver tabela (até 1.5x)
        extended_max = int(max_chars * 1.5)
        if len(table_with_context) <= extended_max:
            return table_with_context
        else:
            # Tabela é muito grande, pegar só as últimas linhas dela
            return table_with_context[-extended_max:]
    
    # Fallback: comportamento original (filtrar tabelas do contexto de estilo)
    filtered = []
    for ln in lines:
        s = ln.strip()
        # Remove linhas de tabela markdown e separadores
        if (s.startswith('|') and s.endswith('|')) or re.match(r'^\s*\|[\s:|-]+\|[\s:|-]*$', s):
            continue
        # Remove títulos de quadros/tabelas no contexto (para não "puxar" só o fechamento)
        if re.match(r'^#{3,5}\s*[📋🎯].*', s):
            continue
        filtered.append(ln)

    filtered_text = "\n".join(filtered).strip()
    candidate = filtered_text[-max_chars:] if len(filtered_text) > max_chars else filtered_text
    # Fallback: se filtrou demais, usa o fim do texto original
    if len(candidate.split()) < 30:
        candidate = text[-max_chars:] if len(text) > max_chars else text
    return candidate

def titulos_sao_similares(t1, t2, threshold=0.90):
    """Verifica se dois títulos são semanticamente iguais (fuzzy matching)."""
    def normalizar(t):
        return re.sub(r'[^a-z0-9 ]', '', t.lower())
    
    nt1 = normalizar(t1)
    nt2 = normalizar(t2)
    
    if not nt1 or not nt2: return False
    
    # PROTEÇÃO 1: Diferença de tamanho
    nt1_compact = nt1.replace(' ', '')
    nt2_compact = nt2.replace(' ', '')
    len_ratio = min(len(nt1_compact), len(nt2_compact)) / max(len(nt1_compact), len(nt2_compact))
    if len_ratio < 0.8: return False
    
    # PROTEÇÃO 2: Verificação por palavras exclusivas
    palavras1 = set(nt1.split())
    palavras2 = set(nt2.split())
    diferenca = palavras1.symmetric_difference(palavras2)
    if any(len(w) > 3 for w in diferenca): return False
        
    return difflib.SequenceMatcher(None, nt1_compact, nt2_compact).ratio() > threshold

def limpar_inicio_redundante(texto_novo, texto_acumulado):
    """Remove título no início do novo chunk se similar ao do acumulado."""
    if not texto_acumulado.strip(): return texto_novo

    ultimas_linhas = texto_acumulado.strip().split('\n')[-30:]
    ultimo_titulo = None
    for linha in reversed(ultimas_linhas):
        if linha.strip().startswith('##'):
            ultimo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
            break
    
    if not ultimo_titulo: return texto_novo

    linhas_novas = texto_novo.strip().split('\n')
    for i, linha in enumerate(linhas_novas[:10]):
        if linha.strip().startswith('##'):
            novo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
            if titulos_sao_similares(ultimo_titulo, novo_titulo):
                print(f"Scissors Título duplicado na junção: '{novo_titulo}' ≈ '{ultimo_titulo}'")
                return '\n'.join(linhas_novas[i+1:])
    return texto_novo

def detectar_secoes_duplicadas(texto):
    """v2.11: Detecta seções duplicadas por títulos (Fuzzy Matching) - Inclui H3 e normaliza '(Continuação)'"""
    print("Magnifying glass tilt left Detectando seções duplicadas (fuzzy)...")
    
    linhas = texto.split('\n')
    titulos_vistos = []
    secoes_duplicadas = []
    
    for i, linha in enumerate(linhas):
        linha_strip = linha.strip()
        # Detecta H2 (##) e H3 (###)
        if linha_strip.startswith('##'):
            # Remove numeração e emojis
            titulo_normalizado = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', linha_strip)
            titulo_normalizado = re.sub(r'[📋📊🗂]', '', titulo_normalizado).strip()
            # Remove "(Continuação)" para comparação
            titulo_para_comparar = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_normalizado, flags=re.IGNORECASE).strip()
            
            duplicado = False
            for t_visto, linha_visto in titulos_vistos:
                if titulos_sao_similares(titulo_para_comparar, t_visto):
                    print(f"Warning Duplicado (fuzzy): '{linha_strip[:50]}...' ≈ '{t_visto[:50]}...'")
                    secoes_duplicadas.append({
                        'titulo': titulo_normalizado,
                        'primeira_linha': linha_visto,
                        'duplicada_linha': i
                    })
                    duplicado = True
                    break
            
            if not duplicado:
                titulos_vistos.append((titulo_para_comparar, i))
    
    if secoes_duplicadas:
        print(f"Cross mark {len(secoes_duplicadas)} seções duplicadas detectadas!")
    else:
        print("Check mark button Nenhuma seção duplicada detectada")
    
    return secoes_duplicadas

def remover_secoes_duplicadas(texto, mode="APOSTILA"):
    """v2.17: Remove seções duplicadas com LIMIAR ADAPTATIVO por modo.
    
    Usa LIMIAR_SECOES diferenciado - mais cuidado pois professor pode repetir propositalmente.
    - FIDELIDADE: 0.70 (mais conservador)
    - APOSTILA: 0.60 (mais agressivo)
    """
    from difflib import SequenceMatcher
    
    # Limiares adaptativos por camada de deduplicação
    # Seções duplicadas: mais cuidado, professor pode repetir propositalmente
    LIMIAR_SECOES = 0.70 if mode == "FIDELIDADE" else 0.60
    
    secoes_dup = detectar_secoes_duplicadas(texto)
    if not secoes_dup: return texto
    
    print("Broom Removendo seções duplicadas (Smart Dedupe v2.14)...")
    linhas = texto.split('\n')
    
    # Rastreia o ÚLTIMO segmento adicionado para cada título (para evitar diluição na comparação)
    ultimo_segmento_visto = {}  # titulo_normalizado -> último texto adicionado
    linhas_para_remover = set()
    linhas_para_adicionar_separador = set()
    
    for dup in secoes_dup:
        # --- 1. Extrair Conteúdo Original ---
        idx_orig = dup['primeira_linha']
        header_orig = linhas[idx_orig].strip()
        match_orig = re.match(r'^(#+)', header_orig)
        nivel_orig = len(match_orig.group(1)) if match_orig else 2
        
        # Extrai conteúdo da seção original
        content_orig = []
        for i in range(idx_orig + 1, len(linhas)):
            line = linhas[i].strip()
            if line.startswith('#'):
                match_now = re.match(r'^(#+)', line)
                if match_now and len(match_now.group(1)) <= nivel_orig: break
            content_orig.append(linhas[i])
        text_orig = "\n".join(content_orig)
        
        titulo_key = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', header_orig)
        titulo_key = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_key, flags=re.IGNORECASE).strip()
        
        # Inicializa o rastreamento se for a primeira vez
        if titulo_key not in ultimo_segmento_visto:
            ultimo_segmento_visto[titulo_key] = text_orig

        # --- 2. Extrair Conteúdo Duplicado ---
        idx_dup = dup['duplicada_linha']
        header_dup = linhas[idx_dup].strip()
        match_dup = re.match(r'^(#+)', header_dup)
        nivel_dup = len(match_dup.group(1)) if match_dup else 2
        
        fim_dup_idx = len(linhas)
        content_dup = []
        for i in range(idx_dup + 1, len(linhas)):
            line = linhas[i].strip()
            if line.startswith('#'):
                match_now = re.match(r'^(#+)', line)
                if match_now and len(match_now.group(1)) <= nivel_dup:
                    fim_dup_idx = i
                    break
            content_dup.append(linhas[i])
        text_dup = "\n".join(content_dup)
        
        # --- 3. Comparar Conteúdo (Lógica Janelada v2.16 MELHORADA) ---
        # Compara APENAS com o último segmento conhecido dessa seção
        texto_referencia = ultimo_segmento_visto.get(titulo_key, text_orig)
        
        len_dup = len(text_dup.strip())
        len_ref = len(texto_referencia.strip())
        
        # v2.16: NOVIDADE - Verificar se o título é 100% idêntico após normalização
        # Se for, força a mesclagem independente do conteúdo (evita redundância de temas)
        titulo_dup_key = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', header_dup)
        titulo_dup_key = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_dup_key, flags=re.IGNORECASE).strip()
        
        titulos_identicos = (titulo_key.lower() == titulo_dup_key.lower())
        
        # Lógica de decisão baseada em tamanho
        if len_dup < 50:
            # Duplicado curto demais = lixo, deletar
            sim = 1.0
            print(f"   ⚠️  Seção duplicada muito curta ({len_dup} chars) - marcando para remoção")
        elif titulos_identicos and len_dup > 100:
            # v2.16: Títulos idênticos = força mesclagem, mantém conteúdo novo sob o mesmo tópico
            sim = 0.5  # Valor que força mesclagem (remove header, mantém conteúdo)
            print(f"   🔄  Títulos 100% idênticos - forçando mesclagem de conteúdo")
        elif len_ref < 50 and len_dup >= 200:
            # Original curto, mas duplicado substancial = original estava incompleto
            # Manter o duplicado como novo conteúdo
            sim = 0.0
            print(f"   ℹ️  Original curto ({len_ref}c), duplicado substancial ({len_dup}c) - mantendo novo conteúdo")
        else:
            sim = SequenceMatcher(None, texto_referencia, text_dup).ratio()
            
        print(f"   Similaridade: {sim:.1%} | Linha {idx_dup} | '{titulo_key[:40]}...'")

        
        if sim > LIMIAR_SECOES:  # Usa limiar de seções (0.70 Fidelidade / 0.60 Apostila)
            print(f"   🗑️  Removendo SEÇÃO INTEIRA (Duplicata confirmada)")
            for i in range(idx_dup, fim_dup_idx):
                linhas_para_remover.add(i)
        else:
            print(f"   🔗 Mesclando conteúdo (Nova informação detectada)")
            linhas_para_remover.add(idx_dup)
            if idx_dup + 1 < len(linhas):
                linhas_para_adicionar_separador.add(idx_dup + 1)
            
            # ATUALIZA o último segmento visto para a próxima comparação
            ultimo_segmento_visto[titulo_key] = text_dup

    # --- 4. Reconstrução ---
    linhas_limpas = []
    for i, linha in enumerate(linhas):
        if i in linhas_para_remover:
            continue
        if i in linhas_para_adicionar_separador:
            linhas_limpas.append("") 
        linhas_limpas.append(linha)
        
    print(f"Check mark button {len(linhas_para_remover)} linhas removidas")
    return '\n'.join(linhas_limpas)


def remover_paragrafos_duplicados(texto: str, min_chars: int = 80) -> str:
    """
    v2.17: Remove parágrafos duplicados dentro do documento.
    
    Lógica:
    - Divide o texto em blocos (separados por linhas em branco).
    - Normaliza cada bloco (lowercase, sem pontuação extra).
    - Mantém apenas a primeira ocorrência de cada bloco normalizado.
    - Ignora blocos muito curtos (< min_chars) para não afetar listas.
    - Preserva tabelas e headers intactos.
    
    Args:
        texto: Texto markdown completo
        min_chars: Tamanho mínimo do parágrafo para considerar na deduplicação
    
    Returns:
        Texto com parágrafos duplicados removidos
    """
    import unicodedata
    
    print("🔄 Removendo parágrafos duplicados (v2.17)...")
    
    # Dividir em blocos por linhas em branco duplas
    blocos = re.split(r'\n\s*\n', texto)
    
    vistos = set()
    blocos_limpos = []
    removidos = 0
    
    for bloco in blocos:
        bloco_stripped = bloco.strip()
        
        # Preservar headers, tabelas e blocos curtos sem verificação
        if (bloco_stripped.startswith('#') or 
            bloco_stripped.startswith('|') or 
            bloco_stripped.startswith('```') or
            bloco_stripped.startswith('> [!') or
            len(bloco_stripped) < min_chars):
            blocos_limpos.append(bloco)
            continue
        
        # Normalizar para comparação
        # Remove pontuação, múltiplos espaços, lowercase
        normalizado = bloco_stripped.lower()
        normalizado = re.sub(r'[^\w\s]', '', normalizado)  # Remove pontuação
        normalizado = re.sub(r'\s+', ' ', normalizado).strip()  # Normaliza espaços
        # Remove acentos para comparação mais robusta
        normalizado = unicodedata.normalize('NFKD', normalizado)
        normalizado = normalizado.encode('ASCII', 'ignore').decode('ASCII')
        
        # Hash do conteúdo normalizado
        bloco_hash = hash(normalizado)
        
        if bloco_hash in vistos:
            removidos += 1
            continue  # Pula duplicata
        
        vistos.add(bloco_hash)
        blocos_limpos.append(bloco)
    
    if removidos > 0:
        print(f"   ✅ {removidos} parágrafos duplicados removidos")
    else:
        print(f"   ℹ️  Nenhum parágrafo duplicado encontrado")
    
    return '\n\n'.join(blocos_limpos)


def remover_titulos_orfaos(texto: str, similaridade_minima: float = 0.85) -> str:
    """
    v2.17: Remove linhas que são variações de títulos H2 já existentes.
    
    Detecta e remove:
    - Linhas bold que repetem um H2 (ex: **3. Execuções...**)
    - H3 que são cópias de H2 (ex: ### Execuções... quando já existe ## Execuções...)
    - Numeração inconsistente (ex: "3. Tema" como texto simples)
    
    Args:
        texto: Texto markdown
        similaridade_minima: Threshold para considerar como duplicata (0-1)
    
    Returns:
        Texto limpo
    """
    from difflib import SequenceMatcher
    
    print("🧹 Removendo títulos órfãos (v2.17)...")
    
    linhas = texto.split('\n')
    
    # 1. Extrair todos os títulos H2 existentes (normalizados)
    titulos_h2 = []
    for linha in linhas:
        if linha.strip().startswith('## '):
            # Extrair apenas o texto do título (sem ## e numeração)
            titulo_limpo = re.sub(r'^##\s*\d+(?:\.\d+)*\.?\s*', '', linha.strip())
            titulo_limpo = titulo_limpo.strip().lower()
            if titulo_limpo:
                titulos_h2.append(titulo_limpo)
    
    if not titulos_h2:
        print("   ℹ️  Nenhum H2 encontrado para comparação")
        return texto
    
    # 2. Identificar linhas órfãs para remoção
    linhas_para_remover = set()
    
    for i, linha in enumerate(linhas):
        stripped = linha.strip()
        
        # Ignorar linhas vazias, headers reais, tabelas
        if not stripped or stripped.startswith('## ') or stripped.startswith('|'):
            continue
        
        # Detectar padrões de "título órfão"
        texto_candidato = None
        
        # Padrão 1: Linha bold com numeração (ex: **3. Execuções...**)
        match_bold = re.match(r'^\*\*\d+\.?\s*(.+?)\*\*\s*$', stripped)
        if match_bold:
            texto_candidato = match_bold.group(1).strip().lower()
        
        # Padrão 2: Numeração simples no início (ex: "3. Execuções Envolvendo...")
        elif re.match(r'^\d+\.\s+[A-Z]', stripped):
            texto_candidato = re.sub(r'^\d+\.\s*', '', stripped).strip().lower()
        
        # Padrão 3: H3 que pode ser duplicata de H2
        elif stripped.startswith('### '):
            texto_candidato = re.sub(r'^###\s*\d*\.?\s*', '', stripped).strip().lower()
        
        if texto_candidato:
            # Comparar com todos os H2
            for h2 in titulos_h2:
                sim = SequenceMatcher(None, texto_candidato, h2).ratio()
                if sim >= similaridade_minima:
                    linhas_para_remover.add(i)
                    break
    
    # 3. Reconstruir sem as linhas órfãs
    if linhas_para_remover:
        print(f"   ✅ {len(linhas_para_remover)} títulos órfãos removidos")
        linhas_limpas = [l for i, l in enumerate(linhas) if i not in linhas_para_remover]
        return '\n'.join(linhas_limpas)
    
    print("   ℹ️  Nenhum título órfão detectado")
    return texto


def _split_long_paragraphs_markdown(
    texto: str,
    *,
    max_paragraph_chars: int = 900,
    skip_timestamped: bool = False,
) -> tuple[str, int]:
    """
    Quebra parágrafos muito longos em Markdown (apenas texto "plain"), preservando
    blocos estruturais como títulos, listas, tabelas, citações e code fences.

    Returns:
        tuple: (texto_novo, qtd_paragrafos_quebrados)
    """
    if not texto:
        return texto, 0

    try:
        max_paragraph_chars = int(max_paragraph_chars)
    except Exception:
        max_paragraph_chars = 900

    if max_paragraph_chars <= 0:
        return texto, 0

    sentence_boundary_re = re.compile(r'(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÂÊÔÃÕÜ0-9“"(\[])')
    timestamp_re = re.compile(r"\[\d{1,2}:\d{2}(?::\d{2})?\]")
    abbrev_re = re.compile(
        r'(?:\b(?:art|arts|dr|dra|sr|sra|etc|cf|n|no|nº|inc|par|fls|p|pp|ex)\.)$',
        re.IGNORECASE,
    )

    def _is_special_paragraph(paragraph_lines: list[str]) -> bool:
        for ln in paragraph_lines:
            s = ln.strip()
            if not s:
                continue
            if s.startswith(("```", "~~~")):
                return True
            if s.startswith("#"):
                return True
            if s.startswith(">"):
                return True
            if s.startswith("|") and s.endswith("|"):
                return True
            if re.match(r"^\s*(?:[-*+]|\\d+\\.)\\s+", ln):
                return True
        return False

    def _split_into_sentences(text: str) -> list[str]:
        parts = [p.strip() for p in sentence_boundary_re.split(text.strip()) if p.strip()]
        if len(parts) <= 1:
            return parts

        merged: list[str] = []
        i = 0
        while i < len(parts):
            cur = parts[i]
            if i + 1 < len(parts) and abbrev_re.search(cur.rstrip()):
                cur = f"{cur} {parts[i + 1]}"
                i += 2
                merged.append(cur)
                continue
            merged.append(cur)
            i += 1
        return merged

    def _group_sentences(sentences: list[str]) -> list[str]:
        # Parágrafos dinâmicos: 2–4 frases, tamanho confortável para leitura.
        target_min = max(220, min(320, max_paragraph_chars // 3))
        target_max = max(420, min(650, max_paragraph_chars - 200))

        paras: list[str] = []
        cur: list[str] = []
        cur_len = 0

        def flush():
            nonlocal cur, cur_len
            if cur:
                paras.append(" ".join(cur).strip())
            cur = []
            cur_len = 0

        for s in sentences:
            s = s.strip()
            if not s:
                continue
            add_len = len(s) + (1 if cur else 0)
            if not cur:
                cur = [s]
                cur_len = len(s)
                continue
            if cur_len + add_len <= target_max:
                cur.append(s)
                cur_len += add_len
                continue
            if cur_len < target_min:
                cur.append(s)
                cur_len += add_len
                flush()
                continue
            flush()
            cur = [s]
            cur_len = len(s)
        flush()

        return [p for p in paras if p]

    def _fallback_word_chunk(text: str) -> list[str]:
        words = [w for w in re.split(r"\s+", text.strip()) if w]
        if not words:
            return []
        target_max = max(420, min(650, max_paragraph_chars - 200))
        paras: list[str] = []
        cur: list[str] = []
        cur_len = 0
        for w in words:
            add_len = len(w) + (1 if cur else 0)
            if cur and cur_len + add_len > target_max:
                paras.append(" ".join(cur))
                cur = [w]
                cur_len = len(w)
                continue
            cur.append(w)
            cur_len += add_len
        if cur:
            paras.append(" ".join(cur))
        return paras

    lines = texto.split("\n")
    out_lines: list[str] = []
    in_fence = False
    paragraph_lines: list[str] = []
    changed_paragraphs = 0

    def flush_paragraph():
        nonlocal changed_paragraphs, paragraph_lines, out_lines
        if not paragraph_lines:
            return
        if in_fence or _is_special_paragraph(paragraph_lines):
            out_lines.extend(paragraph_lines)
            paragraph_lines = []
            return

        joined = " ".join([l.strip() for l in paragraph_lines]).strip()
        if skip_timestamped:
            if timestamp_re.search(joined):
                out_lines.extend(paragraph_lines)
                paragraph_lines = []
                return
            first = paragraph_lines[0].strip()
            if re.match(r"^\\*\\*[^*]{1,30}\\*\\*:\\s+", first):
                out_lines.extend(paragraph_lines)
                paragraph_lines = []
                return
        if len(joined) <= max_paragraph_chars:
            out_lines.extend(paragraph_lines)
            paragraph_lines = []
            return

        sentences = _split_into_sentences(joined)
        if len(sentences) <= 1:
            new_paras = _fallback_word_chunk(joined)
        else:
            new_paras = _group_sentences(sentences)

        if len(new_paras) <= 1:
            out_lines.extend(paragraph_lines)
            paragraph_lines = []
            return

        changed_paragraphs += 1
        for idx, p in enumerate(new_paras):
            if idx > 0:
                out_lines.append("")
            out_lines.append(p)
        paragraph_lines = []

    for ln in lines:
        stripped = ln.strip()
        if stripped.startswith(("```", "~~~")):
            flush_paragraph()
            out_lines.append(ln)
            in_fence = not in_fence
            continue

        if in_fence:
            out_lines.append(ln)
            continue

        if stripped == "":
            flush_paragraph()
            out_lines.append(ln)
            continue

        paragraph_lines.append(ln)

    flush_paragraph()

    return "\n".join(out_lines), changed_paragraphs


def aplicar_correcoes_automaticas(texto: str, *, mode: str | None = None) -> tuple:
    """
    v2.19: Aplica correções automáticas baseadas em padrões comuns de erro.
    Portado de format_transcription_gemini.py.
    
    Correções aplicadas:
    1. Remove saudações duplicadas ("Olá, sejam bem-vindos...")
    2. Remove apresentações repetidas ("Eu sou o professor...")
    3. Padroniza nome do professor (variações → nome canônico)
    4. Corrige itens de lista malformados ("3. \\n Requisitos" → "3. Requisitos")
    5. Remove linhas em branco excessivas
    
    Returns:
        tuple: (texto_corrigido, lista_de_correcoes)
    """
    from difflib import SequenceMatcher
    
    print(f"{Fore.CYAN}🔧 Auto-Fix Pass (v2.19)...")
    
    correcoes = []
    texto_original = texto
    
    # 1. Remover saudações duplicadas (apenas mantém a primeira)
    saudacoes_pattern = r'(?:Olá|Oi),?\s*(?:sejam?\s+)?(?:bem[- ]?vindos?(?:\s+e\s+bem[- ]?vindas?)?)[.,!]?'
    matches = list(re.finditer(saudacoes_pattern, texto, re.IGNORECASE))
    if len(matches) > 1:
        for match in reversed(matches[1:]):
            start = texto.rfind('\n', 0, match.start()) + 1
            end = texto.find('\n', match.end())
            if end == -1: end = len(texto)
            linha = texto[start:end].strip()
            if len(linha) < 150:
                texto = texto[:start] + texto[end+1:]
                correcoes.append(f"Saudação duplicada removida")
    
    # 2. Remover apresentações repetidas do professor
    apresentacao_pattern = r'Eu sou o professor\s+\w+(?:\s+\w+)?'
    matches = list(re.finditer(apresentacao_pattern, texto, re.IGNORECASE))
    if len(matches) > 1:
        for match in reversed(matches[1:]):
            start = texto.rfind('\n', 0, match.start()) + 1
            end = texto.find('\n', match.end())
            if end == -1: end = len(texto)
            linha = texto[start:end].strip()
            if len(linha) < 200:
                texto = texto[:start] + texto[end+1:]
                correcoes.append(f"Apresentação duplicada removida")
    
    # 3. Padronizar nome do professor
    nome_match = re.search(r'professor\s+(\w+(?:\s+\w+)?)', texto, re.IGNORECASE)
    if nome_match:
        nome_canonico = nome_match.group(1)
        variacoes_pattern = rf'\bprofessor\s+(\w+(?:\s+\w+)?)\b'
        for m in re.finditer(variacoes_pattern, texto, re.IGNORECASE):
            nome_atual = m.group(1)
            if nome_atual.lower() != nome_canonico.lower():
                sim = SequenceMatcher(None, nome_canonico.lower(), nome_atual.lower()).ratio()
                if sim > 0.6 and sim < 1.0:
                    texto = texto.replace(f"professor {nome_atual}", f"professor {nome_canonico}")
                    texto = texto.replace(f"Professor {nome_atual}", f"Professor {nome_canonico}")
                    correcoes.append(f"Nome padronizado: '{nome_atual}' → '{nome_canonico}'")
    
    # 4. Corrigir itens de lista vazios ou malformados
    texto_temp = re.sub(r'(\d+\.)\s*\n\s*((?:Requisitos|Preenchimento|Fundamento|Artigo|Lei))', r'\1 \2', texto)
    if texto_temp != texto:
        texto = texto_temp
        correcoes.append("Itens de lista malformados corrigidos")
    
    # 5. Remover linhas em branco excessivas
    texto_limpo = re.sub(r'\n{4,}', '\n\n\n', texto)
    if texto_limpo != texto:
        texto = texto_limpo
        correcoes.append("Linhas em branco excessivas removidas")

    # 6. Limpar placeholders de tabelas (evita vazamento de exemplos do prompt)
    #    Exemplos proibidos: "...", "Art. X", "Lei Y", "Art. X, Lei Y"
    def _is_table_separator_line(line: str) -> bool:
        return bool(re.match(r'^\s*\|[\s:|-]+\|[\s:|-]*$', line.strip()))

    linhas = texto.split('\n')
    substituicoes = 0
    for i, line in enumerate(linhas):
        l = line.strip()
        if not (l.startswith('|') and l.endswith('|')):
            continue
        if _is_table_separator_line(l):
            continue

        # Só mexe se parecer placeholder
        if ('...' not in l and '…' not in l and 'Art. X' not in l and 'Lei Y' not in l):
            continue

        cells = [c.strip() for c in l.split('|')[1:-1]]
        new_cells = []
        changed = False
        for c in cells:
            c_clean = c.strip()
            if c_clean in {'...', '…'}:
                new_cells.append('—')
                changed = True
            elif re.search(r'\bArt\.\s*X\b', c_clean):
                new_cells.append('—')
                changed = True
            elif re.search(r'\bLei\s*Y\b', c_clean):
                new_cells.append('—')
                changed = True
            elif 'Art. X' in c_clean or 'Lei Y' in c_clean:
                # Cobrir combinações como "Art. X, Lei Y"
                new_cells.append('—')
                changed = True
            else:
                new_cells.append(c)

        if changed:
            substituicoes += 1
            linhas[i] = '| ' + ' | '.join(new_cells) + ' |'

    if substituicoes > 0:
        texto = '\n'.join(linhas)
        correcoes.append(f"Placeholders de tabela substituídos por '—' ({substituicoes} linha(s))")

    # 7. Quebrar parágrafos muito longos (APOSTILA/FIDELIDADE)
    mode_norm = (mode or "").strip().upper()
    if mode_norm in {"APOSTILA", "FIDELIDADE"}:
        import os

        env_key = "IUDEX_APOSTILA_MAX_PARAGRAPH_CHARS" if mode_norm == "APOSTILA" else "IUDEX_FIDELIDADE_MAX_PARAGRAPH_CHARS"
        default_max = "500" if mode_norm == "APOSTILA" else "1200"  # v2.41: APOSTILA 900→500 (mais granular, alinhado com format_transcription_gemini)
        try:
            max_chars = int(os.getenv(env_key, default_max))
        except Exception:
            max_chars = int(default_max)

        texto_split, changed = _split_long_paragraphs_markdown(
            texto,
            max_paragraph_chars=max_chars,
            skip_timestamped=(mode_norm == "FIDELIDADE"),
        )
        if changed > 0 and texto_split != texto:
            texto = texto_split
            correcoes.append(f"Parágrafos longos quebrados ({mode_norm}: {changed} parágrafo(s))")
    
    if correcoes:
        print(f"   ✅ {len(correcoes)} correções aplicadas:")
        for c in correcoes[:3]:
            print(f"      - {c}")
        if len(correcoes) > 3:
            print(f"      ... e mais {len(correcoes) - 3}")
    else:
        print(f"   ℹ️  Nenhuma correção necessária")
    
    return texto, correcoes


def normalize_headings(texto):
    """
    v1.0: Normaliza títulos semanticamente similares para uma versão única.
    - Agrupa títulos por similaridade
    - Escolhe o título mais descritivo de cada grupo
    - Remove sufixos como "(Continuação)"
    """
    from difflib import SequenceMatcher
    
    print("🔤 Normalizando títulos similares...")
    linhas = texto.split('\n')
    
    # 1. Extrair todos os títulos com info de nível e posição
    titulos = []
    for i, linha in enumerate(linhas):
        stripped = linha.strip()
        if stripped.startswith('##'):
            match = re.match(r'^(#+)\s*', stripped)
            nivel = len(match.group(1)) if match else 2
            # Extrai título limpo (sem # e sem numeração)
            titulo_limpo = re.sub(r'^#{2,4}\s*\d+(?:\.\d+)*\.?\s*', '', stripped).strip()
            # Remove "(Continuação)" para comparação
            titulo_base = re.sub(r'\s*\(Continuação\)\s*$', '', titulo_limpo, flags=re.IGNORECASE).strip()
            titulos.append({
                'linha': i,
                'nivel': nivel,
                'original': stripped,
                'limpo': titulo_limpo,
                'base': titulo_base
            })
    
    if not titulos:
        return texto
    
    # 2. Agrupar títulos similares (mesmo nível + similaridade > 0.85)
    grupos = []
    usados = set()
    
    for i, t1 in enumerate(titulos):
        if i in usados:
            continue
        grupo = [t1]
        usados.add(i)
        
        for j, t2 in enumerate(titulos):
            if j in usados or j <= i:
                continue
            if t1['nivel'] == t2['nivel']:
                sim = SequenceMatcher(None, t1['base'].lower(), t2['base'].lower()).ratio()
                if sim > 0.85:
                    grupo.append(t2)
                    usados.add(j)
        
        if len(grupo) > 1:
            grupos.append(grupo)
    
    if not grupos:
        # Nada para normalizar, mas ainda remove "(Continuação)"
        texto_limpo = re.sub(r'\s*\(Continuação\)\s*(?=\n|$)', '', texto, flags=re.IGNORECASE)
        return texto_limpo
    
    # 3. Para cada grupo, escolher o "melhor" título (mais curto entre os mais longos)
    #    Lógica: preferir títulos sem "(Continuação)" e com descrição completa
    substituicoes = {}
    for grupo in grupos:
        # Ordenar por: não ter "(Continuação)" primeiro, depois por comprimento (prefer médio)
        candidatos = sorted(grupo, key=lambda x: (
            '(Continuação)' in x['limpo'],  # False (0) vem antes de True (1)
            abs(len(x['limpo']) - 40)  # Preferir títulos de ~40 chars (nem muito curto nem muito longo)
        ))
        
        melhor = candidatos[0]['limpo']
        print(f"   📝 Grupo de {len(grupo)} títulos similares → padronizando para: '{melhor[:50]}...'")
        
        for t in grupo:
            if t['limpo'] != melhor:
                substituicoes[t['linha']] = (t['nivel'], melhor)
    
    # 4. Aplicar substituições
    novas_linhas = []
    for i, linha in enumerate(linhas):
        if i in substituicoes:
            nivel, novo_titulo = substituicoes[i]
            # Preservar numeração existente se houver
            match_num = re.match(r'^(#+\s*\d+(?:\.\d+)*\.?\s*)', linhas[i].strip())
            if match_num:
                prefixo = match_num.group(1)
                # Remove a numeração do novo título se ele tiver uma
                novo_titulo = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', novo_titulo)
                novas_linhas.append(f"{prefixo}{novo_titulo}")
            else:
                novas_linhas.append(f"{'#' * nivel} {novo_titulo}")
        else:
            # Remove "(Continuação)" de qualquer título restante
            if linha.strip().startswith('##'):
                linha = re.sub(r'\s*\(Continuação\)\s*(?=\n|$)', '', linha, flags=re.IGNORECASE)
            novas_linhas.append(linha)
    
    print(f"   ✅ {len(substituicoes)} títulos normalizados, '(Continuação)' removidos")
    return '\n'.join(novas_linhas)

PROMPT_STRUCTURE_REVIEW = """Você é um revisor especializado em estrutura de documentos jurídicos educacionais.

## ESTRUTURA DE MAPEAMENTO INICIAL (Referência - se disponível):
{estrutura_mapeada}

## TAREFA
Revise a ESTRUTURA (headers/títulos) do documento abaixo, COMPARANDO com o mapeamento acima (se disponível). Sua missão é harmonizar o mapeamento planejado com o CONTEÚDO REAL da aula.

## ✅ O QUE VOCÊ DEVE FAZER:

### 1. COMPARAR E REFINAR TÍTULOS (CRÍTICO)
Verifique se los títulos refletem os tópicos do mapeamento. Se um título for genérico, torne-o descritivo.
- ERRADO: "### Questão" ou "### Tópico"
- CORRETO: "### Questão 1: Responsabilidade Civil" (Descritivo conforme mapa/conteúdo)
- **CRÍTICO:** Evite títulos idênticos em seções "irmãs". Diferencie-os pelo conteúdo específico de cada uma.

### 2. VALIDAR HIERARQUIA E PROMOÇÃO DE TÓPICOS
Confirme se a estrutura segue lógica consistente (##, ###, ####). 
- **PROMOÇÃO:** Se um sub-subtópico (ex: 9.19.5) for extenso e tratar de um tema central (ex: Execução Fiscal), PROMOVA-O a um nível superior (ex: ### 9.20) para evitar fragmentação excessiva e respeitar o limite de níveis.

### 3. RENUMERAÇÃO SEQUENCIAL OBRIGATÓRIA
Se você criar, deletar ou promover uma seção, você DEVE renumerar TODAS as seções subsequentes daquela mesma hierarquia para manter a sequência lógica (ex: se 9.20 foi criado, o antigo 9.20 vira 9.21, e assim por diante).

### 4. MESCLAR QUESTÕES DUPLICADAS
Se duas seções têm o mesmo número de questão na mesma área, MESCLE-AS.
- ERRADO: "2.1. Questão 1: TAC" + "2.2. Questão 1: TAC" 
- CORRETO: "2.1. Questão 1: TAC" (Única, com todo o conteúdo unificado)

### 5. PRIORIDADE DO CONTEÚDO REAL (DECIDIR ESTRUTURA)
Se houver conflito entre mapeamento e documento, escolha a estrutura que melhor reflete o CONTEÚDO REAL. O mapeamento é apenas um guia.

### 6. LIMPEZA TÉCNICA (SINTAXE MARKDOWN)
Corrija problemas detalhados de formatação:
- **Tabelas:** Alinhar colunas e adicionar separadores faltantes.
- **Listas:** Corrigir bullets ou numeração mal formatada.
- **Espaçamento:** Padronizar linhas em branco entre seções (mínimo uma linha).
- **Headers Vazios:** Remover títulos sem conteúdo abaixo.
- **Metadados:** Remover headers ou tags inline como "[TIPO: ...]", "**[TIPO: ...]**", ou "[BLOCO 0X]".

## ❌ O QUE VOCÊ NÃO DEVE FAZER:
1. **NÃO ALTERE O CONTEÚDO** dos parágrafos - apenas os títulos e a organização.
2. **NUNCA RESUMA** ou encurte o texto (mesmo tamanho de entrada/saída é obrigatório).
3. **NÃO INVENTE** fatos jurídicos.
4. **NÃO REMOVA** trechos técnicos ou exemplos.

## REGRAS CRÍTICAS DE HIERARQUIA:
- Use **MÁXIMO 3** níveis de hierarquia (##, ###, ####).
- Nunca use # (H1) para subtópicos (apenas para o título principal do documento).
- Preserve a ordem cronológica geral.
- **ANTI-FRAGMENTAÇÃO (CRÍTICO):** Se há 4+ seções ## consecutivas que tratam de aspectos do MESMO tema, **REBAIXE-AS** para ### subtópicos de um ## tema-mãe. Exemplo: "## Citação", "## Intimação", "## Notificação" dentro de Atos de Comunicação → devem virar "## Atos de Comunicação" com "### Citação", "### Intimação", "### Notificação".
- **MARCOS LEGAIS como subtópicos:** Súmulas, Teses e Artigos explicados em profundidade devem ser ### (não ##).

## DOCUMENTO PARA REVISAR:
{documento}

## 📝 RELATÓRIO ESPERADO:
Ao final do documento, inclua um bloco de comentário indicando:
- Quantos títulos foram refinados
- Promoções de seções e renumerações realizadas
- Discrepâncias com o mapeamento (se houver)

Formato:
<!-- RELATÓRIO: X títulos refinados | Y seções promovidas/renumeradas | Discrepâncias: [Nenhuma/Lista] -->

## RESPOSTA:
Retorne o documento COMPLETO E INTEGRAL (mesmo tamanho do original) com os títulos/headers corrigidos e o relatório no final. NÃO RESUMA."""

def renumerar_secoes(texto):
    """
    v1.0: Renumeração Sequencial Determinística.
    
    Esta função é uma camada de segurança extra, aplicada APÓS o AI Review.
    Ela percorre todos os headers numerados e corrige qualquer duplicação ou
    sequência quebrada, garantindo que os números sejam estritamente sequenciais
    dentro de cada nível de hierarquia.
    
    Exemplo de correção:
    - Input:  9.20, 9.21, 9.21, 9.35, 9.36  (duplicação e pulo)
    - Output: 9.20, 9.21, 9.22, 9.23, 9.24  (sequencial)
    """
    logger.info("🔢 Executando Renumeração Sequencial Determinística...")
    
    linhas = texto.split('\n')
    novas_linhas = []
    
    # Contadores por nível de hierarquia (ex: {2: 9, 3: 44} -> ## 9.x, ### 9.44.x)
    # Estrutura: {header_level: {parent_prefix: next_number}}
    # Ex: para ### 9.20.1 -> level=3, parent_prefix="9.20", next_number=2 (para o próximo .x)
    contadores = {}
    
    # Regex para detectar headers numerados: ### 9.20.1. Título ou ## 5. Título
    header_pattern = re.compile(r'^(#{1,4})\s+([\d.]+\.?)\s*(.*)$')
    
    for linha in linhas:
        match = header_pattern.match(linha)
        
        if match:
            hashes = match.group(1)       # "###"
            numero = match.group(2)       # "9.20.1." ou "9.20.1"
            titulo = match.group(3)       # "Título..."
            
            nivel = len(hashes)
            numero_limpo = numero.rstrip('.')
            partes = numero_limpo.split('.')
            
            # Determina o prefixo pai e o sufixo atual
            if len(partes) == 1:
                # Nível raiz: ## 1. ou ## 9.
                prefixo_pai = ""
                sufixo_atual = int(partes[0])
            else:
                # Subnível: ### 9.20. ou #### 9.20.1.
                prefixo_pai = '.'.join(partes[:-1])
                sufixo_atual = int(partes[-1])
            
            # Inicializa contador se não existir
            chave = (nivel, prefixo_pai)
            if chave not in contadores:
                contadores[chave] = sufixo_atual  # Começa do número encontrado
            else:
                contadores[chave] += 1
            
            novo_sufixo = contadores[chave]
            
            # Reconstrói o número
            if prefixo_pai:
                novo_numero = f"{prefixo_pai}.{novo_sufixo}."
            else:
                novo_numero = f"{novo_sufixo}."
            
            # Reconstrói a linha
            nova_linha = f"{hashes} {novo_numero} {titulo}"
            novas_linhas.append(nova_linha)
            
            # Log se houve mudança
            if numero_limpo != novo_numero.rstrip('.'):
                logger.info(f"   🔄 {numero_limpo} → {novo_numero.rstrip('.')}")
        else:
            novas_linhas.append(linha)
    
    logger.info("   ✅ Renumeração concluída.")
    return '\n'.join(novas_linhas)


def audit_heading_levels(texto: str, *, apply_fixes: bool = False) -> tuple[str, list[str]]:
    """
    v2.34: Auditoria determinística de hierarquia.

    Regras:
    - H2 com numeração decimal (ex.: "## 18.2 ...") deve ser subtópico (H3/H4).
    - H4 sem H3 anterior (desde o último H2) é inconsistente.

    Returns:
        tuple: (texto_atualizado, issues)
    """
    if not texto:
        return texto, []

    lines = texto.split("\n")
    issues: list[str] = []
    new_lines: list[str] = []
    in_fence = False
    saw_h2 = False
    saw_h3_since_h2 = False

    heading_re = re.compile(r'^(#{2,4})\s+(.+)$')
    decimal_re = re.compile(r'^(\d+(?:\.\d+)+)\.?\s*(.+)$')

    for line in lines:
        stripped = line.strip()
        if stripped.startswith(("```", "~~~")):
            in_fence = not in_fence
            new_lines.append(line)
            continue

        if in_fence:
            new_lines.append(line)
            continue

        m = heading_re.match(stripped)
        if not m:
            new_lines.append(line)
            continue

        hashes = m.group(1)
        level = len(hashes)
        raw_title = m.group(2).strip()

        if level == 2:
            saw_h2 = True
            saw_h3_since_h2 = False
        elif level == 3:
            saw_h3_since_h2 = True

        dec = decimal_re.match(raw_title)
        if level == 2 and dec:
            num = dec.group(1)
            title = dec.group(2).strip()
            depth = num.count(".") + 1
            issues.append(f"Subtópico numerado em H2: '{raw_title}'")
            if apply_fixes and saw_h2:
                target_level = 3 if depth == 2 else 4
                # Evitar H4 sem H3 no bloco atual.
                if target_level == 4 and not saw_h3_since_h2:
                    target_level = 3
                new_lines.append(f"{'#' * target_level} {num}. {title}".strip())
                if target_level == 3:
                    saw_h3_since_h2 = True
                continue

        if level == 4 and not saw_h3_since_h2:
            issues.append(f"H4 sem H3 anterior: '{raw_title}'")
            if apply_fixes:
                new_lines.append(f"### {raw_title}")
                saw_h3_since_h2 = True
                continue

        new_lines.append(line)

    return "\n".join(new_lines), issues


_TABLE_HEADING_RE = re.compile(
    r'^(#{3,5})\s*(?:[📋🎯]\s*)?(.*)$',
    re.IGNORECASE,
)

_HEADING_RE = re.compile(r'^(#{2,4})\s+(.+)$')
_HEADING_NUMBER_RE = re.compile(r'^(\d+(?:\.\d+)*)(?:\.)?\s*(.+)$')

_STOPWORDS_PT = {
    "para", "pela", "pelo", "como", "mais", "menos", "sobre", "entre", "depois", "antes", "quando",
    "onde", "outra", "outro", "outros", "outras", "seu", "sua", "seus", "suas", "que", "porque",
    "pois", "isso", "essa", "esse", "esta", "este", "estas", "estes", "nao", "não", "sim", "com",
    "sem", "dos", "das", "nos", "nas", "por", "pro", "pra", "uma", "uns", "umas", "como", "pela",
    "pelo", "sobre", "sob", "entre", "na", "no", "em", "ao", "aos", "as", "os", "de", "da", "do",
    "das", "dos", "e", "a", "o", "um", "uma", "que", "ser", "sao", "são",
}


def _keyword_set(texto: str) -> set[str]:
    tokens = re.findall(r"[A-Za-zÀ-ÿ0-9]+", (texto or "").lower())
    return {
        t for t in tokens
        if len(t) >= 4 and t not in _STOPWORDS_PT and not t.isdigit()
    }


def _keyword_similarity(a: str, b: str) -> float:
    sa = _keyword_set(a)
    sb = _keyword_set(b)
    if not sa or not sb:
        return 0.0
    return len(sa & sb) / max(1, len(sa | sb))


def _extract_headings(lines: list[str]) -> list[dict]:
    headings: list[dict] = []
    in_fence = False
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith(("```", "~~~")):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        m = _HEADING_RE.match(stripped)
        if not m:
            continue
        level = len(m.group(1))
        raw_title = m.group(2).strip()
        # Ignorar headings de tabelas/quadros para não quebrar seções.
        if level >= 3:
            lower_title = raw_title.lower()
            if any(tok in lower_title for tok in ("tabela", "quadro", "síntese", "sintese", "pegadinha", "banca", "📋", "🎯")):
                continue
        number = ""
        title = raw_title
        nm = _HEADING_NUMBER_RE.match(raw_title)
        if nm:
            number = nm.group(1)
            title = nm.group(2).strip()
        headings.append(
            {
                "line": idx,
                "level": level,
                "number": number,
                "title": title,
                "raw": raw_title,
            }
        )
    return headings


_CONVERSATIONAL_HEADING_PREFIXES = (
    "já ",
    "na prova",
    "para quem",
    "minha proposta",
    "bom dia",
    "gente ",
    "pessoal ",
    "vamos ",
    "então ",
    "logo ",
)

_TECHNICAL_HEADING_TOKENS = (
    "licita",
    "contrat",
    "lei",
    "decreto",
    "súmula",
    "sumula",
    "tcu",
    "stj",
    "stf",
    "juris",
    "administra",
    "governan",
    "execução",
    "execucao",
    "fiscal",
    "responsabil",
    "constituci",
    "nulidade",
    "compet",
    "proced",
    "auditoria",
    "fidelidade",
)


def _normalize_heading_title(raw_title: str) -> str:
    title = (raw_title or "").strip()
    nm = _HEADING_NUMBER_RE.match(title)
    if nm:
        return nm.group(2).strip()
    return title


def _contains_technical_signal(title: str) -> bool:
    t = (title or "").lower()
    return any(tok in t for tok in _TECHNICAL_HEADING_TOKENS)


def _heading_quality_flags(title: str, level: int) -> list[str]:
    flags: list[str] = []
    normalized = re.sub(r"\s+", " ", (title or "").strip())
    lower = normalized.lower()
    words = re.findall(r"[A-Za-zÀ-ÿ0-9]+", normalized)

    if not normalized:
        flags.append("empty")
        return flags

    if len(normalized) > 110:
        flags.append("too_long_chars")
    if len(words) > 20:
        flags.append("too_long_words")
    if any(lower.startswith(pfx) for pfx in _CONVERSATIONAL_HEADING_PREFIXES):
        flags.append("conversational_prefix")
    if level == 2 and "na prova de" in lower:
        flags.append("exam_phrase_h2")
    if level in (2, 3) and not _contains_technical_signal(lower):
        flags.append("missing_technical_signal")

    return flags


def _extract_h2_h3_heading_entries(text: str) -> list[dict]:
    lines = (text or "").splitlines()
    entries: list[dict] = []
    for idx, line in enumerate(lines):
        m = _HEADING_RE.match(line.strip())
        if not m:
            continue
        level = len(m.group(1))
        if level not in (2, 3):
            continue
        raw_title = m.group(2).strip()
        entries.append(
            {
                "line_idx": idx,
                "level": level,
                "raw_title": raw_title,
                "title": _normalize_heading_title(raw_title),
            }
        )
    return entries


def enforce_fidelity_heading_guard(
    original_text: str,
    revised_text: str,
    *,
    freeze_h2_h3: bool = True,
) -> tuple[str, dict]:
    """
    Garante estabilidade de títulos no modo FIDELIDADE.

    - Compara H2/H3 entre texto original e revisado.
    - Opcionalmente congela H2/H3 (sempre restaura título original quando mudou).
    - Aplica rollback seletivo para títulos degradados (frase corrida/conversacional etc.).
    - Retorna texto corrigido + telemetria de drift.
    """
    orig_entries = _extract_h2_h3_heading_entries(original_text or "")
    rev_entries = _extract_h2_h3_heading_entries(revised_text or "")
    lines = (revised_text or "").splitlines()

    changed_count = 0
    restored_count = 0
    degraded_count = 0
    diffs: list[dict] = []

    for idx, (orig, rev) in enumerate(zip(orig_entries, rev_entries), start=1):
        original_title = (orig.get("raw_title") or "").strip()
        revised_title = (rev.get("raw_title") or "").strip()
        if not original_title or not revised_title:
            continue
        if original_title == revised_title:
            continue

        changed_count += 1
        flags = _heading_quality_flags(revised_title, int(rev.get("level") or 2))
        degraded = len(flags) > 0
        if degraded:
            degraded_count += 1

        should_restore = freeze_h2_h3 or degraded
        if should_restore:
            prefix = "#" * int(rev.get("level") or 2)
            target_line = int(rev.get("line_idx") or 0)
            if 0 <= target_line < len(lines):
                lines[target_line] = f"{prefix} {original_title}"
                restored_count += 1

        diffs.append(
            {
                "index": idx,
                "level": int(rev.get("level") or 2),
                "original": original_title,
                "revised": revised_title,
                "restored": bool(should_restore),
                "quality_flags": flags,
            }
        )

    telemetry = {
        "freeze_h2_h3": bool(freeze_h2_h3),
        "headers_changed_count": changed_count,
        "headers_restored_count": restored_count,
        "headers_degraded_count": degraded_count,
        "headers_diff": diffs,
    }
    return "\n".join(lines), telemetry


# ---------------------------------------------------------------------------
# Sanitização de títulos na estrutura mapeada (v2.47)
# ---------------------------------------------------------------------------

_GREETING_TITLE_PREFIXES = (
    "bom dia", "boa tarde", "boa noite", "já ", "pessoal ",
    "gente ", "olha ", "obrigado", "obrigada",
)

_CANONICAL_LABEL_L1 = "Introdução e Contextualização"
_CANONICAL_LABEL_SUB = "Abertura"

_MAX_MAPPED_TITLE_WORDS = 8
_MAX_MAPPED_TITLE_CHARS = 70


def _sanitize_mapped_structure(estrutura: str) -> str:
    """Valida e corrige títulos de estrutura que são trechos literais de fala.

    Reutiliza ``_heading_quality_flags`` e ``_CONVERSATIONAL_HEADING_PREFIXES``
    para detectar títulos degradados no mapeamento.

    Regras (alinhadas com PROMPT_MAPEAMENTO regra 8):
    - Títulos > 8 palavras ou > 70 chars → rótulo canônico
    - Prefixos conversacionais (saudações, logística) → rótulo canônico
    - Preserva âncoras ``| ABRE: "..." | FECHA: "..."`` intactas
    """
    if not estrutura:
        return estrutura

    lines = estrutura.split('\n')
    fixed_lines: list[str] = []
    sanitized_count = 0

    for line in lines:
        stripped = line.strip()
        # Detecta linhas numeradas: "1. Título", "   1.1. Subtítulo"
        m = re.match(r'^(\s*\d+(?:\.\d+)*\.?\s+)(.*)', stripped)
        if not m:
            fixed_lines.append(line)
            continue

        prefix_num = m.group(1)
        rest = m.group(2).strip()

        # Separa âncoras ABRE/FECHA preservando literalmente (incluindo aspas)
        anchor_part = ""
        title = rest
        anchor_idx = rest.find("| ABRE:")
        if anchor_idx >= 0:
            title = rest[:anchor_idx].strip()
            anchor_part = " " + rest[anchor_idx:]

        # Deriva nível: "1." → level 2, "1.1." → level 3
        parts_count = len([p for p in prefix_num.strip().rstrip('.').split('.') if p.strip().isdigit()])
        level = min(parts_count + 1, 4)  # 1. → 2, 1.1. → 3, 1.1.1. → 4

        # Aplica heading quality flags existentes
        flags = _heading_quality_flags(title, level)

        needs_fix = False
        # Flags de qualidade indicam problema
        if "conversational_prefix" in flags or "too_long_chars" in flags or "too_long_words" in flags:
            needs_fix = True
        # Backup: checa limites alinhados com o prompt (8 palavras, 70 chars)
        words = re.findall(r'[A-Za-zÀ-ÿ0-9]+', title)
        if len(title) > _MAX_MAPPED_TITLE_CHARS or len(words) > _MAX_MAPPED_TITLE_WORDS:
            needs_fix = True

        if needs_fix:
            is_level1 = parts_count == 1
            title_lower = title.lower()

            if any(title_lower.startswith(pfx) for pfx in _GREETING_TITLE_PREFIXES):
                canonical = _CANONICAL_LABEL_L1 if is_level1 else _CANONICAL_LABEL_SUB
            elif is_level1:
                canonical = _CANONICAL_LABEL_L1
            else:
                canonical = _CANONICAL_LABEL_SUB

            sanitized_count += 1
            print(f"{Fore.YELLOW}⚠️  Título sanitizado: '{title[:60]}' → '{canonical}'{Style.RESET_ALL}")
            fixed_lines.append(f"{prefix_num}{canonical}{anchor_part}")
        else:
            fixed_lines.append(line)

    if sanitized_count:
        print(f"{Fore.CYAN}🔧 {sanitized_count} título(s) de estrutura sanitizado(s){Style.RESET_ALL}")

    return '\n'.join(fixed_lines)


def _extract_table_blocks(lines: list[str], start: int, end: int) -> list[dict]:
    blocks: list[dict] = []
    i = start
    while i < end:
        line = lines[i].strip()
        m = _TABLE_HEADING_RE.match(line)
        if m:
            heading_level = len(m.group(1))
            heading_text = m.group(2).strip()
            # Only treat as table header if it looks like table/quadros.
            if not any(tok in heading_text.lower() for tok in ("tabela", "quadro", "síntese", "sintese", "pegadinha", "banca")):
                i += 1
                continue
            block_start = i
            i += 1
            has_table_rows = False
            while i < end:
                nxt = lines[i].strip()
                if _HEADING_RE.match(nxt):
                    break
                if nxt.startswith("|"):
                    has_table_rows = True
                    i += 1
                    continue
                if has_table_rows and nxt == "":
                    i += 1
                    continue
                if has_table_rows and nxt and not nxt.startswith("|"):
                    break
                i += 1
            block_end = i
            block_text = "\n".join(lines[block_start:block_end]).strip()
            blocks.append(
                {
                    "start": block_start,
                    "end": block_end,
                    "heading_level": heading_level,
                    "heading_text": heading_text,
                    "text": block_text,
                }
            )
            continue
        i += 1
    return blocks


def reatribuir_tabelas_por_topico(
    texto: str,
    *,
    apply_fixes: bool = True,
    min_similarity: float = 0.08,
    margin: float = 0.08,
) -> tuple[str, list[str]]:
    """
    v2.34: Reatribui tabelas que parecem ter sido vinculadas ao tópico errado.

    Heurística:
    - Avalia similaridade de palavras-chave entre tabela e título atual vs título pai.
    - Se a tabela está em subtópico (numeração decimal ou nível >=3) e o título pai é mais similar,
      move a tabela para antes do heading do subtópico.
    """
    if not texto:
        return texto, []

    lines = texto.split("\n")
    headings = _extract_headings(lines)
    if not headings:
        return texto, []

    # Construir intervalos de seção por heading
    sections: list[dict] = []
    for idx, h in enumerate(headings):
        start = h["line"] + 1
        end = headings[idx + 1]["line"] if idx + 1 < len(headings) else len(lines)
        parent_idx = None
        for j in range(idx - 1, -1, -1):
            if headings[j]["level"] < h["level"]:
                parent_idx = j
                break
        sections.append(
            {
                "heading_index": idx,
                "start": start,
                "end": end,
                "parent_index": parent_idx,
            }
        )

    moves: list[dict] = []
    issues: list[str] = []

    for sec in sections:
        h = headings[sec["heading_index"]]
        parent_idx = sec["parent_index"]
        if parent_idx is None:
            continue
        parent = headings[parent_idx]
        # Considerar apenas subtópicos (decimais ou nível >=3)
        if "." not in (h.get("number") or "") and h.get("level", 2) < 3:
            continue
        table_blocks = _extract_table_blocks(lines, sec["start"], sec["end"])
        if not table_blocks:
            continue
        for block in table_blocks:
            def _context_slice(start_line: int, end_line: int, *, tail: bool = False, max_lines: int = 3) -> str:
                chunk = lines[start_line:end_line]
                nonempty = [ln.strip() for ln in chunk if ln.strip()]
                if not nonempty:
                    return ""
                if tail:
                    return " ".join(nonempty[-max_lines:])
                return " ".join(nonempty[:max_lines])

            # Contexto do subtópico (linhas antes da tabela) e do pai (linhas imediatamente anteriores ao subtópico).
            current_context = _context_slice(sec["start"], block["start"], tail=True)
            parent_context = _context_slice(
                sections[parent_idx]["start"],
                h["line"],
                tail=True,
            )
            current_score = _keyword_similarity(f"{h['title']} {current_context}", block["text"])
            parent_score = _keyword_similarity(f"{parent['title']} {parent_context}", block["text"])
            if parent_score >= min_similarity and (parent_score - current_score) >= margin:
                # Move a tabela para imediatamente antes do heading do subtópico
                moves.append(
                    {
                        "start": block["start"],
                        "end": block["end"],
                        "insert_at": h["line"],
                        "from": h["title"],
                        "to": parent["title"],
                        "heading_text": block["heading_text"],
                        "scores": (current_score, parent_score),
                    }
                )

    if not moves or not apply_fixes:
        if moves:
            for m in moves:
                issues.append(
                    f"Tabela sugerida para mover ('{m['heading_text'][:40]}...'): '{m['from']}' → '{m['to']}'"
                )
        return texto, issues

    # Aplicar movimentos de baixo para cima para preservar índices
    moves.sort(key=lambda m: m["start"], reverse=True)
    for m in moves:
        block_lines = lines[m["start"]:m["end"]]
        del lines[m["start"]:m["end"]]
        insert_at = m["insert_at"]
        if insert_at > m["start"]:
            insert_at = max(0, insert_at - (m["end"] - m["start"]))
        for offset, bl in enumerate(block_lines):
            lines.insert(insert_at + offset, bl)
        issues.append(
            f"Tabela reatribuída: '{m['heading_text'][:40]}...' de '{m['from']}' → '{m['to']}'"
        )

    return "\n".join(lines), issues


def coletar_candidatos_reatribuicao_tabelas(
    texto: str,
    *,
    min_similarity: float = 0.08,
    margin: float = 0.08,
    max_candidates: int = 5,
) -> list[dict]:
    """
    v2.34: Coleta casos ambíguos para reatribuição de tabelas via IA.
    """
    if not texto:
        return []
    lines = texto.split("\n")
    headings = _extract_headings(lines)
    if not headings:
        return []

    sections: list[dict] = []
    for idx, h in enumerate(headings):
        start = h["line"] + 1
        end = headings[idx + 1]["line"] if idx + 1 < len(headings) else len(lines)
        parent_idx = None
        for j in range(idx - 1, -1, -1):
            if headings[j]["level"] < h["level"]:
                parent_idx = j
                break
        sections.append(
            {
                "heading_index": idx,
                "start": start,
                "end": end,
                "parent_index": parent_idx,
            }
        )

    def _context_slice(start_line: int, end_line: int, *, tail: bool = False, max_lines: int = 3) -> str:
        chunk = lines[start_line:end_line]
        nonempty = [ln.strip() for ln in chunk if ln.strip()]
        if not nonempty:
            return ""
        if tail:
            return " ".join(nonempty[-max_lines:])
        return " ".join(nonempty[:max_lines])

    candidates: list[dict] = []
    for sec in sections:
        h = headings[sec["heading_index"]]
        parent_idx = sec["parent_index"]
        if parent_idx is None:
            continue
        parent = headings[parent_idx]
        if "." not in (h.get("number") or "") and h.get("level", 2) < 3:
            continue
        table_blocks = _extract_table_blocks(lines, sec["start"], sec["end"])
        if not table_blocks:
            continue
        for block in table_blocks:
            current_context = _context_slice(sec["start"], block["start"], tail=True)
            parent_context = _context_slice(sections[parent_idx]["start"], h["line"], tail=True)
            current_score = _keyword_similarity(f"{h['title']} {current_context}", block["text"])
            parent_score = _keyword_similarity(f"{parent['title']} {parent_context}", block["text"])
            # Ambíguo: scores próximos ou ambos baixos, mas há match mínimo com um lado.
            if max(current_score, parent_score) < min_similarity:
                continue
            if abs(parent_score - current_score) < margin:
                candidates.append(
                    {
                        "start": block["start"],
                        "end": block["end"],
                        "insert_at": h["line"],
                        "current_title": h["title"],
                        "parent_title": parent["title"],
                        "current_context": current_context,
                        "parent_context": parent_context,
                        "table_text": block["text"],
                    }
                )
        if len(candidates) >= max_candidates:
            break

    return candidates[:max_candidates]

def deterministic_structure_fix(text):
    """
    v1.1: Reorganização Estrutural Determinística (Regex).
    Adaptativo: Detecta se o documento usa H1 ou apenas H2 como nível principal.
    """
    print(f"{Fore.CYAN}🧩 Executando Reorganização Estrutural Determinística...")
    
    lines = text.split('\n')
    
    # Detecção de Hierarquia
    has_h1 = any(re.match(r'^#\s+', line) for line in lines)
    header_level_regex = r'^#\s+' if has_h1 else r'^##\s+'
    print(f"   ℹ️  Nível principal detectado: {'H1 (#)' if has_h1 else 'H2 (##)'}")

    # Estruturas de dados (Preserva ordem de inserção)
    content_map = {
        "PREAMBULO": [],
        "DISCIPLINAS": {}, 
        "ENCERRAMENTO": []
    }
    
    current_area = "PREAMBULO"
    current_block = []
    disciplinas_order = [] 
    
    # Regex Adaptativo
    # Captura o texto do cabeçalho principal (seja H1 ou H2)
    # Exclui "Questão" ou "Q." para não quebrar simulados dentro de uma área
    re_disciplina = re.compile(rf'{header_level_regex}(?!Questão|Q\.)([^0-9\.]+.*)', re.IGNORECASE)
    re_encerramento = re.compile(rf'{header_level_regex}(?:ENCERRAMENTO|CONSIDERAÇÕES|CONCLUSÃO)', re.IGNORECASE)
    
    def flush_block(area, block_lines):
        if not block_lines: return
        block_text = '\n'.join(block_lines)
        
        if area == "PREAMBULO":
            content_map["PREAMBULO"].append(block_text)
        elif area == "ENCERRAMENTO":
            content_map["ENCERRAMENTO"].append(block_text)
        else:
            if area not in content_map["DISCIPLINAS"]:
                content_map["DISCIPLINAS"][area] = []
                disciplinas_order.append(area)
            content_map["DISCIPLINAS"][area].append(block_text)

    for line in lines:
        # 1. Detectar mudança de disciplina macro
        match_disc = re_disciplina.match(line)
        if match_disc:
            flush_block(current_area, current_block)
            current_block = []
            
            raw_area = match_disc.group(1).strip().upper()
            
            # Normalização de nome de área
            if "DIREITO" not in raw_area and len(raw_area) < 50:
                 # Adiciona prefixo se parecer nome de matéria jurídica comum
                 if any(x in raw_area for x in ["CIVIL", "PENAL", "TRABALHO", "ADMINISTRATIVO", "CONSTITUCIONAL"]):
                     current_area = f"DIREITO {raw_area}"
                 else:
                     current_area = raw_area
            else:
                 current_area = raw_area
                 
            # IMPORTANTE: Se estamos operando em modo H2, essa linha é um Header que queremos manter?
            # Se for modo H1, recriamos "# AREA". 
            # Se for modo H2, recriamos "# AREA" (upcast) ou mantemos "## AREA"?
            # Para padronizar Apostilas, vamos promover tudo a H1 na reconstrução.
            continue 
            
        # 2. Detectar Encerramento
        if re_encerramento.match(line):
            flush_block(current_area, current_block)
            current_block = []
            current_area = "ENCERRAMENTO"
            continue

        current_block.append(line)
        
    flush_block(current_area, current_block)
    
    # Reconstrução
    final_output = []
    
    # Preambulo
    if content_map["PREAMBULO"]:
        final_output.append("# ORIENTAÇÕES GERAIS / INTRODUÇÃO")
        final_output.extend(content_map["PREAMBULO"])
        final_output.append("")

    # Disciplinas / Tópicos Principais
    for area in disciplinas_order:
        area_clean = area.replace("#", "").strip()
        final_output.append(f"# {area_clean}")
        for block in content_map["DISCIPLINAS"][area]:
            final_output.append(block)
        final_output.append("")
        
    # Encerramento
    if content_map["ENCERRAMENTO"]:
        final_output.append("# CONSIDERAÇÕES FINAIS")
        final_output.extend(content_map["ENCERRAMENTO"])
        
    num_identified = len(disciplinas_order)
    print(f"   ✅ Reorganizado: {num_identified} seções principais identificadas.")
    
    # Fallback: Se não identificou nada (tudo preambulo), retorna original para não estragar
    if num_identified == 0 and len(content_map["PREAMBULO"]) > 0:
        print("   ⚠️ Nenhuma estrutura detectada. Mantendo original.")
        return text
        
    return '\n'.join(final_output)

PROMPT_STRUCTURE_REVIEW_LITE = """Você é um revisor editorial especializado em ESTRUTURA e FORMATAÇÃO de documentos educacionais. Você receberá:
1. Uma **Estrutura de Mapeamento Inicial** (planejada antes da formatação)
2. O **Documento Processado** (resultado da formatação por chunks)

Sua tarefa é analisar ambos e garantir que os títulos estejam **hierarquicamente corretos e alinhados com o conteúdo real**, sem jamais alterar a ordem cronológica.

---

## 📋 ESTRUTURA DE MAPEAMENTO INICIAL (Referência):
{estrutura_mapeada}

---

## ✅ O QUE VOCÊ DEVE FAZER:
1. **NÃO REESCREVER TÍTULOS EXISTENTES:** preserve o texto dos headings já presentes no documento. Ajuste apenas nível/hierarquia quando estritamente necessário.
2. **Validar Hierarquia:** Confirme que a estrutura (##, ###, ####) segue uma lógica consistente (ex: seções > subseções > detalhes).
3. **Decidir a Melhor Estrutura:** Se houver conflito entre mapeamento e documento, escolha a estrutura que melhor reflete o CONTEÚDO REAL do texto.
4. **Subtópicos Órfãos:** Se detectar headers como "A.", "B.", "C." isolados como tópicos principais, converta-os em subníveis do tópico anterior (ex: ## para ###).
5. **Corrigir Sintaxe Markdown:** Tabelas, listas, espaçamento.
6. **Remover Vazios:** Títulos sem conteúdo abaixo.
7. **NUNCA alterar conteúdo de parágrafos** (somente forma/sintaxe).

## 🔴 REGRAS CRÍTICAS DE HIERARQUIA:
- Use **MÁXIMO 3** níveis de hierarquia (##, ###, ####). Nunca use # (H1) para subtópicos.
- **ANTI-FRAGMENTAÇÃO (CRÍTICO):** Se há 4+ seções ## consecutivas que tratam de aspectos do MESMO tema, **REBAIXE-AS** para ### subtópicos de um ## tema-mãe. Exemplo: "## Citação", "## Intimação", "## Notificação" dentro de Atos de Comunicação → devem virar "## Atos de Comunicação" com "### Citação", "### Intimação", "### Notificação".
- **MARCOS LEGAIS como subtópicos:** Súmulas, Teses de Repercussão Geral e Artigos explicados em profundidade devem ser ### (não ##).
- Preserve a ordem cronológica geral.

## 📌 EXEMPLOS DE CORREÇÃO:

**Subtópicos Órfãos → Hierarquia Correta:**
- ANTES:
  ```
  ## A. Requisitos do Dano
  ## B. Nexo Causal
  ```
- DEPOIS:
  ```
  ### A. Requisitos do Dano
  ### B. Nexo Causal
  ```

**Numeração Duplicada → Sequencial:**
- ANTES: `### 9.20`, `### 9.20`, `### 9.35`
- DEPOIS: `### 9.20`, `### 9.21`, `### 9.22`

## ❌ O QUE VOCÊ NÃO DEVE FAZER:
1. **NÃO MOVA** blocos de texto. A ordem deve permanecer 100% cronológica.
2. **NÃO MESCLE** seções que apareçam em momentos diferentes da aula.
3. **NÃO RESUMA** nem altere o corpo dos parágrafos.
4. **NÃO ADICIONE** conteúdo novo.

## 📝 RELATÓRIO ESPERADO:
Ao final do documento, inclua um bloco de comentário (que será removido) indicando:
- Quantos níveis de heading foram ajustados (sem reescrever o texto dos títulos)
- Se a estrutura final segue o mapeamento ou foi adaptada
- Discrepâncias encontradas (se houver)

Formato:
<!-- RELATÓRIO: X níveis ajustados | Estrutura: [MAPEAMENTO/ADAPTADA] | Discrepâncias: [Nenhuma/Lista] -->

---

## 📄 DOCUMENTO PARA REVISAR:
{documento}

---

## RESPOSTA:
Retorne o documento COMPLETO com a formatação corrigida e o relatório no final."""

async def ai_structure_review_lite(texto, client, model, estrutura_mapeada=None, metrics=None):
    """
    v2.3: Revisão LEVE de formatação Markdown com VALIDAÇÃO CRUZADA.
    Compara o documento processado com a estrutura de mapeamento inicial.
    NÃO reescreve texto de títulos; valida hierarquia/sintaxe e reporta discrepâncias.
    NÃO reorganiza nem mescla conteúdo.

    Melhorias v2.3:
    - Melhor tratamento de rate limits
    - Integração com MetricsCollector
    - Opção de contexto total (sem split/truncate) para máxima fidelidade
    """
    from difflib import SequenceMatcher
    import asyncio
    import time
    import json

    print(f"{Fore.MAGENTA}  🧹 Revisão Leve de Formatação (IA - Modo Fidelidade v2.3)...{Style.RESET_ALL}")

    start_time = time.time()

    # v2.45: Modo totalidade de janela ativo por padrão absoluto
    # para preservar contexto máximo na revisão leve.
    use_full_context = True
    split_threshold = int(os.getenv("IUDEX_SPLIT_REVIEW_THRESHOLD", "400000"))
    max_doc_chars = 800000

    # v2.3: Se documento muito grande, dividir em partes e processar em paralelo
    if not use_full_context and len(texto) > split_threshold:
        print(f"{Fore.CYAN}   🔀 Documento grande ({len(texto)//1000}k chars), dividindo em partes paralelas...{Style.RESET_ALL}")

        # Dividir em partes de ~350k chars cada, com overlap de 10k para contexto
        part_size = 350000
        overlap = 10000
        parts = []
        idx = 0
        while idx < len(texto):
            end = min(idx + part_size, len(texto))
            # Tentar cortar em quebra de linha
            if end < len(texto):
                newline_pos = texto.rfind('\n', idx + part_size - 5000, end)
                if newline_pos > idx:
                    end = newline_pos + 1
            parts.append(texto[idx:end])
            idx = end - overlap if end < len(texto) else end

        print(f"   📦 Dividido em {len(parts)} partes para processamento paralelo")

        # Processar partes em paralelo com semaphore
        semaphore = asyncio.Semaphore(2)  # Max 2 paralelos para evitar rate limit

        async def process_part(part_idx: int, part_text: str) -> tuple:
            async with semaphore:
                try:
                    # Passar parte para processamento (recursivo, mas parte será < threshold)
                    result = await ai_structure_review_lite(
                        part_text, client, model, estrutura_mapeada, metrics
                    )
                    return (part_idx, result, None)
                except Exception as e:
                    return (part_idx, part_text, e)  # Retorna original em caso de erro

        tasks = [process_part(i, p) for i, p in enumerate(parts)]
        results = await asyncio.gather(*tasks)

        # Ordenar e mesclar resultados
        results_sorted = sorted(results, key=lambda x: x[0])
        merged_parts = []
        for part_idx, result, error in results_sorted:
            if error:
                print(f"{Fore.YELLOW}   ⚠️ Erro na parte {part_idx + 1}: {error}. Usando original.{Style.RESET_ALL}")
            merged_parts.append(result)

        # Remover overlaps duplicados na mesclagem
        final_text = merged_parts[0]
        for i in range(1, len(merged_parts)):
            # Encontrar ponto de sobreposição
            overlap_start = final_text[-overlap*2:] if len(final_text) > overlap*2 else final_text
            next_part = merged_parts[i]

            # Buscar melhor ponto de corte
            best_cut = 0
            for line in overlap_start.split('\n'):
                if line.strip() and line.strip() in next_part[:overlap*2]:
                    pos = next_part.find(line.strip())
                    if pos >= 0:
                        best_cut = pos + len(line.strip())
                        break

            final_text += next_part[best_cut:].lstrip()

        duration = time.time() - start_time
        print(f"{Fore.GREEN}   ✅ Revisão paralela concluída ({len(parts)} partes, {duration:.1f}s).{Style.RESET_ALL}")
        return final_text

    # Processamento single-shot
    if not use_full_context and len(texto) > max_doc_chars:
        print(f"{Fore.YELLOW}   ⚠️ Documento muito longo ({len(texto)} chars), truncando para {max_doc_chars//1000}k...{Style.RESET_ALL}")
        texto_para_revisao = texto[:max_doc_chars] + "\n\n[... documento truncado para revisão ...]"
    else:
        texto_para_revisao = texto
    
    # Preparar estrutura mapeada (se disponível)
    if estrutura_mapeada:
        estrutura_str = estrutura_mapeada if use_full_context else estrutura_mapeada[:50000]
        print(f"{Fore.CYAN}   📋 Usando estrutura de mapeamento inicial ({len(estrutura_mapeada)} chars) para validação cruzada.{Style.RESET_ALL}")
    else:
        estrutura_str = "[Estrutura de mapeamento não disponível - analisar documento para inferir estrutura ideal]"
        print(f"{Fore.YELLOW}   ℹ️  Sem mapeamento inicial, IA irá inferir estrutura ideal do próprio documento.{Style.RESET_ALL}")
    
    def call_gemini():
        return client.models.generate_content(
            model=model,
            contents=PROMPT_STRUCTURE_REVIEW_LITE.format(
                estrutura_mapeada=estrutura_str,
                documento=texto_para_revisao
            ),
            config=types.GenerateContentConfig(
                max_output_tokens=65536,  # Máximo permitido
                thinking_config=types.ThinkingConfig(
                    include_thoughts=False, 
                    thinking_level="HIGH"  # HIGH para análise estrutural profunda
                ),
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                ]
            )
        )
    
    # Retry com backoff exponencial
    max_retries = 3
    response = None
    for tentativa in range(max_retries):
        try:
            response = await asyncio.to_thread(call_gemini)
            resultado = response.text.replace('```markdown', '').replace('```', '').strip()
            break
        except Exception as e:
            if tentativa < max_retries - 1:
                wait_time = 2 ** tentativa
                print(f"{Fore.YELLOW}   ⚠️ Tentativa {tentativa + 1} falhou: {e}. Aguardando {wait_time}s...{Style.RESET_ALL}")
                await asyncio.sleep(wait_time)
            else:
                print(f"{Fore.RED}   ⚠️ Erro na revisão leve após {max_retries} tentativas: {e}. Mantendo original.{Style.RESET_ALL}")
                return texto
    
    # Métricas
    duration = time.time() - start_time
    if metrics and response:
        try:
            usage = response.usage_metadata
            metrics.record_call(
                provider="gemini",
                prompt_tokens=getattr(usage, 'prompt_token_count', 0) or 0,
                completion_tokens=getattr(usage, 'candidates_token_count', 0) or 0,
                duration=duration,
                model=model,
                cached_tokens_in=getattr(usage, 'cached_content_token_count', 0) or 0,
            )
        except:
            pass  # Silently ignore metrics errors
    
    # Extrair relatório (suporta JSON e texto)
    relatorio_json_match = re.search(r'<!--\s*RELATÓRIO_JSON:\s*(\{.+?\})\s*-->', resultado, re.IGNORECASE | re.DOTALL)
    relatorio_match = re.search(r'<!--\s*RELATÓRIO:\s*(.+?)\s*-->', resultado, re.IGNORECASE)
    
    if relatorio_json_match:
        try:
            relatorio_data = json.loads(relatorio_json_match.group(1))
            print(f"{Fore.CYAN}   📊 Relatório (JSON): {relatorio_data}{Style.RESET_ALL}")
        except:
            print(f"{Fore.CYAN}   📊 Relatório: {relatorio_json_match.group(1)}{Style.RESET_ALL}")
        resultado = re.sub(r'<!--\s*RELATÓRIO_JSON:.+?-->\s*', '', resultado, flags=re.IGNORECASE | re.DOTALL).strip()
    elif relatorio_match:
        relatorio = relatorio_match.group(1)
        print(f"{Fore.CYAN}   📊 Relatório da IA: {relatorio}{Style.RESET_ALL}")
        resultado = re.sub(r'<!--\s*RELATÓRIO:.+?-->\s*', '', resultado, flags=re.IGNORECASE).strip()
    
    # Validação: resultado deve ter pelo menos 80% do tamanho (padronizado)
    if len(resultado) < len(texto) * 0.80:
        print(f"{Fore.YELLOW}   ⚠️ Revisão retornou texto muito curto ({len(resultado)} vs {len(texto)}). Mantendo original.{Style.RESET_ALL}")
        return texto
    
    # Verificar se a ordem dos headers foi preservada (usando similaridade)
    headers_original = re.findall(r'^(#{1,4})\s+(.+?)$', texto, re.MULTILINE)
    headers_revisado = re.findall(r'^(#{1,4})\s+(.+?)$', resultado, re.MULTILINE)
    
    if len(headers_original) > 5 and len(headers_revisado) > 5:
        # Usar similaridade ao invés de igualdade exata
        similares = sum(1 for (_, h1), (_, h2) in zip(headers_original[:10], headers_revisado[:10]) 
                       if SequenceMatcher(None, h1.strip(), h2.strip()).ratio() > 0.6)
        if similares < 6:
            print(f"{Fore.YELLOW}   ⚠️ Ordem dos headers parece alterada ({similares}/10 similares). Mantendo original.{Style.RESET_ALL}")
            return texto
    
    # 📝 Relatório de Alterações nos Títulos
    alteracoes = []
    for i, ((lvl_orig, h_orig), (lvl_rev, h_rev)) in enumerate(zip(headers_original[:50], headers_revisado[:50])):
        if h_orig.strip() != h_rev.strip():
            alteracoes.append(f"   • '{h_orig[:40]}...' → '{h_rev[:40]}...'")
    
    if alteracoes:
        print(f"{Fore.CYAN}   📝 Títulos Refinados ({len(alteracoes)}):{Style.RESET_ALL}")
        for alt in alteracoes[:10]:  # Mostrar no máximo 10
            print(alt)
        if len(alteracoes) > 10:
            print(f"   ... e mais {len(alteracoes) - 10} alterações.")
    else:
        print(f"{Fore.GREEN}   ℹ️  Nenhum título foi alterado (estrutura já estava OK).{Style.RESET_ALL}")
    
    print(f"{Fore.GREEN}   ✅ Formatação revisada (modo leve v2.2, {duration:.1f}s).{Style.RESET_ALL}")
    return resultado

async def ai_structure_review(texto, client, model, estrutura_mapeada=None, metrics=None):
    """
    v2.2: Revisão semântica de estrutura usando IA com VALIDAÇÃO CRUZADA.
    Compara o documento com a estrutura de mapeamento inicial.
    Corrige: questões duplicadas, subtópicos órfãos, fragmentação excessiva.
    
    Melhorias v2.2:
    - Contexto total por padrão na revisão semântica (APOSTILA)
    - Validação de ordem dos headers
    - Integração com MetricsCollector
    - Suporte a relatório JSON
    """
    from google.genai import types
    from difflib import SequenceMatcher
    import asyncio
    import time
    import json
    
    print(f"{Fore.CYAN}🧠 Revisão Semântica de Estrutura (IA v2.2)...")
    
    start_time = time.time()
    
    # v2.45: Para APOSTILA, usar contexto completo por padrão (sem truncamento).
    # Mantém opção de voltar ao legado via env em cenários extremos.
    use_full_context = _env_truthy("IUDEX_APOSTILA_FULL_CONTEXT", default=True)
    max_doc_chars = 800000
    if not use_full_context and len(texto) > max_doc_chars:
        print(f"   ⚠️ Documento muito longo ({len(texto)} chars), truncando para {max_doc_chars//1000}k...")
        texto_para_revisao = texto[:max_doc_chars] + "\n\n[... documento truncado para revisão estrutural ...]"
    else:
        texto_para_revisao = texto

    # Preparar estrutura mapeada (se disponível)
    if estrutura_mapeada:
        estrutura_str = estrutura_mapeada if use_full_context else estrutura_mapeada[:50000]
        print(f"{Fore.CYAN}   📋 Usando estrutura de mapeamento inicial ({len(estrutura_mapeada)} chars) para validação cruzada.{Style.RESET_ALL}")
    else:
        estrutura_str = "[Estrutura de mapeamento não disponível - analisar documento autonomamente]"
        print(f"{Fore.YELLOW}   ℹ️  Sem mapeamento inicial, IA revisará estrutura autonomamente.{Style.RESET_ALL}")
    
    def call_gemini():
        return client.models.generate_content(
            model=model,
            contents=PROMPT_STRUCTURE_REVIEW.format(
                estrutura_mapeada=estrutura_str,
                documento=texto_para_revisao
            ),
            config=types.GenerateContentConfig(
                max_output_tokens=65536,
                thinking_config=types.ThinkingConfig(
                    include_thoughts=False,
                    thinking_level="HIGH"
                ),
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                ]
            )
        )
    
    # Retry com backoff exponencial
    max_retries = 3
    response = None
    for tentativa in range(max_retries):
        try:
            response = await asyncio.to_thread(call_gemini)
            resultado = response.text.replace('```markdown', '').replace('```', '').strip()
            break
        except Exception as e:
            if tentativa < max_retries - 1:
                wait_time = 2 ** tentativa
                print(f"{Fore.YELLOW}   ⚠️ Tentativa {tentativa + 1} falhou: {e}. Aguardando {wait_time}s...{Style.RESET_ALL}")
                await asyncio.sleep(wait_time)
            else:
                print(f"{Fore.RED}   ⚠️ Erro na revisão por IA após {max_retries} tentativas: {e}. Mantendo estrutura original.{Style.RESET_ALL}")
                return texto
    
    # Métricas
    duration = time.time() - start_time
    if metrics and response:
        try:
            usage = response.usage_metadata
            metrics.record_call(
                provider="gemini",
                prompt_tokens=getattr(usage, 'prompt_token_count', 0) or 0,
                completion_tokens=getattr(usage, 'candidates_token_count', 0) or 0,
                duration=duration,
                model=model,
                cached_tokens_in=getattr(usage, 'cached_content_token_count', 0) or 0,
            )
        except:
            pass  # Silently ignore metrics errors
    
    # Validação: o resultado deve ter pelo menos 80% do tamanho original (padronizado)
    if len(resultado) < len(texto) * 0.80:
        print(f"   ⚠️ Revisão retornou texto muito curto ({len(resultado)} vs {len(texto)}). Mantendo original.")
        return texto
    
    # Extrair relatório (suporta JSON e texto)
    relatorio_json_match = re.search(r'<!--\s*RELATÓRIO_JSON:\s*(\{.+?\})\s*-->', resultado, re.IGNORECASE | re.DOTALL)
    relatorio_match = re.search(r'<!--\s*RELATÓRIO:\s*(.+?)\s*-->', resultado, re.IGNORECASE)
    
    if relatorio_json_match:
        try:
            relatorio_data = json.loads(relatorio_json_match.group(1))
            print(f"{Fore.CYAN}   📊 Relatório (JSON): {relatorio_data}{Style.RESET_ALL}")
        except:
            print(f"{Fore.CYAN}   📊 Relatório: {relatorio_json_match.group(1)}{Style.RESET_ALL}")
        resultado = re.sub(r'<!--\s*RELATÓRIO_JSON:.+?-->\s*', '', resultado, flags=re.IGNORECASE | re.DOTALL).strip()
    elif relatorio_match:
        relatorio = relatorio_match.group(1)
        print(f"{Fore.CYAN}   📊 Relatório da IA: {relatorio}{Style.RESET_ALL}")
        resultado = re.sub(r'<!--\s*RELATÓRIO:.+?-->\s*', '', resultado, flags=re.IGNORECASE).strip()
    
    # Validação de ordem dos headers (novo em v2.2)
    headers_original = re.findall(r'^(#{1,4})\s+(.+?)$', texto, re.MULTILINE)
    headers_revisado = re.findall(r'^(#{1,4})\s+(.+?)$', resultado, re.MULTILINE)
    
    if len(headers_original) > 5 and len(headers_revisado) > 5:
        # Usar similaridade para detectar reordenação
        similares = sum(1 for (_, h1), (_, h2) in zip(headers_original[:10], headers_revisado[:10]) 
                       if SequenceMatcher(None, h1.strip(), h2.strip()).ratio() > 0.6)
        if similares < 6:
            print(f"{Fore.YELLOW}   ⚠️ Ordem dos headers parece alterada ({similares}/10 similares). Mantendo original.{Style.RESET_ALL}")
            return texto
    
    # Contar quantos headers foram alterados
    diff = abs(len(headers_original) - len(headers_revisado))
    
    print(f"{Fore.GREEN}   ✅ Estrutura revisada: {len(headers_original)} → {len(headers_revisado)} headers (Δ{diff}, {duration:.1f}s){Style.RESET_ALL}")
    return resultado

def normalizar_fingerprint(texto, tipo):
    """Normaliza texto para comparação (ex: 'Lei 11.100' -> 'lei 11100')"""
    texto = texto.lower().strip()
    
    if tipo == 'leis':
        nums = re.findall(r'\d+', texto)
        if nums:
            num_full = ''.join(nums)
            if len(num_full) >= 4: return f"lei {num_full}"
            return None
            
    elif tipo == 'sumulas':
        nums = re.findall(r'\d+', texto)
        if nums: return f"súmula {''.join(nums)}"
            
    elif tipo == 'artigos':
        nums = re.findall(r'\d+', texto)
        if nums: return f"artigo {''.join(nums)}"
            
    return re.sub(r'[^\w\s]', '', texto)

def extrair_fingerprints(texto):
    """Extrai 'fingerprints' únicos e normalizados do texto"""
    fingerprints = {'leis': set(), 'sumulas': set(), 'artigos': set(), 'julgados': set()}
    
    lei_pattern = re.compile(r'\b(?:lei|l\.)\s*n?º?\s*([\d\.]+)', re.IGNORECASE)
    sumula_pattern = re.compile(r'\bsúmula\s*(?:vinculante)?\s*n?º?\s*(\d+)', re.IGNORECASE)
    
    for match in lei_pattern.finditer(texto):
        fp = normalizar_fingerprint(f"lei {match.group(1)}", 'leis')
        if fp: fingerprints['leis'].add(fp)
    
    for match in sumula_pattern.finditer(texto):
        fp = normalizar_fingerprint(f"súmula {match.group(1)}", 'sumulas')
        if fp: fingerprints['sumulas'].add(fp)
        
    return fingerprints

def contar_ocorrencias_robust(fingerprints, texto):
    """Conta ocorrências com suporte a formatação jurídica formal"""
    contagens = {}
    texto_lower = texto.lower()
    
    for categoria, items in fingerprints.items():
        for item in items:
            key = f"{categoria}:{item}"
            if categoria == 'leis':
                num_bruto = item.split()[-1] 
                num = re.sub(r'[^\d]', '', num_bruto)
                num_regex = r"\.?".join(list(num))
                pattern = f"lei(?:\\s+|\\.|\\,|nº|n\\.|n\\s|num\\.?)*{num_regex}\\b"
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
            elif categoria == 'sumulas':
                num = item.split()[-1]
                num_regex = r"\.?".join(list(num))
                pattern = f"súmula(?:\\s+|\\.|\\,|vinculante|nº|n\\.|n\\s)*{num_regex}\\b"
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
            else:
                contagens[key] = texto_lower.count(item)
    return contagens

def contar_ocorrencias_robust(fingerprints, texto):
    """Conta ocorrencias com suporte a formatação jurídica formal (Lei nº X)"""
    contagens = {}
    
    # CORREÇÃO 1: Não remover pontuação indiscriminadamente, apenas normalizar espaços
    # Mantemos barras e pontos para evitar fusão de números (11.101/2005)
    texto_lower = texto.lower()
    
    for categoria, items in fingerprints.items():
        for item in items:
            key = f"{categoria}:{item}"
            
            if categoria == 'leis':
                # item ex: "lei 4320" -> extrai "4320"
                # Remove pontuação do item para garantir match limpo no número
                num_bruto = item.split()[-1] 
                num = re.sub(r'[^\d]', '', num_bruto)
                
                # Permite pontos opcionais entre dígitos (ex: 4.320 match com 4320)
                num_regex = r"\.?".join(list(num))
                
                # CORREÇÃO 2: Regex flexível que aceita "n", "nº", "no", "num" no meio
                # Aceita: "Lei 4320", "Lei nº 4.320", "Lei n. 4320"
                # O \W* permite pontos/barras entre Lei e o número
                # Adicionado \b no final para evitar matches parciais (Lei 10 != Lei 100)
                pattern = f"lei(?:\\s+|\\.|\\,|nº|n\\.|n\\s|num\\.?)*{num_regex}\\b"
                
                # Usamos findall no texto original (lower) para pegar variações com pontuação
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
                
            elif categoria == 'sumulas':
                num = item.split()[-1]
                num_regex = r"\.?".join(list(num)) # Súmulas raramente tem ponto, mas por garantia
                
                # Mesma lógica para súmulas (Súmula Vinculante nº 10)
                pattern = f"súmula(?:\\s+|\\.|\\,|vinculante|nº|n\\.|n\\s)*{num_regex}\\b"
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
                
            else:
                # Fallback para outros tipos (busca literal simples)
                contagens[key] = texto_lower.count(item)
                
    return contagens

def verificar_cobertura(texto_original, texto_formatado, arquivo_saida=None):
    """Verifica omissões e duplicações artificiais entre original e formatado"""
    logger.info("🔍 Verificando cobertura e duplicações...")
    
    # Extrai fingerprints do original
    fp_original = extrair_fingerprints(texto_original)
    
    # Conta ocorrências em ambos
    contagem_original = contar_ocorrencias_robust(fp_original, texto_original)
    contagem_formatado = contar_ocorrencias_robust(fp_original, texto_formatado)
    
    omissoes = []
    duplicacoes = []
    
    for key, count_orig in contagem_original.items():
        count_fmt = contagem_formatado.get(key, 0)
        categoria, item = key.split(':', 1)
        
        # Omissão: estava no original mas sumiu
        if count_orig > 0 and count_fmt == 0:
            omissoes.append({
                'categoria': categoria,
                'item': item,
                'original': count_orig,
                'formatado': count_fmt
            })
        
        # Duplicação (agora considerada positiva em materiais didáticos)
        if count_fmt > count_orig:
            duplicacoes.append({
                'categoria': categoria,
                'item': item,
                'original': count_orig,
                'formatado': count_fmt,
                'extra': count_fmt - count_orig
            })
    
    # Gera relatório
    total_items = len([k for k, v in contagem_original.items() if v > 0])
    items_preservados = total_items - len(omissoes)
    cobertura = items_preservados / total_items * 100 if total_items > 0 else 100
    
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    logger.info(f"📊 RELATÓRIO DE VERIFICAÇÃO")
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    logger.info(f"✅ Cobertura: {cobertura:.1f}% ({items_preservados}/{total_items} referências)")
    
    relatorio_txt = []
    relatorio_txt.append(f"RELATÓRIO DE VERIFICAÇÃO DE FIDELIDADE")
    relatorio_txt.append(f"Cobertura: {cobertura:.1f}% ({items_preservados}/{total_items} referências)")
    
    if omissoes:
        logger.warning(f"\n❌ POSSÍVEIS OMISSÕES ({len(omissoes)}):")
        relatorio_txt.append(f"\n❌ POSSÍVEIS OMISSÕES ({len(omissoes)}):")
        for o in omissoes[:15]: 
            msg = f"   - [{o['categoria']}] {o['item']}"
            logger.warning(msg)
            relatorio_txt.append(msg)
        if len(omissoes) > 15:
            logger.warning(f"   ... e mais {len(omissoes) - 15} omissões")
            relatorio_txt.append(f"   ... e mais {len(omissoes) - 15} omissões")
    else:
        logger.info("✅ Nenhuma omissão detectada")
        relatorio_txt.append("\n✅ Nenhuma omissão detectada")
    
    if duplicacoes:
        logger.info(f"\nℹ️ CITAÇÕES REFORÇADAS (Tabelas/Resumos) ({len(duplicacoes)}):")
        relatorio_txt.append(f"\nℹ️ CITAÇÕES REFORÇADAS ({len(duplicacoes)}):")
        for d in duplicacoes[:10]:
            msg = f"   - [{d['categoria']}] {d['item']}: {d['original']}x -> {d['formatado']}x"
            logger.info(msg)
            relatorio_txt.append(msg)
    
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    
    # Salva relatório em arquivo se especificado
    if arquivo_saida:
        relatorio_path = arquivo_saida.replace('.md', '_verificacao.txt')
        with open(relatorio_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(relatorio_txt))
        logger.info(f"📝 Relatório salvo em: {relatorio_path}")

def corrigir_tabelas_prematuras(texto: str, min_chars_apos_tabela: int = 100, min_linhas_apos: int = 2) -> str:
    """
    v2.28: Detecta e corrige tabelas que aparecem antes do conteúdo terminar.

    Problema: O LLM às vezes gera a tabela no meio do tópico, antes de
    terminar de explicar todo o conteúdo.

    Solução: Se houver texto substancial (>min_chars) APÓS uma tabela e ANTES
    do próximo heading, move a tabela para depois desse texto.

    Args:
        texto: Markdown com possíveis tabelas prematuras
        min_chars_apos_tabela: Mínimo de caracteres após tabela para considerar prematura (default: 100)
        min_linhas_apos: Mínimo de linhas de conteúdo após tabela (default: 2)

    Returns:
        Texto com tabelas reposicionadas
    """
    linhas = texto.split('\n')
    resultado = []
    i = 0
    tabelas_corrigidas = 0

    def is_major_heading(line: str) -> bool:
        """Heading H1/H2/H3 que indica novo tópico (e NÃO é título de tabela)."""
        s = (line or "").strip()
        if not s.startswith('#'):
            return False
        level = 0
        for c in s:
            if c == '#':
                level += 1
            else:
                break
        if level <= 3:
            return not is_table_title(s)
        return False

    def is_table_line(line: str) -> bool:
        stripped = line.strip()
        return stripped.startswith('|') and '|' in stripped[1:]

    def is_table_title(line: str) -> bool:
        s = (line or "").strip()
        if not s:
            return False
        lower = s.lower()
        is_heading = bool(re.match(r'^#{3,6}\s+', s))
        starts_with_emoji = s.startswith('📋') or s.startswith('🎯')
        has_emoji = ('📋' in s) or ('🎯' in s)
        has_keyword = (
            'quadro-síntese' in lower
            or 'quadro síntese' in lower
            or 'pegadinha' in lower
            or 'como a banca' in lower
        )
        # Para evitar falsos positivos (texto corrido contendo "quadro"/"pegadinha"),
        # exigimos estrutura típica de "título": heading ou linha iniciando com emoji.
        if starts_with_emoji:
            return True
        if is_heading and (has_emoji or has_keyword):
            return True
        return False

    while i < len(linhas):
        linha = linhas[i]

        # Detectar título de tabela (📋 ou 🎯)
        if is_table_title(linha):
            group_title = linha
            tabelas_linhas = []

            # Capturar um grupo de tabelas consecutivas (ex.: 📋 + 🎯)
            while i < len(linhas) and is_table_title(linhas[i]):
                titulo_tabela = linhas[i]
                tabelas_linhas.append(titulo_tabela)
                i += 1

                # Capturar a tabela (linhas com |) + linhas em branco imediatamente ao redor
                while i < len(linhas) and (is_table_line(linhas[i]) or not linhas[i].strip()):
                    tabelas_linhas.append(linhas[i])
                    i += 1

                # Se houver outra tabela logo em seguida (apenas com espaços/linhas vazias entre),
                # capturamos como parte do mesmo grupo.
                k = i
                while k < len(linhas) and not linhas[k].strip():
                    k += 1
                if k < len(linhas) and is_table_title(linhas[k]):
                    # Preservar as linhas em branco entre as tabelas
                    tabelas_linhas.extend(linhas[i:k])
                    i = k
                    continue
                break

            # Verificar se há conteúdo substancial APÓS o grupo de tabelas e ANTES do próximo major heading
            j = i
            bloco_apos = []
            while j < len(linhas):
                if is_major_heading(linhas[j]):
                    break
                # Se surgir outra tabela não-consecutiva, não atravessar (evita reorder agressivo)
                if is_table_title(linhas[j]):
                    break
                bloco_apos.append(linhas[j])
                j += 1

            conteudo_apos = [l for l in bloco_apos if l.strip() and not is_table_line(l)]
            chars_apos = sum(len(l) for l in conteudo_apos)
            linhas_apos = len([l for l in conteudo_apos if l.strip()])

            # Se há texto substancial após a tabela, é uma tabela prematura
            if chars_apos >= min_chars_apos_tabela and linhas_apos >= min_linhas_apos:
                tabelas_corrigidas += 1
                print(
                    f"{Fore.YELLOW}   🔄 Tabela prematura detectada: '{group_title[:50]}...' "
                    f"({chars_apos} chars, {linhas_apos} linhas de conteúdo após)"
                )

                # Conteúdo primeiro
                resultado.extend(bloco_apos)
                # Depois o grupo de tabelas
                if resultado and resultado[-1].strip():
                    resultado.append('')
                resultado.extend(tabelas_linhas)
                if resultado and resultado[-1].strip():
                    resultado.append('')

                i = j  # Pular o texto já processado
                continue

            # Grupo no lugar certo, adicionar normalmente
            resultado.extend(tabelas_linhas)
            continue

        resultado.append(linha)
        i += 1

    if tabelas_corrigidas > 0:
        print(f"{Fore.GREEN}   ✅ Corrigidas {tabelas_corrigidas} tabelas prematuras")

    return '\n'.join(resultado)


def mover_tabelas_para_fim_de_secao(texto):
    """
    v2.11: Reorganiza tabelas movendo-as para o final do BLOCO ATUAL (H2 ou H3).
    Corrige bug de tabelas sumindo ou ficando muito longe do contexto.
    """
    logger.info("📊 Reorganizando tabelas (Smart Layout)...")
    
    linhas = texto.split('\n')
    resultado = []
    tabelas_pendentes = [] 

    def _is_table_title_line(line: str) -> bool:
        s = (line or "").strip()
        if not s:
            return False
        is_heading = bool(re.match(r'^#{3,5}\s+', s))
        is_bold = s.startswith('**')
        if not (is_heading or is_bold):
            # Alguns modelos às vezes emitem o título sem markdown de heading.
            if s.startswith('📋') or 'quadro-síntese' in s.lower() or 'quadro-sintese' in s.lower():
                return True
            return False

        lowered = s.lower()
        return any(
            x in lowered
            for x in [
                'tabela',
                'resumo',
                'quadro',
                'síntese',
                'sintese',
                'esquema',
                '📋',
                'prova',
                'banca',
                'pegadinha',
                'questão',
                'questao',
                'questões',
                'questoes',
            ]
        )

    def _pop_recent_table_title(result_lines: list, max_lookback_nonempty: int = 12):
        """
        Recupera um título de tabela recente (H3-H5 / bold / 📋) mesmo que
        haja algumas linhas de texto entre o título e a tabela.
        """
        nonempty_seen = 0
        for idx in range(len(result_lines) - 1, -1, -1):
            s = (result_lines[idx] or "").strip()
            if not s:
                continue
            if s.startswith('# ') or s.startswith('## '):
                break
            nonempty_seen += 1
            if nonempty_seen > max_lookback_nonempty:
                break
            if _is_table_title_line(result_lines[idx]):
                return result_lines.pop(idx)
        return None

    def _is_table_separator_line(line: str) -> bool:
        s = (line or "").strip()
        return bool(s) and s.startswith("|") and set(s.replace("|", "").strip()).issubset({"-", ":", " "})

    def _next_nonempty_index(lines: list, start_idx: int) -> int | None:
        j = start_idx
        while j < len(lines):
            if (lines[j] or "").strip():
                return j
            j += 1
        return None

    def _is_table_header_at(lines: list, idx: int) -> bool:
        """Heurística: linha com '|' seguida de uma linha separadora (pula vazias)."""
        if idx < 0 or idx >= len(lines):
            return False
        s = (lines[idx] or "").strip()
        if not s or "|" not in s:
            return False
        nxt = _next_nonempty_index(lines, idx + 1)
        if nxt is None:
            return False
        return _is_table_separator_line(lines[nxt])
    
    i = 0
    while i < len(linhas):
        linha = linhas[i]
        linha_strip = linha.strip()
        
        # 1. DETECTAR SE É UM TÍTULO (apenas H1/H2/H3 delimitam "bloco")
        # Motivo: H4/H5 são frequentemente usados como subtítulos dentro do mesmo assunto.
        # Se flusharmos em qualquer '#', a tabela pode parar antes do assunto terminar.
        if re.match(r'^#{1,3}\s+', linha_strip):
            # Despeja tabelas antes de iniciar o novo tópico
            if tabelas_pendentes:
                resultado.append('') # Espaço antes
                for t_info in tabelas_pendentes:
                    if t_info['titulo']:
                        resultado.append(t_info['titulo'])
                    resultado.extend(t_info['linhas'])
                    resultado.append('') # Espaço depois
                tabelas_pendentes = []
            
            resultado.append(linha)
            i += 1
            continue

        # 2. DETECTAR INÍCIO DE TABELA
        eh_inicio_tabela = False
        if '|' in linha_strip:
            has_separator = False
            for lookahead in range(1, 3): 
                if i + lookahead < len(linhas):
                    prox = linhas[i + lookahead].strip()
                    if set(prox).issubset(set('|- :')): 
                         has_separator = True
                         break
            
            if has_separator or (linha_strip.startswith('|') and linha_strip.endswith('|')):
                eh_inicio_tabela = True

        if eh_inicio_tabela:
            # --- Captura da Tabela ---
            tabela_linhas = []
            titulo_tabela = None
            
            # v2.24+: Recupera título de tabela mesmo com "texto intruso" entre o título e a tabela.
            # Se o LLM inserir explicação após "📋 Quadro-síntese", este trecho remove o título
            # da posição original para reagrupá-lo com a tabela no flush do bloco.
            if resultado:
                titulo_tabela = _pop_recent_table_title(resultado)

            # Captura as linhas da tabela
            j = i
            seen_separator = False
            while j < len(linhas):
                curr = linhas[j].strip()
                if not curr:
                    # Evita "colar" duas tabelas distintas separadas por linhas vazias.
                    # Linhas vazias dentro de tabela são raras; aqui preferimos robustez.
                    # Se a próxima tabela começar após a quebra, ela será capturada no loop externo.
                    nxt = _next_nonempty_index(linhas, j + 1)
                    if nxt is None:
                        break
                    if seen_separator and _is_table_header_at(linhas, nxt):
                        break
                    # Caso contrário, apenas pula vazios (não inclui no buffer da tabela).
                    j += 1
                    continue
                if '|' in curr:
                    # Se já vimos o separador, e encontrarmos um novo header+separador,
                    # tratamos como início de OUTRA tabela (não continuação).
                    if seen_separator and _is_table_header_at(linhas, j):
                        break
                    tabela_linhas.append(linhas[j])
                    if _is_table_separator_line(linhas[j]):
                        seen_separator = True
                    j += 1
                else:
                    break # Texto normal

            if len(tabela_linhas) > 0:
                tabelas_pendentes.append({
                    'titulo': titulo_tabela,
                    'linhas': tabela_linhas
                })
                i = j # Pula as linhas processadas
                continue
            else:
                if titulo_tabela:
                    resultado.append(titulo_tabela)
        
        resultado.append(linha)
        i += 1
    
    # 3. FINAL DO DOCUMENTO
    if tabelas_pendentes:
        resultado.append('')
        for t_info in tabelas_pendentes:
            if t_info['titulo']:
                resultado.append(t_info['titulo'])
            resultado.extend(t_info['linhas'])
            resultado.append('')
            
    return '\n'.join(resultado)


def mesclar_tabelas_divididas(texto: str) -> str:
    """
    v2.27: Detecta tabelas que foram divididas entre chunks e as mescla.
    
    Padrão detectado:
    | Col1 | Col2 |
    |------|------|
    | A    | B    |
    
    [... linhas em branco ...]
    
    | Col1 | Col2 |      <-- Mesma estrutura = tabela continuada
    |------|------|
    | C    | D    |
    
    A função identifica tabelas consecutivas com mesmo número de colunas
    (separadas apenas por linhas em branco) e remove o header duplicado
    da segunda tabela para criar uma tabela unificada.
    """
    logger.info("📊 Mesclando tabelas divididas (v2.27)...")
    
    lines = texto.split('\n')
    result = []
    i = 0
    tables_merged = 0
    
    def is_table_line(line: str) -> bool:
        stripped = line.strip()
        return stripped.startswith('|') and stripped.endswith('|')
    
    def is_separator_line(line: str) -> bool:
        stripped = line.strip()
        return stripped.startswith('|') and set(stripped.replace('|', '').strip()).issubset({'-', ':', ' '})
    
    def count_columns(line: str) -> int:
        return line.count('|') - 1 if '|' in line else 0
    
    while i < len(lines):
        line = lines[i]
        result.append(line)
        
        # Detectar fim de tabela (linha atual é tabela, próxima não é)
        if is_table_line(line) and not is_separator_line(line):
            # Verificar se a próxima linha não é tabela
            if i + 1 < len(lines) and not is_table_line(lines[i + 1]):
                last_table_cols = count_columns(line)
                
                # Encontrar o header da tabela atual (voltando até achar header + separator)
                last_table_header = None
                for back in range(len(result) - 1, -1, -1):
                    if is_separator_line(result[back]) and back > 0:
                        last_table_header = result[back - 1].strip()
                        break
                
                # Procurar próxima tabela (pulando linhas em branco)
                lookahead = 1
                while i + lookahead < len(lines):
                    next_line = lines[i + lookahead].strip()
                    if next_line == '':
                        lookahead += 1
                        continue
                    if next_line.startswith('#'):
                        # Novo título = tabelas são de seções diferentes, NÃO mesclar
                        break
                    if not is_table_line(lines[i + lookahead]):
                        # Qualquer linha não-tabela (ex.: "📋 Quadro-síntese" em bold, explicação, etc.)
                        # indica que não é continuação direta da mesma tabela.
                        break
                    if is_table_line(lines[i + lookahead]):
                        next_table_cols = count_columns(lines[i + lookahead])
                        next_table_header = lines[i + lookahead].strip()
                        
                        # v2.27: Verificar se é continuação (mesmo número de colunas E mesmo header)
                        headers_match = last_table_header == next_table_header if last_table_header else True
                        if last_table_cols == next_table_cols and last_table_cols >= 2 and headers_match:
                            # Verificar se a próxima linha é separator (header + separator = pular ambos)
                            skip_count = 0
                            if i + lookahead + 1 < len(lines) and is_separator_line(lines[i + lookahead + 1]):
                                # Pular header duplicado e separator
                                skip_count = 2
                                tables_merged += 1
                                print(f"   🔗 Mesclando tabela em linha {i + lookahead} ({next_table_cols} colunas)")
                            
                            if skip_count > 0:
                                # Remover linhas em branco que já foram adicionadas ao result
                                while result and result[-1].strip() == '':
                                    result.pop()
                                # Avançar para depois do header+separator da segunda tabela
                                i += lookahead + skip_count - 1
                    break
        
        i += 1
    
    if tables_merged > 0:
        print(f"   ✅ {tables_merged} tabela(s) mesclada(s)")
    else:
        print(f"   ℹ️  Nenhuma tabela dividida detectada")
    
    return '\n'.join(result)

def garantir_titulo_tabela_banca(texto: str) -> str:
    """
    Garante que a tabela "Como a banca cobra / pegadinhas" tenha título visível.
    Evita que a tabela fique colada ao quadro-síntese sem o subtítulo.
    """
    lines = texto.split('\n')
    output = []

    def _is_banca_header(line: str) -> bool:
        return (line or "").strip().lower().startswith('| como a banca cobra |')

    def _is_banca_title(line: str) -> bool:
        s = (line or "").strip().lower()
        return s.startswith('#') and ('🎯' in s or 'banca cobra' in s or 'pegadinha' in s)

    def _is_heading(line: str) -> bool:
        return (line or "").strip().startswith('#')

    for line in lines:
        if _is_banca_header(line):
            has_title = False
            nonempty_seen = 0
            for back in range(len(output) - 1, -1, -1):
                s = (output[back] or "").strip()
                if not s:
                    continue
                nonempty_seen += 1
                if _is_banca_title(output[back]):
                    has_title = True
                    break
                if _is_heading(output[back]) or nonempty_seen >= 8:
                    break
            if not has_title:
                if output and output[-1].strip():
                    output.append('')
                output.append('#### 🎯 Tabela — Como a banca cobra / pegadinhas')
                output.append('')
        output.append(line)

    return '\n'.join(output)


def _similaridade_palavras(texto_a: str, texto_b: str) -> float:
    """
    v2.33: Calcula similaridade entre dois textos baseado em overlap de palavras.
    Retorna valor entre 0.0 (nenhuma similaridade) e 1.0 (idênticos).
    """
    if not texto_a or not texto_b:
        return 0.0

    # Normalizar: lowercase e remover pontuação
    def normalizar(t):
        t = t.lower()
        t = re.sub(r'[^\w\s]', '', t)
        return set(w for w in t.split() if len(w) > 2)  # Ignorar palavras muito curtas

    palavras_a = normalizar(texto_a)
    palavras_b = normalizar(texto_b)

    if not palavras_a or not palavras_b:
        return 0.0

    intersecao = palavras_a & palavras_b
    uniao = palavras_a | palavras_b

    return len(intersecao) / len(uniao) if uniao else 0.0


def _buscar_ancora_no_texto(texto_lower: str, titulo: str, transcricao_completa: str) -> int:
    """
    v2.33: Busca inteligente de âncora quando o modelo não forneceu citação verbatim.

    Estratégias:
    1. Buscar palavras-chave do título no texto
    2. Buscar frases de transição comuns próximas às palavras-chave

    Returns:
        Posição no texto ou -1 se não encontrar
    """
    # Extrair palavras significativas do título (ignorar palavras comuns)
    STOPWORDS = {'de', 'da', 'do', 'das', 'dos', 'em', 'na', 'no', 'nas', 'nos',
                 'para', 'por', 'com', 'sem', 'sobre', 'entre', 'até', 'como',
                 'uma', 'um', 'uns', 'umas', 'aos', 'às', 'e', 'ou', 'que', 'se'}

    titulo_lower = titulo.lower()
    titulo_clean = re.sub(r'[^\w\s]', ' ', titulo_lower)
    palavras_titulo = [w for w in titulo_clean.split() if len(w) > 3 and w not in STOPWORDS]

    if not palavras_titulo:
        return -1

    # Estratégia 1: Buscar sequência de 2-3 palavras-chave consecutivas
    for n_palavras in [3, 2]:
        if len(palavras_titulo) >= n_palavras:
            busca = ' '.join(palavras_titulo[:n_palavras])
            pos = texto_lower.find(busca)
            if pos != -1:
                # Voltar até início da frase/linha
                while pos > 0 and transcricao_completa[pos - 1] not in '.\n':
                    pos -= 1
                    if pos < len(transcricao_completa) - 200:  # Limite de 200 chars para trás
                        break
                return pos

    # Estratégia 2: Buscar frases de transição + primeira palavra-chave
    FRASES_TRANSICAO = [
        'vamos agora', 'passemos para', 'vamos falar', 'vamos tratar',
        'o próximo tema', 'o próximo ponto', 'agora vamos',
        'entrando no', 'entrando em', 'passando para', 'passando ao',
        'quanto ao', 'quanto à', 'em relação ao', 'em relação à',
        'no que tange', 'no que diz respeito', 'sobre o tema',
        'começando por', 'iniciando com', 'primeiro tema',
        'vamos começar', 'vamos iniciar'
    ]

    for frase in FRASES_TRANSICAO:
        pos_transicao = texto_lower.find(frase)
        if pos_transicao != -1:
            # Verificar se alguma palavra-chave do título está próxima (até 200 chars depois)
            zona = texto_lower[pos_transicao:pos_transicao + 200]
            for palavra in palavras_titulo[:2]:
                if palavra in zona:
                    # Voltar até início da linha
                    while pos_transicao > 0 and transcricao_completa[pos_transicao - 1] not in '\n':
                        pos_transicao -= 1
                    return pos_transicao

    # Estratégia 3: Buscar apenas a primeira palavra-chave significativa
    if palavras_titulo:
        palavra_principal = max(palavras_titulo[:3], key=len) if len(palavras_titulo) >= 3 else palavras_titulo[0]
        pos = texto_lower.find(palavra_principal)
        if pos != -1:
            # Voltar até início da linha
            while pos > 0 and transcricao_completa[pos - 1] not in '\n':
                pos -= 1
            return pos

    return -1


def limpar_estrutura_para_review(mapping: str) -> str:
    """
    v2.25: Remove metadados de âncora (ABRE/FECHA) do mapeamento para uso em ai_structure_review.

    Transforma:
        1. Introdução | ABRE: "frase" | FECHA: "frase"
    Em:
        1. Introdução
    """
    if not mapping:
        return mapping
    return re.sub(r'\s*\|\s*(?:ABRE|FECHA):\s*["\'][^"\']*["\']', '', mapping)
    
def filtrar_niveis_excessivos(estrutura: str, max_nivel: int = 3) -> str:
    """
    v2.41: Remove itens da estrutura mais profundos que max_nivel.
    Ex.: se max_nivel=3, remove 1.1.1.1.
    Portado de format_transcription_gemini.py.
    """
    if not estrutura:
        return estrutura
    linhas = estrutura.strip().split('\n')
    filtradas = []
    removidos = 0
    for linha in linhas:
        match = re.match(r'^(\d+(?:\.\d+)*)', linha.strip())
        if match:
            partes = [p for p in match.group(1).split('.') if p.isdigit()]
            if len(partes) <= max_nivel:
                filtradas.append(linha)
            else:
                removidos += 1
        else:
            filtradas.append(linha)
    if removidos:
        print(f"{Fore.CYAN}✂️  Filtrados {removidos} itens com nível > {max_nivel}{Style.RESET_ALL}")
    return '\n'.join(filtradas)


def _sample_evenly(items: list[str], limit: int) -> list[str]:
    """Seleciona itens distribuídos ao longo da lista preservando início e fim."""
    if limit <= 0 or len(items) <= limit:
        return list(items)
    if limit == 1:
        return [items[0]]

    step = (len(items) - 1) / (limit - 1)
    selected = []
    used = set()
    for i in range(limit):
        idx = int(round(i * step))
        idx = max(0, min(len(items) - 1, idx))
        if idx not in used:
            selected.append(items[idx])
            used.add(idx)

    # Se o arredondamento gerar menos itens únicos, completa com faltantes na ordem.
    if len(selected) < limit:
        for item in items:
            if item not in selected:
                selected.append(item)
            if len(selected) >= limit:
                break
    return selected[:limit]


def _extract_outline_key(line: str) -> str | None:
    """
    Extrai chave numérica de outline (ex.: '1.2.3') de uma linha.
    Retorna None para linhas sem numeração hierárquica detectável.
    """
    if not line:
        return None
    m = re.match(r'^\s*(\d+(?:\.\d+)*)\.?\s+', line.strip())
    return m.group(1) if m else None


def _sample_with_parents(items: list[str], limit: int) -> list[str]:
    """
    Amostra distribuída com fechamento pai-filho:
    se um item filho for selecionado, inclui seus pais quando presentes.
    """
    if limit <= 0 or len(items) <= limit:
        return list(items)

    key_to_idx: dict[str, int] = {}
    idx_to_key: dict[int, str] = {}
    for idx, line in enumerate(items):
        key = _extract_outline_key(line)
        if not key:
            continue
        # Mantém a primeira ocorrência para preservar ordem natural do outline.
        if key not in key_to_idx:
            key_to_idx[key] = idx
            idx_to_key[idx] = key

    def _closure(sampled_indexes: set[int]) -> set[int]:
        expanded = set(sampled_indexes)
        for idx in list(sampled_indexes):
            key = idx_to_key.get(idx)
            if not key:
                continue
            parts = key.split(".")
            for depth in range(len(parts) - 1, 0, -1):
                parent_key = ".".join(parts[:depth])
                parent_idx = key_to_idx.get(parent_key)
                if parent_idx is not None:
                    expanded.add(parent_idx)
        return expanded

    def _sample_indexes_evenly(indexes: list[int], sample_limit: int) -> list[int]:
        if sample_limit <= 0 or len(indexes) <= sample_limit:
            return list(indexes)
        if sample_limit == 1:
            return [indexes[0]]
        step = (len(indexes) - 1) / (sample_limit - 1)
        picked: list[int] = []
        used_pos = set()
        for i in range(sample_limit):
            pos = int(round(i * step))
            pos = max(0, min(len(indexes) - 1, pos))
            if pos not in used_pos:
                picked.append(indexes[pos])
                used_pos.add(pos)
        if len(picked) < sample_limit:
            for idx in indexes:
                if idx not in picked:
                    picked.append(idx)
                if len(picked) >= sample_limit:
                    break
        return picked[:sample_limit]

    # Ajusta a amostra-base até caber junto com ancestrais.
    base_limit = min(limit, len(items))
    selected_indexes: set[int] = set()
    all_indexes = list(range(len(items)))
    while base_limit > 0:
        base_indexes = set(_sample_indexes_evenly(all_indexes, base_limit))
        expanded = _closure(base_indexes)
        if len(expanded) <= limit:
            selected_indexes = expanded
            break
        base_limit -= 1

    if not selected_indexes:
        selected_indexes = {0}

    # Preenche vagas remanescentes distribuindo no restante sem quebrar o fechamento já construído.
    remaining_slots = limit - len(selected_indexes)
    if remaining_slots > 0:
        remaining = [i for i in range(len(items)) if i not in selected_indexes]
        if remaining:
            sampled_remaining = _sample_indexes_evenly(remaining, remaining_slots)
            for idx in sampled_remaining:
                selected_indexes.add(idx)
                if len(selected_indexes) >= limit:
                    break

    ordered_indexes = sorted(selected_indexes)[:limit]
    return [items[i] for i in ordered_indexes]


def simplificar_estrutura_se_necessario(
    estrutura: str,
    max_linhas: int = 120,
    max_nivel: int = 3,
) -> str:
    """
    v2.42: Se a estrutura tiver mais de max_linhas itens, preserva níveis até max_nivel
    para evitar prompt bloat nos chunks.
    Portado de format_transcription_gemini.py.
    """
    if not estrutura:
        return estrutura
    max_nivel = max(1, min(3, int(max_nivel or 3)))
    linhas = [l for l in estrutura.strip().split('\n') if l.strip()]
    if len(linhas) <= max_linhas:
        return estrutura

    print(
        f"{Fore.CYAN}📉 Estrutura longa ({len(linhas)} itens). "
        f"Simplificando para níveis 1-{max_nivel}, máx {max_linhas}...{Style.RESET_ALL}"
    )
    nivel1 = []
    nivel2 = []
    nivel3 = []
    for l in linhas:
        s = l.strip()
        if re.match(r'^\d+\.\s', s):
            nivel1.append(l)
        elif re.match(r'^\d+\.\d+\.?\s', s):
            nivel2.append(l)
        elif re.match(r'^\d+\.\d+\.\d+\.?\s', s):
            nivel3.append(l)

    retained = set(nivel1 + nivel2)
    if max_nivel >= 3:
        retained.update(nivel3)

    if len(retained) < 5:
        return estrutura  # fallback

    nova = []
    vistos = set()
    for l in linhas:
        if l in vistos:
            continue
        if l in retained:
            nova.append(l)
            vistos.add(l)
    if len(nova) > max_linhas:
        # Evita viés para o começo e preserva coerência pai-filho.
        nova = _sample_with_parents(nova, max_linhas)

    print(
        f"{Fore.GREEN}✅ Estrutura simplificada: {len(linhas)} → {len(nova)} linhas "
        f"(níveis 1-{max_nivel}).{Style.RESET_ALL}"
    )
    return '\n'.join(nova)


def dividir_sequencial(transcricao_completa, chars_por_parte=25000, estrutura_global=None):
    """
    v2.26: Divide documento em chunks SEQUENCIAIS com preferência por âncoras verbatim.
    
    Melhorias v2.26:
    - Log de cobertura de âncoras (quantas foram encontradas)
    - Validação FECHA (verifica se chunk termina onde esperado)
    - Flag 'instituto_continua' quando um instituto é partido
    
    Args:
        transcricao_completa: Texto bruto completo
        chars_por_parte: Tamanho alvo de cada chunk
        estrutura_global: String com a estrutura mapeada (opcional, com âncoras ABRE/FECHA)
    
    Returns:
        Lista de dicts com 'inicio', 'fim', e metadados de continuidade
    """
    chunks = []
    tamanho_total = len(transcricao_completa)
    inicio = 0
    texto_lower = transcricao_completa.lower()
    
    # v2.26: Estrutura para rastrear âncoras e cobertura
    ancoras_info = []  # Lista de dicts com 'titulo', 'abre_frase', 'fecha_frase', 'abre_pos', 'fecha_pos'
    ancoras_encontradas = 0
    ancoras_totais = 0
    ancoras_nao_encontradas = []
    
    # v2.25: Extrair âncoras verbatim (ABRE/FECHA) da estrutura
    pontos_de_corte = []  # Lista de posições absolutas onde cortar
    
    if estrutura_global:
        # Regex para capturar: NUMERO. Título | ABRE: "frase" | FECHA: "frase"
        anchor_pattern = re.compile(
            r'^\s*(\d+(?:\.\d+)*)\.\s*([^|]+)\|\s*ABRE:\s*["\']([^"\']+)["\']\s*\|\s*FECHA:\s*["\']([^"\']+)["\']',
            re.MULTILINE | re.IGNORECASE
        )
        
        for match in anchor_pattern.finditer(estrutura_global):
            numero = match.group(1)
            titulo = match.group(2).strip()
            frase_abre = match.group(3).strip().lower()
            frase_fecha = match.group(4).strip().lower()
            ancoras_totais += 1

            if len(frase_abre) < 10:
                continue  # Âncora muito curta, pular

            # v2.33: Detectar âncora "fake" (modelo usou título em vez de citação verbatim)
            similaridade = _similaridade_palavras(titulo, frase_abre)
            ancora_fake = similaridade > 0.6  # Mais de 60% de overlap = provavelmente fake

            if ancora_fake:
                print(f"{Fore.YELLOW}   ⚠️  Âncora fake detectada (sim={similaridade:.0%}): '{frase_abre[:30]}...'")

            # Buscar a frase ABRE no texto
            pos_abre = texto_lower.find(frase_abre) if not ancora_fake else -1
            pos_fecha = None

            if pos_abre == -1 and not ancora_fake:
                # Tentar busca fuzzy com as primeiras 5 palavras
                palavras = frase_abre.split()[:5]
                frase_curta = ' '.join(palavras)
                pos_abre = texto_lower.find(frase_curta)
                if pos_abre != -1:
                    print(f"{Fore.YELLOW}   📍 Âncora parcial: '{frase_curta}' @ {pos_abre}")

            if pos_abre == -1 and frase_abre and not ancora_fake:
                # Fallback: busca tolerante a quebras de linha (whitespace-insensitive)
                try:
                    pattern = r'\s+'.join(re.escape(w) for w in frase_abre.split())
                    m = re.search(pattern, transcricao_completa, flags=re.IGNORECASE)
                    if m:
                        pos_abre = m.start()
                        print(f"{Fore.YELLOW}   📍 Âncora com whitespace-flex: '{frase_abre[:40]}...' @ {pos_abre}")
                except re.error:
                    pass

            # v2.33: Fallback inteligente para âncoras fake ou não encontradas
            if pos_abre == -1:
                pos_abre = _buscar_ancora_no_texto(texto_lower, titulo, transcricao_completa)
                if pos_abre != -1:
                    metodo = "busca por título" if ancora_fake else "fallback inteligente"
                    print(f"{Fore.CYAN}   🔍 Âncora via {metodo}: '{titulo[:30]}...' @ {pos_abre}")

            if pos_abre != -1:
                # Voltar até o início da linha/parágrafo
                while pos_abre > 0 and transcricao_completa[pos_abre - 1] not in '\n':
                    pos_abre -= 1
                pontos_de_corte.append(pos_abre)
                ancoras_encontradas += 1
                if not ancora_fake:
                    print(f"{Fore.GREEN}   📍 Âncora ABRE: '{titulo[:30]}...' @ {pos_abre}")

                # v2.26: Buscar FECHA para validação
                if frase_fecha.lower() != 'fim':
                    pos_fecha = texto_lower.find(frase_fecha)
                    if pos_fecha != -1:
                        print(f"{Fore.CYAN}   📍 Âncora FECHA: '{frase_fecha[:30]}...' @ {pos_fecha}")
            else:
                ancoras_nao_encontradas.append(f"{numero}. {titulo}")
                print(f"{Fore.RED}   ❌ Âncora não encontrada: '{titulo[:40]}...'")

            ancoras_info.append({
                'numero': numero,
                'titulo': titulo,
                'abre_frase': frase_abre,
                'fecha_frase': frase_fecha,
                'abre_pos': pos_abre if pos_abre != -1 else None,
                'fecha_pos': pos_fecha
            })
        
        # Ordenar e remover duplicatas
        pontos_de_corte = sorted(set(pontos_de_corte))
        
        # v2.26: Log de cobertura de âncoras
        if ancoras_totais > 0:
            cobertura = (ancoras_encontradas / ancoras_totais) * 100
            cor = Fore.GREEN if cobertura >= 80 else (Fore.YELLOW if cobertura >= 50 else Fore.RED)
            print(f"{cor}   📊 Cobertura de âncoras: {ancoras_encontradas}/{ancoras_totais} ({cobertura:.0f}%)")
            if ancoras_nao_encontradas:
                print(f"{Fore.YELLOW}   ⚠️  Não localizadas: {', '.join(ancoras_nao_encontradas[:5])}" + 
                      (f" (+{len(ancoras_nao_encontradas)-5} mais)" if len(ancoras_nao_encontradas) > 5 else ""))
    
    # Fallback: extrair âncoras antigas (primeiras 3 palavras do título)
    ancoras_fallback = []
    if estrutura_global and not pontos_de_corte:
        for line in estrutura_global.split('\n'):
            line = line.strip()
            # Remover metadados de âncora para extrair só o título
            line_clean = re.sub(r'\s*\|.*$', '', line)
            match = re.match(r'^\d+(?:\.\d+)*\.?\s+(.+)', line_clean)
            if match:
                titulo = match.group(1).strip()
                palavras = [w for w in titulo.split() if len(w) > 3][:3]
                if palavras:
                    ancoras_fallback.append(' '.join(palavras).lower())
    
    # v2.26: Construir mapa de intervalos de institutos (para detectar cortes no meio)
    intervalos_institutos = []
    for i, info in enumerate(ancoras_info):
        if info['abre_pos'] is not None:
            fim_instituto = tamanho_total  # Default: até o fim
            # O fim do instituto é o início do próximo (ou fim do texto)
            for j in range(i + 1, len(ancoras_info)):
                if ancoras_info[j]['abre_pos'] is not None:
                    fim_instituto = ancoras_info[j]['abre_pos']
                    break
            intervalos_institutos.append({
                'titulo': info['titulo'],
                'inicio': info['abre_pos'],
                'fim': fim_instituto
            })
    
    while inicio < tamanho_total:
        fim_ideal = min(inicio + chars_por_parte, tamanho_total)
        fim = fim_ideal
        instituto_continua = False
        instituto_nome = None
        
        if fim < tamanho_total:
            bloco = transcricao_completa[inicio:fim]
            melhor_ponto = None
            
            # ESTRATÉGIA 1: Usar pontos de corte de âncoras verbatim
            if pontos_de_corte:
                # Encontrar o ponto de corte mais próximo do fim_ideal (nos últimos 30%)
                limite_inferior = inicio + int(chars_por_parte * 0.7)
                for ponto in pontos_de_corte:
                    if limite_inferior <= ponto < fim_ideal:
                        melhor_ponto = ponto
                        print(f"{Fore.GREEN}   ✂️  Cortando em âncora verbatim @ {ponto}")
                        break
            
            # ESTRATÉGIA 2: Fallback para âncoras antigas (primeiras 3 palavras)
            if melhor_ponto is None and ancoras_fallback:
                zona_busca = bloco[int(chars_por_parte * 0.7):]
                zona_offset = int(chars_por_parte * 0.7)
                
                for ancora in ancoras_fallback:
                    pos = zona_busca.lower().find(ancora)
                    if pos != -1:
                        ponto_corte = zona_offset + pos
                        while ponto_corte > 0 and bloco[ponto_corte] != '\n':
                            ponto_corte -= 1
                        if ponto_corte > chars_por_parte * 0.7:
                            melhor_ponto = inicio + ponto_corte
                            break
            
            # ESTRATÉGIA 3: Fallback para fim de parágrafo
            if melhor_ponto is None:
                ultimo_paragrafo = bloco.rfind('\n\n')
                if ultimo_paragrafo != -1 and ultimo_paragrafo > chars_por_parte * 0.8:
                    melhor_ponto = inicio + ultimo_paragrafo + 2
            
            if melhor_ponto:
                fim = melhor_ponto
            
            # v2.26: Verificar se estamos cortando no meio de um instituto
            for intervalo in intervalos_institutos:
                # Se o chunk começa dentro de um instituto e termina antes do fim dele
                if intervalo['inicio'] <= inicio < intervalo['fim'] and fim < intervalo['fim']:
                    instituto_continua = True
                    instituto_nome = intervalo['titulo']
                    print(f"{Fore.YELLOW}   ⚠️  Instituto '{instituto_nome[:30]}...' será continuado no próximo chunk")
                    break
        
        chunks.append({
            'inicio': inicio, 
            'fim': fim,
            'instituto_continua': instituto_continua,
            'instituto_nome': instituto_nome
        })
        inicio = fim
        
    return chunks

def dividir_por_blocos_markdown(
    texto: str,
    *,
    max_chars: int = 25000,
    block_prefix_pattern: Optional[str] = None,
    split_overlap_chars: int = 300,
) -> list:
    """
    Divide por blocos naturais em Markdown (v2.32).

    Detecta headings "## Bloco XX — ..." (ou outros prefixos) e agrupa blocos inteiros até atingir `max_chars`.
    Se um bloco exceder `max_chars`, ele é subdividido via `chunk_texto_seguro`.

    Retorna o mesmo formato de `dividir_sequencial`: lista de dicts {inicio, fim, ...}.
    """
    texto = texto or ""
    if not texto.strip():
        return []

    # Encontrar blocos por headings (prefixo configurável)
    prefix = block_prefix_pattern or os.getenv("IUDEX_HEARING_BLOCK_PREFIX_REGEX", r"Bloco|Ato|Parte")
    try:
        block_regex = re.compile(rf'(?m)^##\s+(?:{prefix})\b', flags=re.IGNORECASE)
    except re.error:
        block_regex = re.compile(r'(?m)^##\s+Bloco\b', flags=re.IGNORECASE)
    matches = list(block_regex.finditer(texto))
    if len(matches) < 2:
        # Poucos blocos → não vale chunking por bloco
        return []

    block_ranges: list[tuple[int, int]] = []
    for idx, m in enumerate(matches):
        start = m.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(texto)
        block_ranges.append((start, end))

    chunks: list[dict] = []
    cur_start: Optional[int] = None
    cur_end: Optional[int] = None
    cur_len = 0
    cur_blocks: list[int] = []

    def _flush():
        nonlocal cur_start, cur_end, cur_len, cur_blocks
        if cur_start is None or cur_end is None:
            return
        chunks.append(
            {
                "inicio": cur_start,
                "fim": cur_end,
                "block_ids": cur_blocks[:],
                "instituto_continua": False,
                "instituto_nome": None,
            }
        )
        cur_start = None
        cur_end = None
        cur_len = 0
        cur_blocks = []

    for block_idx, (b_start, b_end) in enumerate(block_ranges):
        block_len = b_end - b_start
        if block_len > max_chars:
            # Se há chunk em andamento, fecha antes de subdividir bloco grande
            _flush()
            block_text = texto[b_start:b_end]
            overlap_chars = max(0, int(split_overlap_chars))
            parts = chunk_texto_seguro(block_text, max_chars=max_chars, overlap_chars=overlap_chars)
            pos = b_start
            for part in parts:
                part_len = len(part)
                end_pos = min(len(texto), pos + part_len)
                chunks.append(
                    {
                        "inicio": pos,
                        "fim": end_pos,
                        "block_ids": [block_idx],
                        "instituto_continua": False,
                        "instituto_nome": None,
                    }
                )
                pos = end_pos
            continue

        if cur_start is None:
            cur_start = b_start
            cur_end = b_end
            cur_len = block_len
            cur_blocks = [block_idx]
            continue

        if cur_len + block_len > max_chars and cur_blocks:
            _flush()
            cur_start = b_start
            cur_end = b_end
            cur_len = block_len
            cur_blocks = [block_idx]
            continue

        cur_end = b_end
        cur_len += block_len
        cur_blocks.append(block_idx)

    _flush()

    # Garantir contiguidade (sem gaps). Se houver gaps, fallback para dividir_sequencial.
    expected = 0
    for c in chunks:
        if c["inicio"] != expected:
            return []
        expected = c["fim"]
    if expected != len(texto):
        return []

    return chunks

def validar_chunks(chunks, texto_completo):
    """Valida se não houve perda de texto entre chunks"""
    esperado = 0
    for c in chunks:
        if c['inicio'] != esperado:
             print(f"{Fore.RED}⚠️ GAP detectado! Chunk começa em {c['inicio']} mas devia ser {esperado}")
        esperado = c['fim']
    
    if esperado != len(texto_completo):
        print(f"{Fore.RED}⚠️ Texto incompleto! Processado: {esperado}, Total: {len(texto_completo)}")
    else:
        print(f"{Fore.GREEN}✅ Divisão de chunks validada (Bytes match).")

# Removed old detectar_secoes_duplicadas & remover_secoes_duplicadas since they are replaced by robust versions
# The new versions are placed in the helper section we just updated.

def remover_duplicacoes_literais(texto):
    """Remove parágrafos individuais duplicados (v2.7 logic)"""
    from difflib import SequenceMatcher
    paragrafos = texto.split('\n\n')
    paragrafos_limpos = []
    dup_count = 0
    
    for i, para in enumerate(paragrafos):
        if i == 0:
            paragrafos_limpos.append(para)
            continue
        
        if len(para.strip()) < 80 or para.strip().startswith('#'):
            paragrafos_limpos.append(para)
            continue
        
        is_duplicate = False
        para_norm = ' '.join(para.split()).lower()
        
        # Check against last 3 paragraphs (v2.7 logic)
        for j in range(max(0, len(paragrafos_limpos) - 3), len(paragrafos_limpos)):
            para_ant = paragrafos_limpos[j]
            para_ant_norm = ' '.join(para_ant.split()).lower()
            
            ratio = SequenceMatcher(None, para_norm, para_ant_norm).ratio()
            
            if ratio > 0.95:
                is_duplicate = True
                dup_count += 1
                break
        
        if not is_duplicate:
            paragrafos_limpos.append(para)
    
    if dup_count > 0:
         print(f"⚠️  {dup_count} parágrafos duplicados removidos (Literal Dedup)")
    
    return '\n\n'.join(paragrafos_limpos)


def numerar_titulos(texto):
    """Adiciona numeração sequencial aos títulos"""
    linhas = texto.split('\n')
    linhas_numeradas = []
    
    contador_h2 = 0
    contador_h3 = 0
    contador_h4 = 0
    
    titulo_pattern = re.compile(r'^(#{2,4})\s+(?:\d+(?:\.\d+)*\.?\s+)?(.+)$')
    
    for linha in linhas:
        match = titulo_pattern.match(linha)
        
        if match:
            nivel = len(match.group(1))
            texto_titulo = match.group(2).strip()
            
            # Não numera títulos de resumo/quadros
            if any(keyword in texto_titulo.lower() for keyword in ['resumo', 'quadro', 'esquema', '📋', '📊', '🗂️']):
                linhas_numeradas.append(linha)
                continue
            
            if nivel == 2:
                contador_h2 += 1
                contador_h3 = 0
                contador_h4 = 0
                nova_linha = f"## {contador_h2}. {texto_titulo}"
            elif nivel == 3:
                contador_h3 += 1
                contador_h4 = 0
                nova_linha = f"### {contador_h2}.{contador_h3}. {texto_titulo}"
            elif nivel == 4:
                contador_h4 += 1
                nova_linha = f"#### {contador_h2}.{contador_h3}.{contador_h4}. {texto_titulo}"
            else:
                nova_linha = linha
            
            linhas_numeradas.append(nova_linha)
        else:
            linhas_numeradas.append(linha)
    
    return '\n'.join(linhas_numeradas)


# ==================== METRICS COLLECTOR (v2.10) ====================
class MetricsCollector:
    """Tracks API usage, timing, and cost estimation for optimization."""
    
    # Preços Gemini 3 Flash Preview (Dezembro 2025) - USD per 1M tokens
    PRICE_INPUT = 0.50   # $0.50 (Gemini 3 Flash Preview Input)
    PRICE_OUTPUT = 3.00  # $3.00 (Gemini 3 Flash Preview Output)
    
    # OpenAI GPT-5 Mini (Estimado)
    PRICE_OPENAI_INPUT = 0.15
    PRICE_OPENAI_OUTPUT = 0.60
    
    def __init__(self, provider="gemini"):
        self.provider = provider
        self.reset()
    
    def reset(self):
        self.api_calls = 0
        self.gemini_calls = 0
        self.openai_calls = 0
        self.total_prompt_tokens = 0
        self.total_completion_tokens = 0
        self.total_time_seconds = 0.0
        self.call_times = []
        self.cache_hits = 0
        self.adaptive_splits = 0
    
    def set_provider(self, provider: str):
        """Updates the provider for cost calculation."""
        self.provider = provider
    
    def record_call(
        self,
        provider: str,
        prompt_tokens: int,
        completion_tokens: int,
        duration: float,
        model: Optional[str] = None,
        cached_tokens_in: Optional[int] = None,
    ):
        """Records a single API call."""
        self.api_calls += 1
        if provider == "gemini":
            self.gemini_calls += 1
        elif provider == "openai":
            self.openai_calls += 1
        self.total_prompt_tokens += prompt_tokens
        self.total_completion_tokens += completion_tokens
        self.total_time_seconds += duration
        self.call_times.append(duration)
        _record_llm_usage(
            provider=provider,
            model=model,
            tokens_in=prompt_tokens,
            tokens_out=completion_tokens,
            cached_tokens_in=cached_tokens_in,
        )
    
    def record_cache_hit(self):
        self.cache_hits += 1
    
    def record_adaptive_split(self):
        self.adaptive_splits += 1
    
    def estimate_cost(self, provider="gemini") -> float:
        """Estimates total USD cost based on recorded tokens and provider pricing."""
        
        # Pricing Tables (USD per 1M tokens)
        PRICING = {
            "gemini": {
                "input": 0.075,       # Gemini 1.5 Pro
                "cached_input": 0.01875, # 25% of input
                "output": 0.30
            },
            "openai": {
                "input": 0.25,        # GPT-5 Mini
                "cached_input": 0.025, # 10x discount
                "output": 2.00
            }
        }
        
        # Fallback to gemini pricing if provider unknown
        prices = PRICING.get(provider, PRICING["gemini"])
        
        # Calculate cost
        # Note: cached_prompt_tokens tracking would be ideal, but for now we assume standard ratio or 0 if not tracked
        # Here we use total_prompt_tokens for standard input cost calculation
        
        # Separate cached vs uncached if available (future proofing)
        # For current implementation, we assume all prompts are uncached for conservative estimate
        # unless specific metric is added.
        
        cost = (
            (self.total_prompt_tokens * prices["input"] / 1_000_000) +
            (self.total_completion_tokens * prices["output"] / 1_000_000)
        )
        return cost
    
    def get_summary(self) -> str:
        """Returns a formatted summary string."""
        avg_time = (self.total_time_seconds / self.api_calls) if self.api_calls > 0 else 0
        cost = self.estimate_cost(self.provider)
        
        return f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 MÉTRICAS DE EXECUÇÃO
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
   📡 Total de Chamadas API: {self.api_calls}
      - Gemini: {self.gemini_calls}
      - OpenAI: {self.openai_calls}
      - Cache Hits: {self.cache_hits}
   ✂️ Divisões Adaptativas: {self.adaptive_splits}
   🎯 Tokens Usados:
      - Prompt: {self.total_prompt_tokens:,}
      - Completion: {self.total_completion_tokens:,}
      - Total: {self.total_prompt_tokens + self.total_completion_tokens:,}
   ⏱️ Tempo Total: {self.total_time_seconds:.1f}s (média: {avg_time:.2f}s/chamada)
   💰 Custo Estimado: ${cost:.4f} USD
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

# Global metrics instance
metrics = MetricsCollector()

# ==================== CHECKPOINT SYSTEM ====================
def get_checkpoint_path(video_name, folder):
    return Path(folder) / f"{video_name}.checkpoint.json"

def save_checkpoint(video_name, folder, results, segments_info, current_idx):
    path = get_checkpoint_path(video_name, folder)
    data = {
        'video_name': video_name,
        'current_idx': current_idx,
        'total_segments': len(segments_info),
        'results': results,
        'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
    }
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_checkpoint(video_name, folder):
    path = get_checkpoint_path(video_name, folder)
    if path.exists():
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️ Erro ao carregar checkpoint: {e}")
    return None

def delete_checkpoint(video_name, folder):
    path = get_checkpoint_path(video_name, folder)
    if path.exists(): 
        path.unlink()
        print(f"🧹 Checkpoint removido: {path.name}")

def get_hil_output_path(video_name, folder, mode_suffix):
    return Path(folder) / f"{video_name}_{mode_suffix}_HIL.md"

def save_hil_output(formatted_text, video_name, folder, mode_suffix, reason=None):
    """Salva o texto formatado para revisão humana (HIL)."""
    path = get_hil_output_path(video_name, folder, mode_suffix)
    reason_note = f" | motivo: {reason}" if reason else ""
    header = f"<!-- HIL_CHECKPOINT{reason_note} | {time.strftime('%Y-%m-%d %H:%M:%S')} -->\n\n"
    try:
        with open(path, 'w', encoding='utf-8') as f:
            f.write(header)
            f.write(formatted_text or "")
        print(f"{Fore.YELLOW}⏸️  HIL checkpoint salvo: {path.name}")
    except Exception as e:
        print(f"{Fore.YELLOW}⚠️ Falha ao salvar HIL checkpoint: {e}")

class VomoMLX:
    # GPT-5 Mini: 400k tokens input, 128k output
    MAX_CHUNK_SIZE = 100000  
    CHUNK_OVERLAP = 3000     # 5k overlap
    # Map structure settings (v2.31)
    MAP_MAX_SINGLE_CHARS = 350_000
    MAP_CHUNK_CHARS = 150_000
    MAP_CHUNK_OVERLAP_CHARS = 8_000
    MAP_MAX_LINES_PER_CHUNK = 60
    RAW_CONTEXT_OVERLAP_CHARS = 1200

    # ==================== PARALELIZAÇÃO (v2.40) ====================
    # Número de chunks processados em paralelo. Default=1 (sequencial).
    # Valores > 1 aceleram mas podem reduzir consistência de estilo entre chunks.
    # Recomendado: 2-3 para balanço velocidade/qualidade.
    PARALLEL_CHUNKS = int(os.getenv("IUDEX_PARALLEL_CHUNKS", "1"))

    # ==================== CHUNKING DE ÁUDIO LONGO (v2.32/v2.34) ====================
    # Áudios maiores que este limite serão divididos em partes para evitar
    # problemas de memória/processamento do MLX-Whisper com arquivos muito longos
    AUDIO_MAX_DURATION_SECONDS = 2 * 60 * 60  # 2 horas
    AUDIO_CHUNK_OVERLAP_SECONDS = 45  # v2.34: 45s de overlap (era 30s) - mais seguro para frases longas
    # NOTA: Diarização em chunking pode resetar speaker IDs entre chunks.
    # Para diarização consistente em áudios longos, use AssemblyAI ou diarize o áudio inteiro separadamente.

    # ==================== IDIOMAS SUPORTADOS ====================
    # Expandido para incluir os principais idiomas suportados por Whisper e AssemblyAI
    SUPPORTED_LANGUAGES = {
        "auto": None,   # Whisper detecta automaticamente
        "pt": "pt",     # Português
        "en": "en",     # Inglês
        "es": "es",     # Espanhol
        "fr": "fr",     # Francês
        "de": "de",     # Alemão
        "it": "it",     # Italiano
        "ja": "ja",     # Japonês
        "ko": "ko",     # Coreano
        "zh": "zh",     # Chinês
        "ru": "ru",     # Russo
        "ar": "ar",     # Árabe
        "hi": "hi",     # Hindi
        "nl": "nl",     # Holandês
        "pl": "pl",     # Polonês
        "tr": "tr",     # Turco
        "sv": "sv",     # Sueco
        "da": "da",     # Dinamarquês
        "fi": "fi",     # Finlandês
        "no": "no",     # Norueguês
        "uk": "uk",     # Ucraniano
    }

    # ==================== INITIAL PROMPTS POR MODO (v2.29) ====================
    # Contexto do Whisper ajustado ao tipo de áudio para melhor reconhecimento de termos
    # Chave externa: (modo, idioma). Fallback: só modo (assume pt).
    INITIAL_PROMPTS = {
        "APOSTILA": "Esta é uma transcrição de aula jurídica em português brasileiro sobre direito administrativo, constitucional, civil, penal e processual.",
        "FIDELIDADE": "Esta é uma transcrição de aula jurídica em português brasileiro sobre direito administrativo, constitucional, civil, penal e processual.",
        "AUDIENCIA": "Esta é uma transcrição de audiência judicial em português brasileiro. Termos forenses, procedimentos processuais e linguagem jurídica formal.",
        "REUNIAO": "Esta é uma transcrição de reunião profissional em português brasileiro.",
        "DEPOIMENTO": "Esta é uma transcrição de depoimento judicial em português brasileiro. Termos forenses e linguagem jurídica formal.",
    }

    INITIAL_PROMPTS_I18N: dict[tuple[str, str], str] = {
        # Inglês
        ("APOSTILA", "en"): "This is a transcription of a legal lecture in English about administrative, constitutional, civil, criminal and procedural law.",
        ("FIDELIDADE", "en"): "This is a transcription of a legal lecture in English about administrative, constitutional, civil, criminal and procedural law.",
        ("AUDIENCIA", "en"): "This is a transcription of a court hearing in English. Forensic terms, procedural law and formal legal language.",
        ("REUNIAO", "en"): "This is a transcription of a professional meeting in English.",
        ("DEPOIMENTO", "en"): "This is a transcription of a legal deposition in English. Forensic terms and formal legal language.",
        # Espanhol
        ("APOSTILA", "es"): "Esta es una transcripción de una clase jurídica en español sobre derecho administrativo, constitucional, civil, penal y procesal.",
        ("FIDELIDADE", "es"): "Esta es una transcripción de una clase jurídica en español sobre derecho administrativo, constitucional, civil, penal y procesal.",
        ("AUDIENCIA", "es"): "Esta es una transcripción de una audiencia judicial en español. Términos forenses, procedimientos procesales y lenguaje jurídico formal.",
        ("REUNIAO", "es"): "Esta es una transcripción de una reunión profesional en español.",
        ("DEPOIMENTO", "es"): "Esta es una transcripción de una declaración judicial en español. Términos forenses y lenguaje jurídico formal.",
        # Francês
        ("APOSTILA", "fr"): "Ceci est une transcription d'un cours juridique en français sur le droit administratif, constitutionnel, civil, pénal et procédural.",
        ("FIDELIDADE", "fr"): "Ceci est une transcription d'un cours juridique en français sur le droit administratif, constitutionnel, civil, pénal et procédural.",
        ("AUDIENCIA", "fr"): "Ceci est une transcription d'une audience judiciaire en français. Termes forensiques, procédures judiciaires et langage juridique formel.",
        ("REUNIAO", "fr"): "Ceci est une transcription d'une réunion professionnelle en français.",
        ("DEPOIMENTO", "fr"): "Ceci est une transcription d'une déposition judiciaire en français. Termes forensiques et langage juridique formel.",
        # Alemão
        ("APOSTILA", "de"): "Dies ist eine Transkription einer juristischen Vorlesung auf Deutsch über Verwaltungs-, Verfassungs-, Zivil-, Straf- und Verfahrensrecht.",
        ("FIDELIDADE", "de"): "Dies ist eine Transkription einer juristischen Vorlesung auf Deutsch über Verwaltungs-, Verfassungs-, Zivil-, Straf- und Verfahrensrecht.",
        ("AUDIENCIA", "de"): "Dies ist eine Transkription einer Gerichtsverhandlung auf Deutsch. Forensische Begriffe, Verfahrensrecht und formale juristische Sprache.",
        ("REUNIAO", "de"): "Dies ist eine Transkription eines professionellen Meetings auf Deutsch.",
        ("DEPOIMENTO", "de"): "Dies ist eine Transkription einer gerichtlichen Aussage auf Deutsch. Forensische Begriffe und formale juristische Sprache.",
    }

    # ==================== MODULAR PROMPT COMPONENTS (v2.22) ====================
    # These components are composed by _build_system_prompt to allow partial customization.
    
    # --- APOSTILA MODE ---
    PROMPT_HEAD_APOSTILA = """# DIRETRIZES DE REDAÇÃO: MANUAL JURÍDICO DIDÁTICO (MODO APOSTILA)

## PAPEL
VOCÊ É UM EXCELENTÍSSIMO REDATOR JURÍDICO E DIDÁTICO.
- **Tom:** doutrinário, impessoal, estilo manual de Direito.
- **Pessoa:** 3ª pessoa ou construções impessoais ("O professor explica...", "A doutrina define...").
- **Estilo:** prosa densa, porém com parágrafos curtos e didáticos.
- **Objetivo:** transformar a aula em texto de apostila/manual, sem alterar conteúdo nem inventar informações.

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias.
4. **NÃO CRIE PARÁGRAFOS LONGOS**. Máximo 3-6 linhas visuais por parágrafo.

## ❌ PRESERVE OBRIGATORIAMENTE
- **IDENTIFICAÇÃO DE FALANTES**: Se houver SPEAKER A/B/C ou similar, identifique o professor pelo contexto (quando ele se apresentar: "Eu sou o professor João", "Meu nome é Maria"). Substitua "SPEAKER X" pelo nome identificado. Se não identificar, use "Professor" ou "Palestrante".
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados, Temas de Repercussão Geral, Recursos Repetitivos. **NUNCA OMITA NÚMEROS DE TEMAS OU SÚMULAS**.
- **JURISPRUDÊNCIA**: Se o texto citar "Tema 424", "RE 123", "ADI 555", **MANTENHA O NÚMERO**. Não generalize para "jurisprudência do STJ".
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios.
- **Referências**: leis, artigos, jurisprudência (STF/STJ), autores, casos citados.
- **Ênfases intencionais** e **Observações pedagógicas**.

## 🎯 PRESERVAÇÃO ESPECIAL: DICAS DE PROVA E EXAMINADORES (CRÍTICO)
Aulas presenciais frequentemente contêm informações valiosas sobre:
1. **Referências a Examinadores**: Nomes de examinadores de concursos, suas preferências, posicionamentos ou temas favoritos. **PRESERVE INTEGRALMENTE**.
2. **Dicas de Prova**: Orientações sobre o que costuma cair em provas, pegadinhas comuns, temas recorrentes.
3. **Estratégias de Estudo**: Sugestões do professor sobre priorização, macetes, formas de memorização.
4. **Casos Práticos e Histórias Reais**: Exemplos de situações reais, casos julgados, histórias ilustrativas. **NUNCA RESUMA**.

> ⚠️ **ESSAS INFORMAÇÕES SÃO O DIFERENCIAL DE UMA AULA AO VIVO.** Sua omissão representa perda irreparável de valor didático."""

    PROMPT_STYLE_APOSTILA = """## ✅ DIRETRIZES DE ESTILO E FORMATAÇÃO VISUAL
1. **Correção Gramatical**: Ajuste a linguagem coloquial para o padrão culto.
2. **Limpeza**: Remova gírias, vocativos e cacoetes ("né", "tipo assim", "então", "meu irmão", "cara", "mano", "galera") e vícios de oralidade. Se houver parentesco factual (ex.: "Rodolfo (irmão do professor)"), mantenha a informação de forma formal.
3. **Coesão**: Use conectivos e pontuação adequada para tornar o texto fluido.
4. **Legibilidade Visual** (OBRIGATÓRIO):
   - **PARÁGRAFOS CURTOS**: máximo **3-5 linhas visuais** por parágrafo. **QUEBRE SEMPRE.**
   - **RECUOS COM MARCADORES**: Use `>` para citações, destaques ou observações importantes.
   - **NEGRITO MODERADO**: Destaque conceitos-chave com **negrito**, mas sem exagero.
   - **ITÁLICO**: Use para termos em latim, expressões estrangeiras ou ênfase leve.
5. **Formatação Didática** (use generosamente para legibilidade):
   - **Bullet points** (`-` ou `*`) para enumerar elementos, requisitos ou características.
   - **Listas numeradas** (`1.`, `2.`) para etapas, correntes doutrinárias ou exemplos ordenados.
   - **Marcadores relacionais** como `→` para consequências lógicas.
   - **Subseções** (###, ####) para organizar subtópicos dentro de um mesmo tema.

## 🎨 FORMATAÇÃO VISUAL AVANÇADA
Para garantir legibilidade superior:
1. **Após cada conceito importante**, quebre o parágrafo e inicie outro.
2. **Use listas** sempre que houver enumeração de mais de 2 itens.
3. **Use citações recuadas** (`>`) para destacar teses jurídicas, pontos polêmicos, observações práticas e dicas de prova.
4. **Separe visualmente** diferentes aspectos de um mesmo tema com subseções.
5. **Questões e Exercícios**: Se o professor ditar uma questão, exercício ou caso hipotético, **isole-o** em um bloco de citação:
   > **Questão:** O prazo para agravo de petição é de...
   - Separe claramente o enunciado da questão da explicação/gabarito subsequente.
6. **Destaques com Emojis** (use com moderação para facilitar escaneamento visual):
   - 💡 **Dica de Prova** ou **Observação Pedagógica**: Quando o professor der uma dica específica para provas ou concursos.
   - ⚠️ **Atenção** ou **Cuidado**: Para alertas, pegadinhas ou pontos polêmicos.
   - 📌 **Ponto Importante**: Para conceitos-chave que merecem destaque especial.
   - Exemplo: `> 💡 **Dica de Prova:** Esse tema caiu 3 vezes na PGM-Rio.`

## 💎 PILAR 1: ESTILO (VOZ ATIVA E DIRETA)
> 🚫 **PROIBIDO VOZ PASSIVA EXCESSIVA:** "Anunciou-se", "Informou-se".
> ✅ **PREFIRA VOZ ATIVA:** "O professor explica...", "A doutrina define...", "O Art. 37 estabelece..."."""

    PROMPT_STRUCTURE_APOSTILA = """## 📝 ESTRUTURA HIERÁRQUICA (CRÍTICO)

### REGRA DE OURO: TÓPICOS-MÃE COM SUBTÓPICOS
Organize o conteúdo em **hierarquia pai→filho**. Se o professor aborda múltiplos aspectos de um mesmo tema, eles devem ser **subtópicos** (###) de um **tópico-mãe** (##), NUNCA tópicos ## separados.

### NÍVEIS DE HIERARQUIA (MÁXIMO 3):
| Nível | Markdown | Uso | Exemplo |
|-------|----------|-----|---------|
| **Tema principal** | `##` | Mudança real de matéria/assunto | `## 2. Execução Fiscal` |
| **Subtema** | `###` | Aspecto, instituto ou marco legal dentro do tema | `### 2.1. Procedimento da LEF (Lei 6.830/80)` |
| **Detalhamento** | `####` | Detalhe específico, exemplo extenso ou ponto controverso | `#### 2.1.1. Citação por Hora Certa` |

### EXEMPLO DE HIERARQUIA CORRETA:
```
## 2. Execução Fiscal
### 2.1. Procedimento da LEF (Lei 6.830/80)
### 2.2. Súmula 314 do STJ — Citação por Hora Certa
### 2.3. Tema 444 do STJ — Redirecionamento ao Sócio
#### 2.3.1. Requisitos e Prazo
### 2.4. Exceção de Pré-Executividade
## 3. Embargos à Execução
### 3.1. Conceito e Natureza Jurídica
### 3.2. Hipóteses de Cabimento
```

### ❌ ERRADO (tudo como ## sem hierarquia):
```
## 2. Execução Fiscal
## 3. Procedimento da LEF          ← ERRADO! Deveria ser ### 2.1
## 4. Súmula 314 do STJ            ← ERRADO! Deveria ser ### 2.2
## 5. Tema 444 do STJ              ← ERRADO! Deveria ser ### 2.3
## 6. Exceção de Pré-Executividade  ← ERRADO! Deveria ser ### 2.4
```

### REGRAS ADICIONAIS:
- Mantenha a **sequência cronológica** exata das falas.
- **NÃO crie subtópicos para frases soltas** — use títulos APENAS para mudanças reais de assunto.
- Se uma frase parece título mas não inicia seção, use **negrito** no texto, não crie heading.
- **Marcos Legais** como subtópicos: Súmulas, Teses de Repercussão Geral e Artigos de Lei explicados em profundidade devem virar ### subtópicos (ex: `### 2.3. Súmula 314 do STJ`).
- **Anti-fragmentação**: Se o professor trata 4+ aspectos de um tema, TODOS devem ser ### sob um ## tema-mãe.
- Nunca use # (H1) para subtópicos — apenas para o título principal do documento."""

    PROMPT_TABLE_APOSTILA = """## 📊 QUADRO-SÍNTESE (OBRIGATÓRIO)
Ao final de CADA tópico principal (## ou ###), faça um fechamento didático com UM quadro-síntese.
SEMPRE que houver diferenciação de conceitos, prazos, procedimentos, requisitos ou regras, o quadro é OBRIGATÓRIO.

1) Adicione um subtítulo de fechamento **adaptado ao caso concreto**:
- Comece sempre com `#### 📋` (obrigatório para organização interna).
- Depois, use um rótulo contextual específico do tema (evite repetir sempre "Quadro-síntese").
- **Preferência:** use o título original do tópico como base e apenas complemente/especialize quando necessário.
- Exemplo: `#### 📋 Matriz comparativa — Competência tributária municipal`
- Exemplo: `#### 📋 Requisitos essenciais — Improbidade administrativa`

2) Em seguida, gere UMA tabela Markdown (sem placeholders):

| Item (conceito/tema) | Regra/definição (1 frase) | Elementos / requisitos / condições | Base legal / jurisprudência citada | Pegadinha / exemplo / como cai |
| :--- | :--- | :--- | :--- | :--- |

**REGRAS CRÍTICAS (não negocie):**
1. **Sem placeholders:** PROIBIDO usar `"..."`, `"Art. X"`, `"Lei Y"`. Se algo não aparecer no trecho, use `"—"`.
2. **Completude:** 1 linha por item mencionado no bloco (conte mentalmente e confira antes de finalizar).
3. **Concisão:** máximo ~35–45 palavras por célula; frases curtas e diretas.
4. **Compatibilidade:** PROIBIDO usar o caractere `|` dentro de células (isso quebra a tabela). Evite quebras de linha dentro das células.
5. **Sem código:** PROIBIDO blocos de código em células.
6. **Posicionamento:** o quadro vem **APENAS AO FINAL** do bloco concluído (fechamento lógico da seção).
7. **Lastro obrigatório no texto:** cada linha da tabela deve corresponder a conteúdo **já exposto antes** no texto explicativo do mesmo tópico/bloco. **PROIBIDO antecipar** conceito, exceção, fundamento legal ou dica ainda não explicados.

## ⚠️ ORDEM OBRIGATÓRIA: CONTEÚDO PRIMEIRO, TABELA DEPOIS
**NUNCA** gere a tabela antes de terminar TODO o conteúdo explicativo do tópico.
A sequência correta é SEMPRE:
1. TODO o texto explicativo do tópico (parágrafos, exemplos, observações)
2. DEPOIS (e somente depois) o 📋 Quadro-síntese
3. DEPOIS (se aplicável) a 🎯 Tabela de pegadinhas
4. DEPOIS o próximo tópico (## ou ###)

**ERRADO** (tabela no meio do conteúdo):
```
## Tópico X
Explicação inicial...
📋 Quadro-síntese    ← ERRADO!
| ... |
Mais explicação...   ← Deveria estar ANTES da tabela!
```

**CORRETO**:
```
## Tópico X
Explicação inicial...
Mais explicação...   ← TODO conteúdo primeiro
📋 Quadro-síntese    ← Tabela só no final
| ... |
```"""

    PROMPT_TABLE_APOSTILA += """

## 🎯 TABELA 2 (QUANDO APLICÁVEL): COMO A BANCA COBRA / PEGADINHAS
Se (e somente se) o bloco contiver **dicas de prova**, menções a **banca**, **pegadinhas**, “isso cai”, “cuidado”, “tema recorrente” ou exemplos de como a questão aparece:

1) Adicione um subtítulo **adaptado ao caso concreto**:
- Comece sempre com `#### 🎯` (obrigatório para organização interna).
- Depois, use um rótulo contextual de prova/armadilha para o tema.
- **Preferência:** use o título original do tópico/bloco como base e apenas complemente para destacar cobrança, risco ou pegadinha.
- Exemplo: `#### 🎯 Armadilhas de prova — Controle de constitucionalidade`
- Exemplo: `#### 🎯 Como a banca explora o tema — Imunidades tributárias`

2) Gere UMA tabela Markdown:
| Como a banca cobra | Resposta correta (curta) | Erro comum / pegadinha |
| :--- | :--- | :--- |

**REGRAS:**
- Sem placeholders (`...`, `Art. X`, `Lei Y`) → use `—` quando não houver dado no trecho.
- 1 linha por pegadinha/dica/forma de cobrança mencionada.
- Respostas objetivas (1–2 frases curtas por célula).
- PROIBIDO usar `|` dentro de células e evitar quebras de linha dentro das células.
- **Somente com base no já exposto:** não inclua na tabela de pegadinhas conteúdo que não tenha sido explicado antes no mesmo bloco.
- Se não houver material de prova no bloco, **NÃO crie** esta Tabela 2."""

    # --- FIDELIDADE MODE ---
    PROMPT_HEAD_FIDELIDADE = """# DIRETRIZES DE FORMATAÇÃO E REVISÃO (MODO FIDELIDADE)

## PAPEL
VOCÊ É UM EXCELENTÍSSIMO REDATOR TÉCNICO E DIDÁTICO.
- **Tom:** didático, como o professor explicando em aula.
- **Pessoa:** MANTENHA a pessoa original da transcrição (1ª pessoa se for assim na fala).
- **Estilo:** texto corrido, com parágrafos curtos, sem "inventar" doutrina nova.
- **Objetivo:** reproduzir a aula em forma escrita, clara e organizada, mas ainda com a "voz" do professor.

# OBJETIVO
- Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, MANTENDO A FIDELIDADE TOTAL ao conteúdo original.
- **Tamanho:** a saída deve ficar **entre 95% e 115%** do tamanho do trecho de entrada (salvo remoção de muletas e logística).

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias e das falas do professor.
4. **NÃO CRIE MUITOS BULLET POINTS**. PREFIRA UM FORMATO DE MANUAL DIDÁTICO, não checklist.
5. **NÃO USE NEGRITOS EM EXCESSO**. Use apenas para conceitos-chave realmente importantes.

## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados (REsp/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios.
- **Referências**: leis, artigos, jurisprudência, autores, casos citados.
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque).
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico".
- **Encerramento real da aula**: se houver despedida/aviso final/horário no fim do trecho, mantenha (pode organizar como uma seção curta "Encerramento").

## 🎯 PRESERVAÇÃO ESPECIAL: DICAS DE PROVA E EXAMINADORES (CRÍTICO)
Aulas presenciais frequentemente contêm informações valiosas sobre:
1. **Referências a Examinadores**: Nomes de examinadores de concursos, suas preferências, posicionamentos ou temas favoritos. **PRESERVE INTEGRALMENTE**.
2. **Dicas de Prova**: Orientações sobre o que costuma cair em provas, pegadinhas comuns, temas recorrentes.
3. **Estratégias de Estudo**: Sugestões do professor sobre priorização, macetes, formas de memorização.
4. **Casos Práticos e Histórias Reais**: Exemplos de situações reais, casos julgados, histórias ilustrativas. **NUNCA RESUMA**.

> ⚠️ **ESSAS INFORMAÇÕES SÃO O DIFERENCIAL DE UMA AULA AO VIVO.** Sua omissão representa perda irreparável de valor didático."""

    PROMPT_STYLE_FIDELIDADE = """## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Corrija erros gramaticais, regências, ortográficos e de pontuação.
2. **Limpeza Profunda:**
   - **REMOVA** marcadores de oralidade: "né", "tá?", "entende?", "veja bem", "tipo assim".
   - **REMOVA** gírias e vocativos: "meu irmão", "cara", "mano", "galera", "minha gente" (não agregam conteúdo).
     - Se a expressão for PARENTESCO factual ("meu irmão" = irmão do professor), reescreva de forma formal (ex.: "Rodolfo (irmão do professor)").
   - **REMOVA** interações diretas com a turma: "Isso mesmo", "A colega perguntou", "Já estão me vendo?", "Estão ouvindo?".
   - **NÃO REMOVA** o encerramento do professor (ex.: agradecimentos, aviso de horário, "até a próxima", "boa prova") quando estiver no fim do trecho: preserve como um parágrafo final ou uma seção curta "Encerramento".
   - **REMOVA** redundâncias: "subir para cima", "criação nova".
   - **TRANSFORME** perguntas retóricas em afirmações quando possível.
3. **Coesão**: Utilize conectivos para tornar o texto mais fluido. Aplique pontuação adequada.
4. **Legibilidade**:
   - **USE TEXTO CORRIDO NA MEDIDA DO POSSÍVEL.**
   - **PARÁGRAFOS CURTOS**: máximo **3-5 linhas visuais** por parágrafo.
   - **QUEBRE** blocos de texto maciços em parágrafos menores.
   - Seja didático sem perder detalhes e conteúdo.
5. **Linguagem**: Ajuste a linguagem coloquial para português padrão, mantendo o significado original.
6. **Citações**: Use *itálico* para citações curtas e recuo em itálico para citações longas.
7. **Negrito**: Use **negrito** para destacar conceitos-chave (sem exagero).
8. **Formatação Didática** (use com moderação, sem excesso):
   - **Bullet points** para enumerar elementos, requisitos ou características.
   - **Listas numeradas** para etapas, correntes ou exemplos.
   - **Marcadores relacionais** como "→" para consequências lógicas.
9. **Questões e Exercícios**:
   - Se o professor ditar uma questão, exercício ou caso hipotético, **ISOLE-O** em um bloco de citação:
   > **Questão:** O prazo para agravo de petição é de...
   - Separe claramente o enunciado da questão da explicação/gabarito subsequente.
10. **Destaques com Emojis** (use com moderação para facilitar escaneamento visual):
   - 💡 **Dica de Prova** ou **Observação Pedagógica**
   - ⚠️ **Atenção** ou **Cuidado** (pegadinhas, pontos polêmicos)
   - 📌 **Ponto Importante** (conceitos-chave)
   - Exemplo: `> 💡 **Dica de Prova:** Esse tema caiu 3 vezes na PGM-Rio.`"""

    PROMPT_STRUCTURE_FIDELIDADE = """## 📝 ESTRUTURA E TÍTULOS
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis.
- **NÃO crie subtópicos para frases soltas.**
- Use títulos **APENAS** para mudanças reais de assunto.
- Se uma frase parece um título mas não inicia uma nova seção, mantenha como texto normal e use **negrito** se necessário."""

    PROMPT_TABLE_FIDELIDADE = """## 📊 QUADRO-SÍNTESE (CAPTURA COMPLETA)
Ao final de cada **bloco temático relevante**, produza um quadro-síntese didático.

1) Adicione um subtítulo de fechamento **adaptado ao caso concreto**:
- Comece sempre com `#### 📋` (obrigatório para organização interna).
- Depois, use um rótulo contextual específico do tema (evite repetir sempre "Quadro-síntese").
- **Preferência:** use o título original do tópico como base e apenas complemente/especialize quando necessário.
- Exemplo: `#### 📋 Matriz comparativa — Competência tributária municipal`
- Exemplo: `#### 📋 Requisitos essenciais — Improbidade administrativa`

2) Em seguida, gere UMA tabela Markdown (sem placeholders):

| Item (conceito/tema) | Definição/regra (1 frase) | Detalhes (requisitos, exceções, prazos) | Base legal / jurisprudência citada | Dica de prova / ponto polêmico |
| :--- | :--- | :--- | :--- | :--- |

**O QUE DEVE SER CAPTURADO NA TABELA (OBRIGATÓRIO):**
1. **CONCEITOS/INSTITUTOS**: Todo termo técnico definido pelo professor.
2. **FUNDAMENTOS LEGAIS**: Artigos, Leis, Decretos, Súmulas, Enunciados citados.
3. **JURISPRUDÊNCIA**: Julgados do STF, STJ, TCU, tribunais citados (ex: REsp, Acórdão).
4. **PONTOS POLÊMICOS**: Questões controvertidas destacadas pelo professor.
5. **DICAS DE PROVA**: Alertas como "isso cai muito", "cuidado com isso", "a banca gosta de cobrar".
6. **DIVERGÊNCIAS**: Posições doutrinárias conflitantes (ex: "a doutrina majoritária entende X, mas há quem defenda Y").
7. **OBSERVAÇÕES PEDAGÓGICAS**: Destaques do professor sobre importância, frequência em provas, etc.

**REGRAS CRÍTICAS (não negocie):**
1. **Sem placeholders:** PROIBIDO usar `"..."`, `"Art. X"`, `"Lei Y"`. Se algo não aparecer no trecho, use `"—"`.
2. **Completude:** Se o professor mencionou 5 itens no bloco, a tabela DEVE ter 5 linhas (ou mais).
3. **Concisão:** máximo ~35–45 palavras por célula.
4. **Compatibilidade:** PROIBIDO usar o caractere `|` dentro de células. Evite quebras de linha dentro das células.
5. **Sem código:** PROIBIDO blocos de código em células.
6. **Posicionamento:** o quadro vem **APENAS AO FINAL** do bloco concluído, **NUNCA** no meio de explicação.
7. **Lastro obrigatório no texto:** a tabela deve refletir somente itens **já tratados anteriormente** no texto do mesmo bloco temático. **NÃO introduza** informação nova na tabela."""

    PROMPT_TABLE_FIDELIDADE += """

## 🎯 TABELA 2 (QUANDO APLICÁVEL): COMO A BANCA COBRA / PEGADINHAS
Se (e somente se) o bloco contiver **dicas de prova**, menções a **banca**, **pegadinhas**, “isso cai”, “cuidado” ou exemplos de como a questão aparece:

1) Adicione um subtítulo **adaptado ao caso concreto**:
- Comece sempre com `#### 🎯` (obrigatório para organização interna).
- Depois, use um rótulo contextual de prova/armadilha para o tema.
- **Preferência:** use o título original do tópico/bloco como base e apenas complemente para destacar cobrança, risco ou pegadinha.
- Exemplo: `#### 🎯 Armadilhas de prova — Controle de constitucionalidade`
- Exemplo: `#### 🎯 Como a banca explora o tema — Imunidades tributárias`

2) Gere UMA tabela Markdown:
| Como a banca cobra | Resposta correta (curta) | Erro comum / pegadinha |
| :--- | :--- | :--- |

**REGRAS:**
- Sem placeholders (`...`, `Art. X`, `Lei Y`) → use `—` quando não houver dado no trecho.
- 1 linha por pegadinha/dica/forma de cobrança mencionada.
- Respostas objetivas (1–2 frases curtas por célula).
- PROIBIDO usar `|` dentro de células e evitar quebras de linha dentro das células.
- **Somente com base no já exposto:** não inclua na tabela de pegadinhas conteúdo que não tenha sido explicado antes no mesmo bloco.
- Se não houver material de prova no bloco, **NÃO crie** esta Tabela 2."""

    # --- AUDIÊNCIA MODE ---
    PROMPT_HEAD_AUDIENCIA = """# DIRETRIZES DE TRANSCRIÇÃO JURÍDICA (MODO AUDIÊNCIA)

## PAPEL
VOCÊ É UM REDATOR TÉCNICO FORENSE.
- **Tom:** objetivo, fiel e formal.
- **Pessoa:** preserve a pessoa original da fala.
- **Objetivo:** transformar a transcrição em texto legível e coeso, mantendo a fidelidade integral, **SEM RESUMIR**.

## 🎯 OBJETIVO (Fidelidade com clareza)
- **Não resumir:** a saída deve ficar entre **95% e 115%** do tamanho do trecho original (apenas limpeza de oralidade e correções leves).
- **Preservar sequência:** mantenha a ordem cronológica real.
- **Preservar Q&A:** perguntas e respostas devem permanecer em sequência, sem reorganizar.

## ❌ O QUE NÃO FAZER (CRÍTICO)
1. **NÃO RESUMA** nem condense falas ("em síntese", "em resumo", etc.).
2. **NÃO REORGANIZE** por temas; **mantenha cronologia**.
3. **NÃO INVENTE** nomes, cargos, papéis, prazos, datas, valores ou decisões.
4. **NÃO PADRONIZE** vozes diferentes: preserve diferenças entre falas.
5. **NÃO CONVERTA** em narrativa: não transforme depoimentos em “história”.

## ✅ PRESERVE OBRIGATORIAMENTE
- **Identificação de falantes** (SPEAKER 1/2/3, Professor, etc.) quando existir.
- **Timestamps e marcações**: [inaudível], [risos], [interrupção], [sobreposição] e quaisquer timestamps.
- **Números exatos**: datas, valores, artigos/leis, números de processos, prazos, nomes próprios.
- **Negativas e hesitações relevantes** ("não", "talvez", "acho que", "não lembro") quando impactarem sentido.

## 🧷 REGRAS CRÍTICAS DE TRANSCRIÇÃO
1. **NÃO transforme em discurso indireto** (ex.: "o professor disse que…"). Mantenha fala direta.
2. **NÃO transforme em ata resumida**. Preserve a sequência real das falas.
3. **NÃO infira nomes/papéis**: use exatamente os rótulos existentes (ex.: SPEAKER 1/2, Professor).
4. **Uma fala por parágrafo**: não fundir falas de pessoas diferentes no mesmo parágrafo.
5. **Pergunta/Resposta**: mantenha Q&A em blocos consecutivos, sem inserir comentários.
6. **Verbatim decisório**: quando houver trechos explícitos de decisão/encaminhamento (“defiro/indefiro”, “ficou decidido”, “designo”, “intime-se”, etc.), preserve o trecho **literalmente** (pode isolar em citação curta)."""

    PROMPT_STYLE_AUDIENCIA = """## ✅ DIRETRIZES DE ESTILO (sem mudar conteúdo)
1. **Correção leve**: corrija erros gramaticais leves sem alterar o sentido.
2. **Limpeza**: remova muletas orais (“né”, “tá”, “tipo”) quando não forem essenciais.
3. **Pontuação**: ajuste pontuação para legibilidade (sem mudar o que foi dito).
4. **Parágrafos curtos**: 1 fala = 1 parágrafo (quando houver speaker); evite blocos longos.
5. **Não uniformize**: preserve peculiaridades da fala (quando relevantes).
6. **Preservar nomes/dados**: nomes, datas, valores, locais, números de processo, referências jurídicas.
7. **Remova gírias/vocativos**: ex. "meu irmão", "cara", "mano", "galera" (se houver parentesco factual, reescreva de forma formal)."""

    PROMPT_STRUCTURE_AUDIENCIA = """## 📝 ESTRUTURA E TÍTULOS (mínimo necessário)
- **Cronologia**: mantenha a ordem cronológica das falas.
- **Títulos (##/###)**: use apenas quando houver mudança clara de fase:
  - Abertura / Qualificação / Depoimento / Perguntas / Debates / Decisão / Encerramento (exemplos).
- **Q&A**: preserve perguntas e respostas em sequência (sem intercalar resumos).
- **Marcação de falas**: quando houver SPEAKER/participante, mantenha rótulos consistentes."""

    PROMPT_TABLE_AUDIENCIA = """## 📌 QUADROS/TABELAS (somente quando houver conteúdo explícito)
Por padrão, **NÃO** gere quadros-síntese “didáticos” nem tabelas analíticas.

### ✅ EXCEÇÃO (permitida): Registro objetivo de atos/decisões/encaminhamentos
Se (e somente se) existirem trechos **explícitos** de decisão/ato/encaminhamento (ex.: “defiro/indefiro”, “designo”, “intime-se”, “fica consignado”, “prazo de X dias”, “audiência redesignada”, “juntada de documento”):

1) Adicione ao final da fase correspondente:
#### 📌 Registro de atos / decisões / encaminhamentos

2) Gere uma tabela Markdown **curta**:
| Momento (timestamp se houver) | Quem falou/decidiu | Ato/decisão/encaminhamento | Trecho literal (curto) | Prazo/Responsável (se dito) |
| :--- | :--- | :--- | :--- | :--- |

**REGRAS CRÍTICAS:**
- **Sem inferência:** se não houver timestamp/prazo/responsável, use `—`.
- **Trecho literal curto:** copie apenas o mínimo necessário (sem reescrever).
- **Sem `|` dentro das células** e sem quebras de linha nas células.
- Se não houver atos/decisões explícitos, **NÃO crie** esta tabela."""

    # --- REUNIÃO MODE ---
    PROMPT_HEAD_REUNIAO = """# DIRETRIZES DE TRANSCRIÇÃO PROFISSIONAL (MODO REUNIÃO)

## PAPEL
VOCÊ É UM REDATOR DE ATA/REUNIÃO.
- **Tom:** objetivo, formal e direto.
- **Pessoa:** preserve a pessoa original da fala.
- **Objetivo:** registrar a reunião de forma clara e fiel, SEM RESUMIR.

## OBJETIVO
- Transformar a transcrição em texto legível e coeso, mantendo a fidelidade integral.
- **Tamanho:** saída entre **95% e 115%** do trecho original (apenas limpeza de oralidade).

    ## NÃO FAZER
    1. **NÃO RESUMA**. Não omita falas, decisões, encaminhamentos, datas ou valores.
    2. **NÃO ALTERE** a ordem cronológica das falas.
    3. **NÃO INVENTE** informações ausentes.
    4. **NÃO PADRONIZE** falas de participantes diferentes.

    ## REGRAS CRÍTICAS
    1. **NÃO transforme em discurso indireto**. Mantenha fala direta.
    2. **NÃO transforme em ata resumida**. Preserve a sequência real das falas.
    3. **NÃO invente responsáveis, prazos ou decisões**.
    4. **PRESERVE marcações existentes**: [inaudível], [risos], [interrupção] e timestamps.
    5. **NÃO fundir falas de pessoas diferentes no mesmo parágrafo**. Uma fala por parágrafo.
    6. **DESTAQUES VERBATIM (SE EXISTIREM)**: quando houver frases explícitas de decisão/encaminhamento
       (ex.: "ficou definido que..."), você pode isolá-las em bloco de citação ou lista simples,
       copiando o trecho literalmente, sem reescrever e sem reorganizar o conteúdo.

## ✅ PRESERVE OBRIGATORIAMENTE
- **Participantes e identificação** (PARTICIPANTE 1/2/3, nomes, cargos) quando existir.
- **Datas/valores/prazos** e quaisquer números mencionados.
- **Decisões e encaminhamentos explícitos** (não inferir).
- **Ordem cronológica** das falas e a sequência de perguntas/respostas quando houver.
- **Marcações**: [inaudível], [risos], [interrupção] e timestamps."""

    PROMPT_STYLE_REUNIAO = """## ✅ DIRETRIZES DE ESTILO (ata fiel, sem "embelezar")
1. **Correção leve**: corrija erros gramaticais leves sem alterar o sentido.
2. **Limpeza**: remova muletas (“né”, “tá”, “tipo”) quando não forem essenciais.
3. **Pontuação e coesão**: ajuste pontuação para legibilidade, sem mudar conteúdo.
4. **Parágrafos curtos**: uma fala por parágrafo; não fundir participantes diferentes.
5. **Dados críticos**: preserve nomes, cargos, datas, valores, prazos e referências.
6. **Destaques objetivos**: quando houver decisões/encaminhamentos explícitos, destaque-os ao final do tópico com listas/tabela curta, sem inventar.
7. **Remova gírias/vocativos**: ex. "meu irmão", "cara", "mano", "galera" (se houver parentesco factual, reescreva de forma formal)."""

    PROMPT_STRUCTURE_REUNIAO = """## 📝 ESTRUTURA E TÍTULOS (orientado a pauta)
- **Cronologia**: mantenha a ordem cronológica das falas.
- **Títulos (##/###)**: use apenas quando houver mudança clara de pauta/tema.
- **Blocos operacionais**: se houver abertura/encerramento/decisões/encaminhamentos explícitos, você pode criar subtítulos correspondentes.
- **Q&A**: preserve perguntas e respostas em sequência quando houver.
- **Não reorganize** por “assuntos” se a reunião foi caótica: preserve a sequência real."""

    PROMPT_TABLE_REUNIAO = """## 📌 QUADROS/TABELAS (somente quando houver decisões/encaminhamentos explícitos)
Por padrão, **NÃO** gere quadros-síntese “didáticos”.

### ✅ EXCEÇÃO (permitida): Decisões e encaminhamentos (quando explícitos)
Se (e somente se) existirem falas explícitas de decisão/ação (ex.: “ficou definido que…”, “fulano vai…”, “prazo até…”, “enviar documento”, “marcar reunião”):

1) Adicione ao final da pauta/tema correspondente:
#### ✅ Decisões e encaminhamentos

2) Gere uma tabela Markdown curta:
| Item | Decisão/ação (literal curto) | Responsável (se dito) | Prazo (se dito) | Observações |
| :--- | :--- | :--- | :--- | :--- |

**REGRAS CRÍTICAS:**
- **Sem inferência:** se não houver responsável/prazo, use `—`.
- **Literal curto:** não reescreva para “melhorar”; copie o essencial sem inventar.
- **Sem `|` dentro das células** e sem quebras de linha nas células.
- Se não houver decisões/encaminhamentos explícitos, **NÃO crie** esta tabela."""

    # --- DEPOIMENTO MODE ---
    PROMPT_HEAD_DEPOIMENTO = """# DIRETRIZES DE TRANSCRIÇÃO JURÍDICA (MODO DEPOIMENTO)

## PAPEL
VOCÊ É UM REDATOR TÉCNICO FORENSE.
- **Objetivo:** registrar depoimentos com fidelidade total.
- **Tom:** objetivo, sem resumos ou interpretações.

## REGRAS CRÍTICAS
1. **NÃO RESUMA**. Preserve conteúdo integral.
2. **NÃO transforme em discurso indireto**. Mantenha fala direta.
3. **NÃO transforme em ata resumida**. Preserve a sequência real das falas.
4. **PRESERVE** pausas, negações e afirmações relevantes.
5. **MANTENHA** perguntas e respostas em sequência."""

    PROMPT_STYLE_DEPOIMENTO = """## ✅ DIRETRIZES DE ESTILO
1. Corrija apenas erros gramaticais leves.
2. Preserve nomes, datas, valores e qualificações.
3. Se houver identificação de falante, mantenha-a.
4. Remova gírias/vocativos (ex.: "meu irmão", "cara", "mano", "galera") quando não agregarem conteúdo; se houver parentesco factual, reescreva de forma formal."""

    PROMPT_STRUCTURE_DEPOIMENTO = """## 📝 ESTRUTURA
- Mantenha a sequência das falas.
- Use títulos apenas se houver blocos claros (ex.: Depoimento, Esclarecimentos)."""

    PROMPT_TABLE_DEPOIMENTO = """## 📌 OBSERVAÇÃO SOBRE TABELAS
Não gere quadros-síntese automaticamente."""

    # --- SIMULADO / CORREÇÃO DE PROVA ADDON ---
    # Injetado dinamicamente quando o mapeamento detecta tipo SIMULADO ou CORREÇÃO.
    PROMPT_SIMULADO_ADDON = """
## 📝 REGRAS ESPECIAIS: CORREÇÃO DE QUESTÕES / SIMULADO

Este material contém **correção de questões** ou **simulado**. Aplique as regras abaixo ALÉM das regras gerais:

### ESTRUTURA POR QUESTÃO (OBRIGATÓRIO)
Cada questão deve seguir a estrutura:
```
## N. Questão X: [Título descritivo] — [Área do Direito]

### N.1. Enunciado
> [Texto integral da questão em blockquote]

### N.2. Fundamentação / Análise
[Explicação completa do professor: doutrina, jurisprudência, artigos citados]

### N.3. Resposta / Gabarito
[Resposta esperada, espelho de correção, pontuação se mencionada]
```

### REGRAS CRÍTICAS:
1. **PRESERVE O ENUNCIADO INTEGRAL** da questão em blockquote (`>`). NUNCA resuma o enunciado.
2. **SEPARE CLARAMENTE** enunciado, fundamentação e resposta — mesmo que o professor misture na explicação oral.
3. **PRESERVE TODAS as alternativas** (A, B, C, D, E) quando existirem, indicando a correta.
4. **Pontuação e critérios**: Se o professor mencionar pontuação, peso ou critérios de correção, capture em uma linha destacada:
   > 📌 **Pontuação:** X pontos | **Critérios:** ...
5. **Espelho de Correção**: Se o professor detalhar o espelho, formate como lista numerada com os pontos esperados.
6. **Gabarito de Múltipla Escolha**: Se for questão objetiva, destaque:
   > ✅ **Gabarito:** Alternativa **C** — [justificativa curta]
7. **Referências cruzadas**: Se o professor comparar com questões anteriores ou de outras bancas, preserve a referência.
8. **NÃO FUNDA questões diferentes** em uma única seção — cada questão é um bloco ## independente.

### TABELA DE GABARITO (ao final do documento)
Se houver múltiplas questões, gere uma tabela consolidada ao final:

#### 📋 Gabarito Consolidado
| Questão | Área do Direito | Gabarito / Resposta-chave | Fundamento principal |
| :--- | :--- | :--- | :--- |
"""

    # --- SHARED FOOTER (Anti-Duplication) ---
    PROMPT_FOOTER = """## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA.
- **NUNCA formate novamente esse contexto.**
- **NUNCA inclua esse contexto na sua resposta.**
- **NUNCA repita informações que já estão no contexto.**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>.
- **CRÍTICO:** Se o texto começar repetindo a última frase do contexto, **IGNORE A REPETIÇÃO.**"""


    PROMPT_MAPEAMENTO = """Você é um especialista em organização de conteúdo educacional acadêmico.

## ETAPA 1: IDENTIFICAR O TIPO DE CONTEÚDO
Analise a transcrição e determine qual é a **natureza predominante** do material:

| Tipo | Pistas no Texto | Estrutura Ideal |
|------|-----------------|-----------------| 
| **SIMULADO** | "questão 1", "questão 2", "espelho de correção", "correção do simulado", "vamos corrigir" | Organizar por QUESTÕES numeradas |
| **AULA EXPOSITIVA** | explicações contínuas de um tema, teoria, doutrina, sem questões específicas | Organizar por TEMAS/MATÉRIAS |
| **REVISÃO** | "revisão", "resumo", múltiplos temas curtos, "pontos importantes" | Organizar por TÓPICOS de revisão |
| **CORREÇÃO DE PROVA** | "gabarito", "alternativa correta", "item certo/errado" | Organizar por QUESTÕES com gabarito |

## ETAPA 2: EXTRAIR A ESTRUTURA

### Se for SIMULADO ou CORREÇÃO DE PROVA:
```
1. Orientações Gerais / Introdução
2. Questão 1: [Título descritivo] — [Área do Direito]
   2.1. Enunciado e Contexto
   2.2. Fundamentação (Doutrina/Jurisprudência)
   2.3. Pontos do Espelho / Resposta
3. Questão 2: [Título descritivo] — [Área do Direito]
   3.1. ...
[Continue para cada questão]
N. Considerações Finais / Dúvidas
```

### Se for AULA EXPOSITIVA:
```
1. Introdução
2. [Matéria 1: ex. Direito Administrativo]
   2.1. [Subtema]
      2.1.1. [Detalhamento]
3. [Matéria 2: ex. Direito Civil]
   3.1. ...
```

### Se for REVISÃO:
```
1. [Tema 1]
   1.1. Pontos-chave
   1.2. Jurisprudência/Súmulas
2. [Tema 2]
   2.1. ...
```

## REGRAS GERAIS:
1. **MÁXIMO 3 NÍVEIS** de hierarquia (1., 1.1., 1.1.1.)
2. **Seja descritivo** nos títulos — inclua o assunto real, não apenas "Questão 1"
3. **Mantenha a ORDEM** cronológica da transcrição
4. **Mapeie do INÍCIO ao FIM** — não omita partes
5. **Identifique a ÁREA DO DIREITO** de cada bloco quando possível
6. **PREFIRA SUBTÓPICOS (1.1.) a novos tópicos (2.)**: Abra novo tópico de nível 1 SOMENTE quando o macroassunto mudar de verdade (ex.: de Direito Administrativo para Direito Civil). Aspectos, institutos e marcos legais DENTRO do mesmo macroassunto devem ser subtópicos (1.1., 1.2., etc.), NUNCA tópicos de nível 1 separados.
7. **ANTI-FRAGMENTAÇÃO**: Se o professor trata 4+ aspectos de um tema, todos devem ser subtópicos de um único tema-mãe. Exemplo correto: `2. Execução Fiscal` com `2.1. Procedimento`, `2.2. Citação`, `2.3. Exceção de Pré-Executividade`. Exemplo ERRADO: `2. Execução Fiscal`, `3. Procedimento`, `4. Citação`.
8. **TÍTULOS SÃO RÓTULOS, NÃO FALAS**: Os títulos devem ser rótulos descritivos curtos (máx 8 palavras), NUNCA trechos literais da fala do professor. Exemplo ERRADO: "1. Já estávamos conversando aqui antes de começar". Exemplo CORRETO: "1. Introdução e Apresentação".
9. **SAUDAÇÕES E LOGÍSTICA → "Introdução"**: Boas-vindas, apresentação pessoal, ajustes técnicos → agrupar sob "Introdução" ou "Apresentação e Contextualização", nunca com a fala literal como título.

## 🏛️ REGRA ESPECIAL: MARCOS LEGAIS (v2.17)
Quando identificar marcos legais importantes, crie subtópicos específicos:
- **Súmulas** (STF, STJ, Vinculantes): Criar subtópico "X.Y. Súmula [Número] do [Tribunal]."
- **Teses (Repercussão Geral/Repetitivos)**: Criar subtópico "X.Y. Tese/Tema [Número] do STJ/STF."
- **Artigos de Lei Central**: Se um artigo é explicado em profundidade, criar subtópico "X.Y. Art. [Número] da [Lei]."

Exemplo:
```
2. Execução Fiscal
   2.1. Procedimento da LEF (Lei 6.830/80)
   2.2. Súmula 314 do STJ (Citação por Hora Certa)
   2.3. Tema 444 do STJ (Redirecionamento)
```

## 📍 ÂNCORAS VERBATIM (v2.25 — NOVO)
Para CADA tópico de nível 1 e 2, adicione ao final da linha duas âncoras:
- **ABRE:** Frase LITERAL (10-20 palavras) que o professor falou ao INICIAR o tópico.
- **FECHA:** Frase LITERAL que marca a TRANSIÇÃO para o próximo tópico (ou "FIM" se for o último).

Formato: `NÚMERO. Título | ABRE: "frase literal" | FECHA: "frase literal"`

Exemplo:
```
1. Introdução | ABRE: "bom dia pessoal vamos começar a aula de hoje" | FECHA: "então vamos entrar agora no tema principal"
2. Credenciamento | ABRE: "então vamos entrar agora no tema principal que é o credenciamento" | FECHA: "passemos agora para a pré-qualificação"
   2.1. Conceito e Natureza Jurídica | ABRE: "o credenciamento é uma modalidade" | FECHA: "agora vamos ver as hipóteses"
3. Pré-qualificação | ABRE: "passemos agora para a pré-qualificação" | FECHA: "FIM"
```

**IMPORTANTE:**
- Use as palavras EXATAS da transcrição (podem ter erros de fala, ok).
- Se não encontrar frase clara de abertura, use as primeiras 10 palavras do trecho.
- A âncora FECHA de um tópico deve ser igual (ou muito similar) à âncora ABRE do próximo.
- **NÃO use timestamps, labels de falante ou marcadores de formatação** (ex.: `[00:10]`, `**SPEAKER**:` ou `##`).
- Prefira **frases contínuas** do conteúdo falado (8–16 palavras), sem quebras de linha.

## TRANSCRIÇÃO:
{transcricao}

## RESPOSTA:
Primeiro, indique em uma linha: `[TIPO: SIMULADO/EXPOSITIVA/REVISÃO/CORREÇÃO]`
Depois, retorne APENAS a estrutura hierárquica (máx 3 níveis), COM as âncoras ABRE/FECHA para cada item de nível 1 e 2."""


    PROMPT_FIDELIDADE = """# DIRETRIZES DE FORMATAÇÃO E REVISÃO (MODO FIDELIDADE)

## PAPEL
VOCÊ É UM EXCELENTÍSSIMO REDATOR TÉCNICO E DIDÁTICO.
- **Tom:** didático, como o professor explicando em aula.
- **Pessoa:** MANTENHA a pessoa original da transcrição (1ª pessoa se for assim na fala).
- **Estilo:** texto corrido, com parágrafos curtos, sem "inventar" doutrina nova.
- **Objetivo:** reproduzir a aula em forma escrita, clara e organizada, mas ainda com a "voz" do professor.

# OBJETIVO
- Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, MANTENDO A FIDELIDADE TOTAL ao conteúdo original.
- **Tamanho:** a saída deve ficar **entre 95% e 115%** do tamanho do trecho de entrada (salvo remoção de muletas e logística).

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias e das falas do professor.


## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados (REsp/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios.
- **Referências**: leis, artigos, jurisprudência, autores, casos citados.
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque).
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico".

## 🎯 PRESERVAÇÃO ESPECIAL: DICAS DE PROVA E EXAMINADORES (CRÍTICO)
Aulas presenciais frequentemente contêm informações valiosas sobre:
1. **Referências a Examinadores**: Nomes de examinadores de concursos, suas preferências, posicionamentos ou temas favoritos. **PRESERVE INTEGRALMENTE**.
   - Exemplo: "O examinador Fulano costuma cobrar..." → MANTER
   - Exemplo: "Esse tema foi cobrado pelo professor X na prova..." → MANTER
2. **Dicas de Prova**: Orientações sobre o que costuma cair em provas, pegadinhas comuns, temas recorrentes.
   - Exemplo: "Isso cai muito em prova..." → MANTER
   - Exemplo: "Atenção: essa é uma pegadinha clássica..." → MANTER
3. **Estratégias de Estudo**: Sugestões do professor sobre priorização, macetes, formas de memorização.
   - Exemplo: "Gravem isso: na dúvida, marquem..." → MANTER
   - Exemplo: "Para PGM, foquem em..." → MANTER
4. **Casos Práticos e Histórias Reais**: Exemplos de situações reais, casos julgados, histórias ilustrativas.
   - **NUNCA RESUMA** histórias ou exemplos práticos. Preserve na íntegra.

> ⚠️ **ESSAS INFORMAÇÕES SÃO O DIFERENCIAL DE UMA AULA AO VIVO.** Sua omissão representa perda irreparável de valor didático.


## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Corrija erros gramaticais, regências, ortográficos e de pontuação.
2. **Limpeza Profunda:**
   - **REMOVA** marcadores de oralidade: "né", "tá?", "entende?", "veja bem", "tipo assim".
   - **REMOVA** gírias e vocativos: "meu irmão", "cara", "mano", "galera", "minha gente" (se houver parentesco factual, reescreva de forma formal).
   - **REMOVA** interações diretas com a turma: "Isso mesmo", "A colega perguntou", "Já estão me vendo?", "Estão ouvindo?".
   - **REMOVA** redundâncias: "subir para cima", "criação nova".
   - **TRANSFORME** perguntas retóricas em afirmações quando possível.
3. **Coesão**: Utilize conectivos para tornar o texto mais fluido. Aplique pontuação adequada.
4. **Legibilidade**:
   - **PARÁGRAFOS CURTOS**: máximo **3-5 linhas visuais** por parágrafo.
   - **QUEBRE** blocos de texto maciços em parágrafos menores.
   - Seja didático sem perder detalhes e conteúdo.
5. **Formatação Didática** (use com moderação):
   - **Bullet points** para enumerar elementos, requisitos ou características.
   - **Listas numeradas** para etapas, correntes ou exemplos.
   - **Marcadores relacionais** como "→" para consequências lógicas.

## 📝 ESTRUTURA E TÍTULOS
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis.
- **NÃO crie subtópicos para frases soltas.**
- Use títulos **APENAS** para mudanças reais de assunto.

## 📊 QUADRO-SÍNTESE (CAPTURA COMPLETA)
Ao final de cada **bloco temático relevante**, produza um quadro-síntese didático.

1) Adicione um subtítulo de fechamento **adaptado ao caso concreto**:
- Comece sempre com `#### 📋` (obrigatório para organização interna).
- Depois, use um rótulo contextual específico do tema (evite repetir sempre "Quadro-síntese").
- **Preferência:** use o título original do tópico como base e apenas complemente/especialize quando necessário.
- Exemplo: `#### 📋 Matriz comparativa — Competência tributária municipal`
- Exemplo: `#### 📋 Requisitos essenciais — Improbidade administrativa`

2) Em seguida, gere UMA tabela Markdown (sem placeholders):

| Item (conceito/tema) | Definição/regra (1 frase) | Detalhes (requisitos, exceções, prazos) | Base legal / jurisprudência citada | Dica de prova / ponto polêmico |
| :--- | :--- | :--- | :--- | :--- |

**REGRAS CRÍTICAS (não negocie):**
1. **Sem placeholders:** PROIBIDO usar `"..."`, `"Art. X"`, `"Lei Y"`. Se algo não aparecer no trecho, use `"—"`.
2. **Completude:** Se o professor mencionou 5 itens no bloco, a tabela DEVE ter 5 linhas (ou mais).
3. **Concisão:** máximo ~35–45 palavras por célula.
4. **Compatibilidade:** PROIBIDO usar o caractere `|` dentro de células. Evite quebras de linha dentro das células.
5. **Sem código:** PROIBIDO blocos de código em células.
6. **Posicionamento:** o quadro vem **APENAS AO FINAL** do bloco concluído, **NUNCA** no meio de explicação.

## 🎯 TABELA 2 (QUANDO APLICÁVEL): COMO A BANCA COBRA / PEGADINHAS
Se (e somente se) o bloco contiver **dicas de prova**, menções a **banca**, **pegadinhas**, “isso cai”, “cuidado”, “tema recorrente” ou exemplos de como a questão aparece:

1) Adicione um subtítulo **adaptado ao caso concreto**:
- Comece sempre com `#### 🎯` (obrigatório para organização interna).
- Depois, use um rótulo contextual de prova/armadilha para o tema.
- **Preferência:** use o título original do tópico/bloco como base e apenas complemente para destacar cobrança, risco ou pegadinha.
- Exemplo: `#### 🎯 Armadilhas de prova — Controle de constitucionalidade`
- Exemplo: `#### 🎯 Como a banca explora o tema — Imunidades tributárias`

2) Gere UMA tabela Markdown:
| Como a banca cobra | Resposta correta (curta) | Erro comum / pegadinha |
| :--- | :--- | :--- |

**REGRAS:**
- Sem placeholders (`...`, `Art. X`, `Lei Y`) → use `—` quando não houver dado no trecho.
- 1 linha por pegadinha/dica/forma de cobrança mencionada.
- Respostas objetivas (1–2 frases curtas por célula).
- PROIBIDO usar `|` dentro de células e evitar quebras de linha dentro das células.
- Se não houver material de prova no bloco, **NÃO crie** esta Tabela 2.

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA.
- **NUNCA formate novamente esse contexto.**
- **NUNCA inclua esse contexto na sua resposta.**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>.
    """

    PROMPT_APOSTILA_ACTIVE = """# DIRETRIZES DE REDAÇÃO: MANUAL JURÍDICO DIDÁTICO (MODO APOSTILA)

## PAPEL
VOCÊ É UM EXCELENTÍSSIMO REDATOR JURÍDICO E DIDÁTICO.
- **Tom:** doutrinário, impessoal, estilo manual de Direito.
- **Pessoa:** 3ª pessoa ou construções impessoais ("O professor explica...", "A doutrina define...").
- **Estilo:** prosa densa, porém com parágrafos curtos e didáticos.
- **Objetivo:** transformar a aula em texto de apostila/manual, sem alterar conteúdo nem inventar informações.

## 💎 PILAR 1: ESTILO (VOZ ATIVA E DIRETA)
> 🚫 **PROIBIDO VOZ PASSIVA EXCESSIVA:** "Anunciou-se", "Informou-se".
> ✅ **PREFIRA VOZ ATIVA:** "O professor explica...", "A doutrina define...", "O Art. 37 estabelece...".

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias.
4. **NÃO CRIE PARÁGRAFOS LONGOS**. Máximo 3-5 linhas visuais por parágrafo.

## ❌ PRESERVE OBRIGATORIAMENTE
- **IDENTIFICAÇÃO DE FALANTES**: Se houver SPEAKER A/B/C ou similar, identifique o professor pelo contexto (quando ele se apresentar: "Eu sou o professor João", "Meu nome é Maria"). Substitua "SPEAKER X" pelo nome identificado. Se não identificar, use "Professor" ou "Palestrante".
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados, Temas de Repercussão Geral, Recursos Repetitivos. **NUNCA OMITA NÚMEROS DE TEMAS OU SÚMULAS**.
- **JURISPRUDÊNCIA**: Se o texto citar "Tema 424", "RE 123", "ADI 555", **MANTENHA O NÚMERO**. Não generalize para "jurisprudência do STJ".
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios.
- **Referências**: leis, artigos, jurisprudência (STF/STJ), autores, casos citados.
- **Ênfases intencionais** e **Observações pedagógicas**.


## 🎯 PRESERVAÇÃO ESPECIAL: DICAS DE PROVA E EXAMINADORES (CRÍTICO)
Aulas presenciais frequentemente contêm informações valiosas sobre:
1. **Referências a Examinadores**: Nomes de examinadores de concursos, suas preferências, posicionamentos ou temas favoritos. **PRESERVE INTEGRALMENTE**.
   - Exemplo: "O examinador Fulano costuma cobrar..." → MANTER
   - Exemplo: "Esse tema foi cobrado pelo professor X na prova..." → MANTER
2. **Dicas de Prova**: Orientações sobre o que costuma cair em provas, pegadinhas comuns, temas recorrentes.
   - Exemplo: "Isso cai muito em prova..." → MANTER
   - Exemplo: "Atenção: essa é uma pegadinha clássica..." → MANTER
3. **Estratégias de Estudo**: Sugestões do professor sobre priorização, macetes, formas de memorização.
   - Exemplo: "Gravem isso: na dúvida, marquem..." → MANTER
   - Exemplo: "Para PGM, foquem em..." → MANTER
4. **Casos Práticos e Histórias Reais**: Exemplos de situações reais, casos julgados, histórias ilustrativas.
   - **NUNCA RESUMA** histórias ou exemplos práticos. Preserve na íntegra.

> ⚠️ **ESSAS INFORMAÇÕES SÃO O DIFERENCIAL DE UMA AULA AO VIVO.** Sua omissão representa perda irreparável de valor didático.


## ✅ DIRETRIZES DE ESTILO E FORMATAÇÃO VISUAL
1. **Correção Gramatical**: Ajuste a linguagem coloquial para o padrão culto.
2. **Limpeza**: Remova gírias, vocativos e cacoetes ("né", "tipo assim", "então", "meu irmão", "cara", "mano", "galera") e vícios de oralidade. Se houver parentesco factual, reescreva de forma formal.
3. **Coesão**: Use conectivos e pontuação adequada para tornar o texto fluido.
4. **Legibilidade Visual** (OBRIGATÓRIO):
   - **PARÁGRAFOS CURTOS**: máximo **3-5 linhas visuais** por parágrafo. **QUEBRE SEMPRE.**
   - **RECUOS COM MARCADORES**: Use `>` para citações, destaques ou observações importantes.
   - **NEGRITO MODERADO**: Destaque conceitos-chave com **negrito**, mas sem exagero.
   - **ITÁLICO**: Use para termos em latim, expressões estrangeiras ou ênfase leve.
5. **Formatação Didática** (use generosamente para legibilidade):
   - **Bullet points** (`-` ou `*`) para enumerar elementos, requisitos ou características.
   - **Listas numeradas** (`1.`, `2.`) para etapas, correntes doutrinárias ou exemplos ordenados.
   - **Marcadores relacionais** como `→` para consequências lógicas.
   - **Subseções** (###, ####) para organizar subtópicos dentro de um mesmo tema.

## 🎨 FORMATAÇÃO VISUAL AVANÇADA
Para garantir legibilidade superior:
1. **Após cada conceito importante**, quebre o parágrafo e inicie outro.
2. **Use listas** sempre que houver enumeração de mais de 2 itens.
3. **Use citações recuadas** (`>`) para destacar:
   - Teses jurídicas
   - Pontos polêmicos
   - Observações práticas
   - Dicas de prova
4. **Separe visualmente** diferentes aspectos de um mesmo tema com subseções.

## 📝 ESTRUTURA E TÍTULOS
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###, ####) para organizar os tópicos.
- **NÃO crie subtópicos para frases soltas.**
- Use títulos **APENAS** para mudanças reais de assunto.
- **CRIE SUBSEÇÕES** (###) quando o professor abordar aspectos diferentes de um mesmo tema.

## 📊 QUADRO-SÍNTESE (OBRIGATÓRIO)
Ao final de CADA tópico principal (## ou ###), faça um fechamento didático com UM quadro-síntese.
SEMPRE que houver diferenciação de conceitos, prazos, procedimentos, requisitos ou regras, o quadro é OBRIGATÓRIO.

1) Adicione um subtítulo de fechamento **adaptado ao caso concreto**:
- Comece sempre com `#### 📋` (obrigatório para organização interna).
- Depois, use um rótulo contextual específico do tema (evite repetir sempre "Quadro-síntese").
- **Preferência:** use o título original do tópico como base e apenas complemente/especialize quando necessário.
- Exemplo: `#### 📋 Matriz comparativa — Competência tributária municipal`
- Exemplo: `#### 📋 Requisitos essenciais — Improbidade administrativa`

2) Em seguida, gere UMA tabela Markdown (sem placeholders):

| Item (conceito/tema) | Regra/definição (1 frase) | Elementos / requisitos / condições | Base legal / jurisprudência citada | Pegadinha / exemplo / como cai |
| :--- | :--- | :--- | :--- | :--- |

**REGRAS CRÍTICAS (não negocie):**
1. **Sem placeholders:** PROIBIDO usar `"..."`, `"Art. X"`, `"Lei Y"`. Se algo não aparecer no trecho, use `"—"`.
2. **Completude:** 1 linha por item mencionado no bloco (conte mentalmente e confira antes de finalizar).
3. **Concisão:** máximo ~35–45 palavras por célula; frases curtas e diretas.
4. **Compatibilidade:** PROIBIDO usar o caractere `|` dentro de células (isso quebra a tabela). Evite quebras de linha dentro das células.
5. **Sem código:** PROIBIDO blocos de código em células.
6. **Posicionamento:** o quadro vem **APENAS AO FINAL** do bloco concluído (fechamento lógico da seção).

## 🎯 TABELA 2 (QUANDO APLICÁVEL): COMO A BANCA COBRA / PEGADINHAS
Se (e somente se) o bloco contiver **dicas de prova**, menções a **banca**, **pegadinhas**, “isso cai”, “cuidado”, “tema recorrente” ou exemplos de como a questão aparece:

1) Adicione um subtítulo **adaptado ao caso concreto**:
- Comece sempre com `#### 🎯` (obrigatório para organização interna).
- Depois, use um rótulo contextual de prova/armadilha para o tema.
- **Preferência:** use o título original do tópico/bloco como base e apenas complemente para destacar cobrança, risco ou pegadinha.
- Exemplo: `#### 🎯 Armadilhas de prova — Controle de constitucionalidade`
- Exemplo: `#### 🎯 Como a banca explora o tema — Imunidades tributárias`

2) Gere UMA tabela Markdown:
| Como a banca cobra | Resposta correta (curta) | Erro comum / pegadinha |
| :--- | :--- | :--- |

**REGRAS:**
- Sem placeholders (`...`, `Art. X`, `Lei Y`) → use `—` quando não houver dado no trecho.
- 1 linha por pegadinha/dica/forma de cobrança mencionada.
- Respostas objetivas (1–2 frases curtas por célula).
- PROIBIDO usar `|` dentro de células e evitar quebras de linha dentro das células.
- Se não houver material de prova no bloco, **NÃO crie** esta Tabela 2.

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA.
- **NUNCA formate novamente esse contexto.**
- **NUNCA inclua esse contexto na sua resposta.**
- **NUNCA repita informações que já estão no contexto.**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>.
- **CRÍTICO:** Se o texto começar repetindo a última frase do contexto, **IGNORE A REPETIÇÃO.**
"""

    SYSTEM_PROMPT_FORMAT = PROMPT_APOSTILA_ACTIVE  # Default

    def __init__(self, model_size="large-v3-turbo", provider="gemini"):
        """
        MLX-Whisper otimizado para Apple Silicon (M3 Pro)
        Formatação otimizada para Gemini 3 Flash / OpenAI GPT-5 Mini
        """
        print(f"{Fore.CYAN}🚀 Inicializando MLX-Whisper ({model_size}) para Apple Silicon...")
        self.model_name = model_size
        self.provider = provider.lower()
        self.thinking_level = "medium"
        self.use_openai_primary = False
        self._diarization_enabled = False
        self._diarization_required = False

        # v2.30: Override para condition_on_previous_text via env var
        # Em áudios de baixa qualidade, False pode evitar propagação de alucinações
        _cpt_env = _env_truthy("VOMO_CONDITION_PREVIOUS", default=None)
        self._condition_on_previous = _cpt_env if _cpt_env is not None else True

        # Provider Configuration
        if self.provider == "openai":
            print(f"{Fore.GREEN}🧠 Provider: OpenAI (GPT-5 Mini)")
            self.llm_model = "gpt-5-mini-2025-08-07" # Modelo principal
            self.client = OpenAI() # Assumes OPENAI_API_KEY env var
        else:
            print(f"{Fore.BLUE}✨ Provider: Google Gemini")
            self.llm_model = "gemini-3-flash-preview"
            self.client = None
            self._gemini_use_vertex = False
            self._gemini_vertex_project = None
            self._gemini_vertex_location = None
        
        # Carrega variáveis de ambiente (sem sobrescrever env já exportadas pelo caller,
        # ex: uvicorn/serviço). Para forçar override, exporte antes no shell.
        from dotenv import load_dotenv
        load_dotenv(override=False)
        
        # Sync global metrics with provider for cost calculation
        metrics.set_provider(self.provider)
        
        project_id_env = os.getenv("GOOGLE_CLOUD_PROJECT")
        project_id = project_id_env or "gen-lang-client-0727883752"

        # Configuração de Credenciais (Explicit Fallback)
        CREDENTIALS_PATH = "/Users/nicholasjacob/Documents/Aplicativos/Iudex/gen-lang-client-0727883752-f72a632e4ec2.json"
        if os.path.exists(CREDENTIALS_PATH) and not os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDENTIALS_PATH
            print(f"{Fore.CYAN}🔑 Credenciais Vertex carregadas de: {CREDENTIALS_PATH}")
        
        # Estratégia Estrita de Autenticação (Vertex AI Only) - SKIP if OpenAI provider
        if self.provider != "openai":
            api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
            auth_mode = (os.getenv("IUDEX_GEMINI_AUTH") or "auto").strip().lower()
            if auth_mode in ("apikey", "api_key", "key", "dev", "developer", "ai-studio", "aistudio"):
                use_vertex = False
            elif auth_mode in ("vertex", "vertexai", "gcp"):
                use_vertex = True
            else:
                # Auto: prefer Vertex when a project or application creds are available.
                has_vertex_creds = bool(project_id_env) or bool(os.getenv("GOOGLE_APPLICATION_CREDENTIALS"))
                use_vertex = has_vertex_creds or not bool(api_key)

            if use_vertex:
                location = (os.getenv("VERTEX_AI_LOCATION") or "us-central1").strip()
                print(f"{Fore.YELLOW}DEBUG: VERTEX_AI_LOCATION: {location}")
                print(f"{Fore.CYAN}☁️  Conectando via Vertex AI ({project_id})...")
                self._gemini_use_vertex = True
                self._gemini_vertex_project = project_id
                self._gemini_vertex_location = location

                # Use user's preferred auth style if API key is present but Vertex is requested
                if api_key and os.getenv("GOOGLE_APPLICATION_CREDENTIALS") is None:
                    self.client = genai.Client(
                        vertexai=True,
                        api_key=api_key,
                        # project and location might be needed depending on the key type
                    )
                else:
                    self.client = genai.Client(
                        vertexai=True,
                        project=project_id,
                        location=location,
                    )
            else:
                if not api_key:
                    raise RuntimeError("GOOGLE_API_KEY (ou GEMINI_API_KEY) não configurada.")
                print(f"{Fore.CYAN}🔑 Conectando via Google AI Studio (API key)...")
                self._gemini_use_vertex = False
                self._gemini_vertex_project = None
                self._gemini_vertex_location = None
                self.client = genai.Client(api_key=api_key)

            # Teste rápido (best-effort)
            try:
                self.client.models.count_tokens(model=self.llm_model, contents="teste")
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️  Aviso: Falha no teste inicial do modelo {self.llm_model}: {e}")
                # Não explode agora, tenta usar depois.
            print(f"{Fore.GREEN}   ✅ Gemini conectado com sucesso.")

        # Inicializa OpenAI como Fallback Terciário
        self.openai_model = "gpt-5-mini-2025-08-07"
        openai_key = os.getenv("OPENAI_API_KEY")  # Define variable explicitly
        if openai_key:
            self.openai_client = AsyncOpenAI(api_key=openai_key)
            print(f"{Fore.CYAN}🤖 Fallback OpenAI ({self.openai_model}) configurado.")
        else:
            self.openai_client = None
            print(f"{Fore.YELLOW}⚠️  OpenAI Key não encontrada. Fallback GPT desativado.")

        # Inicializa seletor de prompt (Default: Apostila)
        self.prompt_apostila = self.PROMPT_APOSTILA_ACTIVE

        # Gemini não tem client async nativo igual OpenAI, usamos chamadas sync em threads
        self.async_client = None

        # Cache directory
        self.cache_dir = Path(".cache_vomo")
        self.cache_dir.mkdir(exist_ok=True)

    def _resolve_thinking_level(self):
        level = (self.thinking_level or "medium").lower()
        if level == "high":
            return "HIGH"
        if level == "low":
            return "LOW"
        return "MEDIUM"

    def _build_system_prompt(
        self,
        mode: str = "APOSTILA",
        custom_style_override: str = None,
        custom_prompt_scope: str = "tables_only",
        disable_tables: bool = False,
        allow_indirect: bool = False,
        allow_summary: bool = False,
        include_timestamps: bool = True,
    ) -> str:
        """
        v2.22: Composes the system prompt from modular components.
        
        Allows users to provide a custom_style_override that replaces the 
        default STYLE and TABLE components, while preserving HEAD, STRUCTURE, and FOOTER.
        
        Args:
            mode: "APOSTILA", "FIDELIDADE", "AUDIENCIA", "REUNIAO" ou "DEPOIMENTO"
            custom_style_override: Optional custom prompt for style/table layers.
                                   If provided, replaces STYLE+TABLE components.
            allow_indirect: Se True, permite discurso indireto em modos AUDIENCIA/REUNIAO/DEPOIMENTO.
            allow_summary: Se True, permite ata resumida em modos AUDIENCIA/REUNIAO/DEPOIMENTO.
        
        Returns:
            Complete system prompt string.
        """
        if mode == "FIDELIDADE":
            head = self.PROMPT_HEAD_FIDELIDADE
            style = self.PROMPT_STYLE_FIDELIDADE
            structure = self.PROMPT_STRUCTURE_FIDELIDADE
            table = self.PROMPT_TABLE_FIDELIDADE
        elif mode == "AUDIENCIA":
            head = self.PROMPT_HEAD_AUDIENCIA
            style = self.PROMPT_STYLE_AUDIENCIA
            structure = self.PROMPT_STRUCTURE_AUDIENCIA
            table = self.PROMPT_TABLE_AUDIENCIA
        elif mode == "REUNIAO":
            head = self.PROMPT_HEAD_REUNIAO
            style = self.PROMPT_STYLE_REUNIAO
            structure = self.PROMPT_STRUCTURE_REUNIAO
            table = self.PROMPT_TABLE_REUNIAO
        elif mode == "DEPOIMENTO":
            head = self.PROMPT_HEAD_DEPOIMENTO
            style = self.PROMPT_STYLE_DEPOIMENTO
            structure = self.PROMPT_STRUCTURE_DEPOIMENTO
            table = self.PROMPT_TABLE_DEPOIMENTO
        else:  # Default to APOSTILA
            head = self.PROMPT_HEAD_APOSTILA
            style = self.PROMPT_STYLE_APOSTILA
            structure = self.PROMPT_STRUCTURE_APOSTILA
            table = self.PROMPT_TABLE_APOSTILA

        disable_tables = bool(disable_tables)
        if disable_tables:
            table = (
                "## 🚫 TABELAS / EXTRAS (DESABILITADO)\n"
                "- **Não gere tabelas em Markdown** (linhas com `| ... |` e separadores `---`).\n"
                "- **Não inclua** quadro-síntese, pegadinhas, checklists, resumo, fluxograma, mapa mental ou questionário.\n"
                "- Se precisar destacar informações, use **parágrafos** e **listas**.\n"
            )

        if allow_indirect:
            for line in (
                "**NÃO transforme em discurso indireto**. Mantenha fala direta.",
                "**NÃO transforme em discurso indireto** (ex.: \"o juiz disse que...\"). Mantenha fala direta.",
            ):
                if line in head:
                    head = head.replace(
                        line,
                        "**Discurso indireto permitido**. Você pode reescrever falas em estilo indireto, sem inventar conteúdo."
                    )

        if allow_summary and "**NÃO transforme em ata resumida**. Preserve a sequência real das falas." in head:
            head = head.replace(
                "**NÃO transforme em ata resumida**. Preserve a sequência real das falas.",
                "**Ata resumida permitida**. Você pode condensar falas, mantendo decisões, encaminhamentos, nomes, datas, valores e prazos."
            )
            structure = f"{structure}\n- Com ata resumida habilitada, você pode agrupar por pauta/tema e condensar falas, sem inventar informações."

        mode_norm = (mode or "").strip().upper()
        include_timestamps = bool(include_timestamps)
        if mode_norm in {"AUDIENCIA", "REUNIAO", "DEPOIMENTO"} and not include_timestamps:
            # Evita conflito com regras de "preservar timestamps" quando a UI/API pede remoção.
            head = re.sub(
                r"(?m)^-\\s*\\*\\*Timestamps[^\\n]*\\n?",
                "",
                head,
            )
            head = head.replace(
                "- **Marcações**: [inaudível], [risos], [interrupção] e timestamps.",
                "- **Marcações**: [inaudível], [risos], [interrupção] (sem timestamps).",
            )
            head = head.replace(
                "4. **PRESERVE marcações existentes**: [inaudível], [risos], [interrupção] e timestamps.",
                "4. **PRESERVE marcações existentes**: [inaudível], [risos], [interrupção]. **NÃO inclua timestamps**.",
            )
            style = (
                f"{style}\n\n"
                "## ⏱️ TIMESTAMPS (CONFIGURAÇÃO)\n"
                "- **Não incluir timestamps** no texto de saída.\n"
                "- Se houver timestamps no input, remova-os (ex.: `[00:10]`, `[01:02:03]`).\n"
                "- **Não invente** timestamps.\n"
            )
        
        footer = self.PROMPT_FOOTER

        custom_override = (custom_style_override or "").strip()
        scope = (custom_prompt_scope or "tables_only").lower()

        if custom_override and disable_tables and scope == "tables_only":
            print(
                f"{Fore.YELLOW}⚠️  Tabelas/extras desabilitados: ignorando prompt customizado (ele só afeta tabelas/extras neste modo).{Style.RESET_ALL}"
            )
            custom_override = ""

        if custom_override:
            custom_lower = custom_override.lower()

            # Warn only for "structural" headings (#/##/###). ####+ is acceptable for intra-section extras.
            if re.search(r"(^|\n)\s{0,3}#{1,3}\s", custom_override):
                print(
                    f"{Fore.YELLOW}⚠️  Seu prompt customizado contém títulos Markdown (#/##/###). "
                    f"Isso pode interferir na estrutura. Ideal: use no máximo #### para anexos/extras.{Style.RESET_ALL}"
                )
            if any(key in custom_lower for key in ("estrutura", "títulos", "titulos", "sumário", "sumario", "seção", "secao")):
                print(
                    f"{Fore.YELLOW}⚠️  Seu prompt customizado menciona estrutura/títulos/sumário. "
                    f"Para evitar conflitos, restrinja o custom a TABELAS e EXTRAS (resumo/fluxograma/mapa mental/questionário).{Style.RESET_ALL}"
                )

            if scope == "tables_only":
                # tables_only (padrão para TODOS os modos, incluindo FIDELIDADE):
                # custom_prompt afeta SOMENTE tabelas/extras, preservando estilo/estrutura.
                mode_label = {
                    "APOSTILA": "APOSTILA", "AUDIENCIA": "AUDIÊNCIA",
                    "REUNIAO": "REUNIÃO", "FIDELIDADE": "FIDELIDADE",
                }.get(mode.upper(), mode.upper())
                print(f"{Fore.YELLOW}🧩 Usando PROMPT CUSTOMIZADO ({mode_label}: apenas tabelas/extras) ({len(custom_override):,} chars)")
                table_with_custom = (
                    f"{table}\n\n"
                    "## 🧩 PERSONALIZAÇÕES (TABELAS / EXTRAS)\n"
                    "As instruções abaixo são do usuário e se aplicam SOMENTE ao fechamento do tópico:\n"
                    "- Quadros-síntese e tabelas (colunas, critérios, inclusão/omissão de seções de fechamento)\n"
                    "- Anexos ao final do tópico (ex.: resumo, fluxograma, mapa mental, questionário)\n\n"
                    "**REGRAS DE SEGURANÇA (NÃO NEGOCIE):**\n"
                    f"- NÃO altere o tom/estilo do modo {mode.upper()}.\n"
                    "- NÃO altere a estrutura principal (##/###/#### do conteúdo). Se precisar de anexos, use apenas `####` após o bloco de encerramento do tópico.\n"
                    "- NÃO resuma o conteúdo principal; anexos são complementares.\n\n"
                    "### Instruções do usuário\n"
                    f"{custom_override}\n"
                )
                composed = f"{head}\n\n{style}\n\n{structure}\n\n{table_with_custom}\n\n{footer}"
            elif scope == "style_and_tables":
                # Avançado (opt-in explícito): substitui STYLE+TABLE layers
                print(f"{Fore.YELLOW}🎨 Usando PROMPT CUSTOMIZADO avançado de estilo+tabela ({len(custom_override):,} chars)")
                composed = f"{head}\n\n{custom_override}\n\n{structure}\n\n{table}\n\n{footer}"
            else:
                # Fallback seguro para scope desconhecido → tables_only
                composed = f"{head}\n\n{style}\n\n{structure}\n\n{table}\n\n{footer}"
        else:
            # Use default components
            composed = f"{head}\n\n{style}\n\n{structure}\n\n{table}\n\n{footer}"

        # Instrução de idioma de saída (padrão: mesmo idioma do áudio de entrada)
        output_lang = getattr(self, "_output_language", None)
        input_lang = getattr(self, "_current_language", "pt") or "pt"
        effective_lang = output_lang or input_lang

        lang_names = {
            "en": "English",
            "es": "español",
            "fr": "français",
            "de": "Deutsch",
            "pt": "português",
        }

        if effective_lang == "auto":
            # Auto-detect: instruir LLM a manter o idioma do texto de entrada
            composed += (
                "\n\n## IDIOMA DE SAÍDA\n"
                "- O idioma do áudio foi detectado automaticamente.\n"
                "- Identifique o idioma do texto de entrada e escreva TODA a saída nesse MESMO idioma.\n"
                "- Títulos, tabelas, legendas e conteúdo devem estar no idioma original.\n"
                "- NÃO traduza para português se o áudio não for em português.\n"
            )
        elif effective_lang and effective_lang != "pt":
            lang_name = lang_names.get(effective_lang, effective_lang)
            composed += (
                f"\n\n## IDIOMA DE SAÍDA\n"
                f"- O áudio de entrada está em **{lang_name}**.\n"
                f"- Toda a formatação, títulos, tabelas e conteúdo DEVEM ser escritos em **{lang_name}**.\n"
                f"- NÃO traduza para português. Mantenha o idioma original do áudio.\n"
            )

        return composed

    def create_context_cache(self, transcription, global_structure=None):
        """
        v2.2: Wrapper method that calls criar_cache_contexto with class attributes.
        
        Args:
            transcription: Full transcription text
            global_structure: Mapped structure (optional)
        
        Returns:
            Cache object or None
        """
        return criar_cache_contexto(
            client=self.client,
            transcricao_completa=transcription,
            system_prompt=self.prompt_apostila,  # Uses current mode's prompt
            estrutura_global=global_structure,
            model_name=self.llm_model
        )

    def optimize_audio(self, file_path):
        """Extrai áudio otimizado (16kHz mono)"""
        print(f"{Fore.YELLOW}⚡ Verificando áudio...")

        mp3_path = Path(file_path).with_suffix('.mp3')
        if mp3_path.exists():
            print(f"   📂 Usando MP3 existente: {mp3_path.name}")
            return str(mp3_path)

        # Cache baseado em nome do arquivo + tamanho (independente do job ID)
        file_name = Path(file_path).name
        try:
            file_size = os.path.getsize(file_path)
        except OSError:
            file_size = 0
        cache_key = f"{file_name}_{file_size}"
        file_hash = hashlib.md5(cache_key.encode()).hexdigest()[:8]
        output_path = f"temp_{file_hash}.wav"

        if os.path.exists(output_path):
            print(f"   ♻️ Cache encontrado: {output_path} (mesmo arquivo: {file_name})")
            return output_path
        
        print(f"   🔄 Extraindo áudio...")
        enable_loudnorm = str(os.environ.get("IUDEX_AUDIO_LOUDNORM", "1")).strip().lower() not in {"0", "false", "no", "off"}
        ffmpeg_cmd = [
            "ffmpeg",
            "-i",
            file_path,
            "-vn",  # Sem vídeo
            "-sn",  # Sem legendas
            "-dn",  # Sem data streams
            "-map",
            "0:a:0?",
        ]
        if enable_loudnorm:
            ffmpeg_cmd += ["-af", "loudnorm=I=-16:TP=-1.5:LRA=11"]
        ffmpeg_cmd += [
            "-ac",
            "1",  # Mono
            "-ar",
            "16000",  # 16kHz para Whisper
            "-acodec",
            "pcm_s16le",
            output_path,
            "-y",
            "-hide_banner",
            "-loglevel",
            "error",
        ]
        subprocess.run(ffmpeg_cmd, check=True, capture_output=True)
        
        return output_path

    def _resolve_whisper_language(self) -> Optional[str]:
        """Resolve o código de idioma para o Whisper. None = auto-detect."""
        lang = getattr(self, "_current_language", "pt") or "pt"
        return self.SUPPORTED_LANGUAGES.get(lang, lang if lang != "auto" else None)

    def _detect_speech_segments_silero(self, audio_path: str) -> list[dict]:
        """
        Detecta segmentos de fala usando Silero VAD (mais preciso que RMS).

        Returns:
            Lista de dicts {'start': float, 'end': float} com segmentos de fala em segundos.
        """
        try:
            from silero_vad import load_silero_vad, read_audio, get_speech_timestamps

            model = load_silero_vad()
            wav = read_audio(audio_path)

            # Obter timestamps de fala (em segundos)
            speech_timestamps = get_speech_timestamps(
                wav,
                model,
                return_seconds=True,
                min_speech_duration_ms=500,   # Ignora sons < 0.5s
                min_silence_duration_ms=500,  # Une pausas < 0.5s
            )

            return speech_timestamps

        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ Silero VAD falhou: {e}{Style.RESET_ALL}")
            return []

    def _detect_speech_start_silero(self, audio_path: str) -> float:
        """
        Detecta onde a fala começa usando Silero VAD.

        Returns:
            Offset em segundos onde a fala começa (0 se não detectar silêncio inicial)
        """
        segments = self._detect_speech_segments_silero(audio_path)

        if segments and len(segments) > 0:
            first_speech = segments[0].get('start', 0)
            if first_speech > 0:
                print(f"{Fore.YELLOW}   🔇 Silero VAD: Detectado silêncio inicial de {first_speech:.0f}s{Style.RESET_ALL}")
            return first_speech

        return 0.0

    def _detect_speech_start_rms(self, audio_path: str, chunk_seconds: float = 30.0, threshold_db: float = -40.0) -> float:
        """
        Fallback: Detecta onde a fala começa usando RMS (energia do sinal).
        Usado se Silero VAD não estiver disponível.
        """
        import subprocess
        import re

        try:
            probe_cmd = [
                "ffprobe", "-v", "error", "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1", audio_path
            ]
            duration_result = subprocess.run(probe_cmd, capture_output=True, text=True)
            total_duration = float(duration_result.stdout.strip())

            speech_start = 0.0
            max_chunks_to_check = min(20, int(total_duration / chunk_seconds))

            for i in range(max_chunks_to_check):
                offset = i * chunk_seconds
                cmd = [
                    "ffmpeg", "-ss", str(offset), "-t", str(chunk_seconds),
                    "-i", audio_path, "-af", "volumedetect", "-f", "null", "-"
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, stderr=subprocess.STDOUT)
                output = result.stdout + result.stderr if result.stderr else result.stdout

                match = re.search(r"mean_volume:\s*([-\d.]+)\s*dB", output)
                if match:
                    mean_vol = float(match.group(1))
                    if mean_vol > threshold_db:
                        if i > 0:
                            fine_start = max(0, (i - 1) * chunk_seconds)
                            fine_chunk = 5.0
                            for j in range(int(chunk_seconds / fine_chunk) + 2):
                                fine_offset = fine_start + j * fine_chunk
                                if fine_offset >= offset + chunk_seconds:
                                    break
                                cmd_fine = [
                                    "ffmpeg", "-ss", str(fine_offset), "-t", str(fine_chunk),
                                    "-i", audio_path, "-af", "volumedetect", "-f", "null", "-"
                                ]
                                result_fine = subprocess.run(cmd_fine, capture_output=True, text=True, stderr=subprocess.STDOUT)
                                output_fine = result_fine.stdout + (result_fine.stderr or "")
                                match_fine = re.search(r"mean_volume:\s*([-\d.]+)\s*dB", output_fine)
                                if match_fine and float(match_fine.group(1)) > threshold_db:
                                    speech_start = fine_offset
                                    break
                            else:
                                speech_start = offset
                        break

            if speech_start > 0:
                print(f"{Fore.YELLOW}   🔇 RMS VAD: Detectado silêncio inicial de {speech_start:.0f}s{Style.RESET_ALL}")

            return speech_start

        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ RMS VAD falhou: {e}{Style.RESET_ALL}")
            return 0.0

    # ==================== CHUNKING DE ÁUDIO LONGO (v2.32) ====================

    def _get_audio_duration(self, audio_path: str) -> float:
        """
        Obtém a duração do áudio em segundos usando ffprobe.

        v2.34: Melhorada robustez com validação e fallback por tamanho.
        """
        duration = 0.0

        # Método 1: ffprobe
        try:
            cmd = [
                "ffprobe", "-v", "error", "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1", audio_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                duration = float(result.stdout.strip())
                if duration > 0:
                    return duration
        except subprocess.TimeoutExpired:
            print(f"{Fore.YELLOW}   ⚠️ ffprobe timeout para {audio_path}{Style.RESET_ALL}")
        except (ValueError, Exception) as e:
            print(f"{Fore.YELLOW}   ⚠️ Erro ffprobe: {e}{Style.RESET_ALL}")

        # Método 2: wave module (apenas para WAV)
        if audio_path.lower().endswith('.wav'):
            try:
                import wave
                with wave.open(audio_path, 'rb') as wav:
                    frames = wav.getnframes()
                    rate = wav.getframerate()
                    if rate > 0:
                        duration = float(frames) / float(rate)
                        if duration > 0:
                            print(f"{Fore.CYAN}   📏 Duração via wave: {duration/3600:.2f}h{Style.RESET_ALL}")
                            return duration
            except Exception:
                pass

        # Método 3: Estimativa por tamanho do arquivo (fallback)
        # MP3 ~128kbps = ~960KB/min, WAV 16kHz mono = ~1.92MB/min
        try:
            file_size_mb = os.path.getsize(audio_path) / (1024 * 1024)
            if audio_path.lower().endswith('.wav'):
                # WAV 16kHz mono 16-bit = ~1.92 MB/min
                estimated_minutes = file_size_mb / 1.92
            else:
                # MP3 ~128kbps = ~0.96 MB/min (conservador)
                estimated_minutes = file_size_mb / 0.96

            estimated_seconds = estimated_minutes * 60
            if estimated_seconds > self.AUDIO_MAX_DURATION_SECONDS:
                print(f"{Fore.YELLOW}   ⚠️ Duração estimada por tamanho: {estimated_seconds/3600:.1f}h (arquivo: {file_size_mb:.0f}MB){Style.RESET_ALL}")
                return estimated_seconds
        except Exception:
            pass

        return duration

    def _split_audio_into_chunks(self, audio_path: str, chunk_duration: float, overlap: float = 30.0) -> list:
        """
        Divide áudio longo em chunks temporários.

        v2.32: Evita problemas de memória do MLX-Whisper com arquivos muito longos.

        Args:
            audio_path: Caminho do arquivo de áudio
            chunk_duration: Duração de cada chunk em segundos
            overlap: Overlap entre chunks em segundos (para continuidade)

        Returns:
            Lista de dicts: [{'path': str, 'start': float, 'end': float, 'is_temp': bool}]
        """
        import tempfile

        total_duration = self._get_audio_duration(audio_path)
        if total_duration <= 0:
            return [{'path': audio_path, 'start': 0, 'end': 0, 'is_temp': False}]

        # Se áudio é menor que o limite, retorna sem dividir
        if total_duration <= chunk_duration:
            return [{'path': audio_path, 'start': 0, 'end': total_duration, 'is_temp': False}]

        chunks = []
        current_start = 0.0
        chunk_index = 0

        print(f"{Fore.CYAN}   🔪 Dividindo áudio longo ({total_duration/3600:.1f}h) em chunks de {chunk_duration/3600:.1f}h...{Style.RESET_ALL}")

        while current_start < total_duration:
            chunk_end = min(current_start + chunk_duration, total_duration)
            actual_duration = chunk_end - current_start

            # Criar arquivo temporário para o chunk
            base_name = Path(audio_path).stem
            temp_dir = tempfile.gettempdir()
            chunk_path = os.path.join(temp_dir, f"{base_name}_chunk{chunk_index}.wav")

            # Extrair chunk com ffmpeg
            cmd = [
                "ffmpeg", "-y",
                "-ss", str(current_start),
                "-i", audio_path,
                "-t", str(actual_duration),
                "-ar", "16000",
                "-ac", "1",
                "-acodec", "pcm_s16le",
                chunk_path,
                "-hide_banner",
                "-loglevel", "error"
            ]

            try:
                subprocess.run(cmd, check=True, capture_output=True)
                chunks.append({
                    'path': chunk_path,
                    'start': current_start,
                    'end': chunk_end,
                    'is_temp': True
                })
                print(f"{Fore.GREEN}      ✂️ Chunk {chunk_index + 1}: {self._format_timestamp(current_start)} → {self._format_timestamp(chunk_end)}{Style.RESET_ALL}")
            except subprocess.CalledProcessError as e:
                print(f"{Fore.RED}      ❌ Erro ao criar chunk {chunk_index}: {e}{Style.RESET_ALL}")

            # Próximo chunk com overlap
            current_start = chunk_end - overlap
            if current_start >= total_duration - overlap:
                break
            chunk_index += 1

        print(f"{Fore.GREEN}   ✅ Criados {len(chunks)} chunks de áudio{Style.RESET_ALL}")
        return chunks

    def _cleanup_audio_chunks(self, chunks: list):
        """Remove arquivos temporários de chunks."""
        for chunk in chunks:
            if chunk.get('is_temp') and os.path.exists(chunk['path']):
                try:
                    os.unlink(chunk['path'])
                except Exception:
                    pass

    def _merge_chunk_segments(self, all_segments: list, overlap_seconds: float = 30.0) -> list:
        """
        Mescla segmentos de múltiplos chunks, removendo duplicatas do overlap.

        v2.32: Usa fingerprinting de texto para detectar e remover duplicatas.
        v2.34: Melhorado algoritmo de detecção de duplicatas com múltiplas estratégias.
        """
        if not all_segments:
            return []

        merged = []
        last_end_time = 0.0
        recent_texts = []  # Buffer das últimas N frases para detectar duplicatas

        def normalize_for_compare(text: str) -> str:
            """Normaliza texto para comparação (lowercase, sem pontuação extra)."""
            import re
            text = (text or '').strip().lower()
            text = re.sub(r'[^\w\s]', '', text)  # Remove pontuação
            text = re.sub(r'\s+', ' ', text)  # Normaliza espaços
            return text

        def is_duplicate(new_text: str, recent: list) -> bool:
            """Verifica se texto é duplicata de algum texto recente."""
            new_norm = normalize_for_compare(new_text)
            if not new_norm or len(new_norm) < 10:
                return False

            for old_text in recent:
                old_norm = normalize_for_compare(old_text)
                if not old_norm:
                    continue

                # Estratégia 1: Texto exatamente igual
                if new_norm == old_norm:
                    return True

                # Estratégia 2: Um contém o outro (substring)
                if len(new_norm) > 20 and len(old_norm) > 20:
                    if new_norm in old_norm or old_norm in new_norm:
                        return True

                # Estratégia 3: Similaridade alta (Jaccard de palavras)
                new_words = set(new_norm.split())
                old_words = set(old_norm.split())
                if new_words and old_words:
                    intersection = len(new_words & old_words)
                    union = len(new_words | old_words)
                    if union > 0 and intersection / union > 0.8:  # 80% similaridade
                        return True

                # Estratégia 4: Início igual (primeiras N palavras)
                new_start = ' '.join(new_norm.split()[:8])
                old_start = ' '.join(old_norm.split()[:8])
                if len(new_start) > 20 and new_start == old_start:
                    return True

            return False

        for chunk_idx, segments in enumerate(all_segments):
            chunk_start_time = segments[0].get('start', 0) if segments else 0

            for seg in segments:
                seg_start = seg.get('start', 0)
                seg_text = (seg.get('text') or '').strip()

                if not seg_text:
                    continue

                # Para segmentos no período de overlap, verificar duplicatas mais rigorosamente
                in_overlap_zone = seg_start < last_end_time + overlap_seconds * 0.5

                if in_overlap_zone and chunk_idx > 0:
                    # Verificar se é duplicata
                    if is_duplicate(seg_text, recent_texts):
                        continue

                merged.append(seg)
                last_end_time = max(last_end_time, seg.get('end', seg_start))

                # Manter buffer das últimas 10 frases para comparação
                recent_texts.append(seg_text)
                if len(recent_texts) > 10:
                    recent_texts.pop(0)

        print(f"{Fore.CYAN}   🔗 Merge: {sum(len(s) for s in all_segments)} → {len(merged)} segmentos (removidas duplicatas do overlap){Style.RESET_ALL}")
        return merged

    def _transcribe_chunked(self, audio_path: str, *, beam_size: Optional[int] = None, cache_file: str = None, initial_prompt: str = "") -> str:
        """
        Transcreve áudio longo dividindo em chunks de 3h cada.

        v2.32: Evita degradação do MLX-Whisper com arquivos muito longos.
        O Whisper pode gerar apenas pontuação quando processa áudios > 3-4h de uma vez.
        """
        print(f"{Fore.CYAN}   🎬 Iniciando transcrição em chunks (máx {self.AUDIO_MAX_DURATION_SECONDS/3600:.0f}h cada)...{Style.RESET_ALL}")
        start_time = time.time()

        # Dividir áudio em chunks
        chunks = self._split_audio_into_chunks(
            audio_path,
            chunk_duration=self.AUDIO_MAX_DURATION_SECONDS,
            overlap=self.AUDIO_CHUNK_OVERLAP_SECONDS
        )

        if not chunks:
            print(f"{Fore.RED}   ❌ Falha ao dividir áudio em chunks{Style.RESET_ALL}")
            return ""

        all_segments = []
        whisper_lang = self._resolve_whisper_language()
        no_speech_thresh = float(os.getenv("VOMO_NO_SPEECH_THRESHOLD", "0.8"))

        mlx_kwargs = dict(
            path_or_hf_repo=f"mlx-community/whisper-{self.model_name}",
            **({"language": whisper_lang} if whisper_lang else {}),
            temperature=0.0,
            initial_prompt=(initial_prompt or None),
            word_timestamps=True,
            fp16=True,
            no_speech_threshold=no_speech_thresh,
            logprob_threshold=-0.5,
            compression_ratio_threshold=2.2,
            condition_on_previous_text=self._condition_on_previous,
            suppress_tokens=[-1],
            verbose=False,
        )
        if beam_size and beam_size > 1:
            mlx_kwargs["beam_size"] = int(beam_size)

        try:
            for i, chunk in enumerate(chunks):
                chunk_path = chunk['path']
                chunk_start = chunk['start']
                chunk_end = chunk['end']

                print(f"{Fore.CYAN}   📝 Transcrevendo chunk {i+1}/{len(chunks)} ({self._format_timestamp(chunk_start)} → {self._format_timestamp(chunk_end)})...{Style.RESET_ALL}")

                try:
                    result = self._transcribe_with_vad(chunk_path, mlx_kwargs, skip_silence=True)
                except TypeError:
                    mlx_kwargs_copy = dict(mlx_kwargs)
                    mlx_kwargs_copy.pop("beam_size", None)
                    mlx_kwargs_copy.pop("best_of", None)
                    result = self._transcribe_with_vad(chunk_path, mlx_kwargs_copy, skip_silence=True)

                segments = result.get("segments", [])

                # Ajustar timestamps para o offset do chunk (segmentos E words)
                for seg in segments:
                    seg['start'] = seg.get('start', 0) + chunk_start
                    seg['end'] = seg.get('end', 0) + chunk_start
                    # v2.33: Ajustar também timestamps das words individuais
                    if 'words' in seg and seg['words']:
                        for word in seg['words']:
                            if 'start' in word:
                                word['start'] = word['start'] + chunk_start
                            if 'end' in word:
                                word['end'] = word['end'] + chunk_start

                all_segments.append(segments)
                print(f"{Fore.GREEN}      ✅ Chunk {i+1}: {len(segments)} segmentos{Style.RESET_ALL}")

        finally:
            # Limpar arquivos temporários
            self._cleanup_audio_chunks(chunks)

        # Mesclar segmentos de todos os chunks
        merged_segments = self._merge_chunk_segments(all_segments, overlap_seconds=self.AUDIO_CHUNK_OVERLAP_SECONDS)

        # Filtrar segmentos
        segments, filter_stats = self._filter_asr_segments(merged_segments)
        if filter_stats.get("dropped"):
            reasons = ", ".join(
                f"{k}={v}" for k, v in sorted((filter_stats.get("reason_counts") or {}).items())
            )
            print(f"{Fore.YELLOW}   🧹 ASR: removidos {filter_stats['dropped']} segmento(s) suspeitos ({reasons}){Style.RESET_ALL}")

        # Formatar resultado
        lines = []
        current_block = []
        last_timestamp = None

        for segment in segments:
            start = segment['start']
            text = segment['text'].strip()

            if not text:
                continue

            text = self._normalize_raw_text(text)

            if self._should_add_timestamp(start, last_timestamp, interval_seconds=self._get_timestamp_interval_for_mode()):
                if current_block:
                    lines.append(" ".join(current_block))
                    current_block = []

                ts = self._format_timestamp(start)
                current_block.append(f"[{ts}] {text}")
                last_timestamp = start
            else:
                current_block.append(text)

        if current_block:
            lines.append(" ".join(current_block))

        transcript_result = "\n\n".join(lines).strip()
        transcript_result = self._strip_leaked_initial_prompt(transcript_result, initial_prompt)

        elapsed = time.time() - start_time
        print(f"{Fore.GREEN}   ✅ Transcrição chunked concluída em {elapsed:.1f}s ({len(chunks)} chunks, {len(segments)} segmentos){Style.RESET_ALL}")

        # Salvar cache
        if cache_file:
            try:
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump({
                        'transcript': transcript_result,
                        'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
                        'chunks': len(chunks)
                    }, f, ensure_ascii=False, indent=2)
            except Exception as e:
                print(f"Erro ao salvar cache: {e}")

        return transcript_result

    def _transcribe_with_segments_chunked(self, audio_path: str, *, beam_size: Optional[int] = None) -> dict:
        """
        v2.33: Transcreve áudio longo em chunks com suporte a segmentos e diarização.

        Divide o áudio em chunks menores, transcreve cada um, ajusta timestamps,
        e opcionalmente aplica diarização por chunk para evitar estouro de memória.
        """
        import gc

        print(f"{Fore.CYAN}   🎬 Iniciando transcrição chunked com segmentos (máx {self.AUDIO_MAX_DURATION_SECONDS/3600:.0f}h cada)...{Style.RESET_ALL}")
        start_time = time.time()

        # Dividir áudio em chunks
        chunks = self._split_audio_into_chunks(
            audio_path,
            chunk_duration=self.AUDIO_MAX_DURATION_SECONDS,
            overlap=self.AUDIO_CHUNK_OVERLAP_SECONDS
        )

        if not chunks:
            print(f"{Fore.RED}   ❌ Falha ao dividir áudio em chunks{Style.RESET_ALL}")
            return {"text": "", "segments": [], "words": [], "diarization": []}

        # Preparar kwargs do Whisper
        initial_prompt = self._get_whisper_initial_prompt_for_asr(high_accuracy=bool(beam_size and beam_size > 1)) or ""
        whisper_lang = self._resolve_whisper_language()
        mlx_kwargs = dict(
            path_or_hf_repo=f"mlx-community/whisper-{self.model_name}",
            **({"language": whisper_lang} if whisper_lang else {}),
            temperature=0.0,
            initial_prompt=(initial_prompt or None),
            word_timestamps=True,
            fp16=True,
            no_speech_threshold=float(os.getenv("VOMO_NO_SPEECH_THRESHOLD", "0.8")),
            logprob_threshold=-1.0,
            compression_ratio_threshold=float(os.getenv("VOMO_COMPRESSION_THRESHOLD", "2.2")),
            condition_on_previous_text=self._condition_on_previous,
            suppress_tokens=[-1],
            verbose=False,
        )
        if beam_size and beam_size > 1:
            mlx_kwargs["beam_size"] = int(beam_size)
            mlx_kwargs["best_of"] = int(beam_size)

        all_segments = []
        all_diarization = []
        token = self._get_hf_token() if self._diarization_enabled else None
        diarization_pipeline = None

        # Inicializar pipeline de diarização uma vez (se habilitado)
        if self._diarization_enabled and Pipeline and token:
            try:
                self._ensure_diarization_available_or_raise()
                diarization_pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-community-1",
                    token=token
                )
                device = "mps" if torch.backends.mps.is_available() else "cpu"
                diarization_pipeline.to(torch.device(device))
                print(f"{Fore.GREEN}   ✅ Pipeline de diarização inicializado ({device}){Style.RESET_ALL}")
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Erro ao inicializar diarização: {e}{Style.RESET_ALL}")
                diarization_pipeline = None

        try:
            for i, chunk in enumerate(chunks):
                chunk_path = chunk['path']
                chunk_start = chunk['start']
                chunk_end = chunk['end']

                print(f"{Fore.CYAN}   📝 Transcrevendo chunk {i+1}/{len(chunks)} ({self._format_timestamp(chunk_start)} → {self._format_timestamp(chunk_end)})...{Style.RESET_ALL}")

                # Transcrever chunk
                try:
                    result = self._transcribe_with_vad(chunk_path, mlx_kwargs, skip_silence=True)
                except TypeError:
                    mlx_kwargs_copy = dict(mlx_kwargs)
                    mlx_kwargs_copy.pop("beam_size", None)
                    mlx_kwargs_copy.pop("best_of", None)
                    result = self._transcribe_with_vad(chunk_path, mlx_kwargs_copy, skip_silence=True)

                segments = result.get("segments", [])

                # Ajustar timestamps (segmentos e words)
                for seg in segments:
                    seg['start'] = seg.get('start', 0) + chunk_start
                    seg['end'] = seg.get('end', 0) + chunk_start
                    if 'words' in seg and seg['words']:
                        for word in seg['words']:
                            if 'start' in word:
                                word['start'] = word['start'] + chunk_start
                            if 'end' in word:
                                word['end'] = word['end'] + chunk_start

                # Diarização por chunk (se habilitado)
                if diarization_pipeline:
                    try:
                        diarization = diarization_pipeline(chunk_path)
                        for turn, _, speaker in diarization.itertracks(yield_label=True):
                            speaker_id = speaker.split('_')[-1]
                            all_diarization.append({
                                "start": float(turn.start) + chunk_start,
                                "end": float(turn.end) + chunk_start,
                                "speaker_label": f"SPEAKER {int(speaker_id) + 1}"
                            })
                        # Criar segmentos temporários com timestamps originais do chunk para atribuição
                        temp_segments = []
                        for seg in segments:
                            temp_seg = dict(seg)
                            temp_seg['start'] = seg['start'] - chunk_start
                            temp_seg['end'] = seg['end'] - chunk_start
                            temp_segments.append(temp_seg)
                        # Atribuir labels de diarização
                        labeled_temp = self._assign_diarization_labels(temp_segments, diarization)
                        # Copiar labels de volta para segments
                        for seg, labeled in zip(segments, labeled_temp):
                            seg['speaker_label'] = labeled.get('speaker_label', 'SPEAKER 1')
                    except Exception as e:
                        print(f"{Fore.YELLOW}      ⚠️ Diarização chunk {i+1} falhou: {e}{Style.RESET_ALL}")
                        for seg in segments:
                            seg['speaker_label'] = "SPEAKER 1"
                else:
                    # Sem diarização - atribuir SPEAKER 1
                    for seg in segments:
                        seg['speaker_label'] = "SPEAKER 1"

                all_segments.append(segments)
                print(f"{Fore.GREEN}      ✅ Chunk {i+1}: {len(segments)} segmentos{Style.RESET_ALL}")

                # Liberar memória entre chunks
                gc.collect()
                if torch and torch.backends.mps.is_available():
                    try:
                        torch.mps.empty_cache()
                    except Exception:
                        pass

        finally:
            # Limpar arquivos temporários
            self._cleanup_audio_chunks(chunks)

        # Mesclar segmentos de todos os chunks
        merged_segments = self._merge_chunk_segments(all_segments, overlap_seconds=self.AUDIO_CHUNK_OVERLAP_SECONDS)

        # Filtrar segmentos suspeitos
        filtered_segments, filter_stats = self._filter_asr_segments(merged_segments)
        if filter_stats.get("dropped"):
            reasons = ", ".join(
                f"{k}={v}" for k, v in sorted((filter_stats.get("reason_counts") or {}).items())
            )
            print(f"{Fore.YELLOW}   🧹 ASR: removidos {filter_stats['dropped']} segmento(s) suspeitos ({reasons}){Style.RESET_ALL}")

        # Garantir que todos os segmentos têm speaker_label
        for seg in filtered_segments:
            if 'speaker_label' not in seg:
                seg['speaker_label'] = "SPEAKER 1"
            if 'words' not in seg:
                seg['words'] = []

        # Extrair lista flat de words
        all_words = []
        for seg in filtered_segments:
            seg_words = seg.get("words", [])
            speaker = seg.get("speaker_label", "")
            for w in seg_words:
                all_words.append({
                    "word": w.get("word", w.get("text", "")),
                    "start": w.get("start", 0),
                    "end": w.get("end", 0),
                    "speaker": speaker,
                })

        # Gerar texto formatado
        transcript_text = self._segments_to_text(filtered_segments)
        transcript_text = self._strip_leaked_initial_prompt(transcript_text, initial_prompt)

        elapsed = time.time() - start_time
        print(f"{Fore.GREEN}   ✅ Transcrição chunked concluída em {elapsed:.1f}s ({len(chunks)} chunks, {len(filtered_segments)} segmentos){Style.RESET_ALL}")

        return {
            "text": transcript_text,
            "segments": filtered_segments,
            "words": all_words,
            "diarization": all_diarization
        }

    def _transcribe_with_vad(self, audio_path: str, mlx_kwargs: dict, skip_silence: bool = True) -> dict:
        """
        Transcreve áudio com detecção de atividade de voz (VAD).

        Pipeline:
        1. Silero VAD detecta onde há fala (mais preciso)
        2. Fallback para RMS se Silero falhar
        3. Se silêncio inicial > 30s, pula para economizar processamento
        """
        import tempfile

        speech_start = 0.0
        temp_audio = None

        if skip_silence and _env_truthy("VOMO_VAD_SKIP_SILENCE", default=True):
            # Tentar Silero VAD primeiro (mais preciso)
            try:
                speech_start = self._detect_speech_start_silero(audio_path)
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Silero VAD indisponível, usando RMS: {e}{Style.RESET_ALL}")
                speech_start = self._detect_speech_start_rms(audio_path)

            # Se há mais de 30s de silêncio inicial, criar arquivo sem o silêncio
            if speech_start > 30:
                print(f"{Fore.CYAN}   ✂️ Pulando {speech_start:.0f}s de silêncio inicial...{Style.RESET_ALL}")

                # Criar arquivo temporário sem o silêncio
                temp_fd, temp_audio = tempfile.mkstemp(suffix=".wav")
                os.close(temp_fd)

                cmd = [
                    "ffmpeg", "-y", "-ss", str(speech_start), "-i", audio_path,
                    "-ar", "16000", "-ac", "1", "-acodec", "pcm_s16le",
                    temp_audio, "-hide_banner", "-loglevel", "error"
                ]
                subprocess.run(cmd, check=True, capture_output=True)
                audio_path = temp_audio

        try:
            result = mlx_whisper.transcribe(audio_path, **mlx_kwargs)

            # Ajustar timestamps se pulamos silêncio (segmentos E words)
            if speech_start > 0 and result.get("segments"):
                for seg in result["segments"]:
                    seg["start"] = seg.get("start", 0) + speech_start
                    seg["end"] = seg.get("end", 0) + speech_start
                    # v2.34: Ajustar também timestamps das words individuais
                    if seg.get("words"):
                        for word in seg["words"]:
                            if "start" in word:
                                word["start"] = word["start"] + speech_start
                            if "end" in word:
                                word["end"] = word["end"] + speech_start
                if "duration" in result:
                    result["duration"] = result["duration"] + speech_start

            return result

        finally:
            # Limpar arquivo temporário
            if temp_audio and os.path.exists(temp_audio):
                try:
                    os.unlink(temp_audio)
                except Exception:
                    pass

    def transcribe(self, audio_path, *, beam_size: Optional[int] = None):
        """
        MLX-Whisper OTIMIZADO com GPU acelerado + Diarização

        Otimizações v2.0:
        - VAD filtering (pula silêncio)
        - Batched inference (múltiplos chunks GPU)
        - condition_on_previous_text (contexto melhorado)
        - Hallucination suppression (evita texto inventado)
        """
        print(f"{Fore.GREEN}🎙️  Iniciando transcrição OTIMIZADA (MLX GPU)...")
        start_time = time.time()
        
        # Cache de transcrição (separa diarização ON/OFF + hash de parâmetros)
        # Importante: incluir parâmetros que mudam o output (ex.: initial_prompt).
        initial_prompt = self._get_whisper_initial_prompt_for_asr(high_accuracy=bool(beam_size and beam_size > 1)) or ""
        prompt_hash = hashlib.sha256(initial_prompt.encode()).hexdigest()[:8] if initial_prompt else "noprompt"
        clean_enabled = _env_truthy("VOMO_FILTER_ASR_HALLUCINATIONS", default=True)
        params_str = f"{self.model_name}_{self._diarization_enabled}_{self._condition_on_previous}_{prompt_hash}_clean{int(bool(clean_enabled))}"
        if beam_size and beam_size > 1:
            params_str += f"_beam{int(beam_size)}"
        params_hash = hashlib.sha256(params_str.encode()).hexdigest()[:8]
        cache_tag = "DIARIZATION" if self._diarization_enabled else "ASR"
        cache_file = audio_path.replace('.wav', f'_{cache_tag}_{params_hash}.json').replace('.mp3', f'_{cache_tag}_{params_hash}.json')
        
        if os.path.exists(cache_file):
            try:
                print(f"{Fore.CYAN}   📂 Cache encontrado, carregando...")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                return cache_data['transcript']
            except Exception:
                pass

        if not mlx_whisper:
            raise ImportError("mlx_whisper não instalado.")

        # ==================== CHUNKING DE ÁUDIO LONGO (v2.32) ====================
        audio_duration = self._get_audio_duration(audio_path)
        max_duration = self.AUDIO_MAX_DURATION_SECONDS

        # v2.34: Log detalhado para debug
        print(f"{Fore.CYAN}   📏 Duração detectada: {audio_duration/3600:.2f}h (limite: {max_duration/3600:.0f}h){Style.RESET_ALL}")

        if audio_duration > max_duration:
            print(f"{Fore.YELLOW}   ⚠️ Áudio longo detectado ({audio_duration/3600:.1f}h > {max_duration/3600:.0f}h) - ATIVANDO CHUNKING{Style.RESET_ALL}")
            return self._transcribe_chunked(audio_path, beam_size=beam_size, cache_file=cache_file, initial_prompt=initial_prompt)
        elif audio_duration == 0:
            print(f"{Fore.RED}   ❌ AVISO: Duração não detectada! Chunking desabilitado. Arquivo: {audio_path}{Style.RESET_ALL}")

        # ==================== PARÂMETROS OTIMIZADOS ====================
        print("   🔍 Transcrevendo com parâmetros otimizados...")
        
        whisper_lang = self._resolve_whisper_language()

        # v2.31: Threshold mais alto para filtrar silêncio/ruído melhor
        no_speech_thresh = float(os.getenv("VOMO_NO_SPEECH_THRESHOLD", "0.8"))

        mlx_kwargs = dict(
            path_or_hf_repo=f"mlx-community/whisper-{self.model_name}",
            **({"language": whisper_lang} if whisper_lang else {}),
            # === PRECISÃO ===
            temperature=0.0,  # Mais determinístico (desativa sampling estocástico)
            # === CONTEXTO E GLOSSÁRIO (v2.29: contextual por modo) ===
            initial_prompt=(initial_prompt or None),
            # === TIMESTAMPS ===
            word_timestamps=True,
            # === PERFORMANCE ===
            fp16=True,  # Usa float16 (mais rápido na GPU)
            # === QUALIDADE (Hallucination Suppression) - v2.31: thresholds ajustados ===
            no_speech_threshold=no_speech_thresh,  # v2.31: 0.8 (era 0.6) - mais agressivo em silêncio
            logprob_threshold=-0.5,  # v2.31: -0.5 (era -1.0) - rejeita tokens menos confiantes
            compression_ratio_threshold=2.2,  # v2.31: 2.2 (era 2.4) - detecta repetição mais cedo
            # === CONTEXTO ===
            condition_on_previous_text=self._condition_on_previous,  # v2.30: configurável via VOMO_CONDITION_PREVIOUS
            # === SUPRESSÃO DE TOKENS PROBLEMÁTICOS ===
            suppress_tokens=[-1],  # Suprime token de padding
            verbose=False,
        )
        if beam_size and beam_size > 1:
            mlx_kwargs["beam_size"] = int(beam_size)
            # `best_of` só é aceito em algumas implementações; aplicamos best-effort.
            mlx_kwargs["best_of"] = int(beam_size)

        # v2.31: Usar VAD para pular silêncio inicial extenso
        try:
            result = self._transcribe_with_vad(audio_path, mlx_kwargs, skip_silence=True)
        except TypeError:
            # Compatibilidade: algumas versões aceitam `beam_size` mas não `best_of` (ou vice-versa).
            if "best_of" in mlx_kwargs:
                mlx_kwargs.pop("best_of", None)
                try:
                    result = self._transcribe_with_vad(audio_path, mlx_kwargs, skip_silence=True)
                except TypeError:
                    mlx_kwargs.pop("beam_size", None)
                    result = self._transcribe_with_vad(audio_path, mlx_kwargs, skip_silence=True)
            else:
                mlx_kwargs.pop("beam_size", None)
                result = self._transcribe_with_vad(audio_path, mlx_kwargs, skip_silence=True)

        segments, filter_stats = self._filter_asr_segments(result.get("segments", []))
        if filter_stats.get("dropped"):
            reasons = ", ".join(
                f"{k}={v}" for k, v in sorted((filter_stats.get("reason_counts") or {}).items())
            )
            print(f"{Fore.YELLOW}   🧹 ASR: removidos {filter_stats['dropped']} segmento(s) suspeitos ({reasons})")
        
        elapsed = time.time() - start_time
        audio_duration = result.get('duration', 0) if isinstance(result, dict) else 0
        rtf = elapsed / audio_duration if audio_duration > 0 else 0
        
        print(f"{Fore.GREEN}   ✅ Transcrição concluída em {elapsed:.1f}s (RTF: {rtf:.2f}x)")
        
        transcript_result = None
        
        # Diarização (condicional por política)
        if self._diarization_enabled:
            self._ensure_diarization_available_or_raise()
        token = self._get_hf_token()
        if self._diarization_enabled and Pipeline and "torch" in globals() and token:
            try:
                print("   🗣️  Iniciando Diarização (Pyannote)...")
                pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-community-1",
                    token=token
                )
                
                device = "mps" if torch.backends.mps.is_available() else "cpu"
                pipeline.to(torch.device(device))
                diarization = pipeline(audio_path)
                
                transcript_result = self._align_diarization(segments, diarization)
                print(f"{Fore.GREEN}✅ Diarização concluída")
            
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Erro na Diarização: {e}")
        
        if transcript_result is None:
            # Fallback sem diarização - v2.28: pré-formatação condensada
            lines = []
            current_block = []
            last_timestamp = None
            
            for segment in segments:
                start = segment['start']
                text = segment['text'].strip()
                
                if not text:
                    continue
                
                # Normalização leve de ruído
                text = self._normalize_raw_text(text)
                
                # Timestamp a cada 60 segundos
                if self._should_add_timestamp(start, last_timestamp, interval_seconds=self._get_timestamp_interval_for_mode()):
                    # Flush previous block
                    if current_block:
                        lines.append(" ".join(current_block))
                        current_block = []
                    
                    ts = self._format_timestamp(start)
                    current_block.append(f"[{ts}] {text}")
                    last_timestamp = start
                else:
                    current_block.append(text)
            
            # Flush final block
            if current_block:
                lines.append(" ".join(current_block))
            
            transcript_result = "\n\n".join(lines).strip()

        transcript_result = self._strip_leaked_initial_prompt(transcript_result, initial_prompt)
        if _env_truthy("VOMO_ASR_NORMALIZE_TEMAS", default=True):
            try:
                transcript_result, stats = self._normalize_asr_temas_consistency(transcript_result)
                if stats.get("changed", 0) > 0:
                    print(
                        f"{Fore.YELLOW}   🧩 ASR: normalizados {stats['changed']} tema(s) inconsistentes "
                        f"(234→1234: {stats.get('fixed_3_to_4', 0)}, variações: {stats.get('fixed_variants', 0)}){Style.RESET_ALL}"
                    )
            except Exception:
                pass
        
        # Salvar cache
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'transcript': transcript_result,
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
                }, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erro ao salvar cache: {e}")
        
        return transcript_result

    def transcribe_beam_search(self, audio_path):
        """
        Transcrição de ALTA PRECISÃO usando faster-whisper com Beam Search.
        
        Usa beam_size=5 para explorar múltiplos caminhos de frase,
        resultando em transcrições mais precisas para termos jurídicos complexos.
        
        Ativar via: --high-accuracy
        """
        beam_size = self._get_asr_beam_size()
        if not FASTER_WHISPER_AVAILABLE:
            print(f"{Fore.YELLOW}⚠️ faster-whisper não instalado. Tentando Beam Search via MLX ({beam_size})...")
            return self.transcribe(audio_path, beam_size=beam_size)
        
        print(f"{Fore.MAGENTA}🎯 Transcrição ALTA PRECISÃO (Beam Search)...")
        start_time = time.time()
        
        # Cache de transcrição (com hash de parâmetros para invalidação)
        beam_model_size = "large-v3-turbo"
        initial_prompt = self._get_whisper_initial_prompt_for_asr(high_accuracy=True) or ""
        prompt_hash = hashlib.sha256(initial_prompt.encode()).hexdigest()[:8] if initial_prompt else "noprompt"
        cache_params = f"{beam_model_size}_{self._condition_on_previous}_{prompt_hash}_beam{beam_size}"
        beam_hash = hashlib.sha256(cache_params.encode()).hexdigest()[:8]
        cache_file = audio_path.replace('.wav', f'_BEAM_SEARCH_{beam_hash}.json').replace('.mp3', f'_BEAM_SEARCH_{beam_hash}.json')
        
        if os.path.exists(cache_file):
            try:
                print(f"{Fore.CYAN}   📂 Cache Beam Search encontrado, carregando...")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                return cache_data['transcript']
            except Exception:
                pass
        
        # Modelo faster-whisper (CPU/GPU via ctranslate2)
        model_size = "large-v3-turbo"  # Compatível com o modelo MLX
        print(f"   📦 Carregando modelo faster-whisper ({model_size})...")
        
        model = WhisperModel(model_size, device="cpu", compute_type="int8")
        
        whisper_lang = self._resolve_whisper_language()
        segments, info = model.transcribe(
            audio_path,
            language=whisper_lang,
            beam_size=beam_size,           # Explora múltiplos caminhos de frase
            best_of=beam_size,             # Escolhe o melhor de N candidatos
            patience=1.0,          # Preferência por transcrições completas
            length_penalty=1.0,    # Evita cortes abruptos
            temperature=0.0,       # Determinístico
            condition_on_previous_text=self._condition_on_previous,
            no_speech_threshold=float(os.getenv("VOMO_NO_SPEECH_THRESHOLD", "0.8")),
            compression_ratio_threshold=float(os.getenv("VOMO_COMPRESSION_THRESHOLD", "2.2")),
            initial_prompt=(initial_prompt or None),
            word_timestamps=True,
        )

        # Formatar output - v2.28: pré-formatação com line breaks
        lines = []
        last_timestamp = None
        clean_enabled = _env_truthy("VOMO_FILTER_ASR_HALLUCINATIONS", default=True)
        last_key = None
        repeat_run = 0
        
        for segment in segments:
            start = segment.start
            text = segment.text.strip()
            
            if not text:
                continue
            
            # Normalização leve de ruído
            text = self._normalize_raw_text(text)
            if clean_enabled:
                if self._asr_is_noise_only(text) or self._asr_looks_like_hallucination(text):
                    continue
                key = self._asr_repeat_key(text)
                if key and key == last_key and len(key) <= 80:
                    repeat_run += 1
                    if repeat_run >= 2:
                        continue
                else:
                    last_key = key
                    repeat_run = 0
            
            # Timestamp a cada 30 segundos
            if self._should_add_timestamp(start, last_timestamp, interval_seconds=self._get_timestamp_interval_for_mode()):
                ts = self._format_timestamp(start)
                lines.append(f"[{ts}] {text}")
                last_timestamp = start
            else:
                lines.append(text)
        
        transcript_result = "\n".join(lines).strip()
        transcript_result = self._strip_leaked_initial_prompt(transcript_result, initial_prompt)
        if _env_truthy("VOMO_ASR_NORMALIZE_TEMAS", default=True):
            try:
                transcript_result, stats = self._normalize_asr_temas_consistency(transcript_result)
                if stats.get("changed", 0) > 0:
                    print(
                        f"{Fore.YELLOW}   🧩 ASR: normalizados {stats['changed']} tema(s) inconsistentes "
                        f"(234→1234: {stats.get('fixed_3_to_4', 0)}, variações: {stats.get('fixed_variants', 0)}){Style.RESET_ALL}"
                    )
            except Exception:
                pass
        
        elapsed = time.time() - start_time
        audio_duration = info.duration if info else 0
        rtf = elapsed / audio_duration if audio_duration > 0 else 0
        
        print(f"{Fore.GREEN}   ✅ Transcrição Beam Search concluída em {elapsed:.1f}s (RTF: {rtf:.2f}x)")
        
        # Salvar cache
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'transcript': transcript_result,
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
                    'backend': 'faster-whisper',
                    'beam_size': beam_size
                }, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erro ao salvar cache: {e}")
        
        return transcript_result
    
    def transcribe_with_segments(self, audio_path, *, beam_size: Optional[int] = None):
        """
        Transcreve e retorna segmentos com timestamps e speaker_label quando diarização estiver disponível.

        v2.33: Suporte a chunking para áudios longos (> AUDIO_MAX_DURATION_SECONDS).
        """
        if not mlx_whisper:
            raise ImportError("mlx_whisper não instalado.")

        # v2.33/v2.34: Verificar se áudio é longo e precisa de chunking
        audio_duration = self._get_audio_duration(audio_path)
        max_duration = self.AUDIO_MAX_DURATION_SECONDS

        # v2.34: Log detalhado
        print(f"{Fore.CYAN}   📏 Duração (segments): {audio_duration/3600:.2f}h (limite: {max_duration/3600:.0f}h){Style.RESET_ALL}")

        if audio_duration > max_duration:
            print(f"{Fore.YELLOW}   ⚠️ Áudio longo detectado ({audio_duration/3600:.1f}h) - ATIVANDO CHUNKING{Style.RESET_ALL}")
            return self._transcribe_with_segments_chunked(audio_path, beam_size=beam_size)
        elif audio_duration == 0:
            print(f"{Fore.RED}   ❌ AVISO: Duração não detectada! Arquivo: {audio_path}{Style.RESET_ALL}")

        initial_prompt = self._get_whisper_initial_prompt_for_asr(high_accuracy=bool(beam_size and beam_size > 1)) or ""
        whisper_lang = self._resolve_whisper_language()
        mlx_kwargs = dict(
            path_or_hf_repo=f"mlx-community/whisper-{self.model_name}",
            **({"language": whisper_lang} if whisper_lang else {}),
            temperature=0.0,
            initial_prompt=(initial_prompt or None),
            word_timestamps=True,
            fp16=True,
            no_speech_threshold=float(os.getenv("VOMO_NO_SPEECH_THRESHOLD", "0.8")),
            logprob_threshold=-1.0,
            compression_ratio_threshold=float(os.getenv("VOMO_COMPRESSION_THRESHOLD", "2.2")),
            condition_on_previous_text=self._condition_on_previous,
            suppress_tokens=[-1],
            verbose=False,
        )
        if beam_size and beam_size > 1:
            mlx_kwargs["beam_size"] = int(beam_size)
            mlx_kwargs["best_of"] = int(beam_size)
        try:
            # v2.31: usar VAD também no fluxo com segmentos
            result = self._transcribe_with_vad(audio_path, mlx_kwargs, skip_silence=True)
        except TypeError:
            if "best_of" in mlx_kwargs:
                mlx_kwargs.pop("best_of", None)
                try:
                    result = self._transcribe_with_vad(audio_path, mlx_kwargs, skip_silence=True)
                except TypeError:
                    mlx_kwargs.pop("beam_size", None)
                    result = self._transcribe_with_vad(audio_path, mlx_kwargs, skip_silence=True)
            else:
                mlx_kwargs.pop("beam_size", None)
                result = self._transcribe_with_vad(audio_path, mlx_kwargs, skip_silence=True)

        diarization_segments = []
        diarization = None
        labeled_segments = None
        if self._diarization_enabled:
            self._ensure_diarization_available_or_raise()
        token = self._get_hf_token()
        if self._diarization_enabled and Pipeline and "torch" in globals() and token:
            try:
                pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-community-1",
                    token=token
                )
                device = "mps" if torch.backends.mps.is_available() else "cpu"
                pipeline.to(torch.device(device))
                diarization = pipeline(audio_path)
                for turn, _, speaker in diarization.itertracks(yield_label=True):
                    speaker_id = speaker.split('_')[-1]
                    diarization_segments.append({
                        "start": float(turn.start),
                        "end": float(turn.end),
                        "speaker_label": f"SPEAKER {int(speaker_id) + 1}"
                    })
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Erro na diarização (segments): {e}")

        if diarization:
            asr_segments, filter_stats = self._filter_asr_segments(result.get("segments", []))
            if filter_stats.get("dropped"):
                reasons = ", ".join(
                    f"{k}={v}" for k, v in sorted((filter_stats.get("reason_counts") or {}).items())
                )
                print(f"{Fore.YELLOW}   🧹 ASR: removidos {filter_stats['dropped']} segmento(s) suspeitos ({reasons})")
            labeled_segments = self._assign_diarization_labels(asr_segments, diarization)
        else:
            asr_segments, filter_stats = self._filter_asr_segments(result.get("segments", []))
            if filter_stats.get("dropped"):
                reasons = ", ".join(
                    f"{k}={v}" for k, v in sorted((filter_stats.get("reason_counts") or {}).items())
                )
                print(f"{Fore.YELLOW}   🧹 ASR: removidos {filter_stats['dropped']} segmento(s) suspeitos ({reasons})")
            labeled_segments = [
                {
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": seg["text"],
                    "speaker_label": "SPEAKER 1",
                    "words": seg.get("words", []),  # Preservar words do Whisper
                }
                for seg in asr_segments
            ]

        # Extrair lista flat de todas as words com timestamps
        all_words = []
        for seg in labeled_segments:
            seg_words = seg.get("words", [])
            speaker = seg.get("speaker_label", "")
            for w in seg_words:
                all_words.append({
                    "word": w.get("word", w.get("text", "")),
                    "start": w.get("start", 0),
                    "end": w.get("end", 0),
                    "speaker": speaker,
                })

        transcript_text = self._segments_to_text(labeled_segments)
        transcript_text = self._strip_leaked_initial_prompt(transcript_text, initial_prompt)
        return {
            "text": transcript_text,
            "segments": labeled_segments,
            "words": all_words,  # Lista flat de words para o player
            "diarization": diarization_segments
        }

    def transcribe_beam_with_segments(self, audio_path):
        """
        Transcrição Beam Search com retorno de segmentos.

        v2.33/v2.34: Suporte a chunking para áudios longos.
        """
        beam_size = self._get_asr_beam_size()

        # v2.33/v2.34: Verificar se áudio é longo - delegar para versão com chunking
        audio_duration = self._get_audio_duration(audio_path)
        max_duration = self.AUDIO_MAX_DURATION_SECONDS

        # v2.34: Log detalhado
        print(f"{Fore.CYAN}   📏 Duração (beam+segments): {audio_duration/3600:.2f}h (limite: {max_duration/3600:.0f}h){Style.RESET_ALL}")

        if audio_duration > max_duration:
            print(f"{Fore.YELLOW}   ⚠️ Áudio longo detectado ({audio_duration/3600:.1f}h) - ATIVANDO CHUNKING{Style.RESET_ALL}")
            return self._transcribe_with_segments_chunked(audio_path, beam_size=beam_size)
        elif audio_duration == 0:
            print(f"{Fore.RED}   ❌ AVISO: Duração não detectada! Arquivo: {audio_path}{Style.RESET_ALL}")

        if not FASTER_WHISPER_AVAILABLE:
            print(f"{Fore.YELLOW}⚠️ faster-whisper não instalado. Tentando Beam Search via MLX ({beam_size})...")
            return self.transcribe_with_segments(audio_path, beam_size=beam_size)

        model_size = "large-v3-turbo"
        model = WhisperModel(model_size, device="cpu", compute_type="int8")
        initial_prompt = self._get_whisper_initial_prompt_for_asr(high_accuracy=True) or ""
        whisper_lang = self._resolve_whisper_language()
        segments, info = model.transcribe(
            audio_path,
            language=whisper_lang,
            beam_size=beam_size,
            best_of=beam_size,
            patience=1.0,
            length_penalty=1.0,
            temperature=0.0,
            condition_on_previous_text=self._condition_on_previous,
            no_speech_threshold=float(os.getenv("VOMO_NO_SPEECH_THRESHOLD", "0.8")),
            compression_ratio_threshold=float(os.getenv("VOMO_COMPRESSION_THRESHOLD", "2.2")),
            initial_prompt=(initial_prompt or None),
            word_timestamps=True,
        )

        asr_segments = [
            {
                "start": float(seg.start),
                "end": float(seg.end),
                "text": seg.text,
                "words": [
                    {"word": w.word, "start": float(w.start), "end": float(w.end)}
                    for w in (seg.words or [])
                ] if hasattr(seg, 'words') and seg.words else [],
            }
            for seg in segments
        ]

        diarization_segments = []
        diarization = None
        if self._diarization_enabled:
            self._ensure_diarization_available_or_raise()
        token = self._get_hf_token()
        if self._diarization_enabled and Pipeline and "torch" in globals() and token:
            try:
                pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-community-1",
                    token=token
                )
                device = "mps" if torch.backends.mps.is_available() else "cpu"
                pipeline.to(torch.device(device))
                diarization = pipeline(audio_path)
                for turn, _, speaker in diarization.itertracks(yield_label=True):
                    speaker_id = speaker.split('_')[-1]
                    diarization_segments.append({
                        "start": float(turn.start),
                        "end": float(turn.end),
                        "speaker_label": f"SPEAKER {int(speaker_id) + 1}"
                    })
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Erro na diarização (beam segments): {e}")

        if diarization:
            labeled_segments = self._assign_diarization_labels(asr_segments, diarization)
        else:
            labeled_segments = [
                {
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": seg["text"],
                    "speaker_label": "SPEAKER 1",
                    "words": seg.get("words", []),  # Preservar words
                }
                for seg in asr_segments
            ]

        # Extrair lista flat de todas as words com timestamps
        all_words = []
        for seg in labeled_segments:
            seg_words = seg.get("words", [])
            speaker = seg.get("speaker_label", "")
            for w in seg_words:
                all_words.append({
                    "word": w.get("word", ""),
                    "start": w.get("start", 0),
                    "end": w.get("end", 0),
                    "speaker": speaker,
                })

        transcript_text = self._segments_to_text(labeled_segments)
        transcript_text = self._strip_leaked_initial_prompt(transcript_text, initial_prompt)
        return {
            "text": transcript_text,
            "segments": labeled_segments,
            "words": all_words,  # Lista flat de words para o player
            "diarization": diarization_segments
        }

    def _segments_to_text(self, segments, timestamp_interval_seconds=None):
        """
        Converte segmentos em texto pré-formatado para melhor chunking.

        v2.29: Pré-formatação com agrupamento por intervalo de tempo:
        - APOSTILA/FIDELIDADE: agrupa segmentos no mesmo intervalo de 60s em parágrafos
        - AUDIENCIA/REUNIAO: 1 segmento = 1 linha (por utterance)
        - Mudança de speaker → quebra dupla + header
        - Normalizações leves de ruído
        """
        if not segments:
            return ""

        lines = []
        current_speaker = None
        last_timestamp = None

        # v2.29: Determinar intervalo de agrupamento baseado no modo
        mode = getattr(self, "_current_mode", "APOSTILA").upper()
        group_by_interval = mode in ("APOSTILA", "FIDELIDADE")
        interval = timestamp_interval_seconds if timestamp_interval_seconds is not None else self._get_timestamp_interval_for_mode()

        # Buffer para acumular texto no mesmo intervalo de timestamp
        current_paragraph = []
        paragraph_timestamp = None

        for segment in segments:
            speaker_label = segment.get("speaker_label", "")
            text = (segment.get("text") or "").strip()
            start = segment.get("start") or 0

            # Pular segmentos vazios
            if not text:
                continue

            # Normalização leve de ruído
            text = self._normalize_raw_text(text)

            # Mudança de speaker → flush buffer e quebra dupla + header
            if speaker_label and speaker_label != current_speaker:
                # Flush buffer atual antes de mudar de speaker
                if current_paragraph:
                    para_text = " ".join(current_paragraph)
                    ts_str = f"[{self._format_timestamp(paragraph_timestamp)}] " if paragraph_timestamp is not None else ""
                    lines.append(f"{ts_str}{para_text}")
                    current_paragraph = []
                    paragraph_timestamp = None

                if lines:  # Não adiciona linha em branco se for o primeiro
                    lines.append("")  # Linha em branco para separar
                lines.append(f"{speaker_label}")
                current_speaker = speaker_label
                last_timestamp = None  # Reset timestamp para novo speaker

            # v2.29: Agrupamento por intervalo para APOSTILA/FIDELIDADE
            if group_by_interval and interval > 0:
                # Verificar se deve iniciar novo parágrafo
                should_new_paragraph = self._should_add_timestamp(start, last_timestamp, interval_seconds=interval)

                if should_new_paragraph:
                    # Flush buffer atual
                    if current_paragraph:
                        para_text = " ".join(current_paragraph)
                        ts_str = f"[{self._format_timestamp(paragraph_timestamp)}] " if paragraph_timestamp is not None else ""
                        lines.append(f"{ts_str}{para_text}")
                        current_paragraph = []

                    # Iniciar novo parágrafo
                    paragraph_timestamp = start
                    last_timestamp = start

                current_paragraph.append(text)
            else:
                # Modo por utterance (AUDIENCIA, REUNIAO, etc.) - 1 segmento = 1 linha
                if self._should_add_timestamp(start, last_timestamp, interval_seconds=timestamp_interval_seconds):
                    timestamp_str = f"[{self._format_timestamp(start)}] "
                    last_timestamp = start
                else:
                    timestamp_str = ""

                lines.append(f"{timestamp_str}{text}")

        # v2.29: Flush buffer final para modo agrupado
        if current_paragraph:
            para_text = " ".join(current_paragraph)
            ts_str = f"[{self._format_timestamp(paragraph_timestamp)}] " if paragraph_timestamp is not None else ""
            lines.append(f"{ts_str}{para_text}")

        transcript = "\n".join(lines).strip()
        if _env_truthy("VOMO_ASR_NORMALIZE_TEMAS", default=True):
            try:
                transcript, stats = self._normalize_asr_temas_consistency(transcript)
                if stats.get("changed", 0) > 0:
                    print(
                        f"{Fore.YELLOW}   🧩 ASR: normalizados {stats['changed']} tema(s) inconsistentes "
                        f"(234→1234: {stats.get('fixed_3_to_4', 0)}, variações: {stats.get('fixed_variants', 0)}){Style.RESET_ALL}"
                    )
            except Exception:
                pass

        return transcript

    def _normalize_asr_temas_consistency(self, text: str):
        """
        Corrige inconsistências comuns de ASR em referências do tipo "Tema N":
        - Perda do dígito inicial em temas de 4 dígitos (ex.: 234 vs 1234), quando a forma canônica já aparece no texto.
        - Variações 4-dígitos com mesmo sufixo (ex.: 1033 vs 1933), quando uma forma é dominante.

        Regra conservadora: só corrige quando há evidência interna (forma canônica presente).
        """
        import re

        if not text:
            return text, {"changed": 0, "fixed_3_to_4": 0, "fixed_variants": 0}

        # Capture occurrences like: "Tema 1.234", "tema 1234", "tema n° 234"
        pattern = re.compile(r"\b[Tt]ema\b\s*(?:n[º°]?\s*)?(\d{1,4})(?:\.(\d{3}))?\b")
        matches = list(pattern.finditer(text))
        if not matches:
            return text, {"changed": 0, "fixed_3_to_4": 0, "fixed_variants": 0}

        def _digits_from_match(m) -> str:
            g1 = m.group(1) or ""
            g2 = m.group(2) or ""
            digits = re.sub(r"\D+", "", f"{g1}{g2}")
            if digits and 2 <= len(digits) <= 4:
                return digits
            return ""

        themes: list[str] = []
        for m in matches:
            d = _digits_from_match(m)
            if d:
                themes.append(d)

        if not themes:
            return text, {"changed": 0, "fixed_3_to_4": 0, "fixed_variants": 0}

        counts: dict[str, int] = {}
        for d in themes:
            counts[d] = counts.get(d, 0) + 1

        fixed_3_to_4 = 0
        fixed_variants = 0
        out = text

        def _digits_to_optional_thousands_regex(digits: str) -> str:
            digits = re.sub(r"\D+", "", digits or "")
            if not digits:
                return ""
            if len(digits) <= 3:
                return re.escape(digits)
            if len(digits) > 6:
                return re.escape(digits)
            prefix = re.escape(digits[:-3])
            suffix = re.escape(digits[-3:])
            return rf"{prefix}\.?{suffix}"

        # Optional explicit overrides for known ASR confusions (highest priority).
        # Format: VOMO_ASR_TEMA_OVERRIDES="1933=1033,234=1234"
        overrides_raw = (os.getenv("VOMO_ASR_TEMA_OVERRIDES") or "").strip()
        if overrides_raw:
            pairs = []
            for part in overrides_raw.split(","):
                if "=" not in part:
                    continue
                src, dst = part.split("=", 1)
                src = re.sub(r"\D+", "", src.strip())
                dst = re.sub(r"\D+", "", dst.strip())
                if src and dst and 2 <= len(src) <= 6 and 2 <= len(dst) <= 6:
                    pairs.append((src, dst))
            for src, dst in pairs:
                src_re = _digits_to_optional_thousands_regex(src)
                if not src_re:
                    continue
                before = out
                out = re.sub(
                    rf"\b([Tt]ema)\s*(?:n[º°]?\s*)?{src_re}\b",
                    rf"\1 {dst}",
                    out,
                )
                if out != before:
                    fixed_variants += 1

        # Rule A: 3-digit -> 4-digit when the 4-digit variant (prefixed with '1') exists in the same transcript.
        for d in list(counts.keys()):
            if len(d) != 3:
                continue
            target = f"1{d}"
            if target not in counts:
                continue
            before = out
            out = re.sub(
                rf"\b([Tt]ema)\s*(?:n[º°]?\s*)?{re.escape(d)}\b",
                rf"\1 {target}",
                out,
            )
            if out != before:
                fixed_3_to_4 += 1

        # Re-scan after replacements
        matches = list(pattern.finditer(out))
        counts = {}
        for m in matches:
            d = _digits_from_match(m)
            if d:
                counts[d] = counts.get(d, 0) + 1

        # Rule B: unify 4-digit variants that share the same last 3 digits, when one is clearly dominant.
        by_suffix: dict[str, list[str]] = {}
        for d in counts:
            if len(d) == 4:
                by_suffix.setdefault(d[-3:], []).append(d)
        for suffix, variants in by_suffix.items():
            if len(variants) <= 1:
                continue
            ones = [v for v in variants if v.startswith("1")]
            if ones:
                preferred = max(ones, key=lambda v: counts.get(v, 0))
            else:
                preferred = max(variants, key=lambda v: counts.get(v, 0))

            pref_count = counts.get(preferred, 0)
            for v in variants:
                if v == preferred:
                    continue
                v_count = counts.get(v, 0)
                # Only normalize if the preferred is at least 2x more common or the variant is rare (<=1 hit).
                if pref_count >= (2 * v_count) or v_count <= 1:
                    before = out
                    out = re.sub(
                        rf"\b([Tt]ema)\s*(?:n[º°]?\s*)?{re.escape(v)}\b",
                        rf"\1 {preferred}",
                        out,
                    )
                    if out != before:
                        fixed_variants += 1

        changed = fixed_3_to_4 + fixed_variants if out != text else 0
        return out, {"changed": changed, "fixed_3_to_4": fixed_3_to_4, "fixed_variants": fixed_variants}
    
    def _normalize_raw_text(self, text):
        """
        Normalizações leves e seguras no texto raw (sem reescrever frases).
        
        v2.28: Limpeza determinística:
        - Remove whitespace excessivo
        - Converte ruídos em tokens padrão
        - Preserva todo conteúdo semântico
        """
        import re
        
        # Remove whitespace excessivo
        text = re.sub(r'\s+', ' ', text)
        
        # Normaliza traços e reticências
        text = re.sub(r'—|–', '-', text)  # Em-dash/en-dash → hífen
        text = re.sub(r'…', '...', text)  # Ellipsis → três pontos
        
        # Tokens padrão para ruídos (se detectados)
        noise_patterns = [
            (r'\[inaudível\]|\(inaudível\)|\[inaudible\]', '[inaudível]'),
            (r'\[risos\]|\(risos\)|\[laughter\]', '[risos]'),
            (r'\[pausa\]|\(pausa\)|\[pause\]', '[pausa]'),
            (r'\[música\]|\(música\)|\[music\]', '[música]'),
            (r'\[aplausos\]|\(aplausos\)|\[applause\]', '[aplausos]'),
        ]
        for pattern, replacement in noise_patterns:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text.strip()

    def _asr_repeat_key(self, text: str) -> str:
        """Normalização agressiva apenas para detecção de repetição (não para saída)."""
        import re

        if not text:
            return ""
        key = re.sub(r"\s+", " ", str(text)).strip().lower()
        key = re.sub(r"[^\wÀ-ÖØ-öø-ÿ0-9 ]+", "", key)
        key = re.sub(r"\s+", " ", key).strip()
        return key

    def _asr_is_noise_only(self, text: str) -> bool:
        """True quando o segmento é só marcador de ruído (ex.: [música])."""
        import re

        if not text:
            return True
        normalized = re.sub(r"\s+", " ", str(text)).strip()
        if not normalized:
            return True
        # Se não houver nenhum caractere alfanumérico, trata como ruído (ex.: ".", "..." ou "?!")
        if not re.search(r"[\wÀ-ÖØ-öø-ÿ0-9]", normalized):
            return True
        # Depois de _normalize_raw_text, os ruídos ficam padronizados como [xxx]
        noise = r"(?:inaudível|risos|pausa|música|aplausos)"
        return bool(re.fullmatch(rf"(?:\[(?:{noise})\]\s*)+", normalized, flags=re.IGNORECASE))

    def _get_asr_beam_size(self) -> int:
        """
        Beam size padrão para o modo "Alta Precisão" (Beam Search).

        - UI: on/off (alta precisão).
        - Ajuste avançado via env:
          - VOMO_ASR_BEAM_SIZE (ou VOMO_BEAM_SIZE) -> inteiro
        """
        raw = (os.getenv("VOMO_ASR_BEAM_SIZE") or os.getenv("VOMO_BEAM_SIZE") or "").strip()
        try:
            value = int(raw) if raw else 5
        except Exception:
            value = 5
        # Beam search real começa em 2; limite superior conservador para evitar explosão de custo.
        if value < 2:
            value = 2
        if value > 10:
            value = 10
        return value

    def _get_whisper_initial_prompt_for_asr(self, *, high_accuracy: bool) -> Optional[str]:
        """
        Decide o `initial_prompt` do Whisper para a etapa de ASR.

        Notas importantes:
        - Prompt pode melhorar vocabulário, mas pode "vazar" como texto transcrito.
        - Por padrão, evitamos usar prompt no modo normal (rápido).
        - Em alta precisão (Beam Search), usamos o prompt por modo como padrão (best-effort),
          pois o usuário já optou por mais qualidade e o risco é mitigado por stripping.

        Controles:
        - VOMO_WHISPER_INITIAL_PROMPT="..." (sempre tem prioridade)
        - VOMO_WHISPER_USE_MODE_PROMPT=1 (força uso do prompt por modo em TODOS os modos)
        - Se VOMO_WHISPER_USE_MODE_PROMPT não estiver definido:
          - high_accuracy=True -> usa prompt por modo
          - high_accuracy=False -> não usa prompt (default)
        """
        explicit = (os.getenv("VOMO_WHISPER_INITIAL_PROMPT") or "").strip()
        if explicit:
            return explicit

        # Se o usuário não definiu explicitamente, decidimos pelo modo:
        # - Beam (high_accuracy): ON
        # - Normal: OFF
        use_mode_prompt = _env_truthy("VOMO_WHISPER_USE_MODE_PROMPT", default=None)
        if use_mode_prompt is None:
            use_mode_prompt = bool(high_accuracy)

        if not use_mode_prompt:
            return None

        mode_key = getattr(self, "_current_mode", "FIDELIDADE")
        if isinstance(mode_key, str):
            mode_key = mode_key.strip().upper() or "FIDELIDADE"
        else:
            mode_key = "FIDELIDADE"

        lang = getattr(self, "_current_language", "pt") or "pt"
        # Tenta prompt i18n primeiro, depois fallback para pt
        if lang != "pt" and lang != "auto":
            i18n_prompt = self.INITIAL_PROMPTS_I18N.get((mode_key, lang))
            if i18n_prompt:
                return i18n_prompt
        return self.INITIAL_PROMPTS.get(mode_key, self.INITIAL_PROMPTS["FIDELIDADE"])

    def _get_whisper_initial_prompt(self) -> Optional[str]:
        """
        Compat: mantém a API antiga (sem saber se é high_accuracy).
        Por padrão, segue o comportamento do modo normal (não-beam).
        """
        return self._get_whisper_initial_prompt_for_asr(high_accuracy=False)

    def _strip_leaked_initial_prompt(self, text: str, initial_prompt: str) -> str:
        """
        Best-effort: remove `initial_prompt` caso ele tenha vazado como primeira linha da transcrição.

        Estratégia:
        - Tokeniza primeira linha e o prompt, compara sobreposição.
        - Remove SOMENTE quando há alta similaridade (muito conservador).
        """
        import re

        if not text:
            return text
        prompt = (initial_prompt or "").strip()
        if not prompt:
            return text

        first_line, rest = (text.split("\n", 1) + [""])[:2]

        def _tokens(value: str) -> list[str]:
            return re.findall(r"[\wÀ-ÖØ-öø-ÿ0-9]+", (value or "").lower())

        prompt_tokens = _tokens(prompt)
        if len(prompt_tokens) < 8:
            return text
        line_tokens = _tokens(first_line)
        if not line_tokens:
            return text

        prompt_set = set(prompt_tokens)
        overlap = sum(1 for t in prompt_tokens if t in set(line_tokens))

        # Regras conservadoras: muita sobreposição e tamanho similar.
        overlap_ratio = overlap / max(1, len(prompt_tokens))
        size_ok = len(line_tokens) <= (len(prompt_tokens) + 6)
        if overlap_ratio >= 0.9 and size_ok:
            return (rest or "").lstrip()

        return text

    def _asr_has_repeated_ngram_run(self, tokens, *, max_ngram: int = 6) -> bool:
        """Detecta repetições consecutivas extremas (alucinação típica do Whisper)."""
        token_count = len(tokens)
        if token_count < 12:
            return False

        max_ngram = min(max_ngram, token_count // 2)
        for n in range(1, max_ngram + 1):
            min_repeats = 8 if n == 1 else 4  # muito conservador para evitar falsos positivos
            if token_count < n * min_repeats:
                continue

            # Procura por runs consecutivos em qualquer offset.
            for i in range(0, token_count - n * min_repeats + 1):
                phrase = tokens[i : i + n]
                repeats = 1
                while i + (repeats + 1) * n <= token_count and tokens[i + repeats * n : i + (repeats + 1) * n] == phrase:
                    repeats += 1
                if repeats >= min_repeats:
                    return True

        return False

    def _asr_looks_like_hallucination(self, text: str) -> bool:
        """
        Heurísticas conservadoras para filtrar segmentos obviamente quebrados:
        - Repetição extrema (palavra/frase em loop)
        - Sequências numéricas repetidas (ex.: "50 50 50 ...")
        """
        import re

        if not text:
            return True

        normalized = re.sub(r"\s+", " ", str(text)).strip()
        if not normalized:
            return True

        # Tokens "palavra-like" (mantém números; remove pontuação)
        tokens = re.findall(r"[\wÀ-ÖØ-öø-ÿ0-9]+", normalized.lower())
        if len(tokens) < 12:
            return False

        # 1) Repetição consecutiva extrema de n-gramas curtos
        if self._asr_has_repeated_ngram_run(tokens):
            return True

        # 2) Baixa diversidade lexical em sequência longa (ex.: mesmo slogan repetido)
        if len(tokens) >= 25:
            unique_ratio = len(set(tokens)) / max(1, len(tokens))
            if unique_ratio < 0.25:
                return True

        # 3) Segmento quase só números e com pouca variedade (ex.: token IDs ou contagem)
        numeric_tokens = [t for t in tokens if t.isdigit()]
        if numeric_tokens and len(numeric_tokens) / len(tokens) > 0.85:
            if len(numeric_tokens) >= 12 and len(set(numeric_tokens)) <= 3:
                return True
            if len(numeric_tokens) >= 30 and len(set(numeric_tokens)) <= 8:
                return True

        return False

    def _filter_asr_segments(self, segments):
        """
        Remove segmentos claramente inúteis (ruídos/loops) antes de formatar ou diarizar.
        Mantém timestamps originais.
        """
        if not segments:
            return [], {"dropped": 0, "reason_counts": {}}

        clean_enabled = _env_truthy("VOMO_FILTER_ASR_HALLUCINATIONS", default=True)
        if not clean_enabled:
            return segments, {"dropped": 0, "reason_counts": {}}

        dropped = 0
        reason_counts = {}
        cleaned = []

        last_key = None
        repeat_run = 0

        for seg in segments:
            raw_text = (seg.get("text") or "").strip()
            if not raw_text:
                dropped += 1
                reason_counts["empty"] = reason_counts.get("empty", 0) + 1
                continue

            text = self._normalize_raw_text(raw_text)
            if self._asr_is_noise_only(text):
                dropped += 1
                reason_counts["noise_only"] = reason_counts.get("noise_only", 0) + 1
                continue

            key = self._asr_repeat_key(text)
            if key and key == last_key and len(key) <= 80:
                repeat_run += 1
                # Permite no máximo 2 repetições consecutivas de segmentos curtos idênticos.
                if repeat_run >= 2:
                    dropped += 1
                    reason_counts["repeat_loop"] = reason_counts.get("repeat_loop", 0) + 1
                    continue
            else:
                last_key = key
                repeat_run = 0

            if self._asr_looks_like_hallucination(text):
                dropped += 1
                reason_counts["hallucination"] = reason_counts.get("hallucination", 0) + 1
                continue

            new_seg = dict(seg)
            new_seg["text"] = text
            cleaned.append(new_seg)

        return cleaned, {"dropped": dropped, "reason_counts": reason_counts}

    def _assign_diarization_labels(self, segments, diarization_output):
        try:
            from intervaltree import IntervalTree
        except ImportError:
            return self._assign_diarization_labels_fallback(segments, diarization_output)

        tree = IntervalTree()
        for turn, _, speaker in diarization_output.itertracks(yield_label=True):
            tree[turn.start:turn.end] = speaker

        labeled_segments = []
        for segment in segments:
            start, end = segment['start'], segment['end']
            overlaps = tree[start:end]
            if overlaps:
                best_overlap = max(
                    overlaps,
                    key=lambda interval: min(end, interval.end) - max(start, interval.begin)
                )
                speaker_id = best_overlap.data.split('_')[-1]
                best_speaker = f"SPEAKER {int(speaker_id) + 1}"
            else:
                best_speaker = "SPEAKER 0"
            labeled_segments.append({
                "start": float(start),
                "end": float(end),
                "text": segment.get("text", ""),
                "speaker_label": best_speaker,
            })
        return labeled_segments

    def _assign_diarization_labels_fallback(self, segments, diarization_output):
        diarization_segments = [(t.start, t.end, s) for t, _, s in diarization_output.itertracks(yield_label=True)]
        labeled_segments = []
        for segment in segments:
            start, end = segment['start'], segment['end']
            best_speaker = "SPEAKER 0"
            max_overlap = 0
            for d_start, d_end, d_speaker in diarization_segments:
                overlap = max(0, min(end, d_end) - max(start, d_start))
                if overlap > max_overlap:
                    max_overlap = overlap
                    best_speaker = f"SPEAKER {int(d_speaker.split('_')[-1]) + 1}"
            labeled_segments.append({
                "start": float(start),
                "end": float(end),
                "text": segment.get("text", ""),
                "speaker_label": best_speaker,
            })
        return labeled_segments

    def _format_timestamp(self, seconds):
        m, s = divmod(seconds, 60)
        h, m = divmod(m, 60)
        return f"{int(h):02d}:{int(m):02d}:{int(s):02d}" if h > 0 else f"{int(m):02d}:{int(s):02d}"
    
    def _get_timestamp_interval_for_mode(self) -> int:
        """
        Retorna o intervalo de timestamps baseado no modo atual.

        - APOSTILA/FIDELIDADE: 60s (aulas longas, menos interrupções)
        - AUDIENCIA/REUNIAO/LEGENDA: 0 (por utterance, cada segmento tem timestamp)

        Returns:
            int: Intervalo em segundos (0 = por utterance)
        """
        mode = getattr(self, "_current_mode", "APOSTILA").upper()
        if mode in ("APOSTILA", "FIDELIDADE"):
            return 60
        # AUDIENCIA, REUNIAO, LEGENDA, DEPOIMENTO → por utterance
        return 0

    def _should_add_timestamp(self, current_seconds, last_timestamp_seconds, interval_minutes=None, interval_seconds=None):
        """
        Determina se um timestamp deve ser adicionado baseado no intervalo configurado.

        Args:
            current_seconds: Tempo atual em segundos
            last_timestamp_seconds: Último timestamp inserido (None se for o primeiro)
            interval_minutes: Intervalo em minutos entre timestamps (padrão: 20 se nenhum especificado)
            interval_seconds: Intervalo em segundos (tem precedência sobre interval_minutes)
                              Se None e interval_minutes também None, usa intervalo baseado no modo.

        Returns:
            bool: True se deve adicionar timestamp
        """
        if last_timestamp_seconds is None:
            return True  # Sempre adiciona o primeiro timestamp

        # interval_seconds tem precedência
        if interval_seconds is not None:
            target_interval = interval_seconds
        elif interval_minutes is not None:
            target_interval = interval_minutes * 60
        else:
            # Usar intervalo baseado no modo atual
            target_interval = self._get_timestamp_interval_for_mode()
            if target_interval == 0:
                return True  # Por utterance: sempre adiciona

        return (current_seconds - last_timestamp_seconds) >= target_interval

    def _align_diarization(self, segments, diarization_output):
        """
        Alinha segmentos com diarização (Versão Otimizada com IntervalTree)
        
        v2.28: Saída pré-formatada com line breaks e timestamps frequentes
        """
        try:
            from intervaltree import IntervalTree
        except ImportError:
            print("⚠️ intervaltree não instalado, usando fallback O(n)")
            return self._align_diarization_fallback(segments, diarization_output)
        
        lines = []
        current_speaker = None
        current_block = []
        last_timestamp = None
        
        # Pré-computar spatial index
        tree = IntervalTree()
        for turn, _, speaker in diarization_output.itertracks(yield_label=True):
            tree[turn.start:turn.end] = speaker
        
        for segment in segments:
            start, end = segment['start'], segment['end']
            text = segment.get('text', '').strip()
            
            if not text:
                continue
            
            # Normalização leve
            text = self._normalize_raw_text(text)
            
            # Busca O(log n) para speaker
            overlaps = tree[start:end]
            
            if overlaps:
                best_overlap = max(
                    overlaps,
                    key=lambda interval: min(end, interval.end) - max(start, interval.begin)
                )
                speaker_id = best_overlap.data.split('_')[-1]
                best_speaker = f"SPEAKER {int(speaker_id) + 1}"
            else:
                best_speaker = "SPEAKER 1"
            
            # Mudança de speaker
            if best_speaker != current_speaker:
                # Flush previous block
                if current_block:
                    lines.append(" ".join(current_block))
                    current_block = []
                
                if lines:
                    lines.append("")  # Linha em branco extra
                
                lines.append(f"{best_speaker}")
                current_speaker = best_speaker
                last_timestamp = None  # Reset timestamp logic for new speaker? Or keep continuous? 
                # Keeping continuous is usually better for "every 60s" regardless of speaker, 
                # but resetting ensures first line of speaker has timestamp if we want.
                # User asked "timestamps a cada 60 segundos".
                # Let's keep logic simple: Check timestamp interval.
                # If we reset last_timestamp, we force a timestamp at speaker start.
                last_timestamp = None 
            
            # Timestamp a cada 60 segundos
            if self._should_add_timestamp(start, last_timestamp, interval_seconds=self._get_timestamp_interval_for_mode()):
                if current_block:
                    lines.append(" ".join(current_block))
                    current_block = []
                
                ts = self._format_timestamp(start)
                current_block.append(f"[{ts}] {text}")
                last_timestamp = start
            else:
                current_block.append(text)
        
        # Flush final block
        if current_block:
            lines.append(" ".join(current_block))
        
        return "\n\n".join(lines).strip()
    
    def _align_diarization_fallback(self, segments, diarization_output):
        """
        Fallback O(n) caso intervaltree não esteja disponível.
        
        v2.28: Saída pré-formatada condensada com line breaks e timestamps frequentes
        """
        lines = []
        current_speaker = None
        current_block = []
        last_timestamp = None
        diarization_segments = [(t.start, t.end, s) for t, _, s in diarization_output.itertracks(yield_label=True)]

        for segment in segments:
            start, end = segment['start'], segment['end']
            text = segment.get('text', '').strip()
            
            if not text:
                continue
            
            # Normalização leve
            text = self._normalize_raw_text(text)
            
            best_speaker = "SPEAKER 1"
            max_overlap = 0
            
            # Busca O(n)
            for d_start, d_end, d_speaker in diarization_segments:
                overlap = max(0, min(end, d_end) - max(start, d_start))
                if overlap > max_overlap:
                    max_overlap = overlap
                    speaker_id = d_speaker.split('_')[-1]
                    best_speaker = f"SPEAKER {int(speaker_id) + 1}"
            
            # Mudança de speaker
            if best_speaker != current_speaker:
                # Flush previous block
                if current_block:
                    lines.append(" ".join(current_block))
                    current_block = []
                
                if lines:
                    lines.append("")  # Linha em branco extra
                
                lines.append(f"{best_speaker}")
                current_speaker = best_speaker
                last_timestamp = None
            
            # Timestamp a cada 60 segundos
            if self._should_add_timestamp(start, last_timestamp, interval_seconds=self._get_timestamp_interval_for_mode()):
                if current_block:
                    lines.append(" ".join(current_block))
                    current_block = []
                    
                ts = self._format_timestamp(start)
                current_block.append(f"[{ts}] {text}")
                last_timestamp = start
            else:
                current_block.append(text)
        
        # Flush final block
        if current_block:
            lines.append(" ".join(current_block))

        return "\n\n".join(lines).strip()

    def _segment_raw_transcription(self, raw_text):
        lines = raw_text.split('\n')
        speaker_pattern = re.compile(r'^SPEAKER \d+$')
        segments = []
        current_speaker = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if speaker_pattern.match(line):
                if current_speaker:
                    segments.append({'speaker': current_speaker, 'content': "\n".join(current_content)})
                current_speaker = line
                current_content = []
            else:
                if current_speaker: current_content.append(line)
        
        if current_speaker:
            segments.append({'speaker': current_speaker, 'content': "\n".join(current_content)})
            
        # Merge simples
        final_segments = []
        if segments:
            current = segments[0]
            for next_seg in segments[1:]:
                if next_seg['speaker'] == current['speaker']:
                    current['content'] += "\n" + next_seg['content']
                else:
                    final_segments.append(current)
                    current = next_seg
            final_segments.append(current)
            
        return final_segments

    def _smart_chunk_with_overlap(self, text, max_size=None, overlap=None):
        """
        v2.28: Chunking inteligente com overlap adaptativo para tabelas.

        Melhorias:
        - Overlap 30% maior quando chunk contém tabela
        - Nunca corta no meio de tabela
        - Prioriza corte após pares de tabelas (📋 + 🎯)
        """
        max_size = max_size or self.MAX_CHUNK_SIZE
        base_overlap = overlap or self.CHUNK_OVERLAP
        if len(text) <= max_size: return [text]

        chunks = []
        start = 0

        def _is_table_line(line: str) -> bool:
            stripped = line.strip()
            return bool(stripped) and stripped.startswith('|') and '|' in stripped

        def _prev_next_nonempty_lines_around(pos: int, window: int = 5000) -> tuple:
            s_start = max(0, pos - window)
            s_end = min(len(text), pos + window)
            s = text[s_start:s_end]
            p = pos - s_start
            before_lines = s[:p].splitlines()
            after_lines = s[p:].splitlines()
            prev_line = next((l for l in reversed(before_lines) if l.strip()), "")
            next_line = next((l for l in after_lines if l.strip()), "")
            return prev_line, next_line

        def _pos_inside_table_line(pos: int) -> bool:
            if pos <= 0 or pos >= len(text):
                return False
            # `end` (slice stop) é seguro se estiver logo após um '\n'
            if text[pos - 1] == '\n':
                return False
            line_start = text.rfind('\n', 0, pos) + 1
            line_end = text.find('\n', pos)
            if line_end == -1:
                line_end = len(text)
            return _is_table_line(text[line_start:line_end])

        def _table_block_bounds_around(pos: int, window: int = 15000) -> Optional[tuple]:
            s_start = max(0, pos - window)
            s_end = min(len(text), pos + window)
            s = text[s_start:s_end]
            p = pos - s_start

            lines = s.splitlines(keepends=True)
            if not lines:
                return None

            # Encontrar o índice da linha que contém `p` (ou a anterior se `p` cair no separador)
            cumulative = 0
            idx = 0
            for idx, ln in enumerate(lines):
                nxt = cumulative + len(ln)
                if p < nxt:
                    break
                cumulative = nxt
            else:
                idx = len(lines) - 1

            if not _is_table_line(lines[idx]) and idx > 0 and _is_table_line(lines[idx - 1]):
                idx -= 1
            if not _is_table_line(lines[idx]):
                return None

            start_idx = idx
            while start_idx > 0 and _is_table_line(lines[start_idx - 1]):
                start_idx -= 1
            end_idx = idx
            while end_idx + 1 < len(lines) and _is_table_line(lines[end_idx + 1]):
                end_idx += 1

            start_off = sum(len(ln) for ln in lines[:start_idx])
            end_off = sum(len(ln) for ln in lines[:end_idx + 1])
            return (s_start + start_off, s_start + end_off)

        while start < len(text):
            end = start + max_size
            chunk_text = text[start:end] if end <= len(text) else text[start:]

            # v2.28: Detectar se chunk contém tabela para overlap maior
            contem_tabela = '|' in chunk_text and re.search(r'^\s*\|', chunk_text, re.MULTILINE)
            current_overlap = int(base_overlap * 1.3) if contem_tabela else base_overlap

            if end < len(text):
                # Zona de busca para ponto de corte
                search_start = max(0, end - 3000)
                search_zone = text[search_start:end]
                best_break = -1

                # v2.28: Prioridade 1 - Após tabela de pegadinhas completa
                match_pegadinha = re.search(r'\n(?=####?\s*🎯.*\n)', search_zone)
                if match_pegadinha:
                    # Encontrar fim da tabela após o heading
                    pos_heading = match_pegadinha.end()
                    remaining = search_zone[pos_heading:]
                    lines = remaining.split('\n')
                    pos_after_table = pos_heading
                    in_table = False
                    for i, line in enumerate(lines):
                        if line.strip().startswith('|'):
                            in_table = True
                        elif in_table and not line.strip().startswith('|'):
                            # Fim da tabela
                            pos_after_table = pos_heading + sum(len(l)+1 for l in lines[:i])
                            break
                    if pos_after_table > pos_heading:
                        best_break = pos_after_table

                # v2.28: Prioridade 2 - Antes de novo heading ## (bloco temático)
                if best_break == -1:
                    match_heading = list(re.finditer(r'\n(?=##\s+\d)', search_zone))
                    if match_heading:
                        best_break = match_heading[-1].end()  # Último heading encontrado

                # v2.28: Prioridade 3 - Parágrafo duplo (evitando meio de tabela)
                if best_break == -1:
                    # Verificar se estamos no meio de uma tabela
                    last_newlines = list(re.finditer(r'\n\n', search_zone))
                    for match in reversed(last_newlines):
                        pos = match.start()
                        # Verificar se próxima linha não é tabela
                        next_char_pos = match.end()
                        if next_char_pos < len(search_zone):
                            next_line_start = search_zone[next_char_pos:next_char_pos+50]
                            if not next_line_start.strip().startswith('|'):
                                best_break = pos + 2
                                break

                # Fallback: qualquer \n\n
                if best_break == -1:
                    last_break = search_zone.rfind('\n\n')
                    if last_break != -1:
                        best_break = last_break

                if best_break != -1:
                    end = search_start + best_break

                # v2.28: Nunca cortar no meio de uma tabela (separação por linha)
                prev_line, next_line = _prev_next_nonempty_lines_around(end)
                if (_is_table_line(prev_line) and _is_table_line(next_line)) or _pos_inside_table_line(end):
                    bounds = _table_block_bounds_around(end)
                    if bounds:
                        table_start, table_end = bounds
                        min_chunk_size = max(800, int(max_size * 0.2))
                        # Preferir cortar ANTES da tabela; se ficar pequeno demais, cortar APÓS a tabela
                        candidate = table_start
                        if candidate <= start + min_chunk_size and table_end > start + min_chunk_size:
                            candidate = table_end
                        if candidate > start:
                            end = candidate

            # Garantir progresso (evitar loop infinito se `end` voltar demais)
            if end <= start:
                end = min(start + max_size, len(text))
                if end <= start:
                    break

            chunks.append(text[start:end].strip())
            if end >= len(text): break
            next_start = end - current_overlap
            if next_start <= start:
                next_start = end
            start = next_start

        return chunks

    def _merge_chunks_deduplicated(self, chunks):
        if not chunks: return ""
        if len(chunks) == 1: return chunks[0]
        
        merged = chunks[0]
        for i in range(1, len(chunks)):
            current = chunks[i]
            tail = merged[-2000:]
            head = current[:2000]
            matcher = difflib.SequenceMatcher(None, tail, head)
            match = matcher.find_longest_match(0, len(tail), 0, len(head))
            
            if match.size > 200:
                merged += "\n\n" + current[match.b + match.size:]
            else:
                merged += "\n\n" + current
        return merged

    def _get_chunk_cache(self, chunk_text, prompt_hash):
        content_hash = hashlib.sha256(f"{chunk_text}{prompt_hash}".encode()).hexdigest()
        cache_path = self.cache_dir / f"{content_hash}.json"
        if cache_path.exists():
            try:
                with open(cache_path, 'r') as f:
                    return json.load(f)['result']
            except:
                return None
        return None

    def _save_chunk_cache(self, chunk_text, prompt_hash, result):
        content_hash = hashlib.sha256(f"{chunk_text}{prompt_hash}".encode()).hexdigest()
        cache_path = self.cache_dir / f"{content_hash}.json"
        try:
            with open(cache_path, 'w') as f:
                json.dump({'result': result}, f)
        except:
            pass

    def _detect_open_table_state(self, text: str) -> dict:
        """
        v2.27: Detecta se o texto termina com uma tabela/quadro aberto mas não concluído.
        
        Casos detectados:
        1. Título de tabela principal (#### 📋 [rótulo contextual]) sem tabela depois
        2. Tabela iniciada mas incompleta (menos linhas que o esperado)
        
        Returns:
            dict com 'needs_table_continuation' e 'context_hint' se aberto, {} se fechado
        """
        if not text or len(text) < 100:
            return {}
        
        lines = text.strip().splitlines()
        last_50 = lines[-50:] if len(lines) > 50 else lines
        
        # Caso 1: Título de tabela principal (📋) sem tabela
        for i, line in enumerate(last_50):
            if re.match(r'^#{3,5}\s*📋', line):
                # Há um título de quadro, verifica se tabela foi gerada depois
                remaining = last_50[i+1:]
                has_table = any('|' in l and l.strip().startswith('|') for l in remaining)
                if not has_table:
                    section_title = line.strip()
                    return {
                        "needs_table_continuation": True,
                        "open_section_title": section_title,
                        "context_hint": f"\n\n⚠️ **CONTINUAÇÃO OBRIGATÓRIA**: O chunk anterior terminou com o título '{section_title}' mas SEM a tabela correspondente. Você DEVE gerar a tabela Markdown desse bloco ANTES de qualquer novo conteúdo."
                    }
        
        # Caso 2: Última linha é tabela (pode precisar continuação se poucos dados)
        # Apenas logamos, não adicionamos instrução explícita neste caso
        
        # Caso 3: Célula amputada (corte no meio da frase dentro da tabela)
        # Ex: "| Item 1 | O princípio da legalidade define que" (sem pipe final)
        if lines:
            last_line = lines[-1].strip()
            if last_line.startswith('|') and not last_line.endswith('|'):
                return {
                    "needs_table_continuation": True,
                    "open_section_title": "Tabela cortada no meio",
                    "context_hint": f"\n\n⚠️ CONTINUE a tabela de onde parou: '{last_line[-60:]}...' → complete a frase, feche com '|', continue normalmente."
                }
        
        return {}

    async def _format_chunk_async(self, chunk_text, idx, prompt=None, total=1, context="", depth=0, **kwargs):
        """Wrapper de compatibilidade para process_chunk_async que aceita prompt (ignorado)"""
        return await self.process_chunk_async(
            chunk_text,
            idx=idx,
            total=total,
            previous_context=context,
            depth=depth,
            **kwargs
        )

    @retry(
        retry=retry_if_exception_type(Exception), 
        stop=stop_after_attempt(3), 
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    async def process_chunk_async(
        self,
        chunk_text,
        idx=0,
        total=1,
        previous_context="",
        depth=0,
        global_structure=None,
        overlap_text="",
        cached_content=None,
        max_output_tokens_override=None,
        disable_cache=False,
        table_retry=False,
        trunc_retry=False,
    ):
        """Processa um chunk de forma assíncrona com retry, cache e CHUNKING ADAPTATIVO (v2.10 Ported)"""
        
        # Calculate prompt hash for local caching (apenas se não usar context cache)
        prompt_content = f"{chunk_text}_{previous_context}_{overlap_text}_{global_structure}"
        prompt_hash = hashlib.sha256(prompt_content.encode()).hexdigest()
        
        # Check local cache (only if depth 0 to avoid fragment caching issues)
        # Se estiver usando Context Caching, ignoramos o cache local para garantir uso do contexto global
        if not cached_content and not disable_cache:
            cached = self._get_chunk_cache(chunk_text, prompt_hash) if depth == 0 else None
            if cached:
                metrics.record_cache_hit()
                return cached
        
        # Constrói o contexto e prompt
        contexto_estilo = f"Últimos parágrafos formatados:\n{previous_context}" if previous_context else ""
        contexto_raw = f"OVERLAP RAW (somente contexto, NÃO INCLUIR na resposta):\n{overlap_text}" if overlap_text else ""
        if not contexto_estilo and not contexto_raw:
            contexto_estilo = "Inicio do documento."
        
        # SE USAR CONTEXT CACHING: System Prompt já está no cache
        # SE NÃO USAR: System Prompt precisa ir no contents
        
        # Se não tiver cache, montamos o system prompt completo
        system_prompt = self.prompt_apostila
            
        # Adiciona estrutura global se disponível e NÃO estiver no cache (se estiver no cache, já foi incluída na criação)
        if global_structure and not cached_content:
            system_prompt += f"\n\n## ESTRUTURA GLOBAL (GUIA):\n{global_structure}"
            
        # Adiciona contexto anterior (estilo + overlap RAW)
        secao_contexto = ""
        if previous_context or overlap_text:
            blocks = []
            if contexto_estilo:
                blocks.append(contexto_estilo)
            if contexto_raw:
                blocks.append(contexto_raw)
            contexto_bloco = "\n\n".join(blocks).strip()
            secao_contexto = f"""
🔒 CONTEXTO ANTERIOR (SOMENTE REFERÊNCIA DE ESTILO)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{contexto_bloco}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ ATENÇÃO: O bloco acima JÁ FOI FORMATADO anteriormente.
- NÃO formate novamente esse conteúdo
- NÃO inclua esse conteúdo na sua resposta
- Use APENAS como referência de estilo de escrita e continuidade
- Se houver OVERLAP RAW, use apenas para continuidade; não copie nem reformate
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 NOVO TEXTO PARA FORMATAR (comece aqui):
"""

        user_content = f"""{secao_contexto}
<texto_para_formatar>
{chunk_text}
</texto_para_formatar>

        **INSTRUÇÕES FINAIS**:
        - Esta é a parte {idx} de {total if total else '?'} (Profundidade {depth})
        - Formate APENAS o texto entre <texto_para_formatar>
        - Se houver contexto acima, NÃO o reprocesse
        - Retorne APENAS o Markdown formatado do NOVO texto
        - NÃO insira marcadores artificiais de continuação (ex.: `[continua]`, `[continuação]`, `(continua)`)
        """

        def _has_incomplete_table(text: str) -> bool:
            if not text:
                return False
            lines = [ln.rstrip() for ln in text.splitlines() if ln.strip()]
            if not lines:
                return False
            if "|" not in lines[-1]:
                return False
            end = len(lines) - 1
            start = end
            while start >= 0 and "|" in lines[start]:
                start -= 1
            start += 1
            block = lines[start:end + 1]
            if len(block) < 2:
                return False
            header = block[0]
            separator = block[1]
            if not re.search(r'-{3,}', separator):
                return False
            data_rows = block[2:]
            if not data_rows:
                return True
            header_pipes = header.count("|")
            last_pipes = data_rows[-1].count("|")
            return header_pipes >= 2 and last_pipes < header_pipes

        def _has_missing_table(text: str) -> bool:
            """v2.41: Detecta se o output tem títulos de quadro-síntese (📋) sem tabela correspondente."""
            if not text:
                return False
            lines = text.splitlines()
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                # Detecta título de quadro-síntese
                if re.match(r'^#{3,5}\s*📋', line):
                    # Procura tabela nas próximas 5 linhas não-vazias
                    found_table = False
                    for j in range(i + 1, min(i + 8, len(lines))):
                        next_line = lines[j].strip()
                        if next_line.startswith('|') and '|' in next_line[1:]:
                            found_table = True
                            break
                        if next_line.startswith('#'):
                            break  # Novo heading sem tabela
                    if not found_table:
                        return True
                i += 1
            return False

        async def _retry_incomplete_table(result_text: str):
            incomplete_table = _has_incomplete_table(result_text or "")
            missing_table = _has_missing_table(result_text or "")
            needs_retry = incomplete_table or missing_table
            if needs_retry and not table_retry and depth < 2 and len(chunk_text) > 4000:
                reason = "incompleta" if incomplete_table else "ausente (título 📋 sem tabela)"
                retry_tokens = max_output_tokens_override or 32000
                print(f"{Fore.MAGENTA}✂️ Tabela {reason} no Chunk {idx}. Reprocessando com mais tokens...")
                retry_result = await self.process_chunk_async(
                    chunk_text,
                    idx=idx,
                    total=total,
                    previous_context=previous_context,
                    depth=depth,
                    global_structure=global_structure,
                    overlap_text=overlap_text,
                    cached_content=cached_content,
                    max_output_tokens_override=retry_tokens,
                    disable_cache=True,
                    table_retry=True
                )
                if retry_result and not _has_incomplete_table(retry_result):
                    if depth == 0 and not disable_cache:
                        self._save_chunk_cache(chunk_text, prompt_hash, retry_result)
                    return retry_result
                print(f"{Fore.MAGENTA}✂️ Reprocessamento não resolveu. Dividindo chunk...")
                metrics.record_adaptive_split()
                return await self._split_and_retry_async(
                    chunk_text,
                    idx,
                    system_prompt,
                    total,
                    contexto_estilo,
                    depth,
                    max_output_tokens_override=retry_tokens,
                    disable_cache=True
                )
            return None

        def _near_token_limit(completion_tokens: int, max_tokens: int) -> bool:
            try:
                completion_tokens = int(completion_tokens or 0)
                max_tokens = int(max_tokens or 0)
            except Exception:
                return False
            if completion_tokens <= 0 or max_tokens <= 0:
                return False
            return completion_tokens >= int(max_tokens * 0.92)

        def _looks_hard_truncated(text: str) -> bool:
            """
            Heurística conservadora para truncamento "duro" (normalmente por limite de tokens):
            - termina no meio de palavra
            - termina com bracket aberto ou fechamento "sobrando"
            - contém marcador de continuação no final
            """
            s = (text or "").strip()
            if not s:
                return True

            tail = s[-300:]
            if re.search(r"(?i)(?:\\[\\s*(?:continua|continuação|continuacao)\\s*\\]|\\(\\s*(?:continua|continuação|continuacao)\\s*\\))\\s*$", tail):
                return True

            last = s[-1]
            if last in "[({":
                return True

            # Fechamento "sobrando" no tail (ex.: termina com ']' sem haver '[' suficiente)
            tail2 = s[-2000:]
            if last == "]" and tail2.count("[") < tail2.count("]"):
                return True
            if last == ")" and tail2.count("(") < tail2.count(")"):
                return True
            if last == "}" and tail2.count("{") < tail2.count("}"):
                return True

            # Meio de palavra no final (sem pontuação final típica)
            if last.isalnum():
                if not re.search(r"[.!?…][\"”’')\\]]?\\s*$", s):
                    return True

            return False
        
        try:
            # Configuração de Segurança (Block None) e Parâmetros
            max_output_tokens = max_output_tokens_override or 32000  # v2.41: Aumentado de 16k para 32k (alinhado com format_transcription_gemini)
            safety_config = [
                types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
            ]

            if self.use_openai_primary and self.openai_client:
                try:
                    start_time_oai = time.time()
                    openai_kwargs = {}
                    if max_output_tokens_override:
                        openai_kwargs["max_completion_tokens"] = max_output_tokens
                    response = await self.openai_client.chat.completions.create(
                        model=self.openai_model,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_content}
                        ],
                        timeout=180,
                        **openai_kwargs
                    )
                    duration_oai = time.time() - start_time_oai
                    result = response.choices[0].message.content

                    oai_prompt = getattr(response.usage, 'prompt_tokens', 0) if hasattr(response, 'usage') else 0
                    oai_compl = getattr(response.usage, 'completion_tokens', 0) if hasattr(response, 'usage') else 0
                    finish_reason = None
                    try:
                        finish_reason = response.choices[0].finish_reason
                    except Exception:
                        finish_reason = None
                    openai_truncated = (finish_reason == "length")
                    cached_tokens = 0
                    if hasattr(response, 'usage'):
                        details = getattr(response.usage, 'prompt_tokens_details', None)
                        cached_tokens = getattr(details, 'cached_tokens', 0) or 0 if details else 0
                    metrics.record_call(
                        "openai",
                        oai_prompt,
                        oai_compl,
                        duration_oai,
                        model=self.openai_model,
                        cached_tokens_in=cached_tokens,
                    )

                    if contexto_estilo and result:
                        result = remover_eco_do_contexto(result, contexto_estilo)

                    retry_result = await _retry_incomplete_table(result)
                    if retry_result is not None:
                        return retry_result

                    if result:
                        try:
                            result = remover_marcadores_continua(result)
                        except Exception:
                            pass

                    # Se aparenta truncamento por limite, split para preservar conteúdo
                    openai_near_limit = _near_token_limit(oai_compl, max_output_tokens_override) if max_output_tokens_override else False
                    if (
                        not trunc_retry
                        and depth < 2
                        and len(chunk_text) > 4000
                        and _looks_hard_truncated(result or "")
                        and (openai_truncated or openai_near_limit)
                    ):
                        print(
                            f"{Fore.MAGENTA}✂️ Saída aparenta truncada por limite (OpenAI, Chunk {idx}). "
                            "Dividindo chunk..."
                        )
                        metrics.record_adaptive_split()
                        return await self._split_and_retry_async(
                            chunk_text,
                            idx,
                            system_prompt,
                            total,
                            contexto_estilo,
                            depth,
                            max_output_tokens_override=max_output_tokens_override,
                            disable_cache=True,
                        )

                    if depth == 0:
                        self._save_chunk_cache(chunk_text, prompt_hash, result)
                    return result
                except Exception as e_openai:
                    print(f"{Fore.YELLOW}⚠️ Falha no OpenAI primário (Chunk {idx}): {e_openai}. Tentando Gemini...")

            def call_gemini():
                nonlocal cached_content
                gen_config = types.GenerateContentConfig(
                    max_output_tokens=max_output_tokens,
                    temperature=0.1,
                    top_p=0.9,
                    top_k=40,
                    safety_settings=safety_config,
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level=self._resolve_thinking_level(),
                    ),
                )

                # SE USAR CACHE: Passar cached_content e APENAS user_content
                if cached_content:
                    gen_config.cached_content = cached_content.name
                    contents = user_content
                else:
                    # Sem cache: System + User
                    contents = f"{system_prompt}\n\n{user_content}"

                max_retries = int(os.getenv("IUDEX_GEMINI_RETRY_ATTEMPTS", "4"))
                base_sleep = float(os.getenv("IUDEX_GEMINI_RETRY_BASE_SECONDS", "6"))
                attempt = 0
                tried_global = False

                while True:
                    try:
                        return self.client.models.generate_content(
                            model=self.llm_model,
                            contents=contents,
                            config=gen_config,
                        )
                    except Exception as e:
                        msg = str(e)
                        is_model_not_found = (
                            ("404" in msg or "NOT_FOUND" in msg)
                            and "Publisher Model" in msg
                            and "was not found" in msg
                        )
                        if (
                            is_model_not_found
                            and not tried_global
                            and getattr(self, "_gemini_use_vertex", False)
                            and (getattr(self, "_gemini_vertex_location", None) or "").lower() not in ("", "global")
                            and os.getenv("IUDEX_VERTEX_FALLBACK_GLOBAL_ON_NOT_FOUND", "true").lower() in ("1", "true", "yes")
                        ):
                            tried_global = True
                            prev_loc = getattr(self, "_gemini_vertex_location", None)
                            print(
                                f"{Fore.YELLOW}⚠️  Modelo '{self.llm_model}' indisponível em '{prev_loc}'. "
                                f"Tentando Vertex AI em 'global'..."
                            )
                            cached_content = None
                            try:
                                gen_config.cached_content = None
                            except Exception:
                                pass
                            project = (
                                getattr(self, "_gemini_vertex_project", None)
                                or os.getenv("GOOGLE_CLOUD_PROJECT")
                            )
                            self.client = genai.Client(vertexai=True, project=project, location="global")
                            self._gemini_vertex_location = "global"
                            contents = f"{system_prompt}\n\n{user_content}"
                            continue

                        is_rate_limit = (
                            "429" in msg
                            or "RESOURCE_EXHAUSTED" in msg
                            or "rate limit" in msg.lower()
                        )
                        if not is_rate_limit:
                            raise
                        attempt += 1
                        if attempt > max_retries:
                            raise
                        sleep_for = min(base_sleep * (2 ** (attempt - 1)), 60.0)
                        sleep_for += random.uniform(0.2, 1.5)
                        print(
                            f"{Fore.YELLOW}⏳ Rate limit Gemini (Chunk {idx}). "
                            f"Aguardando {sleep_for:.1f}s (tentativa {attempt}/{max_retries})..."
                        )
                        time.sleep(sleep_for)

            # Executa chamada síncrona em thread separada com timeout
            start_time = time.time()
            timeout_seconds = int(os.getenv("IUDEX_GEMINI_TIMEOUT_SECONDS", "120"))
            try:
                response = await asyncio.wait_for(asyncio.to_thread(call_gemini), timeout=timeout_seconds)
            except asyncio.TimeoutError as e:
                print(f"{Fore.YELLOW}⏱️ Timeout no Gemini (Chunk {idx}) após {timeout_seconds}s")
                if depth < 2 and len(chunk_text) > 4000:
                    print(f"{Fore.MAGENTA}✂️ Timeout detectado. Tentando ADAPTIVE CHUNKING...")
                    return await self._split_and_retry_async(
                        chunk_text,
                        idx,
                        system_prompt,
                        total,
                        contexto_estilo,
                        depth,
                        max_output_tokens_override=max_output_tokens_override,
                        disable_cache=disable_cache
                    )
                raise e
            duration = time.time() - start_time
            
            # Extract token counts from response metadata
            prompt_tokens = 0
            completion_tokens = 0
            try:
                if hasattr(response, 'usage_metadata'):
                    usage = response.usage_metadata
                    prompt_tokens = getattr(usage, 'prompt_token_count', 0) or 0
                    completion_tokens = getattr(usage, 'candidates_token_count', 0) or 0
            except: pass
            
            cached_tokens = 0
            if hasattr(response, 'usage_metadata'):
                usage = response.usage_metadata
                cached_tokens = getattr(usage, 'cached_content_token_count', 0) or 0
            metrics.record_call(
                "gemini",
                prompt_tokens,
                completion_tokens,
                duration,
                model=self.llm_model,
                cached_tokens_in=cached_tokens,
            )
            
            try:
                result = response.text
            except ValueError:
                result = "" # Fallback se bloqueado ou vazio
                
            # === SMART STITCHING (v2.10) ===
            if contexto_estilo and result:
                result = remover_eco_do_contexto(result, contexto_estilo)
            retry_result = await _retry_incomplete_table(result)
            if retry_result is not None:
                return retry_result

            if result:
                try:
                    result = remover_marcadores_continua(result)
                except Exception:
                    pass

            # === ADAPTIVE CHUNCHING CHECK (v2.10) ===
            # v2.24: ratio por PALAVRAS (mais robusto que chars) e ignorando metadados do transcript
            # Motivo: chunks brutos podem conter "SPEAKER X" e timestamps [HH:MM], que são removidos na formatação
            # e derrubam artificialmente o ratio len(result)/len(chunk_text), ativando chunking adaptativo indevidamente.
            def _strip_transcript_metadata_for_ratio(text: str) -> str:
                if not text:
                    return ""
                out_lines = []
                for ln in text.splitlines():
                    s = ln.strip()
                    # Remove headers de diarização
                    if re.match(r'^SPEAKER\s+\d+\s*$', s):
                        continue
                    # Remove timestamps no início da linha
                    s = re.sub(r'^\[\d{1,2}:\d{2}(?::\d{2})?\]\s*', '', s)
                    out_lines.append(s)
                return "\n".join(out_lines)

            def _count_words(text: str) -> int:
                if not text:
                    return 0
                return len(re.findall(r'\b\w+\b', text, flags=re.UNICODE))

            mode_now = getattr(self, "_current_mode", "APOSTILA")
            in_words = _count_words(_strip_transcript_metadata_for_ratio(chunk_text))
            out_words = _count_words(result)
            ratio = (out_words / in_words) if in_words > 0 else 1.0
            # Limiares: mais tolerante em APOSTILA (limpeza de oralidade), mais estrito em FIDELIDADE
            threshold = 0.55 if str(mode_now).upper() != "FIDELIDADE" else 0.70
            is_compressed = ratio < threshold
            
            if (is_compressed or not result) and depth < 2 and len(chunk_text) > 4000:
                reason = "compressão excessiva" if is_compressed else "resposta vazia"
                if is_compressed:
                    print(f"\n{Fore.MAGENTA}✂️ ATIVANDO CHUNKING ADAPTATIVO para Chunk {idx} (Motivo: {reason} | Ratio(palavras): {ratio:.2f} | in={in_words} out={out_words} | limiar={threshold:.2f})")
                else:
                    print(f"\n{Fore.MAGENTA}✂️ ATIVANDO CHUNKING ADAPTATIVO para Chunk {idx} (Motivo: {reason})")
                metrics.record_adaptive_split()
                return await self._split_and_retry_async(
                    chunk_text,
                    idx,
                    system_prompt,
                    total,
                    contexto_estilo,
                    depth,
                    max_output_tokens_override=max_output_tokens_override,
                    disable_cache=disable_cache
                )

            # Truncamento por limite de tokens: split para evitar perda (especialmente no final do doc)
            if (
                not trunc_retry
                and depth < 2
                and len(chunk_text) > 4000
                and _near_token_limit(completion_tokens, max_output_tokens)
                and _looks_hard_truncated(result or "")
            ):
                print(
                    f"{Fore.MAGENTA}✂️ Saída aparenta truncada por limite (Gemini, Chunk {idx}). "
                    "Dividindo chunk..."
                )
                metrics.record_adaptive_split()
                return await self._split_and_retry_async(
                    chunk_text,
                    idx,
                    system_prompt,
                    total,
                    contexto_estilo,
                    depth,
                    max_output_tokens_override=max_output_tokens_override,
                    disable_cache=True,
                )

            if depth == 0 and not disable_cache:
                self._save_chunk_cache(chunk_text, prompt_hash, result)
            return result

        except Exception as e:
            # Log full exception details
            import traceback
            print(f"{Fore.RED}❌ Detalhes do erro no Chunk {idx}:")
            traceback.print_exc()
            
            # Fallback Logic can trigger adaptive chunking too if it's a token limit error
            if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                if depth < 2:
                    print(f"{Fore.MAGENTA}✂️ Erro de limite detectado. Tentando ADAPTIVE CHUNKING...")
                    return await self._split_and_retry_async(
                        chunk_text,
                        idx,
                        system_prompt,
                        total,
                        contexto_estilo,
                        depth,
                        max_output_tokens_override=max_output_tokens_override,
                        disable_cache=disable_cache
                    )
            
            print(f"{Fore.YELLOW}⚠️  Falha no Gemini (Chunk {idx}): {e}")
            
            if self.openai_client:
                print(f"{Fore.CYAN}🤖 Tentando fallback para OpenAI ({self.openai_model})...")
                try:
                    start_time_oai = time.time()
                    openai_kwargs = {}
                    if max_output_tokens_override:
                        openai_kwargs["max_completion_tokens"] = max_output_tokens_override
                    response = await self.openai_client.chat.completions.create(
                        model=self.openai_model,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_content}
                        ],
                        timeout=180,
                        **openai_kwargs
                    )
                    duration_oai = time.time() - start_time_oai
                    result = response.choices[0].message.content
                    
                    # Record OpenAI metrics
                    oai_prompt = getattr(response.usage, 'prompt_tokens', 0) if hasattr(response, 'usage') else 0
                    oai_compl = getattr(response.usage, 'completion_tokens', 0) if hasattr(response, 'usage') else 0
                    finish_reason = None
                    try:
                        finish_reason = response.choices[0].finish_reason
                    except Exception:
                        finish_reason = None
                    openai_truncated = (finish_reason == "length")
                    cached_tokens = 0
                    if hasattr(response, 'usage'):
                        details = getattr(response.usage, 'prompt_tokens_details', None)
                        cached_tokens = getattr(details, 'cached_tokens', 0) or 0 if details else 0
                    metrics.record_call(
                        "openai",
                        oai_prompt,
                        oai_compl,
                        duration_oai,
                        model=self.openai_model,
                        cached_tokens_in=cached_tokens,
                    )
                    
                    # Apply cleanup to OpenAI result too
                    if contexto_estilo and result:
                        result = remover_eco_do_contexto(result, contexto_estilo)
                    if result:
                        try:
                            result = remover_marcadores_continua(result)
                        except Exception:
                            pass

                    openai_near_limit = _near_token_limit(oai_compl, max_output_tokens_override) if max_output_tokens_override else False
                    if (
                        not trunc_retry
                        and depth < 2
                        and len(chunk_text) > 4000
                        and _looks_hard_truncated(result or "")
                        and (openai_truncated or openai_near_limit)
                    ):
                        print(
                            f"{Fore.MAGENTA}✂️ Saída aparenta truncada por limite (OpenAI fallback, Chunk {idx}). "
                            "Dividindo chunk..."
                        )
                        metrics.record_adaptive_split()
                        return await self._split_and_retry_async(
                            chunk_text,
                            idx,
                            system_prompt,
                            total,
                            contexto_estilo,
                            depth,
                            max_output_tokens_override=max_output_tokens_override,
                            disable_cache=True,
                        )
                        
                    if depth == 0 and not disable_cache:
                        self._save_chunk_cache(chunk_text, prompt_hash, result)
                    return result
                except Exception as e_openai:
                    print(f"{Fore.RED}❌ Falha também no OpenAI: {e_openai}")
                    raise e_openai
            else:
                print(f"{Fore.RED}❌ Falha no chunk {idx} e sem fallback OpenAI configurado.")
                raise e # Let tenacity handle retry

    async def _split_and_retry_async(
        self,
        text,
        idx,
        prompt,
        total,
        context,
        depth,
        max_output_tokens_override=None,
        disable_cache=False
    ):
        """Divides chunk in half and processes recursively"""
        mid = len(text) // 2
        # Try to find paragraph break near middle
        margin = int(len(text) * 0.2)
        start_search = max(0, mid - margin)
        end_search = min(len(text), mid + margin)
        search_zone = text[start_search:end_search]
        
        split_pos = -1
        last_para = search_zone.rfind('\n\n')
        if last_para != -1:
            split_pos = start_search + last_para + 2
        else:
            split_pos = mid # Fallback hard split
            
        part_a = text[:split_pos]
        part_b = text[split_pos:]
        
        print(f"   -> Dividindo (Sequencial): Parte A ({len(part_a)}c) + Parte B ({len(part_b)}c)")
        
        # ===================================================================
        # SEQUENTIAL EXECUTION (v2.10 Improvement)
        # Process A first, then use A's tail as context for B.
        # Trades speed (~2x slower) for style continuity and coherence.
        # ===================================================================
        
        res_a = await self._format_chunk_async(
            part_a,
            f"{idx}.A",
            prompt,
            total,
            context,
            depth + 1,
            max_output_tokens_override=max_output_tokens_override,
            disable_cache=disable_cache
        )
        
        # Use the tail of A's result as context for B
        context_for_b = res_a[-2000:] if len(res_a) > 2000 else res_a
        
        res_b = await self._format_chunk_async(
            part_b,
            f"{idx}.B",
            prompt,
            total,
            context_for_b,
            depth + 1,
            max_output_tokens_override=max_output_tokens_override,
            disable_cache=disable_cache
        )
        
        return f"{res_a}\n\n{res_b}"

    async def _identify_speaker_async(self, content, professors_info, speaker_label):
        """Identifica speaker com cache e heurística"""
        # Cache simples em memória
        if not hasattr(self, 'speaker_cache'): self.speaker_cache = {}
        if speaker_label in self.speaker_cache: return self.speaker_cache[speaker_label]
        
        prompt = f"""
        Analise o início do texto abaixo e a lista de professores extraída da introdução.
        Identifique quem é o provável professor falando e qual a disciplina.
        
        Falante (Label): {speaker_label if speaker_label else "Desconhecido"}
        
        Lista de Professores (Contexto):
        {professors_info}
        
        Texto (Início):
        {content[:5000]}...
        
        Retorne APENAS um JSON:
        {{
            "nome": "Nome do Professor",
            "disciplina": "Disciplina"
        }}
        """
        
        try:
            def call_gemini():
                return self.client.models.generate_content(
                    model=self.llm_model,
                    contents=f"Você é um assistente que identifica palestrantes.\n\n{prompt}",
                    config=types.GenerateContentConfig(
                        temperature=0.1,
                        max_output_tokens=20000,
                        thinking_config=types.ThinkingConfig(
                            include_thoughts=False,
                            thinking_level="LOW"
                        )
                    )
                )

            response = await asyncio.to_thread(call_gemini)
            _record_genai_usage(response, model=self.llm_model)
            content_json = response.text
            self.speaker_cache[speaker_label] = content_json
            return content_json
        except Exception as e:
            return '{"nome": "Professor", "disciplina": "Disciplina"}'

    def _extract_professors_context(self, full_text):
        """Extrai lista de professores (Deep Scan)"""
        print(f"   🕵️  Extraindo contexto de professores (Scan Completo)...")
        
        intro_context = full_text[:5000]
        keywords = ["meu nome é", "sou o professor", "sou a professora", "aqui é o professor"]
        
        found_contexts = []
        lower_text = full_text.lower()
        
        for keyword in keywords:
            start_idx = 0
            while True:
                idx = lower_text.find(keyword, start_idx)
                if idx == -1: break
                start_ctx = max(0, idx - 500)
                end_ctx = min(len(full_text), idx + 500)
                found_contexts.append(full_text[start_ctx:end_ctx])
                start_idx = idx + len(keyword)
        
        combined_context = intro_context + "\n\n... [TRECHOS] ...\n\n" + "\n\n".join(found_contexts)
        if len(combined_context) > 50000: combined_context = combined_context[:50000]

        try:
            response = self.client.models.generate_content(
                model=self.llm_model,
                contents=f"Extraia professores JSON\n\n{combined_context}",
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="LOW"
                    )
                )
            )
            _record_genai_usage(response, model=self.llm_model)
            return response.text
        except:
            return "{'professores': []}"

    async def _generate_header_async(self, formatted_content, professor_context_json):
        try:
            try:
                prof_ctx = json.loads(professor_context_json)
                prof = prof_ctx.get("nome", "Professor")
                disc = prof_ctx.get("disciplina", "Disciplina")
            except:
                prof = "Professor"
                disc = "Disciplina"
            
            prompt = f"""
            Gere APENAS o título Markdown para esta seção.
            Professor: {prof}
            Disciplina: {disc}
            
            Conteúdo:
            {formatted_content[:1000]}...
            
            FORMATO DE SAÍDA:
            # Prof. {prof} - {disc}
            """
            
            try:
                def call_gemini():
                    return self.client.models.generate_content(
                        model=self.llm_model,
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            max_output_tokens=100,
                            thinking_config=types.ThinkingConfig(
                                include_thoughts=False,
                                thinking_level="LOW"
                            )
                        )
                    )
                response = await asyncio.to_thread(call_gemini)
                _record_genai_usage(response, model=self.llm_model)
                return response.text.strip()
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️  Falha no header (Gemini): {e}")
                
                if self.openai_client:
                     print(f"{Fore.CYAN}🤖 Fallback: Header com OpenAI...")
                     try:
                        response = await self.openai_client.chat.completions.create(
                            model=self.openai_model,
                            messages=[{"role": "user", "content": prompt}],
                            max_tokens=100
                        )
                        _record_openai_usage(response, model=self.openai_model)
                        return response.choices[0].message.content.strip()
                     except Exception as e_openai:
                         print(f"{Fore.RED}❌ Falha também no OpenAI: {e_openai}")
                
                # Fallback final se tudo falhar
                return "## Tópico (Recuperado)"
        except: return f"# {prof} - {disc}\n"

    async def _process_segment_parallel(self, segment, professors_info, idx, system_prompt):
        speaker = segment['speaker']
        content = segment['content']
        
        print(f"\n{Fore.YELLOW}▶ Segmento {idx+1} ({speaker})...")
        chunks = self._smart_chunk_with_overlap(content, max_size=8000, overlap=1500)
        print(f"   {len(chunks)} chunks de ~8k chars (com 1.5k overlap)")
        
        context_task = self._identify_speaker_async(content[:5000], professors_info, speaker)
        
        chunk_tasks = [
            self._format_chunk_async(chunk, j, system_prompt)
            for j, chunk in enumerate(chunks)
        ]
        
        prof_context, *formatted_parts = await asyncio.gather(context_task, *chunk_tasks)
        
        full_content = remover_overlap_duplicado(formatted_parts)
        # Fallback para limpeza fina
        full_content = remover_paragrafos_identicos_consecutivos(full_content)
        header = await self._generate_header_async(full_content[:10000], prof_context)
        
        return f"{header}\n\n{full_content}\n\n---\n\n"

    def _renumber_topics(self, text):
        """Renumera tópicos sequencialmente"""
        lines = text.split('\n')
        new_lines = []
        main_counter = 0
        sub_counter = 0
        
        for line in lines:
            match_main = re.match(r'^##\s+\d+\.?\s+(.+)', line)
            if match_main:
                main_counter += 1
                sub_counter = 0
                new_lines.append(f"## {main_counter}. {match_main.group(1)}")
                continue
            
            match_sub = re.match(r'^###\s+(?:\d+(?:\.\d+)?\.?)?\s*(.+)', line)
            if match_sub:
                sub_counter += 1
                new_lines.append(f"### {main_counter}.{sub_counter} {match_sub.group(1)}")
                continue
                
            new_lines.append(line)
        return "\n".join(new_lines)

    def _fix_omissions(self, raw_transcript, formatted_text, omissions_report):
        """Tenta corrigir as omissões detectadas usando abordagem targeted chunk-by-chunk"""
        print(f"{Fore.CYAN}🔧 Tentando corrigir omissões automaticamente...")
        
        extraction_prompt = """# TAREFA: EXTRAIR CONTEÚDO OMITIDO
Você receberá:
1. RELATÓRIO DE OMISSÕES: Lista do que está faltando
2. TRANSCRIÇÃO BRUTA: Onde encontrar o conteúdo

## SUA MISSÃO:
Localize na transcrição bruta os trechos exatos que correspondem às omissões listadas.
Para cada omissão, extraia o trecho relevante da transcrição.

## RELATÓRIO DE OMISSÕES:
{report}

Retorne APENAS os trechos extraídos, sem comentários adicionais."""

        try:
            # v2.28: Expandir limite para cobrir transcrições longas (Gemini suporta ~1M tokens)
            max_transcript_chars = 500_000
            transcript_excerpt = raw_transcript[:max_transcript_chars]
            if len(raw_transcript) > max_transcript_chars:
                print(f"{Fore.YELLOW}   ⚠️ Transcrição truncada para {max_transcript_chars:,} chars (total: {len(raw_transcript):,})")
            extract_response = self.client.models.generate_content(
                model=self.llm_model,
                contents=f"{extraction_prompt.format(report=omissions_report)}\n\nTRANSCRIÇÃO BRUTA:\n{transcript_excerpt}",
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="HIGH"
                    )
                )
            )
            _record_genai_usage(extract_response, model=self.llm_model)
            
            # Gemini client returns .text, not .choices[0].message.content (OpenAI style)
            missing_content = extract_response.text
            print(f"{Fore.CYAN}   📝 Conteúdo omitido extraído ({len(missing_content)} caracteres)")
            
            chunks = self._smart_chunk_with_overlap(formatted_text, max_size=self.MAX_CHUNK_SIZE)
            fixed_chunks = []
            
            insertion_prompt = """# TAREFA: INSERIR CONTEÚDO FALTANTE
Você receberá:
1. Um TRECHO da apostila formatada
2. CONTEÚDO OMITIDO que precisa ser adicionado

## SUA MISSÃO:
Se o trecho da apostila é onde o conteúdo omitido deveria estar, insira-o naturalmente.
Se não for o local apropriado, retorne o trecho INALTERADO.

## REGRAS:
- Mantenha toda formatação Markdown
- Integre o conteúdo omitido de forma fluida
- Use tom didático e formal (3ª pessoa)
- NÃO remova nada existente

## CONTEÚDO OMITIDO A INSERIR:
{missing}

Retorne o trecho (modificado ou inalterado)."""

            for i, chunk in enumerate(chunks):
                print(f"{Fore.CYAN}   🔧 Processando chunk {i+1}/{len(chunks)}...")
                try:
                    fix_response = self.client.models.generate_content(
                        model=self.llm_model,
                        contents=f"{insertion_prompt.format(missing=missing_content)}\n\n{chunk}",
                        config=types.GenerateContentConfig(
                            thinking_config=types.ThinkingConfig(
                                include_thoughts=False,
                                thinking_level="HIGH"
                            )
                        )
                    )
                    _record_genai_usage(fix_response, model=self.llm_model)
                    fixed_chunks.append(fix_response.text)
                except Exception as chunk_error:
                    print(f"{Fore.YELLOW}   ⚠️  Erro no chunk {i+1}, mantendo original: {chunk_error}")
                    fixed_chunks.append(chunk)
            
            fixed_text = "\n\n".join(fixed_chunks)
            print(f"{Fore.GREEN}   ✅ Texto corrigido gerado ({len(fixed_chunks)} chunks processados)")

            # v2.30: Re-validação heurística leve pós-correção (sem custo LLM extra)
            reval_ok, reval_issues = self._validate_preservation_heuristics(raw_transcript, fixed_text)
            if not reval_ok:
                print(f"{Fore.YELLOW}   ⚠️ Re-validação pós-fix detectou {len(reval_issues)} problemas:")
                for ri in reval_issues[:3]:
                    print(f"      - {ri}")
            else:
                print(f"{Fore.GREEN}   ✅ Re-validação pós-fix aprovada")

            return fixed_text, reval_ok

        except Exception as e:
            print(f"{Fore.RED}   ❌ Falha ao corrigir omissões: {e}")
            return formatted_text, False

    def _validate_preservation_heuristics(self, original_text, formatted_text):
        """Validação Heurística com Tolerância Adaptativa"""
        print(f"\n{Fore.CYAN}🔍 Validação Heurística de Preservação (Adaptativa)...")
        issues = []
        
        # 1. Referências legais e jurisprudenciais (v2.30: padrão expandido)
        _LEGAL_REF_PATTERN = (
            r'(?:Lei\s+(?:Complementar\s+|Ordinária\s+)?|LC\s+|Decreto(?:-Lei)?\s+|DL\s+|MP\s+|Medida\s+Provisória\s+)'
            r'n?º?\s*[\d\.]+(?:/\d+)?'
            r'|Art\.?\s*\d+[°º]?'
            r'|Súmula(?:\s+Vinculante)?\s+\d+'
            r'|(?:REsp|RE|HC|MS|ADI|ADPF|ADC|RCL|Rcl|AgRg|AREsp)\s*n?º?\s*[\d\.\/\-]+'
            r'|Tema\s+(?:de\s+)?(?:Repercussão\s+Geral\s+)?\d+'
            r'|Informativo\s+\d+'
        )
        original_laws = set(re.findall(_LEGAL_REF_PATTERN, original_text, re.IGNORECASE))
        formatted_laws = set(re.findall(_LEGAL_REF_PATTERN, formatted_text, re.IGNORECASE))
        missing_laws = original_laws - formatted_laws
        if missing_laws:
            # Logar as referências específicas para facilitar revisão
            samples = list(missing_laws)[:5]
            issues.append(f"❌ {len(missing_laws)} referências legais/jurisprudenciais omitidas: {', '.join(samples)}")
        else:
            print(f"{Fore.GREEN}   ✅ Referências legais preservadas ({len(original_laws)} encontradas)")
        
        # 2. Comprimento Adaptativo (Lógica do script Gemini)
        palavras_input = len(original_text.split())
        palavras_output = len(formatted_text.split())
        
        if palavras_input == 0: ratio = 0
        else: ratio = palavras_output / palavras_input
        
        # Heurística de Oralidade
        marcadores_oralidade = ['né', 'então', 'tipo', 'aí', 'pessoal', 'galera', 'tá', 'olha', 'gente', 'veja', 'bom']
        input_lower = original_text.lower()
        count_oralidade = sum(input_lower.count(m) for m in marcadores_oralidade)
        densidade_oralidade = count_oralidade / palavras_input if palavras_input > 0 else 0
        
        # Define tolerância baseada na densidade
        if densidade_oralidade > 0.025:  # Muito coloquial (>2.5%)
            tolerancia = 0.45  # Aceita reduzir até 45%
            tipo = "Muito Coloquial"
        elif densidade_oralidade > 0.015:  # Médio
            tolerancia = 0.38
            tipo = "Coloquial"
        elif densidade_oralidade > 0.008:  # Pouca
            tolerancia = 0.30
            tipo = "Pouca Oralidade"
        else:  # Técnico
            tolerancia = 0.22
            tipo = "Técnico/Denso"
            
        limite_minimo = 1.0 - tolerancia
        
        print(f"   📊 Análise: {tipo} (Densidade: {densidade_oralidade:.2%})")
        print(f"      Ratio atual: {ratio:.1%} (Mínimo aceitável: {limite_minimo:.1%})")
        
        if ratio < limite_minimo:
            issues.append(f"⚠️ Texto formatado muito curto ({ratio:.1%}). Esperado no mínimo {limite_minimo:.1%}")
        else:
            print(f"{Fore.GREEN}   ✅ Comprimento aprovado")
        
        if issues:
            print(f"{Fore.RED}━━━ PROBLEMAS ━━━")
            for i in issues: print(f"   {i}")
            return False, issues
        return True, []

    def validate_completeness_full(self, raw_transcript, formatted_text, video_name, global_structure=None):
        """
        v2.16: Validação LLM Full-Context - Envia documento INTEIRO para análise.
        Aproveita a janela de contexto do Gemini 3 Flash (2M tokens).
        """
        print(f"{Fore.YELLOW}🔍 Validação LLM Full-Context (v2.16)...")
        
        # Calcular tamanho aproximado em tokens (estimativa: 4 chars = 1 token)
        total_chars = len(raw_transcript) + len(formatted_text)
        estimated_tokens = total_chars // 4
        print(f"   📊 Tamanho estimado: {estimated_tokens:,} tokens")
        
        # Safety check: Se exceder 1.5M tokens, usar fallback de amostragem
        if estimated_tokens > 1_500_000:
            print(f"{Fore.YELLOW}   ⚠️ Documento muito grande ({estimated_tokens:,} tokens). Usando validação por amostragem.")
            return self._validate_by_sampling(raw_transcript, formatted_text, video_name)
        
        validation_prompt = """# TAREFA DE VALIDAÇÃO DE FIDELIDADE (Full-Context)

Você é um auditor de qualidade para transcrições jurídicas formatadas.

## SEU OBJETIVO
Compare o TEXTO ORIGINAL (transcrição bruta) com o TEXTO FORMATADO (apostila) e identifique:

1. **OMISSÕES GRAVES**: Conceitos jurídicos, leis, súmulas, artigos ou exemplos importantes que estavam no original mas foram omitidos no formatado.
2. **DISTORÇÕES**: Informações que foram alteradas de forma que mude o sentido jurídico.
3. **ESTRUTURA**: Verifique se os tópicos e subtópicos estão organizados de forma lógica e se não há duplicações.

## REGRAS
- NÃO considere como omissão: hesitações, "né", "então", dados repetitivos, conversas paralelas.
- CONSIDERE como omissão: qualquer lei, súmula, artigo, jurisprudência, exemplo prático ou dica de prova.
- Preste atenção especial em: números de leis, prazos, percentuais, valores monetários.
- NÃO faça análise jurídica externa nem verifique a veracidade de leis.
- Sua saída deve refletir apenas divergências entre o texto bruto e o formatado.

## FORMATO DE RESPOSTA (JSON)
{
    "aprovado": true/false,
    "nota_fidelidade": 0-10,
    "omissoes_graves": ["descrição clara do item omitido"],
    "distorcoes": ["descrição clara da distorção"],
    "problemas_estrutura": ["títulos duplicados ou hierarquia quebrada"],
    "observacoes": "comentário geral sobre a qualidade"
}

Retorne APENAS o JSON, sem markdown."""

        structure_context = ""
        if global_structure:
            structure_context = f"\n\n## ESTRUTURA ESPERADA (Mapeamento Inicial):\n{global_structure[:5000]}"
        
        try:
            response = self.client.models.generate_content(
                model=self.llm_model,
                contents=f"""{validation_prompt}{structure_context}

## TEXTO ORIGINAL (Transcrição Bruta):
{raw_transcript}

## TEXTO FORMATADO (Apostila):
{formatted_text}
""",
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    max_output_tokens=8000,
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="HIGH"
                    )
                )
            )
            _record_genai_usage(response, model=self.llm_model)
            
            # === ROBUST JSON PARSING (v2.25) ===
            result = None
            raw_text = response.text.strip()
            
            # Attempt 1: Direct JSON parse
            try:
                result = json.loads(raw_text)
            except json.JSONDecodeError:
                pass
            
            # Attempt 2: Extract JSON from markdown code block
            if result is None:
                import re
                json_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', raw_text)
                if json_match:
                    try:
                        result = json.loads(json_match.group(1).strip())
                    except json.JSONDecodeError:
                        pass
            
            # Attempt 3: Find JSON object in text
            if result is None:
                brace_start = raw_text.find('{')
                brace_end = raw_text.rfind('}')
                if brace_start != -1 and brace_end > brace_start:
                    try:
                        result = json.loads(raw_text[brace_start:brace_end+1])
                    except json.JSONDecodeError:
                        pass
            
            # Attempt 4: Try ast.literal_eval as last resort
            if result is None:
                import ast
                try:
                    brace_start = raw_text.find('{')
                    brace_end = raw_text.rfind('}')
                    if brace_start != -1 and brace_end > brace_start:
                        result = ast.literal_eval(raw_text[brace_start:brace_end+1])
                except (ValueError, SyntaxError):
                    pass
            
            # Attempt 5: Retry with stricter prompt
            if result is None:
                print(f"{Fore.YELLOW}   ⚠️ JSON malformado, tentando retry com prompt estrito...")
                retry_response = self.client.models.generate_content(
                    model=self.llm_model,
                    contents=f"""O seguinte texto deveria ser um JSON válido mas está malformado. 
Corrija-o e retorne APENAS o JSON válido, sem explicações:

{raw_text[:3000]}

O JSON deve ter exatamente esta estrutura:
{{"aprovado": true/false, "nota_fidelidade": 0-10, "omissoes_graves": [], "distorcoes": [], "problemas_estrutura": [], "observacoes": ""}}""",
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        max_output_tokens=2000
                    )
                )
                _record_genai_usage(retry_response, model=self.llm_model)
                try:
                    result = json.loads(retry_response.text.strip())
                except json.JSONDecodeError:
                    print(f"{Fore.RED}   ❌ Retry também falhou, usando valores padrão.")
                    result = {"aprovado": False, "nota_fidelidade": 0, "erro_validacao": True, "requires_manual_review": True, "observacoes": "Parsing JSON falhou após 5 tentativas. Revisão manual necessária."}
            
            if isinstance(result, list):
                if len(result) > 0 and isinstance(result[0], dict):
                    result = result[0]
                else:
                    result = {}

            # Processar resultado
            aprovado = result.get('aprovado', True)
            nota = result.get('nota_fidelidade', 10)
            omissoes = result.get('omissoes_graves', [])
            distorcoes = result.get('distorcoes', [])
            problemas_estrutura = result.get('problemas_estrutura', [])
            observacoes = result.get('observacoes', '')
            
            # Log do resultado
            if aprovado:
                print(f"{Fore.GREEN}   ✅ Validação Full-Context APROVADA (Nota: {nota}/10)")
            else:
                print(f"{Fore.RED}   ❌ Validação Full-Context REPROVADA (Nota: {nota}/10)")
                if omissoes:
                    print(f"{Fore.RED}   📌 Omissões: {len(omissoes)}")
                    for o in omissoes[:3]:
                        print(f"      - {o[:100]}...")
                if distorcoes:
                    print(f"{Fore.RED}   ⚠️ Distorções: {len(distorcoes)}")
                if problemas_estrutura:
                    print(f"{Fore.YELLOW}   🏗️ Problemas de Estrutura: {len(problemas_estrutura)}")
            
            # Retornar relatório completo
            return {
                'aprovado': aprovado,
                'nota': nota,
                'omissoes': omissoes,
                'distorcoes': distorcoes,
                'problemas_estrutura': problemas_estrutura,
                'observacoes': observacoes
            }
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro na validação Full-Context: {e}")
            print(f"{Fore.RED}   ⚠️ ATENÇÃO: Documento requer revisão manual (validação falhou).")
            return {
                'aprovado': False,
                'nota': 0,
                'erro_validacao': True,
                'requires_manual_review': True,
                'omissoes': [],
                'distorcoes': [],
                'problemas_estrutura': [],
                'observacoes': f'ATENÇÃO: Validação falhou ({str(e)}). Revisão manual recomendada.'
            }

    def validate_fidelity_primary(
        self,
        raw_transcript,
        formatted_text,
        video_name,
        modo="APOSTILA",
        include_sources=False,
    ):
        """
        Auditoria de fidelidade primária (preventiva) com fallback para full-context.
        Retorna um relatório compatível com o formato antigo (_fidelidade.json).
        """
        if FIDELITY_AUDIT_AVAILABLE and FIDELITY_AUDIT_ENABLED:
            result = auditar_fidelidade_preventiva(
                self.client,
                raw_transcript,
                formatted_text,
                video_name,
                output_path=None,
                modo=modo,
                include_sources=include_sources,
            )
            compat = (result or {}).get("compat_fidelidade")
            if isinstance(compat, dict) and compat:
                return compat
            return result or {}

        fallback = self.validate_completeness_full(raw_transcript, formatted_text, video_name, None)
        if isinstance(fallback, dict):
            fallback["source"] = "validate_completeness_full"
        return fallback

    def _validate_by_sampling(self, raw_transcript, formatted_text, video_name):
        """Fallback: Validação por amostragem para documentos muito grandes (>1.5M tokens).

        Processa 3 janelas (INÍCIO, MEIO, FIM) de 80k chars cada via LLM e agrega os resultados.
        """
        print(f"{Fore.CYAN}   Usando validação por amostragem (3 janelas)...")

        window_size = 80000
        mid_raw = len(raw_transcript) // 2
        mid_fmt = len(formatted_text) // 2
        half_window = window_size // 2

        windows = [
            ("INÍCIO", raw_transcript[:window_size], formatted_text[:window_size]),
            ("MEIO", raw_transcript[max(0, mid_raw - half_window):mid_raw + half_window],
                      formatted_text[max(0, mid_fmt - half_window):mid_fmt + half_window]),
            ("FIM", raw_transcript[-window_size:], formatted_text[-window_size:]),
        ]

        validation_prompt = """# TAREFA DE VALIDAÇÃO DE FIDELIDADE (Amostragem)

Você é um auditor de qualidade para transcrições jurídicas formatadas.

## SEU OBJETIVO
Compare o TEXTO ORIGINAL (transcrição bruta) com o TEXTO FORMATADO (apostila) e identifique:

1. **OMISSÕES GRAVES**: Conceitos jurídicos, leis, súmulas, artigos ou exemplos importantes que estavam no original mas foram omitidos no formatado.
2. **DISTORÇÕES**: Informações que foram alteradas de forma que mude o sentido jurídico.
3. **ESTRUTURA**: Verifique se os tópicos e subtópicos estão organizados de forma lógica e se não há duplicações.

## REGRAS
- NÃO considere como omissão: hesitações, "né", "então", dados repetitivos, conversas paralelas.
- CONSIDERE como omissão: qualquer lei, súmula, artigo, jurisprudência, exemplo prático ou dica de prova.
- Preste atenção especial em: números de leis, prazos, percentuais, valores monetários.
- NÃO faça análise jurídica externa nem verifique a veracidade de leis.
- Sua saída deve refletir apenas divergências entre o texto bruto e o formatado.

## FORMATO DE RESPOSTA (JSON)
{
    "aprovado": true/false,
    "nota_fidelidade": 0-10,
    "omissoes_graves": ["descrição clara do item omitido"],
    "distorcoes": ["descrição clara da distorção"],
    "problemas_estrutura": ["títulos duplicados ou hierarquia quebrada"],
    "observacoes": "comentário geral sobre a qualidade"
}

Retorne APENAS o JSON, sem markdown."""

        all_results = []
        for label, raw_window, fmt_window in windows:
            print(f"{Fore.CYAN}   📊 Validando janela {label}...")
            try:
                response = self.client.models.generate_content(
                    model=self.llm_model,
                    contents=f"""{validation_prompt}

## JANELA: {label}

## TEXTO ORIGINAL (Transcrição Bruta):
{raw_window}

## TEXTO FORMATADO (Apostila):
{fmt_window}
""",
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        max_output_tokens=4000,
                        thinking_config=types.ThinkingConfig(
                            include_thoughts=False,
                            thinking_level="MEDIUM"
                        )
                    )
                )
                _record_genai_usage(response, model=self.llm_model)

                raw_text = response.text.strip()
                result = None
                try:
                    result = json.loads(raw_text)
                except json.JSONDecodeError:
                    brace_start = raw_text.find('{')
                    brace_end = raw_text.rfind('}')
                    if brace_start != -1 and brace_end > brace_start:
                        try:
                            result = json.loads(raw_text[brace_start:brace_end + 1])
                        except json.JSONDecodeError:
                            pass

                if isinstance(result, dict):
                    all_results.append(result)
                    nota_w = result.get('nota_fidelidade', 10)
                    print(f"{Fore.GREEN}      ✅ {label}: nota {nota_w}/10")
                else:
                    print(f"{Fore.YELLOW}      ⚠️ {label}: resposta inválida, ignorando janela")

            except Exception as e:
                print(f"{Fore.RED}      ❌ {label}: erro — {e}")

        if not all_results:
            print(f"{Fore.RED}   ❌ Nenhuma janela validada. Revisão manual necessária.")
            return {
                'aprovado': False, 'nota': 0,
                'omissoes': [], 'distorcoes': [], 'problemas_estrutura': [],
                'observacoes': 'Validação por amostragem falhou em todas as janelas. Revisão manual obrigatória.',
                'erro_validacao': True, 'requires_manual_review': True,
            }

        # Agregar resultados: aprovado só se TODAS as janelas aprovaram
        aprovado = all(r.get('aprovado', True) for r in all_results)
        notas = [r.get('nota_fidelidade', 10) for r in all_results]
        nota_media = sum(notas) / len(notas)
        omissoes = []
        distorcoes = []
        problemas = []
        for r in all_results:
            omissoes.extend(r.get('omissoes_graves', []) or [])
            distorcoes.extend(r.get('distorcoes', []) or [])
            problemas.extend(r.get('problemas_estrutura', []) or [])

        obs_parts = [r.get('observacoes', '') for r in all_results if r.get('observacoes')]
        observacoes = f"Validação por amostragem ({len(all_results)}/3 janelas). " + " | ".join(obs_parts)

        print(f"{Fore.GREEN if aprovado else Fore.RED}   {'✅' if aprovado else '❌'} Resultado agregado: nota {nota_media:.1f}/10 ({len(all_results)} janelas)")

        return {
            'aprovado': aprovado,
            'nota': round(nota_media, 1),
            'omissoes': omissoes,
            'distorcoes': distorcoes,
            'problemas_estrutura': problemas,
            'observacoes': observacoes,
        }

    async def auto_fix_structure(self, formatted_text: str, problemas: list, global_structure: str = None) -> str:
        """
        v2.17: Corretor IA Ativo - Corrige automaticamente problemas estruturais.
        
        Recebe o texto formatado e a lista de problemas detectados pelo auditor,
        e envia ao LLM para correção automática.
        
        Args:
            formatted_text: Texto markdown com problemas
            problemas: Lista de strings descrevendo os problemas estruturais
            global_structure: Estrutura esperada (opcional)
        
        Returns:
            Texto corrigido
        """
        print(f"{Fore.CYAN}🔧 Corretor IA Ativo (v2.17)...")
        print(f"   📋 Problemas a corrigir: {len(problemas)}")
        for p in problemas[:3]:
            print(f"      - {p[:80]}...")
        
        fix_prompt = """# TAREFA DE CORREÇÃO ESTRUTURAL

Você é um editor de documentos jurídicos formatados em Markdown.

## SEU OBJETIVO
Corrija os problemas estruturais listados abaixo NO TEXTO FORNECIDO.

## PROBLEMAS A CORRIGIR:
{problemas}

## REGRAS DE CORREÇÃO:
1. **Títulos Duplicados**: Se um título H2 aparece duas vezes, REMOVA a segunda ocorrência e mescle o conteúdo sob o primeiro.
2. **Hierarquia Quebrada**: Se um H3 está fora de seu H2 pai correto, MOVA-O para debaixo do H2 apropriado.
3. **Parágrafos Repetidos**: Se o mesmo parágrafo aparece duas vezes, REMOVA a duplicata.
4. **Numeração Inconsistente**: Se a numeração está errada (ex: 1, 2, 3, 3, 4), RENUMERE sequencialmente.

## IMPORTANTE:
- NÃO altere o conteúdo textual, apenas a estrutura.
- NÃO adicione novos conteúdos.
- NÃO remova informações jurídicas (leis, súmulas, artigos).
- Mantenha todas as tabelas intactas.
- Preserve a formatação Markdown (negrito, itálico, listas).

## FORMATO DE RESPOSTA:
Retorne APENAS o texto Markdown corrigido, sem explicações adicionais."""

        structure_hint = ""
        if global_structure:
            structure_hint = f"\n\n## ESTRUTURA ESPERADA:\n{global_structure[:3000]}"
        
        try:
            response = self.client.models.generate_content(
                model=self.llm_model,
                contents=f"""{fix_prompt.format(problemas=chr(10).join(f'- {p}' for p in problemas))}{structure_hint}

## TEXTO A CORRIGIR:
{formatted_text}
""",
                config=types.GenerateContentConfig(
                    max_output_tokens=65000,  # Documento pode ser grande
                    thinking_config=types.ThinkingConfig(
                        include_thoughts=False,
                        thinking_level="HIGH"
                    )
                )
            )
            _record_genai_usage(response, model=self.llm_model)
            
            fixed_text = response.text.strip()
            
            # Validação básica: o texto corrigido não deve ser muito menor
            if len(fixed_text) < len(formatted_text) * 0.7:
                print(f"{Fore.YELLOW}   ⚠️ Texto corrigido muito curto ({len(fixed_text)} vs {len(formatted_text)}). Mantendo original.")
                return formatted_text
            
            # Remover possíveis wrappers de código markdown que o LLM pode adicionar
            if fixed_text.startswith('```markdown'):
                fixed_text = fixed_text[len('```markdown'):].strip()
            if fixed_text.startswith('```'):
                fixed_text = fixed_text[3:].strip()
            if fixed_text.endswith('```'):
                fixed_text = fixed_text[:-3].strip()
            
            print(f"{Fore.GREEN}   ✅ Correção automática aplicada com sucesso")
            return fixed_text
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro no Corretor IA: {e}")
            print(f"{Fore.YELLOW}   ℹ️ Mantendo texto original")
            return formatted_text


    def _generate_audit_report(self, video_name, heuristic_issues, llm_issues):
        with open(f"audit_{video_name}.md", 'w') as f:
            f.write(f"# Auditoria: {video_name}\n")
            f.write(f"Heurística: {heuristic_issues}\n")
            f.write(f"LLM: {llm_issues}\n")

    async def map_structure(self, full_text):
        """Creates a global structure skeleton to guide the formatting."""
        _map_t0 = time.time()
        print(f"{Fore.CYAN}🗺️  Mapeando estrutura global do documento... [start={time.strftime('%H:%M:%S')}]")
        
        full_text = full_text or ""
        max_single = int(os.getenv("IUDEX_MAP_MAX_SINGLE_CHARS", self.MAP_MAX_SINGLE_CHARS))
        map_chunk_chars = int(os.getenv("IUDEX_MAP_CHUNK_CHARS", self.MAP_CHUNK_CHARS))
        map_chunk_overlap = int(os.getenv("IUDEX_MAP_CHUNK_OVERLAP_CHARS", self.MAP_CHUNK_OVERLAP_CHARS))

        # Preferir mapear o documento inteiro em uma chamada quando estiver dentro de um limite seguro.
        # Para transcrições muito longas, usar chunking + merge determinístico para cobrir INÍCIO→FIM.
        if len(full_text) <= max_single:
            input_samples = [full_text]
        else:
            print(f"{Fore.YELLOW}   ℹ️  Transcrição longa ({len(full_text):,} chars). Mapeando em chunks...")
            map_chunking_mode = os.getenv("IUDEX_MAP_CHUNKING_MODE", "auto").strip().lower()
            input_samples = None
            if map_chunking_mode != "safe":
                input_samples = chunk_texto_por_segmentos(
                    full_text,
                    max_chars=map_chunk_chars,
                    overlap_chars=map_chunk_overlap,
                )
                if input_samples:
                    print(f"{Fore.CYAN}   🧩 Chunking por segmentos (map) ativado: {len(input_samples)}")
            if not input_samples:
                input_samples = chunk_texto_seguro(
                    full_text,
                    max_chars=map_chunk_chars,
                    overlap_chars=map_chunk_overlap,
                )
                print(f"{Fore.CYAN}   📦 Chunks para mapeamento: {len(input_samples)}")
        fallback_sample = input_samples[0] if input_samples else full_text[:200000]

        def _extract_map_lines(text: str) -> list[dict]:
            items = []
            if not text:
                return items
            for raw_line in (text or "").splitlines():
                line = raw_line.strip()
                if not line or "| ABRE:" not in line:
                    continue
                m = re.match(
                    r'^\s*(\d+(?:\.\d+)*)\.\s*([^|]+?)\s*\|\s*ABRE:\s*"([^"]+)"\s*\|\s*FECHA:\s*"([^"]+)"\s*$',
                    raw_line,
                    flags=re.IGNORECASE,
                )
                if not m:
                    continue
                number = m.group(1).strip()
                title = m.group(2).strip()
                abre = m.group(3).strip()
                fecha = m.group(4).strip()
                depth = number.count(".") + 1
                items.append(
                    {
                        "depth": depth,
                        "title": title,
                        "abre": abre,
                        "fecha": fecha,
                    }
                )
            return items

        def _normalize_key(value: str) -> str:
            if not value:
                return ""
            value = value.lower()
            value = re.sub(r"[^\w\s]", " ", value, flags=re.UNICODE)
            value = re.sub(r"\s+", " ", value).strip()
            return value

        def _merge_structure_maps(maps: list[str]) -> Optional[str]:
            if not maps:
                return None
            all_items: list[dict] = []
            for part in maps:
                all_items.extend(_extract_map_lines(part or ""))
            if not all_items:
                return None

            # Dedup por âncora ABRE (mais estável) e fallback por título.
            seen: set[str] = set()
            ordered: list[dict] = []
            for item in all_items:
                key = _normalize_key(item.get("abre") or "") or _normalize_key(item.get("title") or "")
                if not key:
                    continue
                if key in seen:
                    continue
                seen.add(key)
                ordered.append(item)

            # Renumeração determinística (máx 3 níveis) para manter sequência limpa.
            h1 = 0
            h2 = 0
            h3 = 0
            out_lines: list[str] = []
            for item in ordered:
                depth = int(item.get("depth") or 1)
                depth = max(1, min(3, depth))
                if depth == 1:
                    h1 += 1
                    h2 = 0
                    h3 = 0
                    number = f"{h1}."
                elif depth == 2:
                    if h1 <= 0:
                        h1 = 1
                    h2 += 1
                    h3 = 0
                    number = f"{h1}.{h2}."
                else:
                    if h1 <= 0:
                        h1 = 1
                    if h2 <= 0:
                        h2 = 1
                    h3 += 1
                    number = f"{h1}.{h2}.{h3}."

                indent = "   " * (depth - 1)
                title = (item.get("title") or "").strip()
                abre = (item.get("abre") or "").strip()
                fecha = (item.get("fecha") or "").strip()
                if not title or not abre or not fecha:
                    continue
                out_lines.append(f'{indent}{number} {title} | ABRE: "{abre}" | FECHA: "{fecha}"')

            return "\n".join(out_lines).strip() or None

        async def _map_one(sample: str, *, part_idx: int, total_parts: int) -> Optional[str]:
            if not sample:
                return None
            # Mantém prompt original; evita inserir marcadores no texto para não contaminar âncoras verbatim.
            prompt = self.PROMPT_MAPEAMENTO.format(transcricao=sample)
            try:
                if self.provider == "openai":
                    response = await self.client.chat.completions.create(
                        model=self.llm_model,
                        messages=[{"role": "system", "content": prompt}],
                        max_completion_tokens=16384,
                    )
                    _record_openai_usage(response, model=self.llm_model)
                    content = response.choices[0].message.content.replace('```markdown', '').replace('```', '')
                    print(f"{Fore.GREEN}   ✅ Estrutura mapeada (OpenAI) [{part_idx}/{total_parts}]")
                    return content

                def call_gemini():
                    return self.client.models.generate_content(
                        model=self.llm_model,
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            max_output_tokens=10000,
                            thinking_config=types.ThinkingConfig(
                                include_thoughts=False,
                                thinking_level="HIGH"
                            )
                        )
                    )

                response = await asyncio.to_thread(call_gemini)
                _record_genai_usage(response, model=self.llm_model)
                content = response.text.replace('```markdown', '').replace('```', '')
                print(f"{Fore.GREEN}   ✅ Estrutura mapeada (Vertex AI) [{part_idx}/{total_parts}]")
                return content
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️  Falha no mapeamento (parte {part_idx}/{total_parts}) via {self.provider}: {e}")
                if self.openai_client:
                    print(f"{Fore.CYAN}🤖 Fallback: Mapeando com OpenAI ({self.openai_model})...")
                    try:
                        response = await self.openai_client.chat.completions.create(
                            model=self.openai_model,
                            messages=[{"role": "system", "content": prompt}],
                            max_completion_tokens=10000
                        )
                        _record_openai_usage(response, model=self.openai_model)
                        content = response.choices[0].message.content.replace('```markdown', '').replace('```', '')
                        print(f"{Fore.GREEN}   ✅ Estrutura mapeada (OpenAI Fallback) [{part_idx}/{total_parts}]")
                        return content
                    except Exception as e_openai:
                        print(f"{Fore.RED}❌ Falha também no OpenAI fallback: {e_openai}")
                        return None
                return None

        try:
            parts = []
            total_parts = len(input_samples)
            for idx, sample in enumerate(input_samples, start=1):
                mapped = await _map_one(sample, part_idx=idx, total_parts=total_parts)
                if mapped:
                    parts.append(mapped)

            if not parts:
                _elapsed = time.time() - _map_t0
                print(f"{Fore.YELLOW}🗺️  map_structure finalizado SEM resultado em {_elapsed:.1f}s")
                return None
            if len(parts) == 1:
                _elapsed = time.time() - _map_t0
                print(f"{Fore.GREEN}🗺️  map_structure OK em {_elapsed:.1f}s (1 parte)")
                return _sanitize_mapped_structure(parts[0])

            merged = _merge_structure_maps(parts)
            if merged:
                _elapsed = time.time() - _map_t0
                print(f"{Fore.CYAN}   🧩 Estrutura global consolidada (chunks merged) em {_elapsed:.1f}s.")
                return _sanitize_mapped_structure(merged)
            # Fallback: concatena (melhor do que perder estrutura)
            _elapsed = time.time() - _map_t0
            print(f"{Fore.CYAN}🗺️  map_structure OK (concat fallback) em {_elapsed:.1f}s")
            return _sanitize_mapped_structure("\n\n".join(parts).strip())
        except Exception as e:
            _elapsed = time.time() - _map_t0
            print(f"{Fore.YELLOW}⚠️  Falha no mapeamento via {self.provider} após {_elapsed:.1f}s: {e}")

            # Fallback Universal
            if self.openai_client:
                print(f"{Fore.CYAN}🤖 Fallback: Mapeando com OpenAI ({self.openai_model})...")
                try:
                    response = await self.openai_client.chat.completions.create(
                        model=self.openai_model,
                        messages=[
                            {"role": "system", "content": self.PROMPT_MAPEAMENTO.format(transcricao=fallback_sample)}
                        ],
                        max_completion_tokens=10000
                    )
                    _record_openai_usage(response, model=self.openai_model)
                    content = response.choices[0].message.content.replace('```markdown', '').replace('```', '')
                    _elapsed = time.time() - _map_t0
                    print(f"{Fore.GREEN}   ✅ Estrutura mapeada (OpenAI Fallback) em {_elapsed:.1f}s.")
                    return _sanitize_mapped_structure(content)
                except Exception as e_openai:
                    _elapsed = time.time() - _map_t0
                    print(f"{Fore.RED}❌ Falha também no OpenAI após {_elapsed:.1f}s: {e_openai}")
                    return None
            else:
                 _elapsed = time.time() - _map_t0
                 print(f"{Fore.RED}   ❌ Erro ao mapear estrutura e sem fallback ({_elapsed:.1f}s).")
                 return None

    async def _ai_reassign_tables(self, texto: str, *, max_tables: int = 3) -> tuple[str, list[str]]:
        """
        v2.34: Fallback de reatribuição de tabelas via IA (decisão binária PARENT/CURRENT).
        """
        candidates = coletar_candidatos_reatribuicao_tabelas(texto, max_candidates=max_tables)
        if not candidates:
            return texto, []

        async def _decide_one(candidate: dict) -> str:
            prompt = (
                "Você é um revisor de estrutura. Decida se a tabela pertence ao TÓPICO ATUAL "
                "ou ao TÓPICO PAI. Responda apenas com 'PARENT' ou 'CURRENT'.\n\n"
                f"TÓPICO PAI: {candidate['parent_title']}\n"
                f"TÓPICO ATUAL: {candidate['current_title']}\n\n"
                f"CONTEXTO PAI (antes do subtópico): {candidate['parent_context']}\n"
                f"CONTEXTO ATUAL (antes da tabela): {candidate['current_context']}\n\n"
                "TABELA:\n"
                f"{candidate['table_text'][:3500]}\n"
            )
            if self.provider == "openai" and self.client:
                response = await self.client.chat.completions.create(
                    model=self.llm_model,
                    messages=[{"role": "system", "content": prompt}],
                    max_completion_tokens=128,
                )
                content = response.choices[0].message.content or ""
                return content.strip().upper()

            from google.genai import types
            def call_gemini():
                return self.client.models.generate_content(
                    model=self.llm_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        max_output_tokens=256,
                        thinking_config=types.ThinkingConfig(
                            include_thoughts=False,
                            thinking_level="LOW",
                        ),
                    ),
                )
            response = await asyncio.to_thread(call_gemini)
            _record_genai_usage(response, model=self.llm_model)
            return (response.text or "").strip().upper()

        lines = texto.split("\n")
        moves: list[dict] = []
        issues: list[str] = []
        for candidate in candidates:
            decision = await _decide_one(candidate)
            if "PARENT" in decision:
                moves.append(candidate)
                issues.append(
                    f"Tabela reatribuída via IA: '{candidate['current_title']}' → '{candidate['parent_title']}'"
                )

        if not moves:
            return texto, []

        moves.sort(key=lambda m: m["start"], reverse=True)
        for m in moves:
            block_lines = lines[m["start"]:m["end"]]
            del lines[m["start"]:m["end"]]
            insert_at = m["insert_at"]
            if insert_at > m["start"]:
                insert_at = max(0, insert_at - (m["end"] - m["start"]))
            for offset, bl in enumerate(block_lines):
                lines.insert(insert_at + offset, bl)
        return "\n".join(lines), issues

    def resolve_diarization_policy(
        self,
        mode: str,
        *,
        diarization: Optional[bool] = None,
        diarization_strict: Optional[bool] = None,
    ) -> tuple[bool, bool]:
        """
        Resolve política de diarização por modo.

        - `AUDIENCIA`/`REUNIAO`/`DEPOIMENTO`: diarização ON por padrão e STRICT (falha se indisponível).
        - `APOSTILA`/`FIDELIDADE`: diarização OFF por padrão; opt-in por configuração.

        Opt-in env para apostilas:
        - `IUDEX_ENABLE_DIARIZATION_APOSTILA=1` (ou `ENABLE_DIARIZATION_APOSTILA=1`)

        Strictness:
        - `--diarization` (forçar ON) torna strict por padrão.
        - `IUDEX_DIARIZATION_STRICT=1` pode forçar strict quando habilitado por env.
        """
        mode_norm = (mode or "").strip().upper()

        if diarization is None:
            apostila_opt_in = bool(
                _env_truthy("IUDEX_ENABLE_DIARIZATION_APOSTILA", False)
                or _env_truthy("ENABLE_DIARIZATION_APOSTILA", False)
            )
            diarization_enabled = mode_norm in {"AUDIENCIA", "REUNIAO", "DEPOIMENTO"} or (
                mode_norm in {"APOSTILA", "FIDELIDADE"} and apostila_opt_in
            )
        else:
            diarization_enabled = bool(diarization)

        if not diarization_enabled:
            return False, False

        strict_env = _env_truthy("IUDEX_DIARIZATION_STRICT", None)
        if diarization_strict is not None:
            diarization_required = bool(diarization_strict)
        elif diarization is True:
            diarization_required = True
        elif mode_norm in {"AUDIENCIA", "REUNIAO", "DEPOIMENTO"}:
            diarization_required = True
        elif strict_env is not None:
            diarization_required = bool(strict_env)
        else:
            diarization_required = False

        return True, diarization_required

    def set_diarization_policy(self, *, enabled: bool, required: bool) -> None:
        self._diarization_enabled = bool(enabled)
        self._diarization_required = bool(required) if enabled else False
        if self._diarization_enabled:
            strict_label = "STRICT" if self._diarization_required else "SOFT"
            print(f"{Fore.CYAN}🗣️  Diarização: ATIVA ({strict_label})")
        else:
            print(f"{Fore.CYAN}🗣️  Diarização: DESATIVADA")

    def _get_hf_token(self) -> Optional[str]:
        return (os.getenv("HUGGING_FACE_TOKEN") or os.getenv("HF_TOKEN") or "").strip() or None

    def _diarization_available(self) -> tuple[bool, str]:
        # Se AssemblyAI está configurado como primário, diarização é feita externamente
        aai_primary = os.getenv("ASSEMBLYAI_PRIMARY", "").strip().lower() in ("1", "true", "yes")
        aai_key = os.getenv("ASSEMBLYAI_API_KEY", "").strip()
        if aai_primary and aai_key:
            return True, "AssemblyAI (externo)"
        # Verificar diarização local (pyannote)
        if Pipeline and self._get_hf_token():
            return True, "pyannote (local)"
        # Fallback: RunPod diarize endpoint
        runpod_diarize = os.getenv("RUNPOD_DIARIZE_ENDPOINT_ID", "").strip()
        runpod_key = os.getenv("RUNPOD_API_KEY", "").strip()
        if runpod_diarize and runpod_key:
            return True, "RunPod (externo)"
        # Fallback: AssemblyAI (mesmo não sendo primário, pode fazer diarização)
        if aai_key:
            return True, "AssemblyAI (externo, não-primário)"
        # Nenhum provider disponível
        if not Pipeline:
            return False, "pyannote.audio não instalado"
        if not self._get_hf_token():
            return False, "HUGGING_FACE_TOKEN/HF_TOKEN não configurado"
        return False, "nenhum provider de diarização disponível"

    def _local_diarization_available(self) -> bool:
        """Verifica se diarização LOCAL (pyannote) está disponível."""
        return bool(Pipeline and self._get_hf_token())

    def _ensure_diarization_available_or_raise(self) -> None:
        ok, reason = self._diarization_available()
        if ok:
            return
        if self._diarization_required:
            raise RuntimeError(
                "Diarização obrigatória, mas indisponível. "
                f"Motivo: {reason}. "
                "Instale `pyannote.audio` e `torch` e configure `HUGGING_FACE_TOKEN`."
            )

    def transcribe_file(
        self,
        audio_path: str,
        *,
        mode: str = "APOSTILA",
        high_accuracy: bool = False,
        diarization: Optional[bool] = None,
        diarization_strict: Optional[bool] = None,
        language: Optional[str] = None,
    ) -> str:
        """
        Transcrição com política de diarização por modo (ponto único de entrada).
        Retorna apenas o texto. Use transcribe_file_full() para obter words também.
        """
        result = self.transcribe_file_full(
            audio_path,
            mode=mode,
            high_accuracy=high_accuracy,
            diarization=diarization,
            diarization_strict=diarization_strict,
            language=language,
        )
        return result["text"]

    def transcribe_file_full(
        self,
        audio_path: str,
        *,
        mode: str = "APOSTILA",
        high_accuracy: bool = False,
        diarization: Optional[bool] = None,
        diarization_strict: Optional[bool] = None,
        language: Optional[str] = None,
    ) -> dict:
        """
        Transcrição com política de diarização por modo.
        Retorna dict com: text, words, segments.

        Returns:
            dict: {
                "text": str,           # Texto formatado com timestamps a cada 60s
                "words": list,         # Lista de {word, start, end, speaker} para player
                "segments": list,      # Segmentos originais
            }
        """
        # Mantém o modo atual
        self._current_mode = (mode or "FIDELIDADE").strip().upper()
        self._current_language = (language or "pt").strip().lower()
        enabled, required = self.resolve_diarization_policy(
            mode, diarization=diarization, diarization_strict=diarization_strict
        )
        self.set_diarization_policy(enabled=enabled, required=required)

        if enabled:
            self._ensure_diarization_available_or_raise()
            if self._local_diarization_available():
                # Diarização local via pyannote
                if high_accuracy:
                    return self.transcribe_beam_with_segments(audio_path)
                return self.transcribe_with_segments(audio_path)
            else:
                # Diarização disponível externamente (RunPod/AAI) — transcrever sem diarização local,
                # o chamador (transcription_service) fará a diarização via provider externo
                print(f"{Fore.YELLOW}⚠️  pyannote indisponível localmente — diarização será feita externamente")
                original_diarization_enabled = self._diarization_enabled
                try:
                    self._diarization_enabled = False
                    if high_accuracy:
                        result = self.transcribe_with_segments(audio_path, beam_size=self._get_asr_beam_size())
                    else:
                        result = self.transcribe_with_segments(audio_path)
                    result["_needs_external_diarization"] = True
                    return result
                finally:
                    self._diarization_enabled = original_diarization_enabled

        # Sem diarização: ainda precisamos obter words para o player
        # Usar transcribe_with_segments mas forçar diarização desabilitada
        original_diarization_enabled = self._diarization_enabled
        try:
            self._diarization_enabled = False  # Forçar desabilitado para não rodar pyannote
            if high_accuracy:
                result = self.transcribe_with_segments(audio_path, beam_size=self._get_asr_beam_size())
            else:
                result = self.transcribe_with_segments(audio_path)
            return result
        finally:
            self._diarization_enabled = original_diarization_enabled

    def renumber_headings(self, text):
        """
        Post-processing: Enforces strictly sequential numbering (1, 2, 3...)
        for H2/H3/H4 headers using a STACK-BASED approach for correct nesting.

        v2.16: Fixed to properly reset child counters when parent level changes.
        v2.41: Added semantic title merge (SequenceMatcher) to fuse near-duplicates
               from chunk boundaries — prevents title inflation.
        """
        from difflib import SequenceMatcher

        print(f"{Fore.CYAN}🔢 Renumerando tópicos sequencialmente (Stack-Based v2.41)...")
        lines = text.split('\n')
        new_lines = []

        # Stack-based counters: [H1_count, H2_count, H3_count, H4_count]
        # Index 0 = H1 (usually title, skip), Index 1 = H2, etc.
        counters = [0, 0, 0, 0, 0]  # Extra slot for safety

        # Keywords to skip numbering (summary tables, etc.)
        skip_keywords = ['resumo', 'quadro', 'tabela', 'síntese', 'esquema', 'bibliografia', 'referências', 'sumário']

        # Emoji pattern to detect decorative headers like "## 📋 Sumário"
        emoji_pattern = re.compile(r'^[\U0001F300-\U0001F9FF]')
        seen_h2_numbers = set()
        level_adjustments = 0

        # v2.41: Semantic merge tracking
        last_h2_text = ""
        last_h3_text = ""
        merge_count = 0

        for line in lines:
            stripped = line.strip()
            
            # Determine header level
            header_match = re.match(r'^(#{1,4})\s+(.*)$', stripped)
            
            if header_match:
                hashes = header_match.group(1)
                level = len(hashes)  # 1 for H1, 2 for H2, etc.
                raw_title = header_match.group(2).strip()

                # Heurística determinística (v2.17):
                # Se o header já contém numeração explícita (ex.: "5.5."), mas veio em nível errado (ex.: "## 5.5"),
                # ajusta o nível para preservar a hierarquia esperada antes de renumerar sequencialmente.
                #
                # Ex.: "## 5.5. Subtópico" deve ser tratado como H3, para virar "### 5.5." após a renumeração stack-based.
                num_match = re.match(r'^(\d+(?:\.\d+)*)(?:\.)?\s+', raw_title)
                if num_match:
                    depth = num_match.group(1).count(".") + 1
                    desired_level = level
                    # Regra: títulos com numeração decimal devem ser subtópicos quando já houve H2.
                    if depth == 1 and level > 2:
                        desired_level = 2
                    elif depth >= 2:
                        # Evitar criar "0.x" quando não há H2 prévio.
                        if counters[2] > 0:
                            desired_level = 3 if depth == 2 else 4
                            # Se não houver H3 prévio, evita H4 direto.
                            if desired_level == 4 and counters[3] == 0:
                                desired_level = 3
                    if desired_level != level and desired_level in (2, 3, 4):
                        level = desired_level
                        hashes = "#" * level
                        level_adjustments += 1
                
                # Clean existing numbers from title (e.g., "1.2.3. Title" -> "Title")
                title_text = re.sub(r'^(\d+(\.\d+)*\.?\s*)+', '', raw_title).strip()

                # Skip H1 (document title) - just clean and pass through
                if level == 1:
                    new_lines.append(f"# {title_text}")
                    continue

                # v2.41: Semantic merge — fuse near-duplicate titles from chunk boundaries
                title_norm = re.sub(r'\s*\(Continuação\)\s*$', '', title_text, flags=re.IGNORECASE).strip().lower()
                if level == 2 and last_h2_text:
                    ratio = SequenceMatcher(None, title_norm, last_h2_text).ratio()
                    if ratio > 0.85:
                        merge_count += 1
                        continue  # skip duplicate — content flows under existing H2
                if level == 3 and last_h3_text:
                    ratio = SequenceMatcher(None, title_norm, last_h3_text).ratio()
                    if ratio > 0.85:
                        merge_count += 1
                        continue
                if level == 2:
                    last_h2_text = title_norm
                    last_h3_text = ""  # reset H3 tracker when H2 changes
                elif level == 3:
                    last_h3_text = title_norm

                # Check if this header should be skipped from numbering
                title_lower = title_text.lower()
                should_skip = (
                    any(k in title_lower for k in skip_keywords) or
                    emoji_pattern.match(title_text)  # Headers starting with emoji
                )

                if should_skip:
                    new_lines.append(f"{'#' * level} {title_text}")
                else:
                    # STACK LOGIC: Increment current level, reset all deeper levels
                    counters[level] += 1
                    for deeper_level in range(level + 1, len(counters)):
                        counters[deeper_level] = 0
                    
                    # Build hierarchical number (e.g., "2.3.1")
                    number_parts = [str(counters[lvl]) for lvl in range(2, level + 1)]
                    hierarchical_number = ".".join(number_parts)
                    
                    new_lines.append(f"{'#' * level} {hierarchical_number}. {title_text}")
                    if level == 2:
                        seen_h2_numbers.add(str(counters[2]))
            else:
                new_lines.append(line)
        
        if level_adjustments:
            print(f"{Fore.YELLOW}   ℹ️  Ajustes de nível aplicados: {level_adjustments}")
        if merge_count:
            print(f"{Fore.YELLOW}   🔄 Títulos duplicados mesclados: {merge_count}")
        print(f"{Fore.GREEN}   ✅ Renumeração concluída: {counters[2]} seções H2, {counters[3]} H3, {counters[4]} H4")
        return '\n'.join(new_lines)


    def check_coverage(self, original, formatted):
        """Checks for missing laws or sumulas using ROBUST fingerprints (v2.10 Ported)."""
        print(f"{Fore.YELLOW}🔍 Verificando fidelidade (Leis/Súmulas/Robust)...")
        
        # Use migrated robust helpers
        fp_original = extrair_fingerprints(original)
        contagem_original = contar_ocorrencias_robust(fp_original, original)
        contagem_formatado = contar_ocorrencias_robust(fp_original, formatted)
        
        omissoes = []
        duplicacoes = []
        
        for key, count_orig in contagem_original.items():
            count_fmt = contagem_formatado.get(key, 0)
            categoria, item = key.split(':', 1)
            
            if count_orig > 0 and count_fmt == 0:
                omissoes.append(f"[{categoria}] {item}")
            if count_fmt > count_orig:
                duplicacoes.append(f"[{categoria}] {item} (+{count_fmt - count_orig})")
        
        report = []
        if omissoes:
            report.append(f"⚠️ {len(omissoes)} POSSÍVEIS OMISSÕES:")
            for o in omissoes[:15]: report.append(f"   - {o}")
            if len(omissoes) > 15: report.append(f"   ... e mais {len(omissoes)-15}")
            
        if duplicacoes:
            report.append(f"\nℹ️ {len(duplicacoes)} CITAÇÕES REFORÇADAS (Agregadas):")
            for d in duplicacoes[:10]: report.append(f"   - {d}")

        if not report:
            print(f"{Fore.GREEN}✅ Verificação OK: Nenhuma omissão de Leis/Súmulas detectada.")
            return "Verificação OK: Nenhuma omissão detectada."
        else:
            msg = "\n".join(report)
            print(f"{Fore.RED}{msg}")
            return msg

    def final_structure_audit(self, formatted_text, global_structure):
        """
        v2.16: Audita a estrutura final comparando com o mapeamento inicial.
        Retorna um relatório de discrepâncias e, opcionalmente, tenta corrigir.
        """
        print(f"{Fore.CYAN}🔍 Auditoria Final de Estrutura (v2.16)...")
        
        if not global_structure:
            print(f"{Fore.YELLOW}   ⚠️ Mapeamento global não disponível, pulando auditoria.")
            return formatted_text, []
        
        # Extract all H2/H3 titles from formatted text
        formatted_headers = []
        for line in formatted_text.split('\n'):
            match = re.match(r'^(#{2,3})\s+(?:\d+(?:\.\d+)*\.?\s*)?(.+)$', line.strip())
            if match:
                level = len(match.group(1))
                title = match.group(2).strip()
                # Normalize
                title_normalized = re.sub(r'\s*\(Continuação\)\s*$', '', title, flags=re.IGNORECASE).strip().lower()
                formatted_headers.append((level, title_normalized, title))
        
        # Extract expected structure from global mapping
        # v2.41: Support both markdown (## Title) and numbered (1. Title, 1.1. Title) formats
        expected_headers = []
        for line in global_structure.split('\n'):
            stripped = line.strip()
            # Try markdown format first: ## Title, ### Title
            md_match = re.match(r'^(#{2,3})\s+(.+)$', stripped)
            if md_match:
                level = len(md_match.group(1))
                title = md_match.group(2).strip()
                # Remove ABRE/FECHA anchors if present
                title = re.sub(r'\s*\|\s*(?:ABRE|FECHA):\s*["\'][^"\']*["\']', '', title).strip()
                title_normalized = title.lower()
                expected_headers.append((level, title_normalized, title))
                continue
            # Try numbered format: 1. Title (level 2), 1.1. Title (level 3), 1.1.1. Title (level 4)
            num_match = re.match(r'^(\d+(?:\.\d+)*)\.\s+(.+)$', stripped)
            if num_match:
                depth = num_match.group(1).count('.') + 1  # 1. = depth 1 → H2, 1.1. = depth 2 → H3
                level = min(depth + 1, 4)  # map to markdown level: depth 1→H2, depth 2→H3, depth 3→H4
                title = num_match.group(2).strip()
                # Remove ABRE/FECHA anchors if present
                title = re.sub(r'\s*\|\s*(?:ABRE|FECHA):\s*["\'][^"\']*["\']', '', title).strip()
                title_normalized = title.lower()
                if level <= 3:  # Only compare H2 and H3 for audit
                    expected_headers.append((level, title_normalized, title))
        
        # Find duplicates in formatted headers
        seen_titles = {}
        duplicates = []
        for idx, (level, title_norm, title_orig) in enumerate(formatted_headers):
            if title_norm in seen_titles:
                duplicates.append({
                    'title': title_orig,
                    'first_occurrence': seen_titles[title_norm],
                    'duplicate_occurrence': idx + 1
                })
            else:
                seen_titles[title_norm] = idx + 1
        
        issues = []
        if duplicates:
            issues.append(f"⚠️ {len(duplicates)} TÍTULOS DUPLICADOS DETECTADOS:")
            for d in duplicates[:5]:
                issues.append(f"   - '{d['title']}' (linhas ~{d['first_occurrence']} e ~{d['duplicate_occurrence']})")
        
        # Check coverage of expected structure
        expected_titles_set = {h[1] for h in expected_headers}
        formatted_titles_set = {h[1] for h in formatted_headers}
        missing_titles = expected_titles_set - formatted_titles_set
        
        if missing_titles:
            issues.append(f"\n⚠️ {len(missing_titles)} TÓPICOS DO MAPEAMENTO NÃO ENCONTRADOS:")
            for t in list(missing_titles)[:5]:
                issues.append(f"   - '{t}'")
        
        if issues:
            report = "\n".join(issues)
            print(f"{Fore.YELLOW}{report}")
            return formatted_text, issues
        else:
            print(f"{Fore.GREEN}   ✅ Estrutura auditada - Sem duplicatas ou omissões de tópicos.")
            return formatted_text, []

    async def format_transcription_async(
        self,
        transcription,
        video_name,
        output_folder,
        mode="APOSTILA",
        custom_prompt=None,
        custom_prompt_scope: str = "tables_only",
        dry_run=False,
        progress_callback=None,
        skip_audit=False,
        skip_fidelity_audit=False,
        skip_sources_audit=False,
        hil_strict=False,
        include_timestamps: bool = True,
        allow_indirect: bool = False,
        allow_summary: bool = False,
        disable_tables: bool = False,
        segment_timeout_seconds: Optional[int] = None,
    ):
        """
        Orquestrador Principal com Checkpoint e Robustez (Sequential Mode)
        
        Args:
            transcription: Texto da transcrição
            video_name: Nome do vídeo
            output_folder: Pasta de saída
            mode: "APOSTILA", "FIDELIDADE", "AUDIENCIA", "REUNIAO" ou "DEPOIMENTO"
            custom_prompt: Campo de customização opcional (controlado por custom_prompt_scope).
            custom_prompt_scope: 'tables_only' (padrão) → afeta apenas tabelas/extras em TODOS os modos;
                                 'style_and_tables' (avançado, opt-in) → substitui STYLE+TABLE layers.
            dry_run: Se True, apenas valida divisão de chunks
            skip_audit: Se True, pula a auditoria jurídica
            skip_sources_audit: Se True, pula a auditoria de fontes integrada
            allow_indirect: Se True, permite discurso indireto em AUDIENCIA/REUNIAO/DEPOIMENTO.
            allow_summary: Se True, permite ata resumida em AUDIENCIA/REUNIAO/DEPOIMENTO.
        """
        async def emit(stage: str, progress: int, message: str):
            if not progress_callback:
                return
            try:
                result = progress_callback(stage, progress, message)
                if asyncio.iscoroutine(result):
                    await result
            except Exception:
                pass

        # v2.23: Store mode for dynamic file naming
        self._current_mode = mode.upper()
        mode_suffix = self._current_mode
        print(f"{Fore.MAGENTA}🧠 Formatando com {self.llm_model} (Sequential Mode)...")
        
        # v2.22: Modular prompt composition
        # custom_prompt now only overrides the STYLE+TABLE layers, not the entire system prompt.
        self.prompt_apostila = self._build_system_prompt(
            mode=mode,
            custom_style_override=custom_prompt,
            custom_prompt_scope=custom_prompt_scope,
            disable_tables=bool(disable_tables),
            allow_indirect=allow_indirect,
            allow_summary=allow_summary,
            include_timestamps=include_timestamps,
        )
        
        if not custom_prompt:
            mode_label = (mode or "APOSTILA").upper()
            pretty_map = {
                "APOSTILA": "APOSTILA",
                "FIDELIDADE": "FIDELIDADE",
                "AUDIENCIA": "AUDIÊNCIA",
                "REUNIAO": "REUNIÃO",
                "DEPOIMENTO": "DEPOIMENTO",
            }
            pretty = pretty_map.get(mode_label, mode_label)
            icon = "📚" if mode_label == "APOSTILA" else "🎨"
            print(f"{Fore.CYAN}{icon} Modo {pretty} ativo (Prompt modular)")
        
        # 0. Context Extraction
        pass
        # professors_info = self._extract_professors_context(transcription)

        await emit("formatting", 60, "Iniciando formatação...")

        # 0.1 Global Structure Mapping (NEW) — with heartbeat to prevent stall appearance
        _map_hb_done = asyncio.Event()
        _map_hb_start = time.time()
        async def _map_heartbeat():
            while not _map_hb_done.is_set():
                try:
                    await asyncio.wait_for(_map_hb_done.wait(), timeout=8)
                except asyncio.TimeoutError:
                    elapsed = time.time() - _map_hb_start
                    await emit("formatting", 64, f"Mapeando estrutura... ({elapsed:.0f}s)")
        _map_hb_task = asyncio.create_task(_map_heartbeat())
        try:
            global_structure = await self.map_structure(transcription)
        finally:
            _map_hb_done.set()
            _map_hb_task.cancel()
            try:
                await _map_hb_task
            except (asyncio.CancelledError, Exception):
                pass
        await emit("formatting", 68, "Estrutura global mapeada")

        # v2.41: Pré-filtro e separação cut vs hierarchy
        if global_structure:
            global_structure = filtrar_niveis_excessivos(global_structure, max_nivel=3)
            simplify_max_lines = _safe_int(os.getenv("IUDEX_MAP_SIMPLIFY_MAX_LINES")) or 120
            simplify_max_depth = _safe_int(os.getenv("IUDEX_MAP_SIMPLIFY_MAX_DEPTH")) or 3
            global_structure = simplificar_estrutura_se_necessario(
                global_structure,
                max_linhas=simplify_max_lines,
                max_nivel=simplify_max_depth,
            )

        # v2.42: Detectar tipo de conteúdo e injetar addon SIMULADO se aplicável
        _tipo_match = re.search(r'\[TIPO:\s*(SIMULADO|CORREÇÃO|CORRECAO)\]', global_structure or '', re.IGNORECASE)
        if _tipo_match:
            _tipo = _tipo_match.group(1).upper()
            print(f"{Fore.MAGENTA}🎯 Tipo detectado: {_tipo} — Injetando regras de questões/simulado no prompt")
            self.prompt_apostila += self.PROMPT_SIMULADO_ADDON
        # Estrutura limpa (sem ABRE/FECHA) para guiar hierarquia nos chunks
        hierarchy_structure = limpar_estrutura_para_review(global_structure) if global_structure else None

        # 1. Sequential Slicing (v2.17: Com âncoras de estrutura)
        mode_norm = (mode or "APOSTILA").strip().upper()
        print(f"🔪 Dividindo em chunks (v2.32)...")

        # Para AUDIÊNCIA/REUNIÃO/DEPOIMENTO: preferir chunking por blocos naturais (## Bloco XX — ...),
        # evitando cortes no meio de um turno/ato.
        chunks_info = []
        if mode_norm in {"AUDIENCIA", "REUNIAO", "DEPOIMENTO"}:
            block_max_chars = int(os.getenv("IUDEX_HEARING_BLOCK_MAX_CHARS", 25000))
            block_overlap = int(os.getenv("IUDEX_HEARING_BLOCK_SPLIT_OVERLAP_CHARS", 300))
            block_prefix = os.getenv("IUDEX_HEARING_BLOCK_PREFIX_REGEX", None)
            chunks_info = dividir_por_blocos_markdown(
                transcription,
                max_chars=block_max_chars,
                block_prefix_pattern=block_prefix,
                split_overlap_chars=block_overlap,
            )

        # Fallback: slicing sequencial com âncoras (aulas/apostilas e quando não há blocos).
        if not chunks_info:
            print(f"   ℹ️  Usando divisão sequencial (com âncoras v2.17)...")
            chunks_info = dividir_sequencial(transcription, chars_por_parte=15000, estrutura_global=global_structure)
        validar_chunks(chunks_info, transcription)
        
        total_segments = len(chunks_info)
        print(f"📊 Total de segmentos sequenciais: {total_segments}")
        await emit("formatting", 72, f"{total_segments} segmentos preparados")
        
        if dry_run:
            print(f"{Fore.YELLOW}🔍 MODO DRY-RUN: Parando antes das chamadas de API.")
            print(f"   Exemplo do Chunk 1: {transcription[chunks_info[0]['inicio']:chunks_info[0]['inicio']+100]}...")
            return "# DRY RUN OUTPUT"
        
        # 2. Checkpoint Loading
        checkpoint_data = load_checkpoint(video_name, output_folder)
        results_map = {} # Map idx -> result
        
        if checkpoint_data:
            print(f"{Fore.CYAN}📂 Retomando via Checkpoint ({checkpoint_data.get('timestamp')})")
            if len(checkpoint_data.get('results', [])) > 0:
                saved_results = checkpoint_data['results']
                # Restore results map
                for idx, res in enumerate(saved_results):
                    if idx < total_segments:
                        results_map[idx] = res
                print(f"   ✅ {len(results_map)} segmentos recuperados.")
        
        # v2.19: Context Caching Setup
        cached_context = None
        if total_segments > 1: # Só cache se tiver múltiplos chunks
            # v2.41: cache deve receber estrutura limpa para orientar hierarquia (H2/H3),
            # mantendo ABRE/FECHA apenas para o corte de chunks.
            cached_context = self.create_context_cache(
                transcription,
                hierarchy_structure or global_structure,
            )
                
        # 3. Sequential Processing Loop
        ordered_results = []
        
        # Restore ordered results from map
        for i in range(len(results_map)):
            ordered_results.append(results_map[i])

        start_idx = len(ordered_results)
        
        # v2.40: Helper function for processing a single chunk
        async def _process_single_chunk(i: int, prev_result: Optional[str] = None) -> str:
            """Process a single chunk with optional context from previous result."""
            info = chunks_info[i]
            chunk_text = transcription[info['inicio']:info['fim']]
            raw_overlap_chars = int(
                os.getenv("IUDEX_RAW_CONTEXT_OVERLAP_CHARS", self.RAW_CONTEXT_OVERLAP_CHARS)
            )
            overlap_raw = ""
            if raw_overlap_chars > 0 and info.get("inicio", 0) > 0:
                start_overlap = max(0, info["inicio"] - raw_overlap_chars)
                overlap_raw = transcription[start_overlap:info["inicio"]].strip()

            # Context Management - use previous result if available
            contexto_estilo = ""
            if prev_result:
                raw_context = _extract_style_context(prev_result, max_chars=2500)
                if len(raw_context.split()) > 30 and "[!WARNING]" not in raw_context:
                    contexto_estilo = raw_context

            # Rate Limit
            await rate_limiter.wait_if_needed_async()

            # Lógica de Estrutura Local (Janela Deslizante)
            # v2.41: Usa hierarchy_structure (sem ABRE/FECHA) para guiar H2/H3
            estrutura_referencia = None
            _struct_source = hierarchy_structure or global_structure
            if _struct_source and not cached_context:
                max_lines = int(
                    os.getenv(
                        "IUDEX_MAP_MAX_LINES_PER_CHUNK",
                        self.MAP_MAX_LINES_PER_CHUNK,
                    )
                )
                itens_estrutura = [ln for ln in _struct_source.split('\n') if ln.strip()]
                if len(itens_estrutura) > 8 and total_segments > 1:
                    ratio = len(itens_estrutura) / total_segments
                    center_idx = int(i * ratio)
                    if len(itens_estrutura) > max_lines:
                        available = max(4, max_lines - 2)
                        half = max(2, available // 2)
                        start_idx_w = max(0, center_idx - half)
                        end_idx_w = min(len(itens_estrutura), start_idx_w + available)
                        start_idx_w = max(0, end_idx_w - available)
                        slice_itens = itens_estrutura[start_idx_w:end_idx_w]
                        if start_idx_w > 0:
                            slice_itens.insert(0, "[... Tópicos anteriores ...]")
                        if end_idx_w < len(itens_estrutura):
                            slice_itens.append("[... Tópicos posteriores ...]")
                        estrutura_referencia = '\n'.join(slice_itens)
                    else:
                        window_size = max(4, int(len(itens_estrutura) * 0.15))
                        start_idx_w = max(0, center_idx - window_size)
                        end_idx_w = min(len(itens_estrutura), center_idx + window_size + 2)
                        slice_itens = itens_estrutura[start_idx_w:end_idx_w]
                        if start_idx_w > 0:
                            slice_itens.insert(0, "[... Tópicos anteriores ...]")
                        if end_idx_w < len(itens_estrutura):
                            slice_itens.append("[... Tópicos posteriores ...]")
                        estrutura_referencia = '\n'.join(slice_itens)
                else:
                    estrutura_referencia = _struct_source

            # v2.26: Contexto de continuidade
            continuidade_nota = ""
            if info.get('instituto_continua') and info.get('instituto_nome'):
                continuidade_nota = f"\n\n⚠️ AVISO: O instituto '{info['instituto_nome']}' continua no próximo chunk. NÃO gere o Quadro-síntese final ainda — ele será completado na próxima parte."

            # v2.27: Tabela aberta
            tabela_aberta_nota = ""
            if prev_result:
                table_state = self._detect_open_table_state(prev_result)
                if table_state.get('needs_table_continuation'):
                    tabela_aberta_nota = table_state.get('context_hint', '')

            contexto_final = contexto_estilo + continuidade_nota + tabela_aberta_nota

            formatted = await self.process_chunk_async(
                chunk_text=chunk_text,
                idx=i+1,
                total=total_segments,
                previous_context=contexto_final,
                depth=0,
                global_structure=estrutura_referencia,
                overlap_text=overlap_raw,
                cached_content=cached_context
            )
            return formatted

        # v2.40: Parallel or Sequential Processing
        parallel_chunks = int(os.getenv("IUDEX_PARALLEL_CHUNKS", self.PARALLEL_CHUNKS))
        try:
            heartbeat_every = float(os.getenv("IUDEX_PROGRESS_HEARTBEAT_SECONDS", "12"))
        except Exception:
            heartbeat_every = 12.0
        if segment_timeout_seconds is None:
            try:
                segment_timeout_seconds = int(os.getenv("IUDEX_FORMAT_SEGMENT_TIMEOUT_SECONDS", "0"))
            except Exception:
                segment_timeout_seconds = 0
        else:
            try:
                segment_timeout_seconds = int(segment_timeout_seconds)
            except Exception:
                segment_timeout_seconds = 0

        if start_idx < total_segments:
            if parallel_chunks <= 1 or total_segments - start_idx <= 2:
                # Sequential mode (original behavior)
                print(f"▶ Iniciando processamento sequencial do segmento {start_idx + 1}...")

                for i in tqdm(range(start_idx, total_segments), desc="Processando Sequencial"):
                    base_msg = f"Formatando segmento {i+1}/{total_segments}..."
                    display_progress = 72
                    if total_segments:
                        progress = 72 + int(((i) / total_segments) * 23)
                        display_progress = min(progress, 95)
                        await emit("formatting", display_progress, base_msg)

                    hb_done = asyncio.Event()
                    hb_task = None
                    hb_start = time.time()
                    if heartbeat_every and heartbeat_every > 0:
                        async def _heartbeat():
                            while not hb_done.is_set():
                                try:
                                    await asyncio.wait_for(hb_done.wait(), timeout=heartbeat_every)
                                except asyncio.TimeoutError:
                                    elapsed = time.time() - hb_start
                                    msg = f"{base_msg} ({elapsed:.0f}s)"
                                    await emit("formatting", display_progress, msg)
                        hb_task = asyncio.create_task(_heartbeat())

                    try:
                        prev_result = ordered_results[-1] if ordered_results else None
                        if segment_timeout_seconds and segment_timeout_seconds > 0:
                            formatted = await asyncio.wait_for(
                                _process_single_chunk(i, prev_result),
                                timeout=segment_timeout_seconds,
                            )
                        else:
                            formatted = await _process_single_chunk(i, prev_result)

                        # Smart stitching
                        if ordered_results:
                            try:
                                formatted = limpar_inicio_redundante(formatted, ordered_results[-1])
                            except Exception:
                                pass

                        ordered_results.append(formatted)
                        if total_segments:
                            progress = 72 + int(((i + 1) / total_segments) * 23)
                            await emit("formatting", min(progress, 95), f"Segmento {i+1}/{total_segments} concluído")

                        save_checkpoint(video_name, output_folder, ordered_results, chunks_info, i + 1)

                    except Exception as e:
                        print(f"{Fore.RED}❌ Falha Fatal no segmento {i+1}: {e}")
                        save_checkpoint(video_name, output_folder, ordered_results, chunks_info, i)
                        raise e
                    finally:
                        hb_done.set()
                        if hb_task:
                            try:
                                await hb_task
                            except Exception:
                                pass
            else:
                # Parallel mode (v2.40): Process in batches with semaphore
                print(f"▶ Iniciando processamento PARALELO ({parallel_chunks} workers) do segmento {start_idx + 1}...")
                semaphore = asyncio.Semaphore(parallel_chunks)

                async def process_with_semaphore(idx: int, prev_res: Optional[str]) -> tuple:
                    async with semaphore:
                        try:
                            if segment_timeout_seconds and segment_timeout_seconds > 0:
                                result = await asyncio.wait_for(
                                    _process_single_chunk(idx, prev_res),
                                    timeout=segment_timeout_seconds,
                                )
                            else:
                                result = await _process_single_chunk(idx, prev_res)
                            return (idx, result, None)
                        except Exception as e:
                            return (idx, None, e)

                # Process in waves: first chunk sequential, then batches
                remaining = list(range(start_idx, total_segments))

                while remaining:
                    batch_size = min(parallel_chunks, len(remaining))
                    batch = remaining[:batch_size]
                    remaining = remaining[batch_size:]

                    await emit("formatting", 72 + int((total_segments - len(remaining) - batch_size) / total_segments * 23),
                              f"Processando batch de {len(batch)} segmentos...")

                    # First chunk in batch gets context from previous result
                    first_idx = batch[0]
                    prev_result = ordered_results[-1] if ordered_results else None

                    # Process batch in parallel
                    tasks = []
                    for j, idx in enumerate(batch):
                        # Only first chunk gets previous context for better stitching
                        ctx = prev_result if j == 0 else None
                        tasks.append(process_with_semaphore(idx, ctx))

                    results = await asyncio.gather(*tasks)

                    # Sort by index and append
                    results_sorted = sorted(results, key=lambda x: x[0])
                    for idx, result, error in results_sorted:
                        if error:
                            print(f"{Fore.RED}❌ Falha no segmento {idx+1}: {error}")
                            save_checkpoint(video_name, output_folder, ordered_results, chunks_info, idx)
                            raise error

                        # Apply stitching
                        if ordered_results:
                            try:
                                result = limpar_inicio_redundante(result, ordered_results[-1])
                            except Exception:
                                pass

                        ordered_results.append(result)
                        print(f"   ✅ Segmento {idx+1}/{total_segments} concluído")

                    # Checkpoint after batch
                    save_checkpoint(video_name, output_folder, ordered_results, chunks_info, batch[-1] + 1)

                await emit("formatting", 95, f"Todos os {total_segments} segmentos processados")
        
        await emit("formatting", 96, "Consolidando resultados...")

        # 4. Final Assembly
        print(f"\n{Fore.CYAN}🧹 Pipeline de Limpeza Final (v2.7)...")
        full_formatted = f"# {video_name}\n\n" + "\n\n".join(ordered_results)
        
        # 4.1 Limpar metadados de mapeamento que vazam para o output
        # Remove linhas como "[TIPO: AULA EXPOSITIVA]" ou "**[TIPO: SIMULADO]**"
        full_formatted = re.sub(r'^#?\s*\*?\*?\[TIPO:.*?\]\*?\*?\s*$', '', full_formatted, flags=re.MULTILINE)
        # Remove marcadores de bloco [BLOCO 01], [BLOCO 02], etc.
        full_formatted = re.sub(r'^\s*\[BLOCO\s*\d+\]\s*$', '', full_formatted, flags=re.MULTILINE)
        # Remove timestamps órfãos [HH:MM] ou [HH:MM:SS] no início de linha
        full_formatted = re.sub(r'^\s*\[\d{1,2}:\d{2}(:\d{2})?\]\s*$', '', full_formatted, flags=re.MULTILINE)
        full_formatted = re.sub(r'\n{3,}', '\n\n', full_formatted)  # Remove linhas em branco extras
        
        # 5. Post-Processing Pipeline (v2.7 features)
        
        print("  Passada 1: Removendo duplicações literais...")
        await emit("formatting", 96, "Passada 1: Removendo duplicações literais...")
        # Uses the newly ported v2.7 logic
        full_formatted = remover_duplicacoes_literais(full_formatted)
        
        limiar_info = '70%' if mode == 'FIDELIDADE' else '60%'
        print(f"  Passada 2: Removendo seções duplicadas (limiar {limiar_info})...")
        await emit("formatting", 96, f"Passada 2: Removendo seções duplicadas ({limiar_info})...")
        full_formatted = remover_secoes_duplicadas(full_formatted, mode=mode)
        
        print("  Passada 2.5: Removendo parágrafos duplicados (v2.17)...")
        await emit("formatting", 97, "Passada 2.5: Removendo parágr. duplicados...")
        full_formatted = remover_paragrafos_duplicados(full_formatted)
        
        print("  Passada 2.6: Removendo títulos órfãos (v2.17)...")
        full_formatted = remover_titulos_orfaos(full_formatted)
        
        print("  Passada 2.7: Mesclando tabelas divididas (v2.27)...")
        await emit("formatting", 97, "Passada 2.7: Mesclando tabelas divididas...")
        full_formatted = mesclar_tabelas_divididas(full_formatted)

        print("  Passada 2.8: Movendo tabelas para fim de seção (v2.41)...")
        await emit("formatting", 97, "Passada 2.8: Reorganizando tabelas...")
        full_formatted = mover_tabelas_para_fim_de_secao(full_formatted)

        print("  Passada 2.9: Ajustando títulos de tabela de banca...")
        await emit("formatting", 97, "Passada 2.9: Ajustando títulos de tabela...")
        full_formatted = garantir_titulo_tabela_banca(full_formatted)
        
        print("  Passada 3: Normalizando títulos similares...")
        await emit("formatting", 97, "Passada 3: Normalizando títulos...")
        full_formatted = normalize_headings(full_formatted)
        
        if mode != "FIDELIDADE":
            print("  Passada 3.5: Reorganização Estrutural Determinística...")
            await emit("formatting", 97, "Passada 3.5: Reorganização Estrutural...")
            full_formatted = deterministic_structure_fix(full_formatted)
        else:
            print(f"{Fore.YELLOW}  ℹ️  Modo FIDELIDADE: Pulando reorganização para preservar linearidade exata.")
        
        title_drift_telemetry = {
            "freeze_h2_h3": False,
            "headers_changed_count": 0,
            "headers_restored_count": 0,
            "headers_degraded_count": 0,
            "headers_diff": [],
        }

        if mode != "FIDELIDADE":
            print("  Passada 4: Revisão Semântica por IA...")
            await emit("formatting", 98, "Passada 4: Revisão Semântica por IA...")
            full_formatted = await ai_structure_review(full_formatted, self.client, self.llm_model, estrutura_mapeada=limpar_estrutura_para_review(global_structure))
        else:
            print(f"{Fore.MAGENTA}  Passada 4: Revisão Leve de Formatação (Modo Fidelidade)...")
            await emit("formatting", 98, "Passada 4: Revisão Leve (Fidelidade)...")
            _fidelity_original_text = full_formatted
            _fidelity_reviewed_text = await ai_structure_review_lite(
                full_formatted,
                self.client,
                self.llm_model,
                estrutura_mapeada=limpar_estrutura_para_review(global_structure),
            )
            full_formatted, title_drift_telemetry = enforce_fidelity_heading_guard(
                _fidelity_original_text,
                _fidelity_reviewed_text,
                freeze_h2_h3=True,
            )
            if title_drift_telemetry.get("headers_restored_count", 0) > 0:
                print(
                    f"{Fore.YELLOW}   ♻️ Heading guard: "
                    f"{title_drift_telemetry['headers_restored_count']} título(s) restaurado(s) "
                    f"de {title_drift_telemetry.get('headers_changed_count', 0)} alterado(s)."
                )
        
        # Passada 4.5: Renumeração Determinística (Camada de Segurança)
        try:
            full_formatted = renumerar_secoes(full_formatted)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Falha na renumeração determinística: {e}. Continuando...")
        
        print(f"\n{Fore.CYAN}🔢 Renumerando tópicos (1..N) (Stack-Based v2.16)...")
        await emit("formatting", 98, "Renumerando tópicos (1..N)...")
        full_formatted = self.renumber_headings(full_formatted)

        # Passada 4.7: Auditoria determinística de hierarquia (subtópicos vs tópicos)
        strict_subtopic_fix = _env_truthy("IUDEX_STRICT_SUBTOPIC_FIX", default=True)
        strict_subtopic_fix = True if strict_subtopic_fix is None else bool(strict_subtopic_fix)
        mode_norm_fix = (mode or "").strip().upper()
        # Em modos de apostila/fidelidade, inconsistência de nível tende a degradar
        # toda a estrutura final; forçamos correção ativa por segurança.
        if mode_norm_fix in {"APOSTILA", "FIDELIDADE"} and not strict_subtopic_fix:
            print(
                f"{Fore.YELLOW}⚠️ IUDEX_STRICT_SUBTOPIC_FIX=0 ignorado para modo {mode_norm_fix}; "
                f"forçando correção de hierarquia.{Style.RESET_ALL}"
            )
            strict_subtopic_fix = True
        fixed_text, level_issues = audit_heading_levels(full_formatted, apply_fixes=strict_subtopic_fix)
        if level_issues:
            print(f"{Fore.YELLOW}⚠️  {len(level_issues)} inconsistências de hierarquia detectadas")
            for issue in level_issues[:5]:
                print(f"{Fore.YELLOW}   - {issue}")
            if strict_subtopic_fix and fixed_text != full_formatted:
                full_formatted = fixed_text
                # Reaplica renumeração para manter sequência coerente após correções de nível.
                full_formatted = self.renumber_headings(full_formatted)
        
        # 5.6 v2.18: Auto-Fix Pass - Correções automáticas finais
        full_formatted, autofix_correcoes = aplicar_correcoes_automaticas(full_formatted, mode=mode)
        
        # 5.5 v2.16: Auditoria Final de Estrutura
        await emit("formatting", 99, "Auditoria Final de Estrutura...")
        full_formatted, audit_issues = self.final_structure_audit(full_formatted, global_structure)
        if level_issues:
            audit_issues = list(audit_issues or [])
            audit_issues.append("\n⚠️ PROBLEMAS DE HIERARQUIA (DETERMINÍSTICO):")
            audit_issues.extend([f"   - {issue}" for issue in level_issues])

        # Passada 4.8: Reatribuição determinística de tabelas por tópico (subtópicos)
        full_formatted, table_reassign_issues = reatribuir_tabelas_por_topico(full_formatted, apply_fixes=True)
        if table_reassign_issues:
            audit_issues = list(audit_issues or [])
            audit_issues.append("\n⚠️ POSSÍVEL REATRIBUIÇÃO DE TABELAS (DETERMINÍSTICO):")
            audit_issues.extend([f"   - {issue}" for issue in table_reassign_issues])

        # Passada 4.9: Reatribuição cirúrgica via IA (fallback)
        ai_reassign_enabled = os.getenv("IUDEX_TABLE_REASSIGN_AI", "").strip().lower() in ("1", "true", "yes")
        if ai_reassign_enabled:
            try:
                full_formatted, ai_reassign_issues = await self._ai_reassign_tables(full_formatted)
                if ai_reassign_issues:
                    audit_issues = list(audit_issues or [])
                    audit_issues.append("\n⚠️ REATRIBUIÇÃO DE TABELAS (IA CIRÚRGICA):")
                    audit_issues.extend([f"   - {issue}" for issue in ai_reassign_issues])
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Falha na reatribuição de tabelas via IA: {e}")
        
        # 6. Validation & Coverage
        print(f"\n{Fore.CYAN}🛡️  Validando cobertura final...")
        await emit("formatting", 99, "Validando cobertura final...")
        coverage_report = self.check_coverage(transcription, full_formatted)
        
        # Save validation report
        report_path = Path(output_folder) / f"{video_name}_validacao.txt"
        with open(report_path, "w", encoding='utf-8') as f:
            f.write(coverage_report)
        print(f"📄 Relatório de validação salvo: {report_path.name}")
        
        # Save audit report if there are issues
        if audit_issues:
            audit_path = Path(output_folder) / f"{video_name}_{mode_suffix}_verificacao.txt"
            with open(audit_path, "w", encoding='utf-8') as f:
                f.write("# AUDITORIA DE ESTRUTURA (v2.16)\n\n")
                f.write(f"Cobertura: {coverage_report}\n\n")
                f.write("## Problemas Estruturais Detectados\n")
                for issue in audit_issues:
                    f.write(f"{issue}\n")
            print(f"📄 Relatório de auditoria salvo: {audit_path.name}")
        
        # 7. v2.16: Validação Full-Context LLM (Backup opcional)
        validation_result = None
        primary_fidelity_written = False
        if FIDELITY_BACKUP_ENABLED:
            print(f"\n{Fore.CYAN}🔬 Validação Full-Context LLM (backup)...")
            validation_result = self.validate_completeness_full(
                transcription, full_formatted, video_name, global_structure
            )

            # Salvar relatório JSON do backup
            validation_report_path = Path(output_folder) / f"{video_name}_{mode_suffix}_fidelidade_backup.json"
            with open(validation_report_path, "w", encoding='utf-8') as f:
                json.dump(validation_result, f, ensure_ascii=False, indent=2)
            print(f"📄 Relatório de fidelidade (backup) salvo: {validation_report_path.name}")

            # Se houver problemas graves, gerar também um markdown legível
            # v2.30: Se houve erro de validação, default é False (não mascara falha)
            _default_aprovado = False if validation_result.get('erro_validacao') else True
            if not validation_result.get('aprovado', _default_aprovado):
                fidelity_md_path = Path(output_folder) / f"{video_name}_{mode_suffix}_REVISAO.md"
                with open(fidelity_md_path, "w", encoding='utf-8') as f:
                    f.write(f"# ⚠️ REVISÃO NECESSÁRIA: {video_name}\n\n")
                    f.write(f"**Nota de Fidelidade:** {validation_result.get('nota', 0)}/10\n\n")
                    if validation_result.get('omissoes'):
                        f.write("## 📌 Omissões Detectadas\n")
                        for o in validation_result['omissoes']:
                            f.write(f"- {o}\n")
                        f.write("\n")
                    if validation_result.get('distorcoes'):
                        f.write("## ⚠️ Distorções Detectadas\n")
                        for d in validation_result['distorcoes']:
                            f.write(f"- {d}\n")
                        f.write("\n")
                    if validation_result.get('problemas_estrutura'):
                        f.write("## 🏗️ Problemas de Estrutura\n")
                        for p in validation_result['problemas_estrutura']:
                            f.write(f"- {p}\n")
                        f.write("\n")
                    if validation_result.get('observacoes'):
                        f.write(f"## Observações\n{validation_result['observacoes']}\n")
                print(f"{Fore.RED}📄 ATENÇÃO: Documento requer revisão! Veja: {fidelity_md_path.name}")

            # 7.1 v2.18: Corretor IA Seguro (Safe Mode) - Corrige problemas estruturais
            # v2.30: Se houve erro de validação, default é False (não mascara falha)
            _default_aprovado2 = False if (validation_result or {}).get('erro_validacao') else True
            if validation_result and not validation_result.get('aprovado', _default_aprovado2):
                print(f"\n{Fore.CYAN}🔁 Iniciando Auto-Fix Loop (Safe Mode)...")

                # Chama o novo corretor seguro
                full_formatted = await self.auto_fix_smart(full_formatted, validation_result, global_structure)

                # Re-validar após correção
                print(f"{Fore.CYAN}🔬 Re-validando após correção automática...")
                revalidation_result = self.validate_completeness_full(
                    transcription, full_formatted, video_name, global_structure
                )

                # Atualizar o relatório com a revalidação
                validation_result = revalidation_result
                validation_report_path = Path(output_folder) / f"{video_name}_{mode_suffix}_fidelidade_backup.json"
                with open(validation_report_path, "w", encoding='utf-8') as f:
                    json.dump(validation_result, f, ensure_ascii=False, indent=2)
                print(f"📄 Relatório de fidelidade (backup) atualizado: {validation_report_path.name}")
        
        def _build_fidelity_report(result):
            if not isinstance(result, dict):
                return ""
            nota = result.get("nota_fidelidade", result.get("nota", 0))
            aprovado = result.get("aprovado", True)
            omissoes = result.get("omissoes_graves", result.get("omissoes", [])) or []
            distorcoes = result.get("distorcoes", []) or []
            estrutura = result.get("problemas_estrutura", []) or []
            observacoes = result.get("observacoes", "") or ""

            def _sanitize(text):
                return str(text).replace("-->", "-- >").strip()

            lines = [
                "# Relatório de Fidelidade (RAW x Formatado)",
                f"Aprovado: {'sim' if aprovado else 'nao'}",
                f"Nota: {nota}/10",
            ]

            if omissoes:
                lines.append(f"Omissões graves: {len(omissoes)}")
                lines.extend([f"- {_sanitize(item)}" for item in omissoes[:20]])
            else:
                lines.append("Omissões graves: 0")

            if distorcoes:
                lines.append(f"Distorções: {len(distorcoes)}")
                lines.extend([f"- {_sanitize(item)}" for item in distorcoes[:20]])
            else:
                lines.append("Distorções: 0")

            if estrutura:
                lines.append(f"Problemas de estrutura: {len(estrutura)}")
                lines.extend([f"- {_sanitize(item)}" for item in estrutura[:20]])
            else:
                lines.append("Problemas de estrutura: 0")

            if observacoes:
                lines.append(f"Observações: {_sanitize(observacoes)}")

            return "\n".join(lines).strip()

        fidelity_report = _build_fidelity_report(validation_result)
        # Relatório salvo apenas em arquivo JSON separado, não incluído no markdown
        # if fidelity_report:
        #     full_formatted += f"\n\n<!-- RELATÓRIO: {fidelity_report} -->"

        # v2.27: Auditoria Preventiva de Fidelidade (antes do DOCX)
        print(f"{Fore.CYAN}📊 [DIAG] Audit Check: AVAILABLE={FIDELITY_AUDIT_AVAILABLE}, ENABLED={FIDELITY_AUDIT_ENABLED}, skip={skip_fidelity_audit}, output_folder={output_folder}")
        if FIDELITY_AUDIT_AVAILABLE and FIDELITY_AUDIT_ENABLED and not skip_fidelity_audit:
            print(f"\n{Fore.CYAN}🔬 Auditoria Preventiva de Fidelidade (v2.27)...")
            await emit("formatting", 99, "Auditoria preventiva de fidelidade...")

            preventive_json = Path(output_folder) / f"{video_name}_{mode_suffix}_AUDITORIA_FIDELIDADE.json"
            preventive_md = Path(output_folder) / f"{video_name}_{mode_suffix}_AUDITORIA_FIDELIDADE.md"

            try:
                preventive_result = auditar_fidelidade_preventiva(
                    self.client,
                    transcription,
                    full_formatted,
                    video_name,
                    str(preventive_json),
                    modo=mode_suffix,
                    include_sources=(SOURCES_AUDIT_ENABLED and not skip_sources_audit),
                )
                if not isinstance(preventive_result, dict):
                    preventive_result = {
                        "aprovado": False,
                        "nota_fidelidade": 0,
                        "gravidade_geral": "CRÍTICA",
                        "erro": f"Auditoria preventiva retornou tipo inválido: {type(preventive_result)}",
                        "recomendacao_hil": {"pausar_para_revisao": True, "motivo": "Resultado inválido", "areas_criticas": ["auditoria_preventiva"]},
                        "omissoes_criticas": [],
                        "distorcoes": [],
                        "alucinacoes": [],
                        "problemas_estruturais": [],
                        "problemas_contexto": [],
                        "metricas": {},
                    }
                try:
                    needs_persist = True
                    try:
                        needs_persist = (not preventive_json.exists()) or preventive_json.stat().st_size == 0
                    except Exception:
                        needs_persist = True
                    if needs_persist:
                        with open(preventive_json, "w", encoding="utf-8") as f:
                            json.dump(preventive_result, f, ensure_ascii=False, indent=2, default=str)
                except Exception as write_err:
                    print(f"{Fore.YELLOW}⚠️ Falha ao salvar JSON da auditoria preventiva: {write_err}")
                # Markdown generation should never abort/overwrite the JSON result.
                try:
                    gerar_relatorio_markdown_completo(preventive_result, str(preventive_md), video_name)
                except Exception as md_err:
                    try:
                        md_fallback = (
                            f"# 🔬 AUDITORIA PREVENTIVA DE FIDELIDADE: {video_name}\n\n"
                            f"**Status:** ⚠️ REQUER REVISÃO\n"
                            f"**Nota de Fidelidade:** 0.0/10\n"
                            f"**Gravidade Geral:** N/A\n\n"
                            f"## ❌ Erro\n\nFalha ao gerar relatório Markdown: {md_err}\n"
                        )
                        with open(preventive_md, "w", encoding="utf-8") as f:
                            f.write(md_fallback)
                    except Exception as md_write_err:
                        print(f"{Fore.YELLOW}⚠️ Falha ao salvar markdown fallback: {md_write_err}")

                compat = (preventive_result or {}).get("compat_fidelidade")
                if isinstance(compat, dict) and compat:
                    fidelity_path = Path(output_folder) / f"{video_name}_{mode_suffix}_fidelidade.json"
                    with open(fidelity_path, "w", encoding="utf-8") as f:
                        json.dump(compat, f, ensure_ascii=False, indent=2)
                    print(f"📄 Relatório de fidelidade (preventiva) salvo: {fidelity_path.name}")
                    primary_fidelity_written = True

                recomendacao = (preventive_result or {}).get("recomendacao_hil", {}) or {}
                if hil_strict and recomendacao.get("pausar_para_revisao"):
                    save_hil_output(
                        full_formatted,
                        video_name,
                        output_folder,
                        mode_suffix,
                        reason="auditoria_preventiva",
                    )
                    raise HILCheckpointException(
                        f"Auditoria preventiva exige revisão humana. Veja: {preventive_md.name}"
                    )
            except HILCheckpointException:
                raise
            except Exception as e:
                import traceback
                print(f"{Fore.YELLOW}⚠️ Falha na auditoria preventiva: {e}. Continuando...")
                print(f"{Fore.RED}Traceback: {traceback.format_exc()}")
                # Save minimal error report so frontend doesn't show "unavailable"
                error_result = {
                    "aprovado": False,
                    "nota_fidelidade": 0,
                    "gravidade_geral": "N/A",
                    "erro": str(e),
                    "recomendacao_hil": {"pausar_para_revisao": False, "motivo": f"Falha na execução da auditoria: {str(e)}", "areas_criticas": []},
                    "compat_fidelidade": {"aprovado": False, "nota": 0, "erro": str(e)},
                    "omissoes_criticas": [],
                    "distorcoes": [],
                    "alucinacoes": [],
                    "problemas_estruturais": [],
                    "problemas_contexto": [],
                    "metricas": {},
                    "observacoes_gerais": f"Erro na auditoria: {str(e)}"
                }
                try:
                    with open(preventive_json, "w", encoding="utf-8") as f:
                        json.dump(error_result, f, ensure_ascii=False, indent=2)
                    gerar_relatorio_markdown_completo(error_result, str(preventive_md), video_name)
                except Exception as write_err:
                    print(f"{Fore.RED}❌ Falha ao salvar relatório de erro: {write_err}")

        elif not skip_fidelity_audit:
            # Audit wanted but unavailable/disabled - write placeholder
            print(f"{Fore.YELLOW}⚠️ Auditoria preventiva indisponível ou desativada. Gerando relatório de status.")
            preventive_json = Path(output_folder) / f"{video_name}_{mode_suffix}_AUDITORIA_FIDELIDADE.json"
            preventive_md = Path(output_folder) / f"{video_name}_{mode_suffix}_AUDITORIA_FIDELIDADE.md"
            
            reason = "Módulo de auditoria não encontrado ou falha na importação." if not FIDELITY_AUDIT_AVAILABLE else "Auditoria desativada por configuração."
            
            placeholder_result = {
                "aprovado": True,
                "nota_fidelidade": 0,
                "gravidade_geral": "INFO",
                "recomendacao_hil": {
                    "pausar_para_revisao": False, 
                    "motivo": f"Auditoria não executada: {reason}", 
                    "areas_criticas": []
                },
                "observacoes_gerais": f"A auditoria preventiva não foi executada. {reason}",
                "compat_fidelidade": {"aprovado": True, "nota": 0},
                "omissoes_criticas": [],
                "distorcoes": [],
                "alucinacoes": [],
                "problemas_estruturais": [],
                "problemas_contexto": [],
                "metricas": {}
            }
            try:
                with open(preventive_json, "w", encoding="utf-8") as f:
                    json.dump(placeholder_result, f, ensure_ascii=False, indent=2)
                
                # Write simple markdown
                md_content = f"# Auditoria Preventiva\n\n**Status:** Não executada\n\n**Motivo:** {reason}"
                with open(preventive_md, "w", encoding="utf-8") as f:
                    f.write(md_content)
                    
            except Exception as e:
                print(f"{Fore.RED}❌ Falha ao salvar placeholder de auditoria: {e}")

        if not primary_fidelity_written and isinstance(validation_result, dict):
            fidelity_path = Path(output_folder) / f"{video_name}_{mode_suffix}_fidelidade.json"
            try:
                with open(fidelity_path, "w", encoding="utf-8") as f:
                    json.dump(validation_result, f, ensure_ascii=False, indent=2)
                print(f"📄 Relatório de fidelidade (fallback) salvo: {fidelity_path.name}")
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Falha ao salvar relatório de fidelidade fallback: {e}")

        # Telemetria de drift de títulos (modo FIDELIDADE)
        if mode_suffix == "FIDELIDADE":
            try:
                drift_path = Path(output_folder) / f"{video_name}_{mode_suffix}_TITLE_DRIFT.json"
                with open(drift_path, "w", encoding="utf-8") as f:
                    json.dump(title_drift_telemetry, f, ensure_ascii=False, indent=2)
                print(f"📄 Telemetria de drift de títulos salva: {drift_path.name}")
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Falha ao salvar telemetria de drift de títulos: {e}")

        # Checkpoint cleanup success
        delete_checkpoint(video_name, output_folder)
        
        # 8. v2.18: Auditoria Jurídica Pós-Processamento
        if AUDIT_AVAILABLE and not skip_audit:
            print(f"\n{Fore.CYAN}🕵️ Auditoria Jurídica Pós-Processamento...")
            audit_report_path = Path(output_folder) / f"{video_name}_{mode_suffix}_AUDITORIA.md"
            audit_content = auditar_consistencia_legal(
                self.client,
                full_formatted,
                str(audit_report_path),
                raw_transcript=transcription,
            )
            
            if audit_content:
                print(f"{Fore.GREEN}   📎 Relatório de auditoria salvo em arquivo separado...")
                # Não incluir no markdown para não poluir a apostila final
                # full_formatted += f"\n\n<!-- RELATÓRIO: {audit_content} -->"
        
        # v2.19: Cleanup manual do cache para economia
        if cached_context:
            try:
                self.client.caches.delete(name=cached_context.name)
                print(f"{Fore.GREEN}🗑️ Cache {cached_context.name} deletado manualmente para economizar custos.")
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Não foi possível deletar o cache: {e}")

        # v2.30: Limpeza final de vocativos/gírias (ex.: "Meu irmão, ...")
        try:
            full_formatted = remover_vocativos_girias(full_formatted)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Falha ao remover vocativos/gírias: {e}")

        # v2.35: Normalizações finais usadas também no preview/API (não só no Word)
        try:
            full_formatted = normalizar_temas_markdown(full_formatted)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Falha ao normalizar Temas: {e}")
        try:
            full_formatted = remover_marcadores_continua(full_formatted)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Falha ao remover marcadores [continua]: {e}")

        await emit("formatting", 100, "Formatação concluída")

        return full_formatted


    async def auto_fix_smart(self, formatted_text, validation_result, global_structure=None):
        """
        v2.18 (SAFE MODE): Corretor Estrutural Seguro.
        Foca EXCLUSIVAMENTE em problemas de estrutura (títulos, duplicatas, hierarquia).
        NÃO altera conteúdo jurídico para evitar alucinações.
        """
        problemas_estrut = validation_result.get('problemas_estrutura', [])
        
        if not problemas_estrut:
            print(f"{Fore.GREEN}   ✅ Nenhum problema estrutural para corrigir.")
            return formatted_text
        
        print(f"{Fore.CYAN}🔧 Auto-Fix Safe: Corrigindo {len(problemas_estrut)} problemas estruturais...")
        
        report = "### PROBLEMAS ESTRUTURAIS:\n" + "\n".join([f"- {p}" for p in problemas_estrut]) + "\n"
            
        global_reference = (
            "## ESTRUTURA DE REFERÊNCIA (Guia):\n" + global_structure
            if global_structure
            else ""
        )

        PROMPT_FIX = f"""Você é um editor técnico de elite.
        
## TAREFA: LIMPEZA ESTRUTURAL (SEM ALTERAR CONTEÚDO)
Você deve corrigir APENAS a formatação e estrutura do documento.

## REGRA DE OURO (SEGURANÇA JURÍDICA):
- **NÃO altere o texto dos parágrafos.**
- **NÃO adicione nem remova informações jurídicas.**
- **NÃO reescreva explicações.**
- Sua permissão é APENAS para Títulos, Hierarquia e Duplicatas exatas.

## INSTRUÇÕES DE CORREÇÃO:
1. **Títulos Duplicados**: Se um título H2 aparece duas vezes, REMOVA a segunda ocorrência e mescle o conteúdo sob o primeiro.
2. **Hierarquia**: Ajuste níveis (H2, H3) para seguir a lógica do conteúdo.
3. **Parágrafos Repetidos**: Delete duplicações EXATAS de parágrafos (copia-cola acidental).
4. **Renumeração**: Garanta sequência lógica (1, 2, 3...) nos títulos.

{global_reference}

## RELATÓRIO DE ERROS:
{report}

## SAÍDA:
Retorne o documento COMPLETO corrigido em Markdown. Sem explicações."""

        try:
            # Call Gemini Async
            def call_gemini():
                return self.client.models.generate_content(
                    model=self.llm_model,
                    contents=f"{PROMPT_FIX}\n\n## TEXTO A CORRIGIR:\n{formatted_text}",
                    config=types.GenerateContentConfig(
                        max_output_tokens=8192,
                        thinking_config=types.ThinkingConfig(
                            include_thoughts=False,
                            thinking_level="HIGH" 
                        )
                    )
                )

            response = await asyncio.to_thread(call_gemini)
            _record_genai_usage(response, model=self.llm_model)
            
            resultado = response.text.replace('```markdown', '').replace('```', '').strip()
            
            # Validação de segurança estrita (0.8 = 20% tolerância vs 30% antigo)
            if len(resultado) < len(formatted_text) * 0.8:
                print(f"{Fore.YELLOW}⚠️ Auto-Fix Safe cortou muito texto (>20%). Abortando por segurança.")
                return formatted_text
                
            print(f"{Fore.GREEN}   ✅ Auto-Fix Estrutural concluído. ({len(formatted_text)} -> {len(resultado)} chars)")
            return resultado
            
        except Exception as e:
            print(f"{Fore.RED}❌ Falha no Auto-Fix Safe: {e}")
            return formatted_text


    def save_as_word(
        self,
        formatted_text,
        video_name,
        output_folder,
        mode=None,
        document_theme="classic",
        document_header=None,
        document_footer=None,
        document_margins="normal",
        document_font_family=None,
        document_font_size=None,
        document_line_height=None,
        document_paragraph_spacing=None,
    ):
        """Salva markdown formatado como documento Word (.docx) com estilo premium"""
        # v2.23: Dynamic mode suffix for file naming
        mode_suffix = mode.upper() if mode else getattr(self, '_current_mode', 'APOSTILA')
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            print(f"{Fore.YELLOW}⚠️ python-docx não instalado. Salvando apenas Markdown.")
            return None

        print(f"{Fore.CYAN}📄 Gerando documento Word profissional...")

        # v2.28: Sanitização do markdown antes de converter
        try:
            formatted_text = sanitizar_markdown_final(formatted_text)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Erro na sanitização: {e}. Continuando com texto original.")

        # v2.28: Corrigir tabelas que aparecem antes do conteúdo terminar
        try:
            formatted_text = corrigir_tabelas_prematuras(formatted_text)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Erro ao corrigir tabelas prematuras: {e}.")

        # 1. Aplicar Smart Layout (opcional, mantido do Vomo para consistência)
        try:
            formatted_text = mover_tabelas_para_fim_de_secao(formatted_text)
        except Exception as e:
            print(f"{Fore.YELLOW}⚠️ Erro ao reorganizar tabelas: {e}. Usando layout padrão.")

        doc = Document()

        theme_norm = (document_theme or "classic").strip().lower()
        margins_norm = (document_margins or "normal").strip().lower()
        theme_presets = {
            "classic": {
                "font": "Arial",
                "title_color": RGBColor(0, 51, 102),
                "heading_color": RGBColor(0, 51, 102),
                "margins": (1, 1, 1.25, 1.25),
                "table": {
                    "default": {"header_bg": "0066CC", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "F0F6FF"},
                    "quadro_sintese": {"header_bg": "0066CC", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "E6F2FF"},
                    "pegadinhas": {"header_bg": "E67E00", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "FFF5E6"},
                },
            },
            "minimal": {
                "font": "Arial",
                "title_color": RGBColor(55, 65, 81),
                "heading_color": RGBColor(55, 65, 81),
                "margins": (0.9, 0.9, 1.0, 1.0),
                "table": {
                    "default": {"header_bg": "F8FAFC", "header_text": RGBColor(31, 41, 55), "alt_row_bg": "FFFFFF"},
                    "quadro_sintese": {"header_bg": "F8FAFC", "header_text": RGBColor(31, 41, 55), "alt_row_bg": "FFFFFF"},
                    "pegadinhas": {"header_bg": "FFF7ED", "header_text": RGBColor(120, 53, 15), "alt_row_bg": "FFFFFF"},
                },
            },
            "executive": {
                "font": "Arial",
                "title_color": RGBColor(17, 24, 39),
                "heading_color": RGBColor(17, 24, 39),
                "margins": (1, 1, 1.1, 1.1),
                "table": {
                    "default": {"header_bg": "111827", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "F3F4F6"},
                    "quadro_sintese": {"header_bg": "0F172A", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "E5E7EB"},
                    "pegadinhas": {"header_bg": "B45309", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "FEF3C7"},
                },
            },
            "academic": {
                "font": "Times New Roman",
                "title_color": RGBColor(55, 65, 81),
                "heading_color": RGBColor(55, 65, 81),
                "margins": (1.25, 1.25, 1.35, 1.35),
                "table": {
                    "default": {"header_bg": "ECEFF4", "header_text": RGBColor(31, 41, 55), "alt_row_bg": "FFFFFF"},
                    "quadro_sintese": {"header_bg": "E2E8F0", "header_text": RGBColor(31, 41, 55), "alt_row_bg": "FFFFFF"},
                    "pegadinhas": {"header_bg": "FFF7ED", "header_text": RGBColor(120, 53, 15), "alt_row_bg": "FFFFFF"},
                },
            },
        }
        theme = theme_presets.get(theme_norm, theme_presets["classic"])
        font_name = theme["font"]
        if document_font_family:
            font_name = str(document_font_family).strip() or font_name
        self._doc_font_name = font_name
        
        # 2. Configurações Globais de Estilo (Arial + Justificado)
        style = doc.styles['Normal']
        font = style.font
        font.name = font_name
        base_font_size = None
        try:
            if document_font_size is not None:
                font_size_val = float(document_font_size)
                # UI usa px; converter para pontos (1px ≈ 0.75pt)
                font.size = Pt(max(8, font_size_val * 0.75))
            else:
                font.size = Pt(11)
        except Exception:
            font.size = Pt(11)
        base_font_size = font.size or Pt(11)
        
        # Garantir Arial em documentos que ignoram o nome da fonte
        r = style.element.rPr.get_or_add_rFonts()
        r.set(qn('w:ascii'), font_name)
        r.set(qn('w:hAnsi'), font_name)
        
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if document_line_height is not None:
            try:
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                style.paragraph_format.line_spacing = float(document_line_height)
            except Exception:
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if document_paragraph_spacing is not None:
            try:
                spacing_pt = float(document_paragraph_spacing) * 0.75
                style.paragraph_format.space_after = Pt(max(0, spacing_pt))
            except Exception:
                pass
        
        # Aplicar Arial também aos títulos e outros estilos
        for style_name in [f'Heading {i}' for i in range(1, 6)] + ['Quote', 'List Bullet', 'List Number']:
            try:
                s = doc.styles[style_name]
                s.font.name = font_name
                r = s.element.rPr.get_or_add_rFonts()
                r.set(qn('w:ascii'), font_name)
                r.set(qn('w:hAnsi'), font_name)
            except KeyError:
                pass

        # Margens
        section = doc.sections[0]
        top_m, bottom_m, left_m, right_m = theme["margins"]
        if margins_norm == "compact":
            top_m, bottom_m, left_m, right_m = (0.9, 0.9, 1.0, 1.0)
        elif margins_norm == "wide":
            top_m, bottom_m, left_m, right_m = (1.25, 1.25, 1.35, 1.35)
        section.top_margin = Inches(top_m)
        section.bottom_margin = Inches(bottom_m)
        section.left_margin = Inches(left_m)
        section.right_margin = Inches(right_m)

        def _add_page_number(paragraph):
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            run._r.append(instrText)
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)

        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_text = (document_header or "").strip() or f"{video_name} — {mode_suffix}"
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(120, 120, 120)
            run.font.name = font_name

        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = ""
        footer_text = (document_footer or "").strip()
        if footer_text:
            run = footer_para.add_run(f"{footer_text} — ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(120, 120, 120)
            run.font.name = font_name
        _add_page_number(footer_para)
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(120, 120, 120)
            run.font.name = font_name
        
        # Título principal
        title = doc.add_heading(video_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = theme["title_color"]
            run.font.name = font_name
        
        # Data de geração e Modo
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')} - Modo: {mode_suffix}")
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_run.font.name = font_name
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Sumário
        doc.add_heading('Sumário', level=1)
        self.create_toc(doc)
        doc.add_page_break()
        
        # Processa conteúdo markdown
        lines = formatted_text.split('\n')
        i = 0
        in_table = False
        table_rows = []
        current_table_cols = None
        current_table_type = "default"  # v2.28: Tipo de tabela atual

        def _is_table_separator(line: str) -> bool:
            return bool(re.match(r'^\s*\|[\s:|-]+\|[\s:|-]*$', line))

        def _count_table_cols(line: str) -> int:
            if '|' not in line:
                return 0
            return max(0, line.count('|') - 1)

        def _looks_like_table_header(idx: int) -> bool:
            if idx + 1 >= len(lines):
                return False
            if '|' not in lines[idx]:
                return False
            return _is_table_separator(lines[idx + 1].strip())

        def _detect_table_type_from_heading(heading_text: str) -> str:
            """v2.28: Detecta tipo de tabela pelo heading anterior."""
            if '📋' in heading_text or 'uadro' in heading_text.lower():
                return "quadro_sintese"
            elif '🎯' in heading_text or 'pegadinha' in heading_text.lower() or 'banca' in heading_text.lower():
                return "pegadinhas"
            return "default"

        while i < len(lines):
            line = lines[i].strip()

            if line in {"<!-- PAGE_BREAK -->", "<!--PAGE_BREAK-->"}:
                doc.add_page_break()
                i += 1
                continue
            
            if in_table:
                if not line:
                    if table_rows:
                        self._add_table_to_doc(doc, table_rows, current_table_type, theme_norm)
                    in_table = False
                    table_rows = []
                    current_table_cols = None
                    i += 1
                    continue

                if '|' in line:
                    if _looks_like_table_header(i):
                        candidate_cols = _count_table_cols(line)
                        if current_table_cols and table_rows and candidate_cols != current_table_cols:
                            self._add_table_to_doc(doc, table_rows, current_table_type, theme_norm)
                            table_rows = []
                            current_table_cols = None

                    is_separator = _is_table_separator(line)
                    if not is_separator:
                        row = [cell.strip() for cell in line.split('|')[1:-1]]
                        table_rows.append(row)
                        if row:
                            current_table_cols = max(current_table_cols or 0, len(row))

                    if i == len(lines) - 1:
                        if table_rows:
                            self._add_table_to_doc(doc, table_rows, current_table_type, theme_norm)
                        in_table = False
                        table_rows = []
                        current_table_cols = None
                    i += 1
                    continue

                if table_rows:
                    self._add_table_to_doc(doc, table_rows, current_table_type, theme_norm)
                in_table = False
                table_rows = []
                current_table_cols = None
                continue

            if not line:
                i += 1
                continue

            # Tabelas
            if '|' in line:
                in_table = True
                table_rows = []
                current_table_cols = None
                # v2.28: Tipo já foi definido pelo heading anterior
                continue

            # Headings
            if line.startswith('##### '):
                h = doc.add_heading('', level=5)
                self._format_inline_markdown(h.paragraphs[0], line[6:])
                # v2.28: Detectar tipo de tabela para heading level 5
                current_table_type = _detect_table_type_from_heading(line[6:])
            elif h_match := re.match(r'^(####|###|##|#)\s+(.*)', line):
                lvl = len(h_match.group(1))
                h_text = h_match.group(2)
                if lvl == 1 and h_text == video_name:
                    i += 1
                    continue
                h = doc.add_heading('', level=lvl)
                self._format_inline_markdown(h, h_text)
                for run in h.runs:
                    run.font.name = font_name
                    run.font.color.rgb = theme["heading_color"]
                # v2.28: Detectar tipo de tabela para heading level 4
                if lvl == 4:
                    current_table_type = _detect_table_type_from_heading(h_text)
            
            # Separadores
            elif line.strip() in ['---', '***', '___']:
                p = doc.add_paragraph()
                p.add_run('_' * 80).font.color.rgb = RGBColor(192, 192, 192)
            
            # Quotes
            elif line.startswith('>'):
                p = doc.add_paragraph(style='Quote')
                p.paragraph_format.left_indent = Cm(4.0)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._format_inline_markdown(p, line[1:].strip())
                for run in p.runs:
                    run.italic = True
                    run.font.size = Pt(10)
            
            # Listas não-ordenadas
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.first_line_indent = Cm(-0.63)
                self._format_inline_markdown(p, line[2:])
                
            # Listas numeradas
            elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.first_line_indent = Cm(-0.63)
                self._format_inline_markdown(p, line)
                
            # Parágrafo normal
            else:
                p = doc.add_paragraph()
                if document_line_height is not None:
                    try:
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        p.paragraph_format.line_spacing = float(document_line_height)
                    except Exception:
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                else:
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_before = Pt(6)
                if document_paragraph_spacing is not None:
                    try:
                        spacing_pt = float(document_paragraph_spacing) * 0.75
                        p.paragraph_format.space_after = Pt(max(0, spacing_pt))
                    except Exception:
                        p.paragraph_format.space_after = Pt(6)
                else:
                    p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.first_line_indent = Cm(1.0)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._format_inline_markdown(p, line)
                for run in p.runs:
                    run.font.size = base_font_size
            
            i += 1
            
        output_file = os.path.join(output_folder, f"{video_name}_{mode_suffix}.docx")
        doc.save(output_file)
        return output_file

    def _format_inline_markdown(self, paragraph, text):
        """Formata markdown inline avançado (bold, italic, code, underline-style)"""
        from docx.shared import Pt, RGBColor
        paragraph.clear()
        font_name = getattr(self, "_doc_font_name", "Arial")
        
        # Regex robusta do format_transcription_gemini.py
        pattern = r'(\*{3}(.+?)\*{3}|_{3}(.+?)_{3}|\*{2}(.+?)\*{2}|_{2}(.+?)_{2}|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|(?<!_)_(?!(?:_|\s))(.+?)(?<!(?:_|\s))_(?!_)|`(.+?)`)'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            if match.start() > last_end:
                run = paragraph.add_run(text[last_end:match.start()])
                run.font.name = font_name
            
            full_match = match.group(0)
            
            if full_match.startswith('***'):
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
                run.font.name = font_name
            elif full_match.startswith('___'):
                run = paragraph.add_run(match.group(3))
                run.bold = True
                run.italic = True
                run.font.name = font_name
            elif full_match.startswith('**'):
                run = paragraph.add_run(match.group(4))
                run.bold = True
                run.font.name = font_name
            elif full_match.startswith('__'):
                run = paragraph.add_run(match.group(5))
                run.bold = True
                run.font.name = font_name
            elif full_match.startswith('*'):
                run = paragraph.add_run(match.group(6))
                run.italic = True
                run.font.name = font_name
            elif full_match.startswith('_'):
                run = paragraph.add_run(match.group(7))
                run.italic = True
                run.font.name = font_name
            elif full_match.startswith('`'):
                run = paragraph.add_run(match.group(8))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(200, 0, 0)
            
            last_end = match.end()
        
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.name = font_name

    def _add_table_to_doc(self, doc, rows, table_type="default", document_theme="classic"):
        """
        v2.28: Adiciona tabela premium ao Word com estilos diferenciados.

        Args:
            doc: Documento Word
            rows: Lista de listas com dados das células
            table_type: Tipo de tabela para estilização diferenciada
                - "quadro_sintese" (📋): Azul, 5 colunas, didático
                - "pegadinhas" (🎯): Laranja, 3 colunas, alerta
                - "default": Azul padrão
        """
        from docx.shared import RGBColor, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if len(rows) < 2: return
        max_cols = max(len(row) for row in rows)
        if max_cols == 0: return

        theme_norm = (document_theme or "classic").strip().lower()
        palettes = {
            "classic": {
                "default": {"header_bg": "0066CC", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "F0F6FF"},
                "quadro_sintese": {"header_bg": "0066CC", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "E6F2FF"},
                "pegadinhas": {"header_bg": "E67E00", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "FFF5E6"},
            },
            "minimal": {
                "default": {"header_bg": "F8FAFC", "header_text": RGBColor(31, 41, 55), "alt_row_bg": "FFFFFF"},
                "quadro_sintese": {"header_bg": "F8FAFC", "header_text": RGBColor(31, 41, 55), "alt_row_bg": "FFFFFF"},
                "pegadinhas": {"header_bg": "FFF7ED", "header_text": RGBColor(120, 53, 15), "alt_row_bg": "FFFFFF"},
            },
            "executive": {
                "default": {"header_bg": "111827", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "F3F4F6"},
                "quadro_sintese": {"header_bg": "0F172A", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "E5E7EB"},
                "pegadinhas": {"header_bg": "B45309", "header_text": RGBColor(255, 255, 255), "alt_row_bg": "FEF3C7"},
            },
            "academic": {
                "default": {"header_bg": "ECEFF4", "header_text": RGBColor(31, 41, 55), "alt_row_bg": "FFFFFF"},
                "quadro_sintese": {"header_bg": "E2E8F0", "header_text": RGBColor(31, 41, 55), "alt_row_bg": "FFFFFF"},
                "pegadinhas": {"header_bg": "FFF7ED", "header_text": RGBColor(120, 53, 15), "alt_row_bg": "FFFFFF"},
            },
        }

        palette = palettes.get(theme_norm, palettes["classic"])
        cores = palette.get(table_type, palette["default"])

        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_data in enumerate(rows):
            for j in range(max_cols):
                cell = table.rows[i].cells[j]
                cell_text = row_data[j] if j < len(row_data) else ""

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._format_inline_markdown(p, cell_text)

                # Header styling
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.color.rgb = cores["header_text"]
                            run.font.size = Pt(10)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), cores["header_bg"])
                    cell._element.get_or_add_tcPr().append(shading_elm)

                # v2.28: Zebra striping (linhas alternadas)
                elif i % 2 == 0:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), cores["alt_row_bg"])
                    cell._element.get_or_add_tcPr().append(shading_elm)

        # v2.28: Ajustar largura das colunas baseado no tipo
        if table_type == "quadro_sintese" and max_cols == 5:
            # Proporções: Item(15%), Definição(25%), Detalhes(25%), Base legal(15%), Dica(20%)
            widths = [Cm(2.5), Cm(4.0), Cm(4.0), Cm(2.5), Cm(3.5)]
            for j, width in enumerate(widths):
                for row in table.rows:
                    if j < len(row.cells):
                        row.cells[j].width = width
        elif table_type == "pegadinhas" and max_cols == 3:
            # Proporções: Como cobra(35%), Resposta(30%), Erro comum(35%)
            widths = [Cm(5.5), Cm(4.5), Cm(5.5)]
            for j, width in enumerate(widths):
                for row in table.rows:
                    if j < len(row.cells):
                        row.cells[j].width = width

    def create_toc(self, doc):
        """Adiciona Sumário nativo do Word"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

def process_single_video(
    video_path,
    dry_run=False,
    mode="APOSTILA",
    skip_formatting=False,
    custom_prompt=None,
    high_accuracy=False,
    diarization: Optional[bool] = None,
    diarization_strict: bool = False,
    skip_audit=False,
    skip_fidelity_audit=False,
    skip_sources_audit=False,
    hil_strict=False,
    resume_hil=False,
    provider="gemini",
    word_only=False,
    auto_apply_fixes=False,
):
    def _is_public_url(value: str) -> bool:
        try:
            parsed = urlparse((value or "").strip())
            return parsed.scheme in {"http", "https"} and bool(parsed.netloc)
        except Exception:
            return False

    def _download_public_media(url: str) -> str:
        """
        Baixa mídia de URL pública (ex.: YouTube) usando `yt-dlp`.

        - Faz cache por hash da URL no diretório configurável.
        - Extrai áudio para MP3 para acelerar o pipeline (FFmpeg ainda fará WAV 16k mono).
        """
        download_dir = os.getenv("IUDEX_URL_DOWNLOAD_DIR", "tmp/url_imports").strip() or "tmp/url_imports"
        Path(download_dir).mkdir(parents=True, exist_ok=True)

        url_norm = (url or "").strip()
        url_hash = hashlib.sha256(url_norm.encode("utf-8")).hexdigest()[:12]
        base = f"url_{url_hash}"

        # Cache: se já existe MP3 baixado, reutiliza.
        cached_mp3 = Path(download_dir) / f"{base}.mp3"
        if cached_mp3.exists() and cached_mp3.stat().st_size > 1024:
            print(f"{Fore.CYAN}🌐 URL cache: usando {cached_mp3.name}")
            return str(cached_mp3)

        ytdlp = (
            (os.getenv("IUDEX_YTDLP_PATH") or "").strip()
            or shutil.which("yt-dlp")
            or shutil.which("yt_dlp")
            or ("/opt/homebrew/bin/yt-dlp" if os.path.exists("/opt/homebrew/bin/yt-dlp") else None)
            or ("/usr/local/bin/yt-dlp" if os.path.exists("/usr/local/bin/yt-dlp") else None)
        )
        if not ytdlp:
            raise RuntimeError(
                "Para baixar vídeos de URL (ex.: YouTube), instale `yt-dlp`.\n"
                "macOS (Homebrew): `brew install yt-dlp`\n"
                "Python: `python3 -m pip install -U yt-dlp`\n"
            )

        outtmpl = str(Path(download_dir) / f"{base}.%(ext)s")
        cmd = [
            ytdlp,
            "--no-playlist",
            "--restrict-filenames",
            "-f",
            "bestaudio/best",
            "--extract-audio",
            "--audio-format",
            "mp3",
            "--audio-quality",
            "0",
            "-o",
            outtmpl,
            url_norm,
        ]
        print(f"{Fore.CYAN}🌐 Baixando URL com yt-dlp...")
        subprocess.run(cmd, check=True)

        # yt-dlp pode gerar .mp3 ou outro ext dependendo de flags; procurar resultado.
        candidates = sorted(Path(download_dir).glob(f"{base}.*"), key=lambda p: p.stat().st_mtime, reverse=True)
        for path in candidates:
            if path.suffix.lower() == ".mp3" and path.stat().st_size > 1024:
                return str(path)
        # Fallback: se não achou mp3, pega o mais recente.
        if candidates:
            return str(candidates[0])
        raise RuntimeError("Falha ao baixar URL: nenhum arquivo gerado.")

    if _is_public_url(video_path):
        video_path = _download_public_media(video_path)

    if not os.path.exists(video_path):
        print(f"{Fore.RED}❌ Arquivo não encontrado: {video_path}")
        return

    folder = os.path.dirname(video_path)
    video_name = Path(video_path).stem
    
    try:
        vomo = VomoMLX(provider=provider)

        if word_only and video_path.lower().endswith('.md'):
            print(f"{Fore.CYAN}📄 Modo --word-only: Gerando Word a partir de MD existente...")
            with open(video_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            vomo.save_as_word(md_content, video_name, folder)
            print(f"{Fore.GREEN}✅ DOCX gerado com sucesso!")
            return
        
        if video_path.lower().endswith(('.txt', '.md')):
            print(f"{Fore.CYAN}📄 Input é arquivo de texto. Pulando transcrição...")
            with open(video_path, 'r', encoding='utf-8') as f:
                transcription = f.read()
        else:
            if dry_run:
                print("⚠️ Dry run não suporta áudio direto ainda. Use arquivo .txt")
                return
            audio = vomo.optimize_audio(video_path)

            diar_enabled, _diar_required = vomo.resolve_diarization_policy(
                mode, diarization=diarization, diarization_strict=diarization_strict
            )
            raw_parts = [video_name, "RAW"]
            if diar_enabled:
                raw_parts.append("DIAR")
            if high_accuracy:
                raw_parts.append("BEAM")
            raw_path = os.path.join(folder, f"{'_'.join(raw_parts)}.txt")

            if os.path.exists(raw_path):
                with open(raw_path, 'r') as f:
                    transcription = f.read()
            else:
                # Escolhe backend de transcrição
                transcription = vomo.transcribe_file(
                    audio,
                    mode=mode,
                    high_accuracy=high_accuracy,
                    diarization=diarization,
                    diarization_strict=diarization_strict,
                )
                with open(raw_path, 'w') as f:
                    f.write(transcription)
        
        if skip_formatting and not resume_hil:
            print(f"{Fore.GREEN}✅ Transcrição RAW concluída: {raw_path}")
            print(f"{Fore.YELLOW}   ℹ️  Formatação pulada (--skip-formatting usado).{Style.RESET_ALL}")
            return

        formatted = None
        mode_suffix = (mode or "APOSTILA").upper()

        if resume_hil:
            hil_path = get_hil_output_path(video_name, folder, mode_suffix)
            if not hil_path.exists():
                print(f"{Fore.RED}❌ HIL checkpoint não encontrado: {hil_path.name}")
                return
            with open(hil_path, 'r', encoding='utf-8') as f:
                formatted = f.read()
            print(f"{Fore.YELLOW}⏯️  Retomando a partir do HIL checkpoint: {hil_path.name}")
            primary_fidelity_written = False
            validation_result = None

            # v3.0: Carregar relatório unificado anterior para comparação
            _previous_unified = None
            if UNIFIED_AUDIT_AVAILABLE:
                _prev_unified_path = Path(folder) / f"{video_name}_{mode_suffix}_UNIFIED_AUDIT.json"
                if _prev_unified_path.exists():
                    try:
                        _previous_unified = UnifiedReport.load_json(str(_prev_unified_path))
                        _prev_active = [f for f in _previous_unified.findings if f.verdict.value != "FALSO_POSITIVO"]
                        print(f"   📊 Relatório unificado anterior: {len(_prev_active)} findings carregados para comparação")
                    except Exception as e:
                        print(f"   ⚠️ Falha ao carregar relatório anterior: {e}")

            # Revalidação preventiva (opcional) após correções manuais
            if FIDELITY_AUDIT_AVAILABLE and FIDELITY_AUDIT_ENABLED and not skip_fidelity_audit:
                print(f"\n{Fore.CYAN}🔬 Revalidando Auditoria Preventiva de Fidelidade...")
                preventive_json = Path(folder) / f"{video_name}_{mode_suffix}_AUDITORIA_FIDELIDADE.json"
                preventive_md = Path(folder) / f"{video_name}_{mode_suffix}_AUDITORIA_FIDELIDADE.md"
                try:
                    preventive_result = auditar_fidelidade_preventiva(
                        vomo.client,
                        transcription,
                        formatted,
                        video_name,
                        str(preventive_json),
                        modo=mode_suffix,
                        include_sources=(SOURCES_AUDIT_ENABLED and not skip_sources_audit),
                    )
                    gerar_relatorio_markdown_completo(preventive_result, str(preventive_md), video_name)
                    compat = (preventive_result or {}).get("compat_fidelidade")
                    if isinstance(compat, dict) and compat:
                        fidelity_path = Path(folder) / f"{video_name}_{mode_suffix}_fidelidade.json"
                        with open(fidelity_path, "w", encoding="utf-8") as f:
                            json.dump(compat, f, ensure_ascii=False, indent=2)
                        print(f"📄 Relatório de fidelidade (preventiva) salvo: {fidelity_path.name}")
                        primary_fidelity_written = True
                    recomendacao = (preventive_result or {}).get("recomendacao_hil", {}) or {}
                    if hil_strict and recomendacao.get("pausar_para_revisao"):
                        save_hil_output(
                            formatted,
                            video_name,
                            folder,
                            mode_suffix,
                            reason="auditoria_preventiva_resumo",
                        )
                        raise HILCheckpointException(
                            f"Auditoria preventiva exige revisão humana. Veja: {preventive_md.name}"
                        )
                except HILCheckpointException:
                    raise
                except Exception as e:
                    print(f"{Fore.YELLOW}⚠️ Falha na auditoria preventiva: {e}. Continuando...")

            # Validação Full-Context (backup)
            if FIDELITY_BACKUP_ENABLED:
                try:
                    print(f"\n{Fore.CYAN}🔬 Validação Full-Context LLM (revalidação/backup)...")
                    validation_result = vomo.validate_completeness_full(
                        transcription, formatted, video_name, None
                    )
                    validation_report_path = Path(folder) / f"{video_name}_{mode_suffix}_fidelidade_backup.json"
                    with open(validation_report_path, "w", encoding='utf-8') as f:
                        json.dump(validation_result, f, ensure_ascii=False, indent=2)
                    print(f"📄 Relatório de fidelidade (backup) salvo: {validation_report_path.name}")

                    # v2.30: Se houve erro de validação, default é False (não mascara falha)
                    _def_aprov = False if validation_result.get('erro_validacao') else True
                    if not validation_result.get('aprovado', _def_aprov):
                        fidelity_md_path = Path(folder) / f"{video_name}_{mode_suffix}_REVISAO.md"
                        with open(fidelity_md_path, "w", encoding='utf-8') as f:
                            f.write(f"# ⚠️ REVISÃO NECESSÁRIA: {video_name}\n\n")
                            f.write(f"**Nota de Fidelidade:** {validation_result.get('nota', 0)}/10\n\n")
                            if validation_result.get('omissoes'):
                                f.write("## 📌 Omissões Detectadas\n")
                                for o in validation_result['omissoes']:
                                    f.write(f"- {o}\n")
                                f.write("\n")
                            if validation_result.get('distorcoes'):
                                f.write("## ⚠️ Distorções Detectadas\n")
                                for d in validation_result['distorcoes']:
                                    f.write(f"- {d}\n")
                                f.write("\n")
                            if validation_result.get('problemas_estrutura'):
                                f.write("## 🏗️ Problemas de Estrutura\n")
                                for p in validation_result['problemas_estrutura']:
                                    f.write(f"- {p}\n")
                                f.write("\n")
                            if validation_result.get('observacoes'):
                                f.write(f"## Observações\n{validation_result['observacoes']}\n")
                        print(f"{Fore.RED}📄 ATENÇÃO: Documento requer revisão! Veja: {fidelity_md_path.name}")
                except Exception as e:
                    print(f"{Fore.YELLOW}⚠️ Erro na revalidação Full-Context: {e}")

            if not primary_fidelity_written and isinstance(validation_result, dict):
                fidelity_path = Path(folder) / f"{video_name}_{mode_suffix}_fidelidade.json"
                try:
                    with open(fidelity_path, "w", encoding="utf-8") as f:
                        json.dump(validation_result, f, ensure_ascii=False, indent=2)
                    print(f"📄 Relatório de fidelidade (fallback) salvo: {fidelity_path.name}")
                except Exception as e:
                    print(f"{Fore.YELLOW}⚠️ Falha ao salvar relatório de fidelidade fallback: {e}")

            # Auditoria Jurídica (opcional)
            if AUDIT_AVAILABLE and not skip_audit:
                print(f"\n{Fore.CYAN}🕵️ Auditoria Jurídica Pós-Processamento...")
                audit_report_path = Path(folder) / f"{video_name}_{mode_suffix}_AUDITORIA.md"
                audit_content = auditar_consistencia_legal(
                    vomo.client,
                    formatted,
                    str(audit_report_path),
                    raw_transcript=transcription,
                )
                if audit_content:
                    print(f"{Fore.GREEN}   📎 Relatório de auditoria salvo em arquivo separado...")
        else:
            formatted = asyncio.run(vomo.format_transcription_async(
                transcription,
                video_name,
                folder,
                mode=mode,
                custom_prompt=custom_prompt,
                dry_run=dry_run,
                skip_audit=skip_audit,
                skip_fidelity_audit=skip_fidelity_audit,
                skip_sources_audit=skip_sources_audit,
                hil_strict=hil_strict,
            ))
        
        # v2.28+: Aplicar as mesmas correções de tabelas do DOCX também no Markdown final (APOSTILA).
        # Caso contrário, o usuário vê no .md uma tabela "no meio do assunto" que só é corrigida no DOCX.
        try:
            mode_label = (mode or "APOSTILA").upper()
        except Exception:
            mode_label = "APOSTILA"
        if formatted and mode_label == "APOSTILA":
            try:
                formatted = corrigir_tabelas_prematuras(formatted)
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Erro ao corrigir tabelas prematuras (MD): {e}.")
            try:
                formatted = mover_tabelas_para_fim_de_secao(formatted)
            except Exception as e:
                print(f"{Fore.YELLOW}⚠️ Erro ao reorganizar tabelas (MD): {e}. Usando layout padrão.")

        with open(os.path.join(folder, f"{video_name}_{mode}.md"), 'w') as f:
            f.write(formatted)
            
        # Validação de Fidelidade (Portado)
        try:
             verificar_cobertura(transcription, formatted, os.path.join(folder, f"{video_name}_{mode}.md"))
        except Exception as e:
             print(f"{Fore.YELLOW}⚠️ Erro na validação de fidelidade: {e}")

        vomo.save_as_word(formatted, video_name, folder, mode=mode)
        
        # v2.24: Auto-Fix Post-Processing (Structural Analysis - HIL Mode)
        md_output_path = os.path.join(folder, f"{video_name}_{mode}.md")
        auto_apply_flag = auto_apply_fixes
        
        if AUTO_FIX_AVAILABLE:
            print(f"\n{Fore.CYAN}🔧 Auto-Fix Structural Analysis (v2.24 HIL)...")
            try:
                issues = analyze_structural_issues(md_output_path)
                if issues['total_issues'] > 0:
                    print(f"{Fore.YELLOW}   ⚠️ {issues['total_issues']} problema(s) estrutural(is) detectado(s).")
                    print(f"   Seções duplicadas: {len(issues['duplicate_sections'])}")
                    print(f"   Parágrafos duplicados: {len(issues['duplicate_paragraphs'])}")
                    
                    # Save suggestions to JSON for HIL review
                    suggestions_path = os.path.join(folder, f"{video_name}_{mode}_SUGESTOES.json")
                    import json
                    with open(suggestions_path, 'w', encoding='utf-8') as f:
                        json.dump(issues, f, indent=2, ensure_ascii=False)
                    print(f"{Fore.CYAN}   📋 Sugestões salvas em: {os.path.basename(suggestions_path)}")
                    
                    if auto_apply_flag:
                        # Auto-apply fixes (only with explicit flag)
                        result = apply_structural_fixes_to_file(md_output_path, issues)
                        if result['fixes_applied']:
                            print(f"{Fore.GREEN}   ✅ {len(result['fixes_applied'])} correção(ões) aplicada(s) automaticamente.")
                            # Regenerate Word document with fixed content
                            with open(md_output_path, 'r', encoding='utf-8') as f:
                                fixed_content = f.read()
                            vomo.save_as_word(fixed_content, video_name, folder, mode=mode)
                    else:
                        print(f"{Fore.YELLOW}   ℹ️  Modo HIL: Revise as sugestões e use --auto-apply-fixes para aplicar.")
                else:
                    print(f"{Fore.GREEN}   ✅ Nenhum problema estrutural detectado.")
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Erro no Auto-Fix: {e}")

        # v3.0: Relatório Unificado (cross-referencing entre camadas)
        if UNIFIED_AUDIT_AVAILABLE:
            try:
                engine = UnifiedAuditEngine(video_name, mode)
                # Ingerir auditoria preventiva (já salva em disco)
                prev_json = Path(folder) / f"{video_name}_{mode}_AUDITORIA_FIDELIDADE.json"
                if prev_json.exists():
                    with open(prev_json, "r", encoding="utf-8") as f:
                        engine.ingest_fidelity(json.load(f))
                # Ingerir backup
                backup_json = Path(folder) / f"{video_name}_{mode}_fidelidade_backup.json"
                if backup_json.exists():
                    with open(backup_json, "r", encoding="utf-8") as f:
                        engine.ingest_backup(json.load(f))
                # Ingerir structural (variável local do bloco acima)
                if AUTO_FIX_AVAILABLE and 'issues' in dir() and isinstance(issues, dict):
                    engine.ingest_structural(issues)
                unified_report = engine.build()
                # Salvar
                unified_json_path = os.path.join(folder, f"{video_name}_{mode}_UNIFIED_AUDIT.json")
                unified_md_path = os.path.join(folder, f"{video_name}_{mode}_UNIFIED_AUDIT.md")
                unified_report.save_json(unified_json_path)
                generate_unified_markdown(unified_report, unified_md_path)
                summary = unified_report.summary
                print(f"\n{Fore.CYAN}📊 Relatório Unificado (v3.0):")
                print(f"   Nota Geral: {unified_report.nota_geral:.1f}/10 (Fidelidade: {unified_report.nota_fidelidade:.1f} | Estrutural: {unified_report.nota_estrutural:.1f})")
                print(f"   Findings: {summary.get('total_findings', 0)} ativos ({summary.get('false_positives_removed', 0)} FP removidos)")
                by_sev = summary.get('by_severity', {})
                sev_parts = [f"{k}: {v}" for k, v in by_sev.items() if v > 0]
                if sev_parts:
                    print(f"   Severidade: {' | '.join(sev_parts)}")
                print(f"   Salvos: {os.path.basename(unified_json_path)} + {os.path.basename(unified_md_path)}")
                # Comparação com relatório anterior (resume-hil)
                if '_previous_unified' in dir() and _previous_unified is not None:
                    try:
                        delta = compare_reports(_previous_unified, unified_report)
                        print(f"   📈 Delta HIL: {delta['resolved_count']} resolvidos | {delta['persistent_count']} persistentes | {delta['new_count']} novos")
                    except Exception:
                        pass
                # HIL unificado
                if hil_strict and unified_report.hil_recommendation.pausar:
                    save_hil_output(formatted, video_name, folder, mode_suffix, reason="unified_audit")
                    raise HILCheckpointException(
                        f"Auditoria unificada requer revisão. Veja: {os.path.basename(unified_md_path)}"
                    )
            except HILCheckpointException:
                raise
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Erro no relatório unificado: {e}")

        print(f"{Fore.GREEN}✨ SUCESSO!")
        
    except HILCheckpointException as e:
        print(f"{Fore.YELLOW}⏸️  HIL checkpoint acionado: {e}")
        return
    except Exception as e:
        print(f"{Fore.RED}❌ Erro: {e}")
        traceback.print_exc()

def _parse_mode(value):
    normalized = value.strip().upper()
    allowed = {"APOSTILA", "FIDELIDADE", "AUDIENCIA", "REUNIAO", "DEPOIMENTO"}
    if normalized not in allowed:
        raise argparse.ArgumentTypeError(f"Modo invalido: {value}. Use {', '.join(sorted(allowed))}.")
    return normalized


def _parse_provider(value):
    normalized = value.strip().lower()
    allowed = {"gemini", "openai"}
    if normalized not in allowed:
        raise argparse.ArgumentTypeError(f"Provider invalido: {value}. Use {', '.join(sorted(allowed))}.")
    return normalized


def _load_custom_prompt(prompt_value):
    if not prompt_value:
        return None
    if prompt_value.endswith('.txt') and os.path.exists(prompt_value):
        with open(prompt_value, 'r', encoding='utf-8') as f:
            custom_prompt = f.read()
        print(f"{Fore.YELLOW}📝 Prompt carregado de arquivo: {prompt_value} ({len(custom_prompt):,} chars)")
        return custom_prompt
    custom_prompt = prompt_value
    print(f"{Fore.YELLOW}📝 Usando prompt direto ({len(custom_prompt):,} chars)")
    return custom_prompt


def _build_arg_parser():
    parser = argparse.ArgumentParser(
        description="MLX Vomo - transcricao e formatacao juridica.",
    )
    parser.add_argument("input_file", nargs="?", help="Arquivo de video/audio ou .txt/.md de entrada.")
    parser.add_argument("--mode", type=_parse_mode, default="FIDELIDADE",
                        help="Modo de formatacao: APOSTILA, FIDELIDADE, AUDIENCIA, REUNIAO, DEPOIMENTO.")
    parser.add_argument("--provider", type=_parse_provider, default="gemini",
                        help="Provider LLM: gemini ou openai.")
    parser.add_argument(
        "--prompt",
        help=(
            "Prompt customizado (texto direto ou arquivo .txt). "
            "Em APOSTILA/AUDIENCIA/REUNIAO: personaliza apenas TABELAS/EXTRAS (resumo/fluxograma/mapa mental/questionário), "
            "preservando tom/estilo/estrutura. Em outros modos: substitui STYLE+TABLE; HEAD/STRUCTURE/FOOTER são preservados."
        ),
    )
    parser.add_argument("--dry-run", action="store_true", help="Executa apenas etapas locais.")
    parser.add_argument("--skip-formatting", action="store_true", help="Pula a formatacao final.")
    parser.add_argument("--high-accuracy", action="store_true", help="Usa beam search na transcricao.")
    diar_group = parser.add_mutually_exclusive_group()
    diar_group.add_argument(
        "--diarization",
        dest="diarization",
        action="store_true",
        help="Força diarização ON (override do padrão por modo).",
    )
    diar_group.add_argument(
        "--no-diarization",
        dest="diarization",
        action="store_false",
        help="Força diarização OFF (override do padrão por modo).",
    )
    parser.set_defaults(diarization=None)
    parser.add_argument(
        "--diarization-strict",
        action="store_true",
        help="Falha se diarização estiver indisponível (útil quando opt-in em APOSTILA/FIDELIDADE).",
    )
    parser.add_argument("--skip-fidelity-audit", action="store_true", help="Desativa auditoria de fidelidade.")
    parser.add_argument("--skip-sources-audit", action="store_true", help="Desativa auditoria de fontes.")
    parser.add_argument("--hil-strict", action="store_true", help="Habilita checkpoint estrito de HIL.")
    parser.add_argument("--resume-hil", action="store_true", help="Retoma a partir do checkpoint HIL.")
    parser.add_argument("--word-only", action="store_true", help="Gera DOCX a partir de um .md existente.")
    parser.add_argument("--auto-apply-fixes", action="store_true", help="Aplica correcoes estruturais automaticamente.")
    parser.add_argument("--no-audit", "--skip-legal-audit", "--skip-audit",
                        dest="skip_audit", action="store_true",
                        help="Desativa auditoria juridica.")
    return parser


if __name__ == "__main__":
    parser = _build_arg_parser()
    args = parser.parse_args()

    if not args.input_file:
        parser.print_usage()
        sys.exit(1)

    custom_prompt = _load_custom_prompt(args.prompt)
    print(f"{Fore.CYAN}🚀 Iniciando MLX Vomo (v2.26)...")
    try:
        process_single_video(
            args.input_file,
            mode=args.mode,
            custom_prompt=custom_prompt,
            dry_run=args.dry_run,
            high_accuracy=args.high_accuracy,
            diarization=args.diarization,
            diarization_strict=args.diarization_strict,
            skip_formatting=args.skip_formatting,
            skip_audit=args.skip_audit,
            skip_fidelity_audit=args.skip_fidelity_audit,
            skip_sources_audit=args.skip_sources_audit,
            hil_strict=args.hil_strict,
            resume_hil=args.resume_hil,
            provider=args.provider,
            word_only=args.word_only,
            auto_apply_fixes=args.auto_apply_fixes,
        )
    except KeyboardInterrupt:
        print(f"\n{Fore.YELLOW}⚠️ Interrupção pelo usuário.")
        sys.exit(0)
    except Exception as e:
        print(f"\n{Fore.RED}❌ Erro fatal: {e}")
        sys.exit(1)
