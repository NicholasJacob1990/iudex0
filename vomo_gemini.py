#!/usr/bin/env python3
"""
VomoGemini - Transcrição e Formatação de Aulas com Gemini 2.5 Flash
Versão Híbrida Otimizada - Combina o melhor de ambas implementações
Otimizado para Apple Silicon (M3 Pro) + Google Gemini API
"""

import os
import sys
import time
import subprocess
import hashlib
import json
import re
from pathlib import Path
from colorama import Fore, init
from tqdm import tqdm
import google.generativeai as genai
from google.generativeai.types import GenerationConfig

init(autoreset=True)

# Imports opcionais
try:
    import mlx_whisper
except ImportError:
    mlx_whisper = None
    print(f"{Fore.YELLOW}⚠️ mlx_whisper não instalado. Use: pip install mlx-whisper")

try:
    from pyannote.audio import Pipeline
    import torch
    HF_TOKEN = os.getenv("HUGGING_FACE_TOKEN")
except ImportError:
    Pipeline = None
    torch = None
    HF_TOKEN = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print(f"{Fore.YELLOW}⚠️ python-docx não disponível. Apenas Markdown será gerado.")

# ==================== CONFIGURAÇÃO ====================

# Preços Gemini 2.5 Flash (Dezembro 2024)
GEMINI_PRICE_INPUT = 0.075   # USD por 1M tokens
GEMINI_PRICE_OUTPUT = 0.30   # USD por 1M tokens

SYSTEM_PROMPT = """# SEU PAPEL
Você é um **revisor sênior de material didático jurídico** com 20 anos de experiência em preparatórios para concursos de procuradoria (PGM/PGE). Seu trabalho é transformar transcrições brutas de aulas em apostilas de alta qualidade, mantendo 100% do conteúdo técnico.

---

# MISSÃO

Transformar a transcrição de aula em uma **apostila didática** com estilo formal, impessoal e expositivo, **PRESERVANDO TODO O CONTEÚDO ORIGINAL**.

⚠️ **REGRA ABSOLUTA: PRESERVE 100% DO CONTEÚDO - NÃO RESUMA, NÃO OMITA NADA.**

---

# EXEMPLOS DE TRANSFORMAÇÃO (Few-Shot)

## Exemplo 1: Transformação de Linguagem
**INPUT (Transcrição bruta):**
> "Então pessoal, a Lei 8.666, né, ela fala lá no artigo 37 que a gente não pode fazer contratação direta, tá? Exceto em alguns casos, tipo... quando tem emergência, sabe?"

**OUTPUT (Apostila formatada):**
> Observa-se que o *Art. 37 da Lei nº 8.666/93* estabelece a vedação à contratação direta como regra geral. Contudo, o legislador previu hipóteses excepcionais, notadamente os casos de emergência, nos termos do *Art. 24, inciso IV* do referido diploma legal.

## Exemplo 2: Preservação de Dicas de Prova
**INPUT:**
> "Isso aqui cai muito em prova, hein! Atenção total nesse artigo 37."

**OUTPUT:**
> Reveste-se de especial importância este dispositivo para fins de certames públicos, sendo recorrentemente objeto de cobrança em provas de procuradorias.

## Exemplo 3: Preservação de Exemplos Completos
**INPUT:**
> "Por exemplo, teve um caso no Rio que o prefeito tentou fazer uma contratação emergencial pra comprar uns equipamentos, mas o TCE glosou porque não tinha a urgência real."

**OUTPUT:**
> A título ilustrativo, cumpre mencionar caso ocorrido no Estado do Rio de Janeiro, no qual determinado prefeito municipal tentou realizar contratação emergencial para aquisição de equipamentos. Contudo, o Tribunal de Contas do Estado glosou o procedimento, fundamentando sua decisão na ausência de urgência real que justificasse a dispensa de licitação.

---

# DIRETRIZES DE FORMATAÇÃO

1. **Estrutura Hierárquica**: Use títulos Markdown (##, ###, ####) para organizar tópicos
2. **Linguagem**: Terceira pessoa impessoal, tempo presente
3. **Citações**: Itálico para leis (*Lei nº 8.666/93*), negrito para conceitos-chave
4. **Tabelas**: Crie tabelas comparativas ao final de tópicos complexos
5. **Preservação Total**: Mantenha TODOS os exemplos, casos, jurisprudência e dicas
6. **Numeração**: Numere os tópicos principais (1., 2., 3.) e subtópicos (1.1, 1.2)

---

# ❌ ERROS A EVITAR

- ❌ Resumir exemplos ou casos práticos
- ❌ Omitir artigos de lei ou números
- ❌ Inventar conteúdo não presente na transcrição
- ❌ Usar primeira pessoa ("eu", "nós")
- ❌ Reorganizar cronologicamente (mantenha ordem da aula)

---

# AUTO-VERIFICAÇÃO (EXECUTE ANTES DE ENVIAR)

□ Mantive TODOS os artigos de lei com números corretos?
□ Mantive TODOS os exemplos e casos completos?
□ Mantive TODAS as dicas de prova?
□ Criei tabelas ao final de tópicos principais?
□ Usei terceira pessoa em TODO o texto?
□ Output tem pelo menos 80% do tamanho do input?

---

# FORMATO DE RESPOSTA

Retorne **APENAS** a apostila em Markdown, começando com # Título da Apostila, sem meta-comentários.
"""


# ==================== CLASSE PRINCIPAL ====================

class VomoGemini:
    def __init__(self, model_size="large-v3-turbo", gemini_model="gemini-2.5-flash-preview-05-20"):
        """
        Transcrição com MLX-Whisper + Formatação com Gemini 2.5 Flash
        
        Args:
            model_size: Modelo Whisper (large-v3-turbo recomendado)
            gemini_model: gemini-2.5-flash-preview-05-20 (65k output)
        """
        print(f"{Fore.CYAN}🚀 Inicializando VomoGemini...")
        print(f"   🎙️  Whisper: {model_size}")
        print(f"   🧠 LLM: {gemini_model}")
        
        self.model_name = model_size
        self.gemini_model = gemini_model
        
        # Configurar Gemini
        from dotenv import load_dotenv
        load_dotenv(override=True)
        
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError(f"{Fore.RED}❌ Configure a variável GEMINI_API_KEY no .env")
        
        genai.configure(api_key=api_key)
        
        self.model = genai.GenerativeModel(
            model_name=gemini_model,
            system_instruction=SYSTEM_PROMPT
        )
        
        # Cache directory
        self.cache_dir = Path(".cache_vomo")
        self.cache_dir.mkdir(exist_ok=True)
        
        print(f"{Fore.GREEN}✅ Inicialização concluída!\n")

    # ==================== TRANSCRIÇÃO ====================
    
    def optimize_audio(self, file_path):
        """Extrai/converte áudio para formato otimizado"""
        print(f"{Fore.YELLOW}⚡ Verificando áudio...{Fore.RESET}")
        
        # Prioriza MP3 se existir
        mp3_path = Path(file_path).with_suffix('.mp3')
        if mp3_path.exists():
            print(f"   📂 Usando MP3 existente: {mp3_path.name}")
            return str(mp3_path)
        
        # Gera WAV otimizado
        file_hash = hashlib.md5(file_path.encode()).hexdigest()[:8]
        output_path = f"temp_{file_hash}.wav"
        
        if os.path.exists(output_path):
            print(f"   📂 WAV já existe: {output_path}")
            return output_path
        
        print(f"   🔄 Extraindo áudio otimizado (16kHz mono)...")
        try:
            subprocess.run([
                'ffmpeg', '-i', file_path,
                '-vn', '-ac', '1', '-ar', '16000',
                output_path, '-y', '-hide_banner', '-loglevel', 'error'
            ], check=True, capture_output=True)
            print(f"   ✅ Áudio pronto: {output_path}")
            return output_path
        except subprocess.CalledProcessError as e:
            print(f"{Fore.RED}❌ Erro no ffmpeg: {e}")
            sys.exit(1)

    def transcribe(self, audio_path):
        """
        Transcrição com MLX-Whisper (GPU otimizado) + Diarização opcional
        """
        if not mlx_whisper:
            raise ImportError("mlx_whisper não disponível. Instale: pip install mlx-whisper")
        
        print(f"{Fore.GREEN}🎙️  Iniciando transcrição (MLX GPU)...{Fore.RESET}")
        start_time = time.time()
        
        # Verifica cache
        cache_file = Path(audio_path).with_suffix('.transcription.json')
        if cache_file.exists():
            try:
                print(f"{Fore.CYAN}   📂 Cache encontrado, carregando...")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                print(f"{Fore.GREEN}   ✅ Transcrição carregada do cache!")
                return cache_data['transcript']
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Erro ao ler cache: {e}")
        
        # Transcrição com MLX-Whisper
        print("   🔍 Transcrevendo com parâmetros otimizados...")
        
        result = mlx_whisper.transcribe(
            audio_path,
            path_or_hf_repo=f"mlx-community/whisper-{self.model_name}",
            language="pt",
            word_timestamps=True,
            fp16=True,
            no_speech_threshold=0.6,
            logprob_threshold=-1.0,
            compression_ratio_threshold=2.4,
            condition_on_previous_text=True,
            verbose=False
        )
        
        elapsed = time.time() - start_time
        audio_duration = result.get('duration', 0) if isinstance(result, dict) else 0
        rtf = elapsed / audio_duration if audio_duration > 0 else 0
        
        print(f"{Fore.GREEN}   ✅ Transcrição concluída em {elapsed:.1f}s (RTF: {rtf:.2f}x)")
        
        # Tenta diarização
        transcript_result = self._try_diarization(audio_path, result)
        
        # Salva cache
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'transcript': transcript_result,
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
                    'duration': audio_duration
                }, f, ensure_ascii=False, indent=2)
            print(f"   💾 Cache salvo: {cache_file.name}")
        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ Erro ao salvar cache: {e}")
        
        return transcript_result

    def _try_diarization(self, audio_path, whisper_result):
        """Tenta diarização com Pyannote (opcional)"""
        if not (Pipeline and torch and HF_TOKEN):
            return self._format_simple(whisper_result)
        
        try:
            print("   🗣️  Iniciando Diarização (Pyannote)...")
            pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                use_auth_token=HF_TOKEN
            )
            
            device = "mps" if torch.backends.mps.is_available() else "cpu"
            pipeline.to(torch.device(device))
            
            diarization = pipeline(audio_path)
            result = self._align_diarization(whisper_result['segments'], diarization)
            print(f"{Fore.GREEN}   ✅ Diarização concluída!")
            return result
            
        except Exception as e:
            print(f"{Fore.YELLOW}   ⚠️ Diarização falhou ({e}), usando transcrição simples")
            return self._format_simple(whisper_result)

    def _format_simple(self, result):
        """Formatação simples sem diarização"""
        output = []
        for segment in result['segments']:
            ts = self._format_timestamp(segment['start'])
            output.append(f"[{ts}] {segment['text'].strip()}")
        return " ".join(output)

    def _format_timestamp(self, seconds):
        """Converte segundos em HH:MM:SS ou MM:SS"""
        m, s = divmod(seconds, 60)
        h, m = divmod(m, 60)
        return f"{int(h):02d}:{int(m):02d}:{int(s):02d}" if h > 0 else f"{int(m):02d}:{int(s):02d}"

    def _align_diarization(self, segments, diarization_output):
        """Alinha transcrição com speakers identificados"""
        output = []
        current_speaker = None
        diar_segs = [(t.start, t.end, s) for t, _, s in diarization_output.itertracks(yield_label=True)]

        for segment in segments:
            start, end = segment['start'], segment['end']
            best_speaker = "SPEAKER 0"
            max_overlap = 0

            for d_start, d_end, d_speaker in diar_segs:
                overlap = max(0, min(end, d_end) - max(start, d_start))
                if overlap > max_overlap:
                    max_overlap = overlap
                    best_speaker = f"SPEAKER {int(d_speaker.split('_')[-1]) + 1}"
            
            if best_speaker != current_speaker:
                output.append(f"\n\n{best_speaker}\n")
                current_speaker = best_speaker
            
            output.append(f"[{self._format_timestamp(start)}] {segment['text'].strip()}")
            
        return " ".join(output)

    # ==================== FORMATAÇÃO COM GEMINI ====================
    
    def format_with_gemini(self, transcription, video_name):
        """
        Formata a transcrição completa usando Gemini 2.5 Flash em UMA ÚNICA CHAMADA
        """
        print(f"\n{Fore.CYAN}{'='*70}")
        print(f"{Fore.CYAN}🧠 FORMATAÇÃO COM GEMINI 2.5 FLASH")
        print(f"{Fore.CYAN}{'='*70}{Fore.RESET}")
        
        trans_size = len(transcription)
        trans_tokens_est = trans_size // 4  # ~4 chars por token
        
        print(f"   📊 Tamanho da transcrição: {trans_size:,} caracteres (~{trans_tokens_est:,} tokens)")
        print(f"   📹 Vídeo: {video_name}")
        
        if trans_tokens_est > 1_000_000:
            print(f"{Fore.YELLOW}   ⚠️ AVISO: Transcrição muito grande, pode exceder limite do Gemini!")
            print(f"   💡 Considere dividir o vídeo em partes menores.{Fore.RESET}")
        
        # Verifica cache
        cache_hash = hashlib.sha256(transcription.encode()).hexdigest()[:16]
        cache_file = self.cache_dir / f"formatted_{cache_hash}.md"
        
        if cache_file.exists():
            try:
                print(f"{Fore.CYAN}   📂 Apostila encontrada em cache!")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cached_result = f.read()
                    # Mostra estimativa mesmo com cache
                    output_tokens_est = len(cached_result) // 4
                    cost = (trans_tokens_est * GEMINI_PRICE_INPUT + output_tokens_est * GEMINI_PRICE_OUTPUT) / 1_000_000
                    print(f"   💰 Custo estimado (se fosse chamada real): ${cost:.4f} USD")
                    return cached_result
            except:
                pass
        
        # Configuração para output massivo (65k tokens)
        config = GenerationConfig(
            temperature=0.2,
            max_output_tokens=65536,  # Limite máximo do Gemini 2.5 Flash
            top_p=0.95,
        )
        
        # Prepara prompt
        user_prompt = f"""# AULA PARA FORMATAR

**Título**: {video_name}

**TRANSCRIÇÃO COMPLETA**:

{transcription}

---

Agora gere a apostila completa em Markdown seguindo rigorosamente as diretrizes do system prompt.
"""
        
        print(f"   🚀 Enviando para Gemini API...")
        start_time = time.time()
        
        try:
            response = self.model.generate_content(
                user_prompt,
                generation_config=config
            )
            
            elapsed = time.time() - start_time
            
            if not response.text:
                raise ValueError("Gemini retornou resposta vazia")
            
            formatted_text = response.text.strip()
            output_tokens_est = len(formatted_text) // 4
            
            # Calcula custo real
            cost = (trans_tokens_est * GEMINI_PRICE_INPUT + output_tokens_est * GEMINI_PRICE_OUTPUT) / 1_000_000
            
            print(f"{Fore.GREEN}   ✅ Formatação concluída em {elapsed:.1f}s")
            print(f"   📝 Tamanho output: {len(formatted_text):,} chars (~{output_tokens_est:,} tokens)")
            print(f"   💰 Custo estimado: ${cost:.4f} USD")
            
            # Salva cache
            try:
                with open(cache_file, 'w', encoding='utf-8') as f:
                    f.write(formatted_text)
                print(f"   💾 Cache salvo: {cache_file.name}")
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Erro ao salvar cache: {e}")
            
            return formatted_text
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro na chamada Gemini: {e}")
            
            # Informações de debug
            if hasattr(e, 'message'):
                print(f"   🔍 Detalhes: {e.message}")
            
            raise e

    # ==================== EXPORTAÇÃO DOCX ====================
    
    def save_docx(self, markdown_text, video_name, output_path):
        """Converte Markdown para DOCX profissional"""
        if not DOCX_AVAILABLE:
            print(f"{Fore.YELLOW}   ⚠️ python-docx não disponível, pulando DOCX")
            return None
        
        print(f"{Fore.CYAN}📄 Gerando DOCX...")
        
        doc = Document()
        
        # Margens
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        
        # Título
        title = doc.add_heading(video_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Data
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')} - Gemini 2.5 Flash")
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Processa Markdown
        lines = markdown_text.split('\n')
        in_table = False
        table_rows = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Tabelas
            if '|' in line and not line.startswith('|--'):
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append([c.strip() for c in line.split('|')[1:-1]])
                continue
            elif in_table:
                self._add_table(doc, table_rows)
                in_table = False
                table_rows = []
            
            # Headers
            if line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:], style='Quote')
                p.paragraph_format.left_indent = Inches(0.5)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                self._format_inline(p, line[2:])
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                self._format_inline(p, line)
        
        # Tabela final
        if table_rows:
            self._add_table(doc, table_rows)
        
        doc.save(output_path)
        print(f"{Fore.GREEN}   ✅ DOCX salvo: {output_path}")
        return output_path

    def _format_inline(self, paragraph, text):
        """Formata bold/italic inline"""
        paragraph.clear()
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            if match.group(0).startswith('***'):
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(0).startswith('**'):
                run = paragraph.add_run(match.group(3))
                run.bold = True
            elif match.group(0).startswith('*'):
                run = paragraph.add_run(match.group(4))
                run.italic = True
            elif match.group(0).startswith('`'):
                run = paragraph.add_run(match.group(5))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(200, 0, 0)
            
            last_end = match.end()
        
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_table(self, doc, rows):
        """Adiciona tabela formatada"""
        if len(rows) < 2:
            return
        
        max_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        
        for i, row in enumerate(rows):
            for j in range(max_cols):
                if j < len(row):
                    cell = table.rows[i].cells[j]
                    self._format_inline(cell.paragraphs[0], row[j])
                    
                    if i == 0:
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '0066CC')
                        cell._element.get_or_add_tcPr().append(shading_elm)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(255, 255, 255)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ==================== PIPELINE COMPLETO ====================
    
    def process_video(self, video_path, output_folder="apostilas", export_docx=True):
        """
        Pipeline completo: Vídeo → Transcrição → Apostila (MD + DOCX)
        """
        print(f"\n{Fore.CYAN}{'='*70}")
        print(f"{Fore.CYAN}🎬 PROCESSANDO VÍDEO")
        print(f"{Fore.CYAN}{'='*70}{Fore.RESET}\n")
        
        video_path = Path(video_path)
        if not video_path.exists():
            print(f"{Fore.RED}❌ Arquivo não encontrado: {video_path}")
            sys.exit(1)
        
        video_name = video_path.stem
        output_folder = Path(output_folder)
        output_folder.mkdir(exist_ok=True)
        
        md_file = output_folder / f"{video_name}.md"
        docx_file = output_folder / f"{video_name}.docx"
        
        print(f"📹 Vídeo: {video_path.name}")
        print(f"📁 Output: {output_folder}\n")
        
        # Etapa 1: Otimizar áudio
        audio_path = self.optimize_audio(str(video_path))
        
        # Etapa 2: Transcrever
        transcription = self.transcribe(audio_path)
        
        # Salva transcrição bruta
        raw_file = output_folder / f"{video_name}_RAW.txt"
        with open(raw_file, 'w', encoding='utf-8') as f:
            f.write(transcription)
        print(f"   💾 Transcrição bruta salva: {raw_file.name}")
        
        # Etapa 3: Formatar com Gemini
        formatted_text = self.format_with_gemini(transcription, video_name)
        
        # Etapa 4: Salvar Markdown
        print(f"\n{Fore.CYAN}💾 Salvando apostila...{Fore.RESET}")
        try:
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(formatted_text)
            
            file_size = md_file.stat().st_size / 1024  # KB
            print(f"{Fore.GREEN}   ✅ Markdown salvo: {md_file.name} ({file_size:.1f} KB)")
            
        except Exception as e:
            print(f"{Fore.RED}❌ Erro ao salvar Markdown: {e}")
            sys.exit(1)
        
        # Etapa 5: Exportar DOCX
        if export_docx and DOCX_AVAILABLE:
            try:
                self.save_docx(formatted_text, video_name, docx_file)
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Erro ao gerar DOCX: {e}")
        
        # Resumo final
        print(f"\n{Fore.GREEN}{'='*70}")
        print(f"{Fore.GREEN}✅ PROCESSAMENTO CONCLUÍDO!")
        print(f"{Fore.GREEN}{'='*70}{Fore.RESET}")
        print(f"📄 Markdown: {md_file.absolute()}")
        if export_docx and DOCX_AVAILABLE:
            print(f"📄 DOCX: {docx_file.absolute()}")
        
        return str(md_file)


# ==================== CLI ====================

def main():
    import argparse
    
    parser = argparse.ArgumentParser(
        description="VomoGemini - Transcrição e Formatação com Gemini 2.5 Flash",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos de uso:
  python vomo_gemini.py video.mp4
  python vomo_gemini.py aula.mp3 -o apostilas_out
  python vomo_gemini.py video.mov --model large-v3 --no-docx
        """
    )
    
    parser.add_argument('video', help='Caminho do arquivo de vídeo/áudio')
    parser.add_argument('-o', '--output', default='apostilas', 
                       help='Pasta de saída (padrão: apostilas)')
    parser.add_argument('--model', default='large-v3-turbo',
                       help='Modelo Whisper (padrão: large-v3-turbo)')
    parser.add_argument('--gemini', default='gemini-2.5-flash-preview-05-20',
                       help='Modelo Gemini (padrão: gemini-2.5-flash-preview-05-20)')
    parser.add_argument('--no-docx', action='store_true',
                       help='Não gerar arquivo DOCX')
    
    args = parser.parse_args()
    
    try:
        vomo = VomoGemini(model_size=args.model, gemini_model=args.gemini)
        vomo.process_video(args.video, args.output, export_docx=not args.no_docx)
        
    except KeyboardInterrupt:
        print(f"\n{Fore.YELLOW}⚠️ Processo interrompido pelo usuário")
        sys.exit(0)
    except Exception as e:
        print(f"\n{Fore.RED}❌ Erro fatal: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
