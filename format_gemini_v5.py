#!/usr/bin/env python3
"""
VERSÃO FINAL: Chunks de 10k com OVERLAP de 1000 chars para preservar contexto
Inclui deduplicação inteligente na junção dos chunks
"""

import os
import sys
import time
import re
import requests
from pathlib import Path
from docx import Document
from typing import List
from difflib import SequenceMatcher

# Configuração
OPENROUTER_API_KEY = "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"
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL_NAME = "google/gemini-2.0-flash-001"
MAX_CHARS_PER_CHUNK = 10_000
OVERLAP_CHARS = 1_000  # Overlap para contexto
OUTPUT_DIR = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ')

PROMPT_FORMATACAO = """REVISOR DE TEXTO - NÃO RESUMIR!

⛔ REGRA: O texto de saída deve ter tamanho IGUAL OU MAIOR que a entrada!

FAÇA:
✅ Corrija gramática e pontuação
✅ Transforme linguagem coloquial em formal
✅ Organize em parágrafos claros
✅ MANTENHA TODO o conteúdo (exemplos, casos, dicas)
✅ Use **negrito** para termos importantes

NÃO FAÇA:
❌ Não resuma
❌ Não omita informações
❌ Não corte exemplos

[Parte {parte}/{total}]

TEXTO ORIGINAL:
{texto}

TEXTO FORMATADO:"""


def extrair_texto(caminho: str) -> str:
    path = Path(caminho)
    if path.suffix.lower() == '.docx':
        doc = Document(caminho)
        return '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    with open(caminho, 'r', encoding='utf-8') as f:
        return f.read()


def dividir_com_overlap(texto: str, max_chars: int = MAX_CHARS_PER_CHUNK, overlap: int = OVERLAP_CHARS) -> List[str]:
    """Divide texto em chunks com overlap para preservar contexto"""
    chunks = []
    start = 0
    
    while start < len(texto):
        end = min(start + max_chars, len(texto))
        
        # Ajusta para não cortar no meio de palavra
        if end < len(texto):
            # Procura por quebra natural (parágrafo ou ponto)
            search_zone = texto[max(0, end-500):end]
            last_para = search_zone.rfind('\n\n')
            last_period = search_zone.rfind('. ')
            
            if last_para != -1:
                end = end - 500 + last_para
            elif last_period != -1:
                end = end - 500 + last_period + 1
        
        chunk = texto[start:end]
        chunks.append(chunk)
        
        if end >= len(texto):
            break
            
        # Próximo chunk começa com overlap
        start = end - overlap
    
    return chunks


def deduplicate_overlap(chunk_anterior: str, chunk_atual: str) -> str:
    """Remove conteúdo duplicado do overlap entre chunks"""
    if not chunk_anterior:
        return chunk_atual
    
    # Pega últimas 200 palavras do chunk anterior
    palavras_anterior = chunk_anterior.split()[-200:]
    tail = ' '.join(palavras_anterior)
    
    # Pega primeiras 200 palavras do chunk atual
    palavras_atual = chunk_atual.split()[:200]
    head = ' '.join(palavras_atual)
    
    # Encontra sobreposição
    matcher = SequenceMatcher(None, tail, head)
    match = matcher.find_longest_match(0, len(tail), 0, len(head))
    
    if match.size > 50:  # Se há sobreposição significativa
        # Remove a parte sobreposta do início do chunk atual
        overlap_chars = match.b + match.size
        palavras_para_pular = len(head[:overlap_chars].split())
        return ' '.join(chunk_atual.split()[palavras_para_pular:])
    
    return chunk_atual


def processar_chunk(texto: str, parte: int, total: int) -> str:
    """Processa via OpenRouter"""
    
    prompt = PROMPT_FORMATACAO.format(texto=texto, parte=parte, total=total)
    
    print(f"   📝 {parte}/{total}: {len(texto):,} chars", end="", flush=True)
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://iudex.app"
    }
    
    payload = {
        "model": MODEL_NAME,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.05,
        "max_tokens": 16000
    }
    
    response = requests.post(OPENROUTER_BASE_URL, headers=headers, json=payload, timeout=300)
    response.raise_for_status()
    
    resultado = response.json()['choices'][0]['message']['content']
    
    ratio = len(resultado) / len(texto) * 100
    status = "✅" if ratio >= 90 else "⚠️"
    print(f" → {len(resultado):,} ({ratio:.0f}%) {status}")
    
    return resultado


def salvar_arquivos(texto: str, nome_base: str):
    """Salva MD e DOCX"""
    md_path = OUTPUT_DIR / f'{nome_base}_FORMATADO_FINAL.md'
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(texto)
    print(f"✅ MD: {md_path}")
    
    docx_path = OUTPUT_DIR / f'{nome_base}_FORMATADO_FINAL.docx'
    doc = Document()
    for linha in texto.split('\n'):
        if linha.strip():
            doc.add_paragraph(linha)
    doc.save(str(docx_path))
    print(f"✅ DOCX: {docx_path}")


def main():
    arquivo = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/Administrativo.txt')
    
    if not arquivo.exists():
        print(f"❌ Não encontrado: {arquivo}")
        return
    
    print(f"\n📄 Lendo: {arquivo.name}")
    texto = extrair_texto(str(arquivo))
    print(f"   {len(texto):,} caracteres\n")
    
    print(f"📊 Chunks de {MAX_CHARS_PER_CHUNK:,} chars c/ overlap de {OVERLAP_CHARS:,}")
    chunks = dividir_com_overlap(texto)
    print(f"   {len(chunks)} partes\n")
    
    resultados = []
    texto_anterior = ""
    
    for i, chunk in enumerate(chunks, 1):
        try:
            resultado = processar_chunk(chunk, i, len(chunks))
            
            # Deduplica overlap
            if i > 1:
                resultado = deduplicate_overlap(texto_anterior, resultado)
            
            resultados.append(resultado)
            texto_anterior = resultado
            
            if i < len(chunks):
                time.sleep(0.3)
                
        except Exception as e:
            print(f" ❌ Erro: {e}")
            resultados.append(chunk)
    
    texto_final = '\n\n'.join(resultados)
    
    print(f"\n💾 Salvando...")
    salvar_arquivos(texto_final, arquivo.stem)
    
    ratio = len(texto_final) / len(texto) * 100
    print(f"\n{'='*60}")
    print(f"🎉 CONCLUÍDO!")
    print(f"   Entrada:  {len(texto):,} chars")
    print(f"   Saída:    {len(texto_final):,} chars ({ratio:.0f}%)")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
