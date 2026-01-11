#!/usr/bin/env python3
"""
Script v2.10 - Formatação de Transcrições com Gemini 2.5 Flash
MELHORIAS: Smart Stitching Anti-Duplicação Cirúrgica

Mudanças v2.10 vs v2.9:
- remover_eco_do_contexto: Remove eco do contexto na resposta da API
- titulos_sao_similares: Fuzzy matching para detecção de títulos duplicados
- limpar_inicio_redundante: Limpeza na junção de chunks
- Injeção dinâmica de ultimo_titulo no prompt para prevenir repetição

Uso: python format_transcription_gemini.py <entrada.txt> [saida]
"""


import os
import sys
import time
import random
import json
import re
import threading
from pathlib import Path
from time import sleep
from difflib import SequenceMatcher
import hashlib

try:
    from audit_module import auditar_consistencia_legal
    AUDIT_AVAILABLE = True
except ImportError:
    AUDIT_AVAILABLE = False
    logger = logging.getLogger(__name__) if 'logging' in locals() else None
    if logger: logger.warning("⚠️  Módulo de auditoria não encontrado.")
    else: print("⚠️  Módulo de auditoria não encontrado.")

try:
    from google import genai
    from google.genai import types
except ImportError:
    print("❌ Erro: Biblioteca google-genai não instalada.")
    print("   Instale com: pip install google-genai")
    sys.exit(1)

try:
    from tqdm import tqdm
except ImportError:
    print("⚠️ Aviso: tqdm não instalado. Progress bar desabilitada.")
    tqdm = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ Aviso: python-docx não disponível. Saída em Word desabilitada.")
    DOCX_AVAILABLE = False

import logging

# =============================================================================
# CONFIGURAÇÕES v2.7
# =============================================================================

CHARS_POR_PARTE = 20000
CONTEXTO_ESTILO = 3000
OUTPUT_TOKEN_LIMIT = 32000
CACHE_TTL = '7200s'
MIN_CHARS_PARA_CACHE = 20000
MAX_RETRIES = 3
MAX_RPM = 10 
# v2.7: FORÇAR delimitadores visíveis para evitar confusão
USE_FANCY_DELIMITERS = True

# Preços API Gemini 2.5 Flash (Dezembro 2025)
PRECO_INPUT_SEM_CACHE = 0.30
PRECO_INPUT_COM_CACHE = 0.03
PRECO_OUTPUT = 2.50

# =============================================================================
# LOGGING
# =============================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S',
    handlers=[
        logging.FileHandler('formatacao.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# =============================================================================
# PROMPTS v2.7 - INSTRUÇÕES ANTI-DUPLICAÇÃO REFORÇADAS
# =============================================================================

PROMPT_FIDELIDADE = """# DIRETRIZES DE FORMATAÇÃO E REVISÃO

## PAPEL
VOCÊ É UM EXCELENTISSIMO REDATOR JURÍDICO E DIDÁTICO

 **Tom:** didático, como o professor explicando em aula.  
- **Pessoa:** manter a pessoa original da transcrição (1ª pessoa se for assim na fala).  
- **Estilo:** texto corrido, com parágrafos curtos, sem “inventar” doutrina nova.  
- **Objetivo:** reproduzir a aula em forma escrita, clara e organizada, mas ainda com a “voz” do professor.


## OBJETIVO
Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, MANTENDO A FIDELIDADE TOTAL ao conteúdo original.



## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias e das falas do professor.
4. **NÃO CRIE MUITOS BULLET POINTS** ou frases curtas demasiadamente. PREFIRA UM FORMATO DE MANUAL DIDÁTICO, não checklist.
5. **NÃO USE NEGRITOS EM EXCESSO**. Use apenas para conceitos-chave realmente importantes.
6. **NÃO RESUMA e NÃO OMITA**. Você pode reescrever frases em português padrão para melhorar a fluidez, preservar a ordem, os detalhes técnicos e os exemplos, mas **REMOVA** pausas excessivas e hesitações.


## ❌ PRESERVE OBRIGATORIAMENTE
- **NÚMEROS EXATOS**: Artigos, Leis, Súmulas, Julgados (REDI/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência, autores, casos citados
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque)
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico"

## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Corrija erros gramaticais, regências, ortográficos e de pontuação, tornando o texto gramaticalmente correto e claro.
2. **Limpeza Profunda:**
   - **REMOVA** marcadores de oralidade: "né", "tá?", "entende?", "veja bem", "tipo assim".
   - **REMOVA** interações diretas com a turma/alunos e logística: "Isso mesmo", "A colega perguntou", "Já estão me vendo?", "Estão ouvindo?", "Como ele disse ali atrás".
   - **REMOVA** redundâncias: "subir para cima", "criação nova".
   - **TRANSFORME** perguntas retóricas em afirmações quando possível (ex: "E o que isso significa?" -> "Isso significa que...").
3. **Coesão**: Utilize conectivos necessários para tornar o texto mais fluido. Aplique a pontuação devida para deixar o texto coeso e coerente.
4. **Legibilidade**:
   - **USE TEXTO CORRIDO NA MEDIDA DO POSSÍVEL**
   - Utilize formatação e estrutura com parágrafos bem definidos, facilitando a leitura e compreensão
   - Evite parágrafos longos (máximo 3-4 linhas visuais)
   - Evite blocos de texto maciços, quebre os blocos de texto em parágrafos menores
   - Seja didático sem perder detalhes e conteúdo
5. **Linguagem**: Ajuste a linguagem coloquial para um português padrão, mantendo o significado original.
6. **Citações**: Use itálico para citações curtas e recuo em itálico para citações longas.
7. -Use **negrito** para destacar conceitos-chave (sem exagero).
8. **Formatação Didática** (use com moderação, sem excesso):
   - **Bullet points** para enumerar elementos, requisitos ou características
   - **Listas numeradas** (1., 2., 3.) para enumerar itens, etapas, correntes ou exemplos
   - **Marcadores relacionais** como "→" para indicar relações, transições, ou consequências lógicas
   - Exemplo: "Processo entre A e B → prova usada contra C"
9. **Questões e Exercícios**:
   - Se o professor ditar uma questão, exercício ou caso hipotético para julgar, **ILHE-O** em um bloco de citação:
   > **Questão:** O prazo para agravo de petição é de...
   - Separe claramente o enunciado da questão da explicação/gabarito subsequente.

## 📝 ESTRUTURA
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis a partir do contexto.

## 🚫 TÍTULOS E SUBTÓPICOS (IMPORTANTE)
- **NÃO críe subtópicos para frases soltas.**
- Use títulos (##, ###) **APENAS** para mudanças reais de assunto.
- Se uma frase parece um título mas não inicia uma nova seção, mantenha como texto normal e use **negrito** se necessário.

## 📊 TABELA DE SÍNTESE (FLEXÍVEL)
Ao final de cada **bloco temático relevante** (ou capítulo), produza uma tabela de síntese (modelo flexível).
Exemplo de estrutura (adapte conforme o conteúdo):

```
### 📋 Tabela de síntese do tópico

| Conceito/Instituto | Definição (conforme a aula) | Fundamento Legal (se citado) | Observações (alertas/exceções/juris) |
| :--- | :--- | :--- | :--- |
| ...  | ...  | Art. X, Lei Y / "—" | ... |
```

***REGRAS CRÍTICAS PARA TABELAS:**
1. **Limite de conteúdo por célula:** máximo ~50 palavras. Se precisar de mais, divida em múltiplas linhas da tabela
2. **PROIBIDO usar blocos de código (```) dentro de células** - use texto simples
3. **NUNCA deixe o título "📋 Resumo do Tópico" sozinho** - se não houver dados para tabela, NÃO escreva o título
4. **POSICIONAMENTO ESTRITO:**
   - A tabela deve vir **APENAS AO FINAL** de um bloco concluído.
   - **PROIBIDO** inserir tabela no meio de uma frase ou interrompendo uma explicação.
   - Se o texto continuar sobre o mesmo assunto, **termine o texto primeiro** e coloque a tabela depois.


## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA
- **NUNCA formate novamente esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- **NUNCA repita informações que já estão no contexto**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>
- Se o texto_para_formatar começar com algo similar ao fim do contexto, NÃO duplique, apenas continue naturalmente
"""

PROMPT_APOSTILA = """# DIRETRIZES DE REDAÇÃO: MANUAL JURÍDICO DIDÁTICO (MODO APOSTILA)
## PAPEL
VOCÊ É UM EXCELENTISSIMO REDATOR JURÍDICO E DIDÁTICO
- **Tom:** doutrinário, impessoal, estilo manual de Direito.  
- **Pessoa:** 3ª pessoa ou construções impessoais (“observa-se”, “entende-se”).  
- **Estilo:** prosa mais densa, porém com parágrafos curtos e didáticos.  
- **Objetivo:** transformar o conteúdo da aula em texto de apostila/livro, sem alterar o conteúdo e sem inventar informações.


## OBJETIVO
Transformar a transcrição em um texto claro, legível e coeso, em Português Padrão, em formato de apostila/manual didático

## 🚫 O QUE NÃO FAZER
1. **NÃO RESUMA**. O tamanho do texto de saída deve ser próximo ao de entrada.
2. **NÃO OMITA** informações, exemplos, casos concretos ou explicações.
3. **NÃO ALTERE** o significado ou a sequência das ideias 


❌ PRESERVE obrigatoriamente:
- **NÚMEROS EXATOS**: Artigos, Leis, Artigos Súmulas, Julgados (REDI/Informativos). **NUNCA OMITA NÚMEROS DE LEIS OU SÚMULAS**.
- **TODO o conteúdo técnico**: exemplos, explicações, analogias, raciocínios
- **Referências**: leis, artigos, jurisprudência, autores, casos citados
- **Ênfases intencionais**: "isso é MUITO importante" (mantenha o destaque)
- **Observações pedagógicas**: "cuidado com isso!", "ponto polêmico"

## ✅ DIRETRIZES DE ESTILO
1. **Correção Gramatical**: Ajuste a linguagem coloquial para o padrão culto.
2. **Limpeza**: Remova gírias, cacoetes ("né", "tipo assim", "então") e vícios de oralidade.
3. **Coesão**: Use conectivos e pontuação adequada para tornar o texto fluido.
4. **Legibilidade**:
   - Use parágrafos bem definidos e curtos (máximo 3-4 linhas visuais).
    - Evite blocos de texto maciços, quebre os blocos de texto em parágrafos menores
   - Use **negrito** para destacar conceitos-chave (sem exagero).
5. **Formatação Didática** (use com moderação, sem excesso):
   - **Bullet points** para enumerar elementos, requisitos ou características
   - **Listas numeradas** (1., 2., 3.) para enumerar itens, etapas, correntes doutrinárias ou exemplos
   - **Marcadores relacionais** como "→" para indicar relações, transições, ou consequências lógicas
   - Exemplo: "Processo entre Pedro e José → prova usada contra Ana"

## 📝 ESTRUTURA
- Mantenha a sequência exata das falas.
- Use Títulos Markdown (##, ###) para organizar os tópicos, se identificáveis a partir do contexto.

## 📊 TABELA DE SÍNTESE (FLEXÍVEL)
Ao final de cada **bloco temático relevante** (ou capítulo), produza uma tabela de síntese (modelo flexível).
Exemplo de estrutura (adapte conforme o conteúdo):

```
### 📋 Tabela de síntese do tópico

| Conceito/Instituto | Definição (conforme a aula) | Fundamento Legal (se citado) | Observações (alertas/exceções/juris) |
| :--- | :--- | :--- | :--- |
| ...  | ...  | Art. X, Lei Y / "—" | ... |
```

***REGRAS CRÍTICAS PARA TABELAS:**
1. **Limite de conteúdo por célula:** máximo ~50 palavras. Se precisar de mais, divida em múltiplas linhas da tabela
2. **PROIBIDO usar blocos de código (```) dentro de células** - use texto simples
3. **NUNCA deixe o título "📋 Resumo do Tópico" sozinho** - se não houver dados para tabela, NÃO escreva o título
4. **POSICIONAMENTO:** A tabela deve vir **APENAS AO FINAL** da explicação completa dos tópicos ou blocos temáticos relevantes 
   - **NUNCA** insira a tabela no meio de uma explicação.
   - **NUNCA** resuma um tópico que você ainda não acabou de explicar no texto.
   - A tabela deve ser o **fechamento** lógico da seção, antes de iniciar um novo título ou tópico (## ou ###).

## ⚠️ REGRA ANTI-DUPLICAÇÃO (CRÍTICA)
Se você receber um CONTEXTO de referência (entre delimitadores ━━━):
- Este contexto é APENAS para você manter o mesmo ESTILO DE ESCRITA
- **NUNCA formate novamente esse contexto**
- **NUNCA inclua esse contexto na sua resposta**
- **NUNCA repita informações que já estão no contexto**
- Formate APENAS o texto que está entre as tags <texto_para_formatar>
- Se o texto_para_formatar começar com algo similar ao fim do contexto, NÃO duplique, apenas continue naturalmente
"""


# =============================================================================
# ESCOLHA O MODO
# =============================================================================

PROMPT_FORMATACAO = PROMPT_FIDELIDADE
# PROMPT_FORMATACAO = PROMPT_APOSTILA

# =============================================================================
# DETECÇÃO DE MODO
# =============================================================================

FIDELIDADE_MODE = "NÃO RESUMA" in PROMPT_FORMATACAO
APOSTILA_MODE = "MANUAL JURÍDICO" in PROMPT_FORMATACAO

if APOSTILA_MODE:
    THRESHOLD_MINIMO = 0.75
    THRESHOLD_CRITICO = 0.65
    MODO_NOME = "APOSTILA"
elif FIDELIDADE_MODE:
    THRESHOLD_MINIMO = 0.75
    THRESHOLD_CRITICO = 0.70
    MODO_NOME = "FIDELIDADE"
else:
    THRESHOLD_MINIMO = 0.70
    THRESHOLD_CRITICO = 0.60
    MODO_NOME = "PADRÃO"

logger.info(f"🎯 Modo: {MODO_NOME} (threshold={THRESHOLD_MINIMO:.0%})")
logger.info(f"🛡️  Anti-duplicação: ATIVADA (v2.7)")

# =============================================================================
# RATE LIMITER
# =============================================================================

class RateLimiter:
    def __init__(self, max_requests_per_minute=MAX_RPM):
        self.max_rpm = max_requests_per_minute
        self.requests = []
        self.lock = threading.Lock()
    
    def wait_if_needed(self):
        with self.lock:
            now = time.time()
            self.requests = [t for t in self.requests if now - t < 60]
            
            if len(self.requests) >= self.max_rpm:
                oldest = min(self.requests)
                wait_time = 60 - (now - oldest) + 0.5
                logger.info(f"⏱️  Rate limit: aguardando {wait_time:.1f}s...")
                sleep(wait_time)
                self.requests = [t for t in self.requests if time.time() - t < 60]
            
            self.requests.append(time.time())

rate_limiter = RateLimiter()

# =============================================================================
# CHECKPOINT/RESUME
# =============================================================================

def get_checkpoint_path(input_file):
    return Path(input_file).with_suffix('.checkpoint.json')

def save_checkpoint(input_file, resultados, chunks_info, secao_atual):
    checkpoint_path = get_checkpoint_path(input_file)
    checkpoint_data = {
        'input_file': str(input_file),
        'secao_atual': secao_atual,
        'total_secoes': len(chunks_info),
        'chunks_info': chunks_info,
        'resultados': resultados,
        'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
        'version': '2.7',
        'modo': MODO_NOME
    }
    with open(checkpoint_path, 'w', encoding='utf-8') as f:
        json.dump(checkpoint_data, f, ensure_ascii=False, indent=2)

def load_checkpoint(input_file):
    checkpoint_path = get_checkpoint_path(input_file)
    if checkpoint_path.exists():
        try:
            with open(checkpoint_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if data.get('version') not in ('2.6', '2.7'):
                    logger.warning("Checkpoint de versão anterior. Reiniciando...")
                    return None
                return data
        except Exception as e:
            logger.error(f"Erro ao carregar checkpoint: {e}")
    return None

def delete_checkpoint(input_file):
    checkpoint_path = get_checkpoint_path(input_file)
    if checkpoint_path.exists():
        checkpoint_path.unlink()
        logger.info("🧹 Checkpoint removido")

# =============================================================================
# DIVISÃO SEQUENCIAL
# =============================================================================

def dividir_sequencial(transcricao_completa):
    """Divide documento em chunks SEQUENCIAIS sem sobreposição"""
    chunks = []
    tamanho_total = len(transcricao_completa)
    inicio = 0
    
    while inicio < tamanho_total:
        fim = min(inicio + CHARS_POR_PARTE, tamanho_total)
        
        if fim < tamanho_total:
            janela = transcricao_completa[max(0, fim - 500):min(tamanho_total, fim + 500)]
            titulo_match = re.search(r'\n(#{2,4}\s+.+)\n', janela)
            
            if titulo_match:
                pos_titulo = janela.find(titulo_match.group(0))
                fim = max(0, fim - 500) + pos_titulo
            else:
                quebra = transcricao_completa.rfind('\n\n', fim - 300, fim + 300)
                if quebra != -1 and quebra > inicio:
                    fim = quebra
                else:
                    quebra = transcricao_completa.rfind('. ', fim - 150, fim + 150)
                    if quebra != -1 and quebra > inicio:
                        fim = quebra + 1
        
        chunks.append({'inicio': inicio, 'fim': fim})
        inicio = fim
    
    return chunks

def validar_chunks(chunks, transcricao_completa):
    """v2.7: Validação rigorosa de chunks sequenciais"""
    logger.info("🔍 Validando chunks sequenciais...")
    
    for i in range(len(chunks)):
        chunk = chunks[i]
        
        # Verifica se início == fim do anterior
        if i > 0:
            anterior = chunks[i-1]
            if chunk['inicio'] != anterior['fim']:
                logger.error(f"❌ Gap/Overlap no chunk {i+1}!")
                logger.error(f"   Anterior termina em: {anterior['fim']}")
                logger.error(f"   Atual começa em: {chunk['inicio']}")
                logger.error(f"   Diferença: {chunk['inicio'] - anterior['fim']} chars")
                
                # Mostra preview
                if chunk['inicio'] < anterior['fim']:
                    overlap_text = transcricao_completa[chunk['inicio']:anterior['fim']]
                    logger.error(f"   OVERLAP: '{overlap_text[:100]}...'")
                
                return False
    
    logger.info(f"✅ {len(chunks)} chunks validados (sequenciais, sem overlap)")
    return True

# =============================================================================
# FUNÇÕES AUXILIARES
# =============================================================================

def limpar_tags_xml(texto):
    texto = re.sub(r'</?[a-z_][\w\-]*>', '', texto, flags=re.IGNORECASE)
    texto = re.sub(r'<[a-z_][\w\-]*\s+[^>]+>', '', texto, flags=re.IGNORECASE)
    return texto

def carregar_transcricao(arquivo):
    try:
        with open(arquivo, 'r', encoding='utf-8-sig') as f:
            conteudo = f.read()
        
        if not conteudo.strip():
            logger.error("Arquivo está vazio.")
            sys.exit(1)
            
        return conteudo
    except FileNotFoundError:
        logger.error(f"Arquivo '{arquivo}' não encontrado.")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Erro ao ler arquivo: {e}")
        sys.exit(1)

def estimar_custo(transcricao, usar_cache, num_chunks=1):
    tokens_in = len(transcricao) // 4
    
    if APOSTILA_MODE:
        tokens_out_estimado = int(tokens_in * 0.65)
    elif FIDELIDADE_MODE:
        tokens_out_estimado = int(tokens_in * 1.00)
    else:
        tokens_out_estimado = int(tokens_in * 0.85)
    
    tokens_prompt = len(PROMPT_FORMATACAO) // 4
    tokens_in_total = tokens_in + (tokens_prompt * num_chunks)
    
    if usar_cache:
        custo_input = (tokens_in * PRECO_INPUT_COM_CACHE + tokens_prompt * num_chunks * PRECO_INPUT_SEM_CACHE) / 1_000_000
        custo_output = (tokens_out_estimado * PRECO_OUTPUT) / 1_000_000
        custo = custo_input + custo_output
    else:
        custo = (tokens_in_total * PRECO_INPUT_SEM_CACHE + tokens_out_estimado * PRECO_OUTPUT) / 1_000_000
    
    logger.info(f"💰 Custo estimado: ${custo:.4f} USD (modo {MODO_NOME})")

# v2.9: Cache REABILITADO com hash inteligente
def criar_cache_contexto(client, transcricao_completa):
    """Cria cache de contexto com hash estável para reutilização"""
    
    # Cache só vale a pena para documentos grandes
    if len(transcricao_completa) < MIN_CHARS_PARA_CACHE:
        logger.info(f"📦 Documento pequeno ({len(transcricao_completa):,} chars), cache não necessário")
        return None
    
    try:
        # Hash do prompt base para cache estável entre execuções
        prompt_hash = hashlib.sha256(PROMPT_FORMATACAO.encode()).hexdigest()[:16]
        cache_name = f"fmt_{prompt_hash}"
        
        # v2.9: Tenta encontrar cache existente válido
        try:
            for c in client.caches.list(page_size=100):
                if c.display_name == cache_name:
                    logger.info(f"♻️  Reusando cache existente: {cache_name} ({c.name})")
                    return c
        except Exception as e:
            logger.warning(f"Cache lookup warning: {e}")

        # Prepara conteúdo para cache: prompt do sistema
        cache_content = f"""{PROMPT_FORMATACAO}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📚 CONTEXTO GLOBAL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Modo: {MODO_NOME}
"""
        
        # Cria cache usando a API do google-genai
        cache = client.caches.create(
            model='gemini-2.5-flash',
            config=types.CreateCachedContentConfig(
                contents=[cache_content],
                ttl=CACHE_TTL,
                display_name=cache_name
            )
        )
        
        logger.info(f"✅ Cache criado: {cache_name} (hash: {prompt_hash}, TTL: {CACHE_TTL})")
        return cache
        
    except Exception as e:
        logger.warning(f"⚠️ Falha ao criar cache: {e}. Continuando sem cache.")
        return None

# =============================================================================
# MAPEAMENTO ESTRUTURAL (v2.8)
# =============================================================================

PROMPT_MAPEAMENTO = """Você é um especialista em organização de conteúdo jurídico.

Analise a transcrição abaixo e extraia a ESTRUTURA DE TÓPICOS do documento.

## INSTRUÇÕES:
1. Identifique os TÓPICOS PRINCIPAIS da aula (temas macro abordados)
2. Se houver múltiplas disciplinas/matérias, cada uma é um tópico de nível 1
3. Se for uma aula sobre um único tema, organize por subtópicos lógicos
4. **MÁXIMO 3 NÍVEIS**: Use apenas 1., 1.1., 1.1.1. (nunca 1.1.1.1.)
5. Seja conciso - apenas títulos, não explicações
6. Mantenha a ORDEM em que aparecem na transcrição
7. Mapeie do início ao fim, sem omitir partes

## FORMATO:
```
1. [Tópico Principal]
   1.1. [Subtópico]
      1.1.1. [Detalhamento]
```

## TRANSCRIÇÃO:
{transcricao}

Retorne APENAS a estrutura hierárquica (máx 3 níveis), sem texto adicional.
"""

def mapear_estrutura(client, transcricao_completa):
    """Analisa o documento completo e extrai a estrutura de tópicos"""
    logger.info("🗺️  Mapeando estrutura do documento...")
    
    rate_limiter.wait_if_needed()
    
    # Gemini 2.5 Flash suporta 1M tokens (~4M chars)
    # Limite de 3.5M chars para deixar margem para output (20k tokens)
    max_chars_mapeamento = 3_500_000
    
    if len(transcricao_completa) > max_chars_mapeamento:
        logger.warning(f"⚠️  Documento EXTREMAMENTE grande ({len(transcricao_completa):,} chars). Cortando final para caber no contexto.")
        # É melhor cortar o final do que picotar o meio para estrutura
        texto_para_mapear = transcricao_completa[:max_chars_mapeamento]
    else:
        texto_para_mapear = transcricao_completa
        logger.info(f"   Mapeando documento completo ({len(transcricao_completa):,} chars)")
    
    prompt = PROMPT_MAPEAMENTO.format(transcricao=texto_para_mapear)
    
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                max_output_tokens=20000,  # Aumentado para documentos grandes
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                ]
            )
        )
        
        
        if not response.text:
            logger.warning("⚠️  Resposta vazia do mapeamento.")
            logger.warning(f"   Response object: {response}")
            if hasattr(response, 'candidates') and response.candidates:
                logger.warning(f"   Finish reason: {response.candidates[0].finish_reason}")
                if hasattr(response.candidates[0], 'safety_ratings'):
                    logger.warning(f"   Safety ratings: {response.candidates[0].safety_ratings}")
            logger.warning("   Continuando sem estrutura prévia.")
            return None
        
        estrutura = response.text.strip()
        
        # Remove markdown code blocks se presentes
        if estrutura.startswith('```'):
            linhas = estrutura.split('\n')
            estrutura = '\n'.join(linhas[1:-1]) if len(linhas) > 2 else estrutura
        
        linhas = [l for l in estrutura.split('\n') if l.strip()]
        
        # Validação da estrutura
        if len(linhas) < 3:
            logger.warning(f"⚠️  Estrutura muito curta ({len(linhas)} linhas). Pode estar incompleta.")
            return None
        
        tem_numeracao = any(re.match(r'^\d+\.', l.strip()) for l in linhas)
        if not tem_numeracao:
            logger.warning("⚠️  Estrutura sem numeração hierárquica. Pode estar mal formatada.")
            return None
        
        logger.info(f"✅ Estrutura mapeada: {len(linhas)} tópicos identificados")
        
        # Log preview (primeiras 10 linhas + total)
        for linha in linhas[:10]:
            logger.info(f"   {linha}")
        if len(linhas) > 10:
            logger.info(f"   ... e mais {len(linhas) - 10} tópicos")
        
        return estrutura
        
    except Exception as e:
        logger.warning(f"⚠️  Falha no mapeamento: {e}. Continuando sem estrutura prévia.")
        return None

def simplificar_estrutura_se_necessario(estrutura, max_linhas=60):
    """
    Se a estrutura for muito longa (> max_linhas), mantém apenas:
    - Nível 1: 1. Assunto
    - Nível 2: 1.1. Subassunto

    Sempre inclui todos os níveis 1, e todos os níveis 2,
    e depois corta no máximo max_linhas, preservando a ordem original.
    """
    if not estrutura:
        return estrutura

    linhas = [l for l in estrutura.strip().split("\n") if l.strip()]
    if len(linhas) <= max_linhas:
        # Já está razoável, mantém até nível 3 (será filtrado depois por filtrar_niveis_excessivos)
        return estrutura

    logger.info(f"📉 Estrutura muito longa ({len(linhas)} itens). Simplificando para níveis 1 e 2, máx {max_linhas} linhas...")

    nivel1 = []
    nivel2 = []

    for l in linhas:
        s = l.strip()
        # 1. Processo do Trabalho
        if re.match(r"^\d+\.\s", s):
            nivel1.append(l)
        # 1.1. Recursos Trabalhistas
        elif re.match(r"^\d+\.\d+\.\s", s):
            nivel2.append(l)

    # Se por algum motivo não identificou quase nada, devolve original para não quebrar
    if len(nivel1) + len(nivel2) < 5:
        logger.warning("⚠️ Simplificação deixou poucos tópicos. Mantendo estrutura original.")
        return estrutura

    # Monta nova estrutura: primeiro todos os níveis 1, depois os níveis 2, respeitando ordem de aparecimento
    nova = []
    vistos = set()

    for l in linhas:
        if l in vistos:
            continue
        if l in nivel1 or l in nivel2:
            nova.append(l)
            vistos.add(l)

    if len(nova) > max_linhas:
        nova = nova[:max_linhas]

    logger.info(f"✅ Estrutura simplificada: {len(linhas)} -> {len(nova)} linhas (níveis 1 e 2).")
    return "\n".join(nova)

def filtrar_niveis_execessivos(estrutura, max_nivel=3):
    """
    Remove itens da estrutura que sejam mais profundos que max_nivel.
    Ex: se max_nivel=3, remove 1.1.1.1
    """
    if not estrutura:
        return estrutura
        
    linhas = estrutura.strip().split('\n')
    linhas_filtradas = []
    itens_removidos = 0
    
    # Regex para validar nível. 
    # Nível 1: \d+\.
    # Nível 2: \d+\.\d+
    # Nível 3: \d+\.\d+\.\d+
    # O regex verifica se tem no máximo (max_nivel-1) pontos internos entre números
    
    for linha in linhas:
        # Conta quantos grupos de números existem
        match = re.match(r'^(\d+(?:\.\d+)*)', linha.strip())
        if match:
            numeracao = match.group(1)
            nivel = numeracao.count('.') + 1
            if linha.strip().endswith('.'): # Se terminar com ponto (1.1.), não conta como nível extra
                 nivel = numeracao.count('.')
            
            # Ajuste robusto: contar números separados por ponto
            partes = [p for p in numeracao.split('.') if p.isdigit()]
            nivel_real = len(partes)
            
            if nivel_real <= max_nivel:
                linhas_filtradas.append(linha)
            else:
                itens_removidos += 1
        else:
            # Linhas sem numeração (títulos soltos?) mantém por segurança ou remove?
            # Vamos manter para não quebrar formatação estranha
            linhas_filtradas.append(linha)
            
    if itens_removidos > 0:
        logger.info(f"✂️  Filtrados {itens_removidos} itens com nível > {max_nivel}")
    
    return '\n'.join(linhas_filtradas)

# =============================================================================
# PROCESSAMENTO
# =============================================================================

def processar_simples(client, transcricao_bruta):
    logger.info("📄 Documento pequeno - processando em requisição única...")
    
    prompt = f"""{PROMPT_FORMATACAO}

<texto_para_formatar>
{transcricao_bruta}
</texto_para_formatar>

Retorne APENAS o Markdown formatado."""
    
    for tentativa in range(MAX_RETRIES):
        try:
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=OUTPUT_TOKEN_LIMIT,
                    temperature=0,
                    safety_settings=[
                        types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                        types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                    ]
                )
            )
            resultado = response.text
            return limpar_tags_xml(resultado)
        except Exception as e:
            if tentativa < MAX_RETRIES - 1:
                wait = (2 ** tentativa) + random.uniform(0, 1)
                logger.warning(f"Erro, retry em {wait:.1f}s...")
                sleep(wait)
            else:
                raise

def processar_chunk(client, cache, texto_chunk, numero, total, contexto_estilo="", estrutura_global=None, ultimo_titulo=None, profundidade=0):
    rate_limiter.wait_if_needed()
    
    # Recursão infinita protection e limite mínimo
    MIN_CHUNK_CHARS = 4000
    if len(texto_chunk) < MIN_CHUNK_CHARS:
        logger.warning(f"⚠️ Chunk {numero} muito pequeno ({len(texto_chunk)} chars). Processando sem dividir.")
    elif profundidade > 2:
        logger.warning(f"⚠️ Chunk {numero}: Profundidade de recursão {profundidade} atingida. Processando sem dividir.")

    # v2.8: Seção de estrutura global
    secao_estrutura = ""
    if estrutura_global:
        secao_estrutura = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 ESTRUTURA GLOBAL DA AULA (SIGA ESTA HIERARQUIA)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{estrutura_global}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ USE esta estrutura para nomear seus títulos (##, ###)
   Os títulos devem corresponder aos tópicos listados acima.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
    
    # v2.10: Aviso sobre último título (anti-duplicação)
    aviso_titulo = ""
    if ultimo_titulo:
        aviso_titulo = f"""
🚫 O bloco anterior TERMINOU no tópico: "{ultimo_titulo}"
   NÃO inicie sua resposta repetindo este título.
   Continue o conteúdo ou inicie o PRÓXIMO subtópico.
"""
    
    # v2.7: Delimitadores MUITO visíveis e instruções reforçadas
    secao_contexto = ""
    if contexto_estilo:
        secao_contexto = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔒 CONTEXTO ANTERIOR (SOMENTE REFERÊNCIA DE ESTILO)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{contexto_estilo}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ ATENÇÃO: O bloco acima JÁ FOI FORMATADO anteriormente.
- NÃO formate novamente esse conteúdo
- NÃO inclua esse conteúdo na sua resposta
- Use APENAS como referência de estilo de escrita
{aviso_titulo}━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 NOVO TEXTO PARA FORMATAR (comece aqui):
"""
    
    # v2.9: Cache Support - Se usar cache, não repete PROMPT_FORMATACAO
    instructions_body = f"""
{secao_estrutura}
{secao_contexto}

<texto_para_formatar>
{texto_chunk}
</texto_para_formatar>

**INSTRUÇÕES FINAIS**:
- Esta é a parte {numero} de {total} (Profundidade: {profundidade})
- Formate APENAS o texto entre <texto_para_formatar>
- Se houver ESTRUTURA GLOBAL acima, use os mesmos nomes de tópicos
- Se houver contexto acima, NÃO o reprocesse
- **ANTI-REPETIÇÃO DE TÍTULOS**: Se o contexto anterior termina com um título (ex: "## Homologação"), NÃO repita esse título no início da sua resposta. Continue o conteúdo diretamente ou inicie o PRÓXIMO tópico diferente.
- Retorne APENAS o Markdown formatado do NOVO texto
"""

    if cache:
        prompt = instructions_body
    else:
        prompt = f"{PROMPT_FORMATACAO}\n{instructions_body}"

    for tentativa in range(MAX_RETRIES):
        try:
            safety_config = [
                types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
            ]
            
            # Configuração dinâmica para suportar cache
            gen_config_args = {
                "max_output_tokens": OUTPUT_TOKEN_LIMIT,
                "temperature": 0.1,
                "safety_settings": safety_config
            }
            if cache:
                gen_config_args['cached_content'] = cache.name

            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(**gen_config_args)
            )
            
            # --- Diagnóstico finishReason e Usage ---
            finish_reason = "UNKNOWN"
            usage_tokens = 0
            
            try:
                if hasattr(response, 'usage_metadata'):
                    usage = response.usage_metadata
                    # Tenta acessar atributos (pode variar entre SDKs/Vertex)
                    # Vertex API retorna candidates_token_count
                    cand_tokens = getattr(usage, 'candidates_token_count', 0)
                    prompt_tokens = getattr(usage, 'prompt_token_count', 0)
                    usage_tokens = cand_tokens
                    logger.info(f"📊 Usage: Prompt={prompt_tokens} | Candidates={cand_tokens}")
                
                if hasattr(response, 'candidates') and response.candidates:
                    cand = response.candidates[0]
                    if hasattr(cand, 'finish_reason'):
                         finish_reason = str(cand.finish_reason) # ex: FinishReason.STOP ou "STOP"
            except Exception as ex_usage:
                logger.warning(f"⚠️ Erro ao ler metadados: {ex_usage}")

            # Captura texto (lidando com caso de .text vazio mas content presente)
            resultado = ""
            try:
                resultado = response.text
            except ValueError:
                # O SDK levanta ValueError se finish_reason for SAFETY ou se não houver text field padrão
                pass
            
            if not resultado and hasattr(response, 'candidates') and response.candidates:
                # Tenta extrair parts[0].text manualmente se .text falhou
                try:
                    parts = response.candidates[0].content.parts
                    if parts:
                        resultado = parts[0].text
                except:
                    pass
            
            if not resultado:
                logger.warning(f"⚠️  Resposta vazia na tentativa {tentativa+1}. Reason: {finish_reason}")
                if tentativa < MAX_RETRIES - 1:
                    sleep(2 * (tentativa + 1))
                    continue
                else:
                    return f"[ERRO SEÇÃO {numero}: RESPOSTA VAZIA]"

            # Validação básica de tamanho (compressão)
            razao = len(resultado) / len(texto_chunk) if len(texto_chunk) > 0 else 0
            
            problema_detectado = False
            msg_problema = ""
            
            compressao_excessiva_severa = False
            
            if razao < THRESHOLD_CRITICO: # < 0.70 por padrão
                problema_detectado = True
                msg_problema = f"Compressão excessiva ({razao:.0%})"
                # Flag para chunking imediato se for muito baixo (ex < 0.65)
                # O usuário pediu < THRESHOLD_CRITICO, vamos ser assertivos.
                compressao_excessiva_severa = True
                
            if problema_detectado:
                logger.warning(f"⚠️ Seção {numero}: {msg_problema}. Reason: {finish_reason}. (Tentativa {tentativa+1}/{MAX_RETRIES})")
                
                # SEÇÃO CRÍTICA: Decisão de Chunking Adaptativo
                
                # Lógica antiga: só no final ou MAX_TOKENS
                # Lógica nova: divide cedo se compressão for severa
                
                deve_dividir = (
                    len(texto_chunk) > MIN_CHUNK_CHARS 
                    and profundidade < 2
                    and (
                        ("MAX_TOKENS" in str(finish_reason)) or 
                        (tentativa == MAX_RETRIES - 1) or
                        (compressao_excessiva_severa) # NOVO: Divide já!
                    )
                )
                
                if deve_dividir:
                    motivo = "MAX_TOKENS" if "MAX_TOKENS" in str(finish_reason) else "COMPRESSÃO"
                    logger.info(f"✂️  ATIVANDO CHUNKING ADAPTATIVO para Seção {numero} (Motivo: {motivo} | Profundidade {profundidade} -> {profundidade+1})")
                    return dividir_e_reprocessar(client, cache, texto_chunk, numero, total, contexto_estilo, estrutura_global, ultimo_titulo, profundidade)
                
                if tentativa < MAX_RETRIES - 1:
                    continue  # Tenta de novo (se não foi severo o suficiente para dividir)
                else:
                    logger.error(f"Seção {numero}: Falha após {MAX_RETRIES} tentativas. Retornando melhor esforço.")
            
            return resultado
            
        except Exception as e:
            erro_msg = str(e)
            is_recoverable = any(code in erro_msg for code in ['503', '429', '500', 'RESOURCE_EXHAUSTED', 'InternalServerError']) or "Resposta vazia" in erro_msg
            
            if tentativa < MAX_RETRIES - 1 and is_recoverable:
                wait = (2 ** (tentativa + 2)) + random.uniform(1, 3)
                if '429' in erro_msg or 'RATE_LIMIT' in erro_msg:
                    wait = 30 + random.uniform(0, 5)
                    logger.warning(f"🛑 Rate Limit (429) detectado na seção {numero}. Pausa longa de {wait:.1f}s...")
                else:
                    logger.warning(f"Erro seção {numero}: {erro_msg}. Retry {tentativa+2}/{MAX_RETRIES} em {wait:.1f}s")
                sleep(wait)
            else:
                logger.error(f"Falha seção {numero}: {erro_msg}")
                return f"\n\n> [!WARNING]\n> Falha ao processar seção {numero}. Texto original:\n\n{texto_chunk}"
    
    return texto_chunk

def dividir_e_reprocessar(client, cache, texto_chunk, numero, total, contexto_estilo, estrutura_global, ultimo_titulo, profundidade):
    """
    Divide um chunk grande em dois menores e processa recursivamente.
    Tenta dividir em quebras de parágrafo (\n\n) próximas ao meio.
    """
    meio = len(texto_chunk) // 2
    
    # Procura quebra ideal (\n\n) num raio de 20% do meio
    margem = int(len(texto_chunk) * 0.2)
    inicio_busca = max(0, meio - margem)
    fim_busca = min(len(texto_chunk), meio + margem)
    
    janela_busca = texto_chunk[inicio_busca:fim_busca]
    pos_relativa = janela_busca.find('\n\n')
    
    if pos_relativa != -1:
        ponto_corte = inicio_busca + pos_relativa + 2 # +2 para incluir os \n\n no primeiro bloco ou pular? Vamos cortar DEPOIS dos \n\n
    else:
        # Tenta quebra simples \n
        pos_relativa_n = janela_busca.find('\n')
        if pos_relativa_n != -1:
            ponto_corte = inicio_busca + pos_relativa_n + 1
        else:
            # Corte seco no espaço mais próximo
            pos_relativa_espaco = janela_busca.find(' ')
            if pos_relativa_espaco != -1:
                ponto_corte = inicio_busca + pos_relativa_espaco + 1
            else:
                ponto_corte = meio # Corte arbitrário
    
    parte_a = texto_chunk[:ponto_corte]
    parte_b = texto_chunk[ponto_corte:]
    
    logger.info(f"   Splitting chunk {numero}: Part A ({len(parte_a)} chars) + Part B ({len(parte_b)} chars)")
    
    # Processa Parte A
    resultado_a = processar_chunk(
        client, cache, parte_a, f"{numero}.A", total, 
        contexto_estilo, estrutura_global, ultimo_titulo, profundidade + 1
    )
    
    # Usa o final de A como contexto para B? Talvez seja excessivo e caro.
    # Vamos manter o contexto original para B por segurança, 
    # ou usar resultado_a[-1000:] como novo contexto_estilo.
    # Usar resultado_a é melhor para continuidade.
    
    novo_contexto = resultado_a[-2000:] if len(resultado_a) > 2000 else resultado_a
    
    # Processa Parte B
    resultado_b = processar_chunk(
        client, cache, parte_b, f"{numero}.B", total, 
        novo_contexto, estrutura_global, None, profundidade + 1 # ultimo_titulo None pois A já tratou disso
    )
    
    return f"{resultado_a}\n\n{resultado_b}"

def extrair_titulos_h2(texto):
    """Extrai todos os títulos de nível 2 (##) do texto"""
    titulos = []
    for linha in texto.split('\n'):
        if linha.strip().startswith('##') and not linha.strip().startswith('###'):
            titulo_limpo = re.sub(r'^##\s*\d+\.\s*', '', linha.strip())
            titulos.append(titulo_limpo.lower())
    return titulos

# =============================================================================
# SMART STITCHING (v2.10) - Anti-Duplicação Cirúrgica
# =============================================================================

def remover_eco_do_contexto(resposta_api, contexto_enviado):
    """
    Remove o início da resposta se for apenas um 'eco' do final do contexto.
    """
    if not contexto_enviado or not resposta_api:
        return resposta_api

    final_contexto = contexto_enviado.strip()[-300:]
    inicio_resposta = resposta_api.strip()[:300]

    matcher = SequenceMatcher(None, final_contexto, inicio_resposta)
    match = matcher.find_longest_match(0, len(final_contexto), 0, len(inicio_resposta))

    if match.size > 50:
        logger.info(f"✂️ Eco detectado! Removendo {match.size} chars repetidos no início.")
        return resposta_api.strip()[match.size:].strip()
    
    return resposta_api

def titulos_sao_similares(t1, t2, threshold=0.90):
    """Verifica se dois títulos são semanticamente iguais (fuzzy matching)."""
    def normalizar(t):
        # Remove apenas caracteres não alfanuméricos, mas MANTÉM tamanho relativo
        return re.sub(r'[^a-z0-9 ]', '', t.lower())
    
    nt1 = normalizar(t1)
    nt2 = normalizar(t2)
    
    if not nt1 or not nt2:
        return False
    
    # PROTEÇÃO 1: Se um título for muito maior que o outro, não são duplicatas
    nt1_compact = nt1.replace(' ', '')
    nt2_compact = nt2.replace(' ', '')
    len_ratio = min(len(nt1_compact), len(nt2_compact)) / max(len(nt1_compact), len(nt2_compact))
    if len_ratio < 0.8:  # Se a diferença de tamanho for > 20%, assume que são diferentes
        return False
    
    # PROTEÇÃO 2: Verificação por palavras - se houver palavras exclusivas significativas
    palavras1 = set(nt1.split())
    palavras2 = set(nt2.split())
    diferenca = palavras1.symmetric_difference(palavras2)
    
    # Se as palavras diferentes forem longas (não apenas 'e', 'do', 'da'), assume diferença real
    if any(len(w) > 3 for w in diferenca):
        return False
        
    return SequenceMatcher(None, nt1_compact, nt2_compact).ratio() > threshold

def limpar_inicio_redundante(texto_novo, texto_acumulado):
    """
    Remove título no início do novo chunk se similar ao último título do texto acumulado.
    """
    if not texto_acumulado.strip():
        return texto_novo

    ultimas_linhas = texto_acumulado.strip().split('\n')[-30:]
    ultimo_titulo = None
    for linha in reversed(ultimas_linhas):
        if linha.strip().startswith('##'):
            ultimo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
            break
    
    if not ultimo_titulo:
        return texto_novo

    linhas_novas = texto_novo.strip().split('\n')
    
    for i, linha in enumerate(linhas_novas[:10]):  # v2.11: Busca mais profunda (era 5)
        if linha.strip().startswith('##'):
            novo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
            
            if titulos_sao_similares(ultimo_titulo, novo_titulo):
                logger.info(f"✂️ Título duplicado na junção: '{novo_titulo}' ≈ '{ultimo_titulo}'")
                return '\n'.join(linhas_novas[i+1:])
    
    return texto_novo

def detectar_secoes_duplicadas(texto):
    """v2.10: Detecta seções duplicadas por títulos (Fuzzy Matching)"""
    logger.info("🔍 Detectando seções duplicadas (fuzzy)...")
    
    linhas = texto.split('\n')
    titulos_vistos = []
    secoes_duplicadas = []
    
    for i, linha in enumerate(linhas):
        if linha.strip().startswith('##') and not linha.strip().startswith('###'):
            titulo_normalizado = re.sub(r'^##\s*\d+\.?\s*', '', linha.strip())
            titulo_normalizado = re.sub(r'[📋📊🗂️]', '', titulo_normalizado).strip()
            
            duplicado = False
            for t_visto, linha_visto in titulos_vistos:
                if titulos_sao_similares(titulo_normalizado, t_visto):
                    logger.warning(f"⚠️  Duplicado (fuzzy): '{linha.strip()}' ≈ '{t_visto}'")
                    secoes_duplicadas.append({
                        'titulo': titulo_normalizado,
                        'primeira_linha': linha_visto,
                        'duplicada_linha': i
                    })
                    duplicado = True
                    break
            
            if not duplicado:
                titulos_vistos.append((titulo_normalizado, i))
    
    if secoes_duplicadas:
        logger.error(f"❌ {len(secoes_duplicadas)} seções duplicadas detectadas!")
    else:
        logger.info("✅ Nenhuma seção duplicada detectada")
    
    return secoes_duplicadas

def remover_secoes_duplicadas(texto):
    """v2.7: Remove seções duplicadas mantendo apenas a primeira ocorrência"""
    secoes_dup = detectar_secoes_duplicadas(texto)
    
    if not secoes_dup:
        return texto
    
    logger.info("🧹 Removendo seções duplicadas...")
    
    linhas = texto.split('\n')
    linhas_para_remover = set()
    
    for dup in secoes_dup:
        # Marca para remoção todas as linhas da seção duplicada
        inicio_remocao = dup['duplicada_linha']
        
        # Encontra onde a seção duplicada termina (próximo ## ou fim do arquivo)
        fim_remocao = len(linhas)
        for i in range(inicio_remocao + 1, len(linhas)):
            if linhas[i].strip().startswith('##') and not linhas[i].strip().startswith('###'):
                fim_remocao = i
                break
        
        logger.info(f"   Removendo linhas {inicio_remocao}-{fim_remocao} (seção duplicada)")
        for i in range(inicio_remocao, fim_remocao):
            linhas_para_remover.add(i)
    
    # Reconstrói texto sem as linhas duplicadas
    linhas_limpas = [linha for i, linha in enumerate(linhas) if i not in linhas_para_remover]
    
    logger.info(f"✅ {len(linhas_para_remover)} linhas removidas")
    return '\n'.join(linhas_limpas)

def remover_duplicacoes_literais(texto):
    """Remove parágrafos individuais duplicados"""
    paragrafos = texto.split('\n\n')
    paragrafos_limpos = []
    dup_count = 0
    
    for i, para in enumerate(paragrafos):
        if i == 0:
            paragrafos_limpos.append(para)
            continue
        
        if len(para.strip()) < 80 or para.strip().startswith('#'):
            paragrafos_limpos.append(para)
            continue
        
        is_duplicate = False
        para_norm = ' '.join(para.split()).lower()
        
        for j in range(max(0, len(paragrafos_limpos) - 3), len(paragrafos_limpos)):
            para_ant = paragrafos_limpos[j]
            para_ant_norm = ' '.join(para_ant.split()).lower()
            
            ratio = SequenceMatcher(None, para_norm, para_ant_norm).ratio()
            
            if ratio > 0.95:
                is_duplicate = True
                dup_count += 1
                break
        
        if not is_duplicate:
            paragrafos_limpos.append(para)
    
    if dup_count > 5:
        logger.warning(f"⚠️  {dup_count} parágrafos duplicados removidos")
    
    return '\n\n'.join(paragrafos_limpos)

def numerar_titulos(texto):
    """Adiciona numeração sequencial aos títulos"""
    linhas = texto.split('\n')
    linhas_numeradas = []
    
    contador_h2 = 0
    contador_h3 = 0
    contador_h4 = 0
    
    titulo_pattern = re.compile(r'^(#{2,4})\s+(?:\d+(?:\.\d+)*\.?\s+)?(.+)$')
    
    # Variaveis para rastrear últimos títulos e evitar repetições
    ultimo_h2_texto = ""
    ultimo_h3_texto = ""
    
    for linha in linhas:
        match = titulo_pattern.match(linha)
        
        if match:
            nivel = len(match.group(1))
            texto_titulo = match.group(2).strip()
            
            # Não numera títulos de resumo/quadros
            if any(keyword in texto_titulo.lower() for keyword in ['resumo', 'quadro', 'esquema', '📋', '📊', '🗂️']):
                linhas_numeradas.append(linha)
                continue
            
            # MERGE INTELIGENTE DE TÍTULOS REPETIDOS (v2.11)
            # Se o título atual for muito similar ao anterior do mesmo nível ("continuação" de chunk), ignoramos o novo
            # para que o texto flua como um único tópico.
            
            from difflib import SequenceMatcher
            eh_duplicado = False
            
            if nivel == 2:
                # Verifica similaridade com último H2
                ratio = SequenceMatcher(None, texto_titulo.lower(), ultimo_h2_texto.lower()).ratio()
                if ratio > 0.9:
                    eh_duplicado = True
                    logger.info(f"🔄 Título H2 mesclado: '{texto_titulo}' ≈ '{ultimo_h2_texto}'")
                else:
                    ultimo_h2_texto = texto_titulo
            elif nivel == 3:
                 # Verifica similaridade com último H3
                ratio = SequenceMatcher(None, texto_titulo.lower(), ultimo_h3_texto.lower()).ratio()
                if ratio > 0.9:
                    eh_duplicado = True
                    logger.info(f"🔄 Título H3 mesclado: '{texto_titulo}' ≈ '{ultimo_h3_texto}'")
                else:
                    ultimo_h3_texto = texto_titulo
            
            if eh_duplicado:
                continue # Pula a linha do título, fundindo o conteúdo
            
            if nivel == 2:
                contador_h2 += 1
                contador_h3 = 0
                contador_h4 = 0
                nova_linha = f"## {contador_h2}. {texto_titulo}"
            elif nivel == 3:
                contador_h3 += 1
                contador_h4 = 0
                nova_linha = f"### {contador_h2}.{contador_h3}. {texto_titulo}"
            elif nivel == 4:
                contador_h4 += 1
                nova_linha = f"#### {contador_h2}.{contador_h3}.{contador_h4}. {texto_titulo}"
            else:
                nova_linha = linha
            
            linhas_numeradas.append(nova_linha)
        else:
            linhas_numeradas.append(linha)
    
    return '\n'.join(linhas_numeradas)

# =============================================================================
# VERIFICAÇÃO DE COBERTURA E DUPLICAÇÕES
# =============================================================================

def normalizar_fingerprint(texto, tipo):
    """Normaliza texto para comparação (ex: 'Lei 11.100' -> 'lei 11100')"""
    texto = texto.lower().strip()
    
    if tipo == 'leis':
        # Mantém apenas 'lei' e números
        nums = re.findall(r'\d+', texto)
        if nums:
            # Reconstrói como 'lei 12345'
            # Filtra leis com menos de 4 dígitos para evitar ruído (ex: lei 10, lei 13)
            num_full = ''.join(nums)
            if len(num_full) >= 4:
                return f"lei {num_full}"
            return None
            
    elif tipo == 'sumulas':
        nums = re.findall(r'\d+', texto)
        if nums:
            return f"súmula {''.join(nums)}"
            
    elif tipo == 'artigos':
        nums = re.findall(r'\d+', texto)
        if nums:
            return f"artigo {''.join(nums)}"
            
    return re.sub(r'[^\w\s]', '', texto)

def extrair_fingerprints(texto):
    """Extrai 'fingerprints' únicos e normalizados do texto"""
    fingerprints = {
        'leis': set(),
        'sumulas': set(),
        'artigos': set(),
        'julgados': set()
    }
    
    # Regex melhorado para capturar variações
    lei_pattern = re.compile(r'\b(?:lei|l\.)\s*n?º?\s*([\d\.]+)', re.IGNORECASE)
    sumula_pattern = re.compile(r'\bsúmula\s*(?:vinculante)?\s*n?º?\s*(\d+)', re.IGNORECASE)
    
    # Extrai e normaliza
    for match in lei_pattern.finditer(texto):
        fp = normalizar_fingerprint(f"lei {match.group(1)}", 'leis')
        if fp: fingerprints['leis'].add(fp)
    
    for match in sumula_pattern.finditer(texto):
        fp = normalizar_fingerprint(f"súmula {match.group(1)}", 'sumulas')
        if fp: fingerprints['sumulas'].add(fp)
        
    return fingerprints

def contar_ocorrencias_robust(fingerprints, texto):
    """Conta ocorrências com suporte a formatação jurídica formal (Lei nº X)"""
    contagens = {}
    
    # CORREÇÃO 1: Não remover pontuação indiscriminadamente, apenas normalizar espaços
    # Mantemos barras e pontos para evitar fusão de números (11.101/2005)
    texto_lower = texto.lower()
    
    for categoria, items in fingerprints.items():
        for item in items:
            key = f"{categoria}:{item}"
            
            if categoria == 'leis':
                # item ex: "lei 4320" -> extrai "4320"
                # Remove pontuação do item para garantir match limpo no número
                num_bruto = item.split()[-1] 
                num = re.sub(r'[^\d]', '', num_bruto)
                
                # Permite pontos opcionais entre dígitos (ex: 4.320 match com 4320)
                num_regex = r"\.?".join(list(num))
                
                # CORREÇÃO 2: Regex flexível que aceita "n", "nº", "no", "num" no meio
                # Aceita: "Lei 4320", "Lei nº 4.320", "Lei n. 4320"
                # O \W* permite pontos/barras entre Lei e o número
                # Adicionado \b no final para evitar matches parciais (Lei 10 != Lei 100)
                pattern = f"lei(?:\\s+|\\.|\\,|nº|n\\.|n\\s|num\\.?)*{num_regex}\\b"
                
                # Usamos findall no texto original (lower) para pegar variações com pontuação
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
                
            elif categoria == 'sumulas':
                num = item.split()[-1]
                num_regex = r"\.?".join(list(num)) # Súmulas raramente tem ponto, mas por garantia
                
                # Mesma lógica para súmulas (Súmula Vinculante nº 10)
                pattern = f"súmula(?:\\s+|\\.|\\,|vinculante|nº|n\\.|n\\s)*{num_regex}\\b"
                matches = re.findall(pattern, texto_lower)
                contagens[key] = len(matches)
                
            else:
                # Fallback para outros tipos (busca literal simples)
                contagens[key] = texto_lower.count(item)
                
    return contagens

def verificar_cobertura(texto_original, texto_formatado, arquivo_saida=None):
    """Verifica omissões e duplicações artificiais entre original e formatado"""
    logger.info("🔍 Verificando cobertura e duplicações...")
    
    # Extrai fingerprints do original
    fp_original = extrair_fingerprints(texto_original)
    
    # Conta ocorrências em ambos
    contagem_original = contar_ocorrencias_robust(fp_original, texto_original)
    contagem_formatado = contar_ocorrencias_robust(fp_original, texto_formatado)
    
    omissoes = []
    duplicacoes = []
    
    for key, count_orig in contagem_original.items():
        count_fmt = contagem_formatado.get(key, 0)
        categoria, item = key.split(':', 1)
        
        # Omissão: estava no original mas sumiu
        if count_orig > 0 and count_fmt == 0:
            omissoes.append({
                'categoria': categoria,
                'item': item,
                'original': count_orig,
                'formatado': count_fmt
            })
        
        # Duplicação (agora considerada positiva em materiais didáticos)
        if count_fmt > count_orig:
            duplicacoes.append({
                'categoria': categoria,
                'item': item,
                'original': count_orig,
                'formatado': count_fmt,
                'extra': count_fmt - count_orig
            })
    
    # Gera relatório
    total_items = len([k for k, v in contagem_original.items() if v > 0])
    items_preservados = total_items - len(omissoes)
    cobertura = items_preservados / total_items * 100 if total_items > 0 else 100
    
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    logger.info(f"📊 RELATÓRIO DE VERIFICAÇÃO")
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    logger.info(f"✅ Cobertura: {cobertura:.1f}% ({items_preservados}/{total_items} referências)")
    
    if omissoes:
        logger.warning(f"\n❌ POSSÍVEIS OMISSÕES ({len(omissoes)}):")
        for o in omissoes[:10]:  # Limita a 10
            logger.warning(f"   - [{o['categoria']}] {o['item']}")
        if len(omissoes) > 10:
            logger.warning(f"   ... e mais {len(omissoes) - 10} omissões")
    else:
        logger.info("✅ Nenhuma omissão detectada")
    
    if duplicacoes:
        logger.info(f"\nℹ️ CITAÇÕES REFORÇADAS (Tabelas/Resumos) ({len(duplicacoes)}):")
        for d in duplicacoes[:10]:
            logger.info(f"   - [{d['categoria']}] {d['item']}: {d['original']}x → {d['formatado']}x (+{d['extra']})")
        if len(duplicacoes) > 10:
            logger.info(f"   ... e mais {len(duplicacoes) - 10} citações extras")
    else:
        logger.info("ℹ️ Nenhuma citação extra detectada")
    
    logger.info(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    
    # Salva relatório em arquivo se especificado
    if arquivo_saida:
        relatorio_path = arquivo_saida.replace('.md', '_verificacao.txt')
        with open(relatorio_path, 'w', encoding='utf-8') as f:
            f.write(f"RELATÓRIO DE VERIFICAÇÃO\n")
            f.write(f"{'='*50}\n\n")
            f.write(f"Cobertura: {cobertura:.1f}% ({items_preservados}/{total_items})\n\n")
            
            if omissoes:
                f.write(f"OMISSÕES ({len(omissoes)}):\n")
                for o in omissoes:
                    f.write(f"  - [{o['categoria']}] {o['item']}\n")
                f.write("\n")
            
            if duplicacoes:
                f.write(f"DUPLICAÇÕES ARTIFICIAIS ({len(duplicacoes)}):\n")
                for d in duplicacoes:
                    f.write(f"  - [{d['categoria']}] {d['item']}: {d['original']}x original → {d['formatado']}x formatado\n")
        
        logger.info(f"📄 Relatório salvo: {relatorio_path}")
    
    return {
        'cobertura': cobertura,
        'omissoes': omissoes,
        'duplicacoes': duplicacoes
    }

# =============================================================================
# V2.10: FUNÇÕES DE PÓS-PROCESSAMENTO ESTRUTURAL (Tabelas e Parágrafos)
# =============================================================================

def mover_tabelas_para_fim_de_secao(texto):
    """
    v2.11: Reorganiza tabelas movendo-as para o final do BLOCO ATUAL (H2 ou H3).
    Corrige bug de tabelas sumindo ou ficando muito longe do contexto.
    """
    logger.info("📊 Reorganizando tabelas (Smart Layout v2.11)...")
    
    linhas = texto.split('\n')
    resultado = []
    tabelas_pendentes = [] 
    
    i = 0
    while i < len(linhas):
        linha = linhas[i]
        linha_strip = linha.strip()
        
        # 1. DETECTAR SE É UM TÍTULO (H1, H2, H3...)
        # Se encontrarmos um novo título, hora de "despejar" as tabelas acumuladas do bloco anterior
        if linha_strip.startswith('#'):
            # Despeja tabelas antes de iniciar o novo tópico
            if tabelas_pendentes:
                resultado.append('') # Espaço antes
                for t_info in tabelas_pendentes:
                    if t_info['titulo']:
                        resultado.append(t_info['titulo'])
                    resultado.extend(t_info['linhas'])
                    resultado.append('') # Espaço depois
                tabelas_pendentes = []
            
            resultado.append(linha)
            i += 1
            continue

        # 2. DETECTAR INÍCIO DE TABELA
        # Critério: Linha tem pipe '|' E parece estrutura de tabela (não apenas citação)
        eh_inicio_tabela = False
        if '|' in linha_strip:
            # Verifica se é uma linha de markdown table válida (tem pipe e chars)
            # E se a próxima linha ou a seguinte tem o separador '---'
            has_separator = False
            for lookahead in range(1, 3): # Olha até 2 linhas pra frente (ignora 1 linha vazia)
                if i + lookahead < len(linhas):
                    prox = linhas[i + lookahead].strip()
                    if set(prox).issubset(set('|- :')): # Só contem caracteres de estrutura de tabela
                         has_separator = True
                         break
            
            if has_separator or (linha_strip.startswith('|') and linha_strip.endswith('|')):
                eh_inicio_tabela = True

        if eh_inicio_tabela:
            # --- Captura da Tabela ---
            tabela_linhas = []
            titulo_tabela = None
            
            # Tenta recuperar o título da tabela que ficou na linha anterior (ou resultado)
            # Verifica se a última linha adicionada ao resultado parece um título de tabela
            if resultado and len(resultado) > 0:
                last_line = resultado[-1].strip()
                # Padrões comuns de título de tabela gerados pela IA
                if (last_line.startswith('###') or last_line.startswith('**')) and \
                   any(x in last_line.lower() for x in ['tabela', 'resumo', 'quadro', 'síntese', 'esquema', '📋']):
                    titulo_tabela = resultado.pop() # Remove do fluxo principal para agrupar com a tabela

            # Captura as linhas da tabela
            j = i
            while j < len(linhas):
                curr = linhas[j].strip()
                # Continua se tiver pipe ou for linha vazia no meio da tabela (mas cuidado com fim)
                if '|' in curr:
                    tabela_linhas.append(linhas[j])
                    j += 1
                elif not curr:
                    # Linha vazia: verifica se a próxima volta a ter pipe
                    if j + 1 < len(linhas) and '|' in linhas[j+1]:
                        tabela_linhas.append(linhas[j]) # Mantém linha vazia interna
                        j += 1
                    else:
                        break # Fim da tabela
                else:
                    break # Texto normal, fim da tabela

            # Verifica se capturou algo útil
            if len(tabela_linhas) > 0:
                tabelas_pendentes.append({
                    'titulo': titulo_tabela,
                    'linhas': tabela_linhas
                })
                i = j # Pula as linhas processadas
                continue
            else:
                # Falso positivo? Devolve o título se tinhamos pego
                if titulo_tabela:
                    resultado.append(titulo_tabela)
        
        # Se não for tabela nem título, adiciona linha normal
        resultado.append(linha)
        i += 1
    
    # 3. FINAL DO DOCUMENTO
    # Se sobraram tabelas no buffer, despeja agora
    if tabelas_pendentes:
        resultado.append('')
        for t_info in tabelas_pendentes:
            if t_info['titulo']:
                resultado.append(t_info['titulo'])
            resultado.extend(t_info['linhas'])
            resultado.append('')
            
    return '\n'.join(resultado)

def quebrar_paragrafos_longos(texto, max_chars=400, max_sentencas=4):
    """
    Quebra parágrafos que excedem limite de chars OU número de sentenças.
    Preserva listas, tabelas, citações e blocos especiais.
    """
    logger.info(f"✂️ Quebrando parágrafos > {max_chars} chars ou > {max_sentencas} sentenças...")
    
    paragrafos = texto.split('\n\n')
    resultado = []
    quebras = 0
    
    for para in paragrafos:
        linha_strip = para.strip()
        
        # PRESERVAR: títulos, listas, tabelas, citações, blocos de código
        if (linha_strip.startswith('#') or 
            linha_strip.startswith('-') or 
            linha_strip.startswith('* ') or 
            linha_strip.startswith('|') or
            linha_strip.startswith('>') or
            linha_strip.startswith('```') or
            re.match(r'^\d+\.', linha_strip)):
            resultado.append(para)
            continue
        
        # Verifica se precisa quebrar
        num_chars = len(para)
        sentencas = re.split(r'(?<=[.!?])\s+', para)
        num_sentencas = len(sentencas)
        
        if num_chars <= max_chars and num_sentencas <= max_sentencas:
            resultado.append(para)
            continue
        
        # QUEBRAR: Agrupa em blocos de até max_sentencas
        quebras += 1
        subparagrafos = []
        bloco_atual = []
        chars_atual = 0
        sentencas_no_bloco = 0
        
        for sentenca in sentencas:
            teste_chars = chars_atual + len(sentenca)
            
            # Se adicionar essa sentença ultrapassar AMBOS os limites ou o limite de sentenças
            # E já temos algo no bloco...
            if (teste_chars > max_chars or sentencas_no_bloco >= max_sentencas) and bloco_atual:
                subparagrafos.append(' '.join(bloco_atual).strip())
                bloco_atual = [sentenca]
                chars_atual = len(sentenca)
                sentencas_no_bloco = 1
            else:
                bloco_atual.append(sentenca)
                chars_atual = teste_chars
                sentencas_no_bloco += 1
        
        # Adiciona último bloco
        if bloco_atual:
            subparagrafos.append(' '.join(bloco_atual).strip())
        
        resultado.append('\n\n'.join(subparagrafos))
        
    if quebras > 0:
        logger.info(f"   ✅ {quebras} parágrafos foram ajustados.")
        
    return '\n\n'.join(resultado)

# =============================================================================
# FLUXO PRINCIPAL
# =============================================================================

def formatar_transcricao(transcricao_completa, usar_cache=True, input_file=None):
    # ⚡ FORÇAR USO DE VERTEX AI - AI Studio desabilitado
    rate_limiter.max_rpm = 60
    # Tenta pegar o projeto do ambiente, se não tiver, usa o hardcoded como fallback
    project_id = os.getenv("GOOGLE_CLOUD_PROJECT", "gen-lang-client-0727883752")
    logger.info(f" Usando Vertex AI (Project: {project_id})")
    logger.info(" Rate Limit: 60 RPM")
    logger.info("🔥 AI Studio DESABILITADO - Usando apenas Vertex AI")
    
    # Verificar se as credenciais do Vertex AI estão configuradas
    if os.getenv("GOOGLE_APPLICATION_CREDENTIALS") is None:
        logger.error("❌ Nenhuma autenticação configurada para Vertex AI.")
        logger.error("Configure a variável de ambiente GOOGLE_APPLICATION_CREDENTIALS:")
        logger.error("  export GOOGLE_APPLICATION_CREDENTIALS='/path/to/service-account.json'")
        sys.exit(1)

    client = genai.Client(
        vertexai=True,
        project=project_id,
        location="us-central1"
    )
    
    # Dry-run mode
    if '--dry-run' in sys.argv:
        logger.info("🔍 MODO DRY-RUN: Validando divisão de chunks")
        chunks = dividir_sequencial(transcricao_completa)
        validar_chunks(chunks, transcricao_completa)
        for i, c in enumerate(chunks):
            print(f"\n  Chunk {i+1}/{len(chunks)}:")
            print(f"    Posição: {c['inicio']:,} → {c['fim']:,} ({c['fim']-c['inicio']:,} chars)")
            inicio_preview = transcricao_completa[c['inicio']:c['inicio']+80].replace('\n', '↵')
            fim_preview = transcricao_completa[max(0, c['fim']-80):c['fim']].replace('\n', '↵')
            print(f"    Início: {inicio_preview}...")
            print(f"    Fim: ...{fim_preview}")
        sys.exit(0)
    
    tamanho_total = len(transcricao_completa)
    logger.info(f"📊 Tamanho: {tamanho_total:,} caracteres")
    
    chunks = dividir_sequencial(transcricao_completa)
    
    # v2.7: Validação rigorosa
    if not validar_chunks(chunks, transcricao_completa):
        logger.error("❌ Chunks inválidos! Abortando.")
        sys.exit(1)
    
    num_partes = len(chunks)
    chunks_info = [{'inicio': c['inicio'], 'fim': c['fim']} for c in chunks]
    
    estimar_custo(transcricao_completa, usar_cache, num_partes)
    
    if num_partes == 1:
        return processar_simples(client, transcricao_completa), None
    
    checkpoint = None
    resultados = []
    inicio_secao = 0
    
    if input_file:
        checkpoint = load_checkpoint(input_file)
        if checkpoint:
            logger.info(f"📁 Checkpoint: seção {checkpoint['secao_atual']}/{checkpoint['total_secoes']}")
            resposta = input("   Continuar? (s/n): ").strip().lower()
            if resposta == 's':
                resultados = checkpoint['resultados']
                inicio_secao = checkpoint['secao_atual']
                chunks_info = checkpoint['chunks_info']
            else:
                delete_checkpoint(input_file)
    
    # v2.9: Cache DESABILITADO (Flash é barato e input < 32k tokens não compensa)
    cache = None
    
    # v2.8: Mapeamento estrutural prévio
    estrutura_global = []
    if num_partes > 1:  # Mapeia sempre que houver múltiplas partes (mesmo retomando, pois checkpoint não salva estrutura)
        estrutura_global = mapear_estrutura(client, transcricao_completa)
        
        # v2.10: Forçar limite de 3 níveis (pedido usuário)
        estrutura_global = filtrar_niveis_execessivos(estrutura_global, max_nivel=3)
        
        # v2.10: Simplifica se for muito grande
        estrutura_global = simplificar_estrutura_se_necessario(estrutura_global)
    
    try:
        iterator = range(inicio_secao, num_partes)
        if tqdm:
            iterator = tqdm(iterator, desc="Formatando", initial=inicio_secao, total=num_partes)
        
        for i in iterator:
            chunk = chunks_info[i]
            texto_chunk = transcricao_completa[chunk['inicio']:chunk['fim']]
            
            # Contexto com validação
            contexto_estilo = ""
            if i > 0 and resultados:
                raw_context = resultados[-1][-CONTEXTO_ESTILO:]
                if len(raw_context.split()) < 50 or "[!WARNING]" in raw_context:
                    logger.warning(f"Contexto chunk {i+1} descartado")
                else:
                    contexto_estilo = raw_context
            
            # v2.10: Extrair último título do chunk anterior para anti-duplicação
            ultimo_titulo = None
            if resultados:
                texto_anterior = resultados[-1]
                for linha in reversed(texto_anterior.split('\n')[-30:]):
                    if linha.strip().startswith('##'):
                        ultimo_titulo = re.sub(r'^#+\s*(?:\d+(?:\.\d+)*\.?)?\s*', '', linha).strip()
                        break
            
            resultado = processar_chunk(
                client, cache, texto_chunk,
                i + 1, num_partes,
                contexto_estilo=contexto_estilo,
                estrutura_global=estrutura_global,
                ultimo_titulo=ultimo_titulo
            )
            
            # v2.10: Smart Stitching - Remove eco do contexto
            if contexto_estilo:
                resultado = remover_eco_do_contexto(resultado, contexto_estilo)
            
            # v2.10: Smart Stitching - Remove título duplicado na fronteira
            texto_acumulado = '\n\n'.join(resultados) if resultados else ""
            resultado = limpar_inicio_redundante(resultado, texto_acumulado)
            
            resultados.append(resultado)
            
            if input_file:
                save_checkpoint(input_file, resultados, chunks_info, i + 1)
            
            if not tqdm:
                logger.info(f"✅ Seção {i+1}/{num_partes}")
        
        # v2.7: Post-processing em múltiplas passadas
        logger.info("🧹 Iniciando limpeza (v2.7)...")
        
        texto_final = '\n\n'.join(resultados)
        
        logger.info("  Passada 1: Removendo duplicações literais...")
        texto_final = remover_duplicacoes_literais(texto_final)
        
        # logger.info("  Passada 2: Detectando seções duplicadas...")
        # texto_final = remover_secoes_duplicadas(texto_final)  # DESATIVADO: causava falsos positivos
        
        
        # v2.10: Reordenação do Pipeline (Tabelas -> Numeração -> Parágrafos)
        
        logger.info("  Passada 2: Reorganizando tabelas (Smart Layout)...")
        texto_final = mover_tabelas_para_fim_de_secao(texto_final)
        
        logger.info("  Passada 3: Numerando títulos...")
        texto_final = numerar_titulos(texto_final)
        
        logger.info("  Passada 4: Ajustando parágrafos longos...")
        texto_final = quebrar_paragrafos_longos(texto_final, max_chars=400, max_sentencas=4)
        
        # Validação final
        palavras_in = len(transcricao_completa.split())
        palavras_out = len(texto_final.split())
        razao = palavras_out / palavras_in if palavras_in > 0 else 1.0
        
        logger.info(f"✅ Validação: {razao:.0%} do original ({palavras_out:,}/{palavras_in:,} palavras)")
        
        if razao < THRESHOLD_CRITICO:
            if FIDELIDADE_MODE:
                logger.error(f"❌ POSSÍVEL PERDA DE CONTEÚDO ({razao:.0%})")
                logger.error(f"   Esperado: >{THRESHOLD_CRITICO:.0%} | Obtido: {razao:.0%}")
            elif APOSTILA_MODE:
                logger.warning(f"⚠️  Texto condensado: {razao:.0%}")
                logger.info(f"   ✅ Esperado no modo {MODO_NOME}")
            
            if razao < 0.30 or (FIDELIDADE_MODE and razao < THRESHOLD_CRITICO):
                resposta = input("\n   Continuar? (s/n): ").strip().lower()
                if resposta != 's':
                    logger.info("Cancelado.")
                    sys.exit(1)
        
        if input_file:
            delete_checkpoint(input_file)
        
        return texto_final, cache, client
    
    except KeyboardInterrupt:
        logger.warning("\n⚠️  Interrompido. Checkpoint salvo.")
        sys.exit(1)

# =============================================================================
# EXPORTAÇÃO WORD
# =============================================================================

def create_toc(doc):
    """Adiciona Sumário (Table of Contents) nativo do Word"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def _format_inline_markdown(paragraph, text):
    """Formata markdown inline (negrito, itálico, código)"""
    paragraph.clear()
    
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        if match.group(0).startswith('***'):
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(0).startswith('**'):
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(0).startswith('*'):
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(0).startswith('`'):
            run = paragraph.add_run(match.group(5))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(200, 0, 0)
        
        last_end = match.end()
    
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

def _add_table_to_doc(doc, rows):
    """Adiciona tabela formatada ao documento Word"""
    if len(rows) < 2:
        return
    
    max_cols = max(len(row) for row in rows)
    if max_cols == 0:
        return
    
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(rows):
        for j in range(max_cols):
            cell = table.rows[i].cells[j]
            cell_text = row_data[j] if j < len(row_data) else ""
            
            # Formata markdown dentro da célula
            _format_inline_markdown(cell.paragraphs[0], cell_text)
            
            # Alinhamento padrão: Esquerda
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '0066CC')
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                 # Conteúdo normal da tabela: Esquerda
                 for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def save_as_word(formatted_text, video_name, output_file):
    """Salva markdown formatado como documento Word (.docx)"""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx não disponível. Salvando apenas Markdown.")
        return None
    
    logger.info("📄 Gerando documento Word...")
    
    doc = Document()
    
    # Margens
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    
    # Título principal
    title = doc.add_heading(video_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Data de geração
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')} - Modo: {MODO_NOME}")
    date_run.italic = True
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Sumário
    doc.add_heading('Sumário', level=1)
    create_toc(doc)
    doc.add_page_break()
    
    # Processa conteúdo markdown
    lines = formatted_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Tabelas
        if '|' in line and not in_table:
            in_table = True
            table_rows = []
        
        if in_table:
            # Ignora linha separadora (ex: |---| ou | :--- | :--- | :--- |)
            is_separator = re.match(r'^\s*\|[\s:|-]+\|[\s:|-]*$', line)
            
            if '|' in line and not is_separator:
                table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
            
            if '|' not in line or i == len(lines) - 1:
                if len(table_rows) > 0:
                    _add_table_to_doc(doc, table_rows)
                in_table = False
                table_rows = []
                if '|' not in line:
                    continue
            i += 1
            continue
        
        # Headings
        if line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# ') and line != f"# {video_name}":
            doc.add_heading(line[2:], level=1)
        # Separadores
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.add_run('_' * 80).font.color.rgb = RGBColor(192, 192, 192)
        # Quotes
        elif line.startswith('>'):
            p = doc.add_paragraph(style='Quote')
            p.paragraph_format.left_indent = Cm(4.0)  # Recuo de 4 cm da margem esquerda
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            run = p.add_run(line[1:].strip())
            run.italic = True
            run.font.size = Pt(10) # Geralmente citações longas têm fonte menor
        # Listas não-ordenadas
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            # Forçar recuo de 1,5cm
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-0.63)
            
            _format_inline_markdown(p, line[2:])
            
        # Listas numeradas
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
            p = doc.add_paragraph(style='Normal')
            # Recuo padronizado de 1,5cm
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-0.63) # Mantém hanging indent para o número
            _format_inline_markdown(p, line)
            
        # Parágrafo normal
        else:
            p = doc.add_paragraph()
            # Espaçamento 1.5 (Line Spacing = 1.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            # Espaçamento antes e após parágrafo (6pt)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6) 
            
            # Recuo de 1ª linha (1cm)
            p.paragraph_format.first_line_indent = Cm(1.0)
            
            # Justificado
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            _format_inline_markdown(p, line)
            
            # Forçar fonte 12pt para texto normal
            for run in p.runs:
                run.font.size = Pt(12)
        
        i += 1
    
    # Flush final (caso o arquivo termine com tabela seguida de linhas em branco)
    if in_table and len(table_rows) > 0:
        _add_table_to_doc(doc, table_rows)
        
    doc.save(output_file)
    logger.info(f"✅ Word salvo: {output_file}")
    return output_file

def salvar_resultado(conteudo, arquivo):
    with open(arquivo, 'w', encoding='utf-8') as f:
        f.write(conteudo)
    logger.info(f"✅ Markdown salvo: {arquivo}")

# =============================================================================
# MAIN
# =============================================================================

def main():
    if len(sys.argv) < 2 or '--help' in sys.argv:
        print("=" * 70)
        print("FORMATADOR v2.7 - ANTI-DUPLICAÇÃO")
        print("=" * 70)
        print("\nUso: python format_transcription_gemini.py <entrada.txt> [saida]")
        print("\nOpções:")
        print("  --dry-run    Valida chunks e mostra preview")
        print("  --help       Mostra esta mensagem")
        print("\n🛡️  CORREÇÕES v2.7:")
        print("  • Detecção agressiva de seções duplicadas")
        print("  • Validação rigorosa de chunks sequenciais")
        print("  • Delimitadores de contexto mais visíveis")
        print("  • Post-processing em múltiplas passadas")
        print("  • Cache desabilitado (debug)")
        print("\nDependências:")
        print("  pip install google-genai python-docx tqdm")
        sys.exit(1)
    
    arquivo_entrada = sys.argv[1]
    
    if len(sys.argv) > 2 and not sys.argv[2].startswith('--'):
        arquivo_saida = sys.argv[2]
    else:
        base = arquivo_entrada.replace('.txt', '_formatada_v2.7')
        arquivo_saida = f"{base}.md"
    
    video_name = Path(arquivo_entrada).stem
    
    logger.info("=" * 60)
    logger.info(f"FORMATADOR v2.7 ANTI-DUPLICAÇÃO - Modo {MODO_NOME}")
    logger.info("=" * 60)
    logger.info(f"📂 Entrada: {arquivo_entrada}")
    logger.info(f"📝 Saída: {arquivo_saida}")
    
    transcricao = carregar_transcricao(arquivo_entrada)
    
    try:
        resultado, cache, client = formatar_transcricao(transcricao, input_file=arquivo_entrada)
    except Exception as e:
        logger.error(f"\n❌ Falha: {e}", exc_info=True)
        sys.exit(1)
    
    salvar_resultado(resultado, arquivo_saida)
    
    # v2.8: Verificação de cobertura e duplicações
    verificar_cobertura(transcricao, resultado, arquivo_saida)

    # v2.9: Auditoria Legal Pós-Processamento
    if AUDIT_AVAILABLE:
        report_path = arquivo_saida.replace('.md', '_RELATORIO_AUDITORIA.md')
        auditar_consistencia_legal(client, resultado, report_path)

    if DOCX_AVAILABLE:
        arquivo_docx = arquivo_saida.replace('.md', '.docx')
        save_as_word(resultado, video_name, arquivo_docx)
    
    tokens_in = len(transcricao) // 4
    tokens_out = len(resultado) // 4
    custo = (tokens_in * PRECO_INPUT_SEM_CACHE + tokens_out * PRECO_OUTPUT) / 1_000_000
    
    logger.info("=" * 60)
    logger.info(f"💰 Custo: ${custo:.4f} USD")
    logger.info(f"✨ Concluído! (v2.8 com verificação)")
    logger.info("=" * 60)

if __name__ == "__main__":
    main()
