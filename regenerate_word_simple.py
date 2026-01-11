#!/usr/bin/env python3
"""
Script ULTRA SIMPLES para gerar documento Word - sem elementos que podem corromper.
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from colorama import Fore, init
from datetime import datetime

init(autoreset=True)


def sanitize_text(text):
    """Remove caracteres problemáticos para XML"""
    if not text:
        return ""
    # Remove caracteres de controle
    result = []
    for char in text:
        code = ord(char)
        # Permite apenas caracteres válidos
        if code == 9 or code == 10 or code == 13 or (code >= 32 and code != 127):
            if code < 65536:  # Basic Multilingual Plane
                result.append(char)
            else:
                result.append(' ')  # Substitui caracteres fora do BMP
        # else: ignora caracteres de controle
    return ''.join(result)


def add_formatted_text(paragraph, text):
    """Adiciona texto com formatação inline básica"""
    text = sanitize_text(text)
    
    # Processa negrito e itálico
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    
    for part in parts:
        if not part:
            continue
        part = sanitize_text(part)
        try:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)
        except:
            # Fallback: texto puro
            clean = re.sub(r'\*+', '', part)
            if clean:
                paragraph.add_run(clean)


def add_simple_table(doc, table_data):
    """Adiciona tabela simples"""
    if not table_data:
        return
    
    # Filtra e limpa dados
    clean_data = []
    for row in table_data:
        if row and any(cell.strip() for cell in row):
            clean_row = [sanitize_text(cell.strip()) for cell in row]
            clean_data.append(clean_row)
    
    if not clean_data:
        return
    
    try:
        num_cols = max(len(row) for row in clean_data)
        table = doc.add_table(rows=len(clean_data), cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(clean_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text
                    # Negrito no cabeçalho
                    if i == 0:
                        for para in row.cells[j].paragraphs:
                            for run in para.runs:
                                run.bold = True
        
        doc.add_paragraph()
    except Exception as e:
        print(f"⚠️ Erro na tabela: {e}")


def markdown_to_word_simple(markdown_text, output_path, title="Apostila de Direito Administrativo"):
    """Converte Markdown para Word - versão SIMPLES e robusta"""
    
    print(f"{Fore.CYAN}📄 Gerando documento Word (versão simplificada)...")
    
    doc = Document()
    
    # Margens básicas
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)
    
    # ============================================
    # CAPA SIMPLES
    # ============================================
    
    for _ in range(5):
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    doc.add_paragraph()
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Material de Estudo para Concursos")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_para.add_run("_" * 40)
    
    for _ in range(4):
        doc.add_paragraph()
    
    concurso_para = doc.add_paragraph()
    concurso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    concurso_run = concurso_para.add_run("PGM/PGE - Procuradorias")
    concurso_run.font.name = 'Arial'
    concurso_run.font.size = Pt(14)
    concurso_run.font.bold = True
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(datetime.now().strftime("Dezembro de 2024"))
    
    doc.add_page_break()
    
    # ============================================
    # CONTEÚDO
    # ============================================
    
    lines = markdown_text.split('\n')
    current_table_data = []
    in_table = False
    line_count = 0
    total_lines = len(lines)
    
    print(f"{Fore.GREEN}   Processando {total_lines:,} linhas...")
    
    for line in lines:
        line_count += 1
        if line_count % 2000 == 0:
            print(f"   Progresso: {line_count:,}/{total_lines:,}")
        
        line_stripped = line.strip()
        
        if not line_stripped:
            if in_table and current_table_data:
                add_simple_table(doc, current_table_data)
                current_table_data = []
                in_table = False
            continue
        
        # Ignora marcadores de chunk
        if '[Fim da transcrição' in line_stripped or '[Continuação' in line_stripped:
            continue
        
        # Detecta tabelas
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            in_table = True
            if not re.match(r'^[\|\-\:\s]+$', line_stripped):
                cells = [cell.strip() for cell in line_stripped.split('|')[1:-1]]
                if cells and any(c for c in cells):
                    current_table_data.append(cells)
            continue
        
        if in_table and current_table_data:
            add_simple_table(doc, current_table_data)
            current_table_data = []
            in_table = False
        
        try:
            line_stripped = sanitize_text(line_stripped)
            
            if line_stripped.startswith('# '):
                doc.add_heading(line_stripped[2:], level=1)
            elif line_stripped.startswith('## '):
                doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith('### '):
                doc.add_heading(line_stripped[4:], level=3)
            elif line_stripped.startswith('#### '):
                doc.add_heading(line_stripped[5:], level=4)
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, line_stripped[2:])
            elif re.match(r'^\d+\.\s', line_stripped):
                p = doc.add_paragraph(style='List Number')
                text = re.sub(r'^\d+\.\s*', '', line_stripped)
                add_formatted_text(p, text)
            elif line_stripped.startswith('---'):
                # Separador
                sep = doc.add_paragraph()
                sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sep.add_run("* * *")
            elif line_stripped.startswith('>'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                add_formatted_text(p, line_stripped[1:].strip())
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, line_stripped)
        except Exception as e:
            # Fallback extremo
            try:
                clean_text = re.sub(r'[^\w\s.,;:!?()-]', '', line_stripped)
                if clean_text.strip():
                    doc.add_paragraph(clean_text)
            except:
                pass
    
    # Finaliza última tabela
    if in_table and current_table_data:
        add_simple_table(doc, current_table_data)
    
    # Salva
    doc.save(output_path)
    print(f"{Fore.GREEN}✅ Documento Word salvo: {output_path}")
    
    # Verifica tamanho do arquivo
    size = os.path.getsize(output_path)
    print(f"{Fore.GREEN}   Tamanho: {size / 1024 / 1024:.2f} MB")


def main():
    md_file = "/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/Administrativo_FORMATADO.md"
    
    if len(sys.argv) >= 2:
        md_file = sys.argv[1]
    
    if not os.path.exists(md_file):
        print(f"{Fore.RED}❌ Arquivo não encontrado: {md_file}")
        sys.exit(1)
    
    print(f"{Fore.CYAN}📂 Lendo: {md_file}")
    
    with open(md_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    print(f"{Fore.GREEN}   ✅ {len(markdown_text):,} caracteres")
    
    output_docx = md_file.replace('.md', '_V2.docx')
    
    markdown_to_word_simple(markdown_text, output_docx)
    
    print(f"\n{Fore.GREEN}{'='*50}")
    print(f"{Fore.GREEN}✅ ARQUIVO: {output_docx}")
    print(f"{Fore.GREEN}{'='*50}")


if __name__ == "__main__":
    main()
