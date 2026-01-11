#!/usr/bin/env python3
"""
VomoUnified - Pipeline de Apostilas Jurídicas com Multi-Engine Support
Suporta: Google Gemini 2.5 Flash e OpenAI GPT-5 Mini
Otimizado para Apple Silicon (M3 Pro)

Uso:
    python vomo_unified.py aula.mp4                    # Gemini (padrão)
    python vomo_unified.py aula.mp4 --engine gpt5      # GPT-5 Mini
    python vomo_unified.py aula.mp4 --engine all       # Ambos (A/B test)
"""

import os
import sys
import time
import subprocess
import hashlib
import json
import re
import argparse
from pathlib import Path
from colorama import Fore, init, Style

# Inicializa cores
init(autoreset=True)

# Imports Condicionais (Bibliotecas de Transcrição)
try:
    import mlx_whisper
except ImportError:
    mlx_whisper = None
    print(f"{Fore.YELLOW}⚠️  mlx_whisper não instalado. Use: pip install mlx-whisper")

try:
    from pyannote.audio import Pipeline
    import torch
    HF_TOKEN_ENV = os.getenv("HUGGING_FACE_TOKEN")
except ImportError:
    Pipeline = None
    torch = None
    HF_TOKEN_ENV = None

# Imports de LLM (Lazy Loading)
try:
    import google.generativeai as genai
    from google.generativeai.types import GenerationConfig
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print(f"{Fore.YELLOW}⚠️  google-generativeai não instalado")

try:
    from openai import OpenAI, APIError
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print(f"{Fore.YELLOW}⚠️  openai não instalado")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ==================== CONFIGURAÇÃO GLOBAL ====================

# Preços (USD por 1M tokens)
GEMINI_PRICE_INPUT = 0.075
GEMINI_PRICE_OUTPUT = 0.30
GPT5_PRICE_INPUT = 0.25
GPT5_PRICE_OUTPUT = 2.00

SYSTEM_PROMPT = """# SEU PAPEL
Você é um **revisor sênior de material didático jurídico** com 20 anos de experiência em preparatórios para concursos de procuradoria (PGM/PGE).

# MISSÃO
Transformar a transcrição bruta em uma **apostila didática** formal, impessoal e expositiva, **PRESERVANDO 100% DO CONTEÚDO TÉCNICO**.

⚠️ **REGRA ABSOLUTA: PRESERVE 100% DO CONTEÚDO - NÃO RESUMA, NÃO OMITA NADA.**

# EXEMPLOS DE TRANSFORMAÇÃO (Few-Shot)

## Exemplo 1: Transformação de Linguagem
**INPUT:** "Então pessoal, a Lei 8.666, né, ela fala lá no artigo 37..."
**OUTPUT:** "Observa-se que o *Art. 37 da Lei nº 8.666/93* estabelece..."

## Exemplo 2: Preservação de Dicas
**INPUT:** "Isso aqui cai muito em prova, hein!"
**OUTPUT:** "Reveste-se de especial importância para fins de certames públicos..."

# DIRETRIZES
1. **Estrutura**: Use títulos Markdown (##, ###) hierárquicos
2. **Linguagem**: Formal, 3ª pessoa, impessoal
3. **Conteúdo**: Mantenha TODOS os exemplos, jurisprudência e dicas
4. **Formatação**: Leis em itálico (*Lei nº 8.666/93*), destaques em negrito
5. **Síntese**: Crie tabelas ao final de tópicos complexos

# PROIBIDO
- Resumir ou cortar exemplos
- Usar 1ª pessoa
- Inventar dados
- Meta-comentários

# FORMATO DE SAÍDA
Apenas o conteúdo da apostila em Markdown, começando com # Título.
"""

# ==================== MÓDULO DE TRANSCRIÇÃO (COMPARTILHADO) ====================

class AudioTranscriber:
    def __init__(self, model_size="large-v3-turbo"):
        self.model_size = model_size
        self.cache_dir = Path(".cache_vomo_unified")
        self.cache_dir.mkdir(exist_ok=True)
        
    def optimize_audio(self, file_path):
        """Converte para WAV 16kHz Mono"""
        path = Path(file_path)
        print(f"{Fore.YELLOW}⚡ Verificando áudio: {path.name}...")
        
        # Check MP3 rápido
        mp3 = path.with_suffix('.mp3')
        if mp3.exists() and mp3 != path:
            print(f"   📂 Usando MP3 existente")
            return str(mp3)

        file_hash = hashlib.md5(str(path).encode()).hexdigest()[:8]
        out_wav = self.cache_dir / f"temp_{file_hash}.wav"
        
        if out_wav.exists():
            print(f"   📂 WAV otimizado já existe")
            return str(out_wav)
            
        print(f"   🔄 Convertendo para WAV otimizado...")
        try:
            subprocess.run([
                'ffmpeg', '-i', str(path),
                '-vn', '-ac', '1', '-ar', '16000',
                str(out_wav), '-y', '-hide_banner', '-loglevel', 'error'
            ], check=True, capture_output=True)
            return str(out_wav)
        except Exception as e:
            print(f"{Fore.RED}❌ Erro no ffmpeg: {e}")
            sys.exit(1)

    def transcribe(self, audio_path):
        """Pipeline MLX-Whisper + Pyannote"""
        if not mlx_whisper:
            raise ImportError("Instale mlx-whisper: pip install mlx-whisper")
            
        print(f"{Fore.GREEN}🎙️  Iniciando Transcrição (MLX GPU)...")
        start = time.time()
        
        # Cache baseado no arquivo de áudio
        cache_name = Path(audio_path).stem
        cache_file = self.cache_dir / f"{cache_name}_transcription.json"
        
        if cache_file.exists():
            print(f"{Fore.CYAN}   📂 Cache de transcrição encontrado!")
            with open(cache_file, 'r', encoding='utf-8') as f:
                return json.load(f)['text']

        # 1. Whisper com parâmetros otimizados
        print("   🔍 Executando Whisper...")
        result = mlx_whisper.transcribe(
            audio_path,
            path_or_hf_repo=f"mlx-community/whisper-{self.model_size}",
            language="pt",
            word_timestamps=True,
            fp16=True,
            no_speech_threshold=0.6,
            logprob_threshold=-1.0,
            compression_ratio_threshold=2.4,
            condition_on_previous_text=True,
            verbose=False
        )
        
        duration = result.get('duration', 0)
        elapsed = time.time() - start
        rtf = elapsed / duration if duration > 0 else 0
        
        print(f"{Fore.GREEN}   ✅ Whisper concluído em {elapsed:.1f}s (RTF: {rtf:.2f}x)")
        
        final_text = self._format_segments(result['segments'])
        
        # 2. Diarização (Opcional)
        if Pipeline and torch and HF_TOKEN_ENV:
            print("   🗣️  Executando Diarização (Pyannote)...")
            try:
                pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-3.1",
                    use_auth_token=HF_TOKEN_ENV
                )
                device = "mps" if torch.backends.mps.is_available() else "cpu"
                pipeline.to(torch.device(device))
                
                diarization = pipeline(audio_path)
                final_text = self._align_speakers(result['segments'], diarization)
                print(f"{Fore.GREEN}   ✅ Diarização concluída")
            except Exception as e:
                print(f"{Fore.YELLOW}   ⚠️ Diarização falhou ({e}), usando texto simples")

        # Salva Cache
        with open(cache_file, 'w', encoding='utf-8') as f:
            json.dump({
                "text": final_text,
                "model": self.model_size,
                "duration": duration,
                "timestamp": time.strftime('%Y-%m-%d %H:%M:%S')
            }, f, ensure_ascii=False, indent=2)
            
        return final_text

    def _format_segments(self, segments):
        """Formata segmentos simples com timestamps"""
        parts = []
        for s in segments:
            ts = self._fmt_ts(s['start'])
            parts.append(f"[{ts}] {s['text'].strip()}")
        return " ".join(parts)
    
    def _fmt_ts(self, sec):
        m, s = divmod(sec, 60)
        h, m = divmod(m, 60)
        return f"{int(h):02d}:{int(m):02d}:{int(s):02d}" if h > 0 else f"{int(m):02d}:{int(s):02d}"

    def _align_speakers(self, segments, diarization):
        """Alinha segmentos com speakers"""
        output = []
        current_spk = None
        spk_turns = list(diarization.itertracks(yield_label=True))
        
        for seg in segments:
            mid = (seg['start'] + seg['end']) / 2
            best_spk = "SPEAKER 0"
            
            for turn, _, spk in spk_turns:
                if turn.start <= mid <= turn.end:
                    best_spk = f"SPEAKER {int(spk.split('_')[-1]) + 1}"
                    break
            
            if best_spk != current_spk:
                output.append(f"\n\n**{best_spk}**:\n")
                current_spk = best_spk
            
            ts = self._fmt_ts(seg['start'])
            output.append(f"[{ts}] {seg['text'].strip()}")
            
        return " ".join(output)

# ==================== ENGINES DE LLM ====================

class EngineGemini:
    def __init__(self, model_name="gemini-2.5-flash"):
        if not GEMINI_AVAILABLE:
            raise ImportError("Instale: pip install google-generativeai")
        
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY não configurada")
        
        self.name = "Gemini"
        self.tag = "gemini"
        self.price_in = GEMINI_PRICE_INPUT
        self.price_out = GEMINI_PRICE_OUTPUT
        self.script_path = Path(__file__).parent / "transcriber_gemini.py"
        
        if not self.script_path.exists():
            raise FileNotFoundError(f"Script não encontrado: {self.script_path}")

    def generate(self, transcription, title):
        print(f"   🚀 Usando {self.name} com chunking strategy...")
        start = time.time()
        
        try:
            # Salva transcrição temporária
            temp_dir = Path(".cache_vomo_unified")
            temp_dir.mkdir(exist_ok=True)
            temp_input = temp_dir / f"temp_{hashlib.md5(title.encode()).hexdigest()[:8]}.txt"
            temp_output = temp_dir / f"temp_{hashlib.md5(title.encode()).hexdigest()[:8]}_formatted"
            
            with open(temp_input, 'w', encoding='utf-8') as f:
                f.write(transcription)
            
            # Executa script externo
            result = subprocess.run(
                [sys.executable, str(self.script_path), str(temp_input), str(temp_output)],
                capture_output=True,
                text=True,
                check=False
            )
            
            if result.returncode != 0:
                print(f"{Fore.RED}   ❌ Erro no script: {result.stderr}")
                return f"# ERRO GEMINI\\n\\n{result.stderr}"
            
            # Lê resultado
            output_md = Path(str(temp_output) + ".md")
            if not output_md.exists():
                # Tenta sem extensão
                output_md = temp_output
            
            if not output_md.exists():
                return f"# ERRO\\n\\nArquivo de saída não encontrado"
            
            with open(output_md, 'r', encoding='utf-8') as f:
                content = f.read()
            
            elapsed = time.time() - start
            
            # Estimativa de custo (aproximada)
            tokens_in = len(transcription) // 4
            tokens_out = len(content) // 4
            cost = (tokens_in * self.price_in + tokens_out * self.price_out) / 1_000_000
            
            print(f"{Fore.GREEN}   ✅ Concluído em {elapsed:.1f}s")
            print(f"   💰 Custo estimado: ${cost:.4f} USD")
            
            # Limpa arquivos temporários
            temp_input.unlink(missing_ok=True)
            output_md.unlink(missing_ok=True)
            
            return content
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro: {e}")
            import traceback
            traceback.print_exc()
            return f"# ERRO GEMINI\\n\\n{str(e)}"

class EngineGPT:
    def __init__(self, model_name="gpt-5-mini-2025-08-07"):
        if not OPENAI_AVAILABLE:
            raise ImportError("Instale: pip install openai")
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY não configurada")
        
        self.client = OpenAI(api_key=api_key)
        self.model_name = model_name
        self.name = "GPT-5 Mini"
        self.tag = "gpt5"
        self.price_in = GPT5_PRICE_INPUT
        self.price_out = GPT5_PRICE_OUTPUT

    def generate(self, transcription, title):
        print(f"   🚀 Enviando para {self.name} (128k output)...")
        start = time.time()
        
        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": f"TÍTULO: {title}\n\nTRANSCRIÇÃO:\n{transcription}"}
                ],
                max_completion_tokens=110000  # GPT-5 uses max_completion_tokens, temperature defaults to 1
            )
            
            elapsed = time.time() - start
            content = response.choices[0].message.content
            
            # Estimativa de custo
            tokens_in = len(transcription) // 4
            tokens_out = len(content) // 4
            cost = (tokens_in * self.price_in + tokens_out * self.price_out) / 1_000_000
            
            print(f"{Fore.GREEN}   ✅ Concluído em {elapsed:.1f}s")
            print(f"   💰 Custo estimado: ${cost:.4f} USD")
            
            return content
            
        except Exception as e:
            print(f"{Fore.RED}   ❌ Erro: {e}")
            return f"# ERRO GPT-5\n\n{str(e)}"

# ==================== EXPORTAÇÃO DOCX ====================

class DocxExporter:
    @staticmethod
    def export(markdown_text, title, output_path):
        if not DOCX_AVAILABLE:
            return None
        
        print(f"{Fore.CYAN}   📄 Gerando DOCX...")
        
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        
        # Título
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        p = doc.add_paragraph()
        r = p.add_run(f"Gerado em {time.strftime('%d/%m/%Y às %H:%M')}")
        r.italic = True
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Processa Markdown
        for line in markdown_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:], style='Quote')
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.save(output_path)
        print(f"{Fore.GREEN}   ✅ DOCX salvo")
        return output_path

# ==================== ORQUESTRADOR ====================

class VomoUnified:
    def __init__(self, args):
        self.args = args
        self.transcriber = AudioTranscriber(model_size=args.whisper_model)
        self.engines = []
        
        # Inicializa Engines Selecionadas
        if args.engine in ['gemini', 'all']:
            try:
                self.engines.append(EngineGemini(model_name=args.gemini_model))
            except Exception as e:
                print(f"{Fore.RED}❌ Falha ao carregar Gemini: {e}")

        if args.engine in ['gpt5', 'all']:
            try:
                self.engines.append(EngineGPT(model_name=args.gpt_model))
            except Exception as e:
                print(f"{Fore.RED}❌ Falha ao carregar GPT: {e}")

        if not self.engines:
            print(f"{Fore.RED}❌ Nenhuma engine disponível. Verifique chaves de API.")
            sys.exit(1)

    def run(self, input_file):
        print(f"\n{Fore.CYAN}{'='*70}")
        print(f"{Fore.CYAN}🎬 VOMO UNIFIED - PROCESSAMENTO MULTI-ENGINE")
        print(f"{Fore.CYAN}{'='*70}\n")
        
        # 1. Transcrição (Única para todas as engines)
        audio_path = self.transcriber.optimize_audio(input_file)
        transcription = self.transcriber.transcribe(audio_path)
        
        file_stem = Path(input_file).stem
        output_dir = Path(self.args.output_dir)
        output_dir.mkdir(exist_ok=True)
        
        # Salva transcrição bruta
        raw_file = output_dir / f"{file_stem}_RAW.txt"
        with open(raw_file, 'w', encoding='utf-8') as f:
            f.write(transcription)
        print(f"   💾 Transcrição bruta: {raw_file.name}")

        # 2. Geração Multi-Engine
        print(f"\n{Fore.CYAN}{'='*70}")
        print(f"{Fore.CYAN}🧠 GERANDO APOSTILAS ({len(self.engines)} engine(s))")
        print(f"{Fore.CYAN}{'='*70}\n")

        for engine in self.engines:
            print(f"👉 Processando com: {Style.BRIGHT}{engine.name}{Style.RESET_ALL}")
            
            # Gera conteúdo
            content = engine.generate(transcription, file_stem)
            
            # Salva Markdown
            md_file = output_dir / f"{file_stem}_{engine.tag}.md"
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(content)
            
            size_kb = md_file.stat().st_size / 1024
            print(f"{Fore.GREEN}   ✅ Markdown: {md_file.name} ({size_kb:.1f} KB)")
            
            # Exporta DOCX se solicitado
            if self.args.export_docx:
                docx_file = output_dir / f"{file_stem}_{engine.tag}.docx"
                DocxExporter.export(content, file_stem, docx_file)
            
            print()
        
        print(f"{Fore.GREEN}{'='*70}")
        print(f"{Fore.GREEN}✅ PROCESSAMENTO CONCLUÍDO")
        print(f"{Fore.GREEN}{'='*70}")
        print(f"📁 Outputs em: {output_dir.absolute()}")

# ==================== CLI ENTRYPOINT ====================

if __name__ == "__main__":
    from dotenv import load_dotenv
    load_dotenv(override=True)

    parser = argparse.ArgumentParser(
        description="Vomo Unified - Multi-Engine Pipeline",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  python vomo_unified.py aula.mp4                    # Gemini (padrão)
  python vomo_unified.py aula.mp4 --engine gpt5      # GPT-5 Mini
  python vomo_unified.py aula.mp4 --engine all       # Ambos (A/B test)
  python vomo_unified.py aula.mp4 --no-docx          # Apenas Markdown
        """
    )
    
    parser.add_argument("input", help="Arquivo de vídeo/áudio")
    parser.add_argument("-o", "--output_dir", default="apostilas_unified",
                       help="Pasta de saída (default: apostilas_unified)")
    
    # Seleção de Engine
    parser.add_argument("--engine", choices=['gemini', 'gpt5', 'all'], default='gemini',
                       help="Engine LLM (default: gemini)")
    
    # Modelos Específicos
    parser.add_argument("--whisper_model", default="large-v3-turbo",
                       help="Modelo Whisper (default: large-v3-turbo)")
    parser.add_argument("--gemini_model", default="gemini-2.5-flash-preview-05-20",
                       help="Modelo Gemini")
    parser.add_argument("--gpt_model", default="gpt-5-mini-2025-08-07",
                       help="Modelo GPT")
    
    # Exportação
    parser.add_argument("--no-docx", dest="export_docx", action="store_false",
                       help="Não gerar arquivos DOCX")

    args = parser.parse_args()
    
    try:
        pipeline = VomoUnified(args)
        pipeline.run(args.input)
    except KeyboardInterrupt:
        print(f"\n{Fore.YELLOW}⚠️ Interrompido pelo usuário")
        sys.exit(0)
    except Exception as e:
        print(f"\n{Fore.RED}❌ Erro fatal: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
