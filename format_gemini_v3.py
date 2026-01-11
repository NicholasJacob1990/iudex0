#!/usr/bin/env python3
"""
Script para processar transcrições - VERSÃO ANTI-RESUMO
Chunks pequenos de 10k chars para GARANTIR preservação do conteúdo
"""

import os
import sys
import time
import re
import requests
from pathlib import Path
from docx import Document
from typing import List

# Configuração
OPENROUTER_API_KEY = "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"
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL_NAME = "google/gemini-2.5-flash"
# CHUNKS DE 10K PARA FORÇAR PRESERVAÇÃO
MAX_CHARS_PER_CHUNK = 10_000
OUTPUT_DIR = Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ')

PROMPT_FORMATACAO = """VOCÊ É UM REVISOR DE TEXTO. SUA ÚNICA FUNÇÃO É FORMATAR, NÃO RESUMIR.

⛔ REGRA ABSOLUTA: O TEXTO DE SAÍDA DEVE TER O MESMO TAMANHO OU MAIOR QUE A ENTRADA!

O que você DEVE fazer:
✅ Corrigir gramática e pontuação
✅ Transformar linguagem coloquial em português formal
✅ Organizar em parágrafos claros
✅ Manter ABSOLUTAMENTE TODO o conteúdo
✅ Manter exemplos, casos, dicas, explicações
✅ Usar negrito para termos importantes

O que você NÃO PODE fazer:
❌ NÃO RESUMA
❌ NÃO OMITA informações
❌ NÃO GENERALIZE
❌ NÃO CORTE exemplos
❌ NÃO REMOVA dicas de prova
❌ NÃO SINTETIZE ideias

Formatação:
- Use texto corrido (parágrafos)
- Negrito para conceitos importantes
- Mantenha primeira pessoa quando presente
- Remove apenas vícios de linguagem ("né", "tipo", "então assim")

{contexto}

TRANSCRIÇÃO PARA FORMATAR (mantenha TODO o conteúdo):

{texto}

TEXTO FORMATADO (mesmo tamanho ou maior):"""


def extrair_texto(caminho: str) -> str:
    path = Path(caminho)
    if path.suffix.lower() == '.docx':
        doc = Document(caminho)
        return '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        with open(caminho, 'r', encoding='utf-8') as f:
            return f.read()


def dividir_em_chunks(texto: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> List[str]:
    """Divide em chunks pequenos em quebras de parágrafo"""
    paragrafos = texto.split('\n\n')
    chunks = []
    chunk_atual = []
    tamanho_atual = 0
    
    for paragrafo in paragrafos:
        tam = len(paragrafo)
        
        if tamanho_atual + tam > max_chars and chunk_atual:
            chunks.append('\n\n'.join(chunk_atual))
            chunk_atual = [paragrafo]
            tamanho_atual = tam
        else:
            chunk_atual.append(paragrafo)
            tamanho_atual += tam
    
    if chunk_atual:
        chunks.append('\n\n'.join(chunk_atual))
    
    return chunks


def processar_chunk(texto: str, parte: int, total: int) -> str:
    """Processa um chunk via OpenRouter"""
    
    contexto = f"[Parte {parte}/{total}]" if total > 1 else ""
    
    prompt = PROMPT_FORMATACAO.format(texto=texto, contexto=contexto)
    
    print(f"   📝 Parte {parte}/{total}: {len(texto):,} chars entrada...")
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://iudex.app"
    }
    
    payload = {
        "model": MODEL_NAME,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.05,  # Muito baixa para ser literal
        "max_tokens": 16000
    }
    
    response = requests.post(OPENROUTER_BASE_URL, headers=headers, json=payload, timeout=300)
    response.raise_for_status()
    
    resultado = response.json()['choices'][0]['message']['content']
    
    # Verifica se resumiu
    ratio = len(resultado) / len(texto) * 100
    status = "✅" if ratio >= 80 else "⚠️"
    print(f"   {status} Saída: {len(resultado):,} chars ({ratio:.0f}%)")
    
    return resultado


def salvar_arquivos(texto: str, nome_base: str):
    """Salva MD e DOCX"""
    md_path = OUTPUT_DIR / f'{nome_base}_FORMATADO_FINAL.md'
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(texto)
    print(f"   ✅ {md_path}")
    
    docx_path = OUTPUT_DIR / f'{nome_base}_FORMATADO_FINAL.docx'
    doc = Document()
    for linha in texto.split('\n'):
        if linha.strip():
            doc.add_paragraph(linha)
    doc.save(str(docx_path))
    print(f"   ✅ {docx_path}")


def main():
    arquivo = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/Administrativo.txt')
    
    if not arquivo.exists():
        print(f"❌ Arquivo não encontrado: {arquivo}")
        return
    
    print(f"\n📄 Lendo: {arquivo.name}")
    texto = extrair_texto(str(arquivo))
    print(f"   {len(texto):,} caracteres")
    
    chunks = dividir_em_chunks(texto)
    print(f"\n📊 {len(chunks)} partes de ~10k chars cada\n")
    
    resultados = []
    
    for i, chunk in enumerate(chunks, 1):
        try:
            resultado = processar_chunk(chunk, i, len(chunks))
            resultados.append(resultado)
            
            if i < len(chunks):
                time.sleep(0.5)
                
        except Exception as e:
            print(f"   ❌ Erro parte {i}: {e}")
            resultados.append(chunk)  # Mantém original em caso de erro
    
    texto_final = '\n\n'.join(resultados)
    
    print(f"\n💾 Salvando...")
    salvar_arquivos(texto_final, arquivo.stem)
    
    ratio = len(texto_final) / len(texto) * 100
    print(f"\n{'='*50}")
    print(f"🎉 CONCLUÍDO!")
    print(f"   Entrada: {len(texto):,} chars")
    print(f"   Saída:   {len(texto_final):,} chars ({ratio:.0f}%)")
    print(f"{'='*50}")


if __name__ == '__main__':
    main()
