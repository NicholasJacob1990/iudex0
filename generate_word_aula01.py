#!/usr/bin/env python3
"""Gera arquivo Word da Aula 01"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ler apostila formatada
md_file = "/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/01_Aula_Inaugural_YouTube_APOSTILA.md"
with open(md_file, 'r', encoding='utf-8') as f:
    formatted_text = f.read()

print(f"📄 Apostila carregada: {len(formatted_text)} caracteres")

# Criar documento Word
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Processar markdown
lines = formatted_text.split('\n')
i = 0

while i < len(lines):
    line = lines[i].strip()
    
    if not line:
        i += 1
        continue
    
    # Título principal (# )
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(12)
        p.space_after = Pt(12)
    
    # Tópico (## )
    elif line.startswith('## '):
        heading = line[3:].strip()
        p = doc.add_paragraph()
        run = p.add_run(heading)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.space_before = Pt(12)
        p.space_after = Pt(6)
    
    # Subtópico (### )
    elif line.startswith('### '):
        subheading = line[4:].strip()
        p = doc.add_paragraph()
        run = p.add_run(subheading)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)
        p.space_before = Pt(6)
        p.space_after = Pt(6)
    
    # Lista com marcador (- )
    elif line.startswith('- '):
        text = line[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        # Processar markdown inline (negrito **)
        text = re.sub(r'\*\*([^*]+)\*\*', lambda m: f'<b>{m.group(1)}</b>', text)
        if '<b>' in text:
            parts = re.split(r'(<b>.*?</b>)', text)
            for part in parts:
                if part.startswith('<b>') and part.endswith('</b>'):
                    run = p.add_run(part[3:-4])
                    run.font.bold = True
                else:
                    p.add_run(part)
        else:
            p.add_run(text)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    # Parágrafo normal
    else:
        p = doc.add_paragraph()
        # Processar markdown inline
        text = re.sub(r'\*\*([^*]+)\*\*', lambda m: f'<b>{m.group(1)}</b>', line)
        if '<b>' in text:
            parts = re.split(r'(<b>.*?</b>)', text)
            for part in parts:
                if part.startswith('<b>') and part.endswith('</b>'):
                    run = p.add_run(part[3:-4])
                    run.font.bold = True
                else:
                    p.add_run(part)
        else:
            p.add_run(text)
        
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Inches(0.5)
    
    i += 1

# Salvar documento
docx_file = "/Users/nicholasjacob/Documents/Aplicativos/Iudex/Aulas_PGM_RJ/01_Aula_Inaugural_YouTube_APOSTILA.docx"
doc.save(docx_file)

print(f"\n✅ Documento Word salvo: {docx_file}")
